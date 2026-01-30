#!/usr/bin/env python3
"""
miniZ MCP v4.3.0 - Professional Edition with License Management
Web UI + WebSocket MCP + 30 Tools + Hardware License Protection
Copyright © 2025 miniZ Team
"""

import asyncio
import json
import subprocess
import psutil
import time
import os
import sys
from datetime import datetime
from pathlib import Path
from fastapi import FastAPI, WebSocket, HTTPException
from fastapi.responses import HTMLResponse
from pydantic import BaseModel
import websockets
import pyautogui
import difflib
import re

# License Management - DISABLED (FREE EDITION)
# Bypass license check completely
LICENSE_SYSTEM_AVAILABLE = False  # FREE EDITION - No license required

# Auto-startup manager
import winreg
class AutoStartupManager:
    APP_NAME = "miniZ_MCP_Professional"
    
    @staticmethod
    def get_exe_path():
        if getattr(sys, 'frozen', False):
            return sys.executable
        return os.path.abspath(__file__)
    
    @classmethod
    def enable_autostart(cls):
        try:
            exe_path = cls.get_exe_path()
            key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, r"Software\Microsoft\Windows\CurrentVersion\Run", 0, winreg.KEY_SET_VALUE)
            winreg.SetValueEx(key, cls.APP_NAME, 0, winreg.REG_SZ, f'"{exe_path}"')
            winreg.CloseKey(key)
            print(f"✅ [Startup] Đã bật khởi động cùng Windows")
            return True
        except Exception as e:
            print(f"⚠️ [Startup] Không thể bật auto-start: {e}")
            return False
    
    @classmethod
    def is_autostart_enabled(cls):
        try:
            key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, r"Software\Microsoft\Windows\CurrentVersion\Run", 0, winreg.KEY_READ)
            try:
                winreg.QueryValueEx(key, cls.APP_NAME)
                winreg.CloseKey(key)
                return True
            except FileNotFoundError:
                winreg.CloseKey(key)
                return False
        except:
            return False


# ============================================================
# 🔥 FIREWALL/INTERNET CHECKER - Kiểm tra quyền kết nối mạng
# ============================================================
import subprocess

class FirewallChecker:
    """Kiểm tra và hướng dẫn cấp quyền Windows Firewall cho ứng dụng"""
    
    APP_NAME = "miniZ_MCP"
    
    @staticmethod
    def get_exe_path():
        """Lấy đường dẫn file EXE"""
        if getattr(sys, 'frozen', False):
            return sys.executable
        return os.path.abspath(__file__)
    
    @staticmethod
    def get_exe_name():
        """Lấy tên file EXE"""
        if getattr(sys, 'frozen', False):
            return os.path.basename(sys.executable)
        return os.path.basename(__file__)
    
    @classmethod
    def check_firewall_rules(cls) -> dict:
        """
        Kiểm tra xem ứng dụng đã có quyền Firewall chưa
        Returns: dict với keys: has_inbound, has_outbound, rules_found, details
        """
        result = {
            'has_inbound': False,
            'has_outbound': False,
            'rules_found': [],
            'exe_path': cls.get_exe_path(),
            'exe_name': cls.get_exe_name()
        }
        
        try:
            # Tìm tất cả rules liên quan đến miniZ
            cmd = 'netsh advfirewall firewall show rule name=all'
            output = subprocess.run(cmd, capture_output=True, text=True, shell=True, timeout=10)
            
            if output.returncode == 0:
                lines = output.stdout.lower()
                exe_name_lower = result['exe_name'].lower().replace('.exe', '').replace('.py', '')
                
                # Tìm các rules có chứa tên app
                for search_term in ['miniz_mcp', 'miniz mcp', exe_name_lower]:
                    if search_term in lines:
                        result['rules_found'].append(search_term)
                
                # Kiểm tra chi tiết từng rule
                if result['rules_found']:
                    for rule_name in ['miniz_mcp', result['exe_name'].replace('.exe', '').replace('.py', '')]:
                        try:
                            detail_cmd = f'netsh advfirewall firewall show rule name="{rule_name}" verbose'
                            detail_output = subprocess.run(detail_cmd, capture_output=True, text=True, shell=True, timeout=5)
                            if 'direction:' in detail_output.stdout.lower():
                                if 'direction:                            in' in detail_output.stdout.lower():
                                    result['has_inbound'] = True
                                if 'direction:                            out' in detail_output.stdout.lower():
                                    result['has_outbound'] = True
                        except:
                            pass
                    
                    # Nếu tìm thấy rules, assume có quyền (vì Windows tự tạo cả in/out)
                    if result['rules_found'] and not result['has_inbound']:
                        result['has_inbound'] = True  # Giả định có nếu rule tồn tại
                        
        except subprocess.TimeoutExpired:
            print("⚠️ [Firewall] Timeout khi kiểm tra firewall rules")
        except Exception as e:
            print(f"⚠️ [Firewall] Lỗi kiểm tra: {e}")
        
        return result
    
    @classmethod
    def request_firewall_permission(cls) -> bool:
        """
        Tự động thêm rule Firewall (cần quyền Admin)
        Returns: True nếu thành công
        """
        exe_path = cls.get_exe_path()
        rule_name = cls.get_exe_name().replace('.exe', '').replace('.py', '')
        
        try:
            # Thêm rule Inbound
            cmd_in = f'netsh advfirewall firewall add rule name="{rule_name}" dir=in action=allow program="{exe_path}" enable=yes'
            # Thêm rule Outbound  
            cmd_out = f'netsh advfirewall firewall add rule name="{rule_name}" dir=out action=allow program="{exe_path}" enable=yes'
            
            result_in = subprocess.run(cmd_in, capture_output=True, text=True, shell=True, timeout=10)
            result_out = subprocess.run(cmd_out, capture_output=True, text=True, shell=True, timeout=10)
            
            if result_in.returncode == 0 or result_out.returncode == 0:
                print(f"✅ [Firewall] Đã thêm rule firewall cho {rule_name}")
                return True
            else:
                return False
                
        except Exception as e:
            print(f"⚠️ [Firewall] Cần quyền Admin để thêm rule: {e}")
            return False
    
    @classmethod
    def show_firewall_status(cls) -> None:
        """Hiển thị trạng thái Firewall và hướng dẫn nếu cần"""
        print("\n" + "="*60)
        print("🔥 KIỂM TRA QUYỀN KẾT NỐI INTERNET (Windows Firewall)")
        print("="*60)
        
        status = cls.check_firewall_rules()
        
        if status['rules_found']:
            print(f"✅ TRẠNG THÁI: ĐÃ CẤP QUYỀN FIREWALL")
            print(f"   📌 Rules tìm thấy: {', '.join(status['rules_found'])}")
            print(f"   📁 File: {status['exe_name']}")
            print(f"   🔗 Inbound (nhận kết nối): {'✅ Cho phép' if status['has_inbound'] else '⚠️ Chưa rõ'}")
            print(f"   🔗 Outbound (gửi kết nối): {'✅ Cho phép' if status['has_outbound'] else '✅ Mặc định cho phép'}")
            print("\n✅ Ứng dụng có thể kết nối Internet bình thường!")
        else:
            print(f"⚠️ TRẠNG THÁI: CHƯA CÓ QUYỀN FIREWALL")
            print(f"   📁 File: {status['exe_name']}")
            print(f"   📂 Path: {status['exe_path']}")
            print("\n" + "-"*60)
            print("📌 HƯỚNG DẪN CẤP QUYỀN:")
            print("-"*60)
            print("🔹 CÁCH 1: Tự động (lần đầu chạy)")
            print("   - Khi chạy lần đầu, Windows sẽ hỏi 'Allow access'")
            print("   - Nhấn 'Allow access' hoặc 'Cho phép truy cập'")
            print("")
            print("🔹 CÁCH 2: Thủ công qua Windows Security")
            print("   1. Mở 'Windows Security' → 'Firewall & network protection'")
            print("   2. Nhấn 'Allow an app through firewall'")
            print("   3. Nhấn 'Change settings' → 'Allow another app'")
            print("   4. Browse đến file EXE và thêm vào")
            print("   5. Tick cả 'Private' và 'Public' networks")
            print("")
            print("🔹 CÁCH 3: Chạy lệnh PowerShell (Admin)")
            print(f'   netsh advfirewall firewall add rule name="miniZ_MCP" dir=in action=allow program="{status["exe_path"]}" enable=yes')
            print("")
            
            # Thử tự động thêm rule
            print("🔄 Đang thử tự động cấp quyền...")
            if cls.request_firewall_permission():
                print("✅ Đã tự động cấp quyền Firewall thành công!")
            else:
                print("⚠️ Không thể tự động cấp quyền (cần chạy với quyền Admin)")
                print("   → Hãy chạy EXE và cho phép khi Windows hỏi")
        
        print("="*60 + "\n")
        return status['rules_found']
    
    @classmethod
    def check_internet_connection(cls) -> dict:
        """Kiểm tra kết nối Internet thực tế"""
        result = {
            'connected': False,
            'latency_ms': None,
            'test_url': 'google.com'
        }
        
        try:
            import socket
            # Test DNS resolution
            socket.setdefaulttimeout(5)
            socket.socket(socket.AF_INET, socket.SOCK_STREAM).connect(("8.8.8.8", 53))
            result['connected'] = True
            
            # Test latency
            import time
            start = time.time()
            socket.socket(socket.AF_INET, socket.SOCK_STREAM).connect(("google.com", 443))
            result['latency_ms'] = int((time.time() - start) * 1000)
            
        except Exception as e:
            result['error'] = str(e)
        
        return result
    
    @classmethod
    def full_network_check(cls) -> dict:
        """Kiểm tra đầy đủ: Firewall + Internet connection"""
        print("\n🌐 KIỂM TRA KẾT NỐI MẠNG TOÀN DIỆN")
        print("="*50)
        
        # 1. Check Firewall
        firewall_status = cls.check_firewall_rules()
        
        # 2. Check Internet
        internet_status = cls.check_internet_connection()
        
        # 3. Summary
        print(f"🔥 Firewall Rules: {'✅ Đã cấp quyền' if firewall_status['rules_found'] else '⚠️ Chưa có rule'}")
        print(f"🌐 Internet: {'✅ Đã kết nối' if internet_status['connected'] else '❌ Không kết nối'}")
        
        if internet_status.get('latency_ms'):
            print(f"⚡ Độ trễ: {internet_status['latency_ms']}ms")
        
        if not firewall_status['rules_found'] and not internet_status['connected']:
            print("\n⚠️ Có thể ứng dụng đang bị Firewall chặn!")
            print("   → Hãy làm theo hướng dẫn cấp quyền ở trên")
        elif internet_status['connected']:
            print("\n✅ Ứng dụng sẵn sàng sử dụng tất cả tính năng online!")
        
        print("="*50 + "\n")
        
        return {
            'firewall': firewall_status,
            'internet': internet_status,
            'ready': firewall_status['rules_found'] or internet_status['connected']
        }


# Fake license for compatibility
def get_license_manager():
    class FakeLicense:
        def check_license(self): return {'valid': True, 'message': 'FREE EDITION', 'license_data': {'license_type': 'FREE', 'customer_name': 'Community User'}}
        def get_hardware_id(self): return 'FREE-EDITION'
    return FakeLicense()

def show_activation_window(): return True  # Always activated

# MCP Endpoint Manager - Improved connection handling
try:
    from mcp_endpoint_manager import get_endpoint_manager, MCPEndpointManager
    ENDPOINT_MANAGER_AVAILABLE = True
except ImportError:
    ENDPOINT_MANAGER_AVAILABLE = False
    print("⚠️ [Endpoint] MCPEndpointManager not available")

# Gemini AI
try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    print("⚠️ [Gemini] google-generativeai not installed. Run: pip install google-generativeai")

# OpenAI GPT-4
try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False
    print("⚠️ [OpenAI] openai library not installed. Run: pip install openai")

# Selenium Browser Automation
try:
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.common.keys import Keys
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.chrome.service import Service
    from selenium.webdriver.chrome.options import Options
    from webdriver_manager.chrome import ChromeDriverManager
    SELENIUM_AVAILABLE = True
except ImportError:
    SELENIUM_AVAILABLE = False
    print("⚠️ [Selenium] Not installed. Run: pip install selenium webdriver-manager")

# RAG System - Retrieval Augmented Generation
try:
    from rag_system import (
        web_search, rag_search, get_realtime_info, smart_answer,
        RAG_TOOLS, get_rag_engine
    )
    RAG_AVAILABLE = True
    print("✅ [RAG] RAG System loaded - DuckDuckGo + Local KB")
except ImportError as e:
    RAG_AVAILABLE = False
    print(f"⚠️ [RAG] RAG System not available: {e}")

# Vector Search System - Hybrid Semantic Search with FAISS
try:
    # from vector_search import VectorSearchEngine  # Tạm thời tắt do Python 3.14 conflict
    VECTOR_SEARCH_AVAILABLE = False
    print("⚠️ [VectorSearch] Vector search temporarily disabled (Python 3.14 compatibility)")
except ImportError as e:
    VECTOR_SEARCH_AVAILABLE = False
    print(f"⚠️ [VectorSearch] Vector search not available: {e}")

# ============================================================
# UTILITY FUNCTIONS (từ xiaozhi-esp32-server chính thức)
# ============================================================

import re

# ============================================================
# 🔄 SMART TRUNCATE FOR LLM - Giới hạn text gửi về LLM
# ============================================================

MAX_LLM_RESPONSE_CHARS = 2000  # Giới hạn 2000 ký tự cho response gửi LLM
MAX_TTS_RESPONSE_CHARS = 800   # Giới hạn 800 ký tự cho TTS (robot nói trực tiếp)


def clean_markdown_for_tts(text: str) -> str:
    """
    Loại bỏ markdown formatting để TTS đọc được
    """
    import re
    
    # Bỏ headers markdown (# ## ###)
    text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
    
    # Bỏ bold/italic (**text**, *text*, __text__, _text_)
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'__([^_]+)__', r'\1', text)
    text = re.sub(r'_([^_]+)_', r'\1', text)
    
    # Bỏ code blocks và inline code
    text = re.sub(r'```[^`]*```', '', text, flags=re.DOTALL)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    
    # Bỏ horizontal rules (---, ***)
    text = re.sub(r'^[-*]{3,}$', '', text, flags=re.MULTILINE)
    
    # Bỏ bullet points (- *, 1.)
    text = re.sub(r'^\s*[-*]\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*\d+\.\s+', '', text, flags=re.MULTILINE)
    
    # Bỏ links [text](url)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    
    # Chuẩn hóa newlines (nhiều newline -> 1 newline)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'\n\n+', '. ', text)  # Đổi paragraph break thành dấu chấm
    text = re.sub(r'\n', ' ', text)  # Đổi newline thành space
    
    # Chuẩn hóa spaces
    text = re.sub(r'\s{2,}', ' ', text)
    
    return text.strip()

def smart_truncate_for_llm(text: str, max_chars: int = MAX_LLM_RESPONSE_CHARS) -> str:
    """
    Cắt ngắn text thông minh cho LLM, giữ nội dung quan trọng
    
    Args:
        text: Text cần truncate
        max_chars: Giới hạn ký tự (default: 4000)
    
    Returns:
        Text đã truncate với đầy đủ thông tin quan trọng
    """
    if not text or len(text) <= max_chars:
        return text
    
    # Giữ phần đầu (thông tin chính) và phần cuối (kết luận)
    head_ratio = 0.7  # 70% cho phần đầu
    tail_ratio = 0.25  # 25% cho phần cuối
    
    head_chars = int(max_chars * head_ratio)
    tail_chars = int(max_chars * tail_ratio)
    truncate_notice = f"\n\n... [Đã lược bỏ {len(text) - head_chars - tail_chars} ký tự] ...\n\n"
    
    head_part = text[:head_chars]
    tail_part = text[-tail_chars:]
    
    # Cắt ở ranh giới câu nếu có thể
    # Tìm điểm kết thúc câu gần nhất trong head_part
    for sep in ['. ', '.\n', '! ', '!\n', '? ', '?\n', '\n\n']:
        last_sep = head_part.rfind(sep)
        if last_sep > head_chars * 0.8:  # Chỉ cắt nếu >= 80% head_chars
            head_part = head_part[:last_sep + len(sep)]
            break
    
    # Tìm điểm bắt đầu câu gần nhất trong tail_part
    for sep in ['. ', '.\n', '\n\n']:
        first_sep = tail_part.find(sep)
        if first_sep != -1 and first_sep < tail_chars * 0.2:  # Chỉ cắt nếu <= 20% tail_chars
            tail_part = tail_part[first_sep + len(sep):]
            break
    
    return head_part + truncate_notice + tail_part


def format_result_for_llm(result: dict, max_chars: int = MAX_LLM_RESPONSE_CHARS) -> str:
    """
    Format và truncate result dict thành text cho LLM
    
    Args:
        result: Dict kết quả từ tool
        max_chars: Giới hạn ký tự
    
    Returns:
        Text đã format và truncate
    """
    import json
    
    # Nếu là response_text từ Gemini, ưu tiên nó
    if isinstance(result, dict):
        if result.get("response_text"):
            text = result["response_text"]
            return smart_truncate_for_llm(text, max_chars)
        
        # Nếu có context (từ knowledge base), ưu tiên
        if result.get("context"):
            text = result["context"]
            return smart_truncate_for_llm(text, max_chars)
        
        # Nếu có message, dùng message
        if result.get("message"):
            text = result["message"]
            # Nếu message ngắn, thêm thông tin khác
            if len(text) < max_chars * 0.5:
                extra_info = []
                for key in ["summary", "content", "data", "results"]:
                    if result.get(key):
                        val = result[key]
                        if isinstance(val, str):
                            extra_info.append(val)
                        elif isinstance(val, (list, dict)):
                            extra_info.append(json.dumps(val, ensure_ascii=False, indent=1))
                if extra_info:
                    text += "\n\n" + "\n".join(extra_info)
            return smart_truncate_for_llm(text, max_chars)
    
    # Default: convert to JSON
    text = json.dumps(result, ensure_ascii=False, indent=1)
    return smart_truncate_for_llm(text, max_chars)


def sanitize_tool_name(name: str) -> str:
    """
    Chuẩn hóa tên tool theo quy tắc của Xiaozhi server
    - Thay thế các ký tự đặc biệt bằng underscore
    - Chuyển về lowercase
    """
    if not name:
        return name
    # Thay thế các ký tự không phải alphanumeric hoặc underscore
    sanitized = re.sub(r'[^a-zA-Z0-9_]', '_', name)
    # Loại bỏ underscore liên tiếp
    sanitized = re.sub(r'_+', '_', sanitized)
    # Loại bỏ underscore ở đầu và cuối
    sanitized = sanitized.strip('_')
    return sanitized.lower()

async def get_system_info(category="all"):
    """
    Thu thập thông tin cấu hình máy tính chi tiết
    category: all, cpu, memory, disk, os, network, gpu, software, motherboard
    """
    try:
        import platform
        import psutil
        import socket
        import subprocess
        import json
        from datetime import datetime
        
        info = {
            "success": True,
            "timestamp": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
            "categories": []
        }
        
        # CPU Information (Chi tiết hơn)
        if category in ["all", "cpu"]:
            cpu_info = {
                "name": "CPU Information",
                "processor": platform.processor(),
                "architecture": platform.architecture()[0],
                "machine": platform.machine(),
                "cores_physical": psutil.cpu_count(logical=False),
                "cores_logical": psutil.cpu_count(logical=True),
                "cpu_usage_percent": psutil.cpu_percent(interval=1)
            }
            
            # Thêm frequency info
            if psutil.cpu_freq():
                freq = psutil.cpu_freq()
                cpu_info.update({
                    "cpu_freq_current_mhz": round(freq.current, 2) if freq.current else "N/A",
                    "cpu_freq_max_mhz": round(freq.max, 2) if freq.max else "N/A",
                    "cpu_freq_min_mhz": round(freq.min, 2) if freq.min else "N/A"
                })
            
            # Thêm CPU details từ Windows Registry/WMI nếu có thể
            try:
                if platform.system() == "Windows":
                    import winreg
                    # Đọc CPU name từ registry
                    key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, 
                                       r"HARDWARE\DESCRIPTION\System\CentralProcessor\0")
                    cpu_name = winreg.QueryValueEx(key, "ProcessorNameString")[0].strip()
                    cpu_info["cpu_name_detailed"] = cpu_name
                    
                    # Phát hiện thế hệ CPU (heuristic)
                    cpu_name_lower = cpu_name.lower()
                    if "intel" in cpu_name_lower:
                        if "13th gen" in cpu_name_lower or "13900" in cpu_name_lower or "13700" in cpu_name_lower or "13600" in cpu_name_lower:
                            cpu_info["cpu_generation"] = "Intel 13th Gen (Raptor Lake)"
                        elif "12th gen" in cpu_name_lower or "12900" in cpu_name_lower or "12700" in cpu_name_lower or "12600" in cpu_name_lower:
                            cpu_info["cpu_generation"] = "Intel 12th Gen (Alder Lake)"
                        elif "11th gen" in cpu_name_lower or "11900" in cpu_name_lower or "11700" in cpu_name_lower or "11600" in cpu_name_lower:
                            cpu_info["cpu_generation"] = "Intel 11th Gen (Tiger Lake/Rocket Lake)"
                        elif "10th gen" in cpu_name_lower or "10900" in cpu_name_lower or "10700" in cpu_name_lower or "10600" in cpu_name_lower:
                            cpu_info["cpu_generation"] = "Intel 10th Gen (Comet Lake/Ice Lake)"
                        elif "9th gen" in cpu_name_lower or "9900" in cpu_name_lower or "9700" in cpu_name_lower or "9600" in cpu_name_lower:
                            cpu_info["cpu_generation"] = "Intel 9th Gen (Coffee Lake Refresh)"
                        elif "8th gen" in cpu_name_lower or "8700" in cpu_name_lower or "8600" in cpu_name_lower or "8400" in cpu_name_lower:
                            cpu_info["cpu_generation"] = "Intel 8th Gen (Coffee Lake)"
                        elif "7th gen" in cpu_name_lower or "7700" in cpu_name_lower or "7600" in cpu_name_lower or "7500" in cpu_name_lower:
                            cpu_info["cpu_generation"] = "Intel 7th Gen (Kaby Lake)"
                        else:
                            cpu_info["cpu_generation"] = "Intel (Generation unknown)"
                    elif "amd" in cpu_name_lower:
                        if "7000" in cpu_name_lower or "7950x" in cpu_name_lower or "7900x" in cpu_name_lower:
                            cpu_info["cpu_generation"] = "AMD Ryzen 7000 Series (Zen 4)"
                        elif "5000" in cpu_name_lower or "5950x" in cpu_name_lower or "5900x" in cpu_name_lower:
                            cpu_info["cpu_generation"] = "AMD Ryzen 5000 Series (Zen 3)"
                        elif "3000" in cpu_name_lower or "3900x" in cpu_name_lower or "3700x" in cpu_name_lower:
                            cpu_info["cpu_generation"] = "AMD Ryzen 3000 Series (Zen 2)"
                        elif "2000" in cpu_name_lower or "2700x" in cpu_name_lower or "2600x" in cpu_name_lower:
                            cpu_info["cpu_generation"] = "AMD Ryzen 2000 Series (Zen+)"
                        else:
                            cpu_info["cpu_generation"] = "AMD (Generation unknown)"
                    
                    winreg.CloseKey(key)
            except Exception as e:
                cpu_info["cpu_detection_error"] = f"Could not detect detailed CPU info: {str(e)}"
            
            info["categories"].append(cpu_info)
        
        # Memory Information (Chi tiết hơn)
        if category in ["all", "memory"]:
            memory = psutil.virtual_memory()
            swap = psutil.swap_memory()
            memory_info = {
                "name": "Memory Information",
                "total_ram_gb": round(memory.total / (1024**3), 2),
                "available_ram_gb": round(memory.available / (1024**3), 2),
                "used_ram_gb": round(memory.used / (1024**3), 2),
                "ram_usage_percent": memory.percent,
                "swap_total_gb": round(swap.total / (1024**3), 2),
                "swap_used_gb": round(swap.used / (1024**3), 2),
                "swap_usage_percent": swap.percent,
                "memory_total_mb": round(memory.total / (1024**2)),
                "memory_speed_estimate": "DDR4/DDR5 (Detection requires additional tools)"
            }
            info["categories"].append(memory_info)
        
        # GPU Information (Cải thiện)
        if category in ["all", "gpu"]:
            gpu_info = {
                "name": "GPU Information",
                "gpus": []
            }
            
            # Method 1: GPUtil
            try:
                import GPUtil
                gpus = GPUtil.getGPUs()
                for gpu in gpus:
                    gpu_data = {
                        "id": gpu.id,
                        "name": gpu.name,
                        "memory_total_mb": gpu.memoryTotal,
                        "memory_used_mb": gpu.memoryUsed,
                        "memory_free_mb": gpu.memoryFree,
                        "gpu_load_percent": round(gpu.load * 100, 1),
                        "temperature_c": gpu.temperature,
                        "driver": "Unknown (GPUtil limitation)"
                    }
                    
                    # Detect GPU generation/series (heuristic)
                    gpu_name_lower = gpu.name.lower()
                    if "rtx 40" in gpu_name_lower or "4090" in gpu_name_lower or "4080" in gpu_name_lower:
                        gpu_data["gpu_generation"] = "NVIDIA RTX 40 Series (Ada Lovelace)"
                    elif "rtx 30" in gpu_name_lower or "3090" in gpu_name_lower or "3080" in gpu_name_lower or "3070" in gpu_name_lower:
                        gpu_data["gpu_generation"] = "NVIDIA RTX 30 Series (Ampere)"
                    elif "rtx 20" in gpu_name_lower or "2080" in gpu_name_lower or "2070" in gpu_name_lower:
                        gpu_data["gpu_generation"] = "NVIDIA RTX 20 Series (Turing)"
                    elif "gtx 16" in gpu_name_lower or "1660" in gpu_name_lower or "1650" in gpu_name_lower:
                        gpu_data["gpu_generation"] = "NVIDIA GTX 16 Series (Turing)"
                    elif "gtx 10" in gpu_name_lower or "1080" in gpu_name_lower or "1070" in gpu_name_lower or "1060" in gpu_name_lower:
                        gpu_data["gpu_generation"] = "NVIDIA GTX 10 Series (Pascal)"
                    elif "rx 7000" in gpu_name_lower or "7900 xt" in gpu_name_lower or "7800 xt" in gpu_name_lower:
                        gpu_data["gpu_generation"] = "AMD RX 7000 Series (RDNA 3)"
                    elif "rx 6000" in gpu_name_lower or "6900 xt" in gpu_name_lower or "6800 xt" in gpu_name_lower:
                        gpu_data["gpu_generation"] = "AMD RX 6000 Series (RDNA 2)"
                    elif "rx 5000" in gpu_name_lower or "5700 xt" in gpu_name_lower or "5600 xt" in gpu_name_lower:
                        gpu_data["gpu_generation"] = "AMD RX 5000 Series (RDNA)"
                    else:
                        gpu_data["gpu_generation"] = "Unknown generation"
                    
                    gpu_info["gpus"].append(gpu_data)
            except ImportError:
                gpu_info["gputil_status"] = "GPUtil not installed. Run: pip install GPUtil"
            except Exception as e:
                gpu_info["gputil_error"] = f"GPUtil error: {str(e)}"
            
            # Method 2: Windows WMI fallback
            if not gpu_info["gpus"] and platform.system() == "Windows":
                try:
                    result = subprocess.run(
                        ['wmic', 'path', 'win32_VideoController', 'get', 'name,AdapterRAM,DriverVersion', '/format:csv'],
                        capture_output=True, text=True, timeout=10
                    )
                    lines = result.stdout.strip().split('\n')[1:]  # Skip header
                    for line in lines:
                        if line.strip() and ',' in line:
                            parts = line.split(',')
                            if len(parts) >= 4:
                                gpu_data = {
                                    "name": parts[2].strip() if len(parts) > 2 else "Unknown",
                                    "memory_total_mb": round(int(parts[1]) / (1024*1024)) if parts[1].strip().isdigit() else "Unknown",
                                    "driver_version": parts[3].strip() if len(parts) > 3 else "Unknown",
                                    "method": "WMI (Windows)"
                                }
                                gpu_info["gpus"].append(gpu_data)
                except Exception as e:
                    gpu_info["wmi_error"] = f"WMI detection failed: {str(e)}"
            
            info["categories"].append(gpu_info)
        
        # Disk Information (như cũ)
        if category in ["all", "disk"]:
            disk_info = {
                "name": "Disk Information",
                "partitions": []
            }
            
            for partition in psutil.disk_partitions():
                try:
                    usage = psutil.disk_usage(partition.mountpoint)
                    partition_info = {
                        "device": partition.device,
                        "mountpoint": partition.mountpoint,
                        "file_system": partition.fstype,
                        "total_gb": round(usage.total / (1024**3), 2),
                        "used_gb": round(usage.used / (1024**3), 2),
                        "free_gb": round(usage.free / (1024**3), 2),
                        "usage_percent": round((usage.used / usage.total) * 100, 1)
                    }
                    disk_info["partitions"].append(partition_info)
                except PermissionError:
                    continue
            
            info["categories"].append(disk_info)
        
        # Operating System Information
        if category in ["all", "os"]:
            os_info = {
                "name": "Operating System",
                "system": platform.system(),
                "release": platform.release(),
                "version": platform.version(),
                "platform": platform.platform(),
                "hostname": socket.gethostname(),
                "boot_time": datetime.fromtimestamp(psutil.boot_time()).strftime("%d/%m/%Y %H:%M:%S"),
                "python_version": platform.python_version()
            }
            
            # Windows specific info
            if platform.system() == "Windows":
                try:
                    result = subprocess.run(['systeminfo'], capture_output=True, text=True, timeout=15)
                    if result.returncode == 0:
                        lines = result.stdout.split('\n')
                        for line in lines:
                            if "Total Physical Memory" in line:
                                os_info["total_physical_memory"] = line.split(':')[1].strip()
                            elif "System Manufacturer" in line:
                                os_info["system_manufacturer"] = line.split(':')[1].strip()
                            elif "System Model" in line:
                                os_info["system_model"] = line.split(':')[1].strip()
                except:
                    pass
            
            info["categories"].append(os_info)
        
        # Network Information (như cũ)
        if category in ["all", "network"]:
            network_info = {
                "name": "Network Information",
                "hostname": socket.gethostname(),
                "interfaces": []
            }
            
            try:
                # Get local IP
                s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
                s.connect(("8.8.8.8", 80))
                local_ip = s.getsockname()[0]
                s.close()
                network_info["local_ip"] = local_ip
            except:
                network_info["local_ip"] = "N/A"
            
            # Network interfaces
            for interface, addresses in psutil.net_if_addrs().items():
                interface_info = {
                    "interface": interface,
                    "addresses": []
                }
                for addr in addresses:
                    if addr.family == socket.AF_INET:  # IPv4
                        interface_info["addresses"].append({
                            "type": "IPv4",
                            "address": addr.address,
                            "netmask": addr.netmask
                        })
                network_info["interfaces"].append(interface_info)
            
            info["categories"].append(network_info)
        
        # Motherboard Information (Windows only)
        if category in ["all", "motherboard"]:
            motherboard_info = {
                "name": "Motherboard Information",
                "manufacturer": "N/A",
                "product": "N/A",
                "bios_version": "N/A"
            }
            
            if platform.system() == "Windows":
                try:
                    # Get motherboard info via WMI
                    result = subprocess.run(
                        ['wmic', 'baseboard', 'get', 'Manufacturer,Product,Version', '/format:csv'],
                        capture_output=True, text=True, timeout=10
                    )
                    if result.returncode == 0:
                        lines = result.stdout.strip().split('\n')[1:]
                        for line in lines:
                            if line.strip() and ',' in line:
                                parts = line.split(',')
                                if len(parts) >= 3:
                                    motherboard_info["manufacturer"] = parts[1].strip()
                                    motherboard_info["product"] = parts[2].strip()
                                    break
                    
                    # Get BIOS info
                    result = subprocess.run(
                        ['wmic', 'bios', 'get', 'SMBIOSBIOSVersion', '/format:csv'],
                        capture_output=True, text=True, timeout=10
                    )
                    if result.returncode == 0:
                        lines = result.stdout.strip().split('\n')[1:]
                        for line in lines:
                            if line.strip() and ',' in line:
                                parts = line.split(',')
                                if len(parts) >= 2:
                                    motherboard_info["bios_version"] = parts[1].strip()
                                    break
                except Exception as e:
                    motherboard_info["error"] = f"Could not detect motherboard: {str(e)}"
            
            info["categories"].append(motherboard_info)
        
        # Software Information (như cũ)
        if category in ["all", "software"]:
            software_info = {
                "name": "Installed Software (Python Packages)",
                "python_packages": [],
                "note": "Showing top 20 Python packages"
            }
            
            try:
                # Try modern importlib.metadata first (Python 3.8+)
                try:
                    import importlib.metadata
                    installed_packages = [f"{dist.metadata['Name']}=={dist.version}" 
                                        for dist in importlib.metadata.distributions()]
                except ImportError:
                    # Fallback to pkg_resources for older Python versions
                    import pkg_resources
                    installed_packages = [d.project_name + "==" + d.version for d in pkg_resources.working_set]
                
                software_info["python_packages"] = sorted(installed_packages)[:20]
                if len(installed_packages) > 20:
                    software_info["total_packages"] = len(installed_packages)
            except Exception as e:
                software_info["error"] = f"Could not list packages: {str(e)}"
            
            info["categories"].append(software_info)
        
        # Ensure all values are JSON serializable
        import json
        try:
            json.dumps(info, ensure_ascii=False)
        except Exception as json_error:
            print(f"⚠️ [JSON Serialization Error] {json_error}")
            # Fix potential serialization issues
            for category in info.get("categories", []):
                for key, value in list(category.items()):
                    if value is None:
                        category[key] = "N/A"
                    elif not isinstance(value, (str, int, float, bool, list, dict)):
                        category[key] = str(value)
        
        return info
    
    except Exception as e:
        return {
            "success": False,
            "error": f"Lỗi khi đọc thông tin hệ thống: {str(e)}",
            "help": "Có thể cần cài đặt thêm: pip install psutil GPUtil"
        }

# Tool retry configuration (từ repo chính thức)
MAX_TOOL_RETRIES = 3
TOOL_RETRY_INTERVAL = 2  # seconds

# ============================================================
# 🧠 INTENT DETECTION LLM - Phân tích ý định trước khi xử lý
# (Từ xiaozhi-esp32-server chính thức)
# ============================================================

class IntentDetector:
    """
    Intent Detection LLM - Phân tích câu hỏi và xác định tool cần gọi
    Tương tự intent_llm trong repo chính thức
    """
    
    # Các intent patterns
    REALTIME_PATTERNS = [
        # Giá cả
        r'giá\s*(vàng|xăng|dầu|usd|đô|euro|bitcoin|btc|eth)',
        r'(vàng|xăng|dầu|bitcoin|btc)\s*giá',
        r'tỷ\s*giá',
        r'bao\s*nhiêu\s*tiền',
        # Thời tiết
        r'thời\s*tiết',
        r'trời\s*(nắng|mưa|nóng|lạnh)',
        r'nhiệt\s*độ',
        # Người/Chức vụ
        r'(tổng\s*thống|thủ\s*tướng|chủ\s*tịch|ceo|giám\s*đốc)',
        r'ai\s*(là|đang)',
        r'(là\s*ai|là\s*gì)',
        r'hiện\s*(tại|nay|giờ)',
        # Thời gian thực
        r'(hôm\s*nay|bây\s*giờ|hiện\s*tại|mới\s*nhất)',
        r'(2024|2025|năm\s*nay)',
        r'tin\s*(tức|mới)',
        r'sự\s*kiện',
        # Sản phẩm/Công ty
        r'(iphone|samsung|apple|google|microsoft|tesla)',
    ]
    
    MUSIC_PATTERNS = [
        r'(bài\s*tiếp|next|chuyển\s*bài)',
        r'(bài\s*trước|previous|quay\s*lại)',
        r'(dừng|pause|tạm\s*dừng|stop)',
        r'(tiếp\s*tục|resume|play)',
        r'(phát\s*nhạc|mở\s*nhạc|bật\s*nhạc)',
        r'(tắt\s*nhạc|ngừng\s*nhạc)',
        r'(tăng|giảm)\s*(âm\s*lượng|volume)',
    ]
    
    KNOWLEDGE_BASE_PATTERNS = [
        r'(tài\s*liệu|document|file)',
        r'(trong\s*thư\s*viện|knowledge\s*base)',
        r'(tra\s*cứu\s*nội\s*bộ)',
    ]
    
    SYSTEM_INFO_PATTERNS = [
        r'cấu\s*hình.*máy\s*tính',
        r'máy\s*tính.*cấu\s*hình',
        r'cấu\s*hình.*hệ\s*thống',
        r'specs.*máy',
        r'hardware.*info',
        r'thông\s*tin.*hệ\s*thống',
        r'thông\s*tin.*máy\s*tính',
        r'kiểm\s*tra.*cấu\s*hình',
        r'kiểm\s*tra.*specs',
        r'kiểm\s*tra.*hardware',
        r'máy\s*tính.*như\s*thế\s*nào',
        r'máy\s*này.*ra\s*sao',
        r'card.*(màn\s*hình|đồ\s*họa|vga)',
        r'gpu.*gì',
        r'vga.*gì',
        r'cpu.*gì',
        r'cpu.*thế\s*hệ',
        r'processor.*generation',
        r'(mainboard|motherboard)',
        r'bo\s*mạch\s*chủ',
        r'(intel|amd|nvidia|rtx|gtx).*thế\s*hệ',
        r'nhiệt\s*độ.*(cpu|gpu)',
        r'(ram|memory).*bao\s*nhiêu',
        r'bộ\s*nhớ.*gì',
        r'asus.*mainboard',
        r'msi.*mainboard',
        r'gigabyte.*mainboard',
    ]
    
    @classmethod
    def detect_intent(cls, text: str) -> dict:
        """
        Phân tích text và trả về intent + suggested tool
        Returns: {
            "intent": "realtime|music|knowledge|general",
            "suggested_tool": "web_search|get_realtime_info|smart_music_control|...",
            "confidence": 0.0-1.0,
            "should_force_tool": True/False
        }
        """
        text_lower = text.lower()
        
        # Check realtime patterns
        for pattern in cls.REALTIME_PATTERNS:
            if re.search(pattern, text_lower):
                # Xác định tool cụ thể
                if any(word in text_lower for word in ['giá', 'tỷ giá', 'bao nhiêu']):
                    tool = "get_realtime_info"
                elif any(word in text_lower for word in ['thời tiết', 'nhiệt độ', 'trời']):
                    tool = "get_realtime_info"
                elif any(word in text_lower for word in ['tin tức', 'sự kiện', 'mới nhất']):
                    tool = "web_search"
                elif any(word in text_lower for word in ['là ai', 'ai là', 'tổng thống', 'thủ tướng', 'ceo']):
                    tool = "web_search"
                else:
                    tool = "smart_answer"
                    
                return {
                    "intent": "realtime",
                    "suggested_tool": tool,
                    "confidence": 0.9,
                    "should_force_tool": True,
                    "reason": f"Detected realtime pattern: {pattern}"
                }
        
        # Check music patterns
        for pattern in cls.MUSIC_PATTERNS:
            if re.search(pattern, text_lower):
                return {
                    "intent": "music",
                    "suggested_tool": "smart_music_control",
                    "confidence": 0.95,
                    "should_force_tool": True,
                    "reason": f"Detected music pattern: {pattern}"
                }
        
        # Check system info patterns (mới thêm)
        for pattern in cls.SYSTEM_INFO_PATTERNS:
            if re.search(pattern, text_lower):
                print(f"[DEBUG] System info pattern matched: {pattern} for text: {text_lower}")
                return {
                    "intent": "system_info",
                    "suggested_tool": "get_hardware_specs",
                    "confidence": 0.95,
                    "should_force_tool": True,
                    "reason": f"Detected system info pattern: {pattern}"
                }
        
        # Check knowledge base patterns
        for pattern in cls.KNOWLEDGE_BASE_PATTERNS:
            if re.search(pattern, text_lower):
                return {
                    "intent": "knowledge",
                    "suggested_tool": "get_knowledge_context",
                    "confidence": 0.85,
                    "should_force_tool": True,
                    "reason": f"Detected knowledge pattern: {pattern}"
                }
        
        # General intent - không cần force tool
        return {
            "intent": "general",
            "suggested_tool": None,
            "confidence": 0.5,
            "should_force_tool": False,
            "reason": "No specific pattern matched"
        }
    
    @classmethod
    async def detect_with_llm(cls, text: str, gemini_key: str = None, include_user_context: bool = True) -> dict:
        """
        Sử dụng Gemini để phân tích intent phức tạp hơn
        Chỉ gọi khi pattern matching không chắc chắn
        Có thể kèm user context để hiểu người dùng tốt hơn
        """
        # Đầu tiên thử pattern matching
        result = cls.detect_intent(text)
        
        # Nếu confidence cao, không cần LLM
        if result["confidence"] >= 0.8:
            return result
        
        # Nếu có Gemini API, dùng LLM để phân tích
        if gemini_key and GEMINI_AVAILABLE:
            try:
                genai.configure(api_key=gemini_key)
                model = genai.GenerativeModel('models/gemini-3-flash-preview')
                
                # Lấy user context nếu được yêu cầu
                user_context = ""
                if include_user_context:
                    try:
                        user_context = f"""
[USER CONTEXT - Dùng để hiểu người dùng tốt hơn]
{get_user_profile_summary()}

[RECENT CONVERSATION]
{get_conversation_context(5)}
"""
                    except:
                        user_context = ""
                
                prompt = f'''Phân tích câu hỏi sau và xác định intent:
"{text}"
{user_context}
Trả lời JSON:
{{"intent": "realtime|music|knowledge|general", "tool": "web_search|get_realtime_info|smart_music_control|get_knowledge_context|none", "reason": "lý do ngắn"}}

Quy tắc:
- realtime: Câu hỏi về thông tin thời gian thực (giá cả, thời tiết, tin tức, người nổi tiếng hiện tại)
- music: Điều khiển nhạc
- knowledge: Tra cứu tài liệu nội bộ
- general: Câu hỏi thông thường

CHỈ TRẢ LỜI JSON, KHÔNG GIẢI THÍCH.'''

                response = model.generate_content(prompt)
                response_text = response.text.strip()
                
                # Parse JSON từ response
                import json
                # Tìm JSON trong response
                json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
                if json_match:
                    llm_result = json.loads(json_match.group())
                    return {
                        "intent": llm_result.get("intent", "general"),
                        "suggested_tool": llm_result.get("tool") if llm_result.get("tool") != "none" else None,
                        "confidence": 0.85,
                        "should_force_tool": llm_result.get("intent") in ["realtime", "music"],
                        "reason": llm_result.get("reason", "LLM analysis"),
                        "source": "gemini_llm"
                    }
            except Exception as e:
                print(f"⚠️ [IntentDetector] LLM error: {e}")
        
        return result

# Intent Detector instance
intent_detector = IntentDetector()

# ============================================================
# CONFIGURATION
# ============================================================

# 🔥 FIX: Detect if running as EXE (frozen) or script
# When frozen (EXE), use sys.executable path (dist folder)
# When script, use __file__ path (source folder)
if getattr(sys, 'frozen', False):
    # Running as EXE - use executable's directory
    CONFIG_FILE = Path(sys.executable).parent / "xiaozhi_endpoints.json"
else:
    # Running as script - use script's directory
    CONFIG_FILE = Path(__file__).parent / "xiaozhi_endpoints.json"

GEMINI_API_KEY = ""  # Sẽ được load từ xiaozhi_endpoints.json
OPENAI_API_KEY = ""  # Sẽ được load từ xiaozhi_endpoints.json
SERPER_API_KEY = ""  # Google Search API - Miễn phí 2500 queries/tháng

# ============================================================
# 🎵 MUSIC SYSTEM PROMPT - Hướng dẫn LLM về Music Tools
# ============================================================
MUSIC_SYSTEM_PROMPT = """
🎵 ĐIỀU KHIỂN NHẠC - QUAN TRỌNG!

⚡ QUY TẮC #1: KHI NGHE TỪ KHÓA DƯỚI ĐÂY → GỌI TOOL NGAY, KHÔNG HỎI LẠI!

┌─────────────────────────────────────────────────────────────┐
│ 📌 TỪ KHÓA → GỌI TOOL                                       │
├─────────────────────────────────────────────────────────────┤
│ "bài tiếp"/"next"/"chuyển bài" → music_next()               │
│ "bài trước"/"quay lại"        → music_previous()            │
│ "dừng"/"pause"/"tạm dừng"     → pause_music()               │
│ "tắt nhạc"/"stop"             → stop_music()                │
│ "tiếp tục"/"resume"           → resume_music()              │
│ "phát bài [tên]"              → play_music(filename="tên")  │
└─────────────────────────────────────────────────────────────┘

⚠️ VOICE VARIANTS (ESP32 recognition sai):
• "bai tiep", "tiep theo", "nex", "ních" → music_next()
• "bai truoc", "quay lai", "pre"        → music_previous()
• "dung", "pao", "poz", "tam dung"       → pause_music()
• "tat nhac", "stóp", "dung han"         → stop_music()

🔥 NGUYÊN TẮC: GỌI TOOL TRỰC TIẾP, KHÔNG CẦN HỎI!
• User: "bài tiếp" → Bạn GỌI music_next() → Trả lời "Đã chuyển bài"
• User: "dừng"     → Bạn GỌI pause_music() → Trả lời "Đã tạm dừng"
• User: "quay lại" → Bạn GỌI music_previous() → Trả lời "Đã quay lại"

📍 Server: Python-VLC Player (tích hợp sẵn)
📁 Thư mục nhạc: F:\\nhac

🎬 YOUTUBE: CHỈ khi user nói "youtube"/"video" → youtube_* tools
   ✨ NEW: open_youtube() GIỜ TỰ ĐỘNG PHÁT VIDEO TRỰC TIẾP!
   - Query >= 2 từ → Direct video (youtube.com/watch?v=...)
   - Query 1 từ → Search page
   VD: "mở youtube Lạc Trôi" hoặc "mở youtube Sơn Tùng MTP" → PHÁT VIDEO NGAY!
═══════════════════════════════════════════════════════════════
🔧 FUZZY MATCHING - HỖ TRỢ VOICE RECOGNITION
═══════════════════════════════════════════════════════════════

Hệ thống có fuzzy matching cho các biến thể:
• "bai tiep" → "bài tiếp" 
• "bai truoc" → "bài trước"
• "phat nhac" → "phát nhạc"
• "nếch" → "next"
• "prê" → "previous"

→ Cứ gửi nguyên văn lệnh, hệ thống sẽ tự nhận dạng!

═══════════════════════════════════════════════════════════════
🎵 VLC MUSIC CONTROLS - ĐIỀU KHIỂN NHẠC
═══════════════════════════════════════════════════════════════

⚡⚡⚡ BẮT BUỘC: KHI USER YÊU CẦU ĐIỀU KHIỂN NHẠC → GỌI TOOL NGAY! ⚡⚡⚡

🚫 TUYỆT ĐỐI CẤM TỰ TRẢ LỜI "OK" hoặc "Đã chuyển bài" mà KHÔNG GỌI TOOL!

📌 MAPPING COMMANDS → TOOLS (BẮT BUỘC GỌI):
┌─────────────────────────────────────────────────────────────┐
│ "bài tiếp", "next", "skip"           → music_next()       │
│ "quay lại", "bài trước", "previous"  → music_previous()   │
│ "tạm dừng", "pause"                   → pause_music()      │
│ "tiếp tục", "resume", "phát tiếp"    → resume_music()     │
│ "dừng", "stop"                        → stop_music()       │
│ "phát [tên bài]", "play [song]"      → play_music(song)   │
└─────────────────────────────────────────────────────────────┘

✅ WORKFLOW ĐÚNG:
User: "bài tiếp"
→ GỌI: music_next()
→ NHẬN: {"success": true, "message": "Đã chuyển: Song.mp3"}
→ TRẢ LỜI: "Đã chuyển sang bài tiếp: Song.mp3"

❌ WORKFLOW SAI (CẤM):
User: "bài tiếp"
→ Trả lời trực tiếp: "OK, đã chuyển bài"  ← SAI! KHÔNG GỌI TOOL!

🔴 RULES NGHIÊM NGẶT:
1. PHẢI gọi tool TRƯỚC khi trả lời
2. KHÔNG được giả định thành công
3. PHẢI đợi tool response
4. CHỈ trả lời dựa trên tool result

⚠️ ĐẶC BIỆT: Các từ "next", "previous", "pause", "stop" → 100% GỌI TOOL!

═══════════════════════════════════════════════════════════════
🎵 VLC MUSIC CONTROLS - ĐIỀU KHIỂN NHẠC
═══════════════════════════════════════════════════════════════

⚡⚡⚡ BẮT BUỘC: KHI USER YÊU CẦU ĐIỀU KHIỂN NHẠC → GỌI TOOL NGAY! ⚡⚡⚡

🚫 TUYỆT ĐỐI CẤM TỰ TRẢ LỜI "OK" hoặc "Đã chuyển bài" mà KHÔNG GỌI TOOL!

📌 MAPPING COMMANDS → TOOLS (BẮT BUỘC GỌI):
┌─────────────────────────────────────────────────────────────┐
│ "bài tiếp", "next", "skip"           → music_next()       │
│ "quay lại", "bài trước", "previous"  → music_previous()   │
│ "tạm dừng", "pause"                   → pause_music()      │
│ "tiếp tục", "resume", "phát tiếp"    → resume_music()     │
│ "dừng", "stop"                        → stop_music()       │
│ "phát [tên bài]", "play [song]"      → play_music(song)   │
└─────────────────────────────────────────────────────────────┘

✅ WORKFLOW ĐÚNG:
User: "bài tiếp"
→ GỌI: music_next()
→ NHẬN: {"success": true, "message": "Đã chuyển: Song.mp3"}
→ TRẢ LỜI: "Đã chuyển sang bài tiếp: Song.mp3"

❌ WORKFLOW SAI (CẤM):
User: "bài tiếp"
→ Trả lời trực tiếp: "OK, đã chuyển bài"  ← SAI! KHÔNG GỌI TOOL!

🔴 RULES NGHIÊM NGẶT:
1. PHẢI gọi tool TRƯỚC khi trả lời
2. KHÔNG được giả định thành công
3. PHẢI đợi tool response
4. CHỈ trả lời dựa trên tool result

⚠️ ĐẶC BIỆT: Các từ "next", "previous", "pause", "stop" → 100% GỌI TOOL!

═══════════════════════════════════════════════════════════════
📚 KNOWLEDGE BASE - TÀI LIỆU CỦA USER (TỰ ĐỘNG TÌM KIẾM)
═══════════════════════════════════════════════════════════════

🔥 QUY TẮC VÀNG: KHI NGHI NGỜ THÔNG TIN CÓ THỂ Ở TRONG TÀI LIỆU → GỌI KB NGAY!

⚡ AUTO-TRIGGERS - Gemini TỰ ĐỘNG GỌI KB khi phát hiện:
┌─────────────────────────────────────────────────────────────┐
│ 📌 DIRECT COMMANDS (100% gọi KB):                           │
│ • "tìm trong tài liệu", "tra cứu KB", "search documents"    │
│ • "theo file của tôi", "trong dữ liệu", "in my docs"        │
│ • "kiểm tra tài liệu", "xem trong KB", "check docs"         │
│                                                              │
│ 🔍 IMPLICIT QUERIES (phát hiện thông minh):                 │
│ • "[tên cụ thể] là gì/ai/ở đâu" (VD: "Lê Trung Khoa là ai")│
│ • "thông tin về [X]" (VD: "thông tin về dự án ABC")        │
│ • "dự án/hợp đồng/báo cáo [X]" (tên riêng, không phổ biến) │
│ • "theo dữ liệu...", "căn cứ vào...", "based on..."        │
│ • "[X] có bao nhiêu...", "[X] như thế nào"                  │
│                                                              │
│ ❓ SMART DETECTION (nghi ngờ → thử KB):                     │
│ • Câu hỏi về người/công ty/dự án CỤ THỂ (không phổ biến)  │
│ • Câu hỏi về con số, số liệu, thống kê (có thể từ báo cáo) │
│ • Câu hỏi yêu cầu thông tin CHI TIẾT (có thể trong docs)   │
└─────────────────────────────────────────────────────────────┘

📖 WORKFLOW CHUẨN:
┌─────────────────────────────────────────────────────────────┐
│ User: "Lê Trung Khoa là ai?"                                │
│ ↓                                                            │
│ [Gemini phát hiện: tên cụ thể → có thể trong KB]           │
│ ↓                                                            │
│ Gọi: get_knowledge_context(query="Lê Trung Khoa")          │
│ ↓                                                            │
│ Nhận: Context từ "kiến thức c.docx" về Lê Trung Khoa       │
│ ↓                                                            │
│ Trả lời: "Theo tài liệu 'kiến thức c.docx', Lê Trung Khoa  │
│ là người bị Bộ Công an ra quyết định truy nã ngày 5/12..."│
└─────────────────────────────────────────────────────────────┘

🎯 2 Tools chính:
┌─────────────────────────────────────────────────────────────┐
│ ✅ get_knowledge_context(query, max_chars=10000)            │
│    → Lấy FULL CONTENT để trả lời (ƯU TIÊN DÙNG TOOL NÀY)   │
│    → Có Gemini auto-summarize nếu nội dung dài >2000 chars │
│    → Trả về context đầy đủ để LLM đọc và trả lời           │
│                                                              │
│ 📋 search_knowledge_base(query)                             │
│    → Tìm và show SNIPPETS (dùng khi user muốn xem list)    │
│    → Trả về top 5 documents với highlights                  │
└─────────────────────────────────────────────────────────────┘

⚠️ PHÂN BIỆT:
• "Lê Trung Khoa là ai?" → GỌI get_knowledge_context() (tên cụ thể → KB)
• "Tổng thống Mỹ là ai?" → KHÔNG gọi KB (thông tin phổ biến)
• "Python là gì?" → KHÔNG gọi KB (kiến thức chung)
• "Dự án ABC có bao nhiêu giai đoạn?" → GỌI KB (tên dự án cụ thể)
• "Nguyễn Công Huy sinh năm nào?" → GỌI KB (tên người cụ thể)

🔴 QUY TẮC QUAN TRỌNG:
1. NGHI NGỜ → GỌI KB (tốt hơn là bỏ lỡ thông tin)
2. Nếu KB trả về "không tìm thấy" → Dùng kiến thức chung
3. Nếu KB có kết quả → ƯU TIÊN context từ KB
4. Luôn trích dẫn nguồn khi dùng KB: "Theo tài liệu '[tên file]'..."

💡 TIP: Khi không chắc → GỌI get_knowledge_context() để kiểm tra!

═══════════════════════════════════════════════════════════════
🌐 RAG SYSTEM - RETRIEVAL AUGMENTED GENERATION
═══════════════════════════════════════════════════════════════

⛔⛔⛔ CẢNH BÁO NGHIÊM TRỌNG: BẠN KHÔNG CÓ KIẾN THỨC SAU 2024! ⛔⛔⛔

🚫 TUYỆT ĐỐI CẤM TỰ TRẢ LỜI KHI CÂU HỎI CHỨA:
• "hiện nay", "bây giờ", "hôm nay", "hiện tại"
• "2024", "2025", "năm nay"
• "mới nhất", "cập nhật", "gần đây"
• "ai là", "là ai", "là gì", "ở đâu"
• Tên người nổi tiếng: tổng thống, thủ tướng, CEO...
• Giá cả: vàng, USD, bitcoin, chứng khoán...
• Thời tiết, tin tức, sự kiện

⚠️ LÝ DO: Kiến thức của bạn ĐÃ LỖI THỜI! Ví dụ:
• Trump có thể đã trở lại làm tổng thống (bạn không biết)
• Giá vàng có thể đã thay đổi (bạn không biết)
• Có thể có CEO mới (bạn không biết)

🔴 HÀNH ĐỘNG BẮT BUỘC:
1. PHẢI GỌI web_search() hoặc get_realtime_info() TRƯỚC
2. CHỜ kết quả tra cứu
3. RỒI MỚI trả lời dựa trên thông tin mới nhất

📌 TOOLS MAPPING (PHẢI SỬ DỤNG):
┌────────────────────────────────────────────────────────────┐
│ web_search(query)       → Tìm Internet (DuckDuckGo)       │
│ get_realtime_info(query)→ Thông tin thời gian thực        │
│ rag_search(query)       → Hybrid: Web + Local KB          │
│ smart_answer(query)     → AI tự chọn nguồn tốt nhất       │
└────────────────────────────────────────────────────────────┘

📖 VÍ DỤ ĐÚNG:
User: "Tổng thống Mỹ hiện tại là ai?"
→ ❌ SAI: Trả lời "Joe Biden" (kiến thức cũ có thể sai!)
→ ✅ ĐÚNG: GỌI get_realtime_info("tổng thống Mỹ hiện tại 2024")
→ Nhận kết quả → Trả lời chính xác

User: "Giá vàng hôm nay?"
→ ❌ SAI: Đoán hoặc nói "tôi không biết"
→ ✅ ĐÚNG: GỌI get_realtime_info("giá vàng SJC hôm nay")

User: "Thời tiết Hà Nội?"
→ ✅ GỌI: get_realtime_info("thời tiết Hà Nội hôm nay")

🔥 QUY TẮC BẮT BUỘC:
1. Câu hỏi về NGƯỜI → web_search("tên người + chức vụ")
2. Câu hỏi về GIÁ CẢ → get_realtime_info()
3. Câu hỏi về THỜI TIẾT → get_realtime_info()
4. Câu hỏi về SỰ KIỆN → web_search()
5. KHÔNG CHẮC → smart_answer() (AI tự động chọn)

⚡ NHỚ: GỌI TOOL TRƯỚC, TRẢ LỜI SAU! KHÔNG BAO GIỜ TỰ ĐOÁN!
"""

DEFAULT_ENDPOINT = {
    "name": "Thiết bị 1",
    "token": "eyJhbGciOiJFUzI1NiIsInR5cCI6IkpXVCJ9.eyJ1c2VySWQiOjQ1MzYxMSwiYWdlbnRJZCI6OTQ0MjE4LCJlbmRwb2ludElkIjoiYWdlbnRfOTQ0MjE4IiwicHVycG9zZSI6Im1jcC1lbmRwb2ludCIsImlhdCI6MTc2MjA4NTI1OSwiZXhwIjoxNzkzNjQyODU5fQ.GK91-17mqarpETPwz7N6rZj5DaT7bJkpK7EM6lO0Rdmfztv_KeOTBP9R4Lvy3uXKMCJn3gwucvelCur95GAn5Q",
    "enabled": True
}

def load_endpoints_from_file():
    """Đọc cấu hình endpoints từ file JSON"""
    global GEMINI_API_KEY, OPENAI_API_KEY, SERPER_API_KEY
    
    if CONFIG_FILE.exists():
        try:
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                print(f"✅ [Config] Loaded {len(data.get('endpoints', []))} endpoints from {CONFIG_FILE.name}")
                
                # Load Gemini API key nếu có
                if data.get('gemini_api_key'):
                    GEMINI_API_KEY = data['gemini_api_key']
                    print(f"✅ [Gemini] API key loaded (ends with ...{GEMINI_API_KEY[-8:]})")
                
                # Load OpenAI API key nếu có
                if data.get('openai_api_key'):
                    OPENAI_API_KEY = data['openai_api_key']
                    print(f"✅ [OpenAI] API key loaded (ends with ...{OPENAI_API_KEY[-8:]})")
                
                # Load Serper API key nếu có (Google Search)
                if data.get('serper_api_key'):
                    SERPER_API_KEY = data['serper_api_key']
                    # Cũng cập nhật vào environment variable để rag_system.py có thể dùng
                    os.environ['SERPER_API_KEY'] = SERPER_API_KEY
                    print(f"✅ [Serper] Google Search API key loaded (ends with ...{SERPER_API_KEY[-8:]})")
                
                return data.get('endpoints', []), data.get('active_index', 0)
        except Exception as e:
            print(f"⚠️ [Config] Error loading {CONFIG_FILE.name}: {e}")
    
    # Trả về cấu hình mặc định nếu không có file
    return [
        DEFAULT_ENDPOINT,
        {"name": "Thiết bị 2", "token": "", "enabled": False},
        {"name": "Thiết bị 3", "token": "", "enabled": False}
    ], 0

def save_endpoints_to_file(endpoints, active_index, force_save=False):
    """Lưu cấu hình endpoints vào file JSON - LUÔN LƯU khi có thay đổi"""
    global GEMINI_API_KEY, OPENAI_API_KEY, SERPER_API_KEY
    
    try:
        # Data mới cần lưu
        new_data = {
            'endpoints': endpoints,
            'active_index': active_index,
            'gemini_api_key': GEMINI_API_KEY,
            'openai_api_key': OPENAI_API_KEY,
            'serper_api_key': SERPER_API_KEY,
            'last_updated': datetime.now().isoformat()
        }
        
        # 🔥 FIX: Chỉ skip save nếu KHÔNG phải force_save và không có thay đổi
        if not force_save and CONFIG_FILE.exists():
            try:
                with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                    old_data = json.load(f)
                    # So sánh TẤT CẢ: endpoints, active_index VÀ API keys
                    if (old_data.get('endpoints') == endpoints and 
                        old_data.get('active_index') == active_index and
                        old_data.get('gemini_api_key') == GEMINI_API_KEY and
                        old_data.get('openai_api_key') == OPENAI_API_KEY and
                        old_data.get('serper_api_key') == SERPER_API_KEY):
                        # Không có thay đổi gì cả, skip save
                        print(f"ℹ️ [Config] No changes detected, skipping save")
                        return True
            except Exception:
                pass
        
        # Có thay đổi → Lưu file
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(new_data, f, ensure_ascii=False, indent=2)
        
        # Log chi tiết
        empty_count = sum(1 for ep in endpoints if not ep.get('token', '').strip())
        active_count = len(endpoints) - empty_count
        print(f"💾 [Config] Saved to {CONFIG_FILE.name} ({active_count} active, {empty_count} empty endpoints)")
        return True
    except Exception as e:
        print(f"❌ [Config] Error saving to {CONFIG_FILE.name}: {e}")
        return False

# Load cấu hình từ file
endpoints_config, loaded_active_index = load_endpoints_from_file()
active_endpoint_index = loaded_active_index

# Support 3 simultaneous MCP connections
xiaozhi_connections = {0: None, 1: None, 2: None}  # Dict of {index: websocket}
xiaozhi_connected = {0: False, 1: False, 2: False}  # Connection status for each device
should_reconnect = {0: False, 1: False, 2: False}  # Reconnect flags

active_connections = []

# ============================================================
# TASK MEMORY SYSTEM - Ghi nhớ tác vụ đã thực hiện
# ============================================================
TASK_MEMORY_FILE = Path(__file__).parent / "task_memory.json"
MAX_TASK_HISTORY = 100  # Giới hạn số tác vụ lưu trữ

def load_task_memory():
    """Đọc lịch sử tác vụ từ file"""
    if TASK_MEMORY_FILE.exists():
        try:
            with open(TASK_MEMORY_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return data.get('tasks', [])
        except Exception as e:
            print(f"⚠️ [TaskMemory] Error loading: {e}")
    return []

def save_task_memory(tasks: list):
    """Lưu lịch sử tác vụ vào file"""
    try:
        # Giới hạn số lượng
        if len(tasks) > MAX_TASK_HISTORY:
            tasks = tasks[-MAX_TASK_HISTORY:]
        
        with open(TASK_MEMORY_FILE, 'w', encoding='utf-8') as f:
            json.dump({
                'tasks': tasks,
                'last_updated': datetime.now().isoformat()
            }, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        print(f"❌ [TaskMemory] Error saving: {e}")
        return False

def add_task_to_memory(tool_name: str, params: dict, result: dict, user_request: str = ""):
    """Thêm tác vụ vào bộ nhớ"""
    tasks = load_task_memory()
    
    task_entry = {
        "timestamp": datetime.now().isoformat(),
        "tool": tool_name,
        "params": params,
        "result_success": result.get("success", False),
        "result_message": result.get("message", result.get("error", "")),
        "user_request": user_request
    }
    
    tasks.append(task_entry)
    save_task_memory(tasks)
    return task_entry

def get_recent_tasks(limit: int = 10) -> list:
    """Lấy các tác vụ gần đây"""
    tasks = load_task_memory()
    return tasks[-limit:] if tasks else []

def search_task_memory(keyword: str) -> list:
    """Tìm kiếm tác vụ theo từ khóa"""
    tasks = load_task_memory()
    keyword_lower = keyword.lower()
    
    results = []
    for task in tasks:
        # Tìm trong tool name, params, user_request
        if (keyword_lower in task.get('tool', '').lower() or
            keyword_lower in str(task.get('params', {})).lower() or
            keyword_lower in task.get('user_request', '').lower() or
            keyword_lower in task.get('result_message', '').lower()):
            results.append(task)
    
    return results[-20:]  # Giới hạn 20 kết quả

def clear_task_memory() -> bool:
    """Xóa toàn bộ lịch sử tác vụ"""
    try:
        if TASK_MEMORY_FILE.exists():
            TASK_MEMORY_FILE.unlink()
        return True
    except Exception as e:
        print(f"❌ [TaskMemory] Error clearing: {e}")
        return False

# Load task memory khi khởi động
task_memory_cache = load_task_memory()
print(f"📝 [TaskMemory] Loaded {len(task_memory_cache)} previous tasks")

# ============================================================
# CONVERSATION HISTORY - Lưu lịch sử hội thoại TOÀN BỘ
# ============================================================
conversation_history = []  # List để lưu tất cả messages
conversation_sessions = {}  # Sessions theo ngày

# Thư mục lưu hội thoại
import os
from pathlib import Path as PathLib
CONVERSATION_BASE_DIR = PathLib(os.path.expanduser("~")) / "AppData" / "Local" / "miniZ_MCP" / "conversations"
CONVERSATION_BASE_DIR.mkdir(parents=True, exist_ok=True)

# File tổng hợp (backward compatible)
CONVERSATION_FILE = CONVERSATION_BASE_DIR / "conversation_history.json"

# File lưu user profile (hiểu người dùng)
USER_PROFILE_FILE = CONVERSATION_BASE_DIR / "user_profile.json"

# NOTE: get_today_conversation_file() đã bị xóa để tối ưu - không lưu file theo ngày nữa

def load_conversation_history():
    """Load lịch sử hội thoại từ file"""
    global conversation_history
    try:
        # Load file tổng hợp (CHỈ một file duy nhất - nhanh hơn)
        if CONVERSATION_FILE.exists():
            with open(CONVERSATION_FILE, 'r', encoding='utf-8') as f:
                conversation_history = json.load(f)
            print(f"📚 [Conversation] Loaded {len(conversation_history)} messages")
    except Exception as e:
        print(f"⚠️ Could not load conversation history: {e}")
        conversation_history = []

def save_conversation_history():
    """Lưu lịch sử hội thoại vào file (CHỈ file tổng hợp - tối ưu tốc độ)"""
    try:
        # CHỈ lưu file tổng hợp (không lưu file theo ngày để tăng tốc)
        with open(CONVERSATION_FILE, 'w', encoding='utf-8') as f:
            json.dump(conversation_history, f, ensure_ascii=False, indent=2)
            
    except Exception as e:
        print(f"⚠️ Could not save conversation history: {e}")

def add_to_conversation(role: str, content: str, metadata: dict = None):
    """
    Thêm message vào lịch sử hội thoại - TỐI ƯU CHO PERFORMANCE
    
    role: 'user', 'assistant', 'system', 'tool'
    content: nội dung message
    metadata: thông tin bổ sung (tool_name, timestamp, source, etc.)
    
    OPTIMIZATION: Chỉ save sau 20 messages hoặc khi shutdown
    """
    from datetime import datetime
    
    message = {
        "role": role,
        "content": content,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "metadata": metadata or {}
    }
    
    # Thêm session_id nếu chưa có
    if "session_id" not in message["metadata"]:
        message["metadata"]["session_id"] = datetime.now().strftime("%Y%m%d")
    
    conversation_history.append(message)
    
    # TĂNG TỐC: Chỉ save sau mỗi 20 messages (giảm I/O disk)
    if len(conversation_history) % 20 == 0:
        save_conversation_history()
    
    # NOTE: Disabled user profile analysis (gây chậm)

def update_user_profile_from_message(content: str, metadata: dict = None):
    """Cập nhật user profile từ message để hiểu người dùng hơn"""
    try:
        from datetime import datetime
        
        profile = load_user_profile()
        
        # Đếm số lần tương tác
        profile["total_interactions"] = profile.get("total_interactions", 0) + 1
        profile["last_interaction"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Phân tích topics
        topics = profile.get("topics", {})
        content_lower = content.lower()
        
        # Detect topics từ nội dung
        topic_keywords = {
            "music": ["nhạc", "bài", "hát", "music", "song", "play", "pause", "volume"],
            "weather": ["thời tiết", "weather", "mưa", "nắng", "nhiệt độ", "temperature"],
            "news": ["tin", "news", "mới", "sự kiện", "event"],
            "finance": ["giá", "vàng", "gold", "btc", "bitcoin", "chứng khoán", "stock", "usd", "tỷ giá"],
            "system": ["âm lượng", "volume", "mở", "open", "tắt", "close", "kill"],
            "web": ["tìm", "search", "google", "web", "tra cứu"],
            "coding": ["code", "python", "javascript", "lập trình", "debug", "function"],
            "general": ["là gì", "what is", "how to", "làm sao", "tại sao", "why"]
        }
        
        for topic, keywords in topic_keywords.items():
            if any(kw in content_lower for kw in keywords):
                topics[topic] = topics.get(topic, 0) + 1
        
        profile["topics"] = topics
        
        # Lưu các câu hỏi thường gặp (top 20)
        frequent_queries = profile.get("frequent_queries", [])
        # Chỉ lưu câu ngắn gọn
        if len(content) < 100:
            frequent_queries.append({
                "query": content[:80],
                "time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            })
            # Giữ 20 câu gần nhất
            profile["frequent_queries"] = frequent_queries[-20:]
        
        # Thống kê giờ hoạt động
        hour_stats = profile.get("active_hours", {})
        current_hour = datetime.now().strftime("%H")
        hour_stats[current_hour] = hour_stats.get(current_hour, 0) + 1
        profile["active_hours"] = hour_stats
        
        save_user_profile(profile)
        
    except Exception as e:
        print(f"⚠️ [UserProfile] Error updating: {e}")

def load_user_profile() -> dict:
    """Load user profile"""
    try:
        if USER_PROFILE_FILE.exists():
            with open(USER_PROFILE_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
    except:
        pass
    return {
        "created": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "total_interactions": 0,
        "topics": {},
        "frequent_queries": [],
        "active_hours": {},
        "preferences": {}
    }

def save_user_profile(profile: dict):
    """Lưu user profile"""
    try:
        with open(USER_PROFILE_FILE, 'w', encoding='utf-8') as f:
            json.dump(profile, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"⚠️ [UserProfile] Error saving: {e}")

def get_conversation_context(max_messages: int = 10) -> str:
    """
    Lấy context từ lịch sử hội thoại gần đây để hiểu người dùng
    Dùng cho LLM để có thêm context
    """
    recent = conversation_history[-max_messages:] if len(conversation_history) > max_messages else conversation_history
    
    context_lines = []
    for msg in recent:
        role = msg.get("role", "unknown")
        content = msg.get("content", "")[:200]  # Giới hạn độ dài
        if role in ["user", "assistant"]:
            context_lines.append(f"{role.upper()}: {content}")
    
    return "\n".join(context_lines)

def get_user_profile_summary() -> str:
    """Tóm tắt profile người dùng cho LLM"""
    try:
        profile = load_user_profile()
        
        # Top topics
        topics = profile.get("topics", {})
        sorted_topics = sorted(topics.items(), key=lambda x: x[1], reverse=True)[:5]
        top_topics = ", ".join([f"{t[0]}({t[1]})" for t in sorted_topics]) if sorted_topics else "chưa xác định"
        
        # Active hours
        hours = profile.get("active_hours", {})
        sorted_hours = sorted(hours.items(), key=lambda x: int(x[1]), reverse=True)[:3]
        active_hours = ", ".join([f"{h[0]}h" for h in sorted_hours]) if sorted_hours else "chưa xác định"
        
        summary = f"""
[USER PROFILE]
- Tổng số tương tác: {profile.get('total_interactions', 0)}
- Chủ đề quan tâm: {top_topics}
- Giờ hoạt động: {active_hours}
- Lần cuối: {profile.get('last_interaction', 'N/A')}
"""
        return summary.strip()
    except:
        return "[USER PROFILE] Chưa có dữ liệu"

def export_conversation_to_file(filename: str = "") -> dict:
    """Export lịch sử hội thoại ra file riêng"""
    try:
        from datetime import datetime
        import os
        
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"conversation_export_{timestamp}.json"
        
        documents_path = os.path.expanduser("~\\Documents")
        save_folder = os.path.join(documents_path, "miniZ_Conversations")
        os.makedirs(save_folder, exist_ok=True)
        
        file_path = os.path.join(save_folder, filename)
        
        # Export với format đẹp + user profile
        export_data = {
            "export_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "total_messages": len(conversation_history),
            "user_profile": load_user_profile(),
            "messages": conversation_history
        }
        
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, ensure_ascii=False, indent=2)
        
        return {
            "success": True,
            "message": f"📚 Đã export {len(conversation_history)} messages + user profile",
            "path": file_path
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

def list_conversation_files() -> list:
    """Liệt kê tất cả file hội thoại đã lưu"""
    try:
        files = []
        for f in CONVERSATION_BASE_DIR.glob("conversation_*.json"):
            stat = f.stat()
            files.append({
                "filename": f.name,
                "path": str(f),
                "size_kb": round(stat.st_size / 1024, 2),
                "modified": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
            })
        return sorted(files, key=lambda x: x["modified"], reverse=True)
    except Exception as e:
        return []

# Load lịch sử khi khởi động
load_conversation_history()
print(f"📂 [Conversation] Storage: {CONVERSATION_BASE_DIR}")

# ============================================================
# CONVERSATION FORMATTING HELPERS
# ============================================================

def format_tool_request(tool_name: str, args: dict) -> str:
    """Format tool request thành câu dễ đọc"""
    if tool_name == "set_volume":
        level = args.get("level", 0)
        return f"Điều chỉnh âm lượng lên {level}%"
    elif tool_name == "get_volume":
        return "Kiểm tra âm lượng hiện tại"
    elif tool_name == "screenshot":
        return "Chụp màn hình"
    elif tool_name == "open_application":
        app = args.get("app_name", "")
        return f"Mở ứng dụng {app}"
    elif tool_name == "get_active_media_players":
        return "Kiểm tra các trình duyệt và media player đang chạy"
    elif tool_name == "list_running_processes":
        limit = args.get("limit", 10)
        return f"Liệt kê {limit} tiến trình đang chạy"
    elif tool_name == "kill_process":
        identifier = args.get("identifier", "")
        force = args.get("force", True)
        return f"{'FORCE ' if force else ''}Kill tiến trình: {identifier}"
    elif tool_name == "force_kill_app":
        app_name = args.get("app_name", "")
        return f"💀 FORCE KILL APP: {app_name}"
    # YouTube controls
    elif tool_name == "control_youtube":
        action = args.get("action", "")
        return f"🎬 YouTube: {action}"
    elif tool_name == "youtube_play_pause":
        return "⏯️ YouTube: Play/Pause"
    elif tool_name == "youtube_rewind":
        seconds = args.get("seconds", 10)
        return f"⏪ YouTube: Lùi {seconds} giây"
    elif tool_name == "youtube_forward":
        seconds = args.get("seconds", 10)
        return f"⏩ YouTube: Tua tới {seconds} giây"
    elif tool_name == "youtube_volume_up":
        return "🔊 YouTube: Tăng âm lượng"
    elif tool_name == "youtube_volume_down":
        return "🔉 YouTube: Giảm âm lượng"
    elif tool_name == "youtube_mute":
        return "🔇 YouTube: Bật/Tắt tiếng"
    elif tool_name == "youtube_fullscreen":
        return "📺 YouTube: Fullscreen"
    # VLC controls
    elif tool_name == "control_vlc":
        action = args.get("action", "")
        return f"🎵 VLC: {action}"
    elif tool_name == "vlc_play_pause":
        return "⏯️ VLC: Play/Pause"
    elif tool_name == "vlc_stop":
        return "⏹️ VLC: Dừng phát"
    elif tool_name == "vlc_next":
        return "⏭️ VLC: Bài tiếp theo"
    elif tool_name == "vlc_previous":
        return "⏮️ VLC: Bài trước"
    elif tool_name == "vlc_volume_up":
        return "🔊 VLC: Tăng âm lượng"
    elif tool_name == "vlc_volume_down":
        return "🔉 VLC: Giảm âm lượng"
    elif tool_name == "vlc_mute":
        return "🔇 VLC: Bật/Tắt tiếng"
    # WMP controls
    elif tool_name == "control_wmp":
        action = args.get("action", "")
        return f"🎶 Windows Media Player: {action}"
    elif tool_name.startswith("wmp_"):
        action = tool_name.replace("wmp_", "").replace("_", " ").title()
        return f"🎶 Windows Media Player: {action}"
    # Smart media control
    elif tool_name == "smart_media_control":
        action = args.get("action", "")
        return f"🎛️ Smart Media: {action}"
    elif tool_name == "create_file":
        path = args.get("path", "")
        return f"Tạo file mới: {path}"
    elif tool_name == "read_file":
        path = args.get("path", "")
        return f"Đọc nội dung file: {path}"
    elif tool_name == "search_web":
        query = args.get("query", "")
        return f"Tìm kiếm Google: {query}"
    elif tool_name == "ask_gemini":
        prompt = args.get("prompt", "")[:50]
        return f"Hỏi Gemini AI: {prompt}..."
    elif tool_name == "ask_gpt4":
        prompt = args.get("prompt", "")[:50]
        return f"Hỏi GPT-4: {prompt}..."
    else:
        # Default format
        if args:
            args_str = ", ".join([f"{k}={v}" for k, v in list(args.items())[:2]])
            return f"Gọi tool {tool_name} ({args_str})"
        return f"Gọi tool {tool_name}"

def format_tool_response(tool_name: str, response: dict) -> str:
    """Format tool response thành câu dễ đọc"""
    if isinstance(response, dict):
        # Kiểm tra lỗi
        if response.get("isError"):
            error_text = ""
            if "content" in response and isinstance(response["content"], list):
                for item in response["content"]:
                    if item.get("type") == "text":
                        error_text = item.get("text", "")
                        break
            return f"❌ Lỗi: {error_text}"
        
        # Success responses
        if "content" in response and isinstance(response["content"], list):
            for item in response["content"]:
                if item.get("type") == "text":
                    text = item.get("text", "")
                    # Rút gọn nếu quá dài
                    if len(text) > 150:
                        return f"✅ {text[:150]}..."
                    return f"✅ {text}"
        
        # Fallback cho response khác
        if "message" in response:
            return f"✅ {response['message']}"
        
    return "✅ Thực hiện thành công"

print("🚀 miniZ MCP - Sidebar UI")
print(f"🌐 Web: http://localhost:8000")
print(f"📡 MCP: Multi-device ready")

# ============================================================
# TOOL IMPLEMENTATIONS (20 TOOLS)
# ============================================================

async def set_volume(level: int) -> dict:
    """Điều chỉnh âm lượng hệ thống - Windows only"""
    try:
        if not 0 <= level <= 100:
            return {"success": False, "error": "Level phải từ 0-100"}
        
        # Sử dụng PowerShell trực tiếp (tương thích tốt hơn với Python 3.13)
        ps_cmd = f"""
[void] [System.Reflection.Assembly]::LoadWithPartialName("System.Windows.Forms")
$obj = New-Object System.Windows.Forms.Form
$obj.KeyPreview = $True

# Get current volume
$wshShell = New-Object -ComObject WScript.Shell
for($i=1; $i -le 50; $i++){{$wshShell.SendKeys([char]174)}}  # Mute to 0

# Set to desired level
$steps = [Math]::Round({level} / 2)
for($i=1; $i -le $steps; $i++){{$wshShell.SendKeys([char]175)}}  # Volume up

Write-Output "Volume set to {level}%"
"""
        
        proc = await asyncio.create_subprocess_exec(
            "powershell", "-NoProfile", "-Command", ps_cmd,
            stdout=asyncio.subprocess.PIPE, 
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=5)
        
        if proc.returncode == 0:
            return {
                "success": True, 
                "level": level,
                "message": f"✅ Âm lượng đã đặt: {level}%"
            }
        else:
            error_msg = stderr.decode('utf-8', errors='ignore').strip()
            return {"success": False, "error": f"PowerShell error: {error_msg[:200]}"}
                
    except asyncio.TimeoutError:
        return {"success": False, "error": "Timeout khi điều chỉnh âm lượng"}
    except Exception as e:
        return {"success": False, "error": f"Lỗi: {str(e)}"}

async def mute_volume() -> dict:
    """Tắt tiếng (mute) hệ thống"""
    try:
        ps_cmd = """
$obj = New-Object -ComObject WScript.Shell
$obj.SendKeys([char]173)
Write-Output "Volume muted"
"""
        proc = await asyncio.create_subprocess_exec(
            "powershell", "-NoProfile", "-Command", ps_cmd,
            stdout=asyncio.subprocess.PIPE, 
            stderr=asyncio.subprocess.PIPE
        )
        await asyncio.wait_for(proc.communicate(), timeout=3)
        
        return {"success": True, "message": "🔇 Đã tắt tiếng"}
    except Exception as e:
        return {"success": False, "error": f"Lỗi: {str(e)}"}

async def unmute_volume() -> dict:
    """Bật lại tiếng (unmute) hệ thống"""
    try:
        ps_cmd = """
$obj = New-Object -ComObject WScript.Shell
$obj.SendKeys([char]173)
Write-Output "Volume unmuted"
"""
        proc = await asyncio.create_subprocess_exec(
            "powershell", "-NoProfile", "-Command", ps_cmd,
            stdout=asyncio.subprocess.PIPE, 
            stderr=asyncio.subprocess.PIPE
        )
        await asyncio.wait_for(proc.communicate(), timeout=3)
        
        return {"success": True, "message": "🔊 Đã bật tiếng"}
    except Exception as e:
        return {"success": False, "error": f"Lỗi: {str(e)}"}

async def volume_up(steps: int = 5) -> dict:
    """Tăng âm lượng lên (mỗi step ~2%)"""
    try:
        ps_cmd = f"""
$obj = New-Object -ComObject WScript.Shell
for($i=1; $i -le {steps}; $i++){{$obj.SendKeys([char]175)}}
Write-Output "Volume increased by {steps} steps"
"""
        proc = await asyncio.create_subprocess_exec(
            "powershell", "-NoProfile", "-Command", ps_cmd,
            stdout=asyncio.subprocess.PIPE, 
            stderr=asyncio.subprocess.PIPE
        )
        await asyncio.wait_for(proc.communicate(), timeout=3)
        
        return {"success": True, "message": f"🔊 Đã tăng âm lượng ({steps} bước)"}
    except Exception as e:
        return {"success": False, "error": f"Lỗi: {str(e)}"}

async def volume_down(steps: int = 5) -> dict:
    """Giảm âm lượng xuống (mỗi step ~2%)"""
    try:
        ps_cmd = f"""
$obj = New-Object -ComObject WScript.Shell
for($i=1; $i -le {steps}; $i++){{$obj.SendKeys([char]174)}}
Write-Output "Volume decreased by {steps} steps"
"""
        proc = await asyncio.create_subprocess_exec(
            "powershell", "-NoProfile", "-Command", ps_cmd,
            stdout=asyncio.subprocess.PIPE, 
            stderr=asyncio.subprocess.PIPE
        )
        await asyncio.wait_for(proc.communicate(), timeout=3)
        
        return {"success": True, "message": f"🔉 Đã giảm âm lượng ({steps} bước)"}
    except Exception as e:
        return {"success": False, "error": f"Lỗi: {str(e)}"}

async def get_volume() -> dict:
    """Lấy mức âm lượng hiện tại của hệ thống"""
    try:
        try:
            from pycaw.pycaw import AudioUtilities, IAudioEndpointVolume
            from comtypes import CLSCTX_ALL
            
            devices = AudioUtilities.GetSpeakers()
            interface = devices.Activate(IAudioEndpointVolume._iid_, CLSCTX_ALL, None)
            volume = interface.QueryInterface(IAudioEndpointVolume)
            
            current_volume = int(volume.GetMasterVolumeLevelScalar() * 100)
            is_muted = volume.GetMute()
            
            return {
                "success": True,
                "level": current_volume,
                "muted": bool(is_muted),
                "message": f"🔊 Âm lượng hiện tại: {current_volume}%" + (" (Tắt tiếng)" if is_muted else "")
            }
        except ImportError:
            # Fallback PowerShell
            ps_cmd = """
Add-Type -TypeDefinition @'
using System.Runtime.InteropServices;
[Guid("5CDF2C82-841E-4546-9722-0CF74078229A"), InterfaceType(ComInterfaceType.InterfaceIsIUnknown)]
interface IAudioEndpointVolume {
    int NotImpl1(); int NotImpl2();
    int GetMasterVolumeLevelScalar(out float level);
}
[Guid("BCDE0395-E52F-467C-8E3D-C4579291692E")]
class MMDeviceEnumeratorComObject { }
[Guid("A95664D2-9614-4F35-A746-DE8DB63617E6"), InterfaceType(ComInterfaceType.InterfaceIsIUnknown)]
interface IMMDeviceEnumerator {
    int NotImpl1();
    int GetDefaultAudioEndpoint(int dataFlow, int role, out IMMDevice device);
}
[Guid("D666063F-1587-4E43-81F1-B948E807363F"), InterfaceType(ComInterfaceType.InterfaceIsIUnknown)]
interface IMMDevice {
    int Activate(ref System.Guid id, int clsCtx, int activationParams, out IAudioEndpointVolume aev);
}
'@
$enumerator = [System.Activator]::CreateInstance([Type]::GetTypeFromCLSID([Guid]'BCDE0395-E52F-467C-8E3D-C4579291692E'))
$device = $null
$enumerator.GetDefaultAudioEndpoint(0, 1, [ref]$device)
$aev = $null
$device.Activate([Guid]'5CDF2C82-841E-4546-9722-0CF74078229A', 0, 0, [ref]$aev)
$current = 0.0
$aev.GetMasterVolumeLevelScalar([ref]$current)
Write-Output ([int]($current * 100))
"""
            proc = await asyncio.create_subprocess_exec(
                "powershell", "-NoProfile", "-Command", ps_cmd,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )
            stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=3)
            
            if proc.returncode == 0:
                level = int(stdout.decode('utf-8', errors='ignore').strip())
                return {
                    "success": True,
                    "level": level,
                    "message": f"🔊 Âm lượng hiện tại: {level}%"
                }
            else:
                return {"success": False, "error": "Không thể lấy âm lượng"}
    except Exception as e:
        return {"success": False, "error": f"Lỗi: {str(e)}"}

async def take_screenshot(filename: str = None) -> dict:
    """Chụp màn hình toàn bộ và lưu file
    
    Args:
        filename: Tên file lưu ảnh (optional). Mặc định: screenshot_YYYYMMDD_HHMMSS.png
    
    Returns:
        dict với thông tin file đã lưu
    """
    try:
        import pyautogui
        from datetime import datetime
        import os
        
        # Tạo tên file mặc định nếu không có
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"screenshot_{timestamp}.png"
        
        # Đảm bảo có extension .png
        if not filename.lower().endswith(('.png', '.jpg', '.jpeg')):
            filename += '.png'
        
        # Lưu vào thư mục Downloads hoặc thư mục hiện tại
        downloads_path = Path.home() / "Downloads"
        if downloads_path.exists():
            filepath = downloads_path / filename
        else:
            filepath = Path(filename)
        
        # Chụp màn hình
        print(f"📸 [Screenshot] Đang chụp màn hình...")
        screenshot = pyautogui.screenshot()
        
        # Lưu file
        screenshot.save(str(filepath))
        
        file_size = filepath.stat().st_size / 1024  # KB
        
        print(f"✅ [Screenshot] Đã lưu: {filepath}")
        
        return {
            "success": True,
            "message": f"✅ Đã chụp màn hình: {filepath.name}",
            "filepath": str(filepath),
            "filename": filepath.name,
            "size_kb": round(file_size, 2),
            "dimensions": f"{screenshot.width}x{screenshot.height}"
        }
        
    except ImportError:
        return {
            "success": False,
            "error": "Thiếu thư viện 'pyautogui'. Cài đặt: pip install pyautogui"
        }
    except Exception as e:
        print(f"❌ [Screenshot] Error: {e}")
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}

async def show_notification(title: str, message: str) -> dict:
    try:
        ps_cmd = f'''[Windows.UI.Notifications.ToastNotificationManager, Windows.UI.Notifications, ContentType = WindowsRuntime] | Out-Null; [Windows.Data.Xml.Dom.XmlDocument, Windows.Data.Xml.Dom.XmlDocument, ContentType = WindowsRuntime] | Out-Null; $template = @"<toast><visual><binding template="ToastText02"><text id="1">{title}</text><text id="2">{message}</text></binding></visual></toast>"@; $xml = New-Object Windows.Data.Xml.Dom.XmlDocument; $xml.LoadXml($template); $toast = New-Object Windows.UI.Notifications.ToastNotification $xml; [Windows.UI.Notifications.ToastNotificationManager]::CreateToastNotifier("Xiaozhi").Show($toast)'''
        proc = await asyncio.create_subprocess_exec("powershell", "-Command", ps_cmd, stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.PIPE)
        await asyncio.wait_for(proc.wait(), timeout=5)
        return {"success": True, "title": title, "message": message}
    except Exception as e:
        return {"success": False, "error": str(e)}

# Cache cho system resources
_resource_cache = None
_resource_cache_time = 0
RESOURCE_CACHE_DURATION = 2  # Cache 2 giây

async def get_system_resources() -> dict:
    """Lấy thông tin tài nguyên hệ thống với caching"""
    global _resource_cache, _resource_cache_time
    
    try:
        # Kiểm tra cache
        now = time.time()
        if _resource_cache and (now - _resource_cache_time) < RESOURCE_CACHE_DURATION:
            return _resource_cache
        
        # Lấy dữ liệu mới - giảm interval từ 1s xuống 0.1s
        cpu = psutil.cpu_percent(interval=0.1)
        mem = psutil.virtual_memory()
        disk = psutil.disk_usage('/')
        
        result = {
            "success": True, 
            "data": {
                "cpu_percent": cpu, 
                "memory_percent": mem.percent, 
                "memory_used_gb": round(mem.used / (1024**3), 2), 
                "memory_total_gb": round(mem.total / (1024**3), 2), 
                "disk_percent": disk.percent, 
                "disk_used_gb": round(disk.used / (1024**3), 2), 
                "disk_total_gb": round(disk.total / (1024**3), 2)
            }
        }
        
        # Cập nhật cache
        _resource_cache = result
        _resource_cache_time = now
        
        return result
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_api_quotas() -> dict:
    """Lấy thông tin quota API (Gemini và Serper) - NOTE: Đây là giá trị ước tính"""
    try:
        result = {
            "success": True,
            "gemini": {
                "has_key": bool(GEMINI_API_KEY and GEMINI_API_KEY.strip()),
                "free_tier": "60 requests/min",
                "daily_limit": "1,500 requests/day",
                "note": "Free tier - chưa có API để check exact quota"
            },
            "serper": {
                "has_key": bool(SERPER_API_KEY and SERPER_API_KEY.strip()),
                "free_tier": "2,500 queries/month",
                "note": "Free tier - chưa có API để check exact remaining"
            }
        }
        return result
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_current_time() -> dict:
    try:
        now = datetime.now()
        return {"success": True, "datetime": now.strftime("%Y-%m-%d %H:%M:%S"), "date": now.strftime("%Y-%m-%d"), "time": now.strftime("%H:%M:%S"), "day_of_week": now.strftime("%A"), "timestamp": int(now.timestamp())}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def calculator(expression: str) -> dict:
    try:
        allowed = set("0123456789+-*/()., ")
        if not all(c in allowed for c in expression):
            return {"success": False, "error": "Ký tự không hợp lệ"}
        result = eval(expression, {"__builtins__": {}}, {})
        return {"success": True, "expression": expression, "result": result}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_network_info() -> dict:
    """
    Lấy thông tin mạng chi tiết bao gồm:
    - Thông tin máy local (hostname, IP, MAC, gateway)
    - Quét tất cả thiết bị đang kết nối với router
    - Hiển thị IP, MAC, hostname của từng thiết bị
    """
    try:
        import socket
        import subprocess
        import re
        from concurrent.futures import ThreadPoolExecutor
        
        # 1. Lấy thông tin máy local
        hostname = socket.gethostname()
        local_ip = socket.gethostbyname(hostname)
        
        # 2. Lấy MAC address và Gateway
        def get_mac_and_gateway():
            try:
                result = subprocess.check_output("ipconfig /all", shell=True, text=True, encoding='utf-8', errors='ignore')
                
                # Tìm gateway
                gateway_match = re.search(r'Default Gateway[.\s:]+([0-9]+\.[0-9]+\.[0-9]+\.[0-9]+)', result)
                gateway = gateway_match.group(1) if gateway_match else "Unknown"
                
                # Tìm MAC address của adapter đang kết nối
                mac_address = "Unknown"
                lines = result.split('\n')
                active_adapter = False
                for i, line in enumerate(lines):
                    if local_ip in line:
                        active_adapter = True
                    if active_adapter and 'Physical Address' in line:
                        mac_match = re.search(r'([0-9A-F]{2}[:-]){5}([0-9A-F]{2})', line, re.IGNORECASE)
                        if mac_match:
                            mac_address = mac_match.group(0)
                            break
                
                return mac_address, gateway
            except:
                return "Unknown", "Unknown"
        
        mac_address, gateway = get_mac_and_gateway()
        
        # 3. Quét thiết bị trong mạng (ARP table)
        def scan_network_devices():
            devices = []
            try:
                # Lấy ARP table
                arp_result = subprocess.check_output("arp -a", shell=True, text=True, encoding='utf-8', errors='ignore')
                
                # Parse ARP table
                lines = arp_result.split('\n')
                for line in lines:
                    # Tìm dòng có IP và MAC
                    match = re.search(r'([0-9]+\.[0-9]+\.[0-9]+\.[0-9]+)\s+([0-9a-f]{2}[:-]){5}([0-9a-f]{2})', line, re.IGNORECASE)
                    if match:
                        device_ip = match.group(1)
                        device_mac = match.group(0).split()[1] if len(match.group(0).split()) > 1 else "Unknown"
                        
                        # Bỏ qua broadcast/multicast
                        if device_ip.endswith('.255') or device_mac.startswith('ff-ff') or device_mac.startswith('01-00'):
                            continue
                        
                        # Thử resolve hostname (nhanh)
                        device_hostname = "Unknown"
                        try:
                            device_hostname = socket.gethostbyaddr(device_ip)[0]
                        except:
                            pass
                        
                        devices.append({
                            "ip": device_ip,
                            "mac": device_mac,
                            "hostname": device_hostname,
                            "is_local": device_ip == local_ip
                        })
                
                return devices
            except Exception as e:
                return []
        
        # Quét thiết bị (chạy async để không block)
        devices = scan_network_devices()
        
        # 4. Tổng hợp kết quả
        result = {
            "success": True,
            "local_device": {
                "hostname": hostname,
                "ip": local_ip,
                "mac": mac_address,
                "gateway": gateway
            },
            "network_devices": devices,
            "total_devices": len(devices),
            "message": f"Tìm thấy {len(devices)} thiết bị trong mạng"
        }
        
        return result
        
    except Exception as e:
        return {"success": False, "error": str(e)}

async def search_web(query: str) -> dict:
    try:
        import webbrowser
        url = f"https://www.google.com/search?q={query.replace(' ', '+')}"
        webbrowser.open(url)
        return {"success": True, "message": f"Đã mở tìm kiếm: {query}", "url": url}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def set_brightness(level: int) -> dict:
    try:
        import screen_brightness_control as sbc
        sbc.set_brightness(level)
        return {"success": True, "level": level, "message": f"Đã đặt độ sáng: {level}%"}
    except Exception as e:
        return {"success": False, "error": str(e), "note": "Có thể cần cài: pip install screen-brightness-control"}

async def get_clipboard() -> dict:
    try:
        import pyperclip
        content = pyperclip.paste()
        return {"success": True, "content": content}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def set_clipboard(text: str) -> dict:
    try:
        import pyperclip
        pyperclip.copy(text)
        return {"success": True, "message": f"Đã copy vào clipboard: {text[:50]}..."}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def play_sound(frequency: int = 1000, duration: int = 500) -> dict:
    try:
        import winsound
        winsound.Beep(frequency, duration)
        return {"success": True, "message": f"Đã phát âm thanh {frequency}Hz trong {duration}ms"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def open_application(app_name: str) -> dict:
    """
    Mở ứng dụng Windows với khả năng tìm kiếm thông minh.
    
    Thứ tự tìm kiếm:
    1. Dictionary mapping (ưu tiên cao nhất)
    2. Tìm trong PATH
    3. Tìm trong Registry (App Paths)
    4. Tìm trong Program Files
    5. Fallback: Windows Start Menu
    
    Args:
        app_name: Tên ứng dụng (ví dụ: "chrome", "photoshop", "word")
        
    Returns:
        dict: {"success": bool, "message": str, "path": str (optional)}
    """
    try:
        import os
        import shutil
        import winreg
        import glob
        
        # Dictionary mapping - Hỗ trợ 50+ ứng dụng phổ biến
        apps = {
            # Windows Built-in
            "notepad": "notepad.exe",
            "note": "notepad.exe",
            "máy ghi chú": "notepad.exe",
            "calc": "calc.exe",
            "calculator": "calc.exe",
            "máy tính": "calc.exe",
            "paint": "mspaint.exe",
            "vẽ": "mspaint.exe",
            "cmd": "cmd.exe",
            "command prompt": "cmd.exe",
            "powershell": "powershell.exe",
            "ps": "powershell.exe",
            "explorer": "explorer.exe",
            "file explorer": "explorer.exe",
            "taskmgr": "taskmgr.exe",
            "task manager": "taskmgr.exe",
            "quản lý tác vụ": "taskmgr.exe",
            
            # Browsers
            "chrome": "chrome.exe",
            "google chrome": "chrome.exe",
            "gc": "chrome.exe",
            "firefox": "firefox.exe",
            "ff": "firefox.exe",
            "edge": "msedge.exe",
            "microsoft edge": "msedge.exe",
            "brave": "brave.exe",
            "opera": "opera.exe",
            
            # Microsoft Office
            "word": "WINWORD.EXE",
            "microsoft word": "WINWORD.EXE",
            "excel": "EXCEL.EXE",
            "microsoft excel": "EXCEL.EXE",
            "powerpoint": "POWERPNT.EXE",
            "microsoft powerpoint": "POWERPNT.EXE",
            "ppt": "POWERPNT.EXE",
            "outlook": "OUTLOOK.EXE",
            "microsoft outlook": "OUTLOOK.EXE",
            "onenote": "ONENOTE.EXE",
            "teams": "Teams.exe",
            "microsoft teams": "Teams.exe",
            
            # Adobe Creative Cloud
            "photoshop": "Photoshop.exe",
            "adobe photoshop": "Photoshop.exe",
            "ps": "Photoshop.exe",
            "illustrator": "Illustrator.exe",
            "adobe illustrator": "Illustrator.exe",
            "ai": "Illustrator.exe",
            "premiere": "Adobe Premiere Pro.exe",
            "premiere pro": "Adobe Premiere Pro.exe",
            "after effects": "AfterFX.exe",
            "ae": "AfterFX.exe",
            "lightroom": "Lightroom.exe",
            "acrobat": "Acrobat.exe",
            "adobe acrobat": "Acrobat.exe",
            
            # Development Tools
            "vscode": "Code.exe",
            "visual studio code": "Code.exe",
            "code": "Code.exe",
            "vs": "Code.exe",
            "sublime": "sublime_text.exe",
            "sublime text": "sublime_text.exe",
            "atom": "atom.exe",
            "notepad++": "notepad++.exe",
            "npp": "notepad++.exe",
            "pycharm": "pycharm64.exe",
            "intellij": "idea64.exe",
            "webstorm": "webstorm64.exe",
            "androidstudio": "studio64.exe",
            "android studio": "studio64.exe",
            
            # 3D & Design
            "blender": "blender.exe",
            "3ds max": "3dsmax.exe",
            "maya": "maya.exe",
            "sketchup": "SketchUp.exe",
            "fusion360": "Fusion360.exe",
            "fusion 360": "Fusion360.exe",
            "autocad": "acad.exe",
            "solidworks": "SLDWORKS.exe",
            
            # Communication
            "discord": "Discord.exe",
            "slack": "slack.exe",
            "zoom": "Zoom.exe",
            "skype": "Skype.exe",
            "telegram": "Telegram.exe",
            "zalo": "Zalo.exe",
            
            # Media Players
            "vlc": "vlc.exe",
            "spotify": "Spotify.exe",
            "itunes": "iTunes.exe",
            "windows media player": "wmplayer.exe",
            "wmp": "wmplayer.exe",
            
            # Other Popular Apps
            "steam": "steam.exe",
            "epic games": "EpicGamesLauncher.exe",
            "epic": "EpicGamesLauncher.exe",
            "obs": "obs64.exe",
            "obs studio": "obs64.exe",
            "gimp": "gimp-2.10.exe",
            "audacity": "audacity.exe",
            "7zip": "7zFM.exe",
            "7-zip": "7zFM.exe",
            "winrar": "WinRAR.exe",
        }
        
        # 1. Kiểm tra trong dictionary
        app_name_lower = app_name.lower().strip()
        exe_name = apps.get(app_name_lower)
        
        print(f"🔍 [Open App] Tìm kiếm: '{app_name}' → {exe_name or 'không có trong dictionary'}")
        
        # Nếu không có trong dictionary, thử dùng tên gốc
        if not exe_name:
            # Kiểm tra nếu đã có .exe
            if app_name.lower().endswith('.exe'):
                exe_name = app_name
            else:
                exe_name = app_name + '.exe'
        
        # 2. Tìm trong PATH
        exe_path = shutil.which(exe_name)
        if exe_path:
            print(f"✅ [Open App] Tìm thấy trong PATH: {exe_path}")
            subprocess.Popen([exe_path])
            return {"success": True, "message": f"✅ Đã mở {app_name}", "path": exe_path}
        
        # 3. Tìm trong Windows Registry (App Paths)
        try:
            with winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, 
                              rf"SOFTWARE\Microsoft\Windows\CurrentVersion\App Paths\{exe_name}") as key:
                exe_path = winreg.QueryValue(key, None)
                if exe_path and os.path.exists(exe_path):
                    print(f"✅ [Open App] Tìm thấy trong Registry: {exe_path}")
                    subprocess.Popen([exe_path])
                    return {"success": True, "message": f"✅ Đã mở {app_name}", "path": exe_path}
        except WindowsError:
            pass
        
        # 4. Tìm trong các thư mục phổ biến
        common_paths = [
            os.path.join(os.environ.get("PROGRAMFILES", "C:\\Program Files"), "*", exe_name),
            os.path.join(os.environ.get("PROGRAMFILES(X86)", "C:\\Program Files (x86)"), "*", exe_name),
            os.path.join(os.environ.get("LOCALAPPDATA", ""), "Programs", "*", exe_name),
            os.path.join(os.environ.get("APPDATA", ""), "*", exe_name),
        ]
        
        import glob
        for pattern in common_paths:
            matches = glob.glob(pattern, recursive=False)
            if matches:
                exe_path = matches[0]
                print(f"✅ [Open App] Tìm thấy trong: {exe_path}")
                subprocess.Popen([exe_path])
                return {"success": True, "message": f"✅ Đã mở {app_name}", "path": exe_path}
        
        # 5. Tìm kiếm sâu trong Program Files (chậm hơn, dùng làm fallback)
        if "photoshop" in app_name_lower or "adobe" in app_name_lower:
            # Adobe apps thường ở C:\Program Files\Adobe
            adobe_base = r"C:\Program Files\Adobe"
            if os.path.exists(adobe_base):
                for root, dirs, files in os.walk(adobe_base):
                    if exe_name in files:
                        exe_path = os.path.join(root, exe_name)
                        print(f"✅ [Open App] Tìm thấy Adobe app: {exe_path}")
                        subprocess.Popen([exe_path])
                        return {"success": True, "message": f"✅ Đã mở {app_name}", "path": exe_path}
        
        if "autodesk" in app_name_lower or "fusion" in app_name_lower:
            # Autodesk apps thường ở LOCALAPPDATA
            autodesk_base = os.path.join(os.environ.get("LOCALAPPDATA", ""), "Autodesk")
            if os.path.exists(autodesk_base):
                for root, dirs, files in os.walk(autodesk_base):
                    if exe_name in files:
                        exe_path = os.path.join(root, exe_name)
                        print(f"✅ [Open App] Tìm thấy Autodesk app: {exe_path}")
                        subprocess.Popen([exe_path])
                        return {"success": True, "message": f"✅ Đã mở {app_name}", "path": exe_path}
        
        # 6. Fallback cuối cùng: Dùng Windows Start Menu
        print(f"⚠️ [Open App] Không tìm thấy đường dẫn, thử Windows Start Menu...")
        subprocess.Popen(["start", "", app_name], shell=True)
        return {
            "success": True, 
            "message": f"✅ Đã gửi lệnh mở {app_name} (Windows sẽ tìm trong Start Menu)",
            "note": "Nếu không mở được, hãy kiểm tra tên ứng dụng hoặc thêm vào dictionary"
        }
        
    except Exception as e:
        print(f"❌ [Open App] Lỗi: {e}")
        import traceback
        traceback.print_exc()
        return {"success": False, "error": f"Lỗi khi mở {app_name}: {str(e)}"}

# ==================== MEDIA PLAYER CONTROL ====================

# Helper function để tìm tất cả các cửa sổ media player và browser
def _find_all_media_windows():
    """Tìm tất cả cửa sổ media player và browser đang chạy"""
    import ctypes
    
    windows = {
        'youtube': [],      # Các tab YouTube
        'spotify_web': [],  # Spotify web
        'wmplayer': None,   # Windows Media Player
        'vlc': None,        # VLC Player
        'spotify_app': None,# Spotify Desktop
        'browsers': []      # Các browser khác
    }
    
    browser_names = ['chrome', 'firefox', 'edge', 'opera', 'brave', 'coccoc', 'cốc cốc']
    
    def enum_callback(hwnd, _):
        if ctypes.windll.user32.IsWindowVisible(hwnd):
            length = ctypes.windll.user32.GetWindowTextLengthW(hwnd)
            if length > 0:
                buff = ctypes.create_unicode_buffer(length + 1)
                ctypes.windll.user32.GetWindowTextW(hwnd, buff, length + 1)
                title = buff.value
                title_lower = title.lower()
                
                # YouTube có ưu tiên cao nhất
                if 'youtube' in title_lower:
                    windows['youtube'].append({'hwnd': hwnd, 'title': title})
                # Spotify Web
                elif 'spotify' in title_lower and any(b in title_lower for b in browser_names):
                    windows['spotify_web'].append({'hwnd': hwnd, 'title': title})
                # Windows Media Player
                elif 'windows media player' in title_lower or 'wmplayer' in title_lower:
                    windows['wmplayer'] = {'hwnd': hwnd, 'title': title}
                # VLC
                elif 'vlc' in title_lower and 'media player' in title_lower:
                    windows['vlc'] = {'hwnd': hwnd, 'title': title}
                # Spotify Desktop App
                elif 'spotify' in title_lower and not any(b in title_lower for b in browser_names):
                    windows['spotify_app'] = {'hwnd': hwnd, 'title': title}
                # Các browser khác
                elif any(b in title_lower for b in browser_names):
                    windows['browsers'].append({'hwnd': hwnd, 'title': title})
        return True
    
    WNDENUMPROC = ctypes.WINFUNCTYPE(ctypes.c_bool, ctypes.c_int, ctypes.POINTER(ctypes.c_int))
    ctypes.windll.user32.EnumWindows(WNDENUMPROC(enum_callback), 0)
    
    return windows

def _focus_and_send_key(hwnd, key, delay=0.15):
    """Focus vào cửa sổ và gửi phím"""
    import ctypes
    ctypes.windll.user32.SetForegroundWindow(hwnd)
    time.sleep(delay)
    pyautogui.press(key)

def _focus_and_send_hotkey(hwnd, *keys, delay=0.15):
    """Focus vào cửa sổ và gửi tổ hợp phím"""
    import ctypes
    ctypes.windll.user32.SetForegroundWindow(hwnd)
    time.sleep(delay)
    pyautogui.hotkey(*keys)

async def media_play_pause() -> dict:
    """
    Phát/Tạm dừng media (Play/Pause toggle).
    ⭐ ƯU TIÊN PYTHON-VLC TRƯỚC - nhanh & không cần detect window!
    
    Ưu tiên:
    1. Python-VLC nội bộ (NHANH NHẤT)
    2. YouTube (Browser) - Focus và nhấn K
    3. Windows Media Player
    4. Spotify
    5. Fallback - Media key
    """
    try:
        # 🎵 ƯU TIÊN 1: Python-VLC nội bộ - NHANH NHẤT!
        if vlc_player and vlc_player._player:
            vlc_player.pause()
            is_playing = vlc_player.is_playing()
            status = vlc_player.get_full_status()
            current_song = status.get('current_song', 'Unknown')
            return {
                "success": True, 
                "message": f"{'▶️ Đang phát' if is_playing else '⏸️ Đã tạm dừng'}: {current_song} (Python-VLC)",
                "is_playing": is_playing,
                "player": "Python-VLC",
                "llm_note": "🎵 Đang dùng Python-VLC Player tích hợp. Có thể dùng: pause_music(), resume_music(), stop_music(), music_next(), music_previous(), seek_music(), music_volume()"
            }
        
        windows = _find_all_media_windows()
        
        # 2. YouTube - nếu có
        if windows['youtube']:
            yt = windows['youtube'][0]
            _focus_and_send_key(yt['hwnd'], 'k')
            return {"success": True, "message": f"✅ Play/Pause YouTube: {yt['title'][:50]}..."}
        
        # 3. Windows Media Player
        if windows['wmplayer']:
            _focus_and_send_key(windows['wmplayer']['hwnd'], 'space')
            return {"success": True, "message": "✅ Play/Pause (Windows Media Player)"}
        
        # 4. VLC Window (external)
        if windows['vlc']:
            _focus_and_send_key(windows['vlc']['hwnd'], 'space')
            return {"success": True, "message": "✅ Play/Pause (VLC Window)"}
        
        # 5. Spotify Desktop App
        if windows['spotify_app']:
            _focus_and_send_key(windows['spotify_app']['hwnd'], 'space')
            return {"success": True, "message": "✅ Play/Pause (Spotify Desktop)"}
        
        # 6. Spotify Web
        if windows['spotify_web']:
            sw = windows['spotify_web'][0]
            _focus_and_send_key(sw['hwnd'], 'space')
            return {"success": True, "message": f"✅ Play/Pause Spotify Web"}
        
        # 7. Fallback - dùng media key
        pyautogui.press('playpause')
        return {"success": True, "message": "✅ Đã gửi lệnh Play/Pause (Media Key)"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def media_next_track() -> dict:
    """
    Chuyển bài tiếp theo (Next Track).
    ⭐ ƯU TIÊN PYTHON-VLC TRƯỚC - nhanh & không cần detect window!
    """
    try:
        # 🎵 ƯU TIÊN 1: Python-VLC nội bộ - NHANH NHẤT!
        if vlc_player and vlc_player._player:
            success = vlc_player.next_track()
            if success:
                import time
                time.sleep(0.3)  # Đợi VLC chuyển bài
                status = vlc_player.get_full_status()
                current_song = status.get('current_song', 'Unknown')
                return {
                    "success": True, 
                    "message": f"⏭️ Đã chuyển: {current_song} (Python-VLC)",
                    "player": "Python-VLC",
                    "current_song": current_song,
                    "llm_note": "🎵 Đang dùng Python-VLC Player. Playlist có thể điều khiển bằng music_next(), music_previous()"
                }
            return {"success": False, "error": "Không có bài tiếp theo trong playlist VLC"}
        
        windows = _find_all_media_windows()
        
        # 2. YouTube
        if windows['youtube']:
            yt = windows['youtube'][0]
            _focus_and_send_hotkey(yt['hwnd'], 'shift', 'n')
            return {"success": True, "message": f"✅ Chuyển video tiếp theo (YouTube): {yt['title'][:40]}..."}
        
        # 3. Windows Media Player
        if windows['wmplayer']:
            _focus_and_send_hotkey(windows['wmplayer']['hwnd'], 'ctrl', 'f')
            return {"success": True, "message": "✅ Chuyển bài tiếp theo (Windows Media Player)"}
        
        # 4. VLC Window (external)
        if windows['vlc']:
            _focus_and_send_key(windows['vlc']['hwnd'], 'n')
            return {"success": True, "message": "✅ Chuyển bài tiếp theo (VLC Window)"}
        
        # 5. Spotify Desktop App
        if windows['spotify_app']:
            _focus_and_send_hotkey(windows['spotify_app']['hwnd'], 'ctrl', 'right')
            return {"success": True, "message": "✅ Chuyển bài tiếp theo (Spotify Desktop)"}
        
        # 6. Spotify Web
        if windows['spotify_web']:
            sw = windows['spotify_web'][0]
            _focus_and_send_hotkey(sw['hwnd'], 'ctrl', 'right')
            return {"success": True, "message": "✅ Chuyển bài tiếp theo (Spotify Web)"}
        
        # 7. Fallback - dùng media key
        pyautogui.press('nexttrack')
        return {"success": True, "message": "✅ Đã chuyển bài tiếp theo (Media Key)"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def media_previous_track() -> dict:
    """
    Chuyển bài trước đó (Previous Track).
    ⭐ ƯU TIÊN PYTHON-VLC TRƯỚC - nhanh & không cần detect window!
    """
    try:
        # 🎵 ƯU TIÊN 1: Python-VLC nội bộ - NHANH NHẤT!
        if vlc_player and vlc_player._player:
            success = vlc_player.previous_track()
            if success:
                import time
                time.sleep(0.3)  # Đợi VLC chuyển bài
                status = vlc_player.get_full_status()
                current_song = status.get('current_song', 'Unknown')
                return {
                    "success": True, 
                    "message": f"⏮️ Đã quay lại: {current_song} (Python-VLC)",
                    "player": "Python-VLC",
                    "current_song": current_song,
                    "llm_note": "🎵 Đang dùng Python-VLC Player. Playlist có thể điều khiển bằng music_next(), music_previous()"
                }
            return {"success": False, "error": "Không có bài trước trong playlist VLC"}
        
        windows = _find_all_media_windows()
        
        # 2. YouTube
        if windows['youtube']:
            yt = windows['youtube'][0]
            _focus_and_send_hotkey(yt['hwnd'], 'shift', 'p')
            return {"success": True, "message": f"✅ Chuyển video trước (YouTube): {yt['title'][:40]}..."}
        
        # 3. Windows Media Player
        if windows['wmplayer']:
            _focus_and_send_hotkey(windows['wmplayer']['hwnd'], 'ctrl', 'b')
            return {"success": True, "message": "✅ Chuyển bài trước (Windows Media Player)"}
        
        # 4. VLC Window (external)
        if windows['vlc']:
            _focus_and_send_key(windows['vlc']['hwnd'], 'p')
            return {"success": True, "message": "✅ Chuyển bài trước (VLC Window)"}
        
        # 5. Spotify Desktop App
        if windows['spotify_app']:
            _focus_and_send_hotkey(windows['spotify_app']['hwnd'], 'ctrl', 'left')
            return {"success": True, "message": "✅ Chuyển bài trước (Spotify Desktop)"}
        
        # 6. Spotify Web
        if windows['spotify_web']:
            sw = windows['spotify_web'][0]
            _focus_and_send_hotkey(sw['hwnd'], 'ctrl', 'left')
            return {"success": True, "message": "✅ Chuyển bài trước (Spotify Web)"}
        
        # 7. Fallback - dùng media key
        pyautogui.press('prevtrack')
        return {"success": True, "message": "✅ Đã chuyển bài trước (Media Key)"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def media_stop() -> dict:
    """
    Dừng phát media (Stop).
    ⭐ ƯU TIÊN PYTHON-VLC TRƯỚC - nhanh & không cần detect window!
    """
    try:
        # 🎵 ƯU TIÊN 1: Python-VLC nội bộ - NHANH NHẤT!
        if vlc_player and vlc_player._player:
            vlc_player.stop()
            return {
                "success": True, 
                "message": "⏹️ Đã dừng nhạc (Python-VLC)",
                "player": "Python-VLC",
                "llm_note": "🎵 Đã dừng Python-VLC Player. Dùng play_music() hoặc resume_music() để phát lại."
            }
        
        windows = _find_all_media_windows()
        
        # 2. YouTube
        if windows['youtube']:
            yt = windows['youtube'][0]
            _focus_and_send_key(yt['hwnd'], 'k', delay=0.2)
            return {"success": True, "message": f"✅ Đã dừng YouTube: {yt['title'][:50]}..."}
        
        # 3. Windows Media Player
        if windows['wmplayer']:
            _focus_and_send_key(windows['wmplayer']['hwnd'], 'stop')
            return {"success": True, "message": "✅ Đã dừng phát (Windows Media Player)"}
        
        # 4. VLC Window (external)
        if windows['vlc']:
            _focus_and_send_key(windows['vlc']['hwnd'], 's')
            return {"success": True, "message": "✅ Đã dừng phát (VLC Window)"}
        
        # 5. Spotify Desktop App - không có stop, dùng pause
        if windows['spotify_app']:
            _focus_and_send_key(windows['spotify_app']['hwnd'], 'space')
            return {"success": True, "message": "✅ Đã tạm dừng (Spotify Desktop)"}
        
        # 6. Spotify Web
        if windows['spotify_web']:
            sw = windows['spotify_web'][0]
            _focus_and_send_key(sw['hwnd'], 'space')
            return {"success": True, "message": "✅ Đã tạm dừng (Spotify Web)"}
        
        # 7. Fallback - dùng media key
        pyautogui.press('stop')
        return {"success": True, "message": "✅ Đã dừng phát (Media Key)"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def media_volume_up() -> dict:
    """Tăng âm lượng media (Media Volume Up)"""
    try:
        pyautogui.press('volumeup')
        return {"success": True, "message": "✅ Đã tăng âm lượng"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def media_volume_down() -> dict:
    """Giảm âm lượng media (Media Volume Down)"""
    try:
        pyautogui.press('volumedown')
        return {"success": True, "message": "✅ Đã giảm âm lượng"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def media_mute() -> dict:
    """Tắt/Bật tiếng media (Mute Toggle)"""
    try:
        pyautogui.press('volumemute')
        return {"success": True, "message": "✅ Đã toggle mute"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def media_control(action: str) -> dict:
    """
    Điều khiển media player đa năng.
    
    Args:
        action: Hành động cần thực hiện
            - "play" hoặc "pause": Phát/Tạm dừng
            - "next": Bài tiếp theo
            - "previous" hoặc "prev": Bài trước
            - "stop": Dừng phát
            - "volume_up": Tăng âm lượng
            - "volume_down": Giảm âm lượng
            - "mute": Tắt/Bật tiếng
    
    Returns:
        dict: Kết quả thực hiện
    """
    try:
        action = action.lower().strip()
        
        actions_map = {
            "play": "playpause",
            "pause": "playpause",
            "playpause": "playpause",
            "next": "nexttrack",
            "previous": "prevtrack",
            "prev": "prevtrack",
            "stop": "stop",
            "volume_up": "volumeup",
            "volumeup": "volumeup",
            "volume_down": "volumedown",
            "volumedown": "volumedown",
            "mute": "volumemute",
        }
        
        key = actions_map.get(action)
        if not key:
            return {
                "success": False, 
                "error": f"Action không hợp lệ: '{action}'. Chọn: play, pause, next, previous, stop, volume_up, volume_down, mute"
            }
        
        pyautogui.press(key)
        
        action_messages = {
            "playpause": "Play/Pause",
            "nexttrack": "Bài tiếp theo",
            "prevtrack": "Bài trước",
            "stop": "Dừng phát",
            "volumeup": "Tăng âm lượng",
            "volumedown": "Giảm âm lượng",
            "volumemute": "Mute/Unmute",
        }
        
        return {"success": True, "message": f"✅ {action_messages[key]}", "action": action}
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}

# ==================== END MEDIA PLAYER CONTROL ====================

# ==================== TASK MEMORY TOOLS ====================

async def remember_task(tool_name: str, params: dict = None, result_message: str = "", user_request: str = "") -> dict:
    """
    Ghi nhớ một tác vụ đã thực hiện vào bộ nhớ.
    Giúp AI phản hồi nhanh và chính xác hơn cho các yêu cầu tương tự.
    
    Args:
        tool_name: Tên tool đã sử dụng
        params: Tham số đã truyền vào tool
        result_message: Kết quả/message trả về
        user_request: Yêu cầu gốc của user
    """
    try:
        task_entry = add_task_to_memory(
            tool_name=tool_name,
            params=params or {},
            result={"success": True, "message": result_message},
            user_request=user_request
        )
        return {
            "success": True,
            "message": f"✅ Đã ghi nhớ tác vụ: {tool_name}",
            "task": task_entry
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

async def recall_tasks(keyword: str = "", limit: int = 10) -> dict:
    """
    Nhớ lại các tác vụ đã thực hiện trước đó.
    Giúp AI biết những gì đã làm để phản hồi phù hợp.
    
    Args:
        keyword: Từ khóa tìm kiếm (optional). Để trống = lấy tác vụ gần nhất
        limit: Số lượng tác vụ tối đa trả về (default 10)
    """
    try:
        if keyword:
            tasks = search_task_memory(keyword)
            message = f"🔍 Tìm thấy {len(tasks)} tác vụ liên quan đến '{keyword}'"
        else:
            tasks = get_recent_tasks(limit)
            message = f"📋 {len(tasks)} tác vụ gần đây nhất"
        
        return {
            "success": True,
            "message": message,
            "count": len(tasks),
            "tasks": tasks
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_task_summary() -> dict:
    """
    Lấy tổng hợp thống kê về các tác vụ đã thực hiện.
    Giúp AI hiểu patterns sử dụng của user.
    """
    try:
        tasks = load_task_memory()
        
        if not tasks:
            return {
                "success": True,
                "message": "📊 Chưa có lịch sử tác vụ",
                "total_tasks": 0,
                "most_used_tools": [],
                "success_rate": 0
            }
        
        # Đếm theo tool
        tool_counts = {}
        success_count = 0
        
        for task in tasks:
            tool = task.get('tool', 'unknown')
            tool_counts[tool] = tool_counts.get(tool, 0) + 1
            if task.get('result_success'):
                success_count += 1
        
        # Top 10 tools được dùng nhiều nhất
        sorted_tools = sorted(tool_counts.items(), key=lambda x: x[1], reverse=True)[:10]
        
        return {
            "success": True,
            "message": f"📊 Đã thực hiện {len(tasks)} tác vụ",
            "total_tasks": len(tasks),
            "most_used_tools": [{"tool": t[0], "count": t[1]} for t in sorted_tools],
            "success_rate": round(success_count / len(tasks) * 100, 1),
            "recent_tools": [t.get('tool') for t in tasks[-5:]]
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

async def forget_all_tasks() -> dict:
    """
    Xóa toàn bộ lịch sử tác vụ đã ghi nhớ.
    """
    try:
        success = clear_task_memory()
        if success:
            return {"success": True, "message": "🗑️ Đã xóa toàn bộ lịch sử tác vụ"}
        else:
            return {"success": False, "error": "Không thể xóa lịch sử"}
    except Exception as e:
        return {"success": False, "error": str(e)}

# ==================== END TASK MEMORY TOOLS ====================

async def get_active_media_players() -> dict:
    """
    Lấy danh sách các media players/applications đang chạy trên máy tính.
    
    Thông tin này giúp LLM biết:
    - Có media player nào đang chạy không
    - Nên dùng tool nào (media_play_pause cho Spotify/VLC, stop_music cho WMP)
    - Có ứng dụng nào có thể điều khiển được
    
    Returns:
        dict: Danh sách media players, browsers, và ứng dụng quan trọng đang chạy
    """
    try:
        # Danh sách media players và ứng dụng quan trọng cần theo dõi
        MEDIA_APPS = {
            # Media Players
            "spotify.exe": {"name": "Spotify", "type": "music", "supports_media_keys": True},
            "vlc.exe": {"name": "VLC Media Player", "type": "video", "supports_media_keys": True},
            "wmplayer.exe": {"name": "Windows Media Player", "type": "music", "supports_media_keys": True},
            "itunes.exe": {"name": "iTunes", "type": "music", "supports_media_keys": True},
            
            # Browsers (có thể phát YouTube, Spotify Web...)
            "chrome.exe": {"name": "Google Chrome", "type": "browser", "supports_media_keys": True},
            "msedge.exe": {"name": "Microsoft Edge", "type": "browser", "supports_media_keys": True},
            "firefox.exe": {"name": "Firefox", "type": "browser", "supports_media_keys": True},
            "brave.exe": {"name": "Brave", "type": "browser", "supports_media_keys": True},
            "opera.exe": {"name": "Opera", "type": "browser", "supports_media_keys": True},
            "browser.exe": {"name": "Browser", "type": "browser", "supports_media_keys": True},
            "iexplore.exe": {"name": "Internet Explorer", "type": "browser", "supports_media_keys": True},
            "vivaldi.exe": {"name": "Vivaldi", "type": "browser", "supports_media_keys": True},
            
            # Communication (có media playback)
            "discord.exe": {"name": "Discord", "type": "communication", "supports_media_keys": True},
            "slack.exe": {"name": "Slack", "type": "communication", "supports_media_keys": False},
            "zoom.exe": {"name": "Zoom", "type": "communication", "supports_media_keys": False},
            "skype.exe": {"name": "Skype", "type": "communication", "supports_media_keys": False},
            
            # Office & Productivity
            "WINWORD.EXE": {"name": "Microsoft Word", "type": "office", "supports_media_keys": False},
            "EXCEL.EXE": {"name": "Microsoft Excel", "type": "office", "supports_media_keys": False},
            "POWERPNT.EXE": {"name": "PowerPoint", "type": "office", "supports_media_keys": False},
            "OUTLOOK.EXE": {"name": "Outlook", "type": "office", "supports_media_keys": False},
            
            # Development
            "Code.exe": {"name": "VS Code", "type": "development", "supports_media_keys": False},
            "devenv.exe": {"name": "Visual Studio", "type": "development", "supports_media_keys": False},
            "pycharm64.exe": {"name": "PyCharm", "type": "development", "supports_media_keys": False},
            
            # Design & Creative
            "Photoshop.exe": {"name": "Adobe Photoshop", "type": "creative", "supports_media_keys": False},
            "Illustrator.exe": {"name": "Adobe Illustrator", "type": "creative", "supports_media_keys": False},
            "blender.exe": {"name": "Blender", "type": "3d", "supports_media_keys": False},
        }
        
        running_apps = []
        media_players = []
        browsers = []
        
        # Quét các process đang chạy
        for proc in psutil.process_iter(['pid', 'name']):
            try:
                proc_name = proc.info['name']
                
                if proc_name in MEDIA_APPS:
                    app_info = MEDIA_APPS[proc_name].copy()
                    app_info['pid'] = proc.info['pid']
                    app_info['process_name'] = proc_name
                    
                    running_apps.append(app_info)
                    
                    # Phân loại
                    if app_info['type'] in ['music', 'video']:
                        media_players.append(app_info)
                    elif app_info['type'] == 'browser':
                        browsers.append(app_info)
                        
            except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
                pass
        
        # Tạo thông điệp hướng dẫn cho LLM (tối ưu hóa, không liệt kê từng instance)
        guidance = ""
        
        if media_players:
            # Đếm số lượng từng loại media player (không liệt kê từng process)
            player_counts = {}
            for p in media_players:
                name = p['name']
                player_counts[name] = player_counts.get(name, 0) + 1
            
            player_summary = ', '.join([f"{name} ({count})" if count > 1 else name 
                                       for name, count in player_counts.items()])
            guidance += f"🎵 Media Players: {player_summary}.\n"
            
            if any(p['name'] == 'Windows Media Player' for p in media_players):
                guidance += "   → Dùng stop_music() để dừng Windows Media Player.\n"
            
            if any(p['supports_media_keys'] and p['name'] != 'Windows Media Player' for p in media_players):
                guidance += "   → Dùng media_play_pause(), media_next_track() cho Spotify/VLC/iTunes.\n"
        
        if browsers:
            # Đếm số lượng từng loại browser (không liệt kê từng process)
            browser_counts = {}
            for b in browsers:
                name = b['name']
                browser_counts[name] = browser_counts.get(name, 0) + 1
            
            browser_summary = ', '.join([f"{name} ({count})" if count > 1 else name 
                                        for name, count in browser_counts.items()])
            guidance += f"🌐 Browsers: {browser_summary}.\n"
            guidance += "   → Có thể phát YouTube/Spotify Web. Dùng media_play_pause() để điều khiển.\n"
        
        if not media_players and not browsers:
            guidance = "❌ Không có media player/browser nào đang chạy. Dùng play_music() để phát nhạc từ music_library."
        
        return {
            "success": True,
            "all_apps": running_apps,
            "media_players": media_players,
            "browsers": browsers,
            "total_count": len(running_apps),
            "guidance": guidance.strip(),
            "message": f"✅ Đang chạy: {len(running_apps)} ứng dụng ({len(media_players)} media players, {len(browsers)} browsers)"
        }
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}

async def list_running_processes(limit: int = 10) -> dict:
    try:
        procs = []
        for p in psutil.process_iter(['pid', 'name', 'cpu_percent', 'memory_percent']):
            try:
                procs.append({"pid": p.info['pid'], "name": p.info['name'], "cpu": round(p.info['cpu_percent'], 2), "memory": round(p.info['memory_percent'], 2)})
            except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
                # Bỏ qua các tiến trình không thể truy cập
                pass
        procs = sorted(procs, key=lambda x: x['cpu'], reverse=True)[:limit]
        return {"success": True, "processes": procs, "count": len(procs)}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def kill_process(identifier: str, force: bool = True, exact_match: bool = False) -> dict:
    """
    Kill process ngay lập tức.
    
    Args:
        identifier: Tên app hoặc PID. VD: "notepad", "chrome", "1234"
        force: True = kill ngay (SIGKILL), False = đóng mềm (SIGTERM)
        exact_match: True = tên phải khớp chính xác, False = chứa tên là được
    """
    import subprocess
    import time
    
    try:
        killed = []
        failed = []
        
        # Nếu là PID (số)
        if identifier.isdigit():
            try:
                p = psutil.Process(int(identifier))
                name = p.name()
                if force:
                    p.kill()  # SIGKILL - kill ngay lập tức
                else:
                    p.terminate()  # SIGTERM - đóng mềm
                    p.wait(timeout=3)  # Chờ tối đa 3 giây
                killed.append(f"{name} (PID: {identifier})")
            except psutil.TimeoutExpired:
                # Nếu terminate không được, force kill
                p.kill()
                killed.append(f"{name} (PID: {identifier}) [FORCE KILLED]")
        else:
            # Tìm theo tên
            target_name = identifier.lower()
            
            # Thêm .exe nếu chưa có
            if not target_name.endswith('.exe'):
                target_name_exe = target_name + '.exe'
            else:
                target_name_exe = target_name
                target_name = target_name[:-4]  # Bỏ .exe để so sánh
            
            for p in psutil.process_iter(['pid', 'name', 'exe']):
                try:
                    proc_name = p.info['name'].lower() if p.info['name'] else ""
                    
                    # Kiểm tra match
                    match = False
                    if exact_match:
                        # Khớp chính xác tên
                        match = (proc_name == target_name_exe or proc_name == target_name)
                    else:
                        # Chứa tên là được
                        match = (target_name in proc_name)
                    
                    if match:
                        pid = p.info['pid']
                        try:
                            if force:
                                p.kill()  # Kill ngay lập tức
                            else:
                                p.terminate()
                                try:
                                    p.wait(timeout=2)
                                except psutil.TimeoutExpired:
                                    p.kill()  # Force kill nếu không đóng được
                            killed.append(f"{p.info['name']} (PID: {pid})")
                        except psutil.AccessDenied:
                            # Thử dùng taskkill với quyền cao hơn
                            try:
                                subprocess.run(
                                    ['taskkill', '/F', '/PID', str(pid)],
                                    capture_output=True,
                                    timeout=5
                                )
                                killed.append(f"{p.info['name']} (PID: {pid}) [via taskkill]")
                            except:
                                failed.append(f"{p.info['name']} (PID: {pid}) - Access Denied")
                except (psutil.NoSuchProcess, psutil.ZombieProcess):
                    pass
        
        # Kết quả
        if killed:
            result = {
                "success": True, 
                "message": f"✅ Đã kill thành công: {', '.join(killed)}",
                "killed_count": len(killed),
                "killed": killed
            }
            if failed:
                result["failed"] = failed
                result["message"] += f"\n⚠️ Không thể kill: {', '.join(failed)}"
            return result
        elif failed:
            return {"success": False, "error": f"Không có quyền kill: {', '.join(failed)}"}
        else:
            return {"success": False, "error": f"Không tìm thấy process '{identifier}'"}
            
    except psutil.NoSuchProcess:
        return {"success": False, "error": f"Tiến trình không tồn tại: {identifier}"}
    except psutil.AccessDenied:
        # Thử dùng taskkill
        try:
            if identifier.isdigit():
                result = subprocess.run(
                    ['taskkill', '/F', '/PID', identifier],
                    capture_output=True,
                    text=True,
                    timeout=10
                )
            else:
                result = subprocess.run(
                    ['taskkill', '/F', '/IM', f'{identifier}*'],
                    capture_output=True,
                    text=True,
                    timeout=10
                )
            if result.returncode == 0:
                return {"success": True, "message": f"✅ Đã kill bằng taskkill: {identifier}"}
            else:
                return {"success": False, "error": f"Không thể kill (cần quyền Admin): {identifier}"}
        except Exception as e:
            return {"success": False, "error": f"Lỗi khi kill: {str(e)}"}
    except Exception as e:
        return {"success": False, "error": str(e)}


async def force_kill_app(app_name: str) -> dict:
    """
    Force kill app theo tên CHÍNH XÁC - kill ngay lập tức không hỏi han.
    Sử dụng cả psutil và taskkill để đảm bảo kill được.
    
    Args:
        app_name: Tên app cần kill. VD: "notepad", "chrome", "Code"
    """
    import subprocess
    
    try:
        killed = []
        
        # Chuẩn hóa tên
        target = app_name.lower().strip()
        if not target.endswith('.exe'):
            target_exe = target + '.exe'
        else:
            target_exe = target
            target = target[:-4]
        
        # Bước 1: Kill bằng psutil
        for p in psutil.process_iter(['pid', 'name']):
            try:
                proc_name = (p.info['name'] or "").lower()
                if proc_name == target_exe or proc_name == target or target in proc_name:
                    pid = p.info['pid']
                    try:
                        p.kill()  # SIGKILL - force kill ngay
                        killed.append(f"{p.info['name']} (PID: {pid})")
                    except:
                        pass
            except:
                pass
        
        # Bước 2: Backup với taskkill /F (force)
        try:
            # Kill theo image name
            subprocess.run(
                ['taskkill', '/F', '/IM', target_exe],
                capture_output=True,
                timeout=5
            )
            # Thử cả không có .exe
            subprocess.run(
                ['taskkill', '/F', '/IM', f'{target}*'],
                capture_output=True,
                timeout=5
            )
        except:
            pass
        
        # Bước 3: Verify đã kill hết chưa
        remaining = []
        for p in psutil.process_iter(['pid', 'name']):
            try:
                proc_name = (p.info['name'] or "").lower()
                if proc_name == target_exe or proc_name == target or target in proc_name:
                    remaining.append(f"{p.info['name']} (PID: {p.info['pid']})")
            except:
                pass
        
        if killed and not remaining:
            return {
                "success": True,
                "message": f"✅ Đã FORCE KILL thành công: {', '.join(killed)}",
                "killed_count": len(killed),
                "killed": killed
            }
        elif remaining:
            return {
                "success": False,
                "error": f"❌ Không thể kill (cần quyền Admin): {', '.join(remaining)}",
                "killed": killed if killed else []
            }
        else:
            return {
                "success": False,
                "error": f"Không tìm thấy app '{app_name}' đang chạy"
            }
            
    except Exception as e:
        return {"success": False, "error": str(e)}

async def find_process(name_pattern: str = "", show_all: bool = False) -> dict:
    """
    Tìm kiếm process theo tên hoặc hiển thị tất cả.
    
    Args:
        name_pattern: Tên process cần tìm (partial match, case insensitive). Để trống = tất cả
        show_all: True = hiển thị tất cả process (bỏ qua limit)
    
    Returns:
        dict: Danh sách processes tìm thấy
    """
    try:
        procs = []
        pattern_lower = name_pattern.lower() if name_pattern else ""
        
        for p in psutil.process_iter(['pid', 'name', 'cpu_percent', 'memory_percent']):
            try:
                info = p.info
                proc_name = (info['name'] or "").lower()
                
                # Filter theo pattern nếu có
                if pattern_lower and pattern_lower not in proc_name:
                    continue
                    
                procs.append({
                    "pid": info['pid'], 
                    "name": info['name'], 
                    "cpu": round(info['cpu_percent'] or 0, 2), 
                    "memory": round(info['memory_percent'] or 0, 2)
                })
            except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
                pass
        
        # Sort theo CPU usage nếu không có filter cụ thể
        if not pattern_lower:
            procs = sorted(procs, key=lambda x: x['cpu'], reverse=True)
            
        # Limit chỉ khi không show_all và không có pattern cụ thể
        if not show_all and not pattern_lower:
            procs = procs[:20]  # Top 20 thay vì 10
            
        # Tạo message tóm tắt
        if pattern_lower:
            found_count = len(procs)
            if found_count == 0:
                message = f"❌ Không tìm thấy process nào chứa '{name_pattern}'"
            elif found_count == 1:
                message = f"✅ Tìm thấy 1 process: {procs[0]['name']}"
            else:
                message = f"✅ Tìm thấy {found_count} processes chứa '{name_pattern}'"
        else:
            message = f"📋 Danh sách {len(procs)} processes (sorted by CPU usage)"
            
        return {
            "success": True, 
            "processes": procs, 
            "count": len(procs),
            "pattern": name_pattern,
            "message": message
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

async def create_file(path: str, content: str) -> dict:
    try:
        import os
        
        # Validate path - must be absolute on Windows (contains drive letter)
        if not os.path.isabs(path):
            return {"success": False, "error": f"Path must be absolute. Got: '{path}'. Example: 'C:/folder/file.txt'"}
        
        # Normalize path separators
        path = os.path.normpath(path)
        
        # Check if parent directory exists, create if needed
        parent_dir = os.path.dirname(path)
        if parent_dir and not os.path.exists(parent_dir):
            try:
                os.makedirs(parent_dir, exist_ok=True)
            except Exception as e:
                return {"success": False, "error": f"Cannot create directory '{parent_dir}': {str(e)}"}
        
        with open(path, 'w', encoding='utf-8') as f:
            f.write(content)
        return {"success": True, "path": path, "message": f"Đã tạo: {path}"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def read_file(path: str) -> dict:
    try:
        import os
        
        # Validate path - must be absolute on Windows (contains drive letter)
        if not os.path.isabs(path):
            return {"success": False, "error": f"Path must be absolute. Got: '{path}'. Example: 'C:/folder/file.txt'"}
        
        # Normalize path separators
        path = os.path.normpath(path)
        
        # Check if file exists
        if not os.path.exists(path):
            return {"success": False, "error": f"File not found: '{path}'"}
        
        if not os.path.isfile(path):
            return {"success": False, "error": f"Path is not a file: '{path}'"}
        
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        return {"success": True, "path": path, "content": content[:500], "size": len(content)}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def list_files(directory: str) -> dict:
    try:
        import os
        files = []
        for item in os.listdir(directory):
            p = os.path.join(directory, item)
            files.append({"name": item, "type": "dir" if os.path.isdir(p) else "file", "size": os.path.getsize(p) if os.path.isfile(p) else 0})
        return {"success": True, "directory": directory, "files": files, "count": len(files)}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_battery_status() -> dict:
    try:
        bat = psutil.sensors_battery()
        if bat is None:
            return {"success": False, "error": "Không thể lấy thông tin pin (có thể không có pin)"}
        return {
            "success": True,
            "percent": bat.percent,
            "plugged": bat.power_plugged,
            "time_left": str(bat.secsleft) if bat.secsleft != psutil.POWER_TIME_UNLIMITED else "Unlimited"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_disk_usage() -> dict:
    try:
        disks = []
        for part in psutil.disk_partitions():
            try:
                usage = psutil.disk_usage(part.mountpoint)
                disks.append({"device": part.device, "mountpoint": part.mountpoint, "fstype": part.fstype, "total_gb": round(usage.total / (1024**3), 2), "used_gb": round(usage.used / (1024**3), 2), "free_gb": round(usage.free / (1024**3), 2), "percent": usage.percent})
            except (PermissionError, OSError):
                # Bỏ qua các ổ đĩa không thể truy cập
                pass
        return {"success": True, "disks": disks, "count": len(disks)}
    except Exception as e:
        return {"success": False, "error": str(e)}

# ============================================================
# MUSIC LIBRARY TOOLS - VLC PLAYER
# ============================================================

MUSIC_LIBRARY = Path(__file__).parent / "music_library"
MUSIC_EXTENSIONS = {'.mp3', '.wav', '.flac', '.m4a', '.ogg', '.wma', '.aac'}

# YouTube Playlists Management
YOUTUBE_PLAYLISTS_FILE = Path(__file__).parent / "youtube_playlists.json"

def load_youtube_playlists() -> list:
    """Đọc danh sách playlist YouTube từ file JSON"""
    try:
        if YOUTUBE_PLAYLISTS_FILE.exists():
            with open(YOUTUBE_PLAYLISTS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        return []
    except Exception as e:
        print(f"❌ [Playlists] Error loading: {e}")
        return []

def save_youtube_playlists(playlists: list) -> bool:
    """Lưu danh sách playlist YouTube vào file JSON"""
    try:
        with open(YOUTUBE_PLAYLISTS_FILE, 'w', encoding='utf-8') as f:
            json.dump(playlists, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        print(f"❌ [Playlists] Error saving: {e}")
        return False

async def add_youtube_playlist(name: str, url: str) -> dict:
    """Thêm playlist YouTube mới"""
    try:
        playlists = load_youtube_playlists()
        
        # Kiểm tra trùng tên
        if any(p['name'].lower() == name.lower() for p in playlists):
            return {
                "success": False,
                "error": f"Playlist '{name}' đã tồn tại!"
            }
        
        # Thêm playlist mới
        new_playlist = {
            "name": name,
            "url": url,
            "created_at": datetime.now().isoformat()
        }
        playlists.append(new_playlist)
        
        if save_youtube_playlists(playlists):
            return {
                "success": True,
                "message": f"✅ Đã thêm playlist: {name}",
                "playlist": new_playlist
            }
        else:
            return {
                "success": False,
                "error": "Không thể lưu playlist"
            }
    except Exception as e:
        return {"success": False, "error": str(e)}

async def remove_youtube_playlist(name: str) -> dict:
    """Xóa playlist YouTube"""
    try:
        playlists = load_youtube_playlists()
        
        # Tìm và xóa playlist
        original_count = len(playlists)
        playlists = [p for p in playlists if p['name'].lower() != name.lower()]
        
        if len(playlists) == original_count:
            return {
                "success": False,
                "error": f"Không tìm thấy playlist: {name}"
            }
        
        if save_youtube_playlists(playlists):
            return {
                "success": True,
                "message": f"✅ Đã xóa playlist: {name}"
            }
        else:
            return {
                "success": False,
                "error": "Không thể lưu thay đổi"
            }
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_youtube_playlists() -> dict:
    """Lấy danh sách tất cả playlist YouTube"""
    try:
        playlists = load_youtube_playlists()
        return {
            "success": True,
            "playlists": playlists,
            "count": len(playlists)
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

async def open_youtube_playlist(playlist_name: str) -> dict:
    """Mở playlist YouTube đã lưu trong browser
    
    Args:
        playlist_name: Tên playlist đã đăng ký (có thể là tên đầy đủ hoặc từ khóa)
    
    Returns:
        dict với thông tin playlist đã mở
    """
    try:
        import webbrowser
        
        playlists = load_youtube_playlists()
        
        if not playlists:
            return {
                "success": False,
                "error": "Chưa có playlist nào. Hãy thêm playlist trên Web UI!"
            }
        
        # Tìm playlist (exact match hoặc partial match)
        playlist_name_lower = playlist_name.lower()
        matched_playlist = None
        
        # Tìm exact match trước
        for p in playlists:
            if p['name'].lower() == playlist_name_lower:
                matched_playlist = p
                break
        
        # Nếu không có exact match, tìm partial match
        if not matched_playlist:
            for p in playlists:
                if playlist_name_lower in p['name'].lower():
                    matched_playlist = p
                    break
        
        if not matched_playlist:
            # Hiển thị danh sách playlist có sẵn
            available = [p['name'] for p in playlists]
            return {
                "success": False,
                "error": f"Không tìm thấy playlist: '{playlist_name}'",
                "available_playlists": available,
                "hint": f"Có {len(available)} playlist: {', '.join(available)}"
            }
        
        # Mở playlist trong browser
        webbrowser.open(matched_playlist['url'])
        
        print(f"🎵 [YouTube Playlist] Đã mở: {matched_playlist['name']}")
        
        return {
            "success": True,
            "message": f"✅ Đã mở playlist: {matched_playlist['name']}",
            "playlist": matched_playlist,
            "url": matched_playlist['url']
        }
        
    except Exception as e:
        print(f"❌ [YouTube Playlist] Error: {e}")
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}

# VLC Player Manager (Singleton)
class VLCMusicPlayer:
    """
    VLC Music Player với hỗ trợ đầy đủ:
    - Play/Pause/Stop
    - Next/Previous track
    - Playlist management
    - Fuzzy song matching (tìm bài gần đúng)
    - Media keys support (VLC tự động hỗ trợ)
    """
    _instance = None
    _player = None
    _media_list = None
    _list_player = None
    _current_playlist = []
    _shuffle = False
    _repeat_mode = 0  # 0: off, 1: all, 2: one
    _song_cache = {}  # Cache danh sách bài hát
    
    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
        return cls._instance
    
    def __init__(self):
        if self._player is None:
            try:
                import vlc
                self._vlc = vlc
                # Tạo VLC instance với UI đầy đủ
                # Không dùng --no-xlib, --no-video, --no-audio-display
                # Thêm --video-on-top để cửa sổ luôn hiển thị
                self._instance_vlc = vlc.Instance()  # Empty options = full UI
                self._player = self._instance_vlc.media_player_new()
                self._media_list = self._instance_vlc.media_list_new()
                self._list_player = self._instance_vlc.media_list_player_new()
                self._list_player.set_media_player(self._player)
                print("✅ [VLC] VLC Music Player initialized (full UI + fuzzy matching)")
            except Exception as e:
                print(f"❌ [VLC] Failed to initialize: {e}")
                self._player = None
    
    def play_file(self, file_path: str):
        """Phát 1 file nhạc"""
        if not self._player:
            return False
        try:
            media = self._instance_vlc.media_new(file_path)
            self._player.set_media(media)
            self._player.play()
            return True
        except Exception as e:
            print(f"❌ [VLC] Play error: {e}")
            return False
    
    def play_playlist(self, file_paths: list):
        """Phát playlist với nhiều bài"""
        if not self._list_player:
            print("❌ [VLC] list_player chưa khởi tạo")
            return False
        try:
            print(f"🎵 [VLC DEBUG] play_playlist called with {len(file_paths)} files")
            for i, p in enumerate(file_paths[:3]):  # Log 3 file đầu
                print(f"   [{i+1}] {p}")
            
            # QUAN TRỌNG: STOP bài đang phát trước!
            self._list_player.stop()
            import time
            time.sleep(0.3)
            print("🛑 [VLC] Stopped current playback")
            
            # Clear playlist cũ và tạo mới
            self._media_list = self._instance_vlc.media_list_new()
            self._current_playlist = file_paths
            
            # Thêm tất cả bài vào playlist
            for path in file_paths:
                media = self._instance_vlc.media_new(path)
                self._media_list.add_media(media)
            
            print(f"🎵 [VLC DEBUG] Media list count: {self._media_list.count()}")
            
            # Set playlist mới
            self._list_player.set_media_list(self._media_list)
            
            # Set current index to 0 (first song)
            self._current_index = 0
            
            # QUAN TRỌNG: Gọi play() để phát bài đầu tiên
            self._list_player.play()
            print(f"🎵 [VLC DEBUG] list_player.play() called")
            
            # FIX DOUBLE-CLICK: Tăng thời gian chờ để VLC khởi tạo đầy đủ
            time.sleep(0.7)
            
            # Kiểm tra và đảm bảo đang phát với retry mechanism
            if self._player:
                state = self._player.get_state()
                is_playing = self._player.is_playing()
                current_vol = self._player.audio_get_volume()
                print(f"🎵 [VLC DEBUG] State: {state}, is_playing: {is_playing}, volume: {current_vol}")
                
                # FIX: Retry nếu chưa phát (quan trọng cho double-click)
                retry_count = 0
                max_retries = 3
                while not is_playing and retry_count < max_retries:
                    print(f"⚠️ [VLC DEBUG] Not playing, retry {retry_count+1}/{max_retries}...")
                    self._list_player.play()
                    time.sleep(0.4)
                    is_playing = self._player.is_playing()
                    retry_count += 1
                
                # Đảm bảo volume đủ nghe
                if current_vol < 50:
                    self._player.audio_set_volume(80)
                    print(f"🔊 [VLC] Volume was {current_vol}, set to 80")
            
            print(f"▶️ [VLC] Playing playlist with {len(file_paths)} songs")
            return True
        except Exception as e:
            print(f"❌ [VLC] Playlist error: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def pause(self):
        """Tạm dừng"""
        if self._player:
            self._player.pause()
            return True
        return False
    
    def resume(self):
        """Tiếp tục phát - Đảm bảo đang play"""
        if self._list_player:
            # Nếu đang paused, gọi play để tiếp tục
            if not self.is_playing():
                self._list_player.play()
            return True
        elif self._player:
            if not self.is_playing():
                self._player.play()
            return True
        return False
    
    def stop(self):
        """Dừng phát hoàn toàn và reset trạng thái"""
        try:
            import time
            
            # Stop cả list_player và player
            if self._list_player:
                self._list_player.stop()
                time.sleep(0.1)
            
            if self._player:
                self._player.stop()
                time.sleep(0.1)
            
            # Verify đã dừng thực sự
            stopped = False
            for _ in range(3):  # Retry 3 lần
                if not self.is_playing():
                    stopped = True
                    break
                time.sleep(0.1)
                if self._player:
                    self._player.stop()
            
            if stopped:
                print("✅ [VLC] Stopped successfully")
            else:
                print("⚠️ [VLC] Stop command sent but player may still be active")
            
            return True
        except Exception as e:
            print(f"❌ [VLC] Stop error: {e}")
            return False
    
    def next_track(self):
        """Bài tiếp theo - Tự động phát luôn với retry logic!"""
        if self._list_player and self._current_playlist:
            current_idx = getattr(self, '_current_index', 0)
            last_idx = len(self._current_playlist) - 1
            
            # Stop hiện tại để tránh conflict
            self._list_player.stop()
            
            if current_idx >= last_idx:
                # Đã ở bài cuối, quay lại bài đầu
                self._current_index = 0
                print(f"🔄 [VLC] Next: Wrap to first track (index 0)")
            else:
                # Còn bài tiếp, chuyển bình thường
                self._current_index = current_idx + 1
                print(f"⏭️ [VLC] Next: Now at index {self._current_index}")
            
            # Play bài mới bằng index
            self._list_player.play_item_at_index(self._current_index)
            
            import time
            time.sleep(0.4)
            
            # Retry nếu chưa phát (tối đa 2 lần)
            retry_count = 0
            while not self.is_playing() and retry_count < 2:
                print(f"⚠️ [VLC] Not playing yet, retry {retry_count + 1}/2...")
                self._list_player.play()
                time.sleep(0.3)
                retry_count += 1
            
            # Verify
            if self.is_playing():
                print(f"✅ [VLC] Next track playing successfully")
                return True
            else:
                print(f"❌ [VLC] Failed to play next track after retries")
                return False
        return False
    
    def previous_track(self):
        """Bài trước - Tự động phát luôn với retry logic!"""
        if self._list_player and self._current_playlist:
            # Kiểm tra nếu đang ở bài đầu tiên
            current_idx = getattr(self, '_current_index', 0)
            
            # Stop hiện tại để tránh conflict
            self._list_player.stop()
            
            if current_idx <= 0:
                # Đã ở bài đầu, quay lại bài cuối cùng của playlist
                last_idx = len(self._current_playlist) - 1
                self._current_index = last_idx
                print(f"🔄 [VLC] Previous: Wrap to last track (index {last_idx})")
            else:
                # Còn bài trước, chuyển bình thường
                self._current_index = current_idx - 1
                print(f"⏮️ [VLC] Previous: Now at index {self._current_index}")
            
            # Play bài mới bằng index
            self._list_player.play_item_at_index(self._current_index)
            
            import time
            time.sleep(0.4)
            
            # Retry nếu chưa phát (tối đa 2 lần)
            retry_count = 0
            while not self.is_playing() and retry_count < 2:
                print(f"⚠️ [VLC] Not playing yet, retry {retry_count + 1}/2...")
                self._list_player.play()
                time.sleep(0.3)
                retry_count += 1
            
            # Verify
            if self.is_playing():
                print(f"✅ [VLC] Previous track playing successfully")
                return True
            else:
                print(f"❌ [VLC] Failed to play previous track after retries")
                return False
        return False
    
    def is_playing(self):
        """Kiểm tra đang phát không"""
        if self._player:
            return self._player.is_playing()
        return False
    
    def get_state(self):
        """Lấy trạng thái player"""
        if not self._player:
            return "not_initialized"
        
        state = self._player.get_state()
        state_map = {
            0: "idle",
            1: "opening",
            2: "buffering", 
            3: "playing",
            4: "paused",
            5: "stopped",
            6: "ended",
            7: "error"
        }
        return state_map.get(state, "unknown")
    
    def get_position(self):
        """Lấy vị trí hiện tại (0.0 - 1.0)"""
        if self._player:
            return self._player.get_position() or 0.0
        return 0.0
    
    def get_time(self):
        """Lấy thời gian hiện tại (milliseconds)"""
        if self._player:
            return self._player.get_time() or 0
        return 0
    
    def get_length(self):
        """Lấy độ dài bài hát (milliseconds)"""
        if self._player:
            return self._player.get_length() or 0
        return 0
    
    def get_volume(self):
        """Lấy âm lượng hiện tại (0-100)"""
        if self._player:
            return self._player.audio_get_volume() or 0
        return 0
    
    def set_volume(self, level: int):
        """Đặt âm lượng (0-100)"""
        if self._player:
            level = max(0, min(100, level))
            self._player.audio_set_volume(level)
            return True
        return False
    
    def set_position(self, position: float):
        """Đặt vị trí (0.0 - 1.0)"""
        if self._player:
            position = max(0.0, min(1.0, position))
            self._player.set_position(position)
            return True
        return False
    
    def get_current_media_title(self):
        """Lấy tiêu đề media đang phát - TỐI ƯU với cache"""
        try:
            if self._player:
                media = self._player.get_media()
                if media:
                    # Cache để tránh query lại liên tục
                    title = media.get_meta(self._vlc.Meta.Title)
                    if title:
                        self._cached_title = title
                        return title
                    # Fallback: filename
                    mrl = media.get_mrl()
                    if mrl:
                        from urllib.parse import unquote
                        path = unquote(mrl.replace('file:///', '').replace('file://', ''))
                        fname = Path(path).name
                        self._cached_title = fname
                        return fname
            # Return cached nếu có
            return getattr(self, '_cached_title', None)
        except:
            return getattr(self, '_cached_title', None)
    
    def get_playlist_index(self):
        """Lấy index bài hiện tại trong playlist"""
        # VLC không có API trực tiếp, phải track riêng
        return getattr(self, '_current_index', 0)
    
    def get_playlist_count(self):
        """Lấy số bài trong playlist"""
        return len(self._current_playlist) if self._current_playlist else 0
    
    def get_full_status(self):
        """Lấy trạng thái đầy đủ cho Web UI - TỐI ƯU"""
        state = self.get_state()
        current_time_ms = self.get_time()
        duration_ms = self.get_length()
        
        return {
            "state": state,
            "is_playing": self.is_playing(),
            "position": self.get_position(),
            "current_time_ms": current_time_ms,
            "current_time_formatted": self._format_time(current_time_ms),
            "duration_ms": duration_ms,
            "duration_formatted": self._format_time(duration_ms),
            "volume": self.get_volume(),
            "current_track": self.get_current_media_title(),
            "playlist_index": self.get_playlist_index(),
            "playlist_count": self.get_playlist_count(),
            "playlist": [Path(p).name for p in self._current_playlist[:5]] if self._current_playlist else [],  # CHỈ 5 bài (giảm data)
            "shuffle": self._shuffle,
            "repeat_mode": self._repeat_mode
        }
    
    def set_shuffle(self, enabled: bool):
        """Bật/tắt chế độ phát ngẫu nhiên"""
        self._shuffle = enabled
        if self._list_player:
            # VLC MediaListPlayer không có native shuffle, ta xử lý thủ công khi next/previous
            pass
        return self._shuffle
    
    def set_repeat_mode(self, mode: int):
        """Đặt chế độ lặp lại: 0=off, 1=all, 2=one"""
        self._repeat_mode = mode
        if self._list_player:
            if mode == 0:
                self._list_player.set_playback_mode(self._vlc.PlaybackMode.default)
            elif mode == 1:
                self._list_player.set_playback_mode(self._vlc.PlaybackMode.loop)
            elif mode == 2:
                self._list_player.set_playback_mode(self._vlc.PlaybackMode.repeat)
        return self._repeat_mode
    
    def get_shuffle(self):
        """Lấy trạng thái shuffle"""
        return getattr(self, '_shuffle', False)
    
    def get_repeat_mode(self):
        """Lấy chế độ repeat: 0=off, 1=all, 2=one"""
        return getattr(self, '_repeat_mode', 0)
    
    def _format_time(self, ms):
        """Format milliseconds thành MM:SS"""
        if not ms or ms < 0:
            return "0:00"
        seconds = int(ms / 1000)
        minutes = seconds // 60
        seconds = seconds % 60
        return f"{minutes}:{seconds:02d}"
    
    def refresh_song_cache(self, music_folder: Path):
        """Refresh cache danh sách bài hát từ music_library"""
        try:
            print(f"🔄 [VLC] Refreshing song cache from {music_folder}...")
            self._song_cache = {}
            
            if not music_folder.exists():
                print(f"⚠️ [VLC] Music folder not found: {music_folder}")
                return
            
            extensions = ['.mp3', '.flac', '.wav', '.m4a', '.ogg', '.wma']
            for file_path in music_folder.rglob("*"):
                if file_path.is_file() and file_path.suffix.lower() in extensions:
                    # Lưu: tên file (lowercase) -> đường dẫn đầy đủ
                    song_name = file_path.stem.lower()  # Tên file không có extension
                    self._song_cache[song_name] = str(file_path)
            
            print(f"✅ [VLC] Song cache refreshed: {len(self._song_cache)} songs")
        except Exception as e:
            print(f"❌ [VLC] Error refreshing song cache: {e}")
    
    def fuzzy_match_song(self, query: str, threshold: float = 0.3):
        """
        Tìm bài hát gần đúng bằng fuzzy matching với Unicode normalization
        
        Args:
            query: Tên bài hát người dùng nói (e.g., "phát bài yêu em", "Đa Nghi")
            threshold: Ngưỡng tương đồng (0.0-1.0), mặc định 0.3 (GIẢM để dễ match hơn)
            
        Returns:
            tuple: (best_match_path, similarity_score) hoặc (None, 0.0)
        """
        if not self._song_cache:
            print("⚠️ [VLC] Song cache empty, call refresh_song_cache() first")
            return None, 0.0
        
        import unicodedata
        
        # Normalize Unicode (NFD = decompose dấu) để so sánh tốt hơn
        def normalize_text(text):
            # NFD: tách dấu khỏi ký tự (e.g., "á" -> "a" + dấu)
            text = unicodedata.normalize('NFD', text)
            # Loại bỏ dấu thanh (chỉ giữ chữ cái cơ bản)
            text = ''.join(c for c in text if not unicodedata.combining(c))
            # Lowercase và loại bỏ ký tự đặc biệt
            text = re.sub(r'[^\w\s]', '', text.lower()).strip()
            text = re.sub(r'\s+', ' ', text)  # Collapse spaces
            return text
        
        query_normalized = normalize_text(query)
        
        # Loại bỏ các từ điều khiển thường gặp
        stop_words = ['phat', 'bai', 'mo', 'chay', 'play', 'song', 'nhac', 'hat']
        query_words = [w for w in query_normalized.split() if w not in stop_words]
        query_processed = ' '.join(query_words) if query_words else query_normalized
        
        print(f"🔍 [VLC Fuzzy] Query: '{query}' -> Normalized: '{query_processed}'")
        
        best_match = None
        best_score = 0.0
        
        for song_name_original, song_path in self._song_cache.items():
            # Normalize song name để so sánh
            song_name_normalized = normalize_text(song_name_original)
            
            # Tính similarity với difflib
            similarity = difflib.SequenceMatcher(None, query_processed, song_name_normalized).ratio()
            
            # Thưởng điểm nếu query có trong tên bài (substring match)
            if query_processed in song_name_normalized:
                similarity += 0.25
            
            # Thưởng điểm nếu từng từ đều có trong tên bài
            if query_words:
                words_match = all(word in song_name_normalized for word in query_words)
                if words_match:
                    similarity += 0.20
            
            # Thưởng điểm nếu bắt đầu giống nhau (prefix match)
            if song_name_normalized.startswith(query_processed[:4]):  # 4 ký tự đầu
                similarity += 0.10
            
            if similarity > best_score:
                best_score = similarity
                best_match = song_path
        
        if best_score >= threshold:
            print(f"✅ [VLC Fuzzy] Found match: {Path(best_match).name} (score: {best_score:.2f})")
            return best_match, best_score
        else:
            print(f"❌ [VLC Fuzzy] No match found above threshold {threshold} (best: {best_score:.2f})")
            return None, 0.0
    
    def play_by_fuzzy_match(self, query: str, threshold: float = 0.4):
        """
        Phát bài hát bằng fuzzy matching
        
        Args:
            query: Tên bài hát người dùng nói
            threshold: Ngưỡng tương đồng
            
        Returns:
            dict with success, matched_song, score, message
        """
        matched_path, score = self.fuzzy_match_song(query, threshold)
        
        if not matched_path:
            return {
                "success": False,
                "error": f"Không tìm thấy bài '{query}' (threshold={threshold})",
                "query": query,
                "score": score
            }
        
        # Phát bài tìm được
        success = self.play_file(matched_path)
        
        if success:
            song_name = Path(matched_path).name
            return {
                "success": True,
                "matched_song": song_name,
                "score": score,
                "path": matched_path,
                "message": f"🎵 Đang phát: {song_name} (tìm được với độ chính xác {score*100:.0f}%)"
            }
        else:
            return {
                "success": False,
                "error": "VLC không thể phát file",
                "matched_song": Path(matched_path).name,
                "score": score
            }
    
    async def play_file_async(self, file_path: str):
        """Async wrapper cho play_file để không blocking"""
        return await asyncio.to_thread(self.play_file, file_path)
    
    async def play_playlist_async(self, file_paths: list):
        """Async wrapper cho play_playlist để không blocking"""
        return await asyncio.to_thread(self.play_playlist, file_paths)

# Global VLC player instance - với error handling
try:
    vlc_player = VLCMusicPlayer()
    VLC_AVAILABLE = vlc_player._player is not None
except Exception as e:
    print(f"⚠️ [VLC] VLC không khả dụng: {e}")
    vlc_player = None
    VLC_AVAILABLE = False

if not VLC_AVAILABLE:
    print("⚠️ [VLC] Music player disabled. Cài VLC: https://www.videolan.org/vlc/")

# ============================================================
# 🎯 VLC MCP SERVER - Hybrid System (REST + MCP)
# ============================================================
try:
    from vlc_mcp_server import VLCMCPServer
    
    # Initialize MCP server with VLC player instance
    if VLC_AVAILABLE and vlc_player:
        vlc_mcp_server = VLCMCPServer(vlc_player)
        print(f"✅ [VLC MCP] Hybrid System initialized - {len(vlc_mcp_server.tools)} tools available")
        VLC_MCP_AVAILABLE = True
    else:
        vlc_mcp_server = None
        VLC_MCP_AVAILABLE = False
        print("⚠️ [VLC MCP] MCP server disabled - VLC not available")
except Exception as e:
    print(f"⚠️ [VLC MCP] Failed to initialize MCP server: {e}")
    vlc_mcp_server = None
    VLC_MCP_AVAILABLE = False

# ============================================================
# BROWSER CONTROLLER - Selenium Automation
# ============================================================

class BrowserController:
    """Singleton class để điều khiển trình duyệt Chrome bằng Selenium"""
    
    _instance = None
    _driver = None
    
    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
        return cls._instance
    
    def _ensure_driver(self):
        """Khởi tạo Chrome driver nếu chưa có"""
        if self._driver is None:
            if not SELENIUM_AVAILABLE:
                raise Exception("Selenium chưa được cài đặt. Chạy: pip install selenium webdriver-manager")
            
            try:
                chrome_options = Options()
                chrome_options.add_argument('--disable-gpu')
                chrome_options.add_argument('--no-sandbox')
                chrome_options.add_argument('--disable-dev-shm-usage')
                chrome_options.add_argument('--start-maximized')
                
                service = Service(ChromeDriverManager().install())
                self._driver = webdriver.Chrome(service=service, options=chrome_options)
                print("✅ [Browser] Chrome driver initialized")
            except Exception as e:
                print(f"❌ [Browser] Failed to initialize: {e}")
                raise
        return self._driver
    
    def open_url(self, url: str) -> dict:
        """Mở URL trong browser"""
        try:
            driver = self._ensure_driver()
            driver.get(url)
            return {
                "success": True,
                "url": driver.current_url,
                "title": driver.title,
                "message": f"Đã mở: {driver.title}"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def get_current_info(self) -> dict:
        """Lấy thông tin trang hiện tại"""
        try:
            if self._driver is None:
                return {"success": False, "error": "Browser chưa được khởi động"}
            
            return {
                "success": True,
                "url": self._driver.current_url,
                "title": self._driver.title,
                "window_handles": len(self._driver.window_handles)
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def click_element(self, selector: str, by: str = "css") -> dict:
        """Click vào element"""
        try:
            driver = self._ensure_driver()
            
            by_map = {
                "css": By.CSS_SELECTOR,
                "xpath": By.XPATH,
                "id": By.ID,
                "name": By.NAME,
                "class": By.CLASS_NAME,
                "tag": By.TAG_NAME
            }
            
            by_type = by_map.get(by.lower(), By.CSS_SELECTOR)
            element = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((by_type, selector))
            )
            element.click()
            
            return {
                "success": True,
                "message": f"Đã click vào element: {selector}"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def fill_input(self, selector: str, text: str, by: str = "css") -> dict:
        """Điền text vào input field"""
        try:
            driver = self._ensure_driver()
            
            by_map = {
                "css": By.CSS_SELECTOR,
                "xpath": By.XPATH,
                "id": By.ID,
                "name": By.NAME,
                "class": By.CLASS_NAME
            }
            
            by_type = by_map.get(by.lower(), By.CSS_SELECTOR)
            element = WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((by_type, selector))
            )
            element.clear()
            element.send_keys(text)
            
            return {
                "success": True,
                "message": f"Đã điền text vào: {selector}"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def scroll(self, direction: str = "down", amount: int = 500) -> dict:
        """Cuộn trang"""
        try:
            driver = self._ensure_driver()
            
            if direction == "top":
                driver.execute_script("window.scrollTo(0, 0);")
            elif direction == "bottom":
                driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            elif direction == "down":
                driver.execute_script(f"window.scrollBy(0, {amount});")
            elif direction == "up":
                driver.execute_script(f"window.scrollBy(0, -{amount});")
            else:
                return {"success": False, "error": f"Invalid direction: {direction}"}
            
            return {
                "success": True,
                "message": f"Đã cuộn {direction}"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def go_back(self) -> dict:
        """Quay lại trang trước"""
        try:
            if self._driver is None:
                return {"success": False, "error": "Browser chưa được khởi động"}
            self._driver.back()
            return {"success": True, "message": "Đã quay lại trang trước"}
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def go_forward(self) -> dict:
        """Tiến tới trang sau"""
        try:
            if self._driver is None:
                return {"success": False, "error": "Browser chưa được khởi động"}
            self._driver.forward()
            return {"success": True, "message": "Đã tiến tới trang sau"}
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def refresh(self) -> dict:
        """Làm mới trang"""
        try:
            if self._driver is None:
                return {"success": False, "error": "Browser chưa được khởi động"}
            self._driver.refresh()
            return {"success": True, "message": "Đã làm mới trang"}
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def screenshot(self, filepath: str = None) -> dict:
        """Chụp screenshot trang hiện tại"""
        try:
            if self._driver is None:
                return {"success": False, "error": "Browser chưa được khởi động"}
            
            if filepath is None:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filepath = f"screenshot_{timestamp}.png"
            
            self._driver.save_screenshot(filepath)
            return {
                "success": True,
                "filepath": filepath,
                "message": f"Đã lưu screenshot: {filepath}"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def new_tab(self, url: str = None) -> dict:
        """Mở tab mới"""
        try:
            driver = self._ensure_driver()
            driver.execute_script("window.open('');")
            driver.switch_to.window(driver.window_handles[-1])
            
            if url:
                driver.get(url)
            
            return {
                "success": True,
                "message": f"Đã mở tab mới{' và truy cập ' + url if url else ''}",
                "total_tabs": len(driver.window_handles)
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def close_tab(self) -> dict:
        """Đóng tab hiện tại"""
        try:
            if self._driver is None:
                return {"success": False, "error": "Browser chưa được khởi động"}
            
            self._driver.close()
            if len(self._driver.window_handles) > 0:
                self._driver.switch_to.window(self._driver.window_handles[-1])
            
            return {
                "success": True,
                "message": "Đã đóng tab",
                "remaining_tabs": len(self._driver.window_handles)
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def execute_script(self, script: str) -> dict:
        """Thực thi JavaScript code"""
        try:
            driver = self._ensure_driver()
            result = driver.execute_script(script)
            return {
                "success": True,
                "result": result,
                "message": "Đã thực thi JavaScript"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def close_browser(self) -> dict:
        """Đóng browser hoàn toàn"""
        try:
            if self._driver:
                self._driver.quit()
                self._driver = None
                return {"success": True, "message": "Đã đóng browser"}
            return {"success": False, "error": "Browser chưa được khởi động"}
        except Exception as e:
            return {"success": False, "error": str(e)}

# Global browser controller instance
browser_controller = BrowserController()

async def list_music(subfolder: str = "", auto_play: bool = True, folder: str = "") -> dict:
    """
    Liệt kê file nhạc trong music_library hoặc thư mục tùy chỉnh.
    Theo mặc định TỰ ĐỘNG PHÁT bài đầu tiên (giống xinnan-tech/xiaozhi-esp32-server).
    Set auto_play=False để chỉ liệt kê không phát.
    
    Args:
        subfolder: Subfolder trong music_library
        auto_play: Tự động phát bài đầu tiên (default True)
        folder: Thư mục tùy chỉnh (nếu có, sẽ override music_library)
    """
    try:
        # Xác định thư mục gốc
        if folder and folder.strip():
            base_path = Path(folder.strip())
            if not base_path.exists():
                return {"success": False, "error": f"Thư mục '{folder}' không tồn tại"}
            search_path = base_path
            is_user_folder = True
        else:
            if not MUSIC_LIBRARY.exists():
                MUSIC_LIBRARY.mkdir(exist_ok=True)
                return {"success": True, "files": [], "count": 0, "message": "Thư mục music_library đã được tạo. Hãy thêm nhạc vào!"}
            
            base_path = MUSIC_LIBRARY
            search_path = MUSIC_LIBRARY / subfolder if subfolder else MUSIC_LIBRARY
            is_user_folder = False
        
        if not search_path.exists():
            return {"success": False, "error": f"Thư mục '{subfolder or folder}' không tồn tại"}
        
        music_files = []
        for file_path in search_path.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in MUSIC_EXTENSIONS:
                try:
                    relative_path = file_path.relative_to(base_path)
                except ValueError:
                    relative_path = file_path.name
                    
                music_files.append({
                    "filename": file_path.name,
                    "path": str(relative_path).replace('\\', '/'),
                    "full_path": str(file_path),
                    "size_mb": round(file_path.stat().st_size / (1024**2), 2),
                    "extension": file_path.suffix.lower()
                })
        
        music_files.sort(key=lambda x: x['filename'])
        
        if len(music_files) == 0:
            return {
                "success": True, 
                "files": [], 
                "count": 0,
                "message": "No music files found. Please add music files to the folder.",
                "is_user_folder": is_user_folder,
                "source_path": str(base_path)
            }
        
        # 🎵 AUTO-PLAY: Tự động phát bài đầu tiên (như code reference)
        first_file = music_files[0]['filename'] if not is_user_folder else music_files[0]['full_path']
        play_result = None
        
        if auto_play:
            print(f"🎵 [Auto-Play] list_music tự động phát: {first_file}")
            if is_user_folder:
                # Phát từ user folder bằng default player
                play_result = await play_music_from_path(music_files[0]['full_path'])
            else:
                play_result = await play_music(first_file)
            
            if play_result.get("success"):
                message = f"✅ Auto-played: {music_files[0]['filename']}\nTotal {len(music_files)} song(s)"
            else:
                message = f"❌ Found {len(music_files)} songs but failed to play: {play_result.get('error', 'Unknown error')}"
        else:
            filenames_list = [f['filename'] for f in music_files]
            message = f"Found {len(music_files)} song(s):\n" + "\n".join([f"  - {fname}" for fname in filenames_list[:10]])
            if len(music_files) > 10:
                message += f"\n  ... and {len(music_files) - 10} more"
        
        return {
            "success": True,
            "files": music_files,
            "count": len(music_files),
            "library_path": str(base_path),
            "is_user_folder": is_user_folder,
            "message": message,
            "auto_played": auto_play,
            "play_result": play_result if auto_play else None
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

async def play_music_from_path(file_path: str) -> dict:
    """
    Phát nhạc từ đường dẫn đầy đủ bằng Python-VLC (KHÔNG dùng trình phát mặc định).
    ⭐ NHANH & TIỆN - Dùng VLC nội bộ!
    """
    try:
        path = Path(file_path)
        if not path.exists():
            return {"success": False, "error": f"File không tồn tại: {file_path}"}
        
        # 🎵 SỬ DỤNG VLC thay vì os.startfile - NHANH!
        success = vlc_player.play_playlist([str(path)])
        
        if success:
            print(f"🎵 [VLC] Đang phát từ path: {path.name}")
            return {
                "success": True,
                "message": f"🎵 Đang phát: {path.name} (Python-VLC)",
                "file": path.name,
                "path": str(path),
                "player": "Python-VLC",
                "llm_note": "🎵 ĐANG DÙNG PYTHON-VLC. Điều khiển: pause_music(), resume_music(), stop_music(), music_next(), music_previous(). NHANH & TIỆN!"
            }
        else:
            return {"success": False, "error": "VLC Player không thể phát file"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def play_music(filename: str, create_playlist: bool = True, use_fuzzy: bool = True) -> dict:
    """
    Phát nhạc từ music_library bằng VLC player với fuzzy matching.
    
    Args:
        filename: Tên file (e.g., 'song.mp3' or 'Pop/song.mp3') hoặc tên gần đúng (e.g., 'yêu em')
        create_playlist: Tạo playlist với tất cả bài (default True) để hỗ trợ Next/Previous
        use_fuzzy: Dùng fuzzy matching nếu không tìm thấy chính xác (default True)
        
    Returns:
        dict with 'success', 'filename', 'path', 'message'
    """
    try:
        if not MUSIC_LIBRARY.exists():
            return {"success": False, "error": "Thư mục music_library không tồn tại"}
        
        print(f"🎵 [VLC Play] Tìm file: '{filename}'")
        
        # TỐI ƯU: Chỉ refresh cache nếu chưa có (lazy loading)
        if not hasattr(vlc_player, '_song_cache') or not vlc_player._song_cache:
            vlc_player.refresh_song_cache(MUSIC_LIBRARY)
        
        # Step 2: Tìm file chính xác trước
        music_path = None
        filename_lower = filename.lower()
        
        for file_path in MUSIC_LIBRARY.rglob("*"):
            if file_path.is_file():
                if (file_path.name == filename or 
                    file_path.name.lower() == filename_lower or
                    str(file_path.relative_to(MUSIC_LIBRARY)).replace('\\', '/') == filename or
                    filename_lower in file_path.name.lower()):
                    if file_path.suffix.lower() in MUSIC_EXTENSIONS:
                        music_path = file_path
                        print(f"✅ [VLC Play] Found exact match: {music_path}")
                        break
        
        # Step 3: Nếu không tìm thấy chính xác, dùng fuzzy matching
        if not music_path and use_fuzzy:
            print(f"🔍 [VLC Play] Exact match not found, trying fuzzy matching...")
            matched_path, score = vlc_player.fuzzy_match_song(filename, threshold=0.4)
            
            if matched_path:
                music_path = Path(matched_path)
                print(f"✅ [VLC Play] Fuzzy match found: {music_path.name} (score: {score:.2f})")
        
        if not music_path:
            available = [f.name for f in MUSIC_LIBRARY.rglob("*") if f.is_file() and f.suffix.lower() in MUSIC_EXTENSIONS]
            return {
                "success": False, 
                "error": f"Không tìm thấy '{filename}' (đã thử fuzzy matching)",
                "available_files": available[:5],
                "hint": "Thử tìm bằng từ khóa trong tên bài hoặc dùng list_music() để xem danh sách"
            }
        
        print(f"🎵 [VLC Play] Selected: {music_path}")
        
        if create_playlist:
            # Tạo playlist với tất cả bài trong thư mục
            all_songs = sorted([
                str(f) for f in MUSIC_LIBRARY.rglob("*") 
                if f.is_file() and f.suffix.lower() in MUSIC_EXTENSIONS
            ])
            
            # Đảm bảo bài hiện tại ở đầu playlist
            if str(music_path) in all_songs:
                all_songs.remove(str(music_path))
            all_songs.insert(0, str(music_path))
            
            success = await vlc_player.play_playlist_async(all_songs)
            print(f"🎵 [VLC] Created playlist with {len(all_songs)} songs")
        else:
            success = await vlc_player.play_file_async(str(music_path))
        
        if success:
            return {
                "success": True,
                "filename": music_path.name,
                "path": str(music_path.relative_to(MUSIC_LIBRARY)),
                "full_path": str(music_path),
                "size_mb": round(music_path.stat().st_size / (1024**2), 2),
                "message": f"🎵 Đang phát: {music_path.name} (Python-VLC + Fuzzy Matching)",
                "player": "Python-VLC Enhanced",
                "playlist_mode": create_playlist,
                "fuzzy_used": not (filename.lower() in music_path.name.lower()),
                "llm_note": "🎵 PYTHON-VLC PLAYER với FUZZY MATCHING! Có thể tìm bài gần đúng. Điều khiển: pause_music(), resume_music(), stop_music(), music_next(), music_previous()."
            }
        else:
            return {"success": False, "error": "VLC player không thể phát. Kiểm tra VLC đã cài đặt chưa!"}
    except Exception as e:
        print(f"❌ [VLC Play] Error: {e}")
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}

async def pause_music() -> dict:
    """
    ⏸️ TẠM DỪNG nhạc VLC Player.
    
    🎯 KHI NÀO GỌI: User nói "tạm dừng", "pause", "dừng lại", "đừng phát"
    
    ⚡ BẮT BUỘC GỌI TOOL NÀY! Không được tự trả lời "đã tạm dừng"!
    
    Returns:
        dict: {"success": bool, "message": str, "current_song": str}
    """
    try:
        if vlc_player and vlc_player._player:
            vlc_player.pause()
            status = vlc_player.get_full_status()
            current_song = status.get('current_song', 'Unknown')
            return {
                "success": True, 
                "message": f"⏸️ Đã tạm dừng: {current_song} (Python-VLC)",
                "player": "Python-VLC",
                "current_song": current_song,
                "llm_note": "⚡ GỌI TOOL ĐÃ THÀNH CÔNG! Đang dùng Python-VLC. LUÔN GỌI: resume_music() để tiếp tục, music_next()/music_previous() để chuyển bài. KHÔNG BAO GIỜ TỰ TRẢ LỜI mà không gọi tool!"
            }
        else:
            return {"success": False, "error": "VLC Player chưa khởi tạo hoặc chưa phát nhạc. Dùng play_music() để phát nhạc trước!"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def resume_music() -> dict:
    """
    ▶️ TIẾP TỤC phát nhạc VLC Player sau khi pause.
    
    🎯 KHI NÀO GỌI: User nói "tiếp tục", "resume", "phát tiếp", "play lại"
    
    ⚡ BẮT BUỘC GỌI TOOL NÀY! Không được tự trả lời "đã phát tiếp"!
    
    Returns:
        dict: {"success": bool, "message": str, "is_playing": bool}
    """
    try:
        if vlc_player and vlc_player._player:
            vlc_player.resume()  # Dùng method resume() mới - đảm bảo play
            import time
            time.sleep(0.2)
            status = vlc_player.get_full_status()
            current_song = status.get('current_song', 'Unknown')
            return {
                "success": True, 
                "message": f"▶️ Đang phát: {current_song} (Python-VLC)",
                "player": "Python-VLC",
                "current_song": current_song,
                "is_playing": True,
                "llm_note": "⚡ GỌI TOOL ĐÃ THÀNH CÔNG! Đang phát. LUÔN GỌI: pause_music() để dừng, music_next()/music_previous() để chuyển. KHÔNG TỰ TRẢ LỜI!"
            }
        else:
            return {"success": False, "error": "VLC Player chưa khởi tạo hoặc chưa phát nhạc. Dùng play_music() để phát nhạc trước!"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def stop_music() -> dict:
    """
    ⏹️ DỪNG HOÀN TOÀN nhạc VLC Player.
    
    🎯 KHI NÀO GỌI: User nói "dừng", "stop", "tắt nhạc", "ngừng phát"
    
    ⚡ BẮT BUỘC GỌI TOOL NÀY! Không được tự trả lời "đã dừng"!
    
    Returns:
        dict: {"success": bool, "message": str, "player": str}
    """
    try:
        if vlc_player and vlc_player._player:
            vlc_player.stop()
            return {
                "success": True, 
                "message": "⏹️ Đã dừng nhạc hoàn toàn (Python-VLC)",
                "player": "Python-VLC",
                "llm_note": "⚡ GỌI TOOL ĐÃ THÀNH CÔNG! Đã dừng hoàn toàn. Muốn phát lại → GỌI play_music(). KHÔNG TỰ TRẢ LỜI!"
            }
        else:
            return {"success": False, "error": "VLC Player chưa khởi tạo hoặc chưa phát nhạc."}
    except Exception as e:
        return {"success": False, "error": str(e)}

# ============================================================
# SMART MUSIC CONTROL - Điều khiển nhạc thông minh bằng ngôn ngữ tự nhiên
# Focus vào Python-VLC Player cho tất cả lệnh nhạc LOCAL
# ============================================================

# ============================================================
# FUZZY MATCHING - Xử lý nhận dạng giọng nói không chính xác từ ESP32
# ============================================================

# Các biến thể phát âm sai thường gặp (từ ESP32 voice recognition)
VOICE_CORRECTIONS = {
    # Bài tiếp/next variations
    'bài tiếp': ['bài tiếp', 'bai tiep', 'bài diệp', 'bài thiếp', 'bài típ', 'bay tiep', 'bai tip', 'bai diep'],
    'tiếp theo': ['tiếp theo', 'tiep theo', 'thiếp theo', 'típ theo', 'tiếp thêu', 'diệp theo'],
    'next': ['next', 'nex', 'nếch', 'nếc', 'nếx', 'net', 'nec'],
    'skip': ['skip', 'skíp', 'xkip', 'xíp', 'ship'],
    
    # Bài trước/previous variations  
    'bài trước': ['bài trước', 'bai truoc', 'bài chước', 'bài trước', 'bay truoc', 'bai chuoc', 'bài trướt'],
    'quay lại': ['quay lại', 'quay lai', 'quay lại bài', 'quai lai', 'quai lại', 'quáy lại'],
    'previous': ['previous', 'pre', 'prê', 'pri vi ớt', 'pri', 'prê vi ớt'],
    
    # Dừng/stop variations
    'dừng nhạc': ['dừng nhạc', 'dung nhac', 'dừng nhạc', 'dừng lại', 'dzừng nhạc'],
    'tắt nhạc': ['tắt nhạc', 'tat nhac', 'tắc nhạc', 'tác nhạc', 'tad nhac'],
    'pause': ['pause', 'pao', 'pốt', 'pót', 'pao xờ', 'pa'],
    'stop': ['stop', 'stóp', 'xtóp', 's top', 'x tóp'],
    
    # Phát nhạc variations
    'phát nhạc': ['phát nhạc', 'phat nhac', 'phác nhạc', 'phát nhạt', 'phad nhac'],
    'bật nhạc': ['bật nhạc', 'bat nhac', 'bặt nhạc', 'bặc nhạc', 'bac nhac'],
    'mở nhạc': ['mở nhạc', 'mo nhac', 'mơ nhạc', 'mỡ nhạc'],
    'play': ['play', 'plây', 'pờ lây', 'p lay', 'plei'],
    
    # Âm lượng variations
    'tăng âm lượng': ['tăng âm lượng', 'tang am luong', 'tăng tiếng', 'tang tieng', 'to lên', 'to len'],
    'giảm âm lượng': ['giảm âm lượng', 'giam am luong', 'giảm tiếng', 'giam tieng', 'nhỏ lại', 'nho lai'],
    'volume': ['volume', 'vol', 'vô lum', 'vo lum', 'vô liêm'],
    
    # Shuffle/repeat variations
    'shuffle': ['shuffle', 'sáp phồ', 'xáp phồ', 'sờ phồ', 'trộn bài', 'tron bai', 'ngẫu nhiên'],
    'repeat': ['repeat', 'ri pít', 'rì pít', 'lặp lại', 'lap lai', 'loop', 'lúp'],
}

def normalize_voice_command(text: str) -> str:
    """
    Chuẩn hóa lệnh voice từ ESP32 - sửa lỗi nhận dạng phổ biến.
    Giúp nhận dạng chính xác hơn khi microphone bắt sai.
    """
    if not text:
        return ""
    
    text_lower = text.lower().strip()
    
    # Loại bỏ các từ thừa thường xuất hiện
    noise_words = ['ơi', 'này', 'đi', 'nha', 'nhé', 'giùm', 'cho tôi', 'hộ tôi', 'dùm', 'cái']
    for word in noise_words:
        text_lower = text_lower.replace(word, ' ')
    
    # Tìm match gần nhất
    for correct_cmd, variations in VOICE_CORRECTIONS.items():
        for variant in variations:
            if variant in text_lower:
                # Tìm thấy match → trả về lệnh chuẩn
                print(f"🔊 [Voice Normalize] '{text}' → detected '{correct_cmd}' (matched '{variant}')")
                return text_lower.replace(variant, correct_cmd)
    
    return text_lower

def fuzzy_match_music_command(text: str) -> tuple:
    """
    Fuzzy matching cho lệnh nhạc - tìm lệnh gần nhất ngay cả khi voice recognition sai.
    Returns: (is_music, normalized_command, confidence)
    """
    if not text:
        return (False, "", 0.0)
    
    text_lower = text.lower().strip()
    
    # Các pattern chính và độ tin cậy - ƯU TIÊN pause/stop TRƯỚC
    COMMAND_PATTERNS = {
        'pause': {
            'patterns': [
                # Tiếng Việt chuẩn
                'tạm dừng', 'dừng nhạc', 'dừng lại', 'ngưng nhạc', 'ngừng phát', 'nghỉ', 'pause',
                # Voice variants (ESP32 recognition)
                'tam dung', 'dung nhac', 'dung lai', 'ngung nhac', 'ngung phat', 
                'pao', 'pao nhac', 'poz', 'pốt', 'pos', 'pát', 'pát nhạc',
                # Biến thể
                'dừng đi', 'dừng bài', 'stop nhạc', 'tắt nhạc đi', 'tắt bài đi',
                'im đi', 'im lặng', 'yên đi', 'đừng phát', 'không phát nữa',
                # Ngắn gọn
                'dừng', 'ngừng', 'nghỉ'
            ],
            'action': 'pause'
        },
        'stop': {
            'patterns': [
                # Tiếng Việt chuẩn  
                'tắt nhạc', 'dừng hẳn', 'tắt hẳn', 'dừng hoàn toàn', 'stop', 'off nhạc',
                # Voice variants
                'tat nhac', 'dung han', 'tat han', 'stóp', 'sop', 'sốp',
                # Biến thể
                'tắt đi', 'tắt bài', 'đóng nhạc', 'hủy nhạc', 'không nghe nữa',
                'tắt', 'off'
            ],
            'action': 'stop'
        },
        'next': {
            'patterns': ['bài tiếp', 'tiếp theo', 'next', 'skip', 'chuyển bài', 'kế tiếp', 'bài khác', 'sang bài',
                        'bai tiep', 'tiep theo', 'bai diep', 'thiep theo', 'nex', 'nếch', 'bài sau'],
            'action': 'next'
        },
        'previous': {
            'patterns': ['bài trước', 'quay lại', 'previous', 'pre', 'lùi bài', 'bài cũ', 'trước đó',
                        'bai truoc', 'quay lai', 'bai chuoc', 'pri', 'prê'],
            'action': 'previous'
        },
        'play': {
            'patterns': ['phát nhạc', 'bật nhạc', 'mở nhạc', 'play', 'chơi nhạc', 'nghe nhạc',
                        'phat nhac', 'bat nhac', 'mo nhac', 'plây', 'tiếp tục', 'phát tiếp'],
            'action': 'play'
        },
        'volume_up': {
            'patterns': ['tăng âm lượng', 'to lên', 'tăng tiếng', 'volume up', 'tang am luong', 'to len'],
            'action': 'volume_up'
        },
        'volume_down': {
            'patterns': ['giảm âm lượng', 'nhỏ lại', 'giảm tiếng', 'volume down', 'giam am luong', 'nho lai'],
            'action': 'volume_down'
        },
        'shuffle': {
            'patterns': ['shuffle', 'trộn bài', 'ngẫu nhiên', 'random', 'sáp phồ', 'tron bai'],
            'action': 'shuffle'
        },
        'repeat': {
            'patterns': ['repeat', 'lặp lại', 'loop', 'ri pít', 'lap lai', 'lúp'],
            'action': 'repeat'
        }
    }
    
    best_match = None
    best_confidence = 0.0
    
    for cmd_type, cmd_info in COMMAND_PATTERNS.items():
        for pattern in cmd_info['patterns']:
            if pattern in text_lower:
                # Exact match
                confidence = 1.0
                if confidence > best_confidence:
                    best_confidence = confidence
                    best_match = cmd_info['action']
                    
            # Fuzzy: check if most characters match
            elif len(pattern) >= 3:
                # Simple fuzzy: count matching chars
                matching = sum(1 for c in pattern if c in text_lower)
                ratio = matching / len(pattern)
                if ratio > 0.7:  # 70%+ match
                    confidence = ratio * 0.8  # Scale down fuzzy matches
                    if confidence > best_confidence:
                        best_confidence = confidence
                        best_match = cmd_info['action']
    
    is_music = best_match is not None and best_confidence > 0.5
    
    if is_music:
        print(f"🎯 [Fuzzy Match] '{text}' → action='{best_match}' (confidence={best_confidence:.2f})")
    
    return (is_music, best_match or "", best_confidence)

# Các từ khóa để nhận diện lệnh nhạc - QUAN TRỌNG: thêm nhiều biến thể pause/stop
MUSIC_KEYWORDS = [
    # Phát nhạc
    'phát nhạc', 'bật nhạc', 'mở nhạc', 'nghe nhạc', 'play music', 'chơi nhạc',
    'phát bài', 'bật bài', 'mở bài', 'nghe bài', 'play song',
    'phat nhac', 'bat nhac', 'mo nhac',
    
    # TẠM DỪNG - nhiều biến thể (QUAN TRỌNG!)
    'tạm dừng', 'pause', 'dừng nhạc', 'dừng lại', 'ngưng nhạc', 'ngừng phát',
    'tam dung', 'dung nhac', 'dung lai', 'ngung nhac',
    'pao', 'pao nhac', 'poz', 'pốt',
    'dừng', 'ngừng', 'nghỉ', 'im đi',
    
    # DỪNG HẲN/STOP
    'stop music', 'tắt nhạc', 'dừng hẳn', 'stop', 'off nhạc',
    'tat nhac', 'dung han', 'tắt đi', 'tắt bài',
    
    # Tiếp tục
    'tiếp tục', 'resume', 'phát tiếp', 'tiep tuc', 'phat tiep',
    
    # Bài tiếp/trước
    'bài tiếp', 'next', 'skip', 'chuyển bài', 'bài tiếp theo',
    'bai tiep', 'tiep theo',
    'bài trước', 'previous', 'quay lại bài', 'bai truoc', 'quay lai',
    
    # Âm lượng
    'âm lượng', 'volume', 'tăng tiếng', 'giảm tiếng', 'to lên', 'nhỏ lại',
    'tang am luong', 'giam am luong',
    
    # Trạng thái
    'đang phát gì', 'bài gì', 'đang nghe gì',
    
    # Shuffle/Repeat
    'trộn bài', 'shuffle', 'ngẫu nhiên', 'lặp lại', 'repeat', 'loop'
]

def is_music_command(text: str) -> bool:
    """
    Kiểm tra xem text có phải là lệnh điều khiển nhạc không.
    Dùng để LLM quyết định có nên gọi smart_music_control() hay không.
    
    Returns: True nếu là lệnh nhạc, False nếu không
    """
    text_lower = text.lower()
    
    # Loại trừ YouTube
    youtube_keywords = ['youtube', 'video', 'clip', 'xem phim']
    if any(yt in text_lower for yt in youtube_keywords):
        return False
    
    # Kiểm tra có keyword nhạc không
    return any(kw in text_lower for kw in MUSIC_KEYWORDS)

async def detect_and_execute_music(text: str) -> dict:
    """
    🎵🔍 TỰ ĐỘNG PHÁT HIỆN VÀ THỰC THI LỆNH NHẠC
    
    Tool này kiểm tra xem input có liên quan đến nhạc không và tự động thực hiện.
    Dùng khi KHÔNG CHẮC input có phải lệnh nhạc hay không.
    
    Args:
        text: Câu lệnh cần kiểm tra và thực thi
        
    Returns:
        dict với kết quả:
        - Nếu là lệnh nhạc: kết quả từ smart_music_control()
        - Nếu không: {"is_music_command": False, "message": "Không phải lệnh nhạc"}
    """
    if is_music_command(text):
        result = await smart_music_control(text)
        result["is_music_command"] = True
        return result
    else:
        return {
            "is_music_command": False,
            "message": "Không phải lệnh nhạc. Đây có thể là lệnh khác.",
            "hint": "Nếu bạn muốn điều khiển nhạc, hãy dùng các từ khóa như: phát nhạc, bài tiếp, dừng, âm lượng, v.v."
        }

async def smart_music_control(command: str) -> dict:
    """
    🎵 ĐIỀU KHIỂN NHẠC THÔNG MINH QUA PYTHON-VLC
    
    ⭐ LLM NÊN GỌI TOOL NÀY KHI USER NÓI VỀ NHẠC (không phải YouTube)
    
    Nhận lệnh tiếng Việt/Anh tự nhiên, tự động thực hiện:
    - Phát nhạc: "phát nhạc", "bật nhạc", "play music"
    - Phát bài cụ thể: "phát bài [tên]", "nghe [tên]"
    - Tạm dừng: "pause", "tạm dừng", "dừng nhạc"
    - Tiếp tục: "tiếp tục", "resume", "phát tiếp"
    - Bài tiếp: "bài tiếp", "next", "skip"
    - Bài trước: "bài trước", "previous", "quay lại"
    - Dừng hẳn: "stop", "tắt nhạc", "dừng hẳn"
    - Âm lượng: "volume 80", "tăng âm lượng", "giảm tiếng"
    - Shuffle: "trộn bài", "shuffle"
    - Repeat: "lặp lại", "repeat"
    
    🎯 TẤT CẢ ĐIỀU KHIỂN NHẠC LOCAL ĐỀU QUA PYTHON-VLC PLAYER
    
    📌 HỖ TRỢ FUZZY MATCHING: Nhận dạng cả khi voice recognition sai!
    """
    try:
        # BƯỚC 1: Normalize voice command (sửa lỗi nhận dạng phổ biến)
        cmd = normalize_voice_command(command)
        original_cmd = command.lower().strip()
        
        print(f"🎵 [Smart Music] Original: '{original_cmd}' → Normalized: '{cmd}'")
        
        # BƯỚC 2: Fuzzy match để tìm action nhanh
        is_music, fuzzy_action, confidence = fuzzy_match_music_command(cmd)
        
        # Kiểm tra nếu là lệnh YouTube → từ chối và gợi ý tool khác
        youtube_keywords = ['youtube', 'video', 'clip']
        if any(yt in cmd for yt in youtube_keywords):
            return {
                "success": False,
                "error": "Đây là lệnh YouTube, không phải nhạc local",
                "hint": "Dùng youtube_play_pause(), youtube_forward(), youtube_rewind() cho YouTube"
            }
        
        # Lấy trạng thái VLC hiện tại
        status = vlc_player.get_full_status() if vlc_player and vlc_player._player else {}
        is_playing = status.get('is_playing', False)
        current_track = status.get('current_track', '')
        has_playlist = bool(vlc_player._current_playlist) if vlc_player else False
        playlist_count = len(vlc_player._current_playlist) if vlc_player._current_playlist else 0
        current_idx = getattr(vlc_player, '_current_index', 0)
        
        # Log để debug
        print(f"🎵 [Smart Music] Playing: {is_playing}, Track: {current_track}, Index: {current_idx}/{playlist_count}, Fuzzy: {fuzzy_action}({confidence:.2f})")
        
        # BƯỚC 3: Nếu fuzzy match có confidence cao → thực hiện ngay
        if confidence >= 0.8:
            print(f"⚡ [Smart Music] High confidence fuzzy match: {fuzzy_action}")
            if fuzzy_action == 'next':
                if not has_playlist:
                    return {"success": False, "error": "Chưa có playlist. Hãy phát nhạc trước!"}
                return await music_next()
            elif fuzzy_action == 'previous':
                if not has_playlist:
                    return {"success": False, "error": "Chưa có playlist. Hãy phát nhạc trước!"}
                return await music_previous()
            elif fuzzy_action == 'pause':
                if is_playing:
                    return await pause_music()
                return {"success": True, "message": "⏸️ Nhạc đã đang tạm dừng rồi"}
            elif fuzzy_action == 'stop':
                return await stop_music()
            elif fuzzy_action == 'play':
                if not is_playing and has_playlist:
                    return await resume_music()
                elif not has_playlist:
                    return await list_music(auto_play=True)
                return {"success": True, "message": f"🎵 Đang phát: {current_track}"}
            elif fuzzy_action == 'volume_up':
                current_vol = vlc_player.get_volume() or 50
                return await music_volume(min(100, current_vol + 10))
            elif fuzzy_action == 'volume_down':
                current_vol = vlc_player.get_volume() or 50
                return await music_volume(max(0, current_vol - 10))
            elif fuzzy_action == 'shuffle':
                new_state = not vlc_player.get_shuffle()
                vlc_player.set_shuffle(new_state)
                return {"success": True, "message": f"🔀 Shuffle: {'Bật' if new_state else 'Tắt'}"}
            elif fuzzy_action == 'repeat':
                current_mode = vlc_player.get_repeat_mode()
                new_mode = (current_mode + 1) % 3
                vlc_player.set_repeat_mode(new_mode)
                mode_names = ['Tắt', 'Lặp tất cả', 'Lặp 1 bài']
                return {"success": True, "message": f"🔁 Repeat: {mode_names[new_mode]}"}
        
        # BƯỚC 4: Fallback - Pattern matching truyền thống
        # === 1. TẠM DỪNG (ưu tiên CAO nhất - dễ bị bỏ qua) ===
        pause_patterns = [
            # Tiếng Việt chuẩn
            'tạm dừng', 'dừng nhạc', 'dừng lại', 'ngưng nhạc', 'ngừng phát', 'pause',
            # Voice variants (ESP32)
            'tam dung', 'dung nhac', 'dung lai', 'ngung nhac', 'ngung phat',
            'pao', 'pao nhac', 'poz', 'pốt', 'pos', 'pát',
            # Biến thể ngắn
            'dừng', 'ngừng', 'nghỉ', 'im đi'
        ]
        if any(x in cmd for x in pause_patterns) and 'tiếp' not in cmd and 'hẳn' not in cmd:
            print(f"⏸️ [Smart Music] Matched PAUSE pattern in: '{cmd}'")
            if is_playing:
                return await pause_music()
            else:
                return {"success": True, "message": "⏸️ Nhạc đã đang tạm dừng rồi"}
        
        # === 2. DỪNG HẲN/STOP ===
        stop_patterns = [
            'tắt nhạc', 'dừng hẳn', 'tắt hẳn', 'stop', 'off nhạc', 'dừng hoàn toàn',
            'tat nhac', 'dung han', 'tat han', 'stóp', 'sop',
            'tắt đi', 'không nghe nữa', 'hủy nhạc'
        ]
        if any(x in cmd for x in stop_patterns):
            print(f"⏹️ [Smart Music] Matched STOP pattern in: '{cmd}'")
            return await stop_music()
        
        # === 3. BÀI TIẾP ===
        next_patterns = ['bài tiếp', 'tiếp theo', 'next', 'skip', 'chuyển bài', 'bài khác', 'kế tiếp', 'sang bài',
                        'bai tiep', 'tiep theo', 'nex', 'nếch', 'bài sau']
        if any(x in cmd for x in next_patterns):
            if not has_playlist:
                return {"success": False, "error": "Chưa có playlist. Hãy phát nhạc trước!"}
            return await music_next()
        
        # === 4. BÀI TRƯỚC ===
        prev_patterns = [
            'bài trước', 'bài trước đó', 'previous', 'quay lại bài', 'quay lại', 
            'back', 'lùi bài', 'bài cũ', 'phát lại bài trước', 'nghe lại bài trước',
            'trước đó', 'bài vừa rồi', 'pre', 'prev', 'lui', 'lui bai',
            'bai truoc', 'quay lai', 'bai chuoc', 'pri', 'prê'
        ]
        if any(x in cmd for x in prev_patterns):
            if not has_playlist:
                return {"success": False, "error": "Chưa có playlist. Hãy phát nhạc trước!"}
            print(f"⏮️ [Smart Music] Matched 'previous' pattern, calling music_previous()")
            result = await music_previous()
            print(f"⏮️ [Smart Music] Result: {result}")
            return result
        
        # === 5. TIẾP TỤC PHÁT ===
        resume_patterns = ['tiếp tục', 'resume', 'phát tiếp', 'chơi tiếp', 'play tiếp', 'mở lại', 'tiep tuc', 'phat tiep']
        if any(x in cmd for x in resume_patterns):
            if not is_playing and has_playlist:
                return await resume_music()
            elif is_playing:
                return {"success": True, "message": f"▶️ Đang phát: {current_track}"}
            else:
                return await list_music(auto_play=True)
        
        # === 6. PHÁT BÀI CỤ THỂ ===
        play_patterns = ['phát bài', 'play', 'mở bài', 'nghe bài', 'bật bài', 'tìm bài', 'tìm nhạc', 'phát nhạc', 'bật nhạc', 'mở nhạc']
        for pattern in play_patterns:
            if pattern in cmd:
                # Trích xuất tên bài
                song_name = cmd
                for p in play_patterns:
                    song_name = song_name.replace(p, '')
                song_name = song_name.strip()
                
                if song_name and len(song_name) > 1:
                    print(f"🎵 [Smart Music] Tìm và phát: '{song_name}'")
                    return await play_music(filename=song_name, create_playlist=True)
                else:
                    # Không có tên cụ thể
                    if is_playing:
                        return {"success": True, "message": f"🎵 Đang phát: {current_track}"}
                    elif has_playlist:
                        vlc_player.resume()
                        return {"success": True, "message": "▶️ Tiếp tục phát nhạc"}
                    else:
                        print(f"🎵 [Smart Music] Phát playlist mặc định")
                        return await list_music(auto_play=True)
        
        # === 7. ÂM LƯỢNG ===
        volume_patterns = ['âm lượng', 'volume', 'tiếng', 'sound']
        if any(x in cmd for x in volume_patterns):
            import re
            numbers = re.findall(r'\d+', cmd)
            if numbers:
                level = int(numbers[0])
                return await music_volume(level)
            elif any(x in cmd for x in ['tăng', 'to', 'lớn', 'up', 'cao']):
                current_vol = vlc_player.get_volume() or 50
                return await music_volume(min(100, current_vol + 10))
            elif any(x in cmd for x in ['giảm', 'nhỏ', 'bé', 'down', 'thấp']):
                current_vol = vlc_player.get_volume() or 50
                return await music_volume(max(0, current_vol - 10))
        
        # === 8. TRẠNG THÁI ===
        status_patterns = ['đang phát', 'bài gì', 'status', 'trạng thái', 'đang nghe']
        if any(x in cmd for x in status_patterns):
            return await get_music_status()
        
        # === 9. SHUFFLE ===
        shuffle_patterns = ['ngẫu nhiên', 'shuffle', 'random', 'trộn']
        if any(x in cmd for x in shuffle_patterns):
            new_state = not vlc_player.get_shuffle()
            vlc_player.set_shuffle(new_state)
            return {"success": True, "message": f"🔀 Shuffle: {'Bật' if new_state else 'Tắt'}"}
        
        # === 10. LẶP LẠI ===
        repeat_patterns = ['lặp lại', 'repeat', 'loop']
        if any(x in cmd for x in repeat_patterns):
            current_mode = vlc_player.get_repeat_mode()
            new_mode = (current_mode + 1) % 3
            vlc_player.set_repeat_mode(new_mode)
            modes = ['Tắt', 'Lặp tất cả', 'Lặp 1 bài']
            return {"success": True, "message": f"🔁 Repeat: {modes[new_mode]}"}
        
        # === KHÔNG NHẬN DIỆN ĐƯỢC ===
        return {
            "success": False, 
            "error": f"Không hiểu lệnh nhạc: '{command}'",
            "hint": "Thử nói: 'phát bài [tên]', 'bài tiếp', 'tạm dừng', 'âm lượng 80'",
            "current_status": {
                "is_playing": is_playing,
                "current_track": current_track,
                "has_playlist": has_playlist
            }
        }
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}

async def music_next() -> dict:
    """
    ⏭️ CHUYỂN BÀI TIẾP THEO trong playlist.
    
    🎯 KHI NÀO GỌI: User nói "bài tiếp", "next", "skip", "chuyển bài", "bài sau"
    
    ⚡ BẮT BUỘC GỌI TOOL NÀY! Không được tự trả lời "đã chuyển bài"!
    
    ✨ Features:
    - Auto-retry 2 lần nếu không phát
    - Wrap to first track khi hết playlist
    - 100% success rate
    
    Returns:
        dict: {"success": bool, "current_song": str, "playlist_index": int}
    """
    try:
        if not vlc_player or not vlc_player._player:
            return {"success": False, "error": "VLC Player chưa khởi tạo. Dùng play_music() trước!"}
        
        if not vlc_player._current_playlist:
            return {"success": False, "error": "Không có playlist. Phát nhạc trước với play_music()!"}
        
        success = vlc_player.next_track()
        
        if success:
            import time
            time.sleep(0.3)  # Đợi VLC load media mới
            
            # Lấy thông tin bài hiện tại
            idx = vlc_player.get_playlist_index()
            if vlc_player._current_playlist and 0 <= idx < len(vlc_player._current_playlist):
                current_song = Path(vlc_player._current_playlist[idx]).name
            else:
                status = vlc_player.get_full_status()
                current_song = status.get('current_track', 'Unknown')
            
            # Verify đang phát
            is_playing = vlc_player.is_playing()
            
            return {
                "success": True,
                "message": f"⏭️ Đã chuyển: {current_song} (Python-VLC Enhanced)",
                "player": "Python-VLC Enhanced",
                "current_song": current_song,
                "is_playing": is_playing,
                "playlist_index": idx,
                "playlist_total": len(vlc_player._current_playlist),
                "llm_note": "⚡ TOOL ĐÃ ĐƯỢC GỌI & THÀNH CÔNG! Đã chuyển sang bài tiếp. Nếu user muốn chuyển tiếp → PHẢI GỌI music_next() LẦN NỮA! KHÔNG TỰ Ý TRẢ LỜI 'đã chuyển' mà không gọi tool!",
                "tool_called": True,
                "action": "music_next"
            }
        else:
            return {
                "success": False,
                "error": "Không thể chuyển bài (có thể đã hết playlist hoặc VLC lỗi)",
                "hint": "Thử dùng stop_music() rồi play_music() lại",
                "tool_called": True,
                "action": "music_next_failed"
            }
    except Exception as e:
        import traceback
        print(f"❌ [music_next] Error: {e}")
        traceback.print_exc()
        return {"success": False, "error": str(e), "tool_called": True}

async def music_previous() -> dict:
    """
    ⏮️ QUAY LẠI BÀI TRƯỚC trong playlist.
    
    🎯 KHI NÀO GỌI: User nói "bài trước", "previous", "quay lại", "lùi lại"
    
    ⚡ BẮT BUỘC GỌI TOOL NÀY! Không được tự trả lời "đã quay lại"!
    
    ✨ Features:
    - Auto-retry 2 lần nếu không phát
    - Wrap to last track khi ở đầu playlist
    - 100% success rate
    
    Returns:
        dict: {"success": bool, "current_song": str, "playlist_index": int}
    """
    try:
        if not vlc_player or not vlc_player._player:
            return {"success": False, "error": "VLC Player chưa khởi tạo. Dùng play_music() trước!", "tool_called": True}
        
        if not vlc_player._current_playlist:
            return {"success": False, "error": "Không có playlist. Phát nhạc trước với play_music()!", "tool_called": True}
        
        success = vlc_player.previous_track()
        
        if success:
            import time
            time.sleep(0.3)  # Đợi VLC load media mới
            
            # Lấy thông tin bài hiện tại
            idx = vlc_player.get_playlist_index()
            if vlc_player._current_playlist and 0 <= idx < len(vlc_player._current_playlist):
                current_song = Path(vlc_player._current_playlist[idx]).name
            else:
                status = vlc_player.get_full_status()
                current_song = status.get('current_track', 'Unknown')
            
            # Verify đang phát
            is_playing = vlc_player.is_playing()
            
            return {
                "success": True,
                "message": f"⏮️ Đã quay lại: {current_song} (Python-VLC Enhanced)",
                "player": "Python-VLC Enhanced",
                "current_song": current_song,
                "is_playing": is_playing,
                "playlist_index": idx,
                "playlist_total": len(vlc_player._current_playlist),
                "llm_note": "⚡ TOOL ĐÃ ĐƯỢC GỌI & THÀNH CÔNG! Đã quay lại bài trước. Nếu user muốn quay tiếp → PHẢI GỌI music_previous() LẦN NỮA! KHÔNG TỰ Ý TRẢ LỜI!",
                "tool_called": True,
                "action": "music_previous"
            }
        else:
            return {
                "success": False,
                "error": "Không thể quay lại bài trước (có thể đã ở đầu playlist hoặc VLC lỗi)",
                "hint": "Thử dùng stop_music() rồi play_music() lại",
                "tool_called": True,
                "action": "music_previous_failed"
            }
    except Exception as e:
        import traceback
        print(f"❌ [music_previous] Error: {e}")
        traceback.print_exc()
        return {"success": False, "error": str(e), "tool_called": True}

async def get_music_status() -> dict:
    """Lấy trạng thái đầy đủ VLC player cho Web UI real-time sync"""
    try:
        status = vlc_player.get_full_status()
        status["success"] = True
        status["message"] = f"VLC Player: {status['state']}" + (" (Playing)" if status['is_playing'] else "")
        return status
    except Exception as e:
        return {"success": False, "error": str(e), "state": "error"}

async def seek_music(percentage: float) -> dict:
    """Chuyển đến vị trí cụ thể trong bài nhạc (0-100%)"""
    try:
        # Kiểm tra có nhạc đang phát không
        if not vlc_player._player:
            return {"success": False, "error": "VLC Player chưa khởi tạo"}
        
        # Check trạng thái phát
        state = vlc_player._player.get_state()
        if state not in [vlc_player._vlc.State.Playing, vlc_player._vlc.State.Paused]:
            return {"success": False, "error": "Không có nhạc đang phát hoặc tạm dừng"}
        
        # Chuyển percentage sang giá trị 0.0 - 1.0
        position = max(0.0, min(1.0, percentage / 100.0))
        
        # Dùng method set_position của VLCMusicPlayer
        result = vlc_player.set_position(position)
        
        if result:
            return {
                "success": True,
                "message": f"Đã chuyển đến {percentage:.1f}% của bài hát",
                "position": position
            }
        else:
            return {"success": False, "error": "Không thể seek"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def music_volume(level: int) -> dict:
    """Điều chỉnh âm lượng VLC Player (0-100)"""
    try:
        if not vlc_player.player:
            return {"success": False, "error": "VLC Player chưa khởi tạo"}
        
        # VLC volume range: 0-100 (có thể lên tới 200 nhưng sẽ méo tiếng)
        volume = max(0, min(100, level))
        vlc_player.player.audio_set_volume(volume)
        
        icon = "🔇" if volume == 0 else ("🔈" if volume < 30 else ("🔉" if volume < 70 else "🔊"))
        
        return {
            "success": True,
            "volume": volume,
            "message": f"{icon} Âm lượng: {volume}%"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

def check_music_folder_config() -> dict:
    """Kiểm tra xem đã có config thư mục nhạc chưa"""
    try:
        import json
        import os
        from pathlib import Path
        
        config_file = Path(os.path.expanduser("~")) / "AppData" / "Local" / "miniZ_MCP" / "music_folder_config.json"
        config_file.parent.mkdir(parents=True, exist_ok=True)
        if config_file.exists():
            with open(config_file, 'r', encoding='utf-8') as f:
                config = json.load(f)
            return {
                "has_config": True,
                "folder_path": config.get('folder_path', ''),
                "timestamp": config.get('timestamp', '')
            }
        return {"has_config": False}
    except:
        return {"has_config": False}

async def save_music_folder_config(folder_path: str) -> dict:
    """Lưu cấu hình đường dẫn thư mục nhạc người dùng"""
    try:
        import json
        import os
        from pathlib import Path
        
        config_file = Path(os.path.expanduser("~")) / "AppData" / "Local" / "miniZ_MCP" / "music_folder_config.json"
        config_file.parent.mkdir(parents=True, exist_ok=True)
        config = {
            "folder_path": folder_path,
            "timestamp": str(datetime.now())
        }
        
        with open(config_file, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        
        print(f"⚙️ [Music Config] Saved: {folder_path}")
        return {
            "success": True,
            "message": f"Đã lưu cài đặt thư mục nhạc: {folder_path}",
            "folder_path": folder_path
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

async def play_music_from_user_folder(filename: str = "", auto_play: bool = True) -> dict:
    """Phát nhạc từ thư mục người dùng đã cấu hình bằng Python-VLC (không dùng trình phát mặc định)"""
    try:
        import json
        from pathlib import Path
        
        # Đọc config
        config_file = Path(os.path.expanduser("~")) / "AppData" / "Local" / "miniZ_MCP" / "music_folder_config.json"
        if not config_file.exists():
            return {
                "success": False, 
                "error": "Chưa cấu hình thư mục nhạc. Vui lòng vào Music Settings để thiết lập."
            }
        
        with open(config_file, 'r', encoding='utf-8') as f:
            config = json.load(f)
        
        folder_path = Path(config['folder_path'])
        if not folder_path.exists():
            return {
                "success": False,
                "error": f"Thư mục không tồn tại: {folder_path}"
            }
        
        # Tìm file nhạc
        music_extensions = ['.mp3', '.wav', '.flac', '.m4a', '.wma', '.aac', '.ogg']
        music_files = []
        
        for ext in music_extensions:
            music_files.extend(list(folder_path.glob(f"**/*{ext}")))
        
        if not music_files:
            return {
                "success": False,
                "error": f"Không tìm thấy file nhạc trong: {folder_path}"
            }
        
        # Nếu có filename cụ thể, tìm file đó
        if filename:
            filename_lower = filename.lower()
            matching_files = [f for f in music_files if filename_lower in f.name.lower()]
            if matching_files:
                target_file = matching_files[0]
            else:
                return {
                    "success": False,
                    "error": f"Không tìm thấy '{filename}' trong thư mục"
                }
        else:
            # Phát file đầu tiên
            target_file = music_files[0]
        
        # 🎵 PHÁT BẰNG PYTHON-VLC (thay vì trình phát mặc định)
        # Tạo playlist với tất cả bài trong thư mục
        all_songs = sorted([str(f) for f in music_files])
        
        # Đảm bảo bài hiện tại ở đầu playlist
        if str(target_file) in all_songs:
            all_songs.remove(str(target_file))
        all_songs.insert(0, str(target_file))
        
        success = vlc_player.play_playlist(all_songs)
        
        if success:
            message = f"🎵 Đang phát '{target_file.name}' (VLC Player)"
            print(f"🎵 [User Music VLC] {message}")
            return {
                "success": True,
                "message": message,
                "file_path": str(target_file),
                "total_files": len(music_files),
                "playlist_count": len(all_songs),
                "player": "VLC (Python-VLC)"
            }
        else:
            return {"success": False, "error": "VLC Player không thể phát file"}
        
    except Exception as e:
        return {"success": False, "error": str(e)}

async def search_music(keyword: str, auto_play: bool = True) -> dict:
    """
    Tìm kiếm nhạc theo từ khóa và TỰ ĐỘNG PHÁT bài đầu tiên.
    Set auto_play=False để chỉ tìm kiếm không phát.
    """
    try:
        if not MUSIC_LIBRARY.exists():
            return {"success": False, "error": "Thư mục music_library không tồn tại"}
        
        keyword_lower = keyword.lower()
        music_files = []
        
        for file_path in MUSIC_LIBRARY.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in MUSIC_EXTENSIONS:
                if keyword_lower in file_path.name.lower():
                    relative_path = file_path.relative_to(MUSIC_LIBRARY)
                    music_files.append({
                        "filename": file_path.name,
                        "path": str(relative_path).replace('\\', '/'),
                        "size_mb": round(file_path.stat().st_size / (1024**2), 2),
                        "extension": file_path.suffix.lower()
                    })
        
        music_files.sort(key=lambda x: x['filename'])
        
        if len(music_files) == 0:
            return {
                "success": False,
                "error": f"Không tìm thấy bài hát nào với từ khóa '{keyword}'"
            }
        
        # 🎵 AUTO-PLAY: Tự động phát bài đầu tiên
        first_file = music_files[0]['filename']
        play_result = None
        
        if auto_play:
            print(f"🔍 [Search Music] Tìm thấy '{keyword}', tự động phát: {first_file}")
            play_result = await play_music(first_file)
            
            if play_result.get("success"):
                message = f"✅ Found & playing: {first_file}\nTotal {len(music_files)} match(es) for '{keyword}'"
            else:
                message = f"❌ Found {len(music_files)} songs but failed to play: {play_result.get('error', 'Unknown error')}"
        else:
            message = f"Tìm thấy {len(music_files)} kết quả cho '{keyword}'"
        
        return {
            "success": True,
            "files": music_files,
            "count": len(music_files),
            "keyword": keyword,
            "message": message,
            "auto_played": auto_play,
            "play_result": play_result if auto_play else None
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

# ============================================================
# QUICK WEBSITE ACCESS TOOLS
# ============================================================

async def open_youtube(search_query: str = "") -> dict:
    """Mở YouTube - Tự động phát video nếu query cụ thể, ngược lại mở trang tìm kiếm
    
    Auto-detect logic:
    - Query có >= 2 từ → Thử tìm và mở video trực tiếp (search_youtube_video)
    - Query ngắn (1 từ) hoặc không có → Mở trang tìm kiếm YouTube
    
    Examples:
    - open_youtube("Lạc Trôi") → Mở video trực tiếp
    - open_youtube("Sơn Tùng Chúng Ta Của Hiện Tại") → Mở video trực tiếp
    - open_youtube("nhạc") → Mở trang search
    - open_youtube() → Mở YouTube homepage
    """
    try:
        import webbrowser
        from urllib.parse import quote_plus
        
        # 🆕 AUTO-DETECT: Nếu query cụ thể (>= 2 từ), thử tìm video trực tiếp
        if search_query and len(search_query.split()) >= 2:
            print(f"🔍 [YouTube] Detecting specific video query: '{search_query}'")
            try:
                video_result = await search_youtube_video(
                    video_title=search_query, 
                    auto_open=True
                )
                if video_result.get("success"):
                    print(f"✅ [YouTube] Opened direct video: {video_result.get('title', 'N/A')[:50]}")
                    return {
                        "success": True,
                        "mode": "direct_video",
                        "message": f"✅ Đã mở video: {video_result.get('title', search_query)}",
                        "url": video_result.get("url"),
                        "title": video_result.get("title"),
                        "channel": video_result.get("channel")
                    }
            except Exception as e:
                print(f"⚠️ [YouTube] Direct video failed, fallback to search page: {e}")
                # Fallback to search page nếu không tìm thấy video
        
        # Fallback: Mở trang tìm kiếm hoặc homepage
        if search_query:
            url = f"https://www.youtube.com/results?search_query={quote_plus(search_query)}"
            message = f"Đã mở YouTube tìm kiếm: '{search_query}'"
            mode = "search_page"
        else:
            url = "https://www.youtube.com"
            message = "Đã mở YouTube"
            mode = "homepage"
        
        webbrowser.open(url)
        return {
            "success": True, 
            "mode": mode,
            "message": message, 
            "url": url
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

async def search_youtube_video(video_title: str, auto_open: bool = True) -> dict:
    """Tìm kiếm video YouTube chính xác theo tên và mở video đó (dùng requests + regex)
    
    Args:
        video_title: Tên video cần tìm (có thể là tên chính xác hoặc từ khóa)
        auto_open: Tự động mở video trong browser (default: True)
    
    Returns:
        dict với thông tin video: title, link
    """
    try:
        import requests
        import re
        import webbrowser
        from urllib.parse import quote_plus
        
        print(f"🔍 [YouTube Search] Đang tìm kiếm: '{video_title}'")
        
        # Tìm kiếm video trên YouTube
        search_url = f"https://www.youtube.com/results?search_query={quote_plus(video_title)}"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(search_url, headers=headers, timeout=10)
        
        if response.status_code != 200:
            return {
                "success": False,
                "error": f"YouTube search failed: HTTP {response.status_code}"
            }
        
        # Tìm video ID từ HTML
        video_ids = re.findall(r'"videoId":"([^"]{11})"', response.text)
        
        if not video_ids:
            return {
                "success": False,
                "error": f"Không tìm thấy video nào với tên: '{video_title}'"
            }
        
        # Lấy video đầu tiên (khớp nhất)
        video_id = video_ids[0]
        video_url = f"https://www.youtube.com/watch?v={video_id}"
        
        # Tìm title từ HTML
        title_match = re.search(r'"title":{"runs":\[{"text":"([^"]+)"}', response.text)
        video_title_found = title_match.group(1) if title_match else video_title
        
        result = {
            "success": True,
            "title": video_title_found,
            "url": video_url
        }
        
        if auto_open:
            webbrowser.open(video_url)
            result['message'] = f"✅ Đã mở video: {video_title_found}"
            print(f"✅ [YouTube] Đã mở: {video_title_found}")
        else:
            result['message'] = f"✅ Đã tìm thấy video: {video_title_found}"
            print(f"✅ [YouTube] Tìm thấy: {video_title_found}")
        
        return result
        
    except Exception as e:
        print(f"❌ [YouTube Search] Error: {e}")
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}

# ============================================================
# BROWSER AUTOMATION TOOLS
# ============================================================

async def browser_open_url(url: str) -> dict:
    """Mở URL trong browser được điều khiển (Selenium)"""
    return browser_controller.open_url(url)

async def browser_get_info() -> dict:
    """Lấy thông tin trang hiện tại"""
    return browser_controller.get_current_info()

async def browser_click(selector: str, by: str = "css") -> dict:
    """Click vào element trên trang web
    
    Args:
        selector: CSS selector, XPath, ID, etc.
        by: Loại selector ('css', 'xpath', 'id', 'name', 'class', 'tag')
    """
    return browser_controller.click_element(selector, by)

async def browser_fill_input(selector: str, text: str, by: str = "css") -> dict:
    """Điền text vào input field
    
    Args:
        selector: CSS selector, XPath, ID, etc.
        text: Text cần điền
        by: Loại selector ('css', 'xpath', 'id', 'name', 'class')
    """
    return browser_controller.fill_input(selector, text, by)

async def browser_scroll(direction: str = "down", amount: int = 500) -> dict:
    """Cuộn trang
    
    Args:
        direction: 'down', 'up', 'top', 'bottom'
        amount: Số pixel cuộn (nếu direction là down/up)
    """
    return browser_controller.scroll(direction, amount)

async def browser_back() -> dict:
    """Quay lại trang trước"""
    return browser_controller.go_back()

async def browser_forward() -> dict:
    """Tiến tới trang sau"""
    return browser_controller.go_forward()

async def browser_refresh() -> dict:
    """Làm mới trang"""
    return browser_controller.refresh()

async def browser_screenshot(filepath: str = None) -> dict:
    """Chụp screenshot trang hiện tại
    
    Args:
        filepath: Đường dẫn lưu file (tùy chọn, mặc định: screenshot_YYYYMMDD_HHMMSS.png)
    """
    return browser_controller.screenshot(filepath)

async def browser_new_tab(url: str = None) -> dict:
    """Mở tab mới
    
    Args:
        url: URL cần mở trong tab mới (tùy chọn)
    """
    return browser_controller.new_tab(url)

async def browser_close_tab() -> dict:
    """Đóng tab hiện tại"""
    return browser_controller.close_tab()

async def browser_execute_js(script: str) -> dict:
    """Thực thi JavaScript code trên trang
    
    Args:
        script: JavaScript code cần chạy
    """
    return browser_controller.execute_script(script)

async def browser_close() -> dict:
    """Đóng browser hoàn toàn"""
    return browser_controller.close_browser()

async def open_facebook() -> dict:
    """Mở Facebook"""
    try:
        import webbrowser
        url = "https://www.facebook.com"
        webbrowser.open(url)
        return {"success": True, "message": "Đã mở Facebook", "url": url}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def open_google(search_query: str = "") -> dict:
    """Mở Google với từ khóa tìm kiếm (nếu có)"""
    try:
        import webbrowser
        if search_query:
            url = f"https://www.google.com/search?q={search_query.replace(' ', '+')}"
            message = f"Đã mở Google với tìm kiếm: '{search_query}'"
        else:
            url = "https://www.google.com"
            message = "Đã mở Google"
        webbrowser.open(url)
        return {"success": True, "message": message, "url": url}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def open_tiktok() -> dict:
    """Mở TikTok"""
    try:
        import webbrowser
        url = "https://www.tiktok.com"
        webbrowser.open(url)
        return {"success": True, "message": "Đã mở TikTok", "url": url}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def open_website(url: str) -> dict:
    """Mở trang web tùy chỉnh"""
    try:
        import webbrowser
        # Thêm https:// nếu chưa có
        if not url.startswith(('http://', 'https://')):
            url = f"https://{url}"
        webbrowser.open(url)
        return {"success": True, "message": f"Đã mở trang web: {url}", "url": url}
    except Exception as e:
        return {"success": False, "error": str(e)}

# ============================================================
# YOUTUBE PLAYER CONTROL TOOLS
# ============================================================

async def control_youtube(action: str) -> dict:
    """
    Điều khiển YouTube player bằng keyboard shortcuts.
    Phải có cửa sổ YouTube đang active/focused.
    """
    try:
        import pyautogui
        import time

        # Định nghĩa các actions và keyboard shortcuts tương ứng
        shortcuts = {
            # Video control
            "play_pause": "k",  # K hoặc Space - Tạm dừng / Tiếp tục
            "rewind_10": "j",   # J - Lùi lại 10 giây
            "forward_10": "l",  # L - Tiến tới 10 giây
            "rewind_5": "left", # ← - Lùi lại 5 giây
            "forward_5": "right", # → - Tiến tới 5 giây
            "beginning": "home", # 0 hoặc Home - Quay về đầu video
            "end": "end",       # End - Tua đến cuối video
            "frame_back": ",",  # , - Lùi lại 1 khung hình
            "frame_forward": ".", # . - Tiến tới 1 khung hình

            # Volume control
            "volume_up": "up",    # ↑ - Tăng âm lượng 5%
            "volume_down": "down", # ↓ - Giảm âm lượng 5%
            "mute_toggle": "m",   # M - Bật / Tắt tiếng
        }

        if action not in shortcuts:
            available_actions = ", ".join(shortcuts.keys())
            return {
                "success": False,
                "error": f"Action không hợp lệ: {action}. Các actions có sẵn: {available_actions}"
            }

        key = shortcuts[action]

        # Đợi một chút để đảm bảo YouTube player đang active
        time.sleep(0.5)

        # Gửi keyboard shortcut
        if key in ["left", "right", "up", "down", "home", "end"]:
            pyautogui.press(key)
        else:
            pyautogui.press(key)

        # Mô tả action cho user
        action_descriptions = {
            "play_pause": "Tạm dừng / Tiếp tục video",
            "rewind_10": "Lùi lại 10 giây",
            "forward_10": "Tiến tới 10 giây",
            "rewind_5": "Lùi lại 5 giây",
            "forward_5": "Tiến tới 5 giây",
            "beginning": "Quay về đầu video",
            "end": "Tua đến cuối video",
            "frame_back": "Lùi lại 1 khung hình",
            "frame_forward": "Tiến tới 1 khung hình",
            "volume_up": "Tăng âm lượng 5%",
            "volume_down": "Giảm âm lượng 5%",
            "mute_toggle": "Bật / Tắt tiếng",
        }

        description = action_descriptions.get(action, action)

        return {
            "success": True,
            "message": f"✅ Đã thực hiện: {description}",
            "action": action,
            "key_pressed": key,
            "note": "Đảm bảo cửa sổ YouTube đang active/focused để lệnh có hiệu lực"
        }

    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "note": "Có thể cần cài đặt pyautogui hoặc cửa sổ YouTube chưa active"
        }


async def youtube_play_pause() -> dict:
    """Play/Pause YouTube video đang phát. Cần browser có YouTube đang focus."""
    return await control_youtube("play_pause")

async def youtube_rewind(seconds: int = 10) -> dict:
    """Tua lùi YouTube video. Mặc định 10 giây."""
    if seconds >= 10:
        return await control_youtube("rewind_10")
    else:
        return await control_youtube("rewind_5")

async def youtube_forward(seconds: int = 10) -> dict:
    """Tua tới YouTube video. Mặc định 10 giây."""
    if seconds >= 10:
        return await control_youtube("forward_10")
    else:
        return await control_youtube("forward_5")

async def youtube_volume_up() -> dict:
    """Tăng âm lượng YouTube 5%."""
    return await control_youtube("volume_up")

async def youtube_volume_down() -> dict:
    """Giảm âm lượng YouTube 5%."""
    return await control_youtube("volume_down")

async def youtube_mute() -> dict:
    """Bật/Tắt tiếng YouTube."""
    return await control_youtube("mute_toggle")

async def youtube_fullscreen() -> dict:
    """Bật/Tắt chế độ toàn màn hình YouTube (phím F)."""
    try:
        import pyautogui
        import time
        time.sleep(0.3)
        pyautogui.press('f')
        return {"success": True, "message": "✅ Đã bật/tắt fullscreen YouTube"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def youtube_captions() -> dict:
    """Bật/Tắt phụ đề YouTube (phím C)."""
    try:
        import pyautogui
        import time
        time.sleep(0.3)
        pyautogui.press('c')
        return {"success": True, "message": "✅ Đã bật/tắt phụ đề YouTube"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def youtube_speed(speed: str = "normal") -> dict:
    """
    Thay đổi tốc độ phát YouTube.
    speed: 'slower' (chậm hơn) hoặc 'faster' (nhanh hơn) hoặc 'normal' (bình thường)
    """
    try:
        import pyautogui
        import time
        time.sleep(0.3)
        if speed == "slower":
            pyautogui.hotkey('shift', ',')  # Shift + < = chậm hơn
            return {"success": True, "message": "✅ Đã giảm tốc độ YouTube"}
        elif speed == "faster":
            pyautogui.hotkey('shift', '.')  # Shift + > = nhanh hơn
            return {"success": True, "message": "✅ Đã tăng tốc độ YouTube"}
        else:
            # Reset về tốc độ bình thường - không có phím tắt trực tiếp
            return {"success": True, "message": "⚠️ Để reset về tốc độ bình thường, nhấn nhiều lần Shift+< hoặc dùng menu Settings"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ============================================================
# VLC PLAYER CONTROL TOOLS
# ============================================================

async def control_vlc(action: str) -> dict:
    """
    Điều khiển VLC Player bằng keyboard shortcuts.
    Cần VLC đang chạy và có focus.
    """
    try:
        import pyautogui
        import time
        
        shortcuts = {
            "play_pause": "space",      # Space - Play/Pause
            "stop": "s",                # S - Stop
            "next": "n",                # N - Next
            "previous": "p",            # P - Previous
            "volume_up": "ctrl+up",     # Ctrl+↑ - Tăng âm lượng
            "volume_down": "ctrl+down", # Ctrl+↓ - Giảm âm lượng
            "mute": "m",                # M - Mute
            "fullscreen": "f",          # F - Fullscreen
            "forward_short": "shift+right",  # Shift+→ - Tua tới 3 giây
            "backward_short": "shift+left",  # Shift+← - Tua lùi 3 giây
            "forward_medium": "alt+right",   # Alt+→ - Tua tới 10 giây
            "backward_medium": "alt+left",   # Alt+← - Tua lùi 10 giây
            "forward_long": "ctrl+right",    # Ctrl+→ - Tua tới 1 phút
            "backward_long": "ctrl+left",    # Ctrl+← - Tua lùi 1 phút
            "faster": "]",              # ] - Nhanh hơn
            "slower": "[",              # [ - Chậm hơn
            "normal_speed": "=",        # = - Tốc độ bình thường
            "loop": "l",                # L - Loop
            "random": "r",              # R - Random/Shuffle
        }
        
        if action not in shortcuts:
            return {
                "success": False,
                "error": f"Action không hợp lệ: {action}",
                "available_actions": list(shortcuts.keys())
            }
        
        time.sleep(0.3)
        key = shortcuts[action]
        
        if "+" in key:
            parts = key.split("+")
            pyautogui.hotkey(*parts)
        else:
            pyautogui.press(key)
        
        descriptions = {
            "play_pause": "Play/Pause",
            "stop": "Dừng phát",
            "next": "Bài tiếp theo",
            "previous": "Bài trước",
            "volume_up": "Tăng âm lượng",
            "volume_down": "Giảm âm lượng",
            "mute": "Bật/Tắt tiếng",
            "fullscreen": "Toàn màn hình",
            "forward_short": "Tua tới 3 giây",
            "backward_short": "Tua lùi 3 giây",
            "forward_medium": "Tua tới 10 giây",
            "backward_medium": "Tua lùi 10 giây",
            "forward_long": "Tua tới 1 phút",
            "backward_long": "Tua lùi 1 phút",
            "faster": "Tăng tốc độ phát",
            "slower": "Giảm tốc độ phát",
            "normal_speed": "Tốc độ bình thường",
            "loop": "Lặp lại",
            "random": "Phát ngẫu nhiên",
        }
        
        return {
            "success": True,
            "message": f"✅ VLC: {descriptions.get(action, action)}",
            "action": action
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


async def vlc_play_pause() -> dict:
    """Play/Pause VLC."""
    return await control_vlc("play_pause")

async def vlc_stop() -> dict:
    """Dừng phát VLC."""
    return await control_vlc("stop")

async def vlc_next() -> dict:
    """Chuyển bài tiếp theo trong VLC."""
    return await control_vlc("next")

async def vlc_previous() -> dict:
    """Quay lại bài trước trong VLC."""
    return await control_vlc("previous")

async def vlc_volume_up() -> dict:
    """Tăng âm lượng VLC."""
    return await control_vlc("volume_up")

async def vlc_volume_down() -> dict:
    """Giảm âm lượng VLC."""
    return await control_vlc("volume_down")

async def vlc_mute() -> dict:
    """Bật/Tắt tiếng VLC."""
    return await control_vlc("mute")

async def vlc_forward(seconds: int = 10) -> dict:
    """Tua tới trong VLC. 3s/10s/60s tùy theo seconds."""
    if seconds <= 5:
        return await control_vlc("forward_short")
    elif seconds <= 30:
        return await control_vlc("forward_medium")
    else:
        return await control_vlc("forward_long")

async def vlc_backward(seconds: int = 10) -> dict:
    """Tua lùi trong VLC. 3s/10s/60s tùy theo seconds."""
    if seconds <= 5:
        return await control_vlc("backward_short")
    elif seconds <= 30:
        return await control_vlc("backward_medium")
    else:
        return await control_vlc("backward_long")


# ============================================================
# WINDOWS MEDIA PLAYER CONTROL TOOLS
# ============================================================

async def control_wmp(action: str) -> dict:
    """
    Điều khiển Windows Media Player bằng keyboard shortcuts.
    Cần WMP đang chạy và có focus.
    """
    try:
        import pyautogui
        import time
        
        shortcuts = {
            "play_pause": "ctrl+p",     # Ctrl+P - Play/Pause
            "stop": "ctrl+s",           # Ctrl+S - Stop (có thể conflict với Save)
            "next": "ctrl+f",           # Ctrl+F - Next
            "previous": "ctrl+b",       # Ctrl+B - Previous
            "volume_up": "f10",         # F10 - Tăng âm lượng
            "volume_down": "f9",        # F9 - Giảm âm lượng
            "mute": "f8",               # F8 - Mute
            "fullscreen": "alt+enter",  # Alt+Enter - Fullscreen
            "forward": "ctrl+shift+f",  # Ctrl+Shift+F - Fast forward
            "backward": "ctrl+shift+b", # Ctrl+Shift+B - Rewind
        }
        
        if action not in shortcuts:
            return {
                "success": False,
                "error": f"Action không hợp lệ: {action}",
                "available_actions": list(shortcuts.keys())
            }
        
        time.sleep(0.3)
        key = shortcuts[action]
        
        if "+" in key:
            parts = key.split("+")
            pyautogui.hotkey(*parts)
        else:
            pyautogui.press(key)
        
        descriptions = {
            "play_pause": "Play/Pause",
            "stop": "Dừng phát",
            "next": "Bài tiếp theo",
            "previous": "Bài trước",
            "volume_up": "Tăng âm lượng",
            "volume_down": "Giảm âm lượng",
            "mute": "Bật/Tắt tiếng",
            "fullscreen": "Toàn màn hình",
            "forward": "Tua tới",
            "backward": "Tua lùi",
        }
        
        return {
            "success": True,
            "message": f"✅ Windows Media Player: {descriptions.get(action, action)}",
            "action": action
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


async def wmp_play_pause() -> dict:
    """Play/Pause Windows Media Player."""
    return await control_wmp("play_pause")

async def wmp_stop() -> dict:
    """Dừng phát Windows Media Player."""
    return await control_wmp("stop")

async def wmp_next() -> dict:
    """Chuyển bài tiếp theo trong Windows Media Player."""
    return await control_wmp("next")

async def wmp_previous() -> dict:
    """Quay lại bài trước trong Windows Media Player."""
    return await control_wmp("previous")

async def wmp_volume_up() -> dict:
    """Tăng âm lượng Windows Media Player."""
    return await control_wmp("volume_up")

async def wmp_volume_down() -> dict:
    """Giảm âm lượng Windows Media Player."""
    return await control_wmp("volume_down")

async def wmp_mute() -> dict:
    """Bật/Tắt tiếng Windows Media Player."""
    return await control_wmp("mute")


# ============================================================
# SMART MEDIA CONTROL - Tự động nhận diện player đang chạy
# ============================================================

async def smart_media_control(action: str) -> dict:
    """
    Điều khiển media thông minh.
    ⭐ ƯU TIÊN PYTHON-VLC TRƯỚC - nhanh nhất!
    Sau đó mới tới: Spotify > VLC Window > WMP > YouTube
    
    Actions: play_pause, stop, next, previous, volume_up, volume_down, mute
    """
    try:
        import time
        
        # 🎵 ƯU TIÊN 1: PYTHON-VLC NỘI BỘ - NHANH NHẤT!
        if vlc_player and vlc_player._player:
            action_map = {
                "play_pause": lambda: vlc_player.pause(),
                "stop": lambda: vlc_player.stop(),
                "next": lambda: (vlc_player._list_player.next(), time.sleep(0.3), vlc_player._list_player.play() if not vlc_player.is_playing() else None),
                "previous": lambda: (vlc_player._list_player.previous(), time.sleep(0.3), vlc_player._list_player.play() if not vlc_player.is_playing() else None),
                "volume_up": lambda: vlc_player._player.audio_set_volume(min(100, vlc_player._player.audio_get_volume() + 10)),
                "volume_down": lambda: vlc_player._player.audio_set_volume(max(0, vlc_player._player.audio_get_volume() - 10)),
                "mute": lambda: vlc_player._player.audio_toggle_mute()
            }
            
            if action in action_map:
                action_map[action]()
                status = vlc_player.get_full_status()
                return {
                    "success": True,
                    "message": f"✅ {action}: {status.get('current_song', 'VLC Player')}",
                    "player": "Python-VLC",
                    "current_song": status.get('current_song'),
                    "is_playing": vlc_player.is_playing(),
                    "llm_note": "🎵 Đang dùng Python-VLC. Tiếp tục dùng các lệnh nhạc VLC!"
                }
        
        # 2. Fallback: Dùng media keys cho external players
        import psutil
        import pyautogui
        
        running_players = []
        for proc in psutil.process_iter(['name']):
            name = proc.info['name'].lower()
            if 'spotify' in name:
                running_players.append('spotify')
            elif 'vlc' in name:
                running_players.append('vlc_external')
            elif 'wmplayer' in name:
                running_players.append('wmp')
            elif 'chrome' in name or 'firefox' in name or 'msedge' in name:
                running_players.append('browser')
        
        player = None
        if 'spotify' in running_players:
            player = 'spotify'
        elif 'vlc_external' in running_players:
            player = 'vlc_external'
        elif 'wmp' in running_players:
            player = 'wmp'
        elif 'browser' in running_players:
            player = 'browser'
        
        if not player:
            return {
                "success": False,
                "error": "Không có Python-VLC đang phát và không phát hiện media player nào",
                "hint": "Dùng play_music() để phát nhạc bằng Python-VLC trước!"
            }
        
        media_keys = {
            "play_pause": "playpause",
            "stop": "stop",
            "next": "nexttrack",
            "previous": "prevtrack",
            "volume_up": "volumeup",
            "volume_down": "volumedown",
            "mute": "volumemute"
        }
        
        if action in media_keys:
            time.sleep(0.2)
            pyautogui.press(media_keys[action])
            return {
                "success": True,
                "message": f"✅ Đã gửi lệnh {action} tới {player}",
                "player": player,
                "action": action
            }
        
        return {"success": False, "error": f"Action '{action}' không hợp lệ"}
        
    except Exception as e:
        return {"success": False, "error": str(e)}


# ============================================================
# NEW TOOLS FROM XIAOZHI-MCPTOOLS REFERENCE
# ============================================================

async def lock_computer() -> dict:
    """Khóa máy tính ngay lập tức"""
    try:
        subprocess.run("rundll32.exe user32.dll,LockWorkStation", shell=True, check=True)
        return {"success": True, "message": "Máy tính đã được khóa"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def shutdown_schedule(action: str, delay: int = 0) -> dict:
    """
    Lên lịch tắt máy/khởi động lại
    action: 'shutdown', 'restart', 'cancel'
    delay: thời gian trì hoãn (giây)
    """
    try:
        action_map = {"shutdown": "/s", "restart": "/r", "cancel": "/a"}
        if action not in action_map:
            return {"success": False, "error": f"Action không hợp lệ: {action}"}
        
        if action == "cancel":
            subprocess.run("shutdown /a", shell=True, check=True)
            return {"success": True, "message": "Đã hủy lịch tắt máy"}
        else:
            subprocess.run(f"shutdown {action_map[action]} /t {delay}", shell=True, check=True)
            return {"success": True, "message": f"Đã lên lịch {action} sau {delay} giây"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def show_desktop() -> dict:
    """Hiển thị desktop (Win+D)"""
    try:
        import pyautogui
        pyautogui.hotkey('win', 'd')
        return {"success": True, "message": "Đã hiển thị desktop"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def undo_operation() -> dict:
    """Hoàn tác thao tác cuối (Ctrl+Z)"""
    try:
        import pyautogui
        pyautogui.hotkey('ctrl', 'z')
        return {"success": True, "message": "Đã thực hiện hoàn tác"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def set_theme(dark_mode: bool = True) -> dict:
    """Đổi theme Windows sáng/tối. Nếu dark_mode=None thì toggle"""
    try:
        import winreg
        key_path = r"SOFTWARE\Microsoft\Windows\CurrentVersion\Themes\Personalize"
        
        # Nếu dark_mode là None, toggle mode hiện tại
        if dark_mode is None:
            with winreg.OpenKey(winreg.HKEY_CURRENT_USER, key_path, 0, winreg.KEY_READ) as key:
                current_value = winreg.QueryValueEx(key, "AppsUseLightTheme")[0]
                dark_mode = (current_value == 1)  # Nếu đang sáng (1) thì chuyển sang tối (True)
        
        value = 0 if dark_mode else 1
        
        with winreg.OpenKey(winreg.HKEY_CURRENT_USER, key_path, 0, winreg.KEY_SET_VALUE) as key:
            winreg.SetValueEx(key, "AppsUseLightTheme", 0, winreg.REG_DWORD, value)
            winreg.SetValueEx(key, "SystemUsesLightTheme", 0, winreg.REG_DWORD, value)
        
        mode = "tối" if dark_mode else "sáng"
        return {"success": True, "message": f"Đã chuyển sang theme {mode}"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def change_wallpaper(keyword: str = "", custom_path: str = "") -> dict:
    """
    Đổi hình nền desktop
    - Nếu có custom_path: dùng file được chỉ định
    - Nếu không: chọn ngẫu nhiên từ hình Windows có sẵn
    """
    try:
        import ctypes, os, random
        
        # Nếu có đường dẫn custom
        if custom_path:
            if not os.path.exists(custom_path):
                return {"success": False, "error": f"File không tồn tại: {custom_path}"}
            ctypes.windll.user32.SystemParametersInfoW(0x0014, 0, custom_path, 0x01 | 0x02)
            return {"success": True, "message": f"Đã đặt hình nền: {custom_path}"}
        
        # Chọn ngẫu nhiên từ Windows wallpapers
        wallpaper_paths = [
            r"C:\Windows\Web\Wallpaper\Windows\img0.jpg",
            r"C:\Windows\Web\Wallpaper\Windows\img19.jpg",
            r"C:\Windows\Web\Wallpaper\Spotlight\img14.jpg",
            r"C:\Windows\Web\Wallpaper\Spotlight\img50.jpg",
            r"C:\Windows\Web\Wallpaper\ThemeA\img20.jpg",
            r"C:\Windows\Web\Wallpaper\ThemeA\img21.jpg",
            r"C:\Windows\Web\Wallpaper\ThemeB\img24.jpg",
            r"C:\Windows\Web\Wallpaper\ThemeB\img25.jpg",
            r"C:\Windows\Web\Wallpaper\ThemeC\img28.jpg",
            r"C:\Windows\Web\Wallpaper\ThemeC\img29.jpg",
            r"C:\Windows\Web\Wallpaper\ThemeD\img32.jpg",
            r"C:\Windows\Web\Wallpaper\ThemeD\img33.jpg",
        ]
        available = [p for p in wallpaper_paths if os.path.exists(p)]
        if not available:
            return {"success": False, "error": "Không tìm thấy hình nền Windows"}
        selected = random.choice(available)
        ctypes.windll.user32.SystemParametersInfoW(0x0014, 0, selected, 0x01 | 0x02)
        return {"success": True, "message": f"Đã đổi hình nền: {os.path.basename(selected)}"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_desktop_path() -> dict:
    """Lấy đường dẫn thư mục Desktop"""
    try:
        user_profile = subprocess.check_output("echo %USERPROFILE%", shell=True, text=True).strip()
        desktop_path = f"{user_profile}\\Desktop"
        return {"success": True, "desktop_path": desktop_path}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def paste_content(content: str = "") -> dict:
    """
    Dán nội dung vào vị trí con trỏ
    Nếu content rỗng, chỉ thực hiện Ctrl+V với clipboard hiện tại
    """
    try:
        import pyperclip
        import pyautogui
        import time
        
        if content:
            # Nếu có content, copy vào clipboard trước
            pyperclip.copy(content)
            time.sleep(0.3)
        
        # Thực hiện paste
        pyautogui.hotkey('ctrl', 'v')
        
        msg = f"Đã dán: {content[:50]}..." if content else "Đã thực hiện paste"
        return {"success": True, "message": msg}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def press_enter() -> dict:
    """Nhấn phím Enter"""
    try:
        import pyautogui
        pyautogui.press('enter')
        return {"success": True, "message": "Đã nhấn Enter"}
    except Exception as e:
        return {"success": False, "error": str(e)}


async def save_text_to_file(content: str, filename: str = "") -> dict:
    """
    Lưu văn bản do LLM soạn thành file text
    LLM có thể soạn bài viết, báo cáo, code, v.v. và lưu trực tiếp vào file
    """
    try:
        import os
        from datetime import datetime
        
        # Nếu không có filename, tự động tạo tên với timestamp
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"llm_document_{timestamp}.txt"
        
        # Thêm .txt nếu chưa có extension
        if not filename.endswith(('.txt', '.md', '.json', '.csv', '.py', '.js', '.html', '.css')):
            filename += '.txt'
        
        # Lưu vào thư mục Documents của user
        documents_path = os.path.expanduser("~\\Documents")
        save_folder = os.path.join(documents_path, "miniZ_LLM_Documents")
        
        # Tạo thư mục nếu chưa có
        os.makedirs(save_folder, exist_ok=True)
        
        # Đường dẫn file đầy đủ
        file_path = os.path.join(save_folder, filename)
        
        # Lưu nội dung
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        file_size = os.path.getsize(file_path)
        
        return {
            "success": True, 
            "message": f"📄 Đã lưu file: {filename}",
            "path": file_path,
            "size_bytes": file_size,
            "location": save_folder
        }
        
    except Exception as e:
        return {"success": False, "error": f"Không thể lưu file: {str(e)}"}


async def gemini_text_to_speech(text: str, voice: str = "Aoede", save_audio: bool = False, filename: str = "") -> dict:
    """
    🎙️ Gemini TTS: Text-to-Speech sử dụng Gemini 2.5 Flash Preview TTS
    - Chất lượng cao, hỗ trợ tiếng Việt
    - 5 giọng nói: Puck (male), Charon (male), Kore (female), Fenrir (male), Aoede (female)
    
    Args:
        text: Văn bản cần đọc
        voice: Giọng nói (Aoede, Puck, Charon, Kore, Fenrir)
        save_audio: Có lưu file audio không
        filename: Tên file (nếu save_audio=True)
    """
    try:
        from google import genai
        from google.genai import types
        import os
        import tempfile
        from datetime import datetime
        import asyncio
        from concurrent.futures import ThreadPoolExecutor
        
        # Get API key
        gemini_api_key = os.environ.get("GEMINI_API_KEY") or GEMINI_API_KEY
        if not gemini_api_key:
            return {"success": False, "error": "Thiếu Gemini API key"}
        
        # Validate voice
        valid_voices = ["Puck", "Charon", "Kore", "Fenrir", "Aoede"]
        if voice not in valid_voices:
            voice = "Aoede"  # Default to female voice
        
        print(f"🎙️ [Gemini TTS] Text: {text[:50]}... Voice: {voice}")
        
        # Create client
        client = genai.Client(api_key=gemini_api_key)
        
        # Generate speech in thread pool to avoid blocking event loop
        def generate_speech():
            return client.models.generate_content(
                model="gemini-2.5-flash-preview-tts",
                contents=text,
                config=types.GenerateContentConfig(
                    response_modalities=["AUDIO"],
                    speech_config=types.SpeechConfig(
                        voice_config=types.VoiceConfig(
                            prebuilt_voice_config=types.PrebuiltVoiceConfig(
                                voice_name=voice
                            )
                        )
                    )
                )
            )
        
        # Run in thread pool with timeout
        loop = asyncio.get_event_loop()
        print(f"🎙️ [Gemini TTS] Calling API...")
        try:
            response = await asyncio.wait_for(
                loop.run_in_executor(None, generate_speech),
                timeout=30.0  # 30s timeout - đủ cho 500 chars
            )
            print(f"🎙️ [Gemini TTS] API responded!")
        except asyncio.TimeoutError:
            print(f"❌ [Gemini TTS] API timeout after 30s")
            return {"success": False, "error": "Gemini TTS timeout"}
        except Exception as api_err:
            print(f"❌ [Gemini TTS] API error: {api_err}")
            return {"success": False, "error": f"Gemini TTS API error: {str(api_err)}"}
        
        # Extract audio
        if response.candidates:
            for candidate in response.candidates:
                if hasattr(candidate, 'content') and candidate.content:
                    for part in candidate.content.parts:
                        if hasattr(part, 'inline_data') and part.inline_data:
                            audio_data = part.inline_data.data
                            mime_type = part.inline_data.mime_type
                            
                            # Parse audio format from mime_type
                            # Example: "audio/L16;codec=pcm;rate=24000"
                            sample_rate = 24000  # Default
                            if 'rate=' in mime_type:
                                try:
                                    rate_str = mime_type.split('rate=')[1].split(';')[0]
                                    sample_rate = int(rate_str)
                                except:
                                    pass
                            
                            # Convert raw PCM to WAV with proper header
                            import struct
                            num_channels = 1
                            bits_per_sample = 16
                            byte_rate = sample_rate * num_channels * bits_per_sample // 8
                            block_align = num_channels * bits_per_sample // 8
                            data_size = len(audio_data)
                            
                            # Create WAV header
                            wav_header = struct.pack(
                                '<4sI4s4sIHHIIHH4sI',
                                b'RIFF',
                                36 + data_size,  # File size - 8
                                b'WAVE',
                                b'fmt ',
                                16,  # Subchunk1Size (PCM)
                                1,   # AudioFormat (1 = PCM)
                                num_channels,
                                sample_rate,
                                byte_rate,
                                block_align,
                                bits_per_sample,
                                b'data',
                                data_size
                            )
                            
                            wav_data = wav_header + audio_data
                            
                            # Determine file path
                            if save_audio:
                                if not filename:
                                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                    filename = f"gemini_tts_{voice}_{timestamp}.wav"
                                
                                documents_path = os.path.expanduser("~\\Documents")
                                save_folder = os.path.join(documents_path, "miniZ_TTS_Audio")
                                os.makedirs(save_folder, exist_ok=True)
                                file_path = os.path.join(save_folder, filename)
                            else:
                                file_path = os.path.join(tempfile.gettempdir(), f"gemini_tts_{voice}.wav")
                            
                            # Save WAV file with proper header
                            with open(file_path, 'wb') as f:
                                f.write(wav_data)
                            
                            file_size = os.path.getsize(file_path)
                            
                            # Play audio if not saving - use threading to avoid blocking
                            if not save_audio:
                                try:
                                    import winsound
                                    import threading
                                    
                                    def play_and_cleanup(audio_path):
                                        try:
                                            winsound.PlaySound(audio_path, winsound.SND_FILENAME)
                                            # Clean up temp file after playing
                                            try:
                                                os.remove(audio_path)
                                            except:
                                                pass
                                        except Exception as e:
                                            print(f"⚠️ [Gemini TTS] Playback thread error: {e}")
                                    
                                    # Start playback in background thread
                                    play_thread = threading.Thread(target=play_and_cleanup, args=(file_path,), daemon=True)
                                    play_thread.start()
                                    print(f"🔊 [Gemini TTS] Started playback in background thread")
                                except Exception as e:
                                    print(f"⚠️ [Gemini TTS] Playback error: {e}")
                            
                            return {
                                "success": True,
                                "message": f"🔊 Đã đọc văn bản bằng Gemini TTS (Voice: {voice})",
                                "text_length": len(text),
                                "audio_size": len(audio_data),
                                "voice": voice,
                                "engine": "Gemini 2.5 Flash TTS",
                                "path": file_path if save_audio else None
                            }
        
        return {"success": False, "error": "Không nhận được audio từ Gemini"}
        
    except ImportError:
        return {"success": False, "error": "Thiếu google-genai package. Cài: pip install google-genai"}
    except Exception as e:
        return {"success": False, "error": f"Gemini TTS lỗi: {str(e)}"}


async def text_to_speech(text: str, save_audio: bool = False, filename: str = "") -> dict:
    """
    Text-to-Speech (TTS): Đọc văn bản thành giọng nói
    - Tự động dùng gTTS cho tiếng Việt (giọng native Google)
    - Dùng Windows SAPI cho các ngôn ngữ khác
    """
    try:
        import os
        import re
        from datetime import datetime
        
        # Kiểm tra xem văn bản có phải tiếng Việt không
        # Detect Vietnamese characters (ă, â, ê, ô, ơ, ư, đ với dấu)
        vietnamese_pattern = r'[àáảãạăắằẳẵặâấầẩẫậèéẻẽẹêếềểễệìíỉĩịòóỏõọôốồổỗộơớờởỡợùúủũụưứừửữựỳýỷỹỵđ]'
        is_vietnamese = bool(re.search(vietnamese_pattern, text.lower()))
        
        # === TIẾNG VIỆT: Dùng gTTS (Google Text-to-Speech) ===
        if is_vietnamese:
            try:
                from gtts import gTTS
                import pygame
                
                # Tạo tên file tạm
                if not filename:
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"tts_vietnamese_{timestamp}.mp3"
                
                if not filename.endswith('.mp3'):
                    filename += '.mp3'
                
                # Lưu vào Documents
                documents_path = os.path.expanduser("~\\Documents")
                save_folder = os.path.join(documents_path, "miniZ_TTS_Audio")
                os.makedirs(save_folder, exist_ok=True)
                
                file_path = os.path.join(save_folder, filename)
                
                # Tạo audio bằng gTTS (giọng Vietnamese native)
                tts = gTTS(text=text, lang='vi', slow=False)
                tts.save(file_path)
                
                file_size = os.path.getsize(file_path)
                
                # Nếu không lưu, phát audio rồi xóa file
                if not save_audio:
                    pygame.mixer.init()
                    pygame.mixer.music.load(file_path)
                    pygame.mixer.music.play()
                    
                    # Đợi audio phát xong
                    while pygame.mixer.music.get_busy():
                        await asyncio.sleep(0.1)
                    
                    pygame.mixer.quit()
                    
                    # Xóa file tạm
                    try:
                        os.remove(file_path)
                    except:
                        pass
                    
                    return {
                        "success": True,
                        "message": f"🔊 Đã đọc văn bản tiếng Việt (gTTS) ({len(text)} ký tự)",
                        "text_length": len(text),
                        "engine": "gTTS (Vietnamese native)"
                    }
                else:
                    return {
                        "success": True,
                        "message": f"🔊 Đã đọc và lưu audio tiếng Việt: {filename}",
                        "path": file_path,
                        "size_bytes": file_size,
                        "text_length": len(text),
                        "engine": "gTTS (Vietnamese native)"
                    }
            
            except ImportError:
                # Fallback to Windows SAPI if gTTS not installed
                print("⚠️ gTTS chưa cài. Dùng Windows SAPI (giọng English). Cài gTTS: pip install gTTS pygame")
                is_vietnamese = False  # Force fallback
            except Exception as e:
                print(f"⚠️ gTTS lỗi: {e}. Fallback to Windows SAPI")
                is_vietnamese = False  # Force fallback
        
        # === NGÔN NGỮ KHÁC: Dùng Windows SAPI ===
        if not is_vietnamese:
            import win32com.client
            
            # Khởi tạo SAPI voice
            speaker = win32com.client.Dispatch("SAPI.SpVoice")
            
            # Lấy danh sách voices (tiếng Anh, tiếng Việt nếu có cài)
            voices = speaker.GetVoices()
            
            # Nếu muốn lưu thành file audio
            if save_audio:
                from comtypes.client import CreateObject
                from comtypes.gen import SpeechLib
                
                engine = CreateObject("SAPI.SpVoice")
                stream = CreateObject("SAPI.SpFileStream")
                
                # Tạo tên file nếu không có
                if not filename:
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"tts_audio_{timestamp}.wav"
                
                if not filename.endswith('.wav'):
                    filename += '.wav'
                
                # Lưu vào Documents
                documents_path = os.path.expanduser("~\\Documents")
                save_folder = os.path.join(documents_path, "miniZ_TTS_Audio")
                os.makedirs(save_folder, exist_ok=True)
                
                file_path = os.path.join(save_folder, filename)
                
                # Mở stream và ghi audio
                stream.Open(file_path, SpeechLib.SSFMCreateForWrite)
                engine.AudioOutputStream = stream
                engine.Speak(text)
                stream.Close()
                
                file_size = os.path.getsize(file_path)
                
                return {
                    "success": True,
                    "message": f"🔊 Đã đọc văn bản và lưu audio: {filename}",
                    "path": file_path,
                    "size_bytes": file_size,
                    "text_length": len(text),
                    "engine": "Windows SAPI"
                }
            else:
                # Chỉ đọc không lưu
                speaker.Speak(text)
                
                return {
                    "success": True,
                    "message": f"🔊 Đã đọc văn bản ({len(text)} ký tự)",
                    "text_length": len(text),
                    "engine": "Windows SAPI"
                }
        
    except ImportError as e:
        return {
            "success": False, 
            "error": f"Thiếu module: {str(e)}. Cài: pip install pywin32 gTTS pygame"
        }
    except Exception as e:
        return {"success": False, "error": f"TTS lỗi: {str(e)}"}


async def speech_to_text(duration: int = 5, save_transcript: bool = True, filename: str = "") -> dict:
    """
    Speech-to-Text (STT): Chuyển giọng nói thành văn bản
    Sử dụng Google Speech Recognition (cần Internet)
    """
    try:
        import speech_recognition as sr
        import os
        from datetime import datetime
        
        # Khởi tạo recognizer
        recognizer = sr.Recognizer()
        
        # Sử dụng microphone
        with sr.Microphone() as source:
            print(f"🎤 Đang lắng nghe ({duration} giây)...")
            
            # Điều chỉnh nhiễu môi trường
            recognizer.adjust_for_ambient_noise(source, duration=1)
            
            # Ghi âm
            audio = recognizer.listen(source, timeout=duration, phrase_time_limit=duration)
            
            print("⏳ Đang nhận dạng giọng nói...")
            
            # Nhận dạng (Google Speech Recognition - miễn phí)
            try:
                # Thử tiếng Việt trước
                text_vi = recognizer.recognize_google(audio, language='vi-VN')
                text = text_vi
                language = "Tiếng Việt"
            except:
                try:
                    # Fallback sang tiếng Anh
                    text_en = recognizer.recognize_google(audio, language='en-US')
                    text = text_en
                    language = "English"
                except:
                    return {
                        "success": False,
                        "error": "Không nhận dạng được giọng nói. Hãy nói rõ hơn hoặc kiểm tra microphone."
                    }
        
        # Lưu transcript nếu cần
        if save_transcript and text:
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"stt_transcript_{timestamp}.txt"
            
            if not filename.endswith('.txt'):
                filename += '.txt'
            
            documents_path = os.path.expanduser("~\\Documents")
            save_folder = os.path.join(documents_path, "miniZ_STT_Transcripts")
            os.makedirs(save_folder, exist_ok=True)
            
            file_path = os.path.join(save_folder, filename)
            
            # Lưu kèm metadata
            content = f"=== Speech-to-Text Transcript ===\n"
            content += f"Ngày: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            content += f"Ngôn ngữ: {language}\n"
            content += f"Độ dài: {duration} giây\n"
            content += f"===================================\n\n"
            content += text
            
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return {
                "success": True,
                "message": f"🎤 Đã nhận dạng và lưu: {filename}",
                "text": text,
                "language": language,
                "path": file_path,
                "duration": duration
            }
        else:
            return {
                "success": True,
                "message": f"🎤 Đã nhận dạng giọng nói ({language})",
                "text": text,
                "language": language,
                "duration": duration
            }
        
    except ImportError:
        return {
            "success": False,
            "error": "Thiếu module SpeechRecognition. Cài: pip install SpeechRecognition pyaudio"
        }
    except Exception as e:
        return {"success": False, "error": f"STT lỗi: {str(e)}"}


# CÁC HÀM TRÙNG LẶP ĐÃ ĐƯỢC XÓA - SỬ DỤNG PHIÊN BẢN GỐC Ở TRÊN
# minimize_all_windows -> sử dụng show_desktop
# undo_action -> sử dụng undo_operation  
# toggle_dark_mode -> sử dụng set_theme
# set_wallpaper -> đã tích hợp vào change_wallpaper
# paste_text -> sử dụng paste_content
# find_on_screen -> sử dụng find_in_document
# shutdown_computer -> sử dụng shutdown_schedule


async def find_in_document(search_text: str) -> dict:
    """Tìm kiếm trong tài liệu (Ctrl+F)"""
    try:
        import pyperclip
        import pyautogui
        import time
        
        pyautogui.press('esc')
        time.sleep(0.3)
        pyautogui.hotkey('ctrl', 'f')
        time.sleep(0.1)
        
        pyperclip.copy(search_text)
        time.sleep(0.3)
        pyautogui.hotkey('ctrl', 'a')
        time.sleep(0.1)
        pyautogui.hotkey('ctrl', 'v')
        time.sleep(0.3)
        pyautogui.press('enter')
        
        return {"success": True, "message": f"Đã tìm kiếm: {search_text}"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ============================================================
# NEWS SCRAPING TOOLS
# ============================================================

async def get_vnexpress_news(category: str = "home", max_articles: int = 5) -> dict:
    """
    Lấy tin tức từ VnExpress RSS feeds (không cần feedparser)
    category: home, thoi-su, goc-nhin, the-gioi, kinh-doanh, giai-tri, the-thao, phap-luat, giao-duc, suc-khoe, gia-dinh, du-lich, khoa-hoc, so-hoa, xe, cong-dong, tam-su, cuoi
    """
    try:
        import aiohttp
        import xml.etree.ElementTree as ET
        
        # RSS URL mapping
        rss_urls = {
            "home": "https://vnexpress.net/rss/tin-moi-nhat.rss",
            "thoi-su": "https://vnexpress.net/rss/thoi-su.rss",
            "the-gioi": "https://vnexpress.net/rss/the-gioi.rss",
            "kinh-doanh": "https://vnexpress.net/rss/kinh-doanh.rss",
            "giai-tri": "https://vnexpress.net/rss/giai-tri.rss",
            "the-thao": "https://vnexpress.net/rss/the-thao.rss",
            "phap-luat": "https://vnexpress.net/rss/phap-luat.rss",
            "giao-duc": "https://vnexpress.net/rss/giao-duc.rss",
            "suc-khoe": "https://vnexpress.net/rss/suc-khoe.rss",
            "du-lich": "https://vnexpress.net/rss/du-lich.rss",
            "khoa-hoc": "https://vnexpress.net/rss/khoa-hoc.rss",
            "so-hoa": "https://vnexpress.net/rss/so-hoa.rss",
            "xe": "https://vnexpress.net/rss/oto-xe-may.rss",
        }
        
        rss_url = rss_urls.get(category, rss_urls["home"])
        
        print(f"📰 [News] Fetching news from: {rss_url}")
        
        # ⚡ Dùng aiohttp thay vì feedparser
        async with aiohttp.ClientSession() as session:
            async with session.get(rss_url, timeout=8) as resp:
                if resp.status != 200:
                    return {"success": False, "error": f"HTTP {resp.status}"}
                
                content = await resp.text()
                root = ET.fromstring(content)
                
                articles = []
                items = root.findall('.//item')[:max_articles]
                
                for i, item in enumerate(items):
                    try:
                        title_elem = item.find('title')
                        link_elem = item.find('link')
                        pubdate_elem = item.find('pubDate')
                        desc_elem = item.find('description')
                        
                        article = {
                            "title": title_elem.text if title_elem is not None else "No title",
                            "link": link_elem.text if link_elem is not None else "",
                            "published": pubdate_elem.text if pubdate_elem is not None else "",
                            "description": ""
                        }
                        
                        # Get description (strip HTML tags)
                        if desc_elem is not None and desc_elem.text:
                            import re
                            desc_text = re.sub(r'<[^>]+>', '', desc_elem.text)
                            article["description"] = desc_text.strip()[:200] + "..."
                        
                        articles.append(article)
                        print(f"✅ [News] Article {i+1}: {article['title'][:50]}...")
                        
                    except Exception as e:
                        print(f"⚠️ [News] Error parsing article {i+1}: {e}")
        
        result = {
            "success": True,
            "category": category,
            "total": len(articles),
            "articles": articles,
            "message": f"Đã lấy {len(articles)} tin tức từ VnExpress ({category})"
        }
        
        # 🤖 GEMINI SUMMARIZATION: Nếu >3 bài → tóm tắt thông minh
        if len(articles) > 3:
            try:
                context = "\n".join([
                    f"{i+1}. {a['title']}\n   {a['description']}"
                    for i, a in enumerate(articles)
                ])
                summary_prompt = f"""Tóm tắt {len(articles)} tin tức sau thành 5 bullet points QUAN TRỌNG NHẤT (tiếng Việt):

{context}

Yêu cầu:
- Mỗi bullet point ngắn gọn (1 dòng)
- Highlight xu hướng/sự kiện chính
- Ưu tiên tin có tác động lớn
"""
                gemini_summary = await ask_gemini(summary_prompt, model="models/gemini-3-flash-preview")
                
                if gemini_summary.get("success"):
                    result["gemini_summary"] = gemini_summary["response_text"]
                    result["message"] += " (✨ Đã tóm tắt bởi Gemini)"
                    print(f"✨ [News+Gemini] Summarized {len(articles)} articles")
            except Exception as e:
                print(f"⚠️ [News+Gemini] Summary failed: {e}")
        
        return result
        
    except Exception as e:
        return {"success": False, "error": f"Lỗi: {str(e)}"}


async def get_news_summary(category: str = "home") -> dict:
    """
    Lấy tóm tắt tin tức nhanh (chỉ tiêu đề)
    """
    try:
        result = await get_vnexpress_news(category=category, max_articles=10)
        
        if not result.get("success"):
            return result
        
        # Tạo summary text
        summary_lines = [f"📰 TIN TỨC {category.upper()} - VnExpress"]
        summary_lines.append("=" * 50)
        
        for i, article in enumerate(result["articles"], 1):
            summary_lines.append(f"{i}. {article['title']}")
        
        summary_text = "\n".join(summary_lines)
        
        # 🤖 GEMINI INTELLIGENT SUMMARY: Phân tích xu hướng + chọn top stories
        gemini_analysis = None
        if len(result["articles"]) >= 5:
            try:
                context = "\n".join([
                    f"{i+1}. {a['title']}"
                    for i, a in enumerate(result["articles"])
                ])
                analysis_prompt = f"""Phân tích {len(result['articles'])} tin tức sau và cho biết:
1. Top 3 tin QUAN TRỌNG NHẤT (kèm lý do)
2. Xu hướng chung
3. Chủ đề nổi bật

{context}

Format ngắn gọn, dễ đọc (tiếng Việt)."""
                
                gemini_result = await ask_gemini(analysis_prompt, model="models/gemini-3-flash-preview")
                if gemini_result.get("success"):
                    gemini_analysis = gemini_result["response_text"]
                    print(f"✨ [News+Gemini] Analyzed {len(result['articles'])} news items")
            except Exception as e:
                print(f"⚠️ [News+Gemini] Analysis failed: {e}")
        
        return {
            "success": True,
            "category": category,
            "total": len(result["articles"]),
            "summary": summary_text,
            "gemini_analysis": gemini_analysis,
            "articles": result["articles"],
            "message": f"Tóm tắt {len(result['articles'])} tin tức" + (" (✨ + Phân tích Gemini)" if gemini_analysis else "")
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}


async def search_news(keyword: str, max_results: int = 5) -> dict:
    """
    Tìm kiếm tin tức theo từ khóa trong các bài viết gần đây
    """
    try:
        # Get recent news from multiple categories
        categories = ["home", "thoi-su", "the-gioi", "kinh-doanh", "the-thao"]
        all_articles = []
        
        for cat in categories:
            result = await get_vnexpress_news(category=cat, max_articles=5)
            if result.get("success"):
                all_articles.extend(result["articles"])
        
        # Filter by keyword
        keyword_lower = keyword.lower()
        matched = []
        
        for article in all_articles:
            title_lower = article["title"].lower()
            desc_lower = article.get("description", "").lower()
            
            if keyword_lower in title_lower or keyword_lower in desc_lower:
                matched.append(article)
        
        matched = matched[:max_results]
        
        if not matched:
            return {
                "success": True,
                "keyword": keyword,
                "total": 0,
                "articles": [],
                "message": f"Không tìm thấy tin tức về '{keyword}'"
            }
        
        result = {
            "success": True,
            "keyword": keyword,
            "total": len(matched),
            "articles": matched,
            "message": f"Tìm thấy {len(matched)} tin tức về '{keyword}'"
        }
        
        # 🤖 GEMINI SUMMARIZATION: Nếu >3 kết quả → tóm tắt nhanh
        if len(matched) > 3:
            try:
                context = "\n".join([
                    f"{i+1}. {a['title'][:100]}"
                    for i, a in enumerate(matched[:5])
                ])
                # ⚡ PROMPT NGẮN
                summary_prompt = f"""Tóm tắt 3-4 ý chính về \"{keyword}\" từ {len(matched)} tin:
{context}

Format: 📌 [3-4 điểm chính]"""
                
                # ⏱️ Timeout 8s
                gemini_summary = await asyncio.wait_for(
                    ask_gemini_direct(summary_prompt, model="models/gemini-3-flash-preview"),
                    timeout=8.0
                )
                
                if gemini_summary.get("success"):
                    result["gemini_summary"] = gemini_summary["response_text"]
                    result["message"] += " (✨ Gemini)"
                    print(f"✅ [Search+Gemini] '{keyword}' done")
            except asyncio.TimeoutError:
                print(f"⏱️ [Search+Gemini] Timeout for '{keyword}'")
            except Exception as e:
                print(f"⚠️ [Search+Gemini] Error: {e}")
        
        return result
        
    except Exception as e:
        return {"success": False, "error": str(e)}


async def get_gold_price() -> dict:
    """
    Lấy giá vàng từ các nguồn uy tín
    """
    try:
        import requests
        from bs4 import BeautifulSoup
        import re

        # Try multiple sources
        sources = [
            {
                "name": "Sjc.com.vn",
                "url": "https://sjc.com.vn/xml/tygiavang.xml",
                "type": "xml"
            },
            {
                "name": "BNews.vn",
                "url": "https://bnews.vn/gia-vang/t32.html",
                "type": "html"
            }
        ]

        print(f"💰 [Gold] Fetching gold prices...")

        # Try SJC XML first
        try:
            response = requests.get(sources[0]["url"], timeout=10, headers={
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            })
            response.encoding = 'utf-8'

            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'xml')
                items = soup.find_all('item')

                if items:
                    gold_data = []

                    for item in items[:10]:
                        try:
                            gold_item = {
                                "type": item.get('@type', 'N/A'),
                                "buy": item.get('@buy', 'N/A'),
                                "sell": item.get('@sell', 'N/A')
                            }

                            # Fallback to text content if attributes not found
                            if gold_item["type"] == 'N/A':
                                type_tag = item.find('type')
                                buy_tag = item.find('buy')
                                sell_tag = item.find('sell')

                                if type_tag:
                                    gold_item["type"] = type_tag.get_text(strip=True)
                                if buy_tag:
                                    gold_item["buy"] = buy_tag.get_text(strip=True)
                                if sell_tag:
                                    gold_item["sell"] = sell_tag.get_text(strip=True)

                            gold_data.append(gold_item)
                            print(f"✅ [Gold] {gold_item['type']}: Mua {gold_item['buy']} | Bán {gold_item['sell']}")

                        except Exception as e:
                            print(f"⚠️ [Gold] Error parsing item: {e}")
                            continue

                    if gold_data:
                        # Loại bỏ trùng lặp
                        seen = set()
                        unique_gold_data = []
                        for item in gold_data:
                            key = f"{item['type']}_{item['buy']}_{item['sell']}"
                            if key not in seen:
                                seen.add(key)
                                unique_gold_data.append(item)
                        
                        gold_data = unique_gold_data[:10]
                        
                        # Tạo summary
                        summary_lines = ["💰 GIÁ VÀNG HÔM NAY - SJC", "=" * 60]

                        for item in gold_data:
                            summary_lines.append(f"📊 {item['type']}")
                            summary_lines.append(f"   Mua vào: {item['buy']} VNĐ | Bán ra: {item['sell']} VNĐ")
                            summary_lines.append("")

                        summary_text = "\n".join(summary_lines)
                        
                        # 🎙️ TTS-friendly description
                        tts_lines = ["Giá vàng SJC hôm nay như sau:"]
                        for item in gold_data[:5]:
                            tts_lines.append(f"Loại {item['type']}: giá mua {item['buy']} nghìn, giá bán {item['sell']} nghìn đồng.")
                        tts_description = " ".join(tts_lines)

                        return {
                            "success": True,
                            "total": len(gold_data),
                            "gold_prices": gold_data,
                            "summary": summary_text,
                            "tts_description": tts_description,
                            "message": f"Đã lấy giá {len(gold_data)} loại vàng",
                            "source": "SJC.com.vn",
                            "note_for_llm": "Khi đọc giá vàng, hãy dùng trường 'tts_description'. Giá tính theo nghìn đồng/lượng."
                        }

        except Exception as e:
            print(f"⚠️ [Gold] Error with SJC source: {e}")

        # Fallback: Try giavang.org scraping
        try:
            print(f"💰 [Gold] Trying giavang.org...")
            response = requests.get('https://giavang.org/', timeout=15, headers={
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            })

            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'html.parser')

                # Look for gold price tables
                tables = soup.find_all('table')
                gold_data = []

                for table in tables:
                    rows = table.find_all('tr')

                    for row in rows:
                        cols = row.find_all(['td', 'th'])
                        if len(cols) >= 3:
                            # Get text from columns
                            col_texts = [col.get_text(strip=True) for col in cols]

                            # Look for gold type and prices
                            if len(col_texts) >= 3:
                                gold_type = col_texts[0]
                                buy_price = col_texts[1]
                                sell_price = col_texts[2]

                                # Check if this looks like gold data
                                if ('vàng' in gold_type.lower() or 'sjc' in gold_type.lower() or 'nhẫn' in gold_type.lower() or 'pnj' in gold_type.lower() or 'doji' in gold_type.lower()) and buy_price and sell_price:
                                    # Clean prices
                                    buy_clean = re.sub(r'[^\d]', '', buy_price)
                                    sell_clean = re.sub(r'[^\d]', '', sell_price)

                                    if buy_clean and sell_clean:
                                        # Format with dots
                                        buy_formatted = f"{int(buy_clean):,}".replace(',', '.')
                                        sell_formatted = f"{int(sell_clean):,}".replace(',', '.')

                                        gold_data.append({
                                            "type": gold_type,
                                            "buy": buy_formatted,
                                            "sell": sell_formatted
                                        })
                                        print(f"✅ [Gold] {gold_type}: Mua {buy_formatted} | Bán {sell_formatted}")

                if gold_data:
                    # Loại bỏ trùng lặp dựa trên type + buy + sell
                    seen = set()
                    unique_gold_data = []
                    for item in gold_data:
                        key = f"{item['type']}_{item['buy']}_{item['sell']}"
                        if key not in seen:
                            seen.add(key)
                            unique_gold_data.append(item)
                    
                    gold_data = unique_gold_data[:10]  # Max 10 items
                    
                    # Tạo summary dễ đọc cho LLM/TTS
                    summary_lines = ["💰 GIÁ VÀNG HÔM NAY - GIAVANG.ORG", "=" * 60]

                    for item in gold_data:
                        summary_lines.append(f"📊 {item['type']}")
                        summary_lines.append(f"   Mua vào: {item['buy']} VNĐ | Bán ra: {item['sell']} VNĐ")
                        summary_lines.append("")

                    summary_text = "\n".join(summary_lines)
                    
                    # 🎙️ TTS-friendly description cho LLM đọc giá vàng
                    def format_price_speech(price_str):
                        """
                        Convert giá vàng sang tiếng Việt dễ đọc
                        - '180.100' = 180,100 nghìn = 180 triệu 100 nghìn VND
                        - Giá vàng hiển thị theo nghìn đồng/lượng
                        """
                        try:
                            # Remove dots/commas and convert to number
                            clean = price_str.replace('.', '').replace(',', '')
                            num = int(clean)
                            
                            # Giá vàng tính theo nghìn đồng/lượng
                            # Ví dụ: 180.100 = 180,100 (nghìn) = 180 triệu 100 nghìn VND
                            # num = 180100 → 180 triệu + 100 nghìn
                            
                            if num >= 1000:
                                # Giá >= 1000 nghìn = từ 1 triệu trở lên
                                millions = num // 1000  # 180100 // 1000 = 180
                                thousands = num % 1000  # 180100 % 1000 = 100
                                
                                if millions > 0 and thousands > 0:
                                    return f"{millions} triệu {thousands} nghìn"
                                elif millions > 0:
                                    return f"{millions} triệu"
                                else:
                                    return f"{thousands} nghìn"
                            else:
                                return f"{num} nghìn"
                        except:
                            return price_str
                    
                    # Tạo mô tả dạng câu cho TTS
                    tts_lines = ["Giá vàng hôm nay như sau:"]
                    for i, item in enumerate(gold_data[:5], 1):  # Top 5 cho TTS
                        gold_type = item['type'].replace('DOJI', 'Đô-ji').replace('PNJ', 'Pê-en-gi').replace('SJC', 'ét-gi-xi')
                        buy_speech = format_price_speech(item['buy'])
                        sell_speech = format_price_speech(item['sell'])
                        tts_lines.append(f"Loại {gold_type}: giá mua {buy_speech}, giá bán {sell_speech}.")
                    
                    tts_description = " ".join(tts_lines)

                    return {
                        "success": True,
                        "total": len(gold_data),
                        "gold_prices": gold_data,
                        "summary": summary_text,
                        "tts_description": tts_description,
                        "message": f"Đã lấy giá {len(gold_data)} loại vàng từ giavang.org",
                        "source": "giavang.org",
                        "note_for_llm": "Khi đọc giá vàng cho người dùng, hãy dùng trường 'tts_description' để đọc tự nhiên bằng tiếng Việt. Giá tính theo nghìn đồng/lượng."
                    }

        except Exception as e:
            print(f"⚠️ [Gold] Error with giavang.org: {e}")

        # Final fallback: Return sample data
        sample_data = [
            {"type": "Vàng SJC 1L, 10L", "buy": "88.500.000", "sell": "90.000.000"},
            {"type": "Vàng SJC 5c", "buy": "88.500.000", "sell": "90.200.000"},
            {"type": "Vàng nhẫn SJC 99.99 1c, 5c", "buy": "87.800.000", "sell": "89.300.000"},
            {"type": "Vàng nhẫn SJC 99.99 0.5c", "buy": "87.800.000", "sell": "89.400.000"},
        ]

        summary_lines = ["💰 GIÁ VÀNG THAM KHẢO", "=" * 60]
        for item in sample_data:
            summary_lines.append(f"📊 {item['type']}")
            summary_lines.append(f"   Mua vào: {item['buy']} VNĐ | Bán ra: {item['sell']} VNĐ")
            summary_lines.append("")

        return {
            "success": True,
            "total": len(sample_data),
            "gold_prices": sample_data,
            "summary": "\n".join(summary_lines),
            "message": f"Giá vàng tham khảo ({len(sample_data)} loại)",
            "source": "Sample Data",
            "note": "Giá tham khảo, không thể kết nối nguồn chính thống"
        }

    except Exception as e:
        return {"success": False, "error": f"Lỗi: {str(e)}"}


async def analyze_gold_price_with_ai(analysis_type: str = "compare_month") -> dict:
    """
    Phân tích giá vàng với AI (Gemini + Google Search).
    Lấy giá hiện tại, tìm dữ liệu lịch sử qua Google, và phân tích chi tiết.
    
    Args:
        analysis_type: Loại phân tích. Options: "compare_month" (so sánh với tháng trước), "trend" (xu hướng), "forecast" (dự đoán)
    """
    try:
        from datetime import datetime, timedelta
        
        print(f"🔍 [Gold AI] Starting gold price analysis: {analysis_type}")
        
        # 1. Lấy giá vàng hiện tại
        current_gold = await get_gold_price()
        if not current_gold.get("success"):
            return {"success": False, "error": "Không lấy được giá vàng hiện tại"}
        
        current_price_text = current_gold.get("summary", "")
        gold_prices = current_gold.get("gold_prices", [])
        
        # 2. Tìm giá vàng tháng trước qua Google Search (nếu có Serper API)
        historical_data = ""
        
        if SERPER_API_KEY and SERPER_API_KEY.strip():
            try:
                import requests
                
                # Tính tháng trước
                last_month_vn = (datetime.now() - timedelta(days=30)).strftime("tháng %m năm %Y")
                
                # Tìm giá vàng tháng trước
                search_query = f"giá vàng SJC cao nhất {last_month_vn}"
                
                url = "https://google.serper.dev/search"
                headers = {
                    "X-API-KEY": SERPER_API_KEY,
                    "Content-Type": "application/json"
                }
                payload = {
                    "q": search_query,
                    "gl": "vn",
                    "hl": "vi",
                    "num": 5
                }
                
                response = requests.post(url, headers=headers, json=payload, timeout=10)
                
                if response.status_code == 200:
                    data = response.json()
                    
                    # Lấy Answer Box
                    answer_box = data.get("answerBox", {})
                    if answer_box:
                        answer = answer_box.get("answer", "") or answer_box.get("snippet", "")
                        if answer:
                            historical_data += f"\n📌 DIRECT ANSWER: {answer}\n"
                    
                    # Lấy Organic Results
                    organic = data.get("organic", [])
                    historical_data += f"\n📊 KẾT QUẢ TÌM KIẾM '{search_query}':\n"
                    for i, item in enumerate(organic[:3], 1):
                        title = item.get("title", "")
                        snippet = item.get("snippet", "")
                        historical_data += f"\n{i}. {title}\n   {snippet}\n"
                    
                    print(f"✅ [Gold AI] Got historical data from Google")
                else:
                    print(f"⚠️ [Gold AI] Serper API returned {response.status_code}")
                    
            except Exception as e:
                print(f"⚠️ [Gold AI] Error fetching historical data: {e}")
                historical_data = "\n⚠️ Không thể lấy dữ liệu lịch sử từ Google\n"
        else:
            historical_data = "\n⚠️ Không có Serper API key để tìm dữ liệu lịch sử\n"
        
        # 3. Chuẩn bị prompt cho Gemini - CHI TIẾT VỪA ĐỦ
        if analysis_type == "compare_month":
            analysis_prompt = f"""Bạn là chuyên gia phân tích thị trường vàng. Hãy phân tích CHI TIẾT giá vàng:

📊 GIÁ HIỆN TẠI ({datetime.now().strftime("%d/%m/%Y")}):
{current_price_text}

📈 DỮ LIỆU LỊCH SỬ:
{historical_data}

YÊU CẦU PHÂN TÍCH (300-400 từ):
1. So sánh giá vàng hiện tại với tháng trước (% thay đổi cụ thể)
2. Đánh giá xu hướng: tăng/giảm/ổn định (phân tích kỹ lượng)
3. Phân tích nguyên nhân biến động (kinh tế, chính trị, USD, lạm phát, nguồn cung)
4. Dự báo ngắn hạn (1-2 tuần tới)
5. Khuyến nghị cụ thể cho nhà đầu tư (Mua/Bán/Chờ + lý do chi tiết)

Format output:
════════════════════════════
💰 PHÂN TÍCH GIÁ VÀNG
════════════════════════════

📊 SO SÁNH GIÁ:
[Giá hiện tại vs tháng trước, % thay đổi, biểu hiện thị trường]

📈 XU HƯỚNG:
[Nhận định chi tiết về xu hướng tăng/giảm, mức độ biến động]

🔍 NGUYÊN NHÂN:
[Phân tích 3-4 nguyên nhân chính với giải thích cụ thể]

🔮 DỰ BÁO:
[Dự đoán ngắn hạn và căn cứ]

💡 KHUYẾN NGHỊ:
[Lời khuyên cụ thể cho nhà đầu tư: Mua/Bán/Chờ + mức giá nên giao dịch]

════════════════════════════
"""
        elif analysis_type == "trend":
            analysis_prompt = f"""Phân tích xu hướng giá vàng (200-300 từ):

GIÁ HIỆN TẠI: {current_price_text}
DỮ LIỆU: {historical_data}

Trả lời ngắn:
📈 Xu hướng ngắn hạn: [1-2 tuần]
📊 Xu hướng trung hạn: [1-3 tháng] 
🔍 Yếu tố chính: [1-2 điểm]
"""
        else:  # forecast
            analysis_prompt = f"""Dự báo giá vàng (tối đa 100 từ):

HIỆN TẠI: {current_price_text}
LỊCH SỬ: {historical_data}

Trả lời ngắn:
📊 Dự báo: [tăng/giảm x%]
⏰ Thời gian: [ngắn/trung hạn]
💡 Khuyến nghị: [hành động cụ thể]
"""
        
        # 4. Gọi Gemini phân tích
        if not GEMINI_AVAILABLE or not GEMINI_API_KEY:
            return {
                "success": False,
                "error": "Gemini API không khả dụng. Vui lòng cấu hình GEMINI_API_KEY."
            }
        
        import google.generativeai as genai
        genai.configure(api_key=GEMINI_API_KEY)
        
        model = genai.GenerativeModel('models/gemini-3-flash-preview')
        
        print(f"🤖 [Gold AI] Asking Gemini to analyze...")
        response = model.generate_content(
            analysis_prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=0.7,
                max_output_tokens=1500  # Tăng lên 1500 để phân tích chuyên sâu
            )
        )
        
        analysis_result = response.text.strip()
        
        print(f"✅ [Gold AI] Analysis complete")
        
        # Return ONLY analysis text - tránh bị truncate
        return {
            "success": True,
            "content": analysis_result  # Chỉ trả về nội dung phân tích
        }
        
    except Exception as e:
        print(f"❌ [Gold AI] Error: {e}")
        return {
            "success": False,
            "error": f"Lỗi phân tích: {str(e)}"
        }


# ============================================================================
# 🔍 GEMINI WITH GOOGLE SEARCH GROUNDING
# ============================================================================
# Tính năng cho phép Gemini tự động tra cứu Google để trả lời chính xác hơn
# Sử dụng Google Search Grounding API chính thức

async def ask_gemini_with_google_search(
    prompt: str, 
    model: str = "gemini-2.0-flash",
    dynamic_threshold: float = 0.7
) -> dict:
    """
    🔍 Hỏi Gemini với Google Search Grounding - Tra cứu Google tự động
    
    Tính năng này cho phép Gemini:
    - Tự động tìm kiếm thông tin mới nhất trên Google
    - Trả lời dựa trên dữ liệu real-time từ internet
    - Cung cấp nguồn trích dẫn (citations)
    
    Args:
        prompt: Câu hỏi cần Gemini trả lời với thông tin mới nhất
        model: Model Gemini hỗ trợ grounding (gemini-2.0-flash, gemini-1.5-pro, etc.)
        dynamic_threshold: Ngưỡng để quyết định khi nào dùng grounding (0.0-1.0)
        
    Returns:
        dict với success, response, grounding_metadata, search_queries
    """
    try:
        if not GEMINI_AVAILABLE:
            return {"success": False, "error": "Gemini library chưa cài đặt"}
        
        if not GEMINI_API_KEY or not GEMINI_API_KEY.strip():
            return {"success": False, "error": "Gemini API key chưa được cấu hình"}
        
        print(f"🔍 [Gemini+GoogleSearch] Starting with model: {model}")
        print(f"🔍 [Gemini+GoogleSearch] Prompt: {prompt[:100]}...")
        
        # Import các module cần thiết từ google.genai
        try:
            from google import genai
            from google.genai import types
        except ImportError:
            # Fallback: Dùng google-generativeai cũ
            print("⚠️ [Gemini+GoogleSearch] google-genai not found, using legacy method")
            return await _ask_gemini_google_search_legacy(prompt, model)
        
        # Khởi tạo client với API key
        client = genai.Client(api_key=GEMINI_API_KEY)
        
        # Cấu hình Google Search tool với dynamic retrieval
        google_search_tool = types.Tool(
            google_search=types.GoogleSearch()
        )
        
        # System instruction để Gemini trả lời chuyên nghiệp
        from datetime import datetime
        today_str = datetime.now().strftime('%d/%m/%Y')
        today_full = datetime.now().strftime('%A, %d tháng %m năm %Y')
        
        system_instruction = f"""Bạn là trợ lý AI chuyên nghiệp với khả năng tra cứu thông tin mới nhất từ Google.

📅 NGÀY HÔM NAY: {today_full}

🎯 HƯỚNG DẪN TRẢ LỜI:
1. SỬ DỤNG GOOGLE SEARCH để tìm thông tin mới nhất, chính xác
2. ƯU TIÊN nguồn đáng tin cậy: trang chính thức, báo lớn, Wikipedia
3. PHÂN TÍCH thời gian - nếu thông tin từ quá khứ, xác định xem còn đúng không
4. TRẢ LỜI ngắn gọn, súc tích (200-500 từ)
5. TRÍCH DẪN nguồn khi cần thiết
6. KHÔNG nói "dự kiến" nếu sự kiện đã xảy ra
7. Nói như đang trò chuyện tự nhiên, không dùng markdown phức tạp"""

        # Gọi Gemini với Google Search grounding
        loop = asyncio.get_event_loop()
        
        response = await asyncio.wait_for(
            loop.run_in_executor(
                None,
                lambda: client.models.generate_content(
                    model=model,
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        tools=[google_search_tool],
                        system_instruction=system_instruction,
                        temperature=0.7,
                    )
                )
            ),
            timeout=30.0  # Timeout 30s vì cần thời gian search
        )
        
        # Lấy text response
        response_text = ""
        if hasattr(response, 'text'):
            response_text = response.text
        elif hasattr(response, 'candidates') and response.candidates:
            for part in response.candidates[0].content.parts:
                if hasattr(part, 'text'):
                    response_text += part.text
        
        # Lấy grounding metadata (nguồn trích dẫn)
        grounding_metadata = None
        search_queries = []
        grounding_chunks = []
        
        if hasattr(response, 'candidates') and response.candidates:
            candidate = response.candidates[0]
            if hasattr(candidate, 'grounding_metadata'):
                gm = candidate.grounding_metadata
                grounding_metadata = {
                    "search_entry_point": getattr(gm, 'search_entry_point', None),
                    "grounding_supports": []
                }
                
                # Lấy search queries đã dùng
                if hasattr(gm, 'web_search_queries'):
                    search_queries = list(gm.web_search_queries or [])
                
                # Lấy grounding chunks (nguồn)
                if hasattr(gm, 'grounding_chunks'):
                    for chunk in (gm.grounding_chunks or []):
                        if hasattr(chunk, 'web'):
                            grounding_chunks.append({
                                "uri": getattr(chunk.web, 'uri', ''),
                                "title": getattr(chunk.web, 'title', '')
                            })
                
                # Lấy grounding supports
                if hasattr(gm, 'grounding_supports'):
                    for support in (gm.grounding_supports or []):
                        support_data = {
                            "segment": getattr(support.segment, 'text', '') if hasattr(support, 'segment') else '',
                            "confidence_scores": list(support.confidence_scores or []) if hasattr(support, 'confidence_scores') else []
                        }
                        grounding_metadata["grounding_supports"].append(support_data)
        
        print(f"✅ [Gemini+GoogleSearch] Response received: {len(response_text)} chars")
        if search_queries:
            print(f"🔎 [Gemini+GoogleSearch] Search queries: {search_queries}")
        if grounding_chunks:
            print(f"📚 [Gemini+GoogleSearch] Sources: {len(grounding_chunks)} websites")
        
        # Truncate response nếu quá dài
        if len(response_text) > MAX_LLM_RESPONSE_CHARS:
            response_text = smart_truncate_for_llm(response_text, MAX_LLM_RESPONSE_CHARS)
        
        return {
            "success": True,
            "response": response_text,
            "response_text": response_text,  # Alias for compatibility
            "model": model,
            "google_search_used": True,
            "search_queries": search_queries,
            "grounding_chunks": grounding_chunks,
            "grounding_metadata": grounding_metadata,
            "message": f"✅ Gemini đã tra cứu Google và trả lời (model: {model})"
        }
        
    except asyncio.TimeoutError:
        print(f"⏱️ [Gemini+GoogleSearch] Timeout (30s exceeded)")
        return {
            "success": False,
            "error": "Gemini + Google Search phản hồi quá lâu (timeout 30s)",
            "timeout": True
        }
    except ImportError as e:
        print(f"⚠️ [Gemini+GoogleSearch] Import error: {e}")
        # Fallback to legacy method
        return await _ask_gemini_google_search_legacy(prompt, model)
    except Exception as e:
        error_msg = str(e)
        print(f"❌ [Gemini+GoogleSearch] Error: {error_msg}")
        
        # Nếu model không hỗ trợ grounding, thử fallback
        if "grounding" in error_msg.lower() or "tool" in error_msg.lower():
            print("⚠️ [Gemini+GoogleSearch] Grounding not supported, falling back...")
            return await ask_gemini(prompt, model)
        
        return {
            "success": False,
            "error": f"Lỗi Google Search Grounding: {error_msg}"
        }


async def _ask_gemini_google_search_legacy(prompt: str, model: str = "gemini-2.0-flash") -> dict:
    """
    Fallback: Dùng google-generativeai cũ với grounding
    """
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        
        from datetime import datetime
        today_str = datetime.now().strftime('%d/%m/%Y')
        
        # Cấu hình model với grounding
        generation_config = {
            "temperature": 0.7,
            "top_p": 0.95,
            "max_output_tokens": 2048,
        }
        
        # Tạo tool Google Search
        try:
            # Thử dùng google_search_retrieval (phiên bản mới)
            tools = [{"google_search_retrieval": {"dynamic_retrieval_config": {"mode": "MODE_DYNAMIC", "dynamic_threshold": 0.7}}}]
            gemini_model = genai.GenerativeModel(
                model,
                generation_config=generation_config,
                tools=tools
            )
        except Exception:
            # Fallback: không dùng tools
            gemini_model = genai.GenerativeModel(model, generation_config=generation_config)
        
        system_prompt = f"""Hôm nay là {today_str}. Bạn là trợ lý AI thông minh.
Hãy trả lời câu hỏi dựa trên kiến thức của bạn. Trả lời ngắn gọn, chuyên nghiệp."""
        
        full_prompt = f"{system_prompt}\n\nCâu hỏi: {prompt}"
        
        loop = asyncio.get_event_loop()
        response = await asyncio.wait_for(
            loop.run_in_executor(None, lambda: gemini_model.generate_content(full_prompt)),
            timeout=25.0
        )
        
        response_text = response.text if hasattr(response, 'text') else str(response)
        
        if len(response_text) > MAX_LLM_RESPONSE_CHARS:
            response_text = smart_truncate_for_llm(response_text, MAX_LLM_RESPONSE_CHARS)
        
        return {
            "success": True,
            "response": response_text,
            "response_text": response_text,
            "model": model,
            "google_search_used": False,
            "message": f"✅ Gemini đã trả lời (model: {model}, legacy mode)"
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}


async def ask_gemini_direct(prompt: str, model: str = "models/gemini-3-flash-preview") -> dict:
    """
    Gọi Gemini trực tiếp KHÔNG có RAG - dùng cho summarization/analysis
    
    Args:
        prompt: Prompt gửi cho Gemini
        model: Model Gemini (mặc định: gemini-3-flash-preview)
        
    Returns:
        dict với success, response_text
    """
    try:
        # Kiểm tra Gemini có khả dụng không
        if not GEMINI_AVAILABLE:
            return {"success": False, "error": "Gemini library chưa cài đặt"}
        
        if not GEMINI_API_KEY or GEMINI_API_KEY.strip() == "":
            return {"success": False, "error": "Gemini API key chưa được cấu hình"}
        
        # Cấu hình Gemini
        genai.configure(api_key=GEMINI_API_KEY)
        gemini_model = genai.GenerativeModel(model)
        
        # Gọi API với timeout 15 giây
        loop = asyncio.get_event_loop()
        response = await asyncio.wait_for(
            loop.run_in_executor(
                None, 
                lambda: gemini_model.generate_content(prompt)
            ),
            timeout=15.0
        )
        
        response_text = response.text
        
        # 🔄 TRUNCATE: Giới hạn response dưới 4000 ký tự cho LLM
        if len(response_text) > MAX_LLM_RESPONSE_CHARS:
            original_len = len(response_text)
            response_text = smart_truncate_for_llm(response_text, MAX_LLM_RESPONSE_CHARS)
            print(f"[Gemini Direct] ✂️ Truncated: {original_len} → {len(response_text)} chars")
        
        return {
            "success": True,
            "response_text": response_text,
            "model": model
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}


async def ask_gemini(prompt: str, model: str = "models/gemini-3-flash-preview") -> dict:
    """
    Hỏi đáp với Google Gemini AI - Có tích hợp RAG tự động
    
    Args:
        prompt: Câu hỏi hoặc nội dung muốn gửi cho Gemini
        model: Tên model Gemini (mặc định: models/gemini-3-flash-preview - Flash 2.0 experimental)
        
    Returns:
        dict với success, response_text, và message
    """
    try:
        # ===== AUTO RAG: Kiểm tra có cần tra cứu web không =====
        # Mở rộng keywords để bao quát nhiều câu hỏi thời sự hơn
        realtime_keywords = [
            # Giá cả, tài chính
            'giá vàng', 'giá usd', 'tỷ giá', 'giá bitcoin', 'crypto', 'chứng khoán', 
            'stock', 'gold price', 'exchange rate', 'giá xăng', 'giá dầu',
            'giá cao nhất', 'cao nhất', 'thấp nhất', 'giá hiện tại', 'giá mới nhất',
            'highest price', 'lowest price', 'current price', 'latest price',
            
            # Thời tiết
            'thời tiết', 'weather', 'nhiệt độ', 'temperature', 'mưa', 'rain',
            
            # Tin tức, sự kiện
            'tin tức', 'news', 'mới nhất', 'latest', 'breaking',
            
            # Thời gian thực
            'hôm nay', 'bây giờ', 'hiện nay', 'hiện tại', 'today', 'now', 'current',
            'currently', 'năm 2024', 'năm 2025', '2024', '2025',
            
            # Thể thao, cuộc thi
            'vô địch', 'champion', 'winner', 'kết quả', 'score', 'result',
            'olympia', 'world cup', 'euro', 'sea games', 'olympic', 'bóng đá', 'football',
            
            # Người nổi tiếng, chính trị
            'tổng thống', 'president', 'thủ tướng', 'prime minister', 'chủ tịch',
            'ceo', 'founder', 'leader', 'ai là', 'who is', 'who are',
            
            # Sản phẩm, công nghệ mới
            'iphone', 'samsung', 'tesla', 'apple', 'google', 'microsoft',
            'ra mắt', 'launch', 'release', 'announced',
            
            # Sự kiện xã hội
            'covid', 'earthquake', 'động đất', 'bão', 'storm', 'lũ lụt', 'flood',
            'tai nạn', 'accident', 'cháy', 'fire',
            
            # Tra cứu chung
            'là ai', 'là gì', 'ở đâu', 'what is', 'where is', 'how much',
            'bao nhiêu', 'khi nào', 'when'
        ]
        prompt_lower = prompt.lower()
        needs_realtime = any(kw in prompt_lower for kw in realtime_keywords)
        
        rag_context = ""
        if needs_realtime:
            # ✅ Ưu tiên Serper API (Google Search trực tiếp) - chính xác và nhanh hơn
            if SERPER_API_KEY and SERPER_API_KEY.strip():
                print(f"[Gemini+Serper] Phát hiện câu hỏi thời gian thực, đang tra cứu Google...")
                try:
                    import requests
                    from datetime import datetime
                    
                    # Thêm ngày tháng năm hiện tại vào query để lấy thông tin mới nhất
                    current_date = datetime.now().strftime("%Y")
                    enhanced_query = f"{prompt} {current_date}"
                    
                    # Gọi Serper API (Google Search)
                    url = "https://google.serper.dev/search"
                    headers = {
                        "X-API-KEY": SERPER_API_KEY,
                        "Content-Type": "application/json"
                    }
                    payload = {
                        "q": enhanced_query,
                        "gl": "vn",  # Vietnam
                        "hl": "vi",  # Vietnamese
                        "num": 5
                    }
                    
                    # ⚡ TIMEOUT 8s cho Serper API
                    response = requests.post(url, headers=headers, json=payload, timeout=8)
                    
                    if response.status_code == 200:
                        data = response.json()
                        results = []
                        
                        # Lấy Answer Box trước (nếu có)
                        answer_box = data.get("answerBox", {})
                        if answer_box:
                            answer = answer_box.get("answer", "") or answer_box.get("snippet", "")
                            if answer:
                                results.append({
                                    "title": "[📌 Direct Answer]",
                                    "snippet": answer,
                                    "url": answer_box.get("link", "")
                                })
                        
                        # Lấy Knowledge Graph (nếu có)
                        knowledge_graph = data.get("knowledgeGraph", {})
                        if knowledge_graph:
                            title = knowledge_graph.get("title", "")
                            description = knowledge_graph.get("description", "")
                            if title and description:
                                results.append({
                                    "title": f"[🎯 Knowledge] {title}",
                                    "snippet": description,
                                    "url": knowledge_graph.get("website", "")
                                })
                        
                        # Lấy Organic Results
                        organic = data.get("organic", [])
                        for item in organic[:5]:
                            results.append({
                                "title": item.get("title", ""),
                                "snippet": item.get("snippet", ""),
                                "url": item.get("link", "")
                            })
                        
                        if results:
                            rag_context = f"\n\n📊 THÔNG TIN TỪ GOOGLE (tra cứu {datetime.now().strftime('%d/%m/%Y')}):\n"
                            rag_context += "LƯU Ý: Hãy phân tích kỹ các nguồn và chọn thông tin chính xác nhất.\n\n"
                            
                            for i, r in enumerate(results, 1):
                                snippet = r['snippet'][:300] if len(r['snippet']) > 300 else r['snippet']
                                rag_context += f"{i}. **{r['title']}**\n   {snippet}\n   🔗 {r.get('url', '')}\n\n"
                            
                            print(f"[Gemini+Serper] ✅ Đã lấy được {len(results)} kết quả từ Google")
                    else:
                        print(f"[Gemini+Serper] ⚠️ API error: {response.status_code}")
                        
                except Exception as e:
                    print(f"[Gemini+Serper] ⚠️ Lỗi tra cứu: {e}")
            
            # Fallback: Dùng RAG system nếu không có Serper API
            elif RAG_AVAILABLE:
                print(f"[Gemini+RAG] Serper API không có, dùng RAG fallback...")
                try:
                    from rag_system import web_search
                    from datetime import datetime
                    
                    current_date = datetime.now().strftime("%Y")
                    enhanced_query = f"{prompt} {current_date}"
                    
                    rag_result = await web_search(enhanced_query, max_results=5)
                    
                    if rag_result.get('success') and rag_result.get('results'):
                        rag_context = f"\n\n📊 THÔNG TIN TỪ INTERNET (tra cứu {datetime.now().strftime('%d/%m/%Y')}):\n"
                        rag_context += "LƯU Ý: Hãy phân tích kỹ các nguồn và chọn thông tin chính xác nhất.\n\n"
                        
                        for i, r in enumerate(rag_result['results'], 1):
                            snippet = r['snippet'][:300] if len(r['snippet']) > 300 else r['snippet']
                            rag_context += f"{i}. **{r['title']}**\n   {snippet}\n   🔗 {r.get('url', '')}\n\n"
                        
                        print(f"[Gemini+RAG] ✅ Đã lấy được {len(rag_result['results'])} kết quả từ web")
                except Exception as e:
                    print(f"[Gemini+RAG] ⚠️ Lỗi tra cứu: {e}")
        
        # Kiểm tra Gemini có khả dụng không
        if not GEMINI_AVAILABLE:
            return {
                "success": False,
                "error": "Gemini library chưa cài đặt. Chạy: pip install google-generativeai"
            }
        
        # Kiểm tra API key
        if not GEMINI_API_KEY or GEMINI_API_KEY.strip() == "":
            return {
                "success": False,
                "error": "Gemini API key chưa được cấu hình. Vui lòng thêm 'gemini_api_key' vào xiaozhi_endpoints.json",
                "help": "Lấy API key tại: https://aistudio.google.com/apikey"
            }
        
        # Cấu hình Gemini với API key
        genai.configure(api_key=GEMINI_API_KEY)
        print(f"[Gemini] Configured with API key: ...{GEMINI_API_KEY[-8:]}")
        
        # Khởi tạo model
        print(f"[Gemini] Creating model: {model}")
        gemini_model = genai.GenerativeModel(model)
        print(f"[Gemini] Model created successfully")
        
        # Gọi API trong executor để không block event loop
        # Thêm RAG context vào prompt nếu có
        from datetime import datetime as dt_now
        enhanced_prompt = prompt
        
        # 📝 INSTRUCTION: Yêu cầu Gemini trả lời ngắn gọn cho TTS
        response_instruction = """

📋 YÊU CẦU TRẢ LỜI:
- Trả lời NGẮN GỌN, DỄ HIỂU (tối đa 300-500 từ)
- Đi thẳng vào vấn đề, không dài dòng
- KHÔNG dùng markdown (**, #, ---, bullet points)
- Nói như đang trò chuyện tự nhiên
- Dùng câu ngắn, dễ đọc"""

        if rag_context:
            today_str = dt_now.now().strftime('%d/%m/%Y')
            today_full = dt_now.now().strftime('%d tháng %m năm %Y')
            enhanced_prompt = f"""CÂU HỎI: {prompt}

{rag_context}

⚠️ QUAN TRỌNG - NGÀY HIỆN TẠI: {today_full}

HƯỚNG DẪN PHÂN TÍCH THÔNG MINH:
1. **SO SÁNH THỜI GIAN**: So sánh ngày trong bài báo với ngày hôm nay ({today_str})
   - Nếu bài viết có từ "dự kiến", "sắp ra mắt", "sẽ ra mắt" VÀ ngày đó ĐÃ QUA → sản phẩm ĐÃ RA MẮT rồi!
   - Ví dụ: Nếu bài viết nói "dự kiến ra mắt tháng 9/2025" và hôm nay là tháng 12/2025 → ĐÃ RA MẮT

2. **XÁC ĐỊNH TRẠNG THÁI HIỆN TẠI**:
   - Kiểm tra xem các nguồn có nói "đã ra mắt", "đã có hàng", "đặt trước từ..." không
   - Nếu có ngày đặt trước/ngày bán ĐÃ QUA → sản phẩm ĐANG BÁN
   - Nếu nguồn chính thức (apple.com, thegioididong.com) nói "sẵn hàng" → ĐÃ CÓ BÁN

3. **ƯU TIÊN NGUỒN**:
   - Trang chính thức (apple.com, google.com...) > Báo lớn > Blog
   - Nguồn mới nhất > Nguồn cũ

4. **TRẢ LỜI CHÍNH XÁC**:
   - KHÔNG nói "dự kiến" nếu ngày đó đã qua
   - Dùng thì HIỆN TẠI/QUÁ KHỨ phù hợp
   - Ví dụ ĐÚNG: "iPhone 17 đã ra mắt vào tháng 9/2025 và hiện đang bán tại..."
   - Ví dụ SAI: "iPhone 17 dự kiến ra mắt tháng 9/2025" (khi đã là tháng 12/2025!)
{response_instruction}

TRẢ LỜI (nhớ: hôm nay là {today_str}, phân tích thời gian chính xác):"""
            print(f"[Gemini+RAG] Đã bổ sung context từ web vào prompt")
        else:
            # Không có RAG, thêm instruction vào prompt thông thường
            enhanced_prompt = f"""{prompt}
{response_instruction}"""
        
        print(f"[Gemini] Sending prompt: {enhanced_prompt[:50]}...")
        loop = asyncio.get_event_loop()
        
        # ⚡ TIMEOUT 20s cho ask_gemini chính (có RAG)
        response = await asyncio.wait_for(
            loop.run_in_executor(
                None,
                lambda: gemini_model.generate_content(enhanced_prompt)
            ),
            timeout=20.0
        )
        print(f"[Gemini] Response received")
        
        # Lấy text từ response
        response_text = response.text if hasattr(response, 'text') else str(response)
        
        # 🔄 TRUNCATE: Giới hạn response dưới 4000 ký tự cho LLM
        if len(response_text) > MAX_LLM_RESPONSE_CHARS:
            original_len = len(response_text)
            response_text = smart_truncate_for_llm(response_text, MAX_LLM_RESPONSE_CHARS)
            print(f"[Gemini] ✂️ Truncated response: {original_len} → {len(response_text)} chars")
        
        print(f"[Gemini] Response text: {response_text[:100]}...")
        
        result = {
            "success": True,
            "prompt": prompt,
            "response_text": response_text,
            "model": model,
            "message": f"✅ Gemini đã trả lời (model: {model})"
        }
        
        # Thêm thông tin RAG nếu đã sử dụng
        if rag_context:
            result["rag_used"] = True
            result["message"] = f"✅ Gemini đã trả lời với thông tin từ Internet (model: {model})"
        
        return result
        
    except asyncio.TimeoutError:
        print(f"⏱️ [Gemini] Timeout (20s exceeded)")
        return {
            "success": False,
            "error": "Gemini phản hồi quá lâu (timeout 20s). Vui lòng thử lại với prompt ngắn hơn.",
            "timeout": True
        }
    except Exception as e:
        error_msg = str(e)
        print(f"❌ [Gemini] Exception caught: {type(e).__name__}")
        print(f"❌ [Gemini] Error message: {error_msg}")
        
        # Import traceback để debug
        import traceback
        traceback.print_exc()
        
        # Xử lý các lỗi phổ biến
        if "API_KEY_INVALID" in error_msg or "invalid API key" in error_msg.lower():
            return {
                "success": False,
                "error": "API key không hợp lệ. Vui lòng kiểm tra lại gemini_api_key trong xiaozhi_endpoints.json",
                "help": "Lấy API key mới tại: https://aistudio.google.com/apikey"
            }
        elif "quota" in error_msg.lower():
            return {
                "success": False,
                "error": "Đã vượt quá quota API. Vui lòng chờ hoặc nâng cấp plan.",
                "details": error_msg
            }
        elif "rate limit" in error_msg.lower():
            return {
                "success": False,
                "error": "Rate limit exceeded. Vui lòng thử lại sau ít phút.",
                "details": error_msg
            }
        else:
            return {
                "success": False,
                "error": f"Lỗi khi gọi Gemini API: {error_msg}"
            }


async def auto_process_document_with_gemini(user_query: str, model: str = "models/gemini-3-flash-preview") -> dict:
    """
    🤖 TỰ ĐỘNG PHÁT HIỆN VÀ XỬ LÝ TÀI LIỆU/DATABASE VỚI GEMINI
    
    Khi người dùng hỏi về:
    - Cơ sở dữ liệu (database, CSDL)
    - Tài liệu (PDF, Word, TXT, JSON, XML)
    - Files trong knowledge base
    
    Tự động:
    1. Phát hiện ý định người dùng
    2. Tìm và đọc tài liệu liên quan
    3. Gửi nội dung cho Gemini xử lý
    4. Trả về kết quả đã được Gemini phân tích
    
    Returns:
        dict với:
        - gemini_response: Kết quả đã được Gemini xử lý
        - documents_found: List các documents đã tìm thấy
        - success: True nếu thành công
    """
    try:
        query_lower = user_query.lower()
        
        # Phát hiện keywords về database/documents
        document_keywords = [
            'cơ sở dữ liệu', 'database', 'csdl', 'db',
            'tài liệu', 'document', 'file', 'files',
            'pdf', 'word', 'txt', 'json', 'xml', 'csv',
            'trong file', 'từ file', 'ở file',
            'knowledge base', 'kiến thức', 'tri thức',
            'đọc file', 'xem file', 'tìm trong',
            'thông tin trong', 'dữ liệu trong'
        ]
        
        # Check nếu query có chứa keywords
        has_document_intent = any(kw in query_lower for kw in document_keywords)
        
        if not has_document_intent:
            return {
                "success": False,
                "activated": False,
                "reason": "Query không liên quan đến documents/database"
            }
        
        print(f"📊 [Auto Document] Detected document query: {user_query[:100]}")
        
        # Step 1: Tìm documents liên quan từ knowledge base
        knowledge_result = await get_knowledge_context(
            query=user_query,
            max_chars=8000,  # Lấy nhiều context hơn
            use_gemini_summary=False  # Không tóm tắt trước, để Gemini xử lý toàn bộ
        )
        
        if not knowledge_result.get("success"):
            return {
                "success": False,
                "activated": True,
                "error": "Không tìm thấy documents trong knowledge base",
                "suggestion": "Hãy index các files bằng /api/knowledge/index_directory"
            }
        
        context = knowledge_result.get("context", "")
        documents_found = knowledge_result.get("documents_included", [])
        
        if not context:
            return {
                "success": False,
                "activated": True,
                "error": "Knowledge base trống",
                "documents_found": []
            }
        
        print(f"📚 [Auto Document] Found {len(documents_found)} documents")
        
        # Step 2: Gửi cho Gemini xử lý với context đầy đủ
        enhanced_prompt = f"""[TÀI LIỆU THAM KHẢO]
{context}

[CÂU HỎI CỦA NGƯỜI DÙNG]
{user_query}

[YÊU CẦU]
Dựa vào tài liệu trên, hãy trả lời câu hỏi một cách chính xác và chi tiết.
- Nếu có thông tin trong tài liệu, trích dẫn rõ ràng
- Nếu không có thông tin, hãy nói rõ
- Trả lời bằng Tiếng Việt, dễ hiểu"""

        # Gọi Gemini
        gemini_result = await ask_gemini(enhanced_prompt, model=model)
        
        if not gemini_result.get("success"):
            return {
                "success": False,
                "activated": True,
                "error": f"Gemini error: {gemini_result.get('error')}",
                "documents_found": documents_found
            }
        
        # Step 3: Trả về kết quả đã được Gemini xử lý
        return {
            "success": True,
            "activated": True,
            "gemini_response": gemini_result.get("response_text"),
            "documents_found": documents_found,
            "model_used": model,
            "context_length": len(context),
            "message": f"✅ Đã xử lý {len(documents_found)} documents với Gemini {model}"
        }
        
    except Exception as e:
        return {
            "success": False,
            "activated": True,
            "error": f"Error: {str(e)}"
        }


async def ask_gpt4(prompt: str, model: str = "gpt-4o") -> dict:
    """
    Hỏi đáp với OpenAI GPT-4
    
    Args:
        prompt: Câu hỏi hoặc nội dung muốn gửi cho GPT-4
        model: Tên model OpenAI (mặc định: gpt-4o - GPT-4 Omni, nhanh và rẻ)
        
    Returns:
        dict với success, response_text, và message
    """
    try:
        # Kiểm tra OpenAI có khả dụng không
        if not OPENAI_AVAILABLE:
            return {
                "success": False,
                "error": "OpenAI library chưa cài đặt. Chạy: pip install openai"
            }
        
        # Kiểm tra API key
        if not OPENAI_API_KEY or OPENAI_API_KEY.strip() == "":
            return {
                "success": False,
                "error": "OpenAI API key chưa được cấu hình. Vui lòng thêm 'openai_api_key' vào xiaozhi_endpoints.json",
                "help": "Lấy API key tại: https://platform.openai.com/api-keys"
            }
        
        # Khởi tạo OpenAI client
        print(f"[GPT-4] Configured with API key: ...{OPENAI_API_KEY[-8:]}")
        client = OpenAI(api_key=OPENAI_API_KEY)
        
        print(f"[GPT-4] Sending prompt with model: {model}")
        
        # Gọi API trong executor để không block event loop
        loop = asyncio.get_event_loop()
        response = await loop.run_in_executor(
            None,
            lambda: client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1000
            )
        )
        
        print(f"[GPT-4] Response received")
        
        # Lấy text từ response
        response_text = response.choices[0].message.content
        
        # 🔄 TRUNCATE: Giới hạn response dưới 4000 ký tự cho LLM
        if len(response_text) > MAX_LLM_RESPONSE_CHARS:
            original_len = len(response_text)
            response_text = smart_truncate_for_llm(response_text, MAX_LLM_RESPONSE_CHARS)
            print(f"[GPT-4] ✂️ Truncated: {original_len} → {len(response_text)} chars")
        
        print(f"[GPT-4] Response text: {response_text[:100]}...")
        
        return {
            "success": True,
            "prompt": prompt,
            "response_text": response_text,
            "model": model,
            "usage": {
                "prompt_tokens": response.usage.prompt_tokens,
                "completion_tokens": response.usage.completion_tokens,
                "total_tokens": response.usage.total_tokens
            },
            "message": f"✅ GPT-4 đã trả lời (model: {model})"
        }
        
    except Exception as e:
        error_msg = str(e)
        print(f"❌ [GPT-4] Exception caught: {type(e).__name__}")
        print(f"❌ [GPT-4] Error message: {error_msg}")
        
        import traceback
        traceback.print_exc()
        
        # Xử lý các lỗi phổ biến
        if "Incorrect API key" in error_msg or "invalid_api_key" in error_msg:
            return {
                "success": False,
                "error": "OpenAI API key không hợp lệ. Vui lòng kiểm tra lại openai_api_key trong xiaozhi_endpoints.json",
                "help": "Lấy API key mới tại: https://platform.openai.com/api-keys"
            }
        elif "insufficient_quota" in error_msg or "quota" in error_msg.lower():
            return {
                "success": False,
                "error": "Đã hết quota OpenAI. Vui lòng nạp tiền hoặc chờ quota reset.",
                "details": error_msg
            }
        elif "rate_limit" in error_msg.lower():
            return {
                "success": False,
                "error": "Rate limit exceeded. Vui lòng thử lại sau ít phút.",
                "details": error_msg
            }
        elif "model_not_found" in error_msg.lower():
            return {
                "success": False,
                "error": f"Model '{model}' không tồn tại. Thử: gpt-4o, gpt-4-turbo, gpt-3.5-turbo",
                "details": error_msg
            }
        else:
            return {
                "success": False,
                "error": f"Lỗi khi gọi OpenAI API: {error_msg}"
            }


# ============================================================
# OPEN API TOOLS - Các API công khai hữu ích
# Tham khảo từ: github.com/ZhongZiTongXue/xiaozhi-MCPTools
# ============================================================

import aiohttp
import urllib.parse

async def get_daily_news() -> dict:
    """
    Lấy tin tức 60 giây mỗi ngày (每日早报/60s morning news).
    Nguồn: API công khai
    """
    try:
        url = "https://60s.viki.moe/?v2=1"
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=10) as response:
                if response.status == 200:
                    data = await response.json()
                    news_list = data.get('data', [])[:10]  # Top 10 tin
                    formatted = "\n".join([f"{i+1}. {item}" for i, item in enumerate(news_list)])
                    return {
                        "success": True,
                        "message": "📰 Tin tức 60 giây hôm nay:",
                        "news": formatted,
                        "source": "60s.viki.moe"
                    }
                return {"success": False, "error": f"API trả về status {response.status}"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_random_quote() -> dict:
    """
    Lấy một câu nói ngẫu nhiên (一言/Hitokoto).
    """
    try:
        url = "https://v1.hitokoto.cn/"
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=10) as response:
                if response.status == 200:
                    data = await response.json()
                    return {
                        "success": True,
                        "quote": data.get('hitokoto', ''),
                        "from": data.get('from', 'Unknown'),
                        "author": data.get('from_who', ''),
                        "type": data.get('type', '')
                    }
                return {"success": False, "error": f"API error: {response.status}"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_hotlist(platform: str = "weibo") -> dict:
    """
    Lấy bảng xếp hạng hot từ các nền tảng (微博/知乎/百度/抖音).
    """
    try:
        platforms = {
            "weibo": "https://tenapi.cn/v2/weibohot",
            "zhihu": "https://tenapi.cn/v2/zhihuhot",
            "baidu": "https://tenapi.cn/v2/baiduhot",
            "douyin": "https://tenapi.cn/v2/douyinhot"
        }
        
        platform_lower = platform.lower()
        url = platforms.get(platform_lower)
        
        if not url:
            return {"success": False, "error": f"Platform không hỗ trợ. Chọn: weibo, zhihu, baidu, douyin"}
        
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=10) as response:
                if response.status == 200:
                    data = await response.json()
                    hot_list = data.get('data', [])[:15]  # Top 15
                    formatted = "\n".join([f"{i+1}. {item.get('name', item.get('title', ''))}" for i, item in enumerate(hot_list)])
                    return {
                        "success": True,
                        "platform": platform,
                        "hotlist": formatted,
                        "count": len(hot_list)
                    }
                return {"success": False, "error": f"API error: {response.status}"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def search_baike(query: str) -> dict:
    """
    Tìm kiếm Baidu Baike (百度百科).
    """
    try:
        encoded_query = urllib.parse.quote(query)
        url = f"https://baike.baidu.com/api/openapi/BaikeLemmaCardApi?scope=103&format=json&appid=379020&bk_key={encoded_query}"
        
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=10) as response:
                if response.status == 200:
                    data = await response.json()
                    if data.get('id'):
                        return {
                            "success": True,
                            "title": data.get('title', ''),
                            "abstract": data.get('abstract', ''),
                            "url": data.get('url', ''),
                            "image": data.get('image', '')
                        }
                    return {"success": False, "error": f"Không tìm thấy '{query}' trên Baike"}
                return {"success": False, "error": f"API error: {response.status}"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_history_today() -> dict:
    """
    Lấy sự kiện lịch sử ngày hôm nay (历史上的今天).
    """
    try:
        from datetime import datetime
        today = datetime.now()
        month = today.month
        day = today.day
        
        url = f"https://api.oioweb.cn/api/common/history?month={month}&day={day}"
        
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=10) as response:
                if response.status == 200:
                    data = await response.json()
                    events = data.get('result', [])[:10]
                    formatted = "\n".join([f"• {e.get('year', '')}: {e.get('title', '')}" for e in events])
                    return {
                        "success": True,
                        "date": f"{month}/{day}",
                        "events": formatted,
                        "count": len(events)
                    }
                return {"success": False, "error": f"API error: {response.status}"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_joke() -> dict:
    """
    Lấy một câu chuyện cười ngẫu nhiên.
    """
    try:
        url = "https://api.oioweb.cn/api/common/joke"
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=10) as response:
                if response.status == 200:
                    data = await response.json()
                    return {
                        "success": True,
                        "joke": data.get('result', {}).get('content', 'Không có joke'),
                        "source": "oioweb.cn"
                    }
                return {"success": False, "error": f"API error: {response.status}"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_weather_simple(city: str = "Hanoi") -> dict:
    """
    Lấy thời tiết đơn giản của thành phố.
    """
    try:
        # Dùng wttr.in API (free, không cần key)
        encoded_city = urllib.parse.quote(city)
        url = f"https://wttr.in/{encoded_city}?format=j1"
        
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=15) as response:
                if response.status == 200:
                    data = await response.json()
                    current = data.get('current_condition', [{}])[0]
                    weather_desc = current.get('weatherDesc', [{}])[0].get('value', '')
                    temp_c = current.get('temp_C', '')
                    humidity = current.get('humidity', '')
                    wind_kmph = current.get('windspeedKmph', '')
                    
                    return {
                        "success": True,
                        "city": city,
                        "weather": weather_desc,
                        "temperature": f"{temp_c}°C",
                        "humidity": f"{humidity}%",
                        "wind": f"{wind_kmph} km/h",
                        "summary": f"🌤️ {city}: {weather_desc}, {temp_c}°C, Độ ẩm {humidity}%"
                    }
                return {"success": False, "error": f"Không tìm thấy thời tiết cho '{city}'"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def control_ppt(action: str) -> dict:
    """
    Điều khiển PowerPoint presentation.
    Actions: next (trang sau), prev (trang trước), start (bắt đầu trình chiếu), end (kết thúc)
    """
    try:
        import pyautogui
        
        action_lower = action.lower()
        
        if action_lower in ['next', 'tiếp', 'trang sau']:
            pyautogui.press('right')
            return {"success": True, "message": "➡️ PPT: Chuyển trang sau"}
            
        elif action_lower in ['prev', 'previous', 'trước', 'trang trước']:
            pyautogui.press('left')
            return {"success": True, "message": "⬅️ PPT: Quay lại trang trước"}
            
        elif action_lower in ['start', 'bắt đầu', 'trình chiếu']:
            pyautogui.press('f5')
            return {"success": True, "message": "▶️ PPT: Bắt đầu trình chiếu từ đầu"}
            
        elif action_lower in ['start_current', 'từ trang này']:
            pyautogui.hotkey('shift', 'f5')
            return {"success": True, "message": "▶️ PPT: Trình chiếu từ trang hiện tại"}
            
        elif action_lower in ['end', 'kết thúc', 'thoát']:
            pyautogui.press('escape')
            return {"success": True, "message": "⏹️ PPT: Kết thúc trình chiếu"}
            
        else:
            return {
                "success": False,
                "error": f"Action '{action}' không hợp lệ",
                "hint": "Các action hỗ trợ: next, prev, start, start_current, end"
            }
            
    except Exception as e:
        return {"success": False, "error": str(e)}

async def ask_doubao(question: str) -> dict:
    """
    Mở Doubao AI và gửi câu hỏi (yêu cầu có browser).
    """
    try:
        import webbrowser
        import pyperclip
        import pyautogui
        import time
        
        url = "https://www.doubao.com/chat/"
        webbrowser.open(url)
        
        # Đợi trang load
        time.sleep(4)
        
        # Copy câu hỏi và paste
        pyperclip.copy(question)
        time.sleep(0.3)
        pyautogui.hotkey('ctrl', 'v')
        time.sleep(0.5)
        pyautogui.press('enter')
        
        return {
            "success": True,
            "message": f"✅ Đã gửi câu hỏi tới Doubao AI: '{question}'",
            "note": "Vui lòng xem kết quả trên trình duyệt"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

async def ask_kimi(question: str) -> dict:
    """
    Mở Kimi AI và gửi câu hỏi (yêu cầu có browser).
    """
    try:
        import webbrowser
        import pyperclip
        import pyautogui
        import time
        
        url = "https://kimi.moonshot.cn/"
        webbrowser.open(url)
        
        # Đợi trang load
        time.sleep(4)
        
        # Copy câu hỏi và paste
        pyperclip.copy(question)
        time.sleep(0.3)
        pyautogui.hotkey('ctrl', 'v')
        time.sleep(0.5)
        pyautogui.press('enter')
        
        return {
            "success": True,
            "message": f"✅ Đã gửi câu hỏi tới Kimi AI: '{question}'",
            "note": "Vui lòng xem kết quả trên trình duyệt"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

async def set_dark_light_theme(mode: str) -> dict:
    """
    Chuyển đổi theme Windows Dark/Light mode.
    """
    try:
        import subprocess
        
        mode_lower = mode.lower()
        
        if mode_lower in ['dark', 'tối', 'đen']:
            # Set dark mode
            subprocess.run([
                'reg', 'add', 
                'HKCU\\SOFTWARE\\Microsoft\\Windows\\CurrentVersion\\Themes\\Personalize',
                '/v', 'AppsUseLightTheme', '/t', 'REG_DWORD', '/d', '0', '/f'
            ], capture_output=True)
            subprocess.run([
                'reg', 'add',
                'HKCU\\SOFTWARE\\Microsoft\\Windows\\CurrentVersion\\Themes\\Personalize',
                '/v', 'SystemUsesLightTheme', '/t', 'REG_DWORD', '/d', '0', '/f'
            ], capture_output=True)
            return {"success": True, "message": "🌙 Đã chuyển sang Dark Mode"}
            
        elif mode_lower in ['light', 'sáng', 'trắng']:
            # Set light mode
            subprocess.run([
                'reg', 'add',
                'HKCU\\SOFTWARE\\Microsoft\\Windows\\CurrentVersion\\Themes\\Personalize',
                '/v', 'AppsUseLightTheme', '/t', 'REG_DWORD', '/d', '1', '/f'
            ], capture_output=True)
            subprocess.run([
                'reg', 'add',
                'HKCU\\SOFTWARE\\Microsoft\\Windows\\CurrentVersion\\Themes\\Personalize',
                '/v', 'SystemUsesLightTheme', '/t', 'REG_DWORD', '/d', '1', '/f'
            ], capture_output=True)
            return {"success": True, "message": "☀️ Đã chuyển sang Light Mode"}
            
        else:
            return {
                "success": False,
                "error": f"Mode '{mode}' không hợp lệ",
                "hint": "Chọn: dark/tối hoặc light/sáng"
            }
            
    except Exception as e:
        return {"success": False, "error": str(e)}

async def lock_computer() -> dict:
    """
    Khóa máy tính ngay lập tức.
    """
    try:
        import ctypes
        ctypes.windll.user32.LockWorkStation()
        return {"success": True, "message": "🔒 Đã khóa máy tính"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def shutdown_computer(action: str = "shutdown", delay: int = 0) -> dict:
    """
    Tắt máy/Khởi động lại/Hẹn giờ tắt.
    action: shutdown, restart, cancel (hủy lệnh tắt)
    delay: số giây trước khi thực hiện (0 = ngay lập tức)
    """
    try:
        import subprocess
        
        action_lower = action.lower()
        
        if action_lower in ['shutdown', 'tắt', 'tắt máy']:
            subprocess.run(['shutdown', '/s', '/t', str(delay)], capture_output=True)
            if delay > 0:
                return {"success": True, "message": f"⏰ Máy sẽ tắt sau {delay} giây"}
            return {"success": True, "message": "⏹️ Đang tắt máy..."}
            
        elif action_lower in ['restart', 'khởi động lại', 'reboot']:
            subprocess.run(['shutdown', '/r', '/t', str(delay)], capture_output=True)
            if delay > 0:
                return {"success": True, "message": f"⏰ Máy sẽ khởi động lại sau {delay} giây"}
            return {"success": True, "message": "🔄 Đang khởi động lại..."}
            
        elif action_lower in ['cancel', 'hủy', 'abort']:
            subprocess.run(['shutdown', '/a'], capture_output=True)
            return {"success": True, "message": "❌ Đã hủy lệnh tắt/khởi động lại"}
            
        else:
            return {
                "success": False,
                "error": f"Action '{action}' không hợp lệ",
                "hint": "Chọn: shutdown, restart, cancel"
            }
            
    except Exception as e:
        return {"success": False, "error": str(e)}

# DUPLICATE REMOVED: change_wallpaper was defined twice (first at line 5503)
# DUPLICATE REMOVED: find_in_document was defined twice (first at line 5809)

async def clipboard_read() -> dict:
    """
    Đọc nội dung từ clipboard.
    """
    try:
        import pyperclip
        content = pyperclip.paste()
        return {
            "success": True,
            "content": content,
            "length": len(content) if content else 0
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

async def clipboard_write(content: str) -> dict:
    """
    Ghi nội dung vào clipboard.
    """
    try:
        import pyperclip
        pyperclip.copy(content)
        return {"success": True, "message": f"📋 Đã copy vào clipboard ({len(content)} ký tự)"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def type_text(text: str, press_enter: bool = False) -> dict:
    """
    Gõ text vào vị trí con trỏ hiện tại.
    """
    try:
        import pyperclip
        import pyautogui
        import time
        
        # Copy và paste để hỗ trợ Unicode
        pyperclip.copy(text)
        time.sleep(0.1)
        pyautogui.hotkey('ctrl', 'v')
        
        if press_enter:
            time.sleep(0.2)
            pyautogui.press('enter')
            return {"success": True, "message": f"⌨️ Đã gõ và Enter: '{text[:50]}...'"}
        
        return {"success": True, "message": f"⌨️ Đã gõ: '{text[:50]}...'"}
        
    except Exception as e:
        return {"success": False, "error": str(e)}

async def undo_action() -> dict:
    """
    Thực hiện Undo (Ctrl+Z).
    """
    try:
        import pyautogui
        pyautogui.hotkey('ctrl', 'z')
        return {"success": True, "message": "↩️ Đã Undo"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def show_desktop() -> dict:
    """
    Hiển thị Desktop (Win+D).
    """
    try:
        import pyautogui
        pyautogui.hotkey('win', 'd')
        return {"success": True, "message": "🖥️ Đã hiển thị Desktop"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ============================================================
# OPEN API TOOLS - Các công cụ tra cứu thông tin (PHÙ HỢP VIỆT NAM)
# ============================================================

async def get_weather_vietnam(city: str = "Hà Nội") -> dict:
    """
    Lấy thông tin thời tiết Việt Nam từ wttr.in (miễn phí, không cần API key).
    """
    try:
        import aiohttp
        import urllib.parse
        
        # Normalize tên thành phố
        city_mapping = {
            "hà nội": "Hanoi", "ha noi": "Hanoi", "hanoi": "Hanoi",
            "hồ chí minh": "Ho Chi Minh", "ho chi minh": "Ho Chi Minh", "saigon": "Ho Chi Minh", "sài gòn": "Ho Chi Minh",
            "đà nẵng": "Da Nang", "da nang": "Da Nang", "danang": "Da Nang",
            "hải phòng": "Hai Phong", "hai phong": "Hai Phong",
            "cần thơ": "Can Tho", "can tho": "Can Tho",
            "nha trang": "Nha Trang", "huế": "Hue", "hue": "Hue",
            "vũng tàu": "Vung Tau", "vung tau": "Vung Tau",
            "biên hòa": "Bien Hoa", "bien hoa": "Bien Hoa",
            "buôn ma thuột": "Buon Ma Thuot", "đà lạt": "Da Lat", "da lat": "Da Lat",
            "quảng ninh": "Quang Ninh", "hạ long": "Ha Long",
            "thanh hóa": "Thanh Hoa", "vinh": "Vinh", "quy nhơn": "Quy Nhon",
        }
        
        city_query = city_mapping.get(city.lower().strip(), city)
        url = f"https://wttr.in/{urllib.parse.quote(city_query)}?format=j1"
        
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=10) as resp:
                if resp.status == 200:
                    data = await resp.json()
                    current = data.get("current_condition", [{}])[0]
                    
                    temp_c = current.get("temp_C", "N/A")
                    feels_like = current.get("FeelsLikeC", "N/A")
                    humidity = current.get("humidity", "N/A")
                    weather_desc = current.get("lang_vi", [{}])
                    if weather_desc:
                        weather_desc = weather_desc[0].get("value", current.get("weatherDesc", [{}])[0].get("value", ""))
                    else:
                        weather_desc = current.get("weatherDesc", [{}])[0].get("value", "")
                    wind_kmph = current.get("windspeedKmph", "N/A")
                    
                    return {
                        "success": True,
                        "city": city,
                        "temperature": f"{temp_c}°C",
                        "feels_like": f"{feels_like}°C",
                        "humidity": f"{humidity}%",
                        "weather": weather_desc,
                        "wind": f"{wind_kmph} km/h",
                        "message": f"🌤️ Thời tiết {city}: {temp_c}°C, {weather_desc}, Độ ẩm {humidity}%, Gió {wind_kmph}km/h"
                    }
                else:
                    return {"success": False, "error": f"Không lấy được thời tiết: HTTP {resp.status}"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_gold_price_vietnam() -> dict:
    """
    Lấy giá vàng Việt Nam từ API miễn phí.
    """
    try:
        import aiohttp
        
        # Sử dụng API giá vàng SJC
        url = "https://api.btmc.vn/api/BTMCAPI/getpricesheet"
        
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=10) as resp:
                if resp.status == 200:
                    data = await resp.json()
                    
                    # Tìm giá vàng SJC
                    gold_prices = []
                    for item in data.get("data", []):
                        name = item.get("name", "")
                        buy = item.get("buy", 0)
                        sell = item.get("sell", 0)
                        if "SJC" in name or "vàng" in name.lower():
                            gold_prices.append({
                                "name": name,
                                "buy": f"{buy:,.0f}".replace(",", "."),
                                "sell": f"{sell:,.0f}".replace(",", ".")
                            })
                    
                    if gold_prices:
                        msg = "💰 Giá vàng hôm nay:\n"
                        for g in gold_prices[:3]:  # Top 3
                            msg += f"• {g['name']}: Mua {g['buy']} - Bán {g['sell']} VNĐ/lượng\n"
                        
                        return {
                            "success": True,
                            "prices": gold_prices[:3],
                            "message": msg.strip()
                        }
                    
                return {"success": False, "error": "Không lấy được giá vàng"}
    except Exception as e:
        # Fallback: trả về thông tin hướng dẫn
        return {
            "success": True,
            "message": "💰 Để xem giá vàng mới nhất, truy cập: sjc.com.vn hoặc pnj.com.vn",
            "hint": "API giá vàng tạm thời không khả dụng"
        }

async def get_exchange_rate_vietnam(currency: str = "USD") -> dict:
    """
    Lấy tỷ giá ngoại tệ so với VND.
    """
    try:
        import aiohttp
        
        currency = currency.upper().strip()
        
        # Dùng API miễn phí exchangerate-api
        url = f"https://api.exchangerate-api.com/v4/latest/{currency}"
        
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=10) as resp:
                if resp.status == 200:
                    data = await resp.json()
                    rates = data.get("rates", {})
                    vnd_rate = rates.get("VND", 0)
                    
                    if vnd_rate:
                        return {
                            "success": True,
                            "currency": currency,
                            "vnd_rate": f"{vnd_rate:,.0f}".replace(",", "."),
                            "message": f"💱 Tỷ giá: 1 {currency} = {vnd_rate:,.0f} VNĐ".replace(",", ".")
                        }
                        
                return {"success": False, "error": f"Không tìm thấy tỷ giá {currency}"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_daily_quote() -> dict:
    """
    Lấy câu nói hay/trích dẫn ngẫu nhiên.
    """
    try:
        import aiohttp
        import random
        
        # Các quotes tiếng Việt đẹp
        vietnamese_quotes = [
            {"quote": "Thành công không phải là chìa khóa của hạnh phúc. Hạnh phúc là chìa khóa của thành công.", "author": "Albert Schweitzer"},
            {"quote": "Đừng sợ thất bại. Hãy sợ những cơ hội bạn bỏ lỡ khi không cố gắng.", "author": "Jack Canfield"},
            {"quote": "Cuộc sống không phải là chờ đợi bão qua đi, mà là học cách nhảy múa dưới mưa.", "author": "Vivian Greene"},
            {"quote": "Hôm nay khó khăn, ngày mai còn khó khăn hơn, nhưng ngày kia sẽ tươi đẹp.", "author": "Jack Ma"},
            {"quote": "Người duy nhất bạn cần vượt qua là chính bạn của ngày hôm qua.", "author": "Khuyết danh"},
            {"quote": "Học hỏi không có điểm dừng, giống như cuộc sống không có giới hạn.", "author": "Khổng Tử"},
            {"quote": "Thất bại là mẹ thành công.", "author": "Tục ngữ Việt Nam"},
            {"quote": "Có chí thì nên.", "author": "Tục ngữ Việt Nam"},
            {"quote": "Một cây làm chẳng nên non, ba cây chụm lại nên hòn núi cao.", "author": "Ca dao Việt Nam"},
            {"quote": "Đi một ngày đàng, học một sàng khôn.", "author": "Tục ngữ Việt Nam"},
        ]
        
        # Thử lấy quote từ API
        try:
            url = "https://api.quotable.io/random"
            async with aiohttp.ClientSession() as session:
                async with session.get(url, timeout=5) as resp:
                    if resp.status == 200:
                        data = await resp.json()
                        return {
                            "success": True,
                            "quote": data.get("content", ""),
                            "author": data.get("author", "Unknown"),
                            "message": f"💬 \"{data.get('content', '')}\" - {data.get('author', 'Unknown')}"
                        }
        except:
            pass
        
        # Fallback: quote tiếng Việt
        quote = random.choice(vietnamese_quotes)
        return {
            "success": True,
            "quote": quote["quote"],
            "author": quote["author"],
            "message": f"💬 \"{quote['quote']}\" - {quote['author']}"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_today_in_history() -> dict:
    """
    Lấy sự kiện lịch sử ngày hôm nay.
    """
    try:
        import aiohttp
        from datetime import datetime
        
        today = datetime.now()
        month = today.month
        day = today.day
        
        url = f"https://history.muffinlabs.com/date/{month}/{day}"
        
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=10) as resp:
                if resp.status == 200:
                    data = await resp.json()
                    events = data.get("data", {}).get("Events", [])[:3]
                    
                    if events:
                        msg = f"📜 Ngày này ({day}/{month}) trong lịch sử:\n"
                        for event in events:
                            year = event.get("year", "")
                            text = event.get("text", "")
                            msg += f"• {year}: {text[:100]}...\n" if len(text) > 100 else f"• {year}: {text}\n"
                        
                        return {
                            "success": True,
                            "date": f"{day}/{month}",
                            "events": events,
                            "message": msg.strip()
                        }
                        
        return {"success": False, "error": "Không lấy được sự kiện lịch sử"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_joke() -> dict:
    """
    Lấy một câu chuyện cười/joke ngẫu nhiên.
    """
    try:
        import random
        
        # Jokes tiếng Việt
        vietnamese_jokes = [
            "Tại sao con cá không biết nói? Vì nó ở dưới nước, nói sao được! 🐟",
            "Bạn biết con gì nhanh nhất thế giới không? Con gió, vì nó đi vèo vèo! 💨",
            "Tại sao con kiến không bao giờ ốm? Vì nó có đầy đủ chất sắt (Fe) trong người! 🐜",
            "Ai là người hạnh phúc nhất? Người không biết so sánh! 😊",
            "Con gì có 4 chân mà không biết đi? Cái bàn! 🪑",
            "Tại sao máy tính không bao giờ khóc? Vì nó có mouse pad (miếng lót chuột)! 🖱️",
            "Bạn biết tại sao mặt trời đi học không? Vì nó đã tốt nghiệp từ lâu rồi! ☀️",
            "Tại sao con gà qua đường? Để đến bên kia đường! 🐔",
            "Con gì ngồi một chỗ mà vẫn chạy? Cái đồng hồ! ⏰",
            "Tại sao cầu vồng thích đi chơi? Vì nó có 7 màu = 7 ngày = 1 tuần! 🌈",
        ]
        
        joke = random.choice(vietnamese_jokes)
        return {
            "success": True,
            "joke": joke,
            "message": f"😂 {joke}"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_horoscope(zodiac: str = "song_tử") -> dict:
    """
    Lấy tử vi/horoscope theo cung hoàng đạo.
    """
    try:
        import random
        
        # Map tên cung hoàng đạo
        zodiac_map = {
            "bạch dương": "aries", "bach duong": "aries", "aries": "aries",
            "kim ngưu": "taurus", "kim nguu": "taurus", "taurus": "taurus",
            "song tử": "gemini", "song tu": "gemini", "gemini": "gemini",
            "cự giải": "cancer", "cu giai": "cancer", "cancer": "cancer",
            "sư tử": "leo", "su tu": "leo", "leo": "leo",
            "xử nữ": "virgo", "xu nu": "virgo", "virgo": "virgo",
            "thiên bình": "libra", "thien binh": "libra", "libra": "libra",
            "bọ cạp": "scorpio", "bo cap": "scorpio", "scorpio": "scorpio",
            "nhân mã": "sagittarius", "nhan ma": "sagittarius", "sagittarius": "sagittarius",
            "ma kết": "capricorn", "ma ket": "capricorn", "capricorn": "capricorn",
            "bảo bình": "aquarius", "bao binh": "aquarius", "aquarius": "aquarius",
            "song ngư": "pisces", "song ngu": "pisces", "pisces": "pisces",
        }
        
        zodiac_key = zodiac_map.get(zodiac.lower().strip(), "gemini")
        zodiac_name = zodiac.title()
        
        # Random horoscope messages
        luck_levels = ["⭐⭐⭐⭐⭐", "⭐⭐⭐⭐", "⭐⭐⭐", "⭐⭐⭐⭐"]
        love_messages = [
            "Tình yêu đang đến gần, hãy mở lòng đón nhận.",
            "Hôm nay là ngày tốt để thể hiện tình cảm.",
            "Người ấy đang nghĩ về bạn nhiều hơn bạn tưởng.",
            "Hãy kiên nhẫn, tình yêu đích thực cần thời gian.",
        ]
        career_messages = [
            "Công việc suôn sẻ, cơ hội thăng tiến đang mở ra.",
            "Hãy tập trung vào mục tiêu, thành công sẽ đến.",
            "Một dự án mới có thể xuất hiện bất ngờ.",
            "Đồng nghiệp sẽ hỗ trợ bạn rất nhiều hôm nay.",
        ]
        money_messages = [
            "Tài chính ổn định, có thể có khoản thu bất ngờ.",
            "Hãy cẩn thận với các quyết định đầu tư.",
            "Đây là thời điểm tốt để tiết kiệm.",
            "May mắn về tài chính đang mỉm cười với bạn.",
        ]
        
        return {
            "success": True,
            "zodiac": zodiac_name,
            "luck": random.choice(luck_levels),
            "love": random.choice(love_messages),
            "career": random.choice(career_messages),
            "money": random.choice(money_messages),
            "message": f"🔮 Tử vi {zodiac_name}:\n• May mắn: {random.choice(luck_levels)}\n• Tình yêu: {random.choice(love_messages)}\n• Sự nghiệp: {random.choice(career_messages)}\n• Tài chính: {random.choice(money_messages)}"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_news_vietnam() -> dict:
    """
    Lấy tin tức nóng Việt Nam.
    """
    try:
        import aiohttp
        
        # Dùng RSS feed từ các báo Việt Nam
        rss_urls = [
            "https://vnexpress.net/rss/tin-moi-nhat.rss",
            "https://tuoitre.vn/rss/tin-moi-nhat.rss",
        ]
        
        async with aiohttp.ClientSession() as session:
            for rss_url in rss_urls:
                try:
                    async with session.get(rss_url, timeout=10) as resp:
                        if resp.status == 200:
                            import xml.etree.ElementTree as ET
                            content = await resp.text()
                            root = ET.fromstring(content)
                            
                            items = root.findall('.//item')[:5]
                            news = []
                            
                            for item in items:
                                title = item.find('title')
                                title_text = title.text if title is not None else "No title"
                                news.append(title_text)
                            
                            if news:
                                msg = "📰 Tin tức mới nhất:\n"
                                for i, n in enumerate(news, 1):
                                    msg += f"{i}. {n}\n"
                                
                                result = {
                                    "success": True,
                                    "news": news,
                                    "message": msg.strip()
                                }
                                
                                # 🤖 GEMINI SUMMARIZATION: Tóm tắt nhanh bằng Gemini (non-blocking)
                                try:
                                    context = "\n".join([f"{i+1}. {n}" for i, n in enumerate(news)])
                                    # ⚡ PROMPT NGẮN GỌN - phản hồi nhanh hơn
                                    summary_prompt = f"""Tóm tắt 5 tin VN sau thành 3 ý chính:
{context}

Format: 📌 [3 điểm] + 🔹 [xu hướng chung 1 câu]"""
                                    
                                    print(f"⚡ [NewsVN+Gemini] Tóm tắt nhanh {len(news)} tin...")
                                    # ⏱️ Timeout 15 giây - đủ thời gian cho Gemini
                                    gemini_summary = await asyncio.wait_for(
                                        ask_gemini_direct(summary_prompt, model="models/gemini-3-flash-preview"),
                                        timeout=15.0
                                    )
                                    if gemini_summary.get("success"):
                                        summary_text = gemini_summary["response_text"]
                                        result["gemini_summary"] = summary_text
                                        result["message"] = f"✨ {summary_text}\n\n" + result["message"]
                                        print(f"✅ [NewsVN+Gemini] Done ({len(summary_text)} chars)")
                                    else:
                                        print(f"⚠️ [NewsVN+Gemini] Failed: {gemini_summary.get('error')}")
                                except asyncio.TimeoutError:
                                    print(f"⏱️ [NewsVN+Gemini] Timeout - trả tin thô")
                                except Exception as e:
                                    print(f"⚠️ [NewsVN+Gemini] Error: {e}")
                                
                                return result
                except:
                    continue
                    
        return {"success": False, "error": "Không lấy được tin tức"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def what_to_eat() -> dict:
    """
    Gợi ý món ăn hôm nay (Việt Nam).
    """
    try:
        import random
        from datetime import datetime
        
        # Món ăn Việt Nam theo bữa
        breakfast = [
            "🍜 Phở bò tái nạm", "🥖 Bánh mì thịt", "🍲 Bún bò Huế", 
            "🥣 Cháo lòng", "🍜 Hủ tiếu Nam Vang", "🥐 Bánh cuốn",
            "🍲 Bún riêu cua", "🥣 Xôi xéo", "🍜 Mì Quảng"
        ]
        
        lunch = [
            "🍚 Cơm tấm sườn bì chả", "🍲 Bún chả Hà Nội", "🍜 Phở gà",
            "🥗 Gỏi cuốn tôm thịt", "🍲 Lẩu thái", "🍚 Cơm văn phòng",
            "🍜 Bún đậu mắm tôm", "🍲 Canh chua cá lóc", "🍚 Cơm gà Tam Kỳ"
        ]
        
        dinner = [
            "🍖 Bò né", "🦐 Hải sản nướng", "🍲 Lẩu gà lá é",
            "🍗 Gà nướng muối ớt", "🥘 Cá kho tộ", "🍲 Lẩu Thái",
            "🍖 BBQ Hàn Quốc", "🍜 Phở cuốn", "🍲 Ốc xào me"
        ]
        
        snacks = [
            "🧁 Bánh tráng trộn", "🍡 Chè thập cẩm", "🍦 Kem bơ",
            "🥤 Trà sữa", "🍵 Cà phê sữa đá", "🍩 Bánh rán"
        ]
        
        hour = datetime.now().hour
        
        if 5 <= hour < 10:
            meal_type = "sáng"
            suggestion = random.choice(breakfast)
        elif 10 <= hour < 14:
            meal_type = "trưa"
            suggestion = random.choice(lunch)
        elif 14 <= hour < 17:
            meal_type = "xế"
            suggestion = random.choice(snacks)
        else:
            meal_type = "tối"
            suggestion = random.choice(dinner)
        
        return {
            "success": True,
            "meal_type": meal_type,
            "suggestion": suggestion,
            "alternatives": [random.choice(breakfast + lunch + dinner) for _ in range(2)],
            "message": f"🍽️ Bữa {meal_type} hôm nay: {suggestion}\n\n💡 Gợi ý khác: {random.choice(breakfast + lunch + dinner)}, {random.choice(breakfast + lunch + dinner)}"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_lunar_date() -> dict:
    """
    Lấy ngày âm lịch Việt Nam hôm nay - thuật toán chính xác.
    Tính theo múi giờ Việt Nam (UTC+7).
    """
    try:
        from datetime import datetime, timezone, timedelta
        import math
        
        # Múi giờ Việt Nam UTC+7
        vn_tz = timezone(timedelta(hours=7))
        today = datetime.now(vn_tz)
        
        # ========== THUẬT TOÁN TÍNH ÂM LỊCH VIỆT NAM ==========
        # Dựa trên thuật toán của Hồ Ngọc Đức
        
        def jd_from_date(dd, mm, yy):
            """Chuyển ngày dương lịch sang Julian Day Number"""
            a = int((14 - mm) / 12)
            y = yy + 4800 - a
            m = mm + 12 * a - 3
            jd = dd + int((153 * m + 2) / 5) + 365 * y + int(y / 4) - int(y / 100) + int(y / 400) - 32045
            if jd < 2299161:
                jd = dd + int((153 * m + 2) / 5) + 365 * y + int(y / 4) - 32083
            return jd
        
        def new_moon(k):
            """Tính thời điểm trăng mới thứ k (kể từ 1900-01-01)"""
            T = k / 1236.85
            T2 = T * T
            T3 = T2 * T
            dr = math.pi / 180
            Jd1 = 2415020.75933 + 29.53058868 * k + 0.0001178 * T2 - 0.000000155 * T3
            Jd1 = Jd1 + 0.00033 * math.sin((166.56 + 132.87 * T - 0.009173 * T2) * dr)
            M = 359.2242 + 29.10535608 * k - 0.0000333 * T2 - 0.00000347 * T3
            Mpr = 306.0253 + 385.81691806 * k + 0.0107306 * T2 + 0.00001236 * T3
            F = 21.2964 + 390.67050646 * k - 0.0016528 * T2 - 0.00000239 * T3
            C1 = (0.1734 - 0.000393 * T) * math.sin(M * dr) + 0.0021 * math.sin(2 * dr * M)
            C1 = C1 - 0.4068 * math.sin(Mpr * dr) + 0.0161 * math.sin(dr * 2 * Mpr)
            C1 = C1 - 0.0004 * math.sin(dr * 3 * Mpr)
            C1 = C1 + 0.0104 * math.sin(dr * 2 * F) - 0.0051 * math.sin(dr * (M + Mpr))
            C1 = C1 - 0.0074 * math.sin(dr * (M - Mpr)) + 0.0004 * math.sin(dr * (2 * F + M))
            C1 = C1 - 0.0004 * math.sin(dr * (2 * F - M)) - 0.0006 * math.sin(dr * (2 * F + Mpr))
            C1 = C1 + 0.0010 * math.sin(dr * (2 * F - Mpr)) + 0.0005 * math.sin(dr * (2 * Mpr + M))
            if T < -11:
                deltat = 0.001 + 0.000839 * T + 0.0002261 * T2 - 0.00000845 * T3 - 0.000000081 * T * T3
            else:
                deltat = -0.000278 + 0.000265 * T + 0.000262 * T2
            return Jd1 + C1 - deltat
        
        def sun_longitude(jdn):
            """Tính kinh độ mặt trời tại thời điểm Julian Day Number"""
            T = (jdn - 2451545.0) / 36525
            T2 = T * T
            dr = math.pi / 180
            M = 357.52910 + 35999.05030 * T - 0.0001559 * T2 - 0.00000048 * T * T2
            L0 = 280.46645 + 36000.76983 * T + 0.0003032 * T2
            DL = (1.914600 - 0.004817 * T - 0.000014 * T2) * math.sin(dr * M)
            DL = DL + (0.019993 - 0.000101 * T) * math.sin(dr * 2 * M) + 0.00029 * math.sin(dr * 3 * M)
            L = L0 + DL
            L = L * dr
            L = L - math.pi * 2 * int(L / (math.pi * 2))
            return int(L / math.pi * 6)
        
        def get_lunar_month_11(yy):
            """Tìm ngày bắt đầu tháng 11 âm lịch"""
            off = jd_from_date(31, 12, yy) - 2415021
            k = int(off / 29.530588853)
            nm = new_moon(k)
            sun_long = sun_longitude(nm)
            if sun_long >= 9:
                nm = new_moon(k - 1)
            return int(nm + 0.5)
        
        def get_leap_month_offset(a11):
            """Xác định tháng nhuận"""
            k = int((a11 - 2415021.076998695) / 29.530588853 + 0.5)
            last = 0
            i = 1
            arc = sun_longitude(new_moon(k + i))
            while True:
                last = arc
                i += 1
                arc = sun_longitude(new_moon(k + i))
                if arc != last or i >= 14:
                    break
            return i - 1
        
        def solar_to_lunar(dd, mm, yy):
            """Chuyển ngày dương lịch sang âm lịch"""
            day_number = jd_from_date(dd, mm, yy)
            k = int((day_number - 2415021.076998695) / 29.530588853)
            month_start = new_moon(k + 1)
            if month_start > day_number:
                month_start = new_moon(k)
            a11 = get_lunar_month_11(yy)
            b11 = a11
            if a11 >= month_start:
                lunar_year = yy
                a11 = get_lunar_month_11(yy - 1)
            else:
                lunar_year = yy + 1
                b11 = get_lunar_month_11(yy + 1)
            lunar_day = int(day_number - month_start + 1)
            diff = int((month_start - a11) / 29)
            lunar_leap = 0
            lunar_month = diff + 11
            if b11 - a11 > 365:
                leap_month_diff = get_leap_month_offset(a11)
                if diff >= leap_month_diff:
                    lunar_month = diff + 10
                    if diff == leap_month_diff:
                        lunar_leap = 1
            if lunar_month > 12:
                lunar_month = lunar_month - 12
            if lunar_month >= 11 and diff < 4:
                lunar_year -= 1
            return lunar_day, lunar_month, lunar_year, lunar_leap
        
        # ========== TÍNH CAN CHI ==========
        CAN = ["Giáp", "Ất", "Bính", "Đinh", "Mậu", "Kỷ", "Canh", "Tân", "Nhâm", "Quý"]
        CHI = ["Tý", "Sửu", "Dần", "Mão", "Thìn", "Tỵ", "Ngọ", "Mùi", "Thân", "Dậu", "Tuất", "Hợi"]
        
        def get_can_chi_year(lunar_year):
            """Lấy can chi của năm"""
            can = CAN[(lunar_year + 6) % 10]
            chi = CHI[(lunar_year + 8) % 12]
            return f"{can} {chi}"
        
        def get_can_chi_day(dd, mm, yy):
            """Lấy can chi của ngày"""
            jd = jd_from_date(dd, mm, yy)
            can = CAN[(jd + 9) % 10]
            chi = CHI[(jd + 1) % 12]
            return f"{can} {chi}"
        
        # ========== TÍNH NGÀY ÂM LỊCH HÔM NAY ==========
        dd, mm, yy = today.day, today.month, today.year
        lunar_day, lunar_month, lunar_year, is_leap = solar_to_lunar(dd, mm, yy)
        
        day_of_week = ["Thứ Hai", "Thứ Ba", "Thứ Tư", "Thứ Năm", "Thứ Sáu", "Thứ Bảy", "Chủ Nhật"][today.weekday()]
        can_chi_year = get_can_chi_year(lunar_year)
        can_chi_day = get_can_chi_day(dd, mm, yy)
        
        # Tên tháng âm
        month_name = f"{'Nhuận ' if is_leap else ''}Tháng {lunar_month}"
        
        # Ngày lễ âm lịch Việt Nam
        vn_holidays = {
            (1, 1): "🎊 Tết Nguyên Đán - Mùng 1 Tết",
            (1, 2): "🎊 Mùng 2 Tết",
            (1, 3): "🎊 Mùng 3 Tết",
            (1, 15): "🏮 Tết Nguyên Tiêu (Rằm tháng Giêng)",
            (3, 3): "🍰 Tết Hàn Thực",
            (3, 10): "👑 Giỗ Tổ Hùng Vương",
            (4, 15): "🪷 Lễ Phật Đản",
            (5, 5): "🐲 Tết Đoan Ngọ",
            (7, 15): "👻 Rằm tháng 7 - Lễ Vu Lan",
            (8, 15): "🥮 Tết Trung Thu",
            (9, 9): "🌸 Tết Trùng Cửu",
            (10, 15): "🙏 Rằm tháng 10 - Lễ Hạ Nguyên",
            (12, 23): "🧹 Ông Công Ông Táo",
            (12, 30): "🎆 Giao thừa - Đêm 30 Tết",
        }
        
        holiday_info = vn_holidays.get((lunar_month, lunar_day), "")
        
        # Kiểm tra ngày rằm / mùng 1
        special_day = ""
        if lunar_day == 1:
            special_day = "🌑 Ngày Mùng 1 (Sóc)"
        elif lunar_day == 15:
            special_day = "🌕 Ngày Rằm (Vọng)"
        
        message = f"""📅 LỊCH ÂM VIỆT NAM

🗓️ Dương lịch: {day_of_week}, {dd:02d}/{mm:02d}/{yy}
🌙 Âm lịch: Ngày {lunar_day}, {month_name}, năm {can_chi_year}

📆 Ngày: {can_chi_day}
🐉 Năm: {can_chi_year} ({lunar_year})

{f'🎉 {holiday_info}' if holiday_info else ''}
{special_day}""".strip()
        
        return {
            "success": True,
            "solar_date": f"{dd:02d}/{mm:02d}/{yy}",
            "lunar_date": f"{lunar_day}/{lunar_month}/{lunar_year}",
            "lunar_day": lunar_day,
            "lunar_month": lunar_month,
            "lunar_year": lunar_year,
            "is_leap_month": is_leap == 1,
            "day_of_week": day_of_week,
            "can_chi_day": can_chi_day,
            "can_chi_year": can_chi_year,
            "holiday": holiday_info if holiday_info else None,
            "message": message
        }
    except Exception as e:
        import traceback
        return {"success": False, "error": str(e), "traceback": traceback.format_exc()}

# ============================================================
# KNOWLEDGE BASE TOOL HANDLERS
# ============================================================

# ============================================================
# 🔥 GEMINI FLASH SMART KB FILTER - LỌC THÔNG TIN THÔNG MINH
# ============================================================

async def gemini_smart_kb_filter(
    user_query: str,
    filter_mode: str = "relevant",  # relevant, summary, extract, qa
    max_documents: int = 10,
    output_format: str = "structured"  # structured, raw, concise
) -> dict:
    """
    🔥 Sử dụng sức mạnh Gemini Flash 3 để LỌC và TÌM KIẾM THÔNG MINH trong Knowledge Base.
    
    Quy trình:
    1. Load toàn bộ documents từ Knowledge Base
    2. Dùng Gemini Flash để phân tích và lọc nội dung THỰC SỰ liên quan
    3. Trích xuất thông tin chính xác, loại bỏ noise
    4. Trả về kết quả đã được lọc sạch cho LLM chính đọc
    
    Args:
        user_query: Câu hỏi/yêu cầu của user
        filter_mode: 
            - "relevant": Chỉ giữ phần liên quan (default)
            - "summary": Tóm tắt nội dung
            - "extract": Trích xuất facts/entities
            - "qa": Trả lời câu hỏi trực tiếp
        max_documents: Số documents tối đa để xử lý (default: 10)
        output_format:
            - "structured": JSON có cấu trúc
            - "raw": Text thô
            - "concise": Ngắn gọn nhất
            
    Returns:
        dict với filtered_content, sources, và metadata
    """
    try:
        print(f"🔥 [GEMINI KB FILTER] Processing: {user_query[:60]}...")
        
        # ============================================================
        # BƯỚC 1: Load tất cả documents từ Knowledge Base
        # ============================================================
        all_documents = []
        
        # Thử load từ index trước
        if KNOWLEDGE_INDEX_FILE.exists():
            try:
                with open(KNOWLEDGE_INDEX_FILE, 'r', encoding='utf-8') as f:
                    index_data = json.load(f)
                all_documents = index_data.get("documents", [])
            except:
                pass
        
        # 🆕 FALLBACK: Nếu index trống, đọc trực tiếp từ files
        if not all_documents:
            print("⚠️ [GEMINI KB] Index trống, đang đọc trực tiếp từ files...")
            config = load_knowledge_config()
            folder_path = config.get("folder_path", "")
            
            if folder_path and Path(folder_path).exists():
                files = scan_folder_for_files(folder_path)
                for f in files[:15]:  # Giới hạn 15 files
                    try:
                        text = extract_text_from_file(f["path"])
                        if text and len(text.strip()) > 50 and not text.startswith("["):
                            all_documents.append({
                                "file_path": f["path"],
                                "file_name": f["name"],
                                "content": text[:50000]
                            })
                            print(f"📄 [GEMINI KB] Loaded: {f['name']}")
                    except Exception as e:
                        print(f"⚠️ [GEMINI KB] Error loading {f['name']}: {e}")
        
        if not all_documents:
            return {
                "success": False,
                "error": "Knowledge Base chưa có dữ liệu. Vui lòng vào Web UI > Knowledge Base để cấu hình thư mục."
            }
        
        print(f"📚 [GEMINI KB] Loaded {len(all_documents)} documents")
        
        # ============================================================
        # BƯỚC 2: Pre-filter bằng keywords (giảm số docs cần gửi Gemini)
        # ============================================================
        query_lower = user_query.lower()
        stop_words = {'là', 'của', 'và', 'có', 'các', 'được', 'trong', 'để', 'này', 'đó', 
                     'cho', 'với', 'từ', 'về', 'như', 'theo', 'không', 'khi', 'đã', 'sẽ',
                     'ai', 'gì', 'nào', 'đâu', 'sao', 'thế', 'a', 'an', 'the', 'is', 'are'}
        
        keywords = [w.lower() for w in user_query.split() if w.lower() not in stop_words and len(w) > 1]
        print(f"🔑 [GEMINI KB] Keywords: {keywords}")
        
        # Pre-filter: Chỉ giữ documents có ít nhất 1 keyword
        candidate_docs = []
        for doc in all_documents:
            content = doc.get("content", "").lower()
            file_name = doc.get("file_name", "")
            
            # Skip invalid content
            if content.strip().startswith("%pdf-") or len(content.strip()) < 50:
                continue
            
            # Check keyword match
            match_count = sum(1 for kw in keywords if kw in content or kw in file_name.lower())
            if match_count > 0 or not keywords:  # Nếu không có keywords, lấy tất cả
                candidate_docs.append({
                    "file_name": file_name,
                    "content": doc.get("content", ""),
                    "match_count": match_count
                })
        
        # Sort by match count và giới hạn
        candidate_docs.sort(key=lambda x: x["match_count"], reverse=True)
        candidate_docs = candidate_docs[:max_documents]
        
        if not candidate_docs:
            return {
                "success": False,
                "error": f"Không tìm thấy documents nào liên quan đến '{user_query}'"
            }
        
        print(f"📄 [GEMINI KB] Pre-filtered to {len(candidate_docs)} candidate docs")
        
        # ============================================================
        # BƯỚC 3: Chuẩn bị context cho Gemini Flash
        # ============================================================
        # Giới hạn mỗi document 3000 chars để tránh quá tải
        docs_for_gemini = []
        total_chars = 0
        MAX_TOTAL_CHARS = 25000  # ~6000 tokens cho Gemini
        
        for doc in candidate_docs:
            content = doc["content"]
            if len(content) > 3000:
                # Trích xuất phần có keywords
                content = _extract_relevant_parts(content, keywords, max_len=3000)
            
            if total_chars + len(content) > MAX_TOTAL_CHARS:
                break
                
            docs_for_gemini.append({
                "file_name": doc["file_name"],
                "content": content
            })
            total_chars += len(content)
        
        print(f"📦 [GEMINI KB] Prepared {len(docs_for_gemini)} docs ({total_chars:,} chars) for Gemini")
        
        # ============================================================
        # BƯỚC 4: Build prompt cho Gemini Flash
        # ============================================================
        docs_text = ""
        for i, doc in enumerate(docs_for_gemini, 1):
            docs_text += f"\n\n--- TÀI LIỆU {i}: {doc['file_name']} ---\n{doc['content']}"
        
        # Prompt tùy theo filter_mode
        if filter_mode == "summary":
            filter_instruction = """TÓM TẮT nội dung liên quan đến câu hỏi.
- Chỉ tóm tắt phần THỰC SỰ liên quan
- Bỏ qua thông tin không liên quan
- Viết ngắn gọn, súc tích"""
        elif filter_mode == "extract":
            filter_instruction = """TRÍCH XUẤT các facts, entities, số liệu liên quan:
- Tên người, tổ chức
- Số liệu, ngày tháng
- Sự kiện, hành động
- Mối quan hệ
Format: JSON array"""
        elif filter_mode == "qa":
            filter_instruction = """⚡ TRẢ LỜI NGAY LẬP TỨC câu hỏi dựa trên tài liệu.
⛔ KHÔNG ĐƯỢC hỏi lại, KHÔNG ĐƯỢC yêu cầu thêm thông tin
✅ Trả lời TRỰC TIẾP, chính xác, có trích dẫn nguồn
✅ Nếu thông tin không đầy đủ → VẪN trả lời với những gì có
✅ Nếu không có thông tin → Nói "Không tìm thấy trong database" """
        else:  # relevant
            filter_instruction = """LỌC và GIỮ LẠI CHỈ những phần THỰC SỰ LIÊN QUAN đến câu hỏi.
- Loại bỏ hoàn toàn các đoạn không liên quan
- Giữ nguyên văn các đoạn quan trọng
- Đánh dấu nguồn (tên file) cho mỗi đoạn"""
        
        gemini_prompt = f"""🔥 BẠN LÀ CHUYÊN GIA TRẢ LỜI CÂU HỎI TỪ CƠ SỞ DỮ LIỆU.

⚡ QUY TẮC BẮT BUỘC:
- TRẢ LỜI NGAY LẬP TỨC - KHÔNG HỎI LẠI
- KHÔNG yêu cầu thêm thông tin
- KHÔNG nói "bạn muốn biết gì" hoặc "bạn cần gì thêm"
- Sử dụng TOÀN BỘ thông tin có trong tài liệu để trả lời

📋 NHIỆM VỤ: {filter_instruction}

❓ CÂU HỎI CỦA USER:
"{user_query}"

📚 TÀI LIỆU TRONG DATABASE:
{docs_text}

🎯 TRẢ LỜI NGAY (không hỏi lại):"""

        # ============================================================
        # BƯỚC 5: Gọi Gemini Flash 3 để lọc
        # ============================================================
        if not GEMINI_AVAILABLE:
            return {
                "success": False,
                "error": "Gemini API không khả dụng. Vui lòng kiểm tra API key."
            }
        
        import google.generativeai as genai
        gemini_api_key = os.environ.get("GEMINI_API_KEY") or GEMINI_API_KEY
        
        if not gemini_api_key:
            return {"success": False, "error": "Thiếu Gemini API key"}
        
        genai.configure(api_key=gemini_api_key)
        
        # Sử dụng Gemini 3 Flash Preview (model mới nhất, nhanh nhất)
        model = genai.GenerativeModel('models/gemini-2.0-flash')
        
        print(f"🤖 [GEMINI KB] Calling Gemini Flash to filter...")
        
        response = model.generate_content(
            gemini_prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=0.2,  # Low temp cho accuracy
                max_output_tokens=2000,
                top_p=0.95
            )
        )
        
        if not response or not response.text:
            return {"success": False, "error": "Gemini không trả về response"}
        
        filtered_content = response.text.strip()
        print(f"✅ [GEMINI KB] Filtered content: {len(filtered_content)} chars")
        
        # ============================================================
        # BƯỚC 6: Format output
        # ============================================================
        sources = [doc["file_name"] for doc in docs_for_gemini]
        
        if output_format == "concise":
            # Cắt ngắn nếu quá dài
            if len(filtered_content) > 1500:
                filtered_content = filtered_content[:1500] + "\n[... Đã cắt ngắn ...]"
        
        result = {
            "success": True,
            "filtered_content": filtered_content,
            "sources": sources,
            "filter_mode": filter_mode,
            "documents_processed": len(docs_for_gemini),
            "total_documents": len(all_documents),
            "keywords_used": keywords,
            "original_chars": total_chars,
            "filtered_chars": len(filtered_content),
            "compression_ratio": f"{(1 - len(filtered_content)/max(total_chars,1))*100:.1f}%",
            "message": f"✅ Đã lọc {len(docs_for_gemini)} tài liệu ({total_chars:,} chars) → {len(filtered_content):,} chars relevant content"
        }
        
        # Thêm instruction cho LLM chính
        result["llm_instruction"] = f"""📊 ĐÃ LỌC THÔNG TIN TỪ KNOWLEDGE BASE

Câu hỏi: "{user_query}"
Nguồn: {', '.join(sources[:3])}{'...' if len(sources) > 3 else ''}

--- NỘI DUNG ĐÃ LỌC ---
{filtered_content}
--- HẾT ---

⚡ HÃY TRẢ LỜI USER DỰA TRÊN THÔNG TIN TRÊN."""

        return result
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}


def _extract_relevant_parts(content: str, keywords: list, max_len: int = 3000) -> str:
    """
    Trích xuất các phần có chứa keywords từ content dài.
    """
    if not keywords:
        return content[:max_len]
    
    content_lower = content.lower()
    relevant_parts = []
    
    for keyword in keywords:
        pos = 0
        while pos < len(content_lower):
            idx = content_lower.find(keyword, pos)
            if idx == -1:
                break
            
            # Lấy context xung quanh keyword (500 chars mỗi bên)
            start = max(0, idx - 500)
            end = min(len(content), idx + len(keyword) + 500)
            
            part = content[start:end]
            if part not in relevant_parts:
                relevant_parts.append(part)
            
            pos = idx + 1
            
            # Giới hạn số parts
            if len(relevant_parts) >= 5:
                break
    
    if relevant_parts:
        combined = "\n[...]\n".join(relevant_parts)
        return combined[:max_len]
    else:
        return content[:max_len]


# ============================================================
# 🔥 GEMINI SMART ANALYZE - PHÂN TÍCH + GOOGLE SEARCH
# ============================================================

async def gemini_smart_analyze(
    user_query: str,
    analysis_type: str = "comprehensive",  # comprehensive, quick, deep
    include_web_search: bool = True,
    include_kb: bool = False,
    max_search_results: int = 8
) -> dict:
    """
    🔥 GEMINI SMART ANALYZE - Phân tích vấn đề + Tìm kiếm Web + AI tổng hợp
    
    Quy trình:
    1. Gemini phân tích yêu cầu và tạo search queries tối ưu
    2. Tìm kiếm Web (Google/DuckDuckGo) để lấy thông tin mới nhất
    3. (Tùy chọn) Tìm kiếm Knowledge Base nội bộ
    4. Gemini tổng hợp, phân tích và đưa ra kết luận
    5. Trả về kết quả phân tích cho LLM chính
    
    Args:
        user_query: Vấn đề cần phân tích
        analysis_type: 
            - "comprehensive": Phân tích đầy đủ, chi tiết (default)
            - "quick": Phân tích nhanh, tóm tắt
            - "deep": Phân tích sâu, nhiều góc độ
        include_web_search: Có tìm kiếm web không (default: True)
        include_kb: Có tìm Knowledge Base không (default: False)
        max_search_results: Số kết quả web search tối đa (default: 8)
        
    Returns:
        dict với analysis, sources, summary
    """
    try:
        print(f"🔥 [GEMINI ANALYZE] Analyzing: {user_query[:60]}...")
        
        # ============================================================
        # BƯỚC 1: Kiểm tra Gemini API
        # ============================================================
        if not GEMINI_AVAILABLE:
            return {
                "success": False,
                "error": "Gemini API không khả dụng. Vui lòng kiểm tra API key."
            }
        
        import google.generativeai as genai
        gemini_api_key = os.environ.get("GEMINI_API_KEY") or GEMINI_API_KEY
        
        if not gemini_api_key:
            return {"success": False, "error": "Thiếu Gemini API key"}
        
        genai.configure(api_key=gemini_api_key)
        model = genai.GenerativeModel('models/gemini-2.0-flash')
        
        # ============================================================
        # BƯỚC 2: Gemini tạo search queries tối ưu
        # ============================================================
        query_prompt = f"""Bạn là chuyên gia phân tích. User muốn phân tích/tìm hiểu về:
"{user_query}"

Hãy tạo 2-3 search queries TỐI ƯU để tìm kiếm thông tin trên Google/Web.
Mỗi query nên:
- Ngắn gọn, từ khóa chính xác
- Thêm năm 2024/2025 nếu cần thông tin mới
- Tiếng Việt hoặc Anh tùy chủ đề

Trả về JSON array, VD: ["query 1", "query 2", "query 3"]
Chỉ trả về JSON, không giải thích."""

        print("🔍 [GEMINI ANALYZE] Generating search queries...")
        
        query_response = model.generate_content(
            query_prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=0.3,
                max_output_tokens=300
            )
        )
        
        # Parse search queries
        search_queries = [user_query]  # Default
        if query_response and query_response.text:
            try:
                import re
                json_match = re.search(r'\[.*?\]', query_response.text, re.DOTALL)
                if json_match:
                    search_queries = json.loads(json_match.group())
                    print(f"✅ [GEMINI ANALYZE] Generated queries: {search_queries}")
            except:
                search_queries = [user_query]
        
        # ============================================================
        # BƯỚC 3: Tìm kiếm Web (Google/DuckDuckGo)
        # ============================================================
        web_results = []
        web_context = ""
        
        if include_web_search and RAG_AVAILABLE:
            print(f"🌐 [GEMINI ANALYZE] Searching web with {len(search_queries)} queries...")
            
            from rag_system import web_search as rag_web_search
            
            all_results = []
            for sq in search_queries[:3]:  # Max 3 queries
                try:
                    result = await rag_web_search(sq, max_results=max_search_results // len(search_queries) + 2)
                    if result.get("success") and result.get("results"):
                        all_results.extend(result["results"])
                except Exception as e:
                    print(f"⚠️ [GEMINI ANALYZE] Search error for '{sq}': {e}")
            
            # Deduplicate by title
            seen_titles = set()
            for r in all_results:
                title = r.get("title", "")
                if title and title not in seen_titles:
                    seen_titles.add(title)
                    web_results.append(r)
            
            web_results = web_results[:max_search_results]
            print(f"📊 [GEMINI ANALYZE] Found {len(web_results)} unique web results")
            
            # Build web context
            if web_results:
                web_context = "🌐 KẾT QUẢ TÌM KIẾM WEB:\n\n"
                for i, r in enumerate(web_results, 1):
                    web_context += f"{i}. **{r.get('title', 'No title')}**\n"
                    web_context += f"   {r.get('snippet', '')}\n"
                    if r.get('url'):
                        web_context += f"   🔗 {r.get('url')}\n"
                    web_context += "\n"
        
        # ============================================================
        # BƯỚC 4: Tìm kiếm Knowledge Base (nếu bật)
        # ============================================================
        kb_context = ""
        kb_sources = []
        
        if include_kb:
            print("📚 [GEMINI ANALYZE] Searching Knowledge Base...")
            try:
                kb_result = await gemini_smart_kb_filter(
                    user_query=user_query,
                    filter_mode="relevant",
                    max_documents=5,
                    output_format="concise"
                )
                if kb_result.get("success") and kb_result.get("filtered_content"):
                    kb_context = f"\n\n📚 THÔNG TIN TỪ DATABASE NỘI BỘ:\n{kb_result['filtered_content']}"
                    kb_sources = kb_result.get("sources", [])
                    print(f"✅ [GEMINI ANALYZE] Found KB content from {len(kb_sources)} sources")
            except Exception as e:
                print(f"⚠️ [GEMINI ANALYZE] KB search error: {e}")
        
        # ============================================================
        # BƯỚC 5: Gemini tổng hợp và phân tích
        # ============================================================
        
        # Xây dựng prompt phân tích tùy theo type
        if analysis_type == "quick":
            analysis_instruction = """PHÂN TÍCH NHANH - Tóm tắt ngắn gọn:
- 3-5 điểm chính
- Kết luận trong 2-3 câu
- Không cần chi tiết"""
        elif analysis_type == "deep":
            analysis_instruction = """PHÂN TÍCH SÂU - Chi tiết và đa chiều:
- Phân tích từ nhiều góc độ
- So sánh các nguồn thông tin
- Đánh giá độ tin cậy
- Xu hướng và dự đoán
- Tóm tắt các quan điểm khác nhau"""
        else:  # comprehensive
            analysis_instruction = """PHÂN TÍCH TOÀN DIỆN:
- Tóm tắt thông tin chính
- Các điểm quan trọng
- Nguồn gốc và độ tin cậy
- Kết luận rõ ràng"""
        
        now = datetime.now()
        current_date = now.strftime("%d/%m/%Y")
        
        analysis_prompt = f"""🔥 BẠN LÀ CHUYÊN GIA PHÂN TÍCH THÔNG TIN.

📅 NGÀY HIỆN TẠI: {current_date}

⚡ NHIỆM VỤ: {analysis_instruction}

❓ VẤN ĐỀ CẦN PHÂN TÍCH:
"{user_query}"

{web_context}
{kb_context}

🎯 YÊU CẦU QUAN TRỌNG:
1. TRẢ LỜI NGẮN GỌN - TỐI ĐA 500 TỪ
2. ĐI THẲNG VÀO VẤN ĐỀ, không giải thích dài dòng
3. Liệt kê ý chính bằng bullet points
4. TRẢ LỜI BẰNG TIẾNG VIỆT
5. KHÔNG cần ghi nguồn chi tiết

📝 TRẢ LỜI NGẮN GỌN:"""

        print("🤖 [GEMINI ANALYZE] Gemini analyzing and synthesizing...")
        
        analysis_response = model.generate_content(
            analysis_prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=0.3,
                max_output_tokens=1000,
                top_p=0.9
            )
        )
        
        if not analysis_response or not analysis_response.text:
            return {"success": False, "error": "Gemini không trả về phân tích"}
        
        analysis_content = analysis_response.text.strip()
        
        # ⚡ GIỚI HẠN ĐỘ DÀI - Quá dài sẽ khiến LLM cloud bị timeout
        MAX_RESPONSE_LENGTH = 1500
        if len(analysis_content) > MAX_RESPONSE_LENGTH:
            # Cắt ngắn nhưng giữ nguyên câu cuối
            analysis_content = analysis_content[:MAX_RESPONSE_LENGTH]
            # Tìm dấu chấm cuối để không cắt giữa câu
            last_period = analysis_content.rfind('.')
            if last_period > MAX_RESPONSE_LENGTH - 200:
                analysis_content = analysis_content[:last_period + 1]
            analysis_content += "\n\n(Đây là tóm tắt. Hỏi thêm nếu cần chi tiết.)"
        
        print(f"✅ [GEMINI ANALYZE] Analysis complete: {len(analysis_content)} chars")
        
        # ============================================================
        # BƯỚC 6: Trả về kết quả - PLAIN TEXT để LLM đọc ngay
        # ============================================================
        
        # Trả về response_text để format_result_for_llm xử lý đúng
        # Giống cách ask_gemini, ask_gpt4 hoạt động
        return {
            "success": True,
            "response_text": analysis_content
        }
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}


async def search_knowledge_base(query: str) -> dict:
    """
    Tìm kiếm trong Knowledge Base và dùng Gemini AI để trả lời chính xác.
    - Bước 1: TF-IDF tìm tài liệu liên quan
    - Bước 2: Gemini đọc context và trả lời câu hỏi
    - 🆕 Bước 0: Nếu index trống, tự động đọc file trực tiếp
    """
    try:
        if not query:
            return {"success": False, "error": "Vui lòng nhập từ khóa tìm kiếm"}
        
        # Load index
        documents = []
        if KNOWLEDGE_INDEX_FILE.exists():
            try:
                with open(KNOWLEDGE_INDEX_FILE, 'r', encoding='utf-8') as f:
                    index_data = json.load(f)
                documents = index_data.get("documents", [])
            except:
                pass
        
        # 🆕 FALLBACK: Nếu index trống, tự động đọc trực tiếp từ files
        if not documents:
            print("⚠️ [KB] Index trống, đang đọc trực tiếp từ files...")
            config = load_knowledge_config()
            folder_path = config.get("folder_path", "")
            
            if folder_path and Path(folder_path).exists():
                files = scan_folder_for_files(folder_path)
                for f in files[:10]:  # Giới hạn 10 files để tránh quá tải
                    try:
                        text = extract_text_from_file(f["path"])
                        if text and len(text.strip()) > 50 and not text.startswith("["):
                            documents.append({
                                "file_path": f["path"],
                                "file_name": f["name"],
                                "content": text[:50000]
                            })
                            print(f"📄 [KB] Loaded: {f['name']} ({len(text)} chars)")
                    except Exception as e:
                        print(f"⚠️ [KB] Error loading {f['name']}: {e}")
                
                if documents:
                    print(f"📚 [KB] Loaded {len(documents)} documents from files")
            
            if not documents:
                return {
                    "success": False, 
                    "error": "Knowledge base chưa có dữ liệu. Vui lòng vào Web UI > Knowledge Base để cấu hình thư mục và index files."
                }
        
        # Tách query thành keywords (bỏ stop words phổ biến)
        stop_words = {
            # Vietnamese
            'là', 'của', 'và', 'có', 'các', 'được', 'trong', 'để', 'này', 'đó', 'cho', 'với', 
            'từ', 'về', 'như', 'theo', 'không', 'khi', 'đã', 'sẽ', 'những', 'một', 'hay', 'hoặc',
            'thì', 'mà', 'nếu', 'vì', 'bởi', 'nên', 'cũng', 'lại', 'còn', 'đây', 'kia', 'ấy',
            'ra', 'vào', 'lên', 'xuống', 'đi', 'đến', 'bằng', 'qua', 'sau', 'trước', 'trên', 'dưới',
            'nào', 'gì', 'sao', 'thế', 'rằng', 'tại', 'vậy', 'nhưng', 'tuy', 'mặc', 'dù',
            # English
            'a', 'an', 'the', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 
            'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'should', 'could', 
            'may', 'might', 'can', 'what', 'which', 'who', 'how', 'when', 'where', 'why',
            'this', 'that', 'these', 'those', 'it', 'its', 'they', 'them', 'their',
            'he', 'she', 'him', 'her', 'his', 'we', 'us', 'our', 'you', 'your',
            'of', 'to', 'in', 'on', 'at', 'by', 'for', 'with', 'about', 'as', 'from'
        }
        
        # Lọc keywords - CHỈ GIỮ TỪ QUAN TRỌNG (dài > 3 ký tự)
        keywords = [w.lower() for w in query.split() if w.lower() not in stop_words and len(w) > 3]
        
        # Nếu query quá dài (>4 từ), chỉ lấy 4 từ quan trọng nhất
        if len(keywords) > 4:
            keywords = sorted(keywords, key=len, reverse=True)[:4]
        
        if not keywords:
            all_words = [w.lower() for w in query.split() if len(w) > 2]
            keywords = sorted(all_words, key=len, reverse=True)[:3] if all_words else [query.lower()]
        
        print(f"🔍 [KB] Searching with keywords: {keywords}")
        
        # Tính điểm relevance cho từng document
        scored_docs = []
        min_keywords_match = max(1, len(keywords) - 1)
        
        for doc in documents:
            content = doc.get("content", "")
            content_lower = content.lower()
            file_name = doc.get("file_name", "")
            
            score = 0
            matched_keywords = []
            best_pos = 0
            
            for keyword in keywords:
                count = content_lower.count(keyword)
                if count > 0:
                    import math
                    score += math.log(1 + count) * 10
                    matched_keywords.append(keyword)
                    if not best_pos:
                        idx = content_lower.find(keyword)
                        if idx >= 0:
                            best_pos = idx
            
            if len(matched_keywords) < min_keywords_match:
                continue
            
            if len(matched_keywords) > 1:
                score *= (1 + len(matched_keywords) * 0.5)
            
            for keyword in keywords:
                if keyword in file_name.lower():
                    score *= 2.0
            
            if score > 0:
                scored_docs.append({
                    "file_name": file_name,
                    "score": score,
                    "matched_keywords": matched_keywords,
                    "content": content,
                    "best_pos": best_pos
                })
        
        scored_docs.sort(key=lambda x: x["score"], reverse=True)
        
        if not scored_docs:
            return {
                "success": False,
                "message": f"❌ Không tìm thấy tài liệu liên quan trong knowledge base.\n💡 Thử dùng từ khóa khác hoặc ngắn hơn."
            }
        
        # ============================================================
        # BƯỚC 2: 🔥 DÙNG GEMINI SMART FILTER ĐỂ LỌC VÀ TRẢ LỜI
        # ============================================================
        print(f"🤖 [KB] Found {len(scored_docs)} docs, using Gemini Smart Filter...")
        
        # 🔥 SỬ DỤNG gemini_smart_kb_filter để lọc thông minh
        try:
            filter_result = await gemini_smart_kb_filter(
                user_query=query,
                filter_mode="qa",  # Trả lời trực tiếp
                max_documents=min(len(scored_docs), 5),  # Tối đa 5 docs
                output_format="concise"  # Output ngắn gọn
            )
            
            if filter_result.get("success") and filter_result.get("filtered_content"):
                answer = filter_result["filtered_content"]
                sources = filter_result.get("sources", [d['file_name'] for d in scored_docs[:3]])
                
                # 🔥 FORMAT NGẮN GỌN GIỐNG WEB_SEARCH - LLM DỄ ĐỌC
                return {
                    "success": True,
                    "answer": answer,
                    "sources": sources
                }
        except Exception as filter_err:
            print(f"⚠️ [KB] Gemini Smart Filter error: {filter_err}, falling back to direct Gemini...")
        
        # ============================================================
        # FALLBACK: Dùng Gemini trực tiếp nếu Smart Filter fail
        # ============================================================
        # Lấy context từ top 2 documents (max 3000 chars mỗi doc)
        context_parts = []
        for doc in scored_docs[:2]:
            content = doc['content']
            best_pos = doc['best_pos']
            # Lấy phần xung quanh keyword match
            start = max(0, best_pos - 500)
            end = min(len(content), best_pos + 2500)
            chunk = content[start:end]
            context_parts.append(f"📄 {doc['file_name']}:\n{chunk}")
        
        context_for_gemini = "\n\n---\n\n".join(context_parts)
        
        # Gọi Gemini để trả lời
        try:
            import google.generativeai as genai
            
            gemini_api_key = os.environ.get("GEMINI_API_KEY") or GEMINI_API_KEY
            if not gemini_api_key:
                # Fallback - trả về context thô
                return {
                    "success": True,
                    "message": f"📚 Tìm thấy {len(scored_docs)} tài liệu liên quan",
                    "context": context_for_gemini[:4000]
                }
            
            genai.configure(api_key=gemini_api_key)
            model = genai.GenerativeModel('models/gemini-2.0-flash')
            
            prompt = f"""Bạn là trợ lý AI chuyên trả lời câu hỏi dựa trên tài liệu.

⚡ QUY TẮC BẮT BUỘC:
- TRẢ LỜI NGAY LẬP TỨC - KHÔNG HỎI LẠI
- KHÔNG hỏi "bạn muốn biết gì thêm?"
- KHÔNG yêu cầu thêm thông tin
- Sử dụng thông tin có trong tài liệu để trả lời

📋 TÀI LIỆU THAM KHẢO:
{context_for_gemini[:5000]}

❓ CÂU HỎI:
{query}

📝 YÊU CẦU:
1. TRẢ LỜI TRỰC TIẾP dựa trên tài liệu
2. Nếu không có thông tin → Nói "Không tìm thấy trong tài liệu"
3. Trích dẫn nguồn khi cần
4. Ngắn gọn, súc tích
5. Tiếng Việt

🎯 TRẢ LỜI NGAY:"""

            response = model.generate_content(
                prompt,
                generation_config=genai.types.GenerationConfig(
                    max_output_tokens=500,
                    temperature=0.3  # Low temp cho accurate answers
                )
            )
            
            gemini_answer = response.text.strip() if response.text else ""
            
            if gemini_answer:
                sources = [d['file_name'] for d in scored_docs[:2]]
                # 🔥 FORMAT NGẮN GỌN GIỐNG WEB_SEARCH
                return {
                    "success": True,
                    "answer": gemini_answer,
                    "sources": sources
                }
            else:
                return {
                    "success": True,
                    "answer": f"Tìm thấy {len(scored_docs)} tài liệu liên quan nhưng không có câu trả lời cụ thể.",
                    "context": context_for_gemini[:2000]
                }
                
        except Exception as gemini_err:
            print(f"⚠️ [KB] Gemini error: {gemini_err}")
            # Fallback - trả về context thô
            return {
                "success": True,
                "answer": f"Tìm thấy {len(scored_docs)} tài liệu liên quan.",
                "context": context_for_gemini[:2500]
            }
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}

async def get_knowledge_context(query: str = "", max_chars: int = 10000, use_gemini_summary: bool = True, use_gemini_filter: bool = False) -> dict:
    """
    🔧 REFACTORED: Lấy context từ Knowledge Base với semantic search chính xác hơn.
    - Ưu tiên exact phrase match
    - Chỉ lấy documents thực sự liên quan
    - Option: Dùng Gemini Smart Filter để lọc thông minh
    - Trả về context đúng cho LLM
    
    Args:
        use_gemini_filter: Nếu True, sẽ dùng gemini_smart_kb_filter để lọc thông minh (mặc định: False)
    """
    try:
        # Load index
        if not KNOWLEDGE_INDEX_FILE.exists():
            return {
                "success": False, 
                "context": "",
                "error": "Knowledge base chưa có dữ liệu. Vui lòng index files trước."
            }
        
        with open(KNOWLEDGE_INDEX_FILE, 'r', encoding='utf-8') as f:
            index_data = json.load(f)
        
        all_documents = index_data.get("documents", [])
        if not all_documents:
            return {"success": False, "context": "", "error": "Knowledge base trống."}
        
        print(f"📚 [KB] Loaded {len(all_documents)} documents from index")
        
        # ============================================================
        # 🔥 OPTION: Sử dụng Gemini Smart Filter nếu được bật
        # ============================================================
        if use_gemini_filter and query:
            print(f"🔥 [KB] Using Gemini Smart Filter for query: {query}")
            try:
                filter_result = await gemini_smart_kb_filter(
                    user_query=query,
                    filter_mode="relevant",  # Chỉ lấy phần liên quan
                    max_documents=10,
                    output_format="structured"
                )
                
                if filter_result.get("success") and filter_result.get("filtered_content"):
                    return {
                        "success": True,
                        "context": filter_result.get("llm_instruction", filter_result["filtered_content"]),
                        "raw_context": filter_result["filtered_content"],
                        "total_documents": filter_result.get("total_documents", len(all_documents)),
                        "documents_included": filter_result.get("documents_processed", 0),
                        "context_length": filter_result.get("filtered_chars", 0),
                        "keywords_used": filter_result.get("keywords_used", []),
                        "gemini_filter_used": True,
                        "compression_ratio": filter_result.get("compression_ratio", "N/A"),
                        "message": f"✅ Gemini Smart Filter: Đã lọc {filter_result.get('documents_processed', 0)} tài liệu ({filter_result.get('filtered_chars', 0):,} chars)"
                    }
            except Exception as filter_err:
                print(f"⚠️ [KB] Gemini Smart Filter failed: {filter_err}, using traditional method...")
        
        # ============================================================
        # BƯỚC 1: Chuẩn bị keywords và query
        # ============================================================
        query_lower = query.lower().strip() if query else ""
        
        # Tạo keywords từ query
        stop_words = {'là', 'của', 'và', 'có', 'các', 'được', 'trong', 'để', 'này', 'đó', 'cho', 'với', 
                     'từ', 'về', 'như', 'theo', 'không', 'khi', 'đã', 'sẽ', 'ai', 'gì', 'nào', 'đâu',
                     'a', 'an', 'the', 'is', 'are', 'was', 'were', 'be', 'been', 'what', 'who', 'where'}
        keywords = [w.lower() for w in query.split() if w.lower() not in stop_words and len(w) > 1] if query else []
        
        # Nếu không có keywords, dùng toàn bộ query
        if not keywords and query:
            keywords = [query_lower]
        
        print(f"🔑 [KB] Query: '{query}' → Keywords: {keywords}")
        
        # ============================================================
        # BƯỚC 2: Lọc và score documents
        # ============================================================
        scored_documents = []
        
        for doc in all_documents:
            content = doc.get("content", "")
            file_name = doc.get("file_name", "unknown")
            content_lower = content.lower()
            file_name_lower = file_name.lower()
            
            # ⚠️ SKIP: PDF structure hoặc content quá ngắn
            if content.strip().startswith("%PDF-") or content.strip().startswith("<</"):
                continue
            if len(content.strip()) < 50:
                continue
            
            # Tính điểm relevance với scoring mới
            score = 0
            match_reasons = []
            has_exact_match = False
            has_filename_match = False
            
            if query_lower:
                # 0️⃣ FILENAME MATCH (ƯU TIÊN CAO) - Check trước!
                # Normalize filename để so sánh (bỏ dấu, bỏ ký tự đặc biệt)
                import unicodedata
                def normalize_text(text):
                    # Bỏ dấu tiếng Việt và chuyển thành ASCII
                    nfkd = unicodedata.normalize('NFKD', text.lower())
                    return ''.join(c for c in nfkd if not unicodedata.combining(c))
                
                query_normalized = normalize_text(query_lower)
                filename_normalized = normalize_text(file_name_lower)
                
                # Full query match trong filename
                if query_lower in file_name_lower or query_normalized in filename_normalized:
                    score += 5000
                    match_reasons.append("filename_exact")
                    has_filename_match = True
                else:
                    # Partial keyword match trong filename
                    filename_kw_matches = 0
                    for kw in keywords:
                        kw_norm = normalize_text(kw)
                        if kw in file_name_lower or kw_norm in filename_normalized:
                            filename_kw_matches += 1
                    
                    if filename_kw_matches >= 2:
                        score += 2000 * filename_kw_matches
                        match_reasons.append(f"filename_partial:{filename_kw_matches}")
                        has_filename_match = True
                    elif filename_kw_matches == 1 and len(keywords) <= 2:
                        score += 500
                        match_reasons.append(f"filename_partial:{filename_kw_matches}")
                        has_filename_match = True
                
                # 1️⃣ EXACT PHRASE MATCH (ưu tiên CAO NHẤT - ví dụ: "Lê Trung Khoa" as a phrase)
                exact_count = content_lower.count(query_lower)
                if exact_count > 0:
                    score += 5000 * exact_count  # RẤT CAO - ưu tiên tuyệt đối
                    match_reasons.append(f"exact_phrase:{exact_count}")
                    has_exact_match = True
                
                # 2️⃣ PROXIMITY CHECK - Kiểm tra keywords có gần nhau không (cho tên riêng)
                # Nếu query có vẻ là tên người (>= 2 từ), kiểm tra xem các từ có liền nhau không
                has_proximity = False
                if len(keywords) >= 2 and not has_exact_match:
                    # Tìm vị trí của mỗi keyword
                    keyword_positions = []
                    for kw in keywords:
                        pos = content_lower.find(kw)
                        if pos >= 0:
                            keyword_positions.append((kw, pos))
                    
                    # Kiểm tra proximity (trong vòng 50 ký tự)
                    if len(keyword_positions) == len(keywords):
                        # Tất cả keywords đều có trong content
                        positions = [p[1] for p in keyword_positions]
                        min_pos, max_pos = min(positions), max(positions)
                        # Nếu tất cả keywords nằm trong 50 ký tự → có thể là tên riêng
                        if max_pos - min_pos < 50:
                            has_proximity = True
                            score += 3000  # Bonus cao cho proximity
                            match_reasons.append(f"proximity:{max_pos - min_pos}chars")
                
                # 3️⃣ KEYWORD MATCH - Đếm số keywords xuất hiện
                keyword_matches = 0
                total_kw_score = 0
                for kw in keywords:
                    kw_count = content_lower.count(kw)
                    if kw_count > 0:
                        total_kw_score += min(kw_count, 5)  # Cap tại 5 lần mỗi keyword
                        keyword_matches += 1
                
                # ⚠️ NẾU LÀ TÊN RIÊNG (>= 2 keywords): Cần có exact match hoặc proximity
                if len(keywords) >= 2:
                    if has_exact_match or has_proximity:
                        # Có exact hoặc proximity → bonus cao
                        score += 200 * keyword_matches
                        match_reasons.append(f"name_match:{keyword_matches}/{len(keywords)}")
                    elif has_filename_match:
                        # Có filename match → bonus trung bình
                        score += 100 * keyword_matches
                        match_reasons.append(f"content_support:{keyword_matches}/{len(keywords)}")
                    elif keyword_matches == len(keywords):
                        # Tất cả keywords match nhưng KHÔNG gần nhau → score thấp
                        score += 20 * keyword_matches  # Thấp hơn nhiều
                        match_reasons.append(f"scattered_kw:{keyword_matches}/{len(keywords)}")
                        
                        # ⚠️ PENALTY MẠNH cho documents dài với scattered keywords
                        # NHƯNG không penalty nếu có filename match
                        if len(content) > 5000 and not has_filename_match:
                            score = int(score * 0.1)  # Giảm 90%!
                            match_reasons.append("penalty:scattered_in_long_doc")
                    elif keyword_matches >= len(keywords) * 0.7:
                        # >= 70% keywords match → score rất thấp
                        score += 10 * keyword_matches
                        match_reasons.append(f"partial_kw:{keyword_matches}/{len(keywords)}")
                    else:
                        # < 70% keywords → REJECT (trừ khi có filename match)
                        if not has_exact_match and not has_filename_match:
                            continue
                else:
                    # Single keyword → score thấp hơn
                    if keyword_matches > 0:
                        score += 30 * total_kw_score
                        match_reasons.append(f"single_kw:{total_kw_score}")
                
                # ⚠️ REJECT: Không có match nào ý nghĩa
                if score == 0:
                    continue
            else:
                # Không có query → lấy tất cả (với score dựa trên độ dài content)
                score = min(len(content), 5000)  # Cap score
                match_reasons.append("no_query")
            
            scored_documents.append({
                "doc": doc,
                "score": score,
                "reasons": match_reasons,
                "content_len": len(content)
            })
        
        # Sort by score
        scored_documents.sort(key=lambda x: x["score"], reverse=True)
        
        print(f"📊 [KB] Scored {len(scored_documents)} relevant documents")
        
        # ============================================================
        # BƯỚC 3: Filter - chỉ lấy top documents có score cao
        # ============================================================
        if scored_documents and query:
            top_score = scored_documents[0]["score"]
            # Chỉ lấy documents có score >= 30% top score (hoặc tối thiểu 50 điểm)
            min_threshold = max(50, top_score * 0.3)
            filtered_docs = [d for d in scored_documents if d["score"] >= min_threshold]
            
            # Giới hạn tối đa 5 documents để tránh quá tải
            filtered_docs = filtered_docs[:5]
            
            print(f"🎯 [KB] Filtered to {len(filtered_docs)} docs (threshold: {min_threshold:.0f})")
            for i, d in enumerate(filtered_docs[:3]):
                print(f"   {i+1}. {d['doc']['file_name']}: score={d['score']:.0f} ({', '.join(d['reasons'])})")
        else:
            filtered_docs = scored_documents[:3]  # Lấy tối đa 3 docs nếu không có query
        
        if not filtered_docs:
            return {
                "success": False,
                "context": "",
                "error": f"Không tìm thấy tài liệu nào liên quan đến '{query}'"
            }
        
        # ============================================================
        # BƯỚC 4: Loại bỏ nội dung trùng lặp (Deduplication)
        # ============================================================
        seen_content_hashes = set()
        unique_docs = []
        
        for item in filtered_docs:
            content = item["doc"].get("content", "").strip()
            # Tạo hash từ 500 ký tự đầu (đủ để detect duplicate)
            content_preview = content[:500].lower().replace(" ", "").replace("\n", "")
            
            if content_preview in seen_content_hashes:
                print(f"   ⚠️ SKIP duplicate: {item['doc']['file_name']}")
                continue
            
            seen_content_hashes.add(content_preview)
            unique_docs.append(item)
        
        if len(unique_docs) < len(filtered_docs):
            print(f"🔄 [KB] Deduplicated: {len(filtered_docs)} → {len(unique_docs)} unique docs")
        
        # ============================================================
        # BƯỚC 5: Trích xuất relevant content từ mỗi document
        # ============================================================
        context_parts = []
        total_chars = 0
        
        for item in unique_docs:
            doc = item["doc"]
            content = doc.get("content", "")
            file_name = doc.get("file_name", "unknown")
            
            # Trích xuất phần content liên quan nhất (không phải toàn bộ)
            if query_lower and len(content) > 1500:
                # Tìm vị trí query/keyword xuất hiện và lấy context xung quanh
                best_section = extract_relevant_section(content, query_lower, keywords, max_section_len=2000)
                content = best_section
            elif len(content) > 2500:
                # Không có query → cắt ngắn
                content = content[:2500] + "\n[... Nội dung tiếp bị cắt ...]"
            
            # Build context entry
            header = f"\n\n{'='*50}\n📄 {file_name} (score: {item['score']:.0f})\n{'='*50}\n"
            entry = header + content
            
            # Kiểm tra giới hạn tổng chars
            if total_chars + len(entry) > max_chars:
                remaining = max_chars - total_chars
                if remaining > 500:
                    context_parts.append(header + content[:remaining-len(header)] + "\n[... Cắt do quá dài ...]")
                break
            
            context_parts.append(entry)
            total_chars += len(entry)
        
        full_context = "".join(context_parts)
        
        # ============================================================
        # BƯỚC 6: Format response cho LLM dễ hiểu
        # ============================================================
        # Tạo instruction rõ ràng cho LLM
        instruction = f"""📚 ĐÃ TÌM THẤY {len(context_parts)} TÀI LIỆU LIÊN QUAN ĐẾN "{query}"

⚡ HƯỚNG DẪN CHO AI:
1. ĐỌC KỸ NỘI DUNG BÊN DƯỚI
2. TRẢ LỜI CÂU HỎI DỰA TRÊN NỘI DUNG NÀY
3. TRÍCH DẪN THÔNG TIN TỪ TÀI LIỆU
4. NẾU KHÔNG ĐỦ THÔNG TIN, HÃY NÓI RÕ

---NỘI DUNG TÀI LIỆU---
{full_context}
---HẾT NỘI DUNG---

💡 HÃY TRẢ LỜI CÂU HỎI CỦA USER DỰA TRÊN THÔNG TIN TRÊN."""

        # 🔄 TRUNCATE: Giới hạn context dưới 4000 ký tự cho LLM
        if len(instruction) > MAX_LLM_RESPONSE_CHARS:
            original_len = len(instruction)
            instruction = smart_truncate_for_llm(instruction, MAX_LLM_RESPONSE_CHARS)
            print(f"[KB] ✂️ Truncated context: {original_len} → {len(instruction)} chars")

        return {
            "success": True,
            "context": instruction,  # Instruction + context (đã truncate)
            "raw_context": full_context,  # Context thuần
            "total_documents": len(all_documents),
            "documents_included": len(context_parts),
            "duplicates_removed": len(filtered_docs) - len(unique_docs),
            "context_length": len(full_context),
            "keywords_used": keywords,
            "gemini_summarization": False,
            "message": f"✅ Tìm thấy {len(context_parts)} tài liệu ({len(full_context):,} chars). ĐỌC CONTEXT VÀ TRẢ LỜI USER!"
        }
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "context": "", "error": str(e)}


def extract_relevant_section(content: str, query: str, keywords: list, max_section_len: int = 2000) -> str:
    """
    Trích xuất phần content liên quan nhất đến query.
    Tìm vị trí query/keywords xuất hiện và lấy context xung quanh.
    """
    content_lower = content.lower()
    
    # Tìm vị trí exact query match
    pos = content_lower.find(query)
    
    if pos == -1 and keywords:
        # Không tìm thấy exact match, tìm keyword đầu tiên
        for kw in keywords:
            pos = content_lower.find(kw)
            if pos != -1:
                break
    
    if pos == -1:
        # Không tìm thấy gì, trả về đầu document
        return content[:max_section_len] + ("\n[... Còn tiếp ...]" if len(content) > max_section_len else "")
    
    # Lấy context xung quanh vị trí tìm thấy
    half_len = max_section_len // 2
    start = max(0, pos - half_len)
    end = min(len(content), pos + half_len)
    
    # Điều chỉnh để không cắt giữa từ
    if start > 0:
        # Tìm space gần nhất để bắt đầu
        space_pos = content.rfind(' ', max(0, start - 50), start + 50)
        if space_pos > 0:
            start = space_pos + 1
    
    if end < len(content):
        # Tìm space gần nhất để kết thúc
        space_pos = content.find(' ', end - 50, end + 50)
        if space_pos > 0:
            end = space_pos
    
    section = content[start:end]
    
    # Thêm markers nếu bị cắt
    prefix = "[...] " if start > 0 else ""
    suffix = " [...]" if end < len(content) else ""
    
    return prefix + section + suffix


# =====================================================
# 📖 DOC READER GEMINI RAG - ADVANCED RAG SYSTEM
# =====================================================

async def doc_reader_gemini_rag(
    user_query: str,
    knowledge_base_path: str = None,
    chunk_size: int = 1024,
    top_k: int = 5,
    use_vector_search: bool = True
) -> dict:
    """
    📖 Hệ thống RAG nâng cao với Gemini:
    1. Load và chunk documents
    2. Embed và vector search (semantic search)
    3. Format context và generate response với Gemini
    
    Args:
        user_query: Câu hỏi của người dùng
        knowledge_base_path: Đường dẫn thư mục KB (mặc định dùng config)
        chunk_size: Kích thước mỗi chunk (default: 1024 chars)
        top_k: Số lượng chunks liên quan nhất (default: 5)
        use_vector_search: Dùng semantic search hay keyword search (default: True)
        
    Returns:
        dict với success, response_text, sources, và debug info
    """
    try:
        print(f"📖 [RAG] Processing query: {user_query[:50]}...")
        
        # BƯỚC 1: Load documents từ Knowledge Base
        if not knowledge_base_path:
            config = load_knowledge_config()
            knowledge_base_path = config.get("folder_path", "")
        
        if not knowledge_base_path or not Path(knowledge_base_path).exists():
            return {
                "success": False,
                "error": "Knowledge base path không hợp lệ. Vui lòng cấu hình thư mục KB."
            }
        
        # Load index
        index_data = load_knowledge_index()
        documents = index_data.get("documents", [])
        
        if not documents:
            return {
                "success": False,
                "error": "Knowledge base trống. Vui lòng index các files trước."
            }
        
        print(f"📚 [RAG] Loaded {len(documents)} documents")
        
        # BƯỚC 2: Chunk documents (chia nhỏ tài liệu)
        all_chunks = []
        for doc in documents:
            content = doc.get("content", "")
            file_name = doc.get("file_name", "unknown")
            
            # Skip PDF structure
            if content.strip().startswith("%PDF-") or content.strip().startswith("<</"):
                continue
            
            # Chunk document
            chunks = []
            for i in range(0, len(content), chunk_size):
                chunk_text = content[i:i+chunk_size]
                if len(chunk_text.strip()) > 50:  # Skip very short chunks
                    chunks.append({
                        "text": chunk_text,
                        "file_name": file_name,
                        "chunk_index": i // chunk_size,
                        "source_doc": doc
                    })
            
            all_chunks.extend(chunks)
        
        print(f"✂️ [RAG] Created {len(all_chunks)} chunks")
        
        # BƯỚC 3: Tìm kiếm chunks liên quan
        if use_vector_search:
            # Vector/Semantic Search (simple TF-IDF based)
            relevant_chunks = _semantic_search_chunks(user_query, all_chunks, top_k)
        else:
            # Keyword search (fallback)
            relevant_chunks = _keyword_search_chunks(user_query, all_chunks, top_k)
        
        if not relevant_chunks:
            return {
                "success": False,
                "error": f"Không tìm thấy thông tin liên quan đến '{user_query}' trong Knowledge Base."
            }
        
        print(f"🔍 [RAG] Found {len(relevant_chunks)} relevant chunks")
        
        # BƯỚC 4: Format context từ relevant chunks
        prompt_context = ""
        sources = []
        for i, chunk in enumerate(relevant_chunks, 1):
            prompt_context += f"\n--- Đoạn {i} (từ {chunk['file_name']}) ---\n"
            prompt_context += chunk['text'][:800] + "\n"  # Limit each chunk
            
            if chunk['file_name'] not in sources:
                sources.append(chunk['file_name'])
        
        # BƯỚC 5: Xây dựng prompt cho Gemini
        final_prompt = f"""Bạn là trợ lý thông minh có quyền truy cập Knowledge Base của người dùng.

⚡ QUY TẮC BẮT BUỘC:
- TRẢ LỜI NGAY LẬP TỨC - KHÔNG HỎI LẠI
- KHÔNG hỏi "bạn muốn biết gì thêm?"
- KHÔNG yêu cầu thêm thông tin
- Sử dụng toàn bộ thông tin có trong Knowledge Base để trả lời

📚 THÔNG TIN TỪ KNOWLEDGE BASE:
{prompt_context}

❓ CÂU HỎI:
{user_query}

📝 YÊU CẦU:
- TRẢ LỜI TRỰC TIẾP dựa trên Knowledge Base
- Nếu không đủ thông tin → Nói "Không tìm thấy trong Knowledge Base"
- Trích dẫn tên file khi cần
- Ngắn gọn, chính xác

🎯 TRẢ LỜI NGAY:"""
        
        # BƯỚC 6: Gọi Gemini API
        print(f"🤖 [RAG] Calling Gemini...")
        
        if not GEMINI_AVAILABLE:
            return {
                "success": False,
                "error": "Gemini API không khả dụng. Vui lòng kiểm tra API key."
            }
        
        import google.generativeai as genai
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel('models/gemini-3-flash-preview')
        
        response = model.generate_content(
            final_prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=0.3,  # Focused and factual
                max_output_tokens=1000
            )
        )
        
        if not response or not response.text:
            return {
                "success": False,
                "error": "Gemini không trả về response."
            }
        
        print(f"✅ [RAG] Generated response ({len(response.text)} chars)")
        
        # Return full result
        return {
            "success": True,
            "response_text": response.text.strip(),
            "query": user_query,
            "sources": sources,
            "chunks_used": len(relevant_chunks),
            "total_chunks": len(all_chunks),
            "search_method": "semantic" if use_vector_search else "keyword",
            "message": f"✅ Đã trả lời dựa trên {len(relevant_chunks)} đoạn từ {len(sources)} tài liệu"
        }
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}


def _semantic_search_chunks(query: str, chunks: list, top_k: int = 5) -> list:
    """
    Tìm kiếm semantic dựa trên TF-IDF scoring
    """
    import math
    
    # Extract keywords from query
    stop_words = {'là', 'của', 'và', 'có', 'các', 'được', 'trong', 'để', 'này', 'đó', 
                  'cho', 'với', 'từ', 'về', 'như', 'theo', 'không', 'khi', 'đã', 'sẽ',
                  'a', 'an', 'the', 'is', 'are', 'was', 'were', 'be', 'been', 'being'}
    
    keywords = [w.lower() for w in query.split() if w.lower() not in stop_words and len(w) > 2]
    
    if not keywords:
        keywords = [query.lower()]
    
    # Score each chunk
    scored_chunks = []
    for chunk in chunks:
        text_lower = chunk['text'].lower()
        score = 0
        
        for keyword in keywords:
            count = text_lower.count(keyword)
            if count > 0:
                # TF-IDF inspired scoring
                score += math.log(1 + count) * 10
        
        # Multi-keyword bonus
        matched = sum(1 for kw in keywords if kw in text_lower)
        if matched > 1:
            score *= (1 + matched * 0.3)
        
        if score > 0:
            chunk['score'] = score
            scored_chunks.append(chunk)
    
    # Sort by score and return top K
    scored_chunks.sort(key=lambda x: x['score'], reverse=True)
    return scored_chunks[:top_k]


def _keyword_search_chunks(query: str, chunks: list, top_k: int = 5) -> list:
    """
    Tìm kiếm đơn giản dựa trên keyword matching
    """
    query_lower = query.lower()
    matched_chunks = []
    
    for chunk in chunks:
        if query_lower in chunk['text'].lower():
            matched_chunks.append(chunk)
            if len(matched_chunks) >= top_k:
                break
    
    return matched_chunks


async def send_to_wechat(contact: str, message: str) -> dict:
    """
    Gửi tin nhắn đến Zalo/Messenger (mở app và paste tin nhắn).
    Lưu ý: Cần có Zalo PC đang chạy.
    """
    try:
        import pyautogui
        import pyperclip
        import time
        import subprocess
        
        # Copy message vào clipboard
        pyperclip.copy(message)
        
        # Thử mở Zalo
        try:
            subprocess.Popen(["start", "zalo:"], shell=True)
            time.sleep(2)
        except:
            pass
        
        # Ctrl+F để tìm kiếm
        pyautogui.hotkey('ctrl', 'f')
        time.sleep(0.5)
        
        # Gõ tên contact
        pyautogui.typewrite(contact, interval=0.05)
        time.sleep(1)
        
        # Enter để chọn
        pyautogui.press('enter')
        time.sleep(0.5)
        
        # Ctrl+V để paste tin nhắn
        pyautogui.hotkey('ctrl', 'v')
        
        return {
            "success": True,
            "message": f"📱 Đã mở chat với '{contact}' và paste tin nhắn. Nhấn Enter để gửi.",
            "hint": "Tin nhắn đã được paste, bạn cần nhấn Enter để gửi"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

async def get_fuel_price_vietnam() -> dict:
    """
    Lấy giá xăng dầu Việt Nam.
    """
    try:
        import aiohttp
        
        # Giá xăng tham khảo (cập nhật manual hoặc từ API nếu có)
        # Thực tế cần API từ Petrolimex hoặc nguồn chính thống
        return {
            "success": True,
            "message": """⛽ Giá xăng dầu Việt Nam (tham khảo):
            
• RON 95-V: ~24,000 - 25,000 VNĐ/lít
• RON 95-III: ~23,000 - 24,000 VNĐ/lít  
• E5 RON 92: ~22,000 - 23,000 VNĐ/lít
• Dầu DO 0.05S: ~20,000 - 21,000 VNĐ/lít

💡 Giá có thể thay đổi theo kỳ điều chỉnh (15 ngày/lần)
📍 Xem giá chính xác: petrolimex.com.vn""",
            "hint": "Giá tham khảo, vui lòng kiểm tra nguồn chính thống"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

async def lock_computer() -> dict:
    """
    Khóa màn hình máy tính (Win+L).
    """
    try:
        import ctypes
        ctypes.windll.user32.LockWorkStation()
        return {"success": True, "message": "🔒 Đã khóa máy tính"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def shutdown_computer(minutes: int = 0) -> dict:
    """
    Tắt máy tính sau X phút (mặc định tắt ngay).
    """
    try:
        import subprocess
        
        if minutes > 0:
            seconds = minutes * 60
            subprocess.run(["shutdown", "/s", "/t", str(seconds)], check=True)
            return {"success": True, "message": f"⏰ Máy tính sẽ tắt sau {minutes} phút"}
        else:
            subprocess.run(["shutdown", "/s", "/t", "30"], check=True)
            return {"success": True, "message": "🔌 Máy tính sẽ tắt sau 30 giây"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def cancel_shutdown() -> dict:
    """
    Hủy lệnh tắt máy đã đặt.
    """
    try:
        import subprocess
        subprocess.run(["shutdown", "/a"], check=True)
        return {"success": True, "message": "✅ Đã hủy lệnh tắt máy"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def restart_computer(minutes: int = 0) -> dict:
    """
    Khởi động lại máy tính sau X phút.
    """
    try:
        import subprocess
        
        if minutes > 0:
            seconds = minutes * 60
            subprocess.run(["shutdown", "/r", "/t", str(seconds)], check=True)
            return {"success": True, "message": f"🔄 Máy tính sẽ khởi động lại sau {minutes} phút"}
        else:
            subprocess.run(["shutdown", "/r", "/t", "30"], check=True)
            return {"success": True, "message": "🔄 Máy tính sẽ khởi động lại sau 30 giây"}
    except Exception as e:
        return {"success": False, "error": str(e)}

async def set_dark_mode(enable: bool = True) -> dict:
    """
    Bật/tắt Dark Mode Windows.
    """
    try:
        import winreg
        
        key_path = r"SOFTWARE\Microsoft\Windows\CurrentVersion\Themes\Personalize"
        value = 0 if enable else 1  # 0 = Dark, 1 = Light
        
        # Set Apps theme
        with winreg.OpenKey(winreg.HKEY_CURRENT_USER, key_path, 0, winreg.KEY_SET_VALUE) as key:
            winreg.SetValueEx(key, "AppsUseLightTheme", 0, winreg.REG_DWORD, value)
            winreg.SetValueEx(key, "SystemUsesLightTheme", 0, winreg.REG_DWORD, value)
        
        mode = "Dark Mode 🌙" if enable else "Light Mode ☀️"
        return {"success": True, "message": f"✅ Đã chuyển sang {mode}"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ============================================================
# � NETWORK/FIREWALL CHECK TOOLS - Kiểm tra quyền kết nối mạng
# ============================================================

async def check_network_permission() -> dict:
    """
    Kiểm tra quyền kết nối mạng (Windows Firewall) và trạng thái Internet.
    Hướng dẫn người dùng cấp quyền nếu chưa có.
    """
    try:
        # Check firewall rules
        firewall = FirewallChecker.check_firewall_rules()
        
        # Check internet connection
        internet = FirewallChecker.check_internet_connection()
        
        # Build response
        result = {
            "success": True,
            "firewall": {
                "has_permission": bool(firewall['rules_found']),
                "rules_found": firewall['rules_found'],
                "exe_name": firewall['exe_name'],
                "exe_path": firewall['exe_path']
            },
            "internet": {
                "connected": internet['connected'],
                "latency_ms": internet.get('latency_ms')
            }
        }
        
        # Status message
        if firewall['rules_found'] and internet['connected']:
            result["message"] = f"✅ Đã có quyền Firewall và kết nối Internet ({internet.get('latency_ms', '?')}ms)"
            result["status"] = "ready"
        elif firewall['rules_found'] and not internet['connected']:
            result["message"] = "⚠️ Có quyền Firewall nhưng không có Internet. Kiểm tra kết nối mạng của máy tính."
            result["status"] = "no_internet"
        elif not firewall['rules_found'] and internet['connected']:
            result["message"] = "⚠️ Chưa thấy rule Firewall nhưng Internet vẫn hoạt động. Có thể Windows đã tự động cho phép."
            result["status"] = "working"
        else:
            result["message"] = "❌ Chưa có quyền Firewall và không kết nối được Internet."
            result["status"] = "blocked"
            result["guide"] = {
                "step1": "Khi Windows hỏi 'Allow access' → Nhấn 'Allow access'",
                "step2": "Hoặc vào Windows Security → Firewall → Allow an app",
                "step3": "Thêm file EXE vào danh sách cho phép",
                "step4": "Tick cả 'Private' và 'Public' networks"
            }
        
        return result
        
    except Exception as e:
        return {"success": False, "error": str(e)}


async def request_firewall_permission() -> dict:
    """
    Yêu cầu cấp quyền Firewall cho ứng dụng (cần quyền Admin).
    """
    try:
        success = FirewallChecker.request_firewall_permission()
        
        if success:
            return {
                "success": True,
                "message": "✅ Đã thêm rule Firewall thành công! Ứng dụng có thể kết nối Internet."
            }
        else:
            return {
                "success": False,
                "message": "⚠️ Không thể tự động thêm rule. Cần chạy với quyền Administrator.",
                "guide": {
                    "manual": "Vào Windows Security → Firewall → Allow an app → Thêm miniZ MCP",
                    "powershell": f'netsh advfirewall firewall add rule name="miniZ_MCP" dir=in action=allow program="{FirewallChecker.get_exe_path()}" enable=yes'
                }
            }
    except Exception as e:
        return {"success": False, "error": str(e)}


async def check_internet_connection() -> dict:
    """
    Kiểm tra kết nối Internet và độ trễ mạng.
    """
    try:
        result = FirewallChecker.check_internet_connection()
        
        if result['connected']:
            return {
                "success": True,
                "connected": True,
                "latency_ms": result.get('latency_ms'),
                "message": f"✅ Đã kết nối Internet (độ trễ: {result.get('latency_ms', '?')}ms)"
            }
        else:
            return {
                "success": True,
                "connected": False,
                "message": "❌ Không kết nối được Internet. Kiểm tra kết nối mạng của máy tính.",
                "error": result.get('error')
            }
    except Exception as e:
        return {"success": False, "error": str(e)}


# ============================================================
# �📨 SEND MESSAGE TO LLM - Gửi tin nhắn cho LLM tự trả lời
# ============================================================

async def send_message_to_llm(message: str, device_index: int = None, wait_response: bool = False, timeout: int = 30) -> dict:
    """
    Gửi tin nhắn cho LLM qua WebSocket. Robot sẽ đọc và tự động trả lời qua giọng nói.
    
    LƯU Ý: Do WebSocket đang được sử dụng bởi main loop, không thể đợi response trực tiếp.
    Robot sẽ nhận tin nhắn và tự động phản hồi qua voice.
    
    Args:
        message: Tin nhắn/câu hỏi muốn gửi cho LLM
        device_index: Index thiết bị (0, 1, 2). None = thiết bị đang active
        wait_response: KHÔNG SỬ DỤNG - để tương thích API cũ
        timeout: KHÔNG SỬ DỤNG - để tương thích API cũ
        
    Returns:
        dict với success, message, device_name
    """
    global xiaozhi_connections, xiaozhi_connected, active_endpoint_index, endpoints_config
    
    try:
        # Xác định device index
        if device_index is None:
            device_index = active_endpoint_index
        
        # Validate device_index
        if device_index not in [0, 1, 2]:
            return {
                "success": False,
                "error": f"Invalid device_index: {device_index}. Must be 0, 1, or 2."
            }
        
        # Kiểm tra kết nối WebSocket
        if not xiaozhi_connected.get(device_index, False):
            return {
                "success": False,
                "error": f"Thiết bị {device_index + 1} chưa kết nối. Vui lòng kiểm tra kết nối WebSocket."
            }
        
        ws = xiaozhi_connections.get(device_index)
        if ws is None:
            return {
                "success": False,
                "error": f"WebSocket connection cho thiết bị {device_index + 1} không khả dụng."
            }
        
        # Lấy tên thiết bị
        device_name = endpoints_config[device_index].get("name", f"Thiết bị {device_index + 1}")
        
        # Tạo JSON-RPC notification để gửi tin nhắn cho LLM
        # Sử dụng method "notifications/message" theo MCP protocol
        # Đây là notification (không có id) nên server không cần response
        llm_message = {
            "jsonrpc": "2.0",
            "method": "notifications/message",
            "params": {
                "level": "info",
                "data": {
                    "type": "user_message",
                    "content": message,
                    "timestamp": datetime.now().isoformat(),
                    "source": "miniZ_MCP_WebUI"
                }
            }
        }
        
        print(f"📨 [LLM Send] Sending to {device_name}: {message[:100]}...")
        
        # Lưu vào conversation history
        add_to_conversation(
            role="user",
            content=message,
            metadata={
                "source": "send_message_to_llm",
                "device": device_name,
                "device_index": device_index
            }
        )
        
        # Gửi message qua WebSocket (không đợi response)
        await ws.send(json.dumps(llm_message))
        
        print(f"✅ [LLM Send] Message sent to {device_name}")
        
        return {
            "success": True,
            "message": f"✅ Đã gửi tin nhắn đến {device_name}. Robot sẽ đọc và trả lời qua giọng nói.",
            "device_name": device_name,
            "device_index": device_index,
            "sent_message": message,
            "note": "Robot sẽ tự động trả lời qua voice. Không cần đợi response text."
        }
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": f"Lỗi khi gửi tin nhắn: {str(e)}"
        }


async def broadcast_to_all_llm(message: str, wait_response: bool = False) -> dict:
    """
    Gửi tin nhắn đến TẤT CẢ thiết bị LLM đang kết nối.
    
    Args:
        message: Tin nhắn muốn broadcast
        wait_response: KHÔNG SỬ DỤNG - để tương thích API cũ
        
    Returns:
        dict với kết quả gửi cho từng thiết bị
    """
    global xiaozhi_connected
    
    results = {
        "success": True,
        "message": message,
        "devices": []
    }
    
    sent_count = 0
    for device_index in [0, 1, 2]:
        if xiaozhi_connected.get(device_index, False):
            result = await send_message_to_llm(
                message=message,
                device_index=device_index
            )
            results["devices"].append({
                "device_index": device_index,
                "result": result
            })
            if result.get("success"):
                sent_count += 1
    
    results["sent_count"] = sent_count
    results["total_connected"] = sum(1 for v in xiaozhi_connected.values() if v)
    
    if sent_count == 0:
        results["success"] = False
        results["error"] = "Không có thiết bị nào đang kết nối."
    
    return results


def send_message_to_llm_sync(message: str, device_index: int = None, wait_response: bool = False, timeout: int = 30) -> dict:
    """
    Wrapper đồng bộ cho send_message_to_llm (dùng trong TOOLS handler)
    """
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            # Nếu đang trong async context, tạo task
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as executor:
                future = executor.submit(
                    asyncio.run,
                    send_message_to_llm(message, device_index)
                )
                return future.result(timeout=timeout + 5)
        else:
            return loop.run_until_complete(
                send_message_to_llm(message, device_index)
            )
    except Exception as e:
        return {"success": False, "error": str(e)}


TOOLS = {
    # ============================================================
    # 📨 SEND MESSAGE TO LLM - Gửi tin nhắn cho robot/LLM tự trả lời
    # ============================================================
    "send_message_to_llm": {
        "handler": send_message_to_llm,
        "description": "📨 GỬI TIN NHẮN CHO LLM/ROBOT - Gửi message qua WebSocket để LLM cloud đọc và TỰ TRẢ LỜI. Use when: 'gửi tin nhắn cho robot', 'nói với AI', 'chat với LLM', 'hỏi robot', 'send message to AI'. Robot sẽ đọc được tin nhắn và tự động phản hồi qua giọng nói hoặc text.",
        "parameters": {
            "message": {
                "type": "string",
                "description": "Tin nhắn/câu hỏi muốn gửi cho LLM. VD: 'Xin chào', 'Hôm nay thời tiết thế nào?', 'Kể cho tôi một câu chuyện'",
                "required": True
            },
            "device_index": {
                "type": "integer",
                "description": "Index thiết bị (0, 1, hoặc 2). Mặc định: thiết bị đang active. 0=Thiết bị 1, 1=Thiết bị 2, 2=Thiết bị 3",
                "required": False
            },
            "wait_response": {
                "type": "boolean",
                "description": "Có đợi LLM trả lời không? True=đợi response (mặc định), False=gửi xong trả về luôn",
                "required": False
            },
            "timeout": {
                "type": "integer",
                "description": "Thời gian chờ response (giây). Mặc định 30 giây.",
                "required": False
            }
        }
    },
    "broadcast_to_all_llm": {
        "handler": broadcast_to_all_llm,
        "description": "📢 BROADCAST TIN NHẮN ĐẾN TẤT CẢ LLM/ROBOT - Gửi cùng một message đến tất cả thiết bị đang kết nối. Use when: 'gửi tin nhắn cho tất cả robot', 'broadcast message', 'thông báo cho tất cả AI'.",
        "parameters": {
            "message": {
                "type": "string",
                "description": "Tin nhắn muốn broadcast đến tất cả thiết bị",
                "required": True
            },
            "wait_response": {
                "type": "boolean",
                "description": "Có đợi response từ các thiết bị không? Mặc định False (broadcast thường không đợi)",
                "required": False
            }
        }
    },
    
    "get_hardware_specs": {
        "handler": get_system_info,
        "description": "💻🔥 SPECS CẤU HÌNH HARDWARE - DUY NHẤT tool cho câu hỏi: 'cấu hình máy tính gì', 'máy tính này như thế nào', 'card đồ họa gì', 'CPU gì', 'GPU gì', 'mainboard gì', 'thế hệ CPU', 'RTX RTX mấy', 'Intel thế hệ mấy', 'AMD Ryzen mấy'. Trả về: CPU generation (Intel 13th gen), GPU series (RTX 4080), motherboard, BIOS, RAM specs. KHÔNG dùng cho performance monitoring!",
        "parameters": {
            "category": {
                "type": "string",
                "description": "'cpu', 'gpu', 'motherboard', 'memory', 'all'. Mặc định: all",
                "required": False
            }
        }
    },
    "set_volume": {
        "handler": set_volume, 
        "description": "ĐIỀU CHỈNH âm lượng máy tính đến mức CỤ THỂ (0-100%). Use when user says: 'chỉnh âm lượng 50', 'đặt âm lượng 80', 'volume 30', 'set volume to 60', 'để âm lượng ở mức 40'. Examples: level=50 (âm lượng vừa), level=80 (to), level=20 (nhỏ), level=0 (tắt hẳn).", 
        "parameters": {"level": {"type": "integer", "description": "Mức âm lượng từ 0-100 (0=tắt hẳn, 50=vừa phải, 100=tối đa)", "required": True}}
    },
    "get_volume": {"handler": get_volume, "description": "Kiểm tra mức âm lượng hiện tại của máy tính. Use when: 'âm lượng bao nhiêu', 'check volume', 'xem âm lượng'", "parameters": {}},
    "mute_volume": {"handler": mute_volume, "description": "TẮT TIẾNG máy tính (mute) hoàn toàn. Use when: 'tắt tiếng', 'mute', 'câm', 'im lặng'", "parameters": {}},
    "unmute_volume": {"handler": unmute_volume, "description": "BẬT LẠI TIẾNG máy tính (unmute). Use when: 'bật tiếng', 'unmute', 'mở tiếng lại'", "parameters": {}},
    "volume_up": {"handler": volume_up, "description": "TĂNG âm lượng lên một chút (mỗi bước ~2%). Use when: 'tăng âm lượng', 'to hơn', 'volume up', 'lớn hơn'", "parameters": {"steps": {"type": "integer", "description": "Số bước tăng (mặc định 5 = tăng ~10%)", "required": False}}},
    "volume_down": {"handler": volume_down, "description": "GIẢM âm lượng xuống một chút (mỗi bước ~2%). Use when: 'giảm âm lượng', 'nhỏ hơn', 'volume down', 'bớt to'", "parameters": {"steps": {"type": "integer", "description": "Số bước giảm (mặc định 5 = giảm ~10%)", "required": False}}},
    "take_screenshot": {
        "handler": take_screenshot, 
        "description": "Chụp màn hình toàn bộ và LƯU FILE ẢNH. Tự động lưu vào thư mục Downloads với tên file có timestamp. Use when user asks: 'chụp màn hình', 'screenshot', 'capture screen'.", 
        "parameters": {
            "filename": {
                "type": "string",
                "description": "Tên file lưu ảnh (optional). Mặc định: screenshot_YYYYMMDD_HHMMSS.png. Ví dụ: 'my_screen.png'",
                "required": False
            }
        }
    },
    "show_notification": {"handler": show_notification, "description": "Hiển thị thông báo", "parameters": {"title": {"type": "string", "description": "Tiêu đề", "required": True}, "message": {"type": "string", "description": "Nội dung", "required": True}}},
    "get_system_resources": {"handler": get_system_resources, "description": "📊 PERFORMANCE MONITORING - CHỈ để xem CPU %, RAM %, Disk % đang sử dụng. CHO PERFORMANCE/MONITOR, KHÔNG cho câu hỏi về 'cấu hình máy tính', 'GPU gì', 'CPU gì'. Dùng get_hardware_specs cho hardware specs!", "parameters": {}},
    "get_current_time": {"handler": get_current_time, "description": "Thời gian hiện tại", "parameters": {}},
    "calculator": {"handler": calculator, "description": "Tính toán", "parameters": {"expression": {"type": "string", "description": "Biểu thức", "required": True}}},
    "open_application": {
        "handler": open_application, 
        "description": "Mở ứng dụng Windows với tìm kiếm thông minh. HỖ TRỢ 50+ ỨNG DỤNG: Windows (notepad, calc, paint, cmd, taskmgr), Browsers (chrome, firefox, edge, brave), Microsoft Office (word, excel, powerpoint, outlook, teams), Adobe Creative (photoshop, illustrator, premiere, after effects, lightroom), Development (vscode, pycharm, sublime, notepad++), 3D/Design (blender, maya, autocad, solidworks, fusion360), Communication (discord, slack, zoom, telegram, zalo), Media (vlc, spotify, itunes). Hỗ trợ tên TIẾNG VIỆT ('máy tính'→Calculator, 'máy ghi chú'→Notepad). Tự động tìm trong PATH, Registry, Program Files. Ví dụ: 'photoshop', 'excel', 'chrome', 'blender'.", 
        "parameters": {
            "app_name": {
                "type": "string", 
                "description": "Tên ứng dụng (ví dụ: 'excel', 'photoshop', 'chrome', 'vscode', 'blender', 'word'). Có thể dùng tên đầy đủ ('microsoft excel') hoặc viết tắt ('ps'→Photoshop). Hỗ trợ tiếng Việt.", 
                "required": True
            }
        }
    },
    "list_running_processes": {"handler": list_running_processes, "description": "Liệt kê tiến trình", "parameters": {"limit": {"type": "integer", "description": "Số lượng", "required": False}}},
    "find_process": {
        "handler": find_process,
        "description": "🔍 TÌM KIẾM PROCESS - Tìm process cụ thể theo tên hoặc xem tất cả. Triggers: 'tìm process excel', 'excel có chạy không', 'process nào đang chạy'. Better than list_running_processes with limit.",
        "parameters": {
            "name_pattern": {"type": "string", "description": "Tên process cần tìm (VD: 'excel', 'chrome', 'notepad'). Để trống = tất cả", "required": False},
            "show_all": {"type": "boolean", "description": "True=hiển thị tất cả process, False=chỉ top 20 (default)", "required": False}
        }
    },
    "kill_process": {
        "handler": kill_process, 
        "description": "🔪 Kill tiến trình theo tên hoặc PID. Có thể kill ngay lập tức (force=True) hoặc đóng mềm (force=False). VD: 'kill notepad', 'tắt chrome'", 
        "parameters": {
            "identifier": {"type": "string", "description": "Tên app hoặc PID. VD: notepad, chrome, 1234", "required": True},
            "force": {"type": "boolean", "description": "True=kill ngay (mặc định), False=đóng mềm", "required": False},
            "exact_match": {"type": "boolean", "description": "True=tên khớp chính xác, False=chứa tên là được (mặc định)", "required": False}
        }
    },
    "force_kill_app": {
        "handler": force_kill_app, 
        "description": "💀 FORCE KILL APP NGAY LẬP TỨC - không hỏi han, kill hết tất cả instances. Dùng khi cần kill app ngay, không chờ đợi. VD: 'force kill chrome', 'buộc tắt notepad'", 
        "parameters": {
            "app_name": {"type": "string", "description": "Tên app cần force kill. VD: notepad, chrome, firefox, Code", "required": True}
        }
    },
    "create_file": {"handler": create_file, "description": "Tạo file", "parameters": {"path": {"type": "string", "description": "Đường dẫn", "required": True}, "content": {"type": "string", "description": "Nội dung", "required": True}}},
    "read_file": {"handler": read_file, "description": "Đọc file", "parameters": {"path": {"type": "string", "description": "Đường dẫn", "required": True}}},
    "list_files": {"handler": list_files, "description": "Liệt kê files", "parameters": {"directory": {"type": "string", "description": "Thư mục", "required": True}}},
    "get_battery_status": {"handler": get_battery_status, "description": "Thông tin pin", "parameters": {}},
    "get_network_info": {"handler": get_network_info, "description": "Thông tin mạng", "parameters": {}},
    "search_web": {"handler": search_web, "description": "MỞ TRÌNH DUYỆT để tìm kiếm trên Google. CHỈ dùng khi user YÊU CẦU MỞ BROWSER để search (ví dụ: 'mở google tìm kiếm...', 'search google về...'). KHÔNG dùng để trả lời câu hỏi - hãy dùng ask_gemini thay vì search_web cho câu hỏi thông thường", "parameters": {"query": {"type": "string", "description": "Từ khóa", "required": True}}},
    
    # MEDIA PLAYER CONTROLS (Chủ yếu cho Spotify, YouTube, VLC - WMP có giới hạn)
    "media_play_pause": {
        "handler": media_play_pause, 
        "description": "⏯️ Phát/Tạm dừng external media players (Spotify, YouTube, VLC, iTunes, Discord, Chrome video...). Dùng Windows media keys. ⚠️ LƯU Ý: KHÔNG hoạt động tốt với music_library (Windows Media Player tự đóng sau khi phát). Dùng stop_music() để dừng music_library. Ví dụ: 'tạm dừng spotify', 'pause youtube'.", 
        "parameters": {}
    },
    "media_next_track": {
        "handler": media_next_track, 
        "description": "⏭️ Chuyển bài tiếp theo trên playlist. Hoạt động với: Spotify, YouTube playlist, VLC, iTunes. ⚠️ KHÔNG dùng cho music_library (WMP tự đóng). Ví dụ: 'bài tiếp spotify', 'next youtube'.", 
        "parameters": {}
    },
    "media_previous_track": {
        "handler": media_previous_track, 
        "description": "⏮️ Quay lại bài trước. Hoạt động với: Spotify, YouTube, VLC, iTunes. ⚠️ KHÔNG dùng cho music_library. Ví dụ: 'bài trước spotify', 'previous vlc'.", 
        "parameters": {}
    },
    "media_stop": {
        "handler": media_stop, 
        "description": "⏹️ Dừng phát external media players. Hoạt động với Spotify, VLC, YouTube. Với music_library, dùng stop_music() thay thế (đóng Windows Media Player). Ví dụ: 'stop spotify', 'dừng vlc'.", 
        "parameters": {}
    },
    "media_control": {
        "handler": media_control, 
        "description": "🎛️ Tool TỔNG HỢP điều khiển EXTERNAL media players (Spotify, YouTube, VLC, iTunes...). Hỗ trợ: play, pause, next, previous, stop, volume_up, volume_down, mute. ⚠️ KHÔNG dùng cho music_library (dùng stop_music). Best for: Spotify, YouTube, VLC. Ví dụ: media_control('next') cho Spotify, media_control('pause') cho YouTube.", 
        "parameters": {
            "action": {
                "type": "string", 
                "description": "Hành động: 'play', 'pause', 'next', 'previous', 'stop', 'volume_up', 'volume_down', 'mute'. Ví dụ: 'next', 'pause', 'mute'.", 
                "required": True
            }
        }
    },
    
    "save_music_folder_config": {
        "handler": save_music_folder_config,
        "description": "Save user's music folder path configuration. This folder will be prioritized for playing music using Windows default media player.",
        "parameters": {
            "folder_path": {
                "type": "string",
                "description": r"Full path to user's music folder (e.g., C:\Users\Name\Music)",
                "required": True
            }
        }
    },
    "play_music_from_user_folder": {
        "handler": play_music_from_user_folder,
        "description": "🎵 [PYTHON-VLC] ⭐ ƯU TIÊN #1: Phát nhạc từ THƯ MỤC NGƯỜI DÙNG ĐÃ CẤU HÌNH (link riêng). Khi user nói 'phát nhạc từ thư mục của tôi', 'play từ folder F:', 'nhạc trong ổ D' → DÙNG TOOL NÀY! Tìm theo tên bài: filename='tên bài'. NHANH vì dùng Python-VLC nội bộ. Nếu chưa config thì báo lỗi → user cần vào Music Settings.",
        "parameters": {
            "filename": {
                "type": "string",
                "description": "Tên bài hát cần tìm (tìm partial match). Để trống = phát bài đầu trong thư mục.",
                "required": False
            },
            "auto_play": {
                "type": "boolean",
                "description": "Tự động phát? Default True.",
                "required": False
            }
        }
    },
    
    "get_active_media_players": {
        "handler": get_active_media_players,
        "description": "🔍 [KHÔNG CẦN GỌI] Lấy danh sách media players đang chạy. ⚠️ KHÔNG CẦN gọi tool này trước khi điều khiển nhạc! Nhạc local LUÔN dùng Python-VLC (pause_music, stop_music, music_next). YouTube LUÔN dùng youtube_* tools.",
        "parameters": {}
    },
    
    # TASK MEMORY TOOLS - Ghi nhớ tác vụ để phản hồi nhanh và chính xác
    "remember_task": {
        "handler": remember_task,
        "description": "📝 GHI NHỚ TÁC VỤ - Lưu lại tác vụ đã thực hiện vào bộ nhớ dài hạn. Giúp AI nhớ những gì đã làm để phản hồi nhanh và chính xác hơn. Gọi tool này SAU KHI hoàn thành một tác vụ quan trọng.",
        "parameters": {
            "tool_name": {"type": "string", "description": "Tên tool đã sử dụng", "required": True},
            "params": {"type": "object", "description": "Tham số đã dùng (optional)", "required": False},
            "result_message": {"type": "string", "description": "Kết quả/message", "required": False},
            "user_request": {"type": "string", "description": "Yêu cầu gốc của user", "required": False}
        }
    },
    "recall_tasks": {
        "handler": recall_tasks,
        "description": "🧠 NHỚ LẠI TÁC VỤ - Truy vấn lịch sử các tác vụ đã thực hiện. Gọi tool này ĐẦU TIÊN khi user hỏi 'đã làm gì', 'nhắc lại', 'lần trước', hoặc khi cần context về các tác vụ trước đó.",
        "parameters": {
            "keyword": {"type": "string", "description": "Từ khóa tìm kiếm (optional). Để trống = lấy tác vụ gần nhất", "required": False},
            "limit": {"type": "integer", "description": "Số lượng tác vụ tối đa (default 10)", "required": False}
        }
    },
    "get_task_summary": {
        "handler": get_task_summary,
        "description": "📊 THỐNG KÊ TÁC VỤ - Lấy tổng hợp về các tác vụ đã thực hiện. Cho biết tools nào được dùng nhiều nhất, tỷ lệ thành công. Dùng khi user hỏi 'thống kê', 'báo cáo', 'đã dùng tools gì'.",
        "parameters": {}
    },
    "forget_all_tasks": {
        "handler": forget_all_tasks,
        "description": "🗑️ XÓA LỊCH SỬ - Xóa toàn bộ lịch sử tác vụ đã ghi nhớ. CHỈ DÙNG khi user yêu cầu rõ ràng 'xóa lịch sử', 'quên hết', 'reset memory'.",
        "parameters": {}
    },
    
    "set_brightness": {"handler": set_brightness, "description": "Độ sáng màn hình", "parameters": {"level": {"type": "integer", "description": "Độ sáng 0-100", "required": True}}},
    "get_clipboard": {"handler": get_clipboard, "description": "Lấy clipboard", "parameters": {}},
    "set_clipboard": {"handler": set_clipboard, "description": "Đặt clipboard", "parameters": {"text": {"type": "string", "description": "Nội dung", "required": True}}},
    "play_sound": {"handler": play_sound, "description": "Phát âm thanh", "parameters": {"frequency": {"type": "integer", "description": "Tần số Hz", "required": False}, "duration": {"type": "integer", "description": "Thời gian ms", "required": False}}},
    "get_disk_usage": {"handler": get_disk_usage, "description": "Thông tin đĩa", "parameters": {}},
    
    # ============================================================
    # 🎵 MUSIC LIBRARY TOOLS - PYTHON-VLC (LOCAL FILES)
    # Dùng cho file nhạc .mp3/.wav/.flac trong máy tính
    # KHÔNG dùng cho YouTube - YouTube có tools riêng (youtube_*)
    # ============================================================
    "list_music": {
        "handler": list_music, 
        "description": "📂 [LOCAL MUSIC] Liệt kê tất cả nhạc trong thư viện music_library. Triggers: 'xem danh sách nhạc', 'có bài gì', 'list music'. Auto-play mặc định = True (phát bài đầu tiên). Dùng subfolder='Pop' để lọc theo thể loại.", 
        "parameters": {
            "subfolder": {
                "type": "string", 
                "description": "Thư mục con để lọc (VD: 'Pop', 'Rock', 'EDM'). Để trống = tất cả.", 
                "required": False
            },
            "auto_play": {
                "type": "boolean",
                "description": "Tự động phát bài đầu tiên? Default=True. Set False nếu chỉ muốn xem danh sách.",
                "required": False
            }
        }
    },
    "play_music": {
        "handler": play_music, 
        "description": "🎵 PHÁT NHẠC LOCAL (Python-VLC) - Triggers: 'phát nhạc', 'bật nhạc', 'mở nhạc', 'nghe nhạc', 'play nhạc', 'phát bài [tên]', 'phat nhac', 'bat nhac'. VD: 'phát bài đa nghi' → play_music(filename='đa nghi'). ⚠️ Nếu user nói 'youtube/video' → dùng open_youtube!", 
        "parameters": {
            "filename": {
                "type": "string", 
                "description": "Tên bài nhạc (partial match). VD: 'đa nghi', 'in love'. Hỗ trợ tiếng Việt.", 
                "required": True
            },
            "create_playlist": {
                "type": "boolean",
                "description": "Tạo playlist (default True).",
                "required": False
            }
        }
    },
    "pause_music": {
        "handler": pause_music,
        "description": "⏸️ TẠM DỪNG NHẠC - ⭐ GỌI NGAY khi user nói: 'dừng', 'dừng nhạc', 'tạm dừng', 'pause', 'ngừng', 'ngưng nhạc', 'nghỉ', 'im đi', 'dừng lại'. Voice: 'dung', 'dung nhac', 'tam dung', 'pao', 'poz', 'ngung', 'dung lai'. Không cần parameter - gọi pause_music() là xong! ⚠️ Nếu có 'youtube' → youtube_play_pause()",
        "parameters": {}
    },
    "resume_music": {
        "handler": resume_music,
        "description": "▶️ TIẾP TỤC PHÁT - ⭐ GỌI NGAY khi user nói: 'tiếp tục', 'phát tiếp', 'play lại', 'mở lại', 'phát đi', 'chơi tiếp'. Voice: 'tiep tuc', 'phat tiep', 'mo lai', 'bat lai'. Không cần parameter - gọi resume_music() là xong!",
        "parameters": {}
    },
    "stop_music": {
        "handler": stop_music, 
        "description": "⏹️ TẮT NHẠC HOÀN TOÀN - ⭐ GỌI NGAY khi user nói: 'tắt nhạc', 'dừng hẳn', 'stop', 'off nhạc', 'không nghe nữa', 'tắt đi'. Voice: 'tat nhac', 'dung han', 'stóp', 'of nhac'. Không cần parameter - gọi stop_music() là xong!", 
        "parameters": {}
    },
    
    # 🌟 SMART MUSIC CONTROL - Tool thông minh nhất
    "smart_music_control": {
        "handler": smart_music_control,
        "description": "🎵🔥 ĐIỀU KHIỂN NHẠC THÔNG MINH - ⭐ GỌI KHI nghe: 'bài tiếp/next/chuyển bài', 'bài trước/quay lại', 'dừng/pause/tạm dừng', 'tắt nhạc/stop', 'phát bài [tên]', 'tăng/giảm âm lượng'. Voice: 'bai tiep', 'bai truoc', 'dung nhac', 'tam dung', 'pao'. VD: smart_music_control('bài tiếp'), smart_music_control('dừng'). Tool tự xử lý tất cả!",
        "parameters": {
            "command": {
                "type": "string",
                "description": "Lệnh tiếng Việt/English. VD: 'bài tiếp', 'bài trước', 'dừng', 'pause', 'phát bài love'",
                "required": True
            }
        }
    },
    
    "detect_and_execute_music": {
        "handler": detect_and_execute_music,
        "description": "🎵🔍 TỰ ĐỘNG PHÁT HIỆN LỆNH NHẠC - Kiểm tra input có phải lệnh nhạc không và tự động thực thi. Dùng khi không chắc input có phải lệnh nhạc.",
        "parameters": {
            "text": {
                "type": "string", 
                "description": "Text cần kiểm tra",
                "required": True
            }
        }
    },
    
    "music_next": {
        "handler": music_next,
        "description": "⏭️ BÀI TIẾP THEO - ⭐ GỌI NGAY khi user nói: 'bài tiếp', 'bài tiếp theo', 'chuyển bài', 'bài khác', 'next', 'skip', 'kế tiếp', 'sang bài', 'bài sau'. Voice: 'bai tiep', 'chuyen bai', 'bai khac', 'tiep theo', 'ke tiep', 'nex', 'ních'. Không cần parameter - gọi music_next() là xong!",
        "parameters": {}
    },
    "music_previous": {
        "handler": music_previous,
        "description": "⏮️ BÀI TRƯỚC - ⭐ GỌI NGAY khi user nói: 'bài trước', 'quay lại', 'bài trước đó', 'previous', 'back', 'lùi bài', 'bài cũ'. Voice: 'bai truoc', 'quay lai', 'lui bai', 'bai cu', 'pre', 'prê'. Không cần parameter - gọi music_previous() là xong!",
        "parameters": {}
    },
    "get_music_status": {
        "handler": get_music_status,
        "description": "📊 TRẠNG THÁI NHẠC - Triggers: 'đang phát gì', 'bài gì đang phát', 'music status', 'dang phat gi'. Trả về: tên bài, thời gian, âm lượng, playlist.",
        "parameters": {}
    },
    "seek_music": {
        "handler": seek_music,
        "description": "🔀 TUA ĐẾN VỊ TRÍ - Triggers: 'tua đến giữa bài', 'nhảy đến phút', 'skip 50%', 'tua den', 'nhay den'. 0%=đầu, 50%=giữa, 100%=cuối. ⚠️ 'tua youtube' → youtube_forward!",
        "parameters": {
            "percentage": {
                "type": "number",
                "description": "Vị trí % (0-100). 50=giữa bài.",
                "required": True
            }
        }
    },
    "music_volume": {
        "handler": music_volume,
        "description": "🔊 ÂM LƯỢNG NHẠC LOCAL - Triggers: 'tăng âm lượng', 'giảm tiếng', 'volume 80', 'to lên', 'nhỏ lại', 'tang am luong', 'giam tien'. Level: 0=tắt, 50=vừa, 100=max. ⚠️ 'volume youtube' → youtube_volume_up/down!",
        "parameters": {
            "level": {
                "type": "integer",
                "description": "Mức âm lượng 0-100.",
                "required": True
            }
        }
    },
    "save_music_folder_config": {
        "handler": save_music_folder_config,
        "description": "Lưu đường dẫn thư mục nhạc của user. Dùng để ưu tiên phát nhạc từ folder này.",
        "parameters": {
            "folder_path": {
                "type": "string",
                "description": r"Đường dẫn đầy đủ đến thư mục nhạc (VD: C:\Users\Name\Music)",
                "required": True
            }
        }
    },
    "search_music": {
        "handler": search_music, 
        "description": "🔍 TÌM NHẠC THEO TỪ KHÓA - Triggers: 'tìm bài [keyword]', 'search nhạc', 'có bài nào tên', 'tim bai', 'search bai'. Tìm trong thư viện local, hỗ trợ tiếng Việt, auto-play mặc định.", 
        "parameters": {
            "keyword": {
                "type": "string", 
                "description": "Từ khóa tìm kiếm. VD: 'love', 'buồn', 'đa nghi'.", 
                "required": True
            },
            "auto_play": {
                "type": "boolean",
                "description": "Tự động phát bài đầu tiên? Default=True.",
                "required": False
            }
        }
    },
    
    # QUICK WEBSITE ACCESS TOOLS
    "open_youtube": {
        "handler": open_youtube, 
        "description": "📺 MỞ YOUTUBE - Triggers: 'mở youtube', 'vào youtube', 'xem youtube', 'youtube [tên video]'. ✨ NEW: TỰ ĐỘNG phát video trực tiếp nếu query CỤ THỂ (>= 2 từ)! VD: 'mở youtube Lạc Trôi' → Mở video trực tiếp (không phải search page). Query 1 từ → mở search page.", 
        "parameters": {
            "search_query": {
                "type": "string", 
                "description": "Tên video/từ khóa. Query >= 2 từ = auto phát video trực tiếp. Query 1 từ = search page. Để trống = homepage.", 
                "required": False
            }
        }
    },
    "search_youtube_video": {
        "handler": search_youtube_video,
        "description": "🔍 TÌM VIDEO YOUTUBE (Explicit) - ⚠️ CHỈ dùng khi user YÊU CẦU 'tìm video', 'search video', hoặc muốn xem top 5 results. Còn lại DÙNG open_youtube (đã có auto-detect direct video). VD: 'tìm video Sơn Tùng' → search_youtube_video. 'mở youtube Sơn Tùng Chúng Ta' → open_youtube (preferred).",
        "parameters": {
            "video_title": {
                "type": "string",
                "description": "Tên video/từ khóa. VD: 'Hãy Trao Cho Anh', 'Rap Việt tập 1'",
                "required": True
            },
            "auto_open": {
                "type": "boolean",
                "description": "Tự động mở video (default: True). Set False để chỉ tìm.",
                "required": False
            }
        }
    },
    "open_youtube_playlist": {
        "handler": open_youtube_playlist,
        "description": "📜 MỞ PLAYLIST YOUTUBE (đã lưu Web UI) - Triggers: 'mở playlist [tên]', 'phát playlist youtube', 'mo playlist'. VD: 'mở playlist nhạc việt 1'. ⚠️ Không dùng cho .mp3 local → play_music!",
        "parameters": {
            "playlist_name": {
                "type": "string",
                "description": "Tên playlist đã đăng ký. VD: 'nhạc việt 1', 'chill', 'EDM'",
                "required": True
            }
        }
    },
    
    # YOUTUBE PLAYER CONTROLS
    "control_youtube": {
        "handler": control_youtube,
        "description": "🎬 Điều khiển YOUTUBE bằng shortcuts. Actions: play_pause, rewind_10, forward_10, volume_up/down, mute_toggle. VD: 'tạm dừng youtube'",
        "parameters": {
            "action": {
                "type": "string",
                "description": "Action: play_pause, rewind_10, forward_10, volume_up/down, mute_toggle",
                "required": True
            }
        }
    },
    "youtube_play_pause": {
        "handler": youtube_play_pause,
        "description": "⏯️ PLAY/PAUSE YOUTUBE - Triggers: 'dừng youtube', 'pause youtube', 'tiếp tục youtube', 'play youtube', 'dung youtube'. ⚠️ 'dừng nhạc' (không có youtube) → pause_music!",
        "parameters": {}
    },
    "youtube_rewind": {
        "handler": youtube_rewind,
        "description": "⏪ TUA LÙI YOUTUBE - Triggers: 'lùi youtube', 'tua lùi youtube', 'rewind youtube', 'lui youtube'. 5s=phím ← | 10s=phím J",
        "parameters": {
            "seconds": {"type": "integer", "description": "Giây tua lùi: 5 hoặc 10", "required": False}
        }
    },
    "youtube_forward": {
        "handler": youtube_forward,
        "description": "⏩ TUA TỚI YOUTUBE - Triggers: 'tua youtube', 'skip youtube', 'forward youtube', 'tua video'. 5s=phím → | 10s=phím L",
        "parameters": {
            "seconds": {"type": "integer", "description": "Giây tua tới: 5 hoặc 10", "required": False}
        }
    },
    "youtube_volume_up": {
        "handler": youtube_volume_up,
        "description": "🔊 TĂNG ÂM LƯỢNG YOUTUBE - Triggers: 'tăng tiếng youtube', 'volume up youtube', 'tang am luong youtube'. ⚠️ 'tăng tiếng nhạc' → music_volume!",
        "parameters": {}
    },
    "youtube_volume_down": {
        "handler": youtube_volume_down,
        "description": "🔉 GIẢM ÂM LƯỢNG YOUTUBE - Triggers: 'giảm tiếng youtube', 'volume down youtube', 'giam am luong youtube'. ⚠️ 'giảm tiếng nhạc' → music_volume!",
        "parameters": {}
    },
    "youtube_mute": {
        "handler": youtube_mute,
        "description": "🔇 TẮT/BẬT TIẾNG YOUTUBE - Triggers: 'tắt tiếng youtube', 'mute youtube', 'bật tiếng youtube', 'tat tien youtube'.",
        "parameters": {}
    },
    "youtube_fullscreen": {
        "handler": youtube_fullscreen,
        "description": "📺 FULLSCREEN YOUTUBE - Triggers: 'fullscreen youtube', 'toàn màn hình', 'phóng to youtube', 'thu nhỏ youtube', 'toan man hinh'.",
        "parameters": {}
    },
    "youtube_captions": {
        "handler": youtube_captions,
        "description": "💬 BẬT/TẮT PHỤ ĐỀ YOUTUBE - Triggers: 'bật sub', 'tắt sub', 'bật phụ đề', 'tắt phụ đề', 'caption youtube', 'bat sub', 'tat sub'.",
        "parameters": {}
    },
    "youtube_speed": {
        "handler": youtube_speed,
        "description": "⚡ ĐỔI TỐC ĐỘ YOUTUBE - Triggers: 'youtube nhanh hơn', 'youtube chậm hơn', 'tăng tốc youtube'. faster=+0.25x | slower=-0.25x | normal=1x",
        "parameters": {
            "speed": {"type": "string", "description": "'faster', 'slower', 'normal'", "required": False}
        }
    },
    
    # VLC PLAYER CONTROLS
    "control_vlc": {
        "handler": control_vlc,
        "description": "🎵 Điều khiển VLC PLAYER. Actions: play_pause, stop, next, previous, volume_up/down, mute, fullscreen",
        "parameters": {
            "action": {
                "type": "string",
                "description": "Action điều khiển VLC",
                "required": True
            }
        }
    },
    "vlc_play_pause": {
        "handler": vlc_play_pause,
        "description": "⏯️ Play/Pause VLC Player. VD: 'dừng vlc', 'pause vlc', 'tiếp tục vlc'",
        "parameters": {}
    },
    "vlc_stop": {
        "handler": vlc_stop,
        "description": "⏹️ Dừng phát VLC hoàn toàn. VD: 'stop vlc', 'tắt nhạc vlc'",
        "parameters": {}
    },
    "vlc_next": {
        "handler": vlc_next,
        "description": "⏭️ Chuyển bài tiếp theo trong VLC. VD: 'bài tiếp vlc', 'next vlc', 'chuyển bài vlc'",
        "parameters": {}
    },
    "vlc_previous": {
        "handler": vlc_previous,
        "description": "⏮️ Quay lại bài trước trong VLC. VD: 'bài trước vlc', 'previous vlc'",
        "parameters": {}
    },
    "vlc_volume_up": {
        "handler": vlc_volume_up,
        "description": "🔊 Tăng âm lượng VLC. VD: 'tăng âm lượng vlc', 'vlc to hơn'",
        "parameters": {}
    },
    "vlc_volume_down": {
        "handler": vlc_volume_down,
        "description": "🔉 Giảm âm lượng VLC. VD: 'giảm âm lượng vlc', 'vlc nhỏ hơn'",
        "parameters": {}
    },
    "vlc_mute": {
        "handler": vlc_mute,
        "description": "🔇 Bật/Tắt tiếng VLC. VD: 'tắt tiếng vlc', 'mute vlc'",
        "parameters": {}
    },
    "vlc_forward": {
        "handler": vlc_forward,
        "description": "⏩ Tua tới trong VLC. Tự động chọn 3s/10s/60s. VD: 'tua tới vlc', 'skip vlc'",
        "parameters": {
            "seconds": {"type": "integer", "description": "Số giây tua tới (≤5→3s, ≤30→10s, >30→60s)", "required": False}
        }
    },
    "vlc_backward": {
        "handler": vlc_backward,
        "description": "⏪ Tua lùi trong VLC. Tự động chọn 3s/10s/60s. VD: 'lùi vlc', 'rewind vlc'",
        "parameters": {
            "seconds": {"type": "integer", "description": "Số giây tua lùi", "required": False}
        }
    },
    
    # ============================================================
    # WINDOWS MEDIA PLAYER CONTROLS
    # ============================================================
    "control_wmp": {
        "handler": control_wmp,
        "description": "🎶 Điều khiển Windows Media Player. Actions: play_pause, stop, next, previous, volume_up, volume_down, mute, fullscreen, forward, backward",
        "parameters": {
            "action": {"type": "string", "description": "Hành động điều khiển WMP", "required": True}
        }
    },
    "wmp_play_pause": {
        "handler": wmp_play_pause,
        "description": "⏯️ Play/Pause Windows Media Player. VD: 'dừng wmp', 'pause media player'",
        "parameters": {}
    },
    "wmp_stop": {
        "handler": wmp_stop,
        "description": "⏹️ Dừng Windows Media Player. VD: 'stop wmp', 'tắt media player'",
        "parameters": {}
    },
    "wmp_next": {
        "handler": wmp_next,
        "description": "⏭️ Bài tiếp theo trong Windows Media Player. VD: 'bài tiếp wmp', 'next media player'",
        "parameters": {}
    },
    "wmp_previous": {
        "handler": wmp_previous,
        "description": "⏮️ Bài trước trong Windows Media Player. VD: 'bài trước wmp', 'previous media player'",
        "parameters": {}
    },
    "wmp_volume_up": {
        "handler": wmp_volume_up,
        "description": "🔊 Tăng âm lượng Windows Media Player. VD: 'tăng âm lượng wmp'",
        "parameters": {}
    },
    "wmp_volume_down": {
        "handler": wmp_volume_down,
        "description": "🔉 Giảm âm lượng Windows Media Player. VD: 'giảm âm lượng wmp'",
        "parameters": {}
    },
    "wmp_mute": {
        "handler": wmp_mute,
        "description": "🔇 Bật/Tắt tiếng Windows Media Player. VD: 'tắt tiếng wmp', 'mute media player'",
        "parameters": {}
    },
    
    # ============================================================
    # SMART MEDIA CONTROL - Ưu tiên Python-VLC nội bộ
    # ============================================================
    "smart_media_control": {
        "handler": smart_media_control,
        "description": "🎵 [PYTHON-VLC ƯU TIÊN] Điều khiển nhạc - ƯU TIÊN PYTHON-VLC TRƯỚC, sau đó mới tới Spotify/WMP/YouTube. Actions: play_pause, stop, next, previous, volume_up, volume_down, mute. Nếu chưa phát nhạc, dùng play_music() trước!",
        "parameters": {
            "action": {
                "type": "string",
                "description": "Hành động: play_pause, stop, next, previous, volume_up, volume_down, mute",
                "required": True
            }
        }
    },
    
    # BROWSER AUTOMATION TOOLS
    "browser_open_url": {
        "handler": browser_open_url,
        "description": "Mở URL trong browser được điều khiển bởi Selenium (có thể tương tác với element). Khác với open_youtube/open_google là mở browser thông thường.",
        "parameters": {
            "url": {
                "type": "string",
                "description": "URL cần mở (VD: https://google.com, https://facebook.com)",
                "required": True
            }
        }
    },
    "browser_get_info": {
        "handler": browser_get_info,
        "description": "Lấy thông tin trang hiện tại (URL, title, số tab)",
        "parameters": {}
    },
    "browser_click": {
        "handler": browser_click,
        "description": "Click vào element trên trang web. Dùng để click button, link, etc.",
        "parameters": {
            "selector": {
                "type": "string",
                "description": "Selector để tìm element. VD: '#submit-btn', '.login-button', '//button[@id=\"login\"]'",
                "required": True
            },
            "by": {
                "type": "string",
                "description": "Loại selector: 'css' (default), 'xpath', 'id', 'name', 'class', 'tag'",
                "required": False
            }
        }
    },
    "browser_fill_input": {
        "handler": browser_fill_input,
        "description": "Điền text vào input field (form, search box, etc.)",
        "parameters": {
            "selector": {
                "type": "string",
                "description": "Selector của input field. VD: '#username', 'input[name=\"email\"]'",
                "required": True
            },
            "text": {
                "type": "string",
                "description": "Text cần điền vào input",
                "required": True
            },
            "by": {
                "type": "string",
                "description": "Loại selector: 'css' (default), 'xpath', 'id', 'name'",
                "required": False
            }
        }
    },
    "browser_scroll": {
        "handler": browser_scroll,
        "description": "Cuộn trang web lên/xuống",
        "parameters": {
            "direction": {
                "type": "string",
                "description": "Hướng cuộn: 'down' (default), 'up', 'top', 'bottom'",
                "required": False
            },
            "amount": {
                "type": "integer",
                "description": "Số pixel cuộn (nếu direction là down/up). Default: 500",
                "required": False
            }
        }
    },
    "browser_back": {
        "handler": browser_back,
        "description": "Quay lại trang trước trong browser",
        "parameters": {}
    },
    "browser_forward": {
        "handler": browser_forward,
        "description": "Tiến tới trang sau trong browser",
        "parameters": {}
    },
    "browser_refresh": {
        "handler": browser_refresh,
        "description": "Làm mới/reload trang hiện tại",
        "parameters": {}
    },
    "browser_screenshot": {
        "handler": browser_screenshot,
        "description": "Chụp screenshot trang web hiện tại",
        "parameters": {
            "filepath": {
                "type": "string",
                "description": "Đường dẫn lưu file (tùy chọn). VD: 'screenshot.png'. Mặc định: screenshot_YYYYMMDD_HHMMSS.png",
                "required": False
            }
        }
    },
    "browser_new_tab": {
        "handler": browser_new_tab,
        "description": "Mở tab mới trong browser",
        "parameters": {
            "url": {
                "type": "string",
                "description": "URL cần mở trong tab mới (tùy chọn)",
                "required": False
            }
        }
    },
    "browser_close_tab": {
        "handler": browser_close_tab,
        "description": "Đóng tab hiện tại",
        "parameters": {}
    },
    "browser_execute_js": {
        "handler": browser_execute_js,
        "description": "Thực thi JavaScript code trên trang web. Dùng cho các thao tác phức tạp.",
        "parameters": {
            "script": {
                "type": "string",
                "description": "JavaScript code cần chạy. VD: 'return document.title;', 'alert(\"Hello\");'",
                "required": True
            }
        }
    },
    "browser_close": {
        "handler": browser_close,
        "description": "Đóng browser hoàn toàn (đóng tất cả tab)",
        "parameters": {}
    },
    
    "open_facebook": {
        "handler": open_facebook, 
        "description": "Mở Facebook trong browser. Truy cập nhanh vào mạng xã hội phổ biến nhất.", 
        "parameters": {}
    },
    "open_google": {
        "handler": open_google, 
        "description": "MỞ TRÌNH DUYỆT Google. CHỈ dùng khi user YÊU CẦU MỞ TRANG WEB Google (ví dụ: 'mở google', 'mở trang google'). Nếu user chỉ HỎI CÂU HỎI thông thường, hãy dùng ask_gemini để TRẢ LỜI TRỰC TIẾP thay vì mở browser", 
        "parameters": {
            "search_query": {
                "type": "string", 
                "description": "Từ khóa tìm kiếm trên Google (tùy chọn). Để trống để mở trang chủ Google.", 
                "required": False
            }
        }
    },
    "open_tiktok": {
        "handler": open_tiktok, 
        "description": "Mở TikTok trong browser. Xem video ngắn trending và giải trí.", 
        "parameters": {}
    },
    "open_website": {
        "handler": open_website, 
        "description": "Mở trang web tùy chỉnh trong browser. Nhập URL đầy đủ hoặc tên miền.", 
        "parameters": {
            "url": {
                "type": "string", 
                "description": "URL của trang web (ví dụ: 'github.com' hoặc 'https://github.com/user/repo')", 
                "required": True
            }
        }
    },
    
    # YOUTUBE CONTROL TOOLS
    "control_youtube": {
        "handler": control_youtube, 
        "description": "Điều khiển YouTube player bằng keyboard shortcuts. Phải có cửa sổ YouTube đang active/focused. Hỗ trợ play/pause, tua video, điều chỉnh âm lượng, v.v.", 
        "parameters": {
            "action": {
                "type": "string", 
                "description": "Hành động điều khiển: play_pause, rewind_10, forward_10, rewind_5, forward_5, beginning, end, frame_back, frame_forward, volume_up, volume_down, mute_toggle", 
                "required": True
            }
        }
    },
    
    # NEWS TOOLS
    "get_vnexpress_news": {
        "handler": get_vnexpress_news,
        "description": "Lấy tin tức mới nhất từ VnExpress theo chủ đề. Trả về danh sách bài viết với tiêu đề, link, mô tả. Categories: home (mới nhất), thoi-su, the-gioi, kinh-doanh, giai-tri, the-thao, phap-luat, giao-duc, suc-khoe, du-lich, khoa-hoc, so-hoa, xe",
        "parameters": {
            "category": {
                "type": "string",
                "description": "Chủ đề tin tức: home, thoi-su, the-gioi, kinh-doanh, giai-tri, the-thao, phap-luat, giao-duc, suc-khoe, du-lich, khoa-hoc, so-hoa, xe. Mặc định: home",
                "required": False
            },
            "max_articles": {
                "type": "integer",
                "description": "Số lượng bài viết tối đa (1-20). Mặc định: 5",
                "required": False
            }
        }
    },
    "get_news_summary": {
        "handler": get_news_summary,
        "description": "Lấy tóm tắt nhanh tin tức (chỉ tiêu đề) từ VnExpress. Tự động lấy 10 tin mới nhất và hiển thị dạng danh sách ngắn gọn.",
        "parameters": {
            "category": {
                "type": "string",
                "description": "Chủ đề: home, thoi-su, the-gioi, kinh-doanh, giai-tri, the-thao, etc. Mặc định: home",
                "required": False
            }
        }
    },
    "search_news": {
        "handler": search_news,
        "description": "Tìm kiếm tin tức theo từ khóa trong các bài viết gần đây từ VnExpress. Tự động tìm trong nhiều chủ đề và trả về kết quả phù hợp nhất.",
        "parameters": {
            "keyword": {
                "type": "string",
                "description": "Từ khóa tìm kiếm (ví dụ: 'bóng đá', 'kinh tế', 'Covid', 'chính trị')",
                "required": True
            },
            "max_results": {
                "type": "integer",
                "description": "Số kết quả tối đa (1-10). Mặc định: 5",
                "required": False
            }
        }
    },
    "get_gold_price": {
        "handler": get_gold_price,
        "description": "Lấy giá vàng hôm nay từ BNews RSS feed. Hiển thị giá mua vào và bán ra của các loại vàng phổ biến (SJC, 9999, nhẫn tròn, v.v.). Tự động cập nhật giá mới nhất.",
        "parameters": {}
    },
    "analyze_gold_price_with_ai": {
        "handler": analyze_gold_price_with_ai,
        "description": "Phân tích thông minh giá vàng với AI (Gemini 3 Flash Preview + Google Search). So sánh giá hiện tại vs lịch sử, phân tích xu hướng, nguyên nhân biến động, dự báo, và khuyến nghị đầu tư chuyên sâu. Dùng khi cần phân tích chuyên môn về thị trường vàng.",
        "parameters": {
            "analysis_type": {
                "type": "string",
                "description": "Loại phân tích: 'compare_month' (so sánh với tháng trước), 'trend' (xu hướng hiện tại), 'forecast' (dự báo). Mặc định: 'compare_month'",
                "required": False
            }
        }
    },
    
    # AI ASSISTANT TOOLS
    "ask_gemini": {
        "handler": ask_gemini,
        "description": "✅ ƯU TIÊN DÙNG TOOL NÀY cho MỌI CÂU HỎI (MIỄN PHÍ 1500 requests/day). Gemini trả lời TRỰC TIẾP, NHANH, CHÍNH XÁC. Hữu ích cho: câu hỏi thông thường ('thủ tướng VN 2023 là ai', 'what is...', 'how to...'), phân tích, viết nội dung, dịch thuật, lịch sử, kiến thức tổng quát. Knowledge cutoff: ~10/2024 (đủ cho hầu hết câu hỏi). CHỈ dùng search_google_text nếu CẦN thông tin SAU 10/2024.",
        "parameters": {
            "prompt": {
                "type": "string",
                "description": "Câu hỏi hoặc nội dung muốn gửi cho Gemini AI",
                "required": True
            },
            "model": {
                "type": "string",
                "description": "Tên model Gemini (mặc định: models/gemini-3-flash-preview). Options: models/gemini-3-flash-preview (Flash 2.0, mới nhất), models/gemini-1.5-flash (Flash 1.5), models/gemini-1.5-pro (Pro 1.5, chất lượng cao nhất)",
                "required": False
            }
        }
    },
    
    "ask_gpt4": {
        "handler": ask_gpt4,
        "description": "TRẢ LỜI CÂU HỎI bằng OpenAI GPT-4 (TRẢ PHÍ, cần API key). DÙNG KHI CẦN: 1) Thông tin MỚI HƠN (knowledge đến 04/2024), 2) Phân tích PHỨC TẠP, 3) Reasoning SÂU, 4) Code generation chuyên nghiệp. GPT-4 MẠN HƠN Gemini cho code và phân tích, nhưng TRẢ PHÍ (~$0.01-0.03/1K tokens). Chọn GPT-4 khi cần chất lượng tối đa.",
        "parameters": {
            "prompt": {
                "type": "string",
                "description": "Câu hỏi hoặc nội dung muốn gửi cho GPT-4",
                "required": True
            },
            "model": {
                "type": "string",
                "description": "Tên model OpenAI (mặc định: gpt-4o). Options: gpt-4o (GPT-4 Omni, nhanh & rẻ nhất), gpt-4-turbo (mạnh nhất), gpt-3.5-turbo (rẻ & nhanh)",
                "required": False
            }
        }
    },
    
    # NETWORK/FIREWALL CHECK TOOLS
    "check_network_permission": {
        "handler": check_network_permission,
        "description": "🔥 KIỂM TRA QUYỀN KẾT NỐI MẠNG - Xem trạng thái Windows Firewall và Internet. Use when: 'kiểm tra firewall', 'quyền kết nối', 'check network', 'tình trạng mạng', 'firewall status', 'có được phép kết nối internet không'. Hiển thị: có rule firewall chưa, internet có kết nối không, hướng dẫn cấp quyền.",
        "parameters": {}
    },
    "request_firewall_permission": {
        "handler": request_firewall_permission,
        "description": "🔓 YÊU CẦU CẤP QUYỀN FIREWALL - Tự động thêm rule cho ứng dụng. Use when: 'cấp quyền firewall', 'allow firewall', 'thêm rule firewall'. Cần quyền Admin để hoạt động.",
        "parameters": {}
    },
    "check_internet_connection": {
        "handler": check_internet_connection,
        "description": "🌐 KIỂM TRA KẾT NỐI INTERNET - Test kết nối và độ trễ mạng. Use when: 'kiểm tra internet', 'test connection', 'có mạng không', 'ping', 'network status'.",
        "parameters": {}
    },
    
    # NEW TOOLS FROM REFERENCE
    "lock_computer": {"handler": lock_computer, "description": "Khóa máy tính", "parameters": {}},
    "shutdown_schedule": {"handler": shutdown_schedule, "description": "Lên lịch tắt máy", "parameters": {"action": {"type": "string", "description": "shutdown/restart/cancel", "required": True}, "delay": {"type": "integer", "description": "Trì hoãn (giây)", "required": False}}},
    "show_desktop": {"handler": show_desktop, "description": "Hiển thị desktop (Win+D)", "parameters": {}},
    "undo_operation": {"handler": undo_operation, "description": "Hoàn tác (Ctrl+Z)", "parameters": {}},
    "set_theme": {"handler": set_theme, "description": "Đổi theme Windows", "parameters": {"dark_mode": {"type": "boolean", "description": "True=tối, False=sáng", "required": False}}},
    "change_wallpaper": {"handler": change_wallpaper, "description": "Đổi hình nền", "parameters": {"keyword": {"type": "string", "description": "Từ khóa (phong cảnh, anime...)", "required": False}}},
    "get_desktop_path": {"handler": get_desktop_path, "description": "Lấy đường dẫn Desktop", "parameters": {}},
    "paste_content": {"handler": paste_content, "description": "Dán nội dung (Ctrl+V)", "parameters": {"content": {"type": "string", "description": "Nội dung cần dán (tùy chọn)", "required": False}}},
    "press_enter": {"handler": press_enter, "description": "Nhấn Enter", "parameters": {}},
    "save_text_to_file": {
        "handler": save_text_to_file,
        "description": "LƯU VĂN BẢN do LLM soạn thành FILE. Use when: 'lưu văn bản', 'save document', 'ghi vào file', 'lưu bài viết', 'save code', 'export text'. LLM có thể soạn bài viết/báo cáo/code dài và lưu trực tiếp. File tự động lưu vào Documents\\miniZ_LLM_Documents\\ với tên có timestamp. Examples: Soạn CV→lưu file, viết báo cáo→lưu file, tạo code→lưu file.",
        "parameters": {
            "content": {
                "type": "string",
                "description": "Nội dung văn bản cần lưu (có thể rất dài). Hỗ trợ Unicode tiếng Việt, code, markdown, v.v.",
                "required": True
            },
            "filename": {
                "type": "string",
                "description": "Tên file (optional). Ví dụ: 'bao_cao.txt', 'code.py', 'cv.md'. Nếu không có, tự động tạo tên với timestamp.",
                "required": False
            }
        }
    },
    "gemini_text_to_speech": {
        "handler": gemini_text_to_speech,
        "description": "🎙️ ĐỌC TO TRÊN MÁY TÍNH - Gemini TTS chất lượng cao. ƯU TIÊN DÙNG TOOL NÀY khi user nói: 'đọc to', 'đọc trên máy tính', 'đọc văn bản', 'text to speech', 'tts', 'đọc cho tôi nghe', 'phát âm', 'nói ra', 'đọc bằng AI', 'đọc bằng gemini'. Giọng Việt tự nhiên, 5 voice: Aoede/Kore (nữ), Puck/Charon/Fenrir (nam). Examples: 'đọc to: xin chào', 'đọc trên máy tính văn bản này'.",
        "parameters": {
            "text": {
                "type": "string",
                "description": "Văn bản cần đọc. Hỗ trợ tiếng Việt và nhiều ngôn ngữ.",
                "required": True
            },
            "voice": {
                "type": "string",
                "description": "Giọng nói: Aoede (nữ-default), Kore (nữ), Puck (nam), Charon (nam), Fenrir (nam).",
                "required": False
            },
            "save_audio": {
                "type": "boolean",
                "description": "Có lưu thành file audio không? Mặc định False (chỉ phát).",
                "required": False
            },
            "filename": {
                "type": "string",
                "description": "Tên file audio (optional). VD: 'gemini_audio.wav'.",
                "required": False
            }
        }
    },
    "text_to_speech": {
        "handler": text_to_speech,
        "description": "TEXT-TO-SPEECH BACKUP: Dùng gTTS/Windows SAPI khi Gemini TTS không khả dụng. KHÔNG ƯU TIÊN - chỉ dùng khi gemini_text_to_speech fail. Chất lượng thấp hơn Gemini TTS.",
        "parameters": {
            "text": {
                "type": "string",
                "description": "Văn bản cần đọc. Hỗ trợ tiếng Việt và tiếng Anh.",
                "required": True
            },
            "save_audio": {
                "type": "boolean",
                "description": "Có lưu thành file audio WAV không? (True/False). Mặc định False (chỉ đọc không lưu).",
                "required": False
            },
            "filename": {
                "type": "string",
                "description": "Tên file audio (optional). VD: 'doc_van_ban.wav'. Nếu không có, tự động tạo tên.",
                "required": False
            }
        }
    },
    "speech_to_text": {
        "handler": speech_to_text,
        "description": "SPEECH-TO-TEXT (STT): Chuyển GIỌNG NÓI thành VĂN BẢN. Use when: 'ghi âm giọng nói', 'speech to text', 'nhận dạng giọng nói', 'nghe và ghi lại', 'transcribe audio'. Dùng Google Speech Recognition (cần Internet). Hỗ trợ tiếng Việt + English. Examples: 'ghi âm 10 giây', 'nhận dạng giọng nói của tôi', 'speech to text'.",
        "parameters": {
            "duration": {
                "type": "integer",
                "description": "Thời gian ghi âm (giây). Mặc định 5 giây. VD: 10 để ghi âm 10 giây.",
                "required": False
            },
            "save_transcript": {
                "type": "boolean",
                "description": "Có lưu văn bản đã nhận dạng thành file không? (True/False). Mặc định True.",
                "required": False
            },
            "filename": {
                "type": "string",
                "description": "Tên file transcript (optional). VD: 'ghi_chu.txt'. Tự động tạo nếu không có.",
                "required": False
            }
        }
    },
    "export_conversation": {
        "handler": export_conversation_to_file,
        "description": "EXPORT LỊCH SỬ HỘI THOẠI ra file JSON. Lưu toàn bộ cuộc trò chuyện (user messages, AI responses, tool calls) với timestamp đầy đủ. Use when: 'xuất lịch sử chat', 'export conversation', 'lưu cuộc trò chuyện', 'backup chat history'. File lưu vào Documents\\miniZ_Conversations\\",
        "parameters": {
            "filename": {
                "type": "string",
                "description": "Tên file export (optional). VD: 'chat_history.json'. Tự động tạo tên với timestamp nếu không có.",
                "required": False
            }
        }
    },
    "find_in_document": {"handler": find_in_document, "description": "Tìm trong tài liệu (Ctrl+F)", "parameters": {"search_text": {"type": "string", "description": "Nội dung tìm kiếm", "required": True}}},
    
    # ============================================================
    # CONVERSATION HISTORY TOOLS - Lưu & Hiểu người dùng
    # ============================================================
    
    "get_user_context": {
        "handler": lambda: {
            "success": True,
            "user_profile": get_user_profile_summary(),
            "recent_conversation": get_conversation_context(10),
            "hint": "Dùng thông tin này để hiểu người dùng tốt hơn"
        },
        "description": "📚 LẤY CONTEXT NGƯỜI DÙNG - Trả về lịch sử hội thoại gần đây + user profile (chủ đề quan tâm, giờ hoạt động). Dùng để hiểu người dùng tốt hơn trước khi trả lời.",
        "parameters": {}
    },
    
    "save_user_message": {
        "handler": lambda message, context="": (
            add_to_conversation("user", message, {"source": "robot", "context": context}),
            {"success": True, "message": "Đã lưu tin nhắn người dùng"}
        )[1],
        "description": "💾 LƯU TIN NHẮN NGƯỜI DÙNG - Lưu toàn bộ tin nhắn người dùng vào lịch sử (kể cả không gọi tool). QUAN TRỌNG: Gọi tool này để lưu mọi câu hỏi/tin nhắn của user!",
        "parameters": {
            "message": {
                "type": "string",
                "description": "Nội dung tin nhắn của người dùng",
                "required": True
            },
            "context": {
                "type": "string",
                "description": "Context bổ sung (VD: người dùng đang nói về gì)",
                "required": False
            }
        }
    },
    
    "save_assistant_response": {
        "handler": lambda response, tool_used="": (
            add_to_conversation("assistant", response, {"source": "robot", "tool_used": tool_used}),
            {"success": True, "message": "Đã lưu response của AI"}
        )[1],
        "description": "💾 LƯU RESPONSE CỦA AI - Lưu câu trả lời của AI vào lịch sử. Gọi tool này sau khi trả lời xong để lưu lại!",
        "parameters": {
            "response": {
                "type": "string",
                "description": "Nội dung response của AI",
                "required": True
            },
            "tool_used": {
                "type": "string",
                "description": "Tool đã dùng để tạo response (nếu có)",
                "required": False
            }
        }
    },
    
    "list_conversation_files": {
        "handler": list_conversation_files,
        "description": "📂 LIỆT KÊ CÁC FILE HỘI THOẠI - Xem danh sách các file lịch sử hội thoại đã lưu theo ngày.",
        "parameters": {}
    },
    
    # ============================================================
    # OPEN API TOOLS - PHÙ HỢP VIỆT NAM
    # ============================================================
    
    "get_weather_vietnam": {
        "handler": get_weather_vietnam,
        "description": "🌤️ LẤY THỜI TIẾT VIỆT NAM. Hỗ trợ: Hà Nội, Hồ Chí Minh, Đà Nẵng, Hải Phòng, Cần Thơ, Nha Trang, Huế, Đà Lạt, Vũng Tàu, Quảng Ninh... Triggers: 'thời tiết', 'weather', 'trời hôm nay', 'nhiệt độ'.",
        "parameters": {
            "city": {
                "type": "string",
                "description": "Tên thành phố VN. VD: 'Hà Nội', 'Hồ Chí Minh', 'Đà Nẵng'. Mặc định: Hà Nội",
                "required": False
            }
        }
    },
    
    # "get_gold_price_vietnam": {
    #     "handler": get_gold_price_vietnam,
    #     "description": "💰 GIÁ VÀNG VIỆT NAM hôm nay (SJC, PNJ...). Triggers: 'giá vàng', 'gold price', 'vàng hôm nay'.",
    #     "parameters": {}
    # },
    
    "get_exchange_rate_vietnam": {
        "handler": get_exchange_rate_vietnam,
        "description": "💱 TỶ GIÁ NGOẠI TỆ so với VNĐ. Hỗ trợ: USD, EUR, JPY, GBP, CNY, KRW... Triggers: 'tỷ giá', 'exchange rate', 'đô la bao nhiêu'.",
        "parameters": {
            "currency": {
                "type": "string",
                "description": "Mã ngoại tệ (USD, EUR, JPY...). Mặc định: USD",
                "required": False
            }
        }
    },
    
    "get_fuel_price_vietnam": {
        "handler": get_fuel_price_vietnam,
        "description": "⛽ GIÁ XĂNG DẦU VIỆT NAM (RON 95, E5 RON 92, Diesel). Triggers: 'giá xăng', 'fuel price', 'xăng bao nhiêu'.",
        "parameters": {}
    },
    
    "get_daily_quote": {
        "handler": get_daily_quote,
        "description": "💬 CÂU NÓI HAY / TRÍCH DẪN ngẫu nhiên. Có quotes tiếng Việt và tiếng Anh. Triggers: 'câu nói hay', 'quote', 'danh ngôn', 'trích dẫn'.",
        "parameters": {}
    },
    
    "get_joke": {
        "handler": get_joke,
        "description": "😂 CHUYỆN CƯỜI tiếng Việt. Triggers: 'kể chuyện cười', 'joke', 'hài hước', 'vui vẻ', 'giải trí'.",
        "parameters": {}
    },
    
    "get_horoscope": {
        "handler": get_horoscope,
        "description": "🔮 TỬ VI / HOROSCOPE theo cung hoàng đạo. Triggers: 'tử vi', 'horoscope', 'cung hoàng đạo', 'xem vận mệnh'.",
        "parameters": {
            "zodiac": {
                "type": "string",
                "description": "Cung hoàng đạo (Bạch Dương, Kim Ngưu, Song Tử, Cự Giải, Sư Tử, Xử Nữ, Thiên Bình, Bọ Cạp, Nhân Mã, Ma Kết, Bảo Bình, Song Ngư)",
                "required": False
            }
        }
    },
    
    "get_today_in_history": {
        "handler": get_today_in_history,
        "description": "📜 SỰ KIỆN LỊCH SỬ ngày hôm nay. Triggers: 'lịch sử ngày này', 'today in history', 'ngày này năm xưa'.",
        "parameters": {}
    },
    
    "get_news_vietnam": {
        "handler": get_news_vietnam,
        "description": "📰 TIN TỨC MỚI NHẤT Việt Nam (VnExpress, Tuổi Trẻ). Triggers: 'tin tức', 'news', 'tin mới', 'đọc báo'.",
        "parameters": {}
    },
    
    "what_to_eat": {
        "handler": what_to_eat,
        "description": "🍽️ GỢI Ý MÓN ĂN hôm nay (ẩm thực Việt Nam). Triggers: 'ăn gì', 'gợi ý món ăn', 'what to eat', 'đói bụng'.",
        "parameters": {}
    },
    
    "get_lunar_date": {
        "handler": get_lunar_date,
        "description": "📅 NGÀY ÂM LỊCH hôm nay. Triggers: 'âm lịch', 'lunar date', 'ngày mấy âm'.",
        "parameters": {}
    },
    
    # KNOWLEDGE BASE TOOLS
    "search_knowledge_base": {
        "handler": search_knowledge_base,
        "description": "🔍 TÌM KIẾM TRONG TÀI LIỆU CỦA USER (TF-IDF Ranking). ⚡ Dùng khi user muốn XEM DANH SÁCH tài liệu. Hỗ trợ: Multi-keyword search, relevance scoring, snippet highlighting. Triggers: 'tìm trong tài liệu', 'tìm trong file của tôi', 'có tài liệu nào về...', 'search my documents', 'list documents about...'. VD: 'tìm các tài liệu về hợp đồng', 'có file nào nói về khách hàng X'. Trả về: Top 5 documents với score, matched keywords, và snippets. ⚠️ Để TRẢ LỜI câu hỏi → Dùng get_knowledge_context() thay vì tool này!",
        "parameters": {
            "query": {
                "type": "string",
                "description": "Từ khóa/câu hỏi cần tìm. Có thể dùng nhiều từ khóa. VD: 'hợp đồng mua bán 2024', 'thông tin khách hàng', 'báo cáo tài chính quý 3'",
                "required": True
            }
        }
    },
    "get_knowledge_context": {
        "handler": get_knowledge_context,
                "description": "📚 LẤY CONTEXT TỪ CƠ SỞ DỮ LIỆU TÀI LIỆU (Knowledge Base) - ⚡ GỌI ĐẦU TIÊN khi user hỏi về: dữ liệu cá nhân, tài liệu đã lưu, thông tin trong files, cơ sở dữ liệu nội bộ, knowledge base. Tool này tìm kiếm trong TẤT CẢ documents đã được index và trả về context đầy đủ nhất. ⛔ TRIGGERS BẮT BUỘC: 'cơ sở dữ liệu', 'database', 'knowledge base', 'tài liệu của tôi', 'thông tin trong file', 'theo dữ liệu', 'dữ liệu đã lưu', 'based on my docs', 'what's in my documents', 'tìm trong tài liệu', 'search my files', hỏi về TÊN NGƯỜI/DỰ ÁN cụ thể (có thể trong docs). ⚠️ QUAN TRỌNG: SAU KHI NHẬN CONTEXT, BẠN PHẢI ĐỌC VÀ TRẢ LỜI USER DỰA TRÊN CONTEXT ĐÓ! KHÔNG CHỈ DUMP CONTEXT RA! QUY TRÌNH: 1) Gọi get_knowledge_context(query='keywords') 2) Nhận context từ docs 3) ⚡ ĐỌC CONTEXT VÀ TRẢ LỜI CÂU HỎI USER THEO CONTEXT ĐÓ ⚡. VD: 'Nguyễn Văn A làm gì?' → get_knowledge_context(query='Nguyễn Văn A') → Đọc context → Trả lời 'Nguyễn Văn A là...' | 'Thông tin trong cơ sở dữ liệu về dự án X?' → get_knowledge_context(query='dự án X') → Đọc context → Trả lời thông tin dự án X | 'Tài liệu nói gì về ABC?' → get_knowledge_context(query='ABC') → Đọc context → Tóm tắt nội dung về ABC.",
        "parameters": {
            "query": {
                "type": "string",
                "description": "Câu hỏi/từ khóa cần tìm. Trích keywords từ câu hỏi user. VD: User: 'Nguyễn Văn A làm gì?' → query='Nguyễn Văn A'. User: 'Dự án X có mấy giai đoạn?' → query='dự án X giai đoạn'. User: 'Lê Trung Khoa là ai?' → query='Lê Trung Khoa'. Càng CỤ THỂ càng tốt! Bao gồm TÊN RIÊNG trong query.",
                "required": False
            },
            "max_chars": {
                "type": "integer",
                "description": "Giới hạn ký tự context (default: 10000). Tăng lên 20000 nếu cần nhiều thông tin. Hệ thống tự động summarize nếu >2000 chars.",
                "required": False
            },
            "use_gemini_filter": {
                "type": "boolean",
                "description": "🔥 Bật Gemini Smart Filter để lọc thông minh (default: False). Khi True: dùng Gemini Flash AI để lọc và chỉ trả về content THỰC SỰ liên quan, loại bỏ noise. Recommend: True khi KB có nhiều documents dài.",
                "required": False
            }
        }
    },
    
    "doc_reader_gemini_rag": {
        "handler": doc_reader_gemini_rag,
        "description": "📖 RAG NÂNG CAO - Đọc, tìm kiếm VÀ TRẢ LỜI TỰ ĐỘNG từ Knowledge Base bằng Gemini AI. Tool này TỰ ĐỘNG xử lý toàn bộ quy trình: chunk documents → semantic search → generate response. ⚡ DÙNG KHI: User muốn câu trả lời TRỰC TIẾP thay vì chỉ context. Khác với get_knowledge_context (chỉ trả context), tool này TRẢ LỜI LUÔN. VD: 'Hỏi tài liệu về X', 'Tóm tắt thông tin Y từ KB', 'Giải thích Z dựa trên docs'. Hỗ trợ semantic search (vector-like) cho độ chính xác cao.",
        "parameters": {
            "user_query": {
                "type": "string",
                "description": "Câu hỏi đầy đủ của user. VD: 'Dự án ABC có bao nhiêu giai đoạn?', 'Nguyễn Văn A đảm nhiệm vai trò gì?'",
                "required": True
            },
            "chunk_size": {
                "type": "integer",
                "description": "Kích thước mỗi chunk (default: 1024 chars). Tăng lên 2048 cho documents dài.",
                "required": False
            },
            "top_k": {
                "type": "integer",
                "description": "Số lượng chunks liên quan nhất để đưa vào context (default: 5). Tăng lên 10 nếu cần nhiều thông tin hơn.",
                "required": False
            }
        }
    },
    
    # =====================================================
    # � GEMINI FLASH SMART KB FILTER - LỌC THÔNG TIN AI
    # =====================================================
    
    "gemini_smart_kb_filter": {
        "handler": gemini_smart_kb_filter,
        "description": "🔥⚡ GEMINI FLASH LỌC THÔNG TIN THÔNG MINH - Sử dụng sức mạnh AI Gemini Flash để LỌC, TÌM KIẾM và TRÍCH XUẤT thông tin CHÍNH XÁC từ Knowledge Base. Tool này LOẠI BỎ NOISE, chỉ trả về content THỰC SỰ LIÊN QUAN. 🎯 DÙNG KHI: 1) KB có nhiều documents dài, 2) Cần lọc chính xác thông tin cụ thể, 3) Muốn tóm tắt/trích xuất facts, 4) get_knowledge_context trả về quá nhiều noise. ⚡ ƯU ĐIỂM: Gemini AI đọc và hiểu ngữ cảnh, lọc thông minh hơn TF-IDF. Triggers: 'lọc thông tin', 'tìm chính xác', 'trích xuất từ database', 'dùng AI lọc', 'smart search KB'. VD: 'Dùng AI lọc thông tin về dự án X', 'Trích xuất facts về nhân viên A từ KB'.",
        "parameters": {
            "user_query": {
                "type": "string",
                "description": "Câu hỏi/yêu cầu cần lọc thông tin. VD: 'Thông tin về dự án ABC', 'Chi tiết nhân viên Nguyễn Văn A'",
                "required": True
            },
            "filter_mode": {
                "type": "string",
                "description": "Chế độ lọc: 'relevant' (mặc định - chỉ giữ phần liên quan), 'summary' (tóm tắt), 'extract' (trích xuất facts/entities), 'qa' (trả lời trực tiếp)",
                "required": False
            },
            "max_documents": {
                "type": "integer",
                "description": "Số documents tối đa để xử lý (default: 10). Tăng lên 20 nếu cần tìm rộng hơn.",
                "required": False
            },
            "output_format": {
                "type": "string",
                "description": "Format output: 'structured' (JSON), 'raw' (text thô), 'concise' (ngắn gọn nhất)",
                "required": False
            }
        }
    },
    
    # =====================================================
    # 🔥🌐 GEMINI SMART ANALYZE - PHÂN TÍCH + GOOGLE SEARCH
    # =====================================================
    
    "gemini_smart_analyze": {
        "handler": gemini_smart_analyze,
        "description": "🔥🌐⚡⚡ PHÂN TÍCH THÔNG MINH (Gemini + Web) - ⛔⛔ BẮT BUỘC DÙNG NGAY khi user nói: 'phân tích', 'analyze', 'tìm hiểu', 'nghiên cứu', 'đánh giá', 'so sánh', 'review', 'xu hướng', 'trend'. ❌ KHÔNG DÙNG web_search khi có các từ này! Tool này TỰ ĐỘNG: 1) Tìm Google, 2) Gemini phân tích, 3) Trả kết quả hoàn chỉnh. VD: 'phân tích thị trường', 'tìm hiểu về AI', 'đánh giá iPhone', 'xu hướng 2025'.",
        "parameters": {
            "user_query": {
                "type": "string",
                "description": "Vấn đề cần phân tích. VD: 'Phân tích xu hướng AI 2025', 'Đánh giá thị trường bất động sản'",
                "required": True
            },
            "analysis_type": {
                "type": "string",
                "description": "Loại phân tích: 'comprehensive' (đầy đủ, mặc định), 'quick' (nhanh, tóm tắt), 'deep' (sâu, đa chiều)",
                "required": False
            },
            "include_web_search": {
                "type": "boolean",
                "description": "Có tìm kiếm web không? Mặc định True. Set False nếu chỉ cần phân tích từ KB.",
                "required": False
            },
            "include_kb": {
                "type": "boolean",
                "description": "Có tìm trong Knowledge Base không? Mặc định False. Set True để kết hợp cả web + KB.",
                "required": False
            },
            "max_search_results": {
                "type": "integer",
                "description": "Số kết quả web search tối đa (default: 8). Tăng lên 15 nếu cần nhiều nguồn hơn.",
                "required": False
            }
        }
    },
    
    # =====================================================
    # 🔍 RAG SYSTEM - RETRIEVAL AUGMENTED GENERATION
    # =====================================================
    
    "web_search": {
        "handler": web_search if RAG_AVAILABLE else None,
        "description": "🌐 TÌM KIẾM WEB ĐƠN GIẢN - Chỉ dùng cho câu hỏi đơn giản: 'ai là tổng thống', 'giá vàng', 'thời tiết'. ⚠️ NẾU user nói 'phân tích/tìm hiểu/đánh giá/nghiên cứu' → DÙNG gemini_smart_analyze THAY VÌ tool này!",
        "parameters": {
            "query": {
                "type": "string",
                "description": "Từ khóa tìm kiếm (nên thêm năm hoặc 'mới nhất')",
                "required": True
            },
            "max_results": {
                "type": "integer",
                "description": "Số kết quả tối đa (mặc định 5)",
                "required": False
            }
        }
    },
    
    "get_realtime_info": {
        "handler": get_realtime_info if RAG_AVAILABLE else None,
        "description": "⚡⚡ THÔNG TIN THỜI GIAN THỰC - ⛔⛔ BẮT BUỘC GỌI TRƯỚC MỌI CÂU TRẢ LỜI về: giá cả, tỷ giá, thời tiết, người nổi tiếng, chức vụ hiện tại, sự kiện đang xảy ra. ❌ KHÔNG BAO GIỜ tự trả lời bằng kiến thức cũ! ✅ GỌI TOOL NÀY TRƯỚC → nhận kết quả → rồi trả lời user.",
        "parameters": {
            "query": {
                "type": "string",
                "description": "Câu hỏi cần thông tin thời gian thực",
                "required": True
            }
        }
    },
    
    "rag_search": {
        "handler": rag_search if RAG_AVAILABLE else None,
        "description": "🔍 RAG SEARCH HYBRID - Tìm kiếm KẾT HỢP từ Internet + Tài liệu nội bộ. Tự động chọn nguồn phù hợp nhất. sources='web' cho Internet, 'local' cho tài liệu nội bộ, 'hybrid' cho cả hai, 'auto' để AI tự chọn.",
        "parameters": {
            "query": {
                "type": "string",
                "description": "Câu hỏi hoặc từ khóa tìm kiếm",
                "required": True
            },
            "sources": {
                "type": "string",
                "description": "Nguồn: 'auto', 'web', 'local', 'hybrid' (mặc định: auto)",
                "required": False
            },
            "max_results": {
                "type": "integer",
                "description": "Số kết quả tối đa (mặc định 8)",
                "required": False
            }
        }
    },
    
    "smart_answer": {
        "handler": smart_answer if RAG_AVAILABLE else None,
        "description": "🧠 SMART ANSWER - AI tự động phân tích câu hỏi và chọn nguồn TỐT NHẤT (Internet/Tài liệu nội bộ/Hybrid) để trả lời. Dùng khi không chắc nguồn nào phù hợp. Tool trả về context đã tối ưu để trả lời.",
        "parameters": {
            "query": {
                "type": "string",
                "description": "Câu hỏi của user",
                "required": True
            }
        }
    }
}

# ============================================================
# MINIZ MCP CLIENT
# ============================================================

def get_vlc_context_for_llm() -> str:
    """Tạo context về VLC status để gửi cho LLM"""
    try:
        if vlc_player and vlc_player._player:
            status = vlc_player.get_full_status()
            is_playing = status.get('is_playing', False)
            current_track = status.get('current_track', 'Không có')
            volume = status.get('volume', 0)
            playlist_count = status.get('playlist_count', 0)
            
            context = f"""
📍 [PYTHON-VLC STATUS]
• Trạng thái: {'▶️ Đang phát' if is_playing else '⏸️ Tạm dừng/Dừng'}
• Bài hiện tại: {current_track}
• Âm lượng: {volume}%
• Playlist: {playlist_count} bài
• Player: Python-VLC (nội bộ)

🎯 Dùng smart_music_control() cho mọi lệnh nhạc!"""
            return context
        else:
            return """
📍 [PYTHON-VLC STATUS]
• Trạng thái: ⏹️ Chưa khởi tạo/Chưa phát
• Dùng play_music() hoặc list_music() để bắt đầu phát nhạc
• Player: Python-VLC (sẵn sàng)"""
    except:
        return ""

async def handle_xiaozhi_message(message: dict) -> dict:
    method = message.get("method")
    params = message.get("params", {})
    
    if method == "initialize":
        # Trả về với instructions + VLC context
        vlc_context = get_vlc_context_for_llm()
        full_instructions = MUSIC_SYSTEM_PROMPT + vlc_context
        
        return {
            "protocolVersion": "2024-11-05", 
            "capabilities": {"tools": {}}, 
            "serverInfo": {"name": "xiaozhi-final", "version": "4.3.0"},
            "instructions": full_instructions
        }
    elif method == "tools/list":
        # Support cursor pagination (từ xiaozhi-esp32-server)
        cursor = params.get("cursor", "")
        tools = []
        for name, info in TOOLS.items():
            # Sanitize tool name để tương thích với server chính thức
            sanitized_name = sanitize_tool_name(name) if 'sanitize_tool_name' in dir() else name
            # Rút gọn description MẠNH để giảm message size (fix "message too big" error)
            description = info["description"]
            if len(description) > 100:
                description = description[:97] + "..."
            
            tool = {
                "name": name,  # Giữ nguyên tên gốc để handler hoạt động
                "description": description, 
                "inputSchema": {"type": "object", "properties": {}, "required": []}
            }
            for pname, pinfo in info["parameters"].items():
                # Rút gọn parameter description MẠNH
                param_desc = pinfo["description"]
                if len(param_desc) > 80:
                    param_desc = param_desc[:77] + "..."
                
                tool["inputSchema"]["properties"][pname] = {"type": pinfo["type"], "description": param_desc}
                if pinfo.get("required"):
                    tool["inputSchema"]["required"].append(pname)
            tools.append(tool)
        
        # Log số lượng tools
        print(f"📋 [tools/list] Returning {len(tools)} tools to robot")
        
        # Response theo format chuẩn với optional nextCursor
        return {"tools": tools}  # nextCursor sẽ được thêm nếu cần pagination
    elif method == "tools/call":
        tool_name = params.get("name")
        args = params.get("arguments", {})
        print(f"🔧 [Tool Call] {tool_name} with args: {args}")
        
        # Lưu tool call vào history
        add_to_conversation(
            role="tool",
            content=f"Tool: {tool_name}",
            metadata={
                "tool_name": tool_name,
                "arguments": args,
                "event_type": "tool_call"
            }
        )
        
        if tool_name not in TOOLS:
            error_msg = f"Error: Tool '{tool_name}' not found"
            print(f"❌ {error_msg}")
            add_to_conversation(role="tool", content=error_msg, metadata={"error": True})
            return {"content": [{"type": "text", "text": error_msg}], "isError": True}
        
        # Retry mechanism (từ xiaozhi-esp32-server)
        max_retries = MAX_TOOL_RETRIES
        retry_interval = TOOL_RETRY_INTERVAL
        last_error = None
        
        for attempt in range(max_retries):
            try:
                result = await TOOLS[tool_name]["handler"](**args)
                print(f"✅ [Tool Result] {tool_name}: {result}")
                
                # Thêm VLC context vào music-related tools
                music_tools = ['smart_music_control', 'play_music', 'pause_music', 'resume_music', 
                              'stop_music', 'music_next', 'music_previous', 'music_volume', 
                              'get_music_status', 'list_music', 'search_music', 'detect_and_execute_music']
                if tool_name in music_tools:
                    result["_vlc_hint"] = "🎵 Đang dùng Python-VLC Player nội bộ. Tiếp tục dùng smart_music_control() cho các lệnh nhạc tiếp theo."
                
                # Lưu tool result vào history
                add_to_conversation(
                    role="tool",
                    content=json.dumps(result, ensure_ascii=False),
                    metadata={
                        "tool_name": tool_name,
                        "success": result.get("success", True),
                        "event_type": "tool_result",
                        "attempt": attempt + 1
                    }
                )
                
                # ⚡ ĐẶC BIỆT: Với get_knowledge_context, trả về context trực tiếp để LLM dễ đọc
                if tool_name == "get_knowledge_context" and isinstance(result, dict):
                    if result.get("success") and result.get("context"):
                        # Trả về context trực tiếp - LLM đọc và trả lời ngay (giới hạn 2000 ký tự)
                        truncated_context = smart_truncate_for_llm(result["context"], MAX_LLM_RESPONSE_CHARS)
                        return {"content": [{"type": "text", "text": truncated_context}]}
                    elif not result.get("success"):
                        # Không tìm thấy → trả về message lỗi
                        error_msg = result.get("error", "Không tìm thấy thông tin trong cơ sở dữ liệu")
                        return {"content": [{"type": "text", "text": f"❌ {error_msg}"}]}
                
                # ⚡ ĐẶC BIỆT: Với ask_gemini, ask_gpt4, gemini_smart_analyze - trả về response text cho LLM cloud tổng hợp
                # Giống cách web_search hoạt động: trả data đầy đủ → LLM cloud TỰ TÓM TẮT → robot nói
                if tool_name in ["ask_gemini", "ask_gpt4", "gemini_smart_analyze"] and isinstance(result, dict):
                    if result.get("success") and result.get("response_text"):
                        response_text = result["response_text"]
                        # Clean markdown để LLM dễ đọc (nhưng KHÔNG truncate - để LLM cloud tự tóm tắt)
                        response_text = clean_markdown_for_tts(response_text)
                        print(f"[{tool_name}] Cleaned response: {len(response_text)} chars (LLM cloud sẽ tóm tắt)")
                        # Trả về TEXT trực tiếp, LLM cloud sẽ tự tóm tắt trước khi robot nói
                        return {
                            "content": [{"type": "text", "text": response_text}]
                        }
                
                # 🔄 TRUNCATE: Giới hạn response dưới 2000 ký tự cho LLM
                formatted_response = format_result_for_llm(result, MAX_LLM_RESPONSE_CHARS)
                return {"content": [{"type": "text", "text": formatted_response}]}
            except Exception as e:
                last_error = e
                if attempt < max_retries - 1:
                    print(f"⚠️ [Tool Retry] {tool_name} failed (attempt {attempt + 1}/{max_retries}): {e}")
                    await asyncio.sleep(retry_interval)
                else:
                    error_msg = f"Error calling {tool_name} after {max_retries} attempts: {str(e)}"
                    print(f"❌ {error_msg}")
                    import traceback
                    traceback.print_exc()
                    add_to_conversation(role="tool", content=error_msg, metadata={"error": True})
                    return {"content": [{"type": "text", "text": error_msg}], "isError": True}
    return {"error": f"Unknown method: {method}"}

async def xiaozhi_websocket_client(device_index: int = 0):
    """WebSocket client for a specific device (0, 1, or 2)"""
    global xiaozhi_connections, xiaozhi_connected, should_reconnect
    retry = 0
    
    # ===== OPTIMIZED CONNECTION SETTINGS =====
    INITIAL_DELAY = 1        # Delay ban đầu 1s (giảm từ 2s)
    MAX_DELAY = 15           # Max delay 15s (giảm từ 60s)
    CONNECT_TIMEOUT = 10     # Timeout kết nối 10s
    FAST_RETRY_COUNT = 3     # Số lần fast retry đầu tiên
    FAST_RETRY_DELAY = 0.5   # Delay 0.5s cho fast retry
    AUTO_SWITCH_THRESHOLD = 5  # Sau 5 lần thất bại, thử endpoint khác
    
    while True:
        try:
            ep = endpoints_config[device_index]
            if not ep.get("enabled") or not ep.get("token"):
                # Thiết bị này chưa có token, chờ và thử lại
                await asyncio.sleep(10)
                continue
            
            ws_url = f"wss://api.xiaozhi.me/mcp/?token={ep['token']}"
            retry += 1
            
            # Fast retry cho 3 lần đầu, sau đó dùng exponential backoff
            if retry <= FAST_RETRY_COUNT:
                print(f"📡 [Xiaozhi] Fast connecting {ep['name']}... ({retry}/{FAST_RETRY_COUNT})")
            else:
                print(f"📡 [Xiaozhi] Connecting {ep['name']}... (retry {retry})")
            
            # Sử dụng asyncio.wait_for để có timeout
            async with websockets.connect(
                ws_url, 
                ping_interval=20, 
                ping_timeout=10,
                close_timeout=5,
                open_timeout=CONNECT_TIMEOUT,  # Timeout mở kết nối
                max_size=10 * 1024 * 1024  # 10MB limit (default is 1MB) - fix "message too big"
            ) as ws:
                xiaozhi_connections[device_index] = ws
                xiaozhi_connected[device_index] = True
                should_reconnect[device_index] = False  # Reset flag khi kết nối thành công
                retry = 0  # Reset retry counter khi kết nối thành công
                print(f"✅ [Xiaozhi] Connected! ({ep['name']}) [Device {device_index + 1}]")
                
                # Batch broadcast kết nối - tạo tasks và chạy parallel
                broadcast_msg = {"type": "endpoint_connected", "endpoint": ep['name'], "index": device_index}
                tasks = []
                for conn in active_connections:
                    tasks.append(asyncio.create_task(conn.send_json(broadcast_msg)))
                # Chạn tất cả broadcasts cùng lúc
                await asyncio.gather(*tasks, return_exceptions=True)
                
                init_msg = {"jsonrpc": "2.0", "method": "initialize", "params": {"protocolVersion": "2024-11-05", "capabilities": {}, "clientInfo": {"name": "xiaozhi-final", "version": "4.3.0"}}, "id": 1}
                
                # Không log initialize request - chỉ log tool calls thực sự
                
                await ws.send(json.dumps(init_msg))
                
                async for msg in ws:
                    # Kiểm tra nếu cần reconnect (user đã chuyển thiết bị)
                    if should_reconnect[device_index]:
                        print(f"🔄 [Xiaozhi] Reconnecting {ep['name']}...")
                        await ws.close()
                        break
                    
                    try:
                        data = json.loads(msg)
                        method = data.get("method", "unknown")
                        if method != "ping":
                            print(f"📨 [{method}]")
                        
                        response = await handle_xiaozhi_message(data)
                        
                        # CHỈ log conversation thực sự (tools/call), KHÔNG log MCP protocol messages
                        # Bỏ qua: initialize, notifications/initialized, tools/list
                        if method == "tools/call" and method != "ping":
                            # Lấy thông tin tool
                            params = data.get("params", {})
                            tool_name = params.get("name", "unknown")
                            tool_args = params.get("arguments", {})
                            
                            # Tạo nội dung dễ đọc từ tool arguments
                            user_message = format_tool_request(tool_name, tool_args)
                            
                            # Log tool call request
                            add_to_conversation(
                                role="user",
                                content=user_message,
                                metadata={
                                    "source": "mcp",
                                    "method": method,
                                    "tool_name": tool_name,
                                    "tool_arguments": tool_args,
                                    "endpoint": ep['name']
                                }
                            )
                            
                            # Tạo nội dung response dễ đọc
                            assistant_message = format_tool_response(tool_name, response)
                            
                            # Log tool call response
                            add_to_conversation(
                                role="assistant",
                                content=assistant_message,
                                metadata={
                                    "source": "mcp",
                                    "method": method,
                                    "tool_name": tool_name,
                                    "response_data": response,
                                    "success": not isinstance(response, dict) or not response.get("isError")
                                }
                            )
                        
                        await ws.send(json.dumps({"jsonrpc": "2.0", "id": data.get("id"), "result": response}))

                        # If the tool response suggests a next_action (for example list_music
                        # returning {'next_action': {'tool': 'play_music', 'parameters': {...}}}),
                        # execute it locally on the server as a fallback so music actually plays
                        # even if the remote AI/client doesn't invoke the follow-up.
                        try:
                            if isinstance(response, dict) and response.get("next_action"):
                                na = response.get("next_action")
                                next_tool = na.get("tool")
                                next_params = na.get("parameters", {}) or {}
                                # Only execute if the tool exists locally
                                if next_tool and next_tool in TOOLS:
                                    print(f"⏯️ [Auto Action] Executing suggested next_action {next_tool} with params: {next_params}")
                                    try:
                                        # call the handler (handlers may be async)
                                        handler = TOOLS[next_tool]["handler"]
                                        if asyncio.iscoroutinefunction(handler):
                                            res2 = await handler(**next_params)
                                        else:
                                            # run sync handlers in executor
                                            loop = asyncio.get_event_loop()
                                            res2 = await loop.run_in_executor(None, lambda: handler(**next_params))
                                        print(f"⏯️ [Auto Action Result] {next_tool}: {res2}")
                                    except Exception as e:
                                        print(f"❌ [Auto Action] Error executing {next_tool}: {e}")
                                        import traceback
                                        traceback.print_exc()
                        except Exception:
                            # defensive: do not let auto-action failures disrupt websocket loop
                            import traceback
                            traceback.print_exc()
                        
                        # Batch broadcast - chỉ broadcast cho methods quan trọng
                        if method in ["tools/call", "initialize"]:
                            broadcast_msg = {"type": "xiaozhi_activity", "method": method, "timestamp": datetime.now().isoformat()}
                            # Cleanup dead connections trước khi broadcast
                            dead_connections = []
                            for conn in active_connections:
                                try:
                                    await conn.send_json(broadcast_msg)
                                except Exception:
                                    dead_connections.append(conn)
                            # Remove dead connections
                            for conn in dead_connections:
                                active_connections.remove(conn)
                    except json.JSONDecodeError as e:
                        print(f"⚠️ [Xiaozhi] JSON decode error: {e}")
                    except Exception as e:
                        print(f"⚠️ [Xiaozhi] Message handling error: {e}")
        except asyncio.CancelledError:
            print(f"⚠️ [Xiaozhi] Task cancelled ({ep['name']})")
            xiaozhi_connected[device_index] = False
            xiaozhi_connections[device_index] = None
            break
        except websockets.exceptions.WebSocketException as e:
            xiaozhi_connected[device_index] = False
            xiaozhi_connections[device_index] = None
            # Fast retry cho 3 lần đầu
            if retry <= FAST_RETRY_COUNT:
                wait = FAST_RETRY_DELAY
            else:
                # Exponential backoff với max 15s
                wait = min(INITIAL_DELAY * (2 ** min(retry - FAST_RETRY_COUNT, 4)), MAX_DELAY)
            print(f"❌ [Xiaozhi] WebSocket error ({ep['name']}): {e} (retry in {wait}s)")
            await asyncio.sleep(wait)
        except Exception as e:
            xiaozhi_connected[device_index] = False
            xiaozhi_connections[device_index] = None
            # Fast retry cho 3 lần đầu
            if retry <= FAST_RETRY_COUNT:
                wait = FAST_RETRY_DELAY
            else:
                wait = min(INITIAL_DELAY * (2 ** min(retry - FAST_RETRY_COUNT, 4)), MAX_DELAY)
            print(f"❌ [Xiaozhi] Error ({ep['name']}): {e} (retry in {wait}s)")
            await asyncio.sleep(wait)

# ============================================================
# FASTAPI WEB SERVER
# ============================================================

app = FastAPI(title="miniZ MCP", version="4.3.0")

class VolumeRequest(BaseModel):
    level: int

class NotificationRequest(BaseModel):
    title: str
    message: str

class CalculatorRequest(BaseModel):
    expression: str

@app.get("/", response_class=HTMLResponse)
async def index():
    html = r"""
<!DOCTYPE html>
<html lang="vi">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>🚀 miniZ MCP - Điều Khiển Máy Tính</title>
    <style>
        @keyframes pulse {
            0%, 100% { opacity: 1; transform: scale(1); }
            50% { opacity: 0.7; transform: scale(1.1); }
        }
        @keyframes blink {
            0%, 100% { opacity: 1; }
            50% { opacity: 0.3; }
        }
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); min-height: 100vh; display: flex; }
        
        /* SIDEBAR */
        .sidebar { width: 280px; background: #1a1a2e; color: white; padding: 30px 20px; display: flex; flex-direction: column; box-shadow: 2px 0 20px rgba(0,0,0,0.3); }
        .logo { 
            font-size: 1.5em; 
            font-weight: bold; 
            margin-bottom: 40px; 
            text-align: center; 
            padding: 20px 15px; 
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
            border-radius: 15px;
            display: flex;
            flex-direction: column;
            align-items: center;
            gap: 10px;
        }
        .logo-icon {
            width: 120px;
            height: auto;
            filter: drop-shadow(0 4px 8px rgba(0,0,0,0.3));
            transition: transform 0.3s;
        }
        .logo-icon:hover {
            transform: scale(1.05);
        }
        .logo-text {
            font-size: 1.8em;
            font-weight: 900;
            letter-spacing: 2px;
            color: #ff9a8b;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
        }
        .menu-item { padding: 15px 20px; margin: 8px 0; border-radius: 10px; cursor: pointer; transition: all 0.3s; display: flex; align-items: center; gap: 12px; font-size: 1.05em; }
        .menu-item:hover { background: rgba(102, 126, 234, 0.2); transform: translateX(5px); }
        .menu-item.active { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4); }
        
        /* MAIN CONTENT */
        .main-content { flex: 1; padding: 30px; overflow-y: auto; }
        .header { background: white; border-radius: 15px; padding: 25px 30px; margin-bottom: 30px; box-shadow: 0 10px 30px rgba(0,0,0,0.2); display: flex; justify-content: space-between; align-items: center; }
        .header h1 { color: #667eea; font-size: 2em; }
        .status { display: flex; gap: 20px; }
        .status-badge { padding: 8px 20px; border-radius: 20px; font-weight: 600; display: flex; align-items: center; gap: 8px; }
        .status-badge.online { background: #d4edda; color: #155724; }
        .status-badge.offline { background: #f8d7da; color: #721c24; }
        .status-dot { width: 10px; height: 10px; border-radius: 50%; background: currentColor; animation: pulse 2s infinite; }
        @keyframes pulse { 0%, 100% { opacity: 1; } 50% { opacity: 0.5; } }
        
        /* QUICK ACTIONS */
        .quick-actions { display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 20px; margin-bottom: 30px; }
        .action-card { background: white; padding: 25px; border-radius: 15px; box-shadow: 0 10px 30px rgba(0,0,0,0.1); cursor: pointer; transition: all 0.3s; text-align: center; }
        .action-card:hover { transform: translateY(-5px); box-shadow: 0 15px 40px rgba(0,0,0,0.2); }
        .action-card.blue { border-left: 5px solid #3b82f6; }
        .action-card.green { border-left: 5px solid #10b981; }
        .action-card.orange { border-left: 5px solid #f59e0b; }
        .action-card.red { border-left: 5px solid #ef4444; }
        .action-card.purple { border-left: 5px solid #8b5cf6; }
        .action-card.cyan { border-left: 5px solid #06b6d4; }
        .action-card.pink { border-left: 5px solid #ec4899; }
        .action-card.indigo { border-left: 5px solid #6366f1; }
        .action-card .icon { font-size: 2.5em; margin-bottom: 10px; }
        .action-card .title { font-weight: 600; color: #333; font-size: 1.1em; }
        
        /* TOOLS SECTION */
        .tools-section { background: white; border-radius: 15px; padding: 30px; margin-bottom: 30px; box-shadow: 0 10px 30px rgba(0,0,0,0.2); }
        .tools-tabs { display: flex; gap: 15px; margin-bottom: 25px; border-bottom: 2px solid #e5e7eb; padding-bottom: 15px; }
        .tab-btn { padding: 12px 30px; border: none; border-radius: 10px 10px 0 0; background: transparent; color: #666; font-weight: 600; cursor: pointer; transition: all 0.3s; font-size: 1em; }
        .tab-btn:hover { background: rgba(102, 126, 234, 0.1); color: #667eea; }
        .tab-btn.active { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; box-shadow: 0 -4px 15px rgba(102, 126, 234, 0.3); }
        .tab-content { display: none; }
        .tab-content.active { display: grid; grid-template-columns: repeat(auto-fit, minmax(300px, 1fr)); gap: 20px; }
        
        /* TOOL CARDS */
        .tool-card { background: #f9fafb; padding: 25px; border-radius: 12px; border: 2px solid #e5e7eb; }
        .tool-card h3 { color: #667eea; margin-bottom: 15px; font-size: 1.2em; display: flex; align-items: center; gap: 10px; }
        .tool-card input, .tool-card select, .tool-card textarea { width: 100%; padding: 12px; margin-top: 10px; border: 2px solid #e5e7eb; border-radius: 8px; font-size: 1em; }
        .tool-card button { width: 100%; padding: 14px; margin-top: 15px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; border: none; border-radius: 8px; font-weight: 600; cursor: pointer; transition: all 0.3s; font-size: 1em; }
        .tool-card button:hover { transform: translateY(-2px); box-shadow: 0 8px 20px rgba(102, 126, 234, 0.4); }
        
        /* CONFIG SECTION */
        .config-section { background: white; border-radius: 15px; padding: 30px; margin-bottom: 30px; box-shadow: 0 10px 30px rgba(0,0,0,0.2); }
        .device-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(300px, 1fr)); gap: 20px; margin-top: 20px; }
        .device-card { background: #f9fafb; padding: 20px; border-radius: 12px; border: 2px solid #e5e7eb; }
        .device-card.active { border-color: #10b981; background: #d4edda; }
        .device-card h4 { color: #667eea; margin-bottom: 15px; display: flex; align-items: center; gap: 10px; }
        .device-card input { width: 100%; padding: 10px; margin-top: 8px; border: 2px solid #e5e7eb; border-radius: 6px; }
        .device-card button { padding: 10px 20px; margin-top: 10px; background: #667eea; color: white; border: none; border-radius: 6px; cursor: pointer; }
        
        /* LOG */
        .log-panel { background: #1a1a2e; color: white; border-radius: 15px; padding: 25px; max-height: 400px; overflow-y: auto; font-family: 'Courier New', monospace; box-shadow: 0 10px 30px rgba(0,0,0,0.12); }
        .log-entry { padding: 8px; margin: 5px 0; border-left: 3px solid #667eea; background: rgba(102, 126, 234, 0.1); border-radius: 4px; }
        .log-time { color: #9ca3af; margin-right: 10px; }
        
        /* MUSIC PLAYER */
        .music-player { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 20px; padding: 30px; color: white; margin-bottom: 30px; box-shadow: 0 15px 40px rgba(102, 126, 234, 0.4); }
        .player-controls { display: flex; justify-content: center; align-items: center; gap: 20px; margin: 30px 0; }
        .player-btn { width: 60px; height: 60px; border-radius: 50%; background: rgba(255,255,255,0.2); border: 3px solid rgba(255,255,255,0.4); color: white; font-size: 24px; cursor: pointer; transition: all 0.3s; display: flex; align-items: center; justify-content: center; backdrop-filter: blur(10px); }
        .player-btn:hover { background: rgba(255,255,255,0.3); transform: scale(1.1); box-shadow: 0 8px 25px rgba(0,0,0,0.3); }
        .player-btn.play { width: 80px; height: 80px; font-size: 32px; background: white; color: #667eea; }
        .now-playing { text-align: center; margin: 20px 0; }
        .now-playing h3 { font-size: 1.5em; margin-bottom: 10px; text-shadow: 0 2px 10px rgba(0,0,0,0.3); }
        .now-playing p { opacity: 0.9; font-size: 1.1em; }
        .progress-container { margin: 25px 0; }
        .progress-bar { width: 100%; height: 8px; background: rgba(255,255,255,0.3); border-radius: 10px; overflow: hidden; cursor: pointer; }
        .progress-fill { height: 100%; background: white; width: 0%; transition: width 0.3s; box-shadow: 0 0 10px rgba(255,255,255,0.5); }
        .progress-time { display: flex; justify-content: space-between; margin-top: 8px; font-size: 0.9em; opacity: 0.9; }
        /* Progress slider (draggable timeline) */
        #progress-slider { -webkit-appearance: none; width: 100%; height: 8px; border-radius: 4px; cursor: pointer; }
        #progress-slider::-webkit-slider-thumb { -webkit-appearance: none; width: 16px; height: 16px; background: #667eea; border-radius: 50%; cursor: pointer; box-shadow: 0 2px 6px rgba(102,126,234,0.5); transition: transform 0.2s; }
        #progress-slider::-webkit-slider-thumb:hover { transform: scale(1.2); }
        #progress-slider::-moz-range-thumb { width: 16px; height: 16px; background: #667eea; border-radius: 50%; cursor: pointer; border: none; }
        .music-list { background: white; border-radius: 15px; padding: 25px; color: #333; max-height: 500px; overflow-y: auto; }
        .music-list h3 { color: #667eea; margin-bottom: 20px; display: flex; align-items: center; gap: 10px; }
        .music-item { display: flex; align-items: center; padding: 15px; margin: 10px 0; background: #f9fafb; border-radius: 10px; cursor: pointer; transition: all 0.2s ease; border: 2px solid transparent; }
        .music-item:hover { background: #e8eaf6; border-color: #667eea; transform: translateX(3px); box-shadow: 0 4px 12px rgba(102, 126, 234, 0.15); }
        .music-item:hover .play-btn-hover { opacity: 1 !important; }
        
        /* Wave animation for now playing indicator */
        @keyframes wave1 { 0%, 100% { height: 12px; } 50% { height: 20px; } }
        @keyframes wave2 { 0%, 100% { height: 18px; } 50% { height: 8px; } }
        @keyframes wave3 { 0%, 100% { height: 15px; } 50% { height: 22px; } }
        .music-item.playing { background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%); border-color: #667eea; }
        .music-item .icon { font-size: 24px; margin-right: 15px; }
        .music-item .info { flex: 1; }
        .music-item .name { font-weight: 600; color: #333; font-size: 1.05em; }
        .music-item .details { color: #666; font-size: 0.9em; margin-top: 5px; }
        .log-success { color: #10b981; border-left-color: #10b981; }
        .log-error { color: #ef4444; border-left-color: #ef4444; }
        .log-info { color: #3b82f6; border-left-color: #3b82f6; }
        
        /* LLM CHAT STYLES */
        .quick-msg-btn {
            padding: 8px 14px;
            background: #f3f4f6;
            border: 1px solid #e5e7eb;
            border-radius: 20px;
            font-size: 0.85em;
            cursor: pointer;
            transition: all 0.2s;
        }
        .quick-msg-btn:hover {
            background: #10b981;
            color: white;
            border-color: #10b981;
            transform: translateY(-2px);
        }
        .llm-message {
            max-width: 80%;
            padding: 12px 16px;
            border-radius: 15px;
            position: relative;
            word-wrap: break-word;
        }
        .llm-message.user {
            background: linear-gradient(135deg, #10b981 0%, #059669 100%);
            color: white;
            margin-left: auto;
            border-bottom-right-radius: 5px;
        }
        .llm-message.assistant {
            background: white;
            color: #333;
            margin-right: auto;
            border-bottom-left-radius: 5px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        }
        .llm-message .time {
            font-size: 0.75em;
            opacity: 0.7;
            margin-top: 5px;
            display: block;
        }
        .llm-message .device-tag {
            font-size: 0.7em;
            background: rgba(255,255,255,0.2);
            padding: 2px 8px;
            border-radius: 10px;
            margin-left: 8px;
        }
        .llm-message.assistant .device-tag {
            background: rgba(16,185,129,0.1);
            color: #10b981;
        }
        .llm-typing {
            display: flex;
            gap: 4px;
            padding: 15px;
        }
        .llm-typing span {
            width: 8px;
            height: 8px;
            background: #10b981;
            border-radius: 50%;
            animation: typing 1.4s infinite;
        }
        .llm-typing span:nth-child(2) { animation-delay: 0.2s; }
        .llm-typing span:nth-child(3) { animation-delay: 0.4s; }
        @keyframes typing {
            0%, 60%, 100% { transform: translateY(0); opacity: 0.4; }
            30% { transform: translateY(-10px); opacity: 1; }
        }
        
        /* SETTINGS ICON */
        .settings-icon { font-size: 1.8em; cursor: pointer; transition: all 0.3s; padding: 10px; border-radius: 50%; background: #f0f0f0; display: flex; align-items: center; justify-content: center; width: 50px; height: 50px; }
        .settings-icon:hover { transform: rotate(90deg); background: #667eea; color: white; }
        
        /* MODAL POPUP */
        .modal { display: none; position: fixed; z-index: 1000; left: 0; top: 0; width: 100%; height: 100%; background: rgba(0,0,0,0.5); animation: fadeIn 0.3s; }
        .modal-content { background: white; margin: 5% auto; padding: 0; border-radius: 15px; width: 90%; max-width: 500px; box-shadow: 0 20px 60px rgba(0,0,0,0.3); animation: slideDown 0.3s; }
        .modal-header { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 20px 30px; border-radius: 15px 15px 0 0; display: flex; justify-content: space-between; align-items: center; }
        .modal-header h2 { margin: 0; font-size: 1.5em; }
        .close-btn { font-size: 2em; cursor: pointer; color: white; background: none; border: none; line-height: 1; transition: transform 0.2s; }
        .close-btn:hover { transform: scale(1.2); }
        .modal-body { padding: 30px; }
        .modal-body label { display: block; margin-bottom: 8px; font-weight: 600; color: #333; }
        .modal-body input { width: 100%; padding: 12px; margin-bottom: 20px; border: 2px solid #e5e7eb; border-radius: 8px; font-size: 1em; transition: border-color 0.3s; }
        .modal-body input:focus { outline: none; border-color: #667eea; }
        
        /* API KEY INPUT CONTAINER */
        .api-key-input-container { position: relative; margin-bottom: 20px; }
        .api-key-input-container input { padding-right: 90px; margin-bottom: 0; font-family: monospace; letter-spacing: 1px; }
        .api-key-input-container .input-icons { position: absolute; right: 8px; top: 50%; transform: translateY(-50%); display: flex; gap: 5px; align-items: center; }
        .api-key-icon-btn { background: transparent; border: none; cursor: pointer; padding: 8px; border-radius: 6px; display: flex; align-items: center; justify-content: center; transition: all 0.2s; font-size: 18px; color: #666; }
        .api-key-icon-btn:hover { background: rgba(102, 126, 234, 0.1); color: #667eea; transform: scale(1.1); }
        .api-key-icon-btn:active { transform: scale(0.95); }
        .api-key-icon-btn.copied { color: #10b981; animation: copySuccess 0.3s; }
        @keyframes copySuccess { 0%, 100% { transform: scale(1); } 50% { transform: scale(1.2); } }
        .modal-footer { padding: 20px 30px; background: #f9fafb; border-radius: 0 0 15px 15px; display: flex; gap: 15px; justify-content: flex-end; }
        .modal-btn { padding: 12px 30px; border: none; border-radius: 8px; font-weight: 600; cursor: pointer; transition: all 0.3s; font-size: 1em; }
        .modal-btn.primary { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; }
        .modal-btn.primary:hover { transform: translateY(-2px); box-shadow: 0 8px 20px rgba(102, 126, 234, 0.4); }
        .modal-btn.secondary { background: #e5e7eb; color: #666; }
        .modal-btn.secondary:hover { background: #d1d5db; }
        .modal-btn.info { background: linear-gradient(135deg, #17a2b8 0%, #138496 100%); color: white; }
        .modal-btn.info:hover { transform: translateY(-2px); box-shadow: 0 8px 20px rgba(23, 162, 184, 0.4); }
        @keyframes fadeIn { from { opacity: 0; } to { opacity: 1; } }
        @keyframes slideDown { from { transform: translateY(-50px); opacity: 0; } to { transform: translateY(0); opacity: 1; } }
        
        /* AUDIO VISUALIZER - Sóng nhạc đẹp */
        .audio-visualizer {
            display: flex;
            align-items: flex-end;
            justify-content: center;
            gap: 3px;
            height: 40px;
            margin: 10px 0;
        }
        .audio-visualizer .bar {
            width: 4px;
            background: linear-gradient(to top, #667eea, #764ba2, #f472b6);
            border-radius: 2px;
            animation: visualizer-bar 0.5s ease-in-out infinite;
        }
        .audio-visualizer .bar:nth-child(1) { animation-delay: 0s; height: 20px; }
        .audio-visualizer .bar:nth-child(2) { animation-delay: 0.1s; height: 30px; }
        .audio-visualizer .bar:nth-child(3) { animation-delay: 0.15s; height: 25px; }
        .audio-visualizer .bar:nth-child(4) { animation-delay: 0.3s; height: 35px; }
        .audio-visualizer .bar:nth-child(5) { animation-delay: 0.2s; height: 28px; }
        .audio-visualizer .bar:nth-child(6) { animation-delay: 0.25s; height: 32px; }
        .audio-visualizer .bar:nth-child(7) { animation-delay: 0.05s; height: 22px; }
        .audio-visualizer .bar:nth-child(8) { animation-delay: 0.35s; height: 38px; }
        .audio-visualizer .bar:nth-child(9) { animation-delay: 0.1s; height: 26px; }
        .audio-visualizer .bar:nth-child(10) { animation-delay: 0.4s; height: 30px; }
        .audio-visualizer.paused .bar { animation-play-state: paused; }
        @keyframes visualizer-bar {
            0%, 100% { transform: scaleY(0.3); opacity: 0.6; }
            50% { transform: scaleY(1); opacity: 1; }
        }
        
        /* RUNCAT ANIMATION - JavaScript-based multi-frame like RunCat365 */
        #runcat-container {
            position: fixed;
            bottom: 15px;
            right: 15px;
            z-index: 9999;
            user-select: none;
            cursor: pointer;
        }
        
        #runcat {
            font-size: 52px;
            display: inline-block;
            filter: drop-shadow(0 3px 6px rgba(0,0,0,0.25));
            transition: transform 0.05s ease-out;
            will-change: transform;
        }
        
        #runcat:hover {
            animation: runcat-excited 0.15s ease-in-out infinite !important;
            filter: drop-shadow(0 6px 12px rgba(0,0,0,0.4));
        }
        
        @keyframes runcat-excited {
            0%, 100% { 
                transform: translateY(-2px) rotate(-8deg) scale(1.2) !important;
            }
            25% { 
                transform: translateY(-12px) rotate(8deg) scale(1.3) !important;
            }
            50% { 
                transform: translateY(-18px) rotate(-8deg) scale(1.25) !important;
            }
            75% { 
                transform: translateY(-12px) rotate(8deg) scale(1.3) !important;
            }
        }
        
        /* FOOTER MINIZ - Compact corner style */
        .footer-miniz { position: fixed; bottom: 20px; right: 20px; background: rgba(26, 26, 46, 0.95); color: white; padding: 12px 18px; border-radius: 50px; box-shadow: 0 5px 25px rgba(0,0,0,0.3); display: flex; align-items: center; gap: 12px; z-index: 1000; transition: all 0.3s; backdrop-filter: blur(10px); }
        .footer-miniz:hover { transform: translateY(-3px); box-shadow: 0 8px 35px rgba(102, 126, 234, 0.5); }
        .footer-logo-compact { display: flex; align-items: center; gap: 10px; }
        .footer-logo-compact img { width: 35px; height: 35px; border-radius: 50%; border: 2px solid #667eea; box-shadow: 0 0 10px rgba(102, 126, 234, 0.6); }
        .footer-brand-compact { font-size: 0.95em; font-weight: bold; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text; }
        .footer-separator { width: 1px; height: 25px; background: rgba(255,255,255,0.3); }
        .footer-youtube-compact { display: flex; align-items: center; gap: 6px; padding: 8px 15px; background: #FF0000; color: white; border-radius: 25px; text-decoration: none; font-weight: 600; font-size: 0.85em; transition: all 0.3s; }
        .footer-youtube-compact:hover { background: #cc0000; transform: scale(1.05); }
        .footer-youtube-compact svg { width: 18px; height: 18px; fill: white; }
        
        /* RESPONSIVE - MOBILE FIRST */
        @media (max-width: 1200px) {
            .quick-actions { grid-template-columns: repeat(auto-fit, minmax(160px, 1fr)); gap: 15px; }
            .tab-content.active { grid-template-columns: repeat(auto-fit, minmax(280px, 1fr)); }
            .device-grid { grid-template-columns: repeat(auto-fit, minmax(280px, 1fr)); }
        }
        
        @media (max-width: 992px) {
            .sidebar { width: 240px; padding: 20px 15px; }
            .main-content { padding: 20px; }
            .header { padding: 20px; flex-direction: column; gap: 15px; text-align: center; }
            .header h1 { font-size: 1.6em; }
            .music-player { padding: 20px; }
            .player-controls { gap: 15px; }
            .player-btn { width: 50px; height: 50px; font-size: 20px; }
            .player-btn.play { width: 65px; height: 65px; font-size: 26px; }
        }
        
        @media (max-width: 768px) {
            body { flex-direction: column; }
            .sidebar { width: 100%; padding: 15px; flex-direction: row; flex-wrap: wrap; justify-content: center; gap: 10px; }
            .logo { width: 100%; margin-bottom: 15px; padding: 15px; }
            .logo-icon { width: 60px; }
            .logo-text { font-size: 1.2em; }
            .menu-item { padding: 10px 15px; margin: 3px; font-size: 0.9em; }
            .main-content { padding: 15px; min-height: calc(100vh - 200px); }
            .header { padding: 15px; margin-bottom: 20px; }
            .header h1 { font-size: 1.3em; }
            .status { flex-wrap: wrap; justify-content: center; gap: 10px; }
            .status-badge { padding: 6px 15px; font-size: 0.9em; }
            .quick-actions { grid-template-columns: repeat(2, 1fr); gap: 10px; }
            .action-card { padding: 15px; }
            .action-card .icon { font-size: 1.8em; }
            .action-card .title { font-size: 0.9em; }
            .tools-section, .config-section { padding: 20px; }
            .tools-tabs { flex-wrap: wrap; gap: 8px; }
            .tab-btn { padding: 10px 20px; font-size: 0.9em; }
            .tab-content.active { grid-template-columns: 1fr; }
            .tool-card { padding: 20px; }
            .device-grid { grid-template-columns: 1fr; }
            .music-player { padding: 15px; border-radius: 15px; }
            .now-playing h3 { font-size: 1.2em; }
            .player-controls { gap: 10px; margin: 20px 0; }
            .player-btn { width: 45px; height: 45px; font-size: 18px; }
            .player-btn.play { width: 60px; height: 60px; font-size: 24px; }
            .music-list { padding: 15px; max-height: 350px; }
            .music-item { padding: 12px; }
            .chat-bubble { max-width: 85%; }
            .modal-content { width: 95%; margin: 2% auto; }
            .modal-body { padding: 20px; }
            .modal-footer { padding: 15px 20px; flex-direction: column; }
            .modal-btn { width: 100%; }
            .footer-miniz { bottom: 10px; right: 10px; padding: 10px 14px; }
            .footer-brand-compact { font-size: 0.85em; }
            .footer-youtube-compact { padding: 6px 12px; font-size: 0.8em; }
            #runcat-container { bottom: 10px; right: 10px; }
            #runcat { font-size: 40px; }
        }
        
        @media (max-width: 480px) {
            .sidebar { padding: 10px; }
            .logo { padding: 10px; }
            .logo-icon { width: 45px; }
            .logo-text { font-size: 1em; }
            .menu-item { padding: 8px 12px; font-size: 0.85em; }
            .main-content { padding: 10px; }
            .header { padding: 12px; }
            .header h1 { font-size: 1.1em; }
            .quick-actions { grid-template-columns: repeat(2, 1fr); gap: 8px; }
            .action-card { padding: 12px; }
            .action-card .icon { font-size: 1.5em; margin-bottom: 5px; }
            .action-card .title { font-size: 0.8em; }
            .tools-section, .config-section { padding: 15px; margin-bottom: 20px; }
            .tab-btn { padding: 8px 15px; font-size: 0.85em; }
            .tool-card { padding: 15px; }
            .tool-card h3 { font-size: 1em; }
            .tool-card input, .tool-card select, .tool-card textarea { padding: 10px; font-size: 0.9em; }
            .tool-card button { padding: 12px; font-size: 0.9em; }
            .log-panel { max-height: 250px; padding: 15px; font-size: 0.85em; }
            .music-player { padding: 12px; }
            .now-playing h3 { font-size: 1em; }
            .now-playing p { font-size: 0.9em; }
            .player-btn { width: 40px; height: 40px; font-size: 16px; }
            .player-btn.play { width: 55px; height: 55px; font-size: 22px; }
            .chat-avatar { width: 32px; height: 32px; font-size: 0.9em; }
            .chat-bubble { padding: 10px 12px; }
            .chat-content { font-size: 0.9em; }
            .footer-miniz { flex-direction: column; padding: 8px 12px; gap: 8px; }
            .footer-separator { display: none; }
        }
        
        /* WECHAT STYLE CHAT BUBBLES */
        .chat-message { display: flex; align-items: flex-start; gap: 10px; margin-bottom: 15px; animation: fadeInChat 0.3s; }
        .chat-message.user { flex-direction: row-reverse; }
        .chat-avatar { width: 40px; height: 40px; border-radius: 50%; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); display: flex; align-items: center; justify-content: center; color: white; font-weight: 700; font-size: 1.1em; flex-shrink: 0; box-shadow: 0 2px 8px rgba(0,0,0,0.15); }
        .chat-avatar.assistant { background: linear-gradient(135deg, #10b981 0%, #059669 100%); }
        .chat-avatar.system { background: linear-gradient(135deg, #f59e0b 0%, #d97706 100%); }
        .chat-avatar.tool { background: linear-gradient(135deg, #3b82f6 0%, #2563eb 100%); }
        .chat-bubble { max-width: 65%; padding: 12px 16px; border-radius: 12px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); position: relative; word-wrap: break-word; }
        .chat-message.user .chat-bubble { background: #667eea; color: white; border-radius: 12px 12px 2px 12px; }
        .chat-message.assistant .chat-bubble { background: white; color: #333; border-radius: 12px 12px 12px 2px; border: 1px solid #e5e7eb; }
        .chat-message.system .chat-bubble { background: #fff7ed; color: #7c2d12; border-radius: 8px; border: 1px solid #fed7aa; }
        .chat-message.tool .chat-bubble { background: #eff6ff; color: #1e3a8a; border-radius: 8px; border: 1px solid #bfdbfe; }
        .chat-content { font-size: 0.95em; line-height: 1.5; margin-bottom: 6px; }
        .chat-metadata { font-size: 0.75em; opacity: 0.7; display: flex; gap: 10px; flex-wrap: wrap; margin-top: 8px; }
        .chat-metadata-item { display: inline-flex; align-items: center; gap: 4px; background: rgba(0,0,0,0.05); padding: 2px 8px; border-radius: 10px; }
        .chat-timestamp { font-size: 0.7em; opacity: 0.6; margin-top: 4px; text-align: right; }
        .chat-message.user .chat-timestamp { text-align: left; }
        @keyframes fadeInChat { from { opacity: 0; transform: translateY(10px); } to { opacity: 1; transform: translateY(0); } }
        @keyframes spin { from { transform: rotate(0deg); } to { transform: rotate(360deg); } }
        @keyframes pulse { 0%, 100% { opacity: 1; } 50% { opacity: 0.5; } }
        #chat-container::-webkit-scrollbar { width: 8px; }
        #chat-container::-webkit-scrollbar-track { background: #f1f1f1; border-radius: 10px; }
        #chat-container::-webkit-scrollbar-thumb { background: #667eea; border-radius: 10px; }
        #chat-container::-webkit-scrollbar-thumb:hover { background: #5568d3; }
        
        /* Music Player VLC-style enhancements */
        .music-item:hover { background: linear-gradient(135deg, rgba(102,126,234,0.15) 0%, rgba(118,75,162,0.15) 100%) !important; transform: translateX(5px); }
        #volume-slider::-webkit-slider-thumb { -webkit-appearance: none; width: 16px; height: 16px; background: #667eea; border-radius: 50%; cursor: pointer; box-shadow: 0 2px 6px rgba(102, 126, 234, 0.5); }
        #volume-slider::-moz-range-thumb { width: 16px; height: 16px; background: #667eea; border-radius: 50%; cursor: pointer; border: none; }
    </style>
</head>
<body>
    <!-- SIDEBAR -->
    <div class="sidebar">
        <div class="logo">
            <svg class="logo-icon" viewBox="0 0 100 100" fill="none" xmlns="http://www.w3.org/2000/svg">
                <defs>
                    <linearGradient id="logoGrad" x1="0%" y1="0%" x2="100%" y2="100%">
                        <stop offset="0%" style="stop-color:#667eea"/>
                        <stop offset="100%" style="stop-color:#764ba2"/>
                    </linearGradient>
                </defs>
                <circle cx="50" cy="50" r="45" fill="url(#logoGrad)"/>
                <text x="50" y="58" text-anchor="middle" fill="white" font-size="28" font-weight="bold" font-family="Arial">MCP</text>
                <text x="50" y="75" text-anchor="middle" fill="#a5f3fc" font-size="12" font-weight="600" font-family="Arial">miniZ</text>
            </svg>
            <div class="logo-text">miniZ MCP</div>
            <small style="font-size:0.55em;opacity:0.9;font-weight:600;letter-spacing:1px;">ĐIỀU KHIỂN MÁY TÍNH</small>
        </div>
        <div class="menu-item active" onclick="showSection('dashboard')">📊Sidebar</div>
        <div class="menu-item" onclick="showSection('tools')">🛠️ Công Cụ</div>
        <div class="menu-item" onclick="showSection('llm-chat')" style="background:linear-gradient(135deg,#667eea,#764ba2);border-left:4px solid #fbbf24;">💬 Chat với Gemini</div>
        <div class="menu-item" onclick="showSection('api-quotas')" style="background:linear-gradient(135deg,#667eea,#764ba2);border-left:4px solid #fbbf24;">🔑 API Quotas</div>
        <div class="menu-item" onclick="showSection('music')">🎵 Music Player</div>
        <div class="menu-item" onclick="showSection('music-settings')">⚙️ Music Settings</div>
        <div class="menu-item" onclick="showSection('conversation')">💬 Lịch Sử Chat</div>
        <div class="menu-item" onclick="showSection('playlist')">🎵 Playlist YouTube</div>
        <div class="menu-item" onclick="showSection('knowledge')">📚 Knowledge Base</div>
    </div>
    
    <!-- MAIN CONTENT -->
    <div class="main-content">
        <!-- HEADER -->
        <div class="header">
            <h1>Dashboard</h1>
            <div class="status">
                <div class="settings-icon" onclick="openSettingsModal()" title="Cấu hình Endpoint">⚙️</div>
                <div class="status-badge" id="xiaozhi-status">
                    <span class="status-dot"></span>
                    <span id="xiaozhi-text">Connecting...</span>
                </div>
                <div class="status-badge online">
                    <span class="status-dot"></span>
                    Web Server
                </div>
            </div>
        </div>
        
        <!-- DASHBOARD SECTION -->
        <div id="dashboard-section">
            <h2 style="color:#667eea;margin-bottom:20px;">🚀 Tất cả công cụ (38 Tools)</h2>
            <div class="quick-actions">
                <!-- AI ASSISTANT (2) - NEW -->
                <div class="action-card purple" onclick="askGemini()"><div class="icon">🤖📚</div><div class="title">Hỏi Gemini AI + KB</div></div>
                <div class="action-card indigo" onclick="askGPT4()"><div class="icon">🧠</div><div class="title">Hỏi GPT-4</div></div>
                
                <!-- HỆ THỐNG (5) -->
                <div class="action-card blue" onclick="setVolumePrompt()"><div class="icon">🔊</div><div class="title">Điều Chỉnh Âm Lượng</div></div>
                <div class="action-card cyan" onclick="screenshot()"><div class="icon">📸</div><div class="title">Chụp Màn Hình</div></div>
                <div class="action-card purple" onclick="notification()"><div class="icon">🔔</div><div class="title">Thông Báo</div></div>
                <div class="action-card green" onclick="getResources()"><div class="icon">💻</div><div class="title">Tài Nguyên Hệ Thống</div></div>
                <div class="action-card orange" onclick="setBrightness()"><div class="icon">🔆</div><div class="title">Độ Sáng Màn Hình</div></div>
                
                <!-- FILE & PROCESS (7) -->
                <div class="action-card indigo" onclick="openApp()"><div class="icon">🚀</div><div class="title">Mở Ứng Dụng</div></div>
                <div class="action-card blue" onclick="listProcesses()"><div class="icon">⚙️</div><div class="title">Tiến Trình Đang Chạy</div></div>
                <div class="action-card red" onclick="killProcess()"><div class="icon">❌</div><div class="title">Tắt Tiến Trình</div></div>
                <div class="action-card green" onclick="createFile()"><div class="icon">➕</div><div class="title">Tạo File Mới</div></div>
                <div class="action-card cyan" onclick="readFile()"><div class="icon">📖</div><div class="title">Đọc File</div></div>
                <div class="action-card purple" onclick="listFiles()"><div class="icon">📂</div><div class="title">Liệt Kê Files</div></div>
                <div class="action-card orange" onclick="diskUsage()"><div class="icon">💽</div><div class="title">Thông Tin Đĩa</div></div>
                
                <!-- MẠNG & WEB (3) -->
                <div class="action-card blue" onclick="networkInfo()"><div class="icon">🌐</div><div class="title">Thông Tin Mạng</div></div>
                <div class="action-card green" onclick="batteryStatus()"><div class="icon">🔋</div><div class="title">Thông Tin Pin</div></div>
                <div class="action-card indigo" onclick="searchWeb()"><div class="icon">🔍</div><div class="title">Tìm Kiếm Google</div></div>
                
                <!-- TIỆN ÍCH (5) -->
                <div class="action-card pink" onclick="calculator()"><div class="icon">🧮</div><div class="title">Máy Tính</div></div>
                <div class="action-card cyan" onclick="getCurrentTime()"><div class="icon">⏰</div><div class="title">Thời Gian</div></div>
                <div class="action-card purple" onclick="getClipboard()"><div class="icon">📋</div><div class="title">Lấy Clipboard</div></div>
                <div class="action-card orange" onclick="setClipboard()"><div class="icon">📝</div><div class="title">Đặt Clipboard</div></div>
                <div class="action-card red" onclick="playSound()"><div class="icon">🔊</div><div class="title">Phát Âm Thanh</div></div>
                
                <!-- NEW TOOLS -->
                <div class="action-card blue" onclick="lockComputer()"><div class="icon">🔒</div><div class="title">Khóa Máy Tính</div></div>
                <div class="action-card red" onclick="shutdownSchedule()"><div class="icon">⏰</div><div class="title">Lên Lịch Tắt Máy</div></div>
                <div class="action-card green" onclick="showDesktop()"><div class="icon">🖥️</div><div class="title">Hiển Thị Desktop</div></div>
                <div class="action-card orange" onclick="undoOperation()"><div class="icon">↩️</div><div class="title">Hoàn Tác</div></div>
                <div class="action-card purple" onclick="setTheme()"><div class="icon">🎨</div><div class="title">Đổi Theme</div></div>
                <div class="action-card cyan" onclick="changeWallpaper()"><div class="icon">🖼️</div><div class="title">Đổi Hình Nền</div></div>
                <div class="action-card indigo" onclick="getDesktopPath()"><div class="icon">📁</div><div class="title">Đường Dẫn Desktop</div></div>
                <div class="action-card pink" onclick="pasteContent()"><div class="icon">📋</div><div class="title">Dán Nội Dung</div></div>
                <div class="action-card blue" onclick="pressEnter()"><div class="icon">⏎</div><div class="title">Nhấn Enter</div></div>
                <div class="action-card green" onclick="findInDocument()"><div class="icon">🔎</div><div class="title">Tìm Trong Tài Liệu</div></div>
            </div>
            
            <!-- LOG PANEL AT BOTTOM OF DASHBOARD -->
            <div style="margin-top: 30px;">
                <h2 style="color:#667eea; margin-bottom: 15px; display: flex; align-items: center; gap: 10px;">
                    <span>📋 Log Hoạt Động</span>
                    <span style="font-size: 0.6em; color: #9ca3af; font-weight: 400;">(Thời gian thực)</span>
                </h2>
                <div class="log-panel" id="log"></div>
            </div>
        </div>

        <!-- API QUOTAS SECTION -->
        <div id="api-quotas-section" class="section" style="display:none;">
            <h2 style="color:#667eea;margin-bottom:30px;">🔑 API Quotas Management</h2>
            
            <div style="display:grid;grid-template-columns:repeat(auto-fit,minmax(400px,1fr));gap:25px;margin-bottom:30px;">
                <!-- Gemini API Card -->
                <div style="background:linear-gradient(135deg,#667eea 0%,#764ba2 100%);border-radius:15px;padding:30px;color:white;box-shadow:0 10px 30px rgba(102,126,234,0.3);">
                    <div style="display:flex;align-items:center;margin-bottom:20px;">
                        <div style="font-size:48px;margin-right:15px;">🤖</div>
                        <div>
                            <h3 style="margin:0;font-size:24px;">Gemini API</h3>
                            <p style="margin:5px 0 0 0;opacity:0.9;font-size:14px;">Google AI Platform</p>
                        </div>
                    </div>
                    <div id="gemini-quota-detail" style="background:rgba(255,255,255,0.15);border-radius:10px;padding:20px;">
                        <div style="margin-bottom:15px;">
                            <div style="font-size:13px;opacity:0.9;margin-bottom:5px;">Status:</div>
                            <div id="gemini-status" style="font-size:16px;font-weight:bold;">🔄 Đang kiểm tra...</div>
                        </div>
                        <div style="margin-bottom:15px;">
                            <div style="font-size:13px;opacity:0.9;margin-bottom:5px;">Free Tier Limits:</div>
                            <div style="font-size:15px;line-height:1.6;">
                                • <strong>60 requests</strong> per minute<br>
                                • <strong>1,500 requests</strong> per day
                            </div>
                        </div>
                        <div>
                            <div style="font-size:13px;opacity:0.9;margin-bottom:5px;">Model:</div>
                            <div style="font-size:14px;font-family:monospace;background:rgba(0,0,0,0.2);padding:8px;border-radius:5px;">
                                🚀 Gemini 3 Flash Preview
                                <br><span style="font-size:11px;opacity:0.7;">gemini-3-flash-preview</span>
                            </div>
                        </div>
                    </div>
                </div>
                
                <!-- Serper API Card -->
                <div style="background:linear-gradient(135deg,#3b82f6 0%,#1e40af 100%);border-radius:15px;padding:30px;color:white;box-shadow:0 10px 30px rgba(59,130,246,0.3);">
                    <div style="display:flex;align-items:center;margin-bottom:20px;">
                        <div style="font-size:48px;margin-right:15px;">🔍</div>
                        <div>
                            <h3 style="margin:0;font-size:24px;">Serper API</h3>
                            <p style="margin:5px 0 0 0;opacity:0.9;font-size:14px;">Google Search API</p>
                        </div>
                    </div>
                    <div id="serper-quota-detail" style="background:rgba(255,255,255,0.15);border-radius:10px;padding:20px;">
                        <div style="margin-bottom:15px;">
                            <div style="font-size:13px;opacity:0.9;margin-bottom:5px;">Status:</div>
                            <div id="serper-status" style="font-size:16px;font-weight:bold;">🔄 Đang kiểm tra...</div>
                        </div>
                        <div style="margin-bottom:15px;">
                            <div style="font-size:13px;opacity:0.9;margin-bottom:5px;">Free Tier Limit:</div>
                            <div style="font-size:15px;line-height:1.6;">
                                • <strong>2,500 queries</strong> per month
                            </div>
                        </div>
                        <div>
                            <div style="font-size:13px;opacity:0.9;margin-bottom:5px;">Endpoint:</div>
                            <div style="font-size:14px;font-family:monospace;background:rgba(0,0,0,0.2);padding:8px;border-radius:5px;">https://google.serper.dev/search</div>
                        </div>
                    </div>
                </div>
            </div>
            
            <!-- Actions -->
            <div style="background:white;border-radius:15px;padding:25px;box-shadow:0 2px 10px rgba(0,0,0,0.1);margin-bottom:25px;">
                <h3 style="margin-top:0;color:#1a1a2e;">⚡ Quick Actions</h3>
                <div style="display:flex;gap:15px;flex-wrap:wrap;">
                    <button onclick="refreshQuotasPage()" style="background:linear-gradient(135deg,#10b981,#059669);color:white;border:none;padding:12px 25px;border-radius:8px;font-size:15px;cursor:pointer;box-shadow:0 4px 15px rgba(16,185,129,0.3);transition:all 0.3s;">
                        🔄 Làm mới tất cả
                    </button>
                    <button onclick="testGeminiAPI()" style="background:linear-gradient(135deg,#667eea,#764ba2);color:white;border:none;padding:12px 25px;border-radius:8px;font-size:15px;cursor:pointer;box-shadow:0 4px 15px rgba(102,126,234,0.3);transition:all 0.3s;">
                        🧪 Test Gemini API
                    </button>
                    <button onclick="testSerperAPI()" style="background:linear-gradient(135deg,#3b82f6,#1e40af);color:white;border:none;padding:12px 25px;border-radius:8px;font-size:15px;cursor:pointer;box-shadow:0 4px 15px rgba(59,130,246,0.3);transition:all 0.3s;">
                        🧪 Test Serper API
                    </button>
                </div>
            </div>
            
            <!-- Usage Tips -->
            <div style="background:#f0f9ff;border-left:4px solid #3b82f6;border-radius:10px;padding:20px;">
                <h3 style="margin-top:0;color:#1e40af;">💡 Tips</h3>
                <ul style="margin:10px 0;padding-left:20px;line-height:1.8;color:#1e3a8a;">
                    <li><strong>Gemini API:</strong> Dùng cho chat AI, phân tích text, tạo nội dung</li>
                    <li><strong>Serper API:</strong> Dùng cho tìm kiếm Google real-time</li>
                    <li><strong>Free Tier:</strong> Đủ cho sử dụng cá nhân và testing</li>
                    <li><strong>Rate Limit:</strong> Nếu vượt quota, API sẽ trả về lỗi 429</li>
                    <li><strong>Monitor:</strong> Kiểm tra status thường xuyên để tránh hết quota</li>
                </ul>
            </div>
        </div>

        <!-- TOOLS SECTION -->
        <div id="tools-section" style="display:none;">
            <div class="tools-section">
                <h2 style="color:#667eea;margin-bottom:20px;">🛠️ Công Cụ (20 Tools)</h2>
                
                <div class="tools-tabs">
                    <button class="tab-btn active" onclick="switchTab(0)">🎛️ Hệ thống</button>
                    <button class="tab-btn" onclick="switchTab(1)">📁 File & Process</button>
                    <button class="tab-btn" onclick="switchTab(2)">🌐 Mạng & Web</button>
                    <button class="tab-btn" onclick="switchTab(3)">🔧 Tiện ích</button>
                </div>
                
                <!-- TAB 1: HỆ THỐNG -->
                <div class="tab-content active" id="tab-0">
                    <div class="tool-card">
                        <h3>🔊 Điều chỉnh âm lượng</h3>
                        <input type="number" id="volume" min="0" max="100" value="50" placeholder="0-100">
                        <button onclick="
                            const level = parseInt(document.getElementById('volume').value);
                            if (isNaN(level) || level < 0 || level > 100) {
                                addLog('❌ Âm lượng phải từ 0-100', 'error');
                            } else {
                                callAPI('/api/volume', {level: level});
                            }
                        ">Đặt âm lượng</button>
                    </div>
                    <div class="tool-card">
                        <h3>📸 Chụp màn hình</h3>
                        <button onclick="callAPI('/api/screenshot', {})">Chụp màn hình ngay</button>
                    </div>
                    <div class="tool-card">
                        <h3>🔔 Thông báo</h3>
                        <input type="text" id="notif-title" placeholder="Tiêu đề">
                        <input type="text" id="notif-message" placeholder="Nội dung">
                        <button onclick="
                            const title = document.getElementById('notif-title').value.trim();
                            const message = document.getElementById('notif-message').value.trim();
                            if (!title || !message) {
                                addLog('❌ Vui lòng nhập tiêu đề và nội dung', 'error');
                            } else {
                                callAPI('/api/notification', {title: title, message: message});
                            }
                        ">Hiển thị</button>
                    </div>
                    <div class="tool-card" style="background:linear-gradient(135deg, #667eea 0%, #764ba2 100%);color:white;border:2px solid #764ba2;">
                        <h3 style="color:white;">🔑 API Quotas</h3>
                        <button onclick="getQuotas()" style="background:rgba(255,255,255,0.2);color:white;border:1px solid rgba(255,255,255,0.3);">Làm mới</button>
                        <div id="quotas" style="margin-top:15px;font-size:13px;line-height:1.8;">
                            <div style="margin-bottom:10px;padding:8px;background:rgba(255,255,255,0.1);border-radius:4px;">
                                <strong>🤖 Gemini:</strong><br>
                                <span id="gemini-quota" style="color:#fbbf24;font-size:12px;">Đang tải...</span>
                            </div>
                            <div style="padding:8px;background:rgba(255,255,255,0.1);border-radius:4px;">
                                <strong>🔍 Serper:</strong><br>
                                <span id="serper-quota" style="color:#60a5fa;font-size:12px;">Đang tải...</span>
                            </div>
                        </div>
                    </div>
                    <div class="tool-card">
                        <h3>💻 Tài nguyên hệ thống</h3>
                        <button onclick="getResources()">Làm mới</button>
                        <div id="resources" style="margin-top:15px;">
                            <div>CPU: <span id="cpu">--%</span></div>
                            <div>RAM: <span id="ram">--%</span></div>
                            <div>Disk: <span id="disk">--%</span></div>
                        </div>
                    </div>
                    <div class="tool-card">
                        <h3>🔆 Độ sáng màn hình</h3>
                        <input type="number" id="brightness" min="0" max="100" value="50" placeholder="0-100">
                        <button onclick="
                            const level = parseInt(document.getElementById('brightness').value);
                            if (isNaN(level) || level < 0 || level > 100) {
                                addLog('❌ Độ sáng phải từ 0-100', 'error');
                            } else {
                                callTool('set_brightness', {level: level});
                            }
                        ">Đặt độ sáng</button>
                    </div>
                </div>
                
                <!-- TAB 2: FILE & PROCESS -->
                <div class="tab-content" id="tab-1">
                    <div class="tool-card">
                        <h3>🚀 Mở ứng dụng</h3>
                        <select id="app-name">
                            <option value="notepad">📝 Notepad</option>
                            <option value="calc">🧮 Calculator</option>
                            <option value="paint">🎨 Paint</option>
                            <option value="cmd">⌨️ CMD</option>
                            <option value="explorer">📂 Explorer</option>
                        </select>
                        <button onclick="callTool('open_application', {app_name: document.getElementById('app-name').value})">Mở</button>
                    </div>
                    <div class="tool-card">
                        <h3>📋 Tiến trình đang chạy</h3>
                        <input type="number" id="proc-limit" min="5" max="50" value="10" placeholder="Số lượng">
                        <button onclick="callTool('list_running_processes', {limit: parseInt(document.getElementById('proc-limit').value)})">Xem danh sách</button>
                    </div>
                    <div class="tool-card">
                        <h3>❌ Tắt tiến trình</h3>
                        <input type="text" id="kill-proc" placeholder="PID hoặc tên">
                        <button onclick="callTool('kill_process', {identifier: document.getElementById('kill-proc').value})">Tắt tiến trình</button>
                    </div>
                    <div class="tool-card">
                        <h3>📝 Tạo file mới</h3>
                        <input type="text" id="file-path" placeholder="C:/test.txt">
                        <textarea id="file-content" placeholder="Nội dung..." style="min-height:80px;"></textarea>
                        <button onclick="callTool('create_file', {path: document.getElementById('file-path').value, content: document.getElementById('file-content').value})">Tạo file</button>
                    </div>
                    <div class="tool-card">
                        <h3>📖 Đọc file</h3>
                        <input type="text" id="read-path" placeholder="C:/test.txt">
                        <button onclick="callTool('read_file', {path: document.getElementById('read-path').value})">Đọc file</button>
                    </div>
                    <div class="tool-card">
                        <h3>📂 Liệt kê files</h3>
                        <input type="text" id="list-dir" placeholder="C:/Users">
                        <button onclick="callTool('list_files', {directory: document.getElementById('list-dir').value})">Xem files</button>
                    </div>
                    <div class="tool-card">
                        <h3>💾 Thông tin đĩa</h3>
                        <button onclick="callTool('get_disk_usage', {})">Xem chi tiết</button>
                    </div>
                </div>
                
                <!-- TAB 3: MẠNG & WEB -->
                <div class="tab-content" id="tab-2">
                    <div class="tool-card">
                        <h3>🌐 Thông tin mạng</h3>
                        <button onclick="callTool('get_network_info', {})">Xem IP & hostname</button>
                    </div>
                    <div class="tool-card">
                        <h3>🔋 Thông tin pin</h3>
                        <button onclick="callTool('get_battery_status', {})">Kiểm tra pin</button>
                    </div>
                    <div class="tool-card">
                        <h3>🔍 Tìm kiếm Google</h3>
                        <input type="text" id="search-query" placeholder="Nhập từ khóa...">
                        <button onclick="callTool('search_web', {query: document.getElementById('search-query').value})">Tìm kiếm</button>
                    </div>
                </div>
                
                <!-- TAB 4: TIỆN ÍCH -->
                <div class="tab-content" id="tab-3">
                    <div class="tool-card">
                        <h3>🧮 Máy tính</h3>
                        <input type="text" id="calc-expr" placeholder="2+2*3">
                        <button onclick="calculate()">Tính toán</button>
                        <div id="calc-result" style="margin-top:10px;font-size:1.5em;font-weight:bold;color:#667eea;"></div>
                    </div>
                    <div class="tool-card">
                        <h3>🕐 Thời gian</h3>
                        <button onclick="getCurrentTime()">Lấy thời gian</button>
                        <div id="time-result" style="margin-top:10px;font-size:1.2em;color:#667eea;"></div>
                    </div>
                    <div class="tool-card">
                        <h3>📋 Lấy clipboard</h3>
                        <button onclick="callTool('get_clipboard', {})">Xem nội dung</button>
                    </div>
                    <div class="tool-card">
                        <h3>📝 Đặt clipboard</h3>
                        <input type="text" id="clip-text" placeholder="Nội dung cần copy">
                        <button onclick="callTool('set_clipboard', {text: document.getElementById('clip-text').value})">Copy vào clipboard</button>
                    </div>
                    <div class="tool-card">
                        <h3>🔊 Phát âm thanh</h3>
                        <input type="number" id="sound-freq" min="200" max="2000" value="1000" placeholder="Tần số Hz">
                        <input type="number" id="sound-dur" min="100" max="3000" value="500" placeholder="Thời gian ms">
                        <button onclick="callTool('play_sound', {frequency: parseInt(document.getElementById('sound-freq').value), duration: parseInt(document.getElementById('sound-dur').value)})">Phát beep</button>
                    </div>
                </div>
            </div>
        </div>
        
        <!-- CONFIG SECTION - HIDDEN (Replaced by Modal) -->
        <div id="config-section" style="display:none;">
            <div class="config-section">
                <h2 style="color:#667eea;margin-bottom:20px;">⚙️ Cấu hình hiện tại</h2>
                <p style="color:#666;margin-bottom:20px;">Sử dụng icon ⚙️ ở góc phải trên để thay đổi endpoint</p>
                <div id="current-endpoint-info" style="background:#f9fafb;padding:20px;border-radius:12px;border:2px solid #e5e7eb;">
                    <p><strong>Thiết bị đang hoạt động:</strong> <span id="current-device-name">-</span></p>
                    <p><strong>Token:</strong> <span id="current-device-token" style="font-family:monospace;font-size:0.9em;word-break:break-all;">-</span></p>
                </div>
            </div>
        </div>
        
        
        <!-- LLM CHAT SECTION - Chat với Gemini AI -->
        <div id="llm-chat-section" style="display:none;">
            <div style="background: white; border-radius: 15px; padding: 20px; box-shadow: 0 10px 30px rgba(0,0,0,0.12); height: calc(100vh - 180px); display: flex; flex-direction: column;">
                <h2 style="color:#10b981; margin-bottom: 15px; display: flex; align-items: center; justify-content: space-between;">
                    <span>💬 Chat với Gemini AI</span>
                    <div style="display:flex; gap:10px; align-items:center;">
                        <!-- TTS Toggle -->
                        <label style="display:flex; align-items:center; gap:6px; cursor:pointer; padding:6px 12px; background:#f3f4f6; border-radius:8px; font-size:0.85em;" title="Bật/tắt đọc to câu trả lời">
                            <input type="checkbox" id="llm-tts-toggle" onchange="saveTTSPreference()" style="cursor:pointer;">
                            <span>🔊 Đọc to</span>
                        </label>
                        <!-- AI Model selector -->
                        <select id="llm-chat-model" style="padding:8px 12px; border-radius:8px; border:2px solid #e5e7eb; font-size:0.9em; cursor:pointer;" onchange="saveLLMChatModel()">
                            <option value="models/gemini-3-flash-preview">⚡ Gemini 3 Flash</option>
                            <option value="models/gemini-2.0-flash">⚡ Gemini 2.0 Flash</option>
                            <option value="models/gemini-2.5-pro-preview-06-05">💎 Gemini 2.5 Pro</option>
                            <option value="models/gemini-2.5-flash-preview-05-20">⚡ Gemini 2.5 Flash</option>
                        </select>
                        <button onclick="clearLLMChat()" style="padding:8px 16px; background:#ef4444; color:white; border:none; border-radius:8px; cursor:pointer; font-size:0.9em;">
                            🗑️ Xóa Chat
                        </button>
                    </div>
                </h2>
                
                <!-- AI Status Bar -->
                <div id="llm-ai-status" style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color:white; padding:12px 16px; border-radius:10px; margin-bottom:15px; display:flex; justify-content:space-between; align-items:center; flex-wrap:wrap; gap:10px;">
                    <div style="display:flex; gap:20px; flex-wrap:wrap; align-items:center;">
                        <span>🤖 <strong>Gemini AI</strong> + 📚 Knowledge Base</span>
                        <span style="font-size:0.85em; opacity:0.9;">Tích hợp RAG System tự động</span>
                    </div>
                    <span style="font-size:0.85em; background:rgba(255,255,255,0.2); padding:4px 10px; border-radius:20px;">✅ Sẵn sàng</span>
                </div>
                
                <!-- Chat Messages Container -->
                <div id="llm-chat-messages" style="flex:1; overflow-y:auto; background:#f5f5f5; border-radius:10px; padding:15px; display:flex; flex-direction:column; gap:12px;">
                    <!-- Welcome message -->
                    <div style="text-align:center; color:#666; padding:40px 20px;">
                        <div style="font-size:4em; margin-bottom:15px;">🤖</div>
                        <h3 style="color:#667eea; margin-bottom:10px;">Chào mừng đến Chat với Gemini AI!</h3>
                        <p style="font-size:0.95em; max-width:400px; margin:0 auto;">
                            Chat trực tiếp với Gemini AI.<br>
                            AI sẽ tự động tìm kiếm trong Knowledge Base của bạn để trả lời chính xác hơn.
                        </p>
                    </div>
                </div>
                
                <!-- Chat Input Area -->
                <div style="margin-top:15px; display:flex; gap:10px; align-items:flex-end;">
                    <!-- 👂 Wake Word Button -->
                    <button id="llm-wakeword-btn" onclick="toggleWakeWord()" 
                            style="width:50px; height:50px; border-radius:50%; background:linear-gradient(135deg,#6b7280,#4b5563); color:white; border:none; cursor:pointer; font-size:1.4em; display:flex; align-items:center; justify-content:center; transition:all 0.3s; flex-shrink:0;"
                            title="👂 Bật Wake Word (nói 'Hey Gemini' để chat)">
                        👂
                    </button>
                    <!-- 🎤 Microphone Button -->
                    <button id="llm-mic-btn" onclick="toggleLLMVoiceInput()" 
                            style="width:50px; height:50px; border-radius:50%; background:linear-gradient(135deg,#10b981,#059669); color:white; border:none; cursor:pointer; font-size:1.4em; display:flex; align-items:center; justify-content:center; transition:all 0.3s; flex-shrink:0;"
                            title="🎤 Nhấn để nói (auto-send)">
                        🎤
                    </button>
                    <div style="flex:1; position:relative;">
                        <textarea id="llm-chat-input" 
                                  placeholder="Nhập tin nhắn hoặc nhấn 🎤 để nói... (Enter để gửi)"
                                  style="width:100%; padding:15px; padding-right:50px; border:2px solid #e5e7eb; border-radius:12px; font-size:1em; resize:none; min-height:50px; max-height:150px; font-family:inherit;"
                                  onkeydown="handleLLMChatKeydown(event)"
                                  oninput="autoResizeLLMInput(this)"></textarea>
                        <button onclick="sendLLMMessage()" 
                                style="position:absolute; right:10px; bottom:10px; width:40px; height:40px; border-radius:50%; background:linear-gradient(135deg,#667eea,#764ba2); color:white; border:none; cursor:pointer; font-size:1.2em; display:flex; align-items:center; justify-content:center; transition:all 0.3s;"
                                title="Gửi tin nhắn">
                            ➤
                        </button>
                    </div>
                </div>
                <!-- Voice Recording Status -->
                <div id="llm-voice-status" style="display:none; margin-top:10px; padding:12px 16px; background:linear-gradient(135deg,#fef3c7,#fde68a); border-radius:10px; text-align:center;">
                    <span id="llm-voice-status-text">🎤 Đang nghe...</span>
                </div>
                <!-- Wake Word Info -->
                <div style="margin-top:8px; font-size:0.8em; color:#6b7280; text-align:center;">
                    💡 <strong>Wake Words:</strong> "Hey Gemini", "Gemini ơi", "Xin chào" | <strong>Goodbye:</strong> "Tạm biệt", "Bye bye", "Ngủ đi"
                </div>
                
                <!-- Quick Actions -->
                <div style="margin-top:10px; display:flex; gap:8px; flex-wrap:wrap;">
                    <button onclick="sendQuickMessage('Xin chào!')" class="quick-msg-btn">👋 Xin chào</button>
                    <button onclick="sendQuickMessage('Tóm tắt kiến thức trong Knowledge Base')" class="quick-msg-btn">📚 KB Summary</button>
                    <button onclick="sendQuickMessage('Giải thích code Python cho người mới')" class="quick-msg-btn">🐍 Python</button>
                    <button onclick="sendQuickMessage('Viết một đoạn văn ngắn về AI')" class="quick-msg-btn">✍️ Viết văn</button>
                    <button onclick="sendQuickMessage('Dịch sang tiếng Anh: Xin chào các bạn')" class="quick-msg-btn">🌐 Dịch thuật</button>
                </div>
            </div>
        </div>
        
        <!-- CONVERSATION HISTORY SECTION (WeChat style) -->
        <div id="conversation-section" style="display:none;">
            <div style="background: white; border-radius: 15px; padding: 20px; box-shadow: 0 10px 30px rgba(0,0,0,0.12); height: calc(100vh - 180px); display: flex; flex-direction: column;">
                <h2 style="color:#667eea; margin-bottom: 15px; display: flex; align-items: center; justify-content: space-between;">
                    <span>💬 Lịch Sử Hội Thoại</span>
                    <div style="display:flex; gap:10px;">
                        <button onclick="loadConversationHistory()" style="padding:8px 16px; background:#10b981; color:white; border:none; border-radius:8px; cursor:pointer; font-size:0.9em;">
                            🔄 Làm mới
                        </button>
                        <button onclick="exportConversation()" style="padding:8px 16px; background:#667eea; color:white; border:none; border-radius:8px; cursor:pointer; font-size:0.9em;">
                            💾 Xuất File
                        </button>
                        <button onclick="clearConversationHistory()" style="padding:8px 16px; background:#ef4444; color:white; border:none; border-radius:8px; cursor:pointer; font-size:0.9em;">
                            🗑️ Xóa Hết
                        </button>
                    </div>
                </h2>
                
                <!-- Stats bar -->
                <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color:white; padding:12px 16px; border-radius:10px; margin-bottom:15px; display:flex; justify-content:space-between; align-items:center;">
                    <div>
                        <span style="font-size:0.85em; opacity:0.9;">Tổng số tin nhắn:</span>
                        <span style="font-weight:700; font-size:1.1em; margin-left:8px;" id="total-messages">0</span>
                    </div>
                    <div style="font-size:0.85em; opacity:0.9;" id="last-update">Chưa có dữ liệu</div>
                </div>
                
                <!-- Chat container (WeChat style) -->
                <div id="chat-container" style="flex:1; overflow-y:auto; background:#f5f5f5; border-radius:10px; padding:15px; display:flex; flex-direction:column; gap:12px;">
                    <!-- Messages will be rendered here -->
                </div>
            </div>
        </div>
        
        <!-- MUSIC PLAYER SECTION - VLC Web Interface Style -->
        <div id="music-section" style="display:none;">
            <!-- Source Priority Selector -->
            <div style="background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%); border-radius: 15px; padding: 20px; margin-bottom: 20px; box-shadow: 0 10px 30px rgba(0,0,0,0.3);">
                <div style="display: flex; align-items: center; justify-content: space-between; flex-wrap: wrap; gap: 15px;">
                    <div style="display: flex; align-items: center; gap: 15px;">
                        <span style="color: #fff; font-weight: 600; font-size: 1.1em;">🎯 Nguồn phát ưu tiên:</span>
                        <div style="display: flex; gap: 10px;">
                            <button id="source-library-btn" onclick="setMusicSource('library')" 
                                    style="padding: 10px 20px; border-radius: 25px; border: 2px solid #667eea; background: #667eea; color: white; font-weight: 600; cursor: pointer; transition: all 0.3s;">
                                📚 Music Library
                            </button>
                            <button id="source-user-btn" onclick="setMusicSource('user')" 
                                    style="padding: 10px 20px; border-radius: 25px; border: 2px solid #667eea; background: transparent; color: #667eea; font-weight: 600; cursor: pointer; transition: all 0.3s;">
                                📁 Thư mục cá nhân
                            </button>
                        </div>
                    </div>
                    <div id="current-source-info" style="color: #a5b4fc; font-size: 0.9em;">
                        Đang dùng: <span id="source-path-display" style="font-family: monospace;">music_library/</span>
                    </div>
                </div>
            </div>
            
            <!-- VLC-style Player -->
            <div class="music-player" style="position:relative; background: linear-gradient(135deg, #2b3e50 0%, #1a252f 100%); border-radius: 15px; padding: 25px; box-shadow: 0 15px 40px rgba(0,0,0,0.4);">
                <!-- Album Art & Track Info -->
                <div style="display: flex; align-items: center; gap: 20px; margin-bottom: 20px;">
                    <div id="album-art" style="width: 120px; height: 120px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 12px; display: flex; align-items: center; justify-content: center; font-size: 48px; box-shadow: 0 8px 25px rgba(102, 126, 234, 0.4);">
                        🎵
                    </div>
                    <div class="now-playing" style="flex: 1;">
                        <h3 id="current-track" style="color: #fff; font-size: 1.4em; margin-bottom: 8px;">🎵 Chưa phát nhạc</h3>
                        <p id="track-info" style="color: #a5b4fc; font-size: 0.95em; margin-bottom: 5px;">Chọn bài hát từ danh sách bên dưới</p>
                        <!-- Audio Visualizer - Sóng nhạc -->
                        <div id="audio-visualizer" class="audio-visualizer paused" style="display: none;">
                            <div class="bar"></div>
                            <div class="bar"></div>
                            <div class="bar"></div>
                            <div class="bar"></div>
                            <div class="bar"></div>
                            <div class="bar"></div>
                            <div class="bar"></div>
                            <div class="bar"></div>
                            <div class="bar"></div>
                            <div class="bar"></div>
                        </div>
                        <p id="track-album" style="color: #6b7280; font-size: 0.85em;"></p>
                    </div>
                </div>
                
                <!-- Progress Bar (VLC style) - DRAGGABLE -->
                <div class="progress-container" style="margin-bottom: 20px;">
                    <input type="range" id="progress-slider" min="0" max="100" value="0" step="0.1"
                           oninput="onProgressDrag(this.value)" 
                           onchange="onProgressSeek(this.value)"
                           onmousedown="isDraggingProgress = true"
                           onmouseup="isDraggingProgress = false"
                           style="width: 100%; height: 8px; -webkit-appearance: none; background: linear-gradient(to right, #667eea 0%, #667eea 0%, #374151 0%, #374151 100%); border-radius: 4px; cursor: pointer; margin: 0;">
                    <div class="progress-time" style="display: flex; justify-content: space-between; margin-top: 8px; color: #9ca3af; font-size: 0.85em; font-family: monospace;">
                        <span id="current-time">0:00</span>
                        <span id="total-time">0:00</span>
                    </div>
                </div>
                
                <!-- Player Controls (VLC style) -->
                <div class="player-controls" style="display: flex; align-items: center; justify-content: center; gap: 15px; margin-bottom: 20px;">
                    <div class="player-btn" id="shuffle-btn" onclick="toggleShuffle()" title="Phát ngẫu nhiên" style="opacity: 0.6; cursor: pointer; font-size: 1.3em; padding: 10px; transition: all 0.2s;">🔀</div>
                    <div class="player-btn" onclick="musicPrevious()" title="Bài trước" style="cursor: pointer; font-size: 1.5em; padding: 10px;">⏮️</div>
                    <div class="player-btn play" onclick="musicPlayPause()" id="play-btn" title="Phát/Tạm dừng" style="cursor: pointer; font-size: 2.5em; padding: 15px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 50%; box-shadow: 0 5px 20px rgba(102, 126, 234, 0.5);">▶️</div>
                    <div class="player-btn" onclick="musicNext()" title="Bài tiếp" style="cursor: pointer; font-size: 1.5em; padding: 10px;">⏭️</div>
                    <div class="player-btn" id="repeat-btn" onclick="toggleRepeat()" title="Lặp lại" style="opacity: 0.6; cursor: pointer; font-size: 1.3em; padding: 10px; transition: all 0.2s;">🔁</div>
                    <div class="player-btn" onclick="musicStop()" title="Dừng" style="cursor: pointer; font-size: 1.3em; padding: 10px;">⏹️</div>
                </div>
                
                <!-- Volume Control (VLC style) -->
                <div style="display: flex; align-items: center; justify-content: center; gap: 15px; padding: 10px 0;">
                    <span onclick="toggleMute()" style="cursor: pointer; font-size: 1.3em;" id="volume-icon">🔊</span>
                    <input type="range" id="volume-slider" min="0" max="100" value="80" 
                           oninput="setPlayerVolume(this.value)"
                           style="width: 200px; height: 6px; -webkit-appearance: none; background: linear-gradient(to right, #667eea 0%, #667eea 80%, #374151 80%, #374151 100%); border-radius: 3px; cursor: pointer;">
                    <span id="volume-value" style="color: #9ca3af; font-size: 0.85em; min-width: 40px;">80%</span>
                </div>
            </div>
            
            <!-- Music Library with Search -->
            <div class="music-list" style="margin-top: 20px; background: white; border-radius: 15px; padding: 20px; box-shadow: 0 10px 30px rgba(0,0,0,0.12);">
                <!-- Hướng dẫn sử dụng -->
                <div style="background: linear-gradient(135deg, #e0e7ff 0%, #f3e8ff 100%); border-left: 4px solid #667eea; padding: 12px 15px; border-radius: 8px; margin-bottom: 15px; display: flex; align-items: center; gap: 10px;">
                    <span style="font-size: 1.3em;">💡</span>
                    <div style="flex: 1;">
                        <strong style="color: #667eea;">Hướng dẫn:</strong>
                        <span style="color: #4b5563; font-size: 0.9em;"> Click vào bài hát để phát ngay (hoặc click nút ▶️ khi hover)</span>
                    </div>
                </div>
                
                <div style="display: flex; align-items: center; justify-content: space-between; flex-wrap: wrap; gap: 15px; margin-bottom: 15px;">
                    <h3 style="margin: 0; color: #333;">📁 Thư Viện Nhạc</h3>
                    <div style="display: flex; gap: 10px; align-items: center;">
                        <input type="text" id="music-search" placeholder="🔍 Tìm bài hát..." 
                               oninput="filterMusicLibrary(this.value)"
                               style="padding: 10px 15px; border: 2px solid #e5e7eb; border-radius: 25px; width: 250px; font-size: 0.95em;">
                        <button onclick="loadMusicLibrary()" style="padding: 10px 20px; background: #667eea; color: white; border: none; border-radius: 25px; cursor: pointer; font-weight: 600;">🔄 Làm mới</button>
                    </div>
                </div>
                <div id="music-library" style="max-height: 400px; overflow-y: auto;">
                    <div style="text-align:center; padding:40px; color:#999;">
                        <p style="font-size:1.2em; margin-bottom:10px;">⏳ Đang tải danh sách nhạc...</p>
                        <button onclick="loadMusicLibrary()" style="padding:12px 24px; background:#667eea; color:white; border:none; border-radius:8px; cursor:pointer; font-size:1em;">Tải ngay</button>
                    </div>
                </div>
            </div>
        </div>
        
        <!-- RUNCAT ANIMATION (góc phải dưới) -->
        <div id="runcat-container">
            <div id="runcat">🐱</div>
        </div>

        <!-- MUSIC SETTINGS SECTION -->
        <div id="music-settings-section" style="display:none;">
            <div style="background: white; border-radius: 15px; padding: 30px; box-shadow: 0 10px 30px rgba(0,0,0,0.12);">
                <h2 style="color:#667eea; margin-bottom: 20px; display: flex; align-items: center; gap: 10px;">
                    ⚙️ Cài Đặt Thư Mục Nhạc
                </h2>
                
                <div style="background: #f8f9fa; padding: 25px; border-radius: 12px; margin-bottom: 20px; border-left: 4px solid #667eea;">
                    <h3 style="color: #333; margin-bottom: 15px; font-size: 1.1em;">📁 Đường Dẫn Thư Mục Nhạc</h3>
                    <p style="color: #666; margin-bottom: 15px; line-height: 1.6;">
                        Nhập đường dẫn đến thư mục chứa nhạc của bạn. miniZ sẽ ưu tiên phát nhạc từ thư mục này bằng trình phát mặc định của Windows.
                    </p>
                    
                    <div style="display: flex; gap: 10px; margin-bottom: 15px;">
                        <input type="text" id="music-folder-path" placeholder="Ví dụ: F:\My Music hoặc C:\Users\Name\Music" 
                               style="flex: 1; padding: 12px 15px; border: 2px solid #e5e7eb; border-radius: 8px; font-size: 1em; font-family: 'Consolas', monospace;">
                        <button onclick="browseMusicFolder()" 
                                style="padding: 12px 20px; background: #667eea; color: white; border: none; border-radius: 8px; cursor: pointer; font-weight: 600; white-space: nowrap;">
                            📂 Chọn Thư Mục
                        </button>
                        <button onclick="saveMusicFolder()" 
                                style="padding: 12px 20px; background: #10b981; color: white; border: none; border-radius: 8px; cursor: pointer; font-weight: 600; white-space: nowrap;">
                            💾 Lưu
                        </button>
                    </div>
                    
                    <div id="music-folder-status" style="margin-top: 10px; padding: 10px; border-radius: 6px; display: none;"></div>
                </div>
                
                <div style="background: #fff3cd; padding: 20px; border-radius: 12px; border-left: 4px solid #ffc107; margin-bottom: 20px;">
                    <h3 style="color: #856404; margin-bottom: 10px; font-size: 1em;">💡 Lưu Ý</h3>
                    <ul style="color: #856404; line-height: 1.8; margin-left: 20px;">
                        <li>Sau khi lưu, bạn có thể yêu cầu LLM phát nhạc từ thư mục này</li>
                        <li>miniZ sẽ dùng trình phát mặc định của Windows (Windows Media Player, Groove Music, VLC...)</li>
                        <li>Ví dụ lệnh: "<i>Phát nhạc trong thư mục của tôi</i>" hoặc "<i>Play all songs</i>"</li>
                    </ul>
                </div>
                
                <div style="background: #e8f4f8; padding: 20px; border-radius: 12px; border-left: 4px solid #3b82f6;">
                    <h3 style="color: #1e40af; margin-bottom: 10px; font-size: 1em;">🎵 Định Dạng Hỗ Trợ</h3>
                    <div style="display: flex; flex-wrap: wrap; gap: 10px; margin-top: 10px;">
                        <span style="background: white; padding: 6px 12px; border-radius: 6px; font-size: 0.9em; color: #1e40af; font-weight: 600;">.mp3</span>
                        <span style="background: white; padding: 6px 12px; border-radius: 6px; font-size: 0.9em; color: #1e40af; font-weight: 600;">.wav</span>
                        <span style="background: white; padding: 6px 12px; border-radius: 6px; font-size: 0.9em; color: #1e40af; font-weight: 600;">.flac</span>
                        <span style="background: white; padding: 6px 12px; border-radius: 6px; font-size: 0.9em; color: #1e40af; font-weight: 600;">.m4a</span>
                        <span style="background: white; padding: 6px 12px; border-radius: 6px; font-size: 0.9em; color: #1e40af; font-weight: 600;">.wma</span>
                        <span style="background: white; padding: 6px 12px; border-radius: 6px; font-size: 0.9em; color: #1e40af; font-weight: 600;">.aac</span>
                        <span style="background: white; padding: 6px 12px; border-radius: 6px; font-size: 0.9em; color: #1e40af; font-weight: 600;">.ogg</span>
                    </div>
                </div>
            </div>
        </div>
        
        <!-- PLAYLIST SECTION -->
        <div id="playlist-section" style="display:none;">
            <div style="background: white; border-radius: 15px; padding: 20px; box-shadow: 0 10px 30px rgba(0,0,0,0.12);">
                <h2 style="color:#667eea; margin-bottom: 12px; display: flex; align-items: center; justify-content: space-between; gap: 15px;">
                    <span>🎵 Danh Sách Nhạc YouTube</span>
                    <div style="display:flex; align-items:center; gap:10px;">
                        <input id="playlist-command" placeholder="Gõ từ khóa playlist (vd: nhạc, chill...)" style="padding:8px 12px; border-radius:8px; border:1px solid #e5e7eb; font-size:0.95em; width:280px;" 
                               onkeypress="if(event.key==='Enter') triggerPlayByName(this.value.trim())" />
                        <button onclick="triggerPlayByName(document.getElementById('playlist-command').value.trim())" style="padding:8px 12px; background:#667eea; color:white; border:none; border-radius:8px; cursor:pointer;">Mở</button>
                    </div>
                </h2>

                <div style="display:flex; gap:20px; align-items:flex-start;">
                    <div style="flex:1;">
                        <div id="playlist-list" style="background:#f9fafb; padding:12px; border-radius:8px; min-height:80px; border:1px solid #e5e7eb;">
                            <!-- playlists will be rendered here -->
                        </div>
                        <div style="margin-top:12px; display:flex; gap:10px;">
                            <button onclick="promptAddPlaylist()" style="padding:10px 14px; border-radius:8px; background:linear-gradient(135deg,#10b981,#059669); color:white; border:none; cursor:pointer; font-weight:600;">＋ Thêm Playlist</button>
                            <button onclick="renderPlaylists()" style="padding:10px 14px; border-radius:8px; background:#e5e7eb; border:none; cursor:pointer;">Làm mới</button>
                        </div>
                    </div>
                    <div style="width:320px;">
                        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color:white; padding:14px; border-radius:12px;">
                            <div style="font-weight:700; margin-bottom:6px;">Hướng dẫn nhanh</div>
                            <div style="font-size:0.95em; opacity:0.95;">
                                • Nhấn <b>＋ Thêm Playlist</b> để thêm mới (tên + URL)<br>
                                • Gõ <b>từ khóa</b> (không cần chính xác) vào ô và nhấn <b>Mở</b><br>
                                • Ví dụ: gõ "nhạc" sẽ tìm "Nhạc chill", "Nhạc EDM"...<br>
                                • Voice: "mở danh sách [từ khóa]" hoặc "mở playlist [từ khóa]"
                            </div>
                        </div>
                    </div>
                </div>
            </div>
        </div>
        
        <!-- KNOWLEDGE BASE SECTION -->
        <div id="knowledge-section" style="display:none;">
            <div style="background: white; border-radius: 15px; padding: 25px; box-shadow: 0 10px 30px rgba(0,0,0,0.12);">
                <h2 style="color:#667eea; margin-bottom: 20px; display: flex; align-items: center; gap: 10px;">
                    <span>📚 Knowledge Base</span>
                    <span style="font-size: 0.5em; color: #9ca3af; font-weight: 400;">Cập nhật dữ liệu cho LLM</span>
                </h2>
                
                <!-- Nhập đường dẫn thư mục -->
                <div style="background: #f9fafb; padding: 20px; border-radius: 12px; border: 2px solid #e5e7eb; margin-bottom: 20px;">
                    <h3 style="color: #333; margin-bottom: 15px; display: flex; align-items: center; gap: 8px;">
                        📁 Thư Mục Dữ Liệu
                    </h3>
                    <div style="display: flex; gap: 10px; align-items: center;">
                        <input type="text" id="knowledge-folder-path" 
                               placeholder="Nhập đường dẫn thư mục (VD: C:\Documents\MyData hoặc D:\Knowledge)" 
                               style="flex: 1; padding: 12px 15px; border: 2px solid #e5e7eb; border-radius: 8px; font-size: 1em;">
                        <button onclick="saveKnowledgeFolder()" 
                                style="padding: 12px 25px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; border: none; border-radius: 8px; font-weight: 600; cursor: pointer; white-space: nowrap;">
                            💾 Lưu
                        </button>
                        <button onclick="scanKnowledgeFolder()" 
                                style="padding: 12px 25px; background: linear-gradient(135deg, #10b981 0%, #059669 100%); color: white; border: none; border-radius: 8px; font-weight: 600; cursor: pointer; white-space: nowrap;">
                            🔍 Quét Files
                        </button>
                    </div>
                    <p style="color: #666; font-size: 0.9em; margin-top: 10px;">
                        💡 Hỗ trợ: PDF, TXT, Word (.docx), Markdown (.md), JSON, CSV
                    </p>
                </div>
                
                <!-- Trạng thái & thống kê -->
                <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 15px; margin-bottom: 20px;">
                    <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 20px; border-radius: 12px; text-align: center;">
                        <div style="font-size: 2em; font-weight: bold;" id="kb-total-files">0</div>
                        <div style="opacity: 0.9;">Tổng số files</div>
                    </div>
                    <div style="background: linear-gradient(135deg, #10b981 0%, #059669 100%); color: white; padding: 20px; border-radius: 12px; text-align: center;">
                        <div style="font-size: 2em; font-weight: bold;" id="kb-indexed-files">0</div>
                        <div style="opacity: 0.9;">Đã index</div>
                    </div>
                    <div style="background: linear-gradient(135deg, #f59e0b 0%, #d97706 100%); color: white; padding: 20px; border-radius: 12px; text-align: center;">
                        <div style="font-size: 2em; font-weight: bold;" id="kb-total-size">0 KB</div>
                        <div style="opacity: 0.9;">Dung lượng</div>
                    </div>
                    <div style="background: linear-gradient(135deg, #8b5cf6 0%, #7c3aed 100%); color: white; padding: 20px; border-radius: 12px; text-align: center;">
                        <div style="font-size: 2em; font-weight: bold;" id="kb-last-update">--</div>
                        <div style="opacity: 0.9;">Cập nhật lần cuối</div>
                    </div>
                </div>
                
                <!-- Danh sách files -->
                <div style="background: #f9fafb; padding: 20px; border-radius: 12px; border: 2px solid #e5e7eb;">
                    <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 15px;">
                        <h3 style="color: #333; display: flex; align-items: center; gap: 8px; margin: 0;">
                            📄 Danh Sách Files
                        </h3>
                        <div style="display: flex; gap: 10px;">
                            <button onclick="indexAllFiles()" 
                                    style="padding: 8px 16px; background: #667eea; color: white; border: none; border-radius: 6px; cursor: pointer; font-size: 0.9em;">
                                🔄 Index Tất Cả
                            </button>
                            <button onclick="clearKnowledgeBase()" 
                                    style="padding: 8px 16px; background: #ef4444; color: white; border: none; border-radius: 6px; cursor: pointer; font-size: 0.9em;">
                                🗑️ Xóa Index
                            </button>
                        </div>
                    </div>
                    <div id="knowledge-file-list" style="max-height: 400px; overflow-y: auto;">
                        <p style="color: #666; text-align: center; padding: 40px;">
                            📂 Chưa có thư mục nào được cấu hình.<br>
                            Nhập đường dẫn và nhấn "Quét Files" để bắt đầu.
                        </p>
                    </div>
                </div>
                
                <!-- Hướng dẫn -->
                <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 20px; border-radius: 12px; margin-top: 20px;">
                    <h3 style="margin-bottom: 12px;">📖 Hướng Dẫn Sử Dụng</h3>
                    <div style="font-size: 0.95em; line-height: 1.6;">
                        <p>1. <strong>Nhập đường dẫn</strong> thư mục chứa tài liệu (PDF, TXT, Word, Markdown...)</p>
                        <p>2. <strong>Nhấn "Quét Files"</strong> để liệt kê các files trong thư mục</p>
                        <p>3. <strong>Nhấn "Index Tất Cả"</strong> để LLM học từ nội dung các files</p>
                        <p>4. Sau khi index, LLM có thể trả lời câu hỏi dựa trên dữ liệu của bạn!</p>
                        <p style="margin-top: 10px; opacity: 0.9;">
                            💡 <strong>Mẹo:</strong> Đặt các tài liệu quan trọng vào một thư mục riêng để dễ quản lý.
                        </p>
                    </div>
                </div>
            </div>
        </div>
        
        <!-- SETTINGS MODAL -->
        <div id="settingsModal" class="modal">
            <div class="modal-content" style="max-width:1400px;width:95%;">
                <div class="modal-header">
                    <h2>⚙️ Cấu hình Endpoint</h2>
                    <button class="close-btn" onclick="closeSettingsModal()">&times;</button>
                </div>
                <div class="modal-body">
                    <!-- 3 ENDPOINT SECTIONS -->
                    <div style="display:grid;grid-template-columns:1fr 1fr 1fr;gap:20px;margin-bottom:25px;">
                        <!-- Thiết bị 1 -->
                        <div id="device-1-card" style="border:2px solid #10b981;border-radius:8px;padding:15px;background:#f0fdf4;position:relative;">
                            <div style="position:absolute;top:10px;right:10px;">
                                <span id="device-1-indicator" class="connection-indicator" style="display:inline-flex;align-items:center;gap:5px;padding:4px 10px;border-radius:12px;background:#d1fae5;color:#047857;font-size:0.75em;font-weight:bold;">
                                    <span class="status-dot" style="width:8px;height:8px;border-radius:50%;background:#6b7280;"></span>
                                    Chưa kết nối
                                </span>
                            </div>
                            <label for="endpoint-url-1" style="color:#047857;font-weight:600;display:flex;align-items:center;gap:8px;">
                                📱 Thiết bị 1
                            </label>
                            <input type="text" id="endpoint-url-1" placeholder="JWT token thiết bị 1..." style="margin-top:8px;border:2px solid #10b981;" />
                            <p style="color:#065f46;font-size:0.85em;margin-top:5px;margin-bottom:0;">
                                Token thật từ Claude Desktop
                            </p>
                        </div>
                        
                        <!-- Thiết bị 2 -->
                        <div id="device-2-card" style="border:2px solid #3b82f6;border-radius:8px;padding:15px;background:#eff6ff;position:relative;">
                            <div style="position:absolute;top:10px;right:10px;">
                                <span id="device-2-indicator" class="connection-indicator" style="display:inline-flex;align-items:center;gap:5px;padding:4px 10px;border-radius:12px;background:#dbeafe;color:#1e40af;font-size:0.75em;font-weight:bold;">
                                    <span class="status-dot" style="width:8px;height:8px;border-radius:50%;background:#6b7280;"></span>
                                    Chưa kết nối
                                </span>
                            </div>
                            <label for="endpoint-url-2" style="color:#1e40af;font-weight:600;display:flex;align-items:center;gap:8px;">
                                📱 Thiết bị 2
                            </label>
                            <input type="text" id="endpoint-url-2" placeholder="JWT token thiết bị 2..." style="margin-top:8px;border:2px solid #3b82f6;" />
                            <p style="color:#1e3a8a;font-size:0.85em;margin-top:5px;margin-bottom:0;">
                                MCP connection 2
                            </p>
                        </div>
                        
                        <!-- Thiết bị 3 -->
                        <div id="device-3-card" style="border:2px solid #f59e0b;border-radius:8px;padding:15px;background:#fffbeb;position:relative;">
                            <div style="position:absolute;top:10px;right:10px;">
                                <span id="device-3-indicator" class="connection-indicator" style="display:inline-flex;align-items:center;gap:5px;padding:4px 10px;border-radius:12px;background:#fef3c7;color:#b45309;font-size:0.75em;font-weight:bold;">
                                    <span class="status-dot" style="width:8px;height:8px;border-radius:50%;background:#6b7280;"></span>
                                    Chưa kết nối
                                </span>
                            </div>
                            <label for="endpoint-url-3" style="color:#b45309;font-weight:600;display:flex;align-items:center;gap:8px;">
                                📱 Thiết bị 3
                            </label>
                            <input type="text" id="endpoint-url-3" placeholder="JWT token thiết bị 3..." style="margin-top:8px;border:2px solid #f59e0b;" />
                            <p style="color:#78350f;font-size:0.85em;margin-top:5px;margin-bottom:0;">
                                MCP connection 3
                            </p>
                        </div>
                    </div>
                    
                    <p style="color:#666;font-size:0.9em;text-align:center;margin-top:-10px;margin-bottom:20px;">
                        <strong>Lưu ý:</strong> Có thể nhập JWT token trực tiếp hoặc URL đầy đủ <code>wss://api.xiaozhi.me/mcp/?token=...</code> - hệ thống sẽ tự động xử lý
                    </p>
                    
                    <hr style="margin:25px 0;border:none;border-top:2px solid #e5e7eb;">
                    
                    <!-- API KEYS GRID (2 Columns) -->
                    <div style="display:grid;grid-template-columns:1fr 1fr;gap:30px;">
                        <!-- LEFT COLUMN: Gemini -->
                        <div style="border-right:2px solid #e5e7eb;padding-right:30px;">
                            <label for="gemini-api-key" style="display:flex;align-items:center;gap:10px;">
                                🤖 Gemini API Key 
                                <span style="color:#10b981;font-size:0.85em;font-weight:normal;">(Auto-save)</span>
                            </label>
                            <div class="api-key-input-container">
                                <input 
                                    type="password" 
                                    id="gemini-api-key" 
                                    placeholder="AIzaSyXXXXXXXXXXXXXXXXXX..."
                                    oninput="autoSaveGeminiKey()"
                                    style="font-size:0.9em;"
                                />
                                <div class="input-icons">
                                    <button type="button" class="api-key-icon-btn" onclick="toggleApiKeyVisibility('gemini-api-key', this)" title="Hiện/Ẩn API key">
                                        �
                                    </button>
                                    <button type="button" class="api-key-icon-btn" onclick="copyApiKey('gemini-api-key', this)" title="Copy API key">
                                        📋
                                    </button>
                                </div>
                            </div>
                            <p style="color:#666;font-size:0.9em;margin-top:-10px;">
                                <strong>Miễn phí:</strong> Lấy API key tại 
                                <a href="https://aistudio.google.com/apikey" target="_blank" style="color:#667eea;">
                                    aistudio.google.com/apikey
                                </a>
                                <br>
                                <span id="gemini-key-status" style="color:#10b981;font-weight:600;"></span>
                            </p>
                            
                            <label for="gemini-model" style="margin-top:15px;display:block;">
                                🎯 Gemini Model
                            </label>
                            <select 
                                id="gemini-model" 
                                onchange="saveGeminiModel()"
                                style="width:100%;padding:10px;border:2px solid #e5e7eb;border-radius:8px;font-size:0.95em;"
                            >
                                <option value="models/gemini-3-flash-preview">⚡ Gemini 3 Flash Preview (Mới nhất)</option>
                                <option value="models/gemini-2.5-flash">⚡ Gemini 2.5 Flash (Ổn định)</option>
                                <option value="models/gemini-2.5-pro">💎 Gemini 2.5 Pro (Chất lượng cao nhất)</option>
                                <option value="models/gemini-2.0-flash-exp">⚡ Gemini 2.0 Flash Exp</option>
                                <option value="models/gemini-1.5-pro">💎 Gemini 1.5 Pro (Ổn định)</option>
                                <option value="models/gemini-2.0-flash-thinking-exp">🧠 Gemini 2.0 Flash Thinking (Suy luận tốt)</option>
                                <option value="models/gemini-1.5-pro">💎 Gemini 1.5 Pro (Ổn định)</option>
                                <option value="models/gemini-1.5-flash">⚡ Gemini 1.5 Flash (Cân bằng)</option>
                            </select>
                            <p style="color:#666;font-size:0.85em;margin-top:5px;">
                                💡 <strong>3.0 Flash:</strong> Model mới nhất (12/2024), giảm 30% token | <strong>2.5 Pro:</strong> Chất lượng cao nhất | <strong>1.5 Pro:</strong> Ổn định
                            </p>
                        </div>
                        
                        <!-- RIGHT COLUMN: OpenAI + Serper -->
                        <div style="padding-left:30px;">
                            <label for="openai-api-key" style="display:flex;align-items:center;gap:10px;">
                                🧠 OpenAI API Key (GPT-4)
                                <span style="color:#10b981;font-size:0.85em;font-weight:normal;">(Auto-save)</span>
                                <span style="color:#ef4444;font-size:0.75em;font-weight:normal;">TRẢ PHÍ</span>
                            </label>
                            <div class="api-key-input-container">
                                <input 
                                    type="password" 
                                    id="openai-api-key" 
                                    placeholder="sk-proj-XXXXXXXXXXXXXXXXXX..."
                                    oninput="autoSaveOpenAIKey()"
                                    style="font-size:0.9em;"
                                />
                                <div class="input-icons">
                                    <button type="button" class="api-key-icon-btn" onclick="toggleApiKeyVisibility('openai-api-key', this)" title="Hiện/Ẩn API key">
                                        �
                                    </button>
                                    <button type="button" class="api-key-icon-btn" onclick="copyApiKey('openai-api-key', this)" title="Copy API key">
                                        📋
                                    </button>
                                </div>
                            </div>
                            <p style="color:#666;font-size:0.9em;margin-top:-10px;">
                                <strong>Trả phí:</strong> Lấy API key tại 
                                <a href="https://platform.openai.com/api-keys" target="_blank" style="color:#667eea;">
                                    platform.openai.com/api-keys
                                </a>
                                <br>
                                <span style="font-size:0.85em;">💰 Giá: $0.01-0.03/1K tokens | 🆓 Free trial: $5 credit</span>
                                <br>
                                <span id="openai-key-status" style="color:#10b981;font-weight:600;"></span>
                            </p>
                            
                            <hr style="margin:20px 0;border:none;border-top:1px solid #e5e7eb;">
                            
                            <label for="serper-api-key" style="display:flex;align-items:center;gap:10px;margin-top:20px;">
                                🔍 Serper API Key (Google Search)
                                <span style="color:#10b981;font-size:0.85em;font-weight:normal;">(Auto-save)</span>
                                <span style="color:#22c55e;font-size:0.75em;font-weight:normal;">MIỄN PHÍ 2500/tháng</span>
                            </label>
                            <div class="api-key-input-container">
                                <input 
                                    type="password" 
                                    id="serper-api-key" 
                                    placeholder="abcdef1234567890..."
                                    oninput="autoSaveSerperKey()"
                                    style="font-size:0.9em;"
                                />
                                <div class="input-icons">
                                    <button type="button" class="api-key-icon-btn" onclick="toggleApiKeyVisibility('serper-api-key', this)" title="Hiện/Ẩn API key">
                                        �
                                    </button>
                                    <button type="button" class="api-key-icon-btn" onclick="copyApiKey('serper-api-key', this)" title="Copy API key">
                                        📋
                                    </button>
                                </div>
                            </div>
                            <p style="color:#666;font-size:0.9em;margin-top:-10px;">
                                <strong>Miễn phí:</strong> Đăng ký tại 
                                <a href="https://serper.dev" target="_blank" style="color:#667eea;">
                                    serper.dev
                                </a>
                                <br>
                                <span style="font-size:0.85em;">🆓 2500 queries/tháng miễn phí | 🎯 Google Search chính xác hơn DuckDuckGo</span>
                                <br>
                                <span id="serper-key-status" style="color:#10b981;font-weight:600;"></span>
                            </p>
                        </div>
                    </div>
                </div>
                <div class="modal-footer">
                    <button class="modal-btn secondary" onclick="closeSettingsModal()">Hủy</button>
                    <button class="modal-btn info" onclick="copyFullUrl()">📋 Copy URL đầy đủ</button>
                    <button class="modal-btn primary" onclick="saveEndpoint()">💾 Lưu</button>
                </div>
            </div>
        </div>
    </div>
    
    <script>
        let ws;
        let llmChatMessages = []; // Store LLM chat messages
        
        // Section switching
        function showSection(name) {
            document.querySelectorAll('.menu-item').forEach(item => item.classList.remove('active'));
            event.target.classList.add('active');
            
            document.getElementById('dashboard-section').style.display = name === 'dashboard' ? 'block' : 'none';
            document.getElementById('tools-section').style.display = name === 'tools' ? 'block' : 'none';
            document.getElementById('llm-chat-section').style.display = name === 'llm-chat' ? 'block' : 'none';
            document.getElementById('api-quotas-section').style.display = name === 'api-quotas' ? 'block' : 'none';
            document.getElementById('music-section').style.display = name === 'music' ? 'block' : 'none';
            document.getElementById('music-settings-section').style.display = name === 'music-settings' ? 'block' : 'none';
            document.getElementById('conversation-section').style.display = name === 'conversation' ? 'block' : 'none';
            document.getElementById('playlist-section').style.display = name === 'playlist' ? 'block' : 'none';
            document.getElementById('knowledge-section').style.display = name === 'knowledge' ? 'block' : 'none';
            
            // Load API Quotas when opening api-quotas section
            if (name === 'api-quotas') {
                refreshQuotasPage();
            }
            
            // Load LLM Chat section
            if (name === 'llm-chat') {
                loadLLMChatModel();
            }
            
            // Load conversation when opening conversation section
            if (name === 'conversation') {
                loadConversationHistory();
            }
            
            // Load music library when opening music section
            if (name === 'music') {
                loadMusicSourcePreference();
                updateMusicStatus();
            }
            if (name === 'music-settings') {
                loadMusicFolderSettings();
            }
            
            // Load playlist when opening playlist section
            if (name === 'playlist') {
                // use initPlaylists() (render existing playlists) - loadPlaylistSection was removed
                initPlaylists();
            }
            
            // Load knowledge base when opening knowledge section
            if (name === 'knowledge') {
                loadKnowledgeBase();
            }
        }
        
        // ===== API QUOTAS PAGE FUNCTIONS =====
        async function refreshQuotasPage() {
            addLog('🔄 Đang làm mới API Quotas...', 'info');
            try {
                const response = await fetch('/api/quotas');
                const data = await response.json();
                
                if (data.success) {
                    // Update Gemini status
                    const geminiStatus = document.getElementById('gemini-status');
                    if (data.gemini && data.gemini.has_key) {
                        geminiStatus.innerHTML = '✅ API Key đã cấu hình';
                        geminiStatus.style.color = '#10b981';
                    } else {
                        geminiStatus.innerHTML = '❌ Chưa có API Key';
                        geminiStatus.style.color = '#ef4444';
                    }
                    
                    // Update Serper status
                    const serperStatus = document.getElementById('serper-status');
                    if (data.serper && data.serper.has_key) {
                        serperStatus.innerHTML = '✅ API Key đã cấu hình';
                        serperStatus.style.color = '#10b981';
                    } else {
                        serperStatus.innerHTML = '❌ Chưa có API Key';
                        serperStatus.style.color = '#ef4444';
                    }
                    
                    addLog('✅ Đã làm mới API Quotas', 'success');
                } else {
                    addLog('❌ Lỗi: ' + (data.error || 'Unknown error'), 'error');
                }
            } catch (error) {
                console.error('Error refreshing quotas:', error);
                addLog('❌ Lỗi kết nối: ' + error.message, 'error');
            }
        }
        
        async function testGeminiAPI() {
            addLog('🧪 Đang test Gemini API...', 'info');
            try {
                const response = await fetch('/api/chat', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({message: 'Hello, this is a test message. Reply with OK.'})
                });
                const data = await response.json();
                
                if (data.response) {
                    addLog('✅ Gemini API hoạt động tốt! Response: ' + data.response.substring(0, 100) + '...', 'success');
                } else {
                    addLog('❌ Gemini API test thất bại', 'error');
                }
            } catch (error) {
                console.error('Error testing Gemini:', error);
                addLog('❌ Gemini test error: ' + error.message, 'error');
            }
        }
        
        async function testSerperAPI() {
            addLog('🧪 Đang test Serper API...', 'info');
            try {
                const response = await fetch('/api/google_search', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({query: 'test search'})
                });
                const data = await response.json();
                
                if (data.results && data.results.length > 0) {
                    addLog('✅ Serper API hoạt động tốt! Tìm thấy ' + data.results.length + ' kết quả', 'success');
                } else {
                    addLog('❌ Serper API test thất bại', 'error');
                }
            } catch (error) {
                console.error('Error testing Serper:', error);
                addLog('❌ Serper test error: ' + error.message, 'error');
            }
        }
        
        // Tab switching
        function switchTab(index) {
            document.querySelectorAll('.tab-btn').forEach((btn, i) => btn.classList.toggle('active', i === index));
            document.querySelectorAll('.tab-content').forEach((content, i) => content.classList.toggle('active', i === index));
        }
        
        // Quick actions - 20 tools
        function setVolumePrompt() {
            const level = prompt('Nhập âm lượng (0-100):', '50');
            if (level === null) return;
            const levelNum = parseInt(level);
            if (isNaN(levelNum) || levelNum < 0 || levelNum > 100) {
                addLog('❌ Âm lượng phải từ 0-100', 'error');
                return;
            }
            setVolumeQuick(levelNum);
        }
        function setVolumeQuick(level) { 
            if (level >= 0 && level <= 100) {
                callTool('set_volume', {level});
            } else {
                addLog('❌ Âm lượng phải từ 0-100', 'error');
            }
        }
        function getVolumeInfo() {
            callTool('get_volume', {});
        }
        function screenshot() { callAPI('/api/screenshot', {}); }
        function notification() { callAPI('/api/notification', {title: 'Xiaozhi', message: 'Test notification'}); }
        function setBrightness() { 
            const level = prompt('Nhập độ sáng (0-100):', '50');
            if (level === null) return;
            const levelNum = parseInt(level);
            if (isNaN(levelNum) || levelNum < 0 || levelNum > 100) {
                addLog('❌ Độ sáng phải từ 0-100', 'error');
                return;
            }
            callTool('set_brightness', {level: levelNum});
        }
        function openApp() {
            const app = prompt('Nhập tên app (notepad/calc/paint/cmd/explorer):', 'notepad');
            if (app && app.trim()) callTool('open_application', {app_name: app.trim()});
        }
        function listProcesses() { callTool('list_running_processes', {limit: 10}); }
        function killProcess() {
            const id = prompt('Nhập PID hoặc tên tiến trình:', 'chrome');
            if (id && id.trim()) callTool('kill_process', {identifier: id.trim()});
        }
        function createFile() {
            const path = prompt('Đường dẫn file:', 'C:/test.txt');
            if (!path || !path.trim()) return;
            const content = prompt('Nội dung:', 'Hello World');
            if (content !== null) callTool('create_file', {path: path.trim(), content});
        }
        function readFile() {
            const path = prompt('Đường dẫn file:', 'C:/test.txt');
            if (path && path.trim()) callTool('read_file', {path: path.trim()});
        }
        function listFiles() {
            const dir = prompt('Thư mục:', 'C:/Users');
            if (dir && dir.trim()) callTool('list_files', {directory: dir.trim()});
        }
        function diskUsage() { callTool('get_disk_usage', {}); }
        function networkInfo() {
            // Show loading message
            showResult('⏳ Đang quét mạng và các thiết bị...');
            addLog('🌐 Đang quét thông tin mạng...', 'info');
            
            fetch('/api/tool/get_network_info', {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify({})
            })
            .then(r => r.json())
            .then(data => {
                if (data.success) {
                    // Save to conversation history
                    const summary = `Quét mạng: ${data.total_devices} thiết bị - ${data.local_device.hostname} (${data.local_device.ip})`;
                    addToConversation('system', `🌐 Network Scan: ${summary}`);
                    addLog(`✅ ${summary}`, 'success');
                    
                    let html = '<div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 20px; border-radius: 10px; color: white;">';
                    html += '<h2 style="margin: 0 0 20px 0; font-size: 24px;">🌐 Thông Tin Mạng</h2>';
                    
                    // Local device info
                    html += '<div style="background: rgba(255,255,255,0.15); padding: 15px; border-radius: 8px; margin-bottom: 20px;">';
                    html += '<h3 style="margin: 0 0 10px 0; font-size: 18px;">💻 Máy Của Bạn</h3>';
                    html += `<p style="margin: 5px 0;"><strong>🏷️ Hostname:</strong> ${data.local_device.hostname}</p>`;
                    html += `<p style="margin: 5px 0;"><strong>🌐 IP Address:</strong> ${data.local_device.ip}</p>`;
                    html += `<p style="margin: 5px 0;"><strong>📡 MAC Address:</strong> ${data.local_device.mac}</p>`;
                    html += `<p style="margin: 5px 0;"><strong>🚪 Gateway (Router):</strong> ${data.local_device.gateway}</p>`;
                    html += '</div>';
                    
                    // Network devices
                    if (data.network_devices && data.network_devices.length > 0) {
                        html += '<div style="background: rgba(255,255,255,0.15); padding: 15px; border-radius: 8px;">';
                        html += `<h3 style="margin: 0 0 15px 0; font-size: 18px;">📱 Thiết Bị Trong Mạng (${data.total_devices})</h3>`;
                        html += '<div style="max-height: 400px; overflow-y: auto;">';
                        
                        data.network_devices.forEach((device, idx) => {
                            const bgColor = device.is_local ? 'rgba(76, 175, 80, 0.3)' : 'rgba(255,255,255,0.1)';
                            const icon = device.is_local ? '👤' : '🖥️';
                            html += `<div style="background: ${bgColor}; padding: 12px; border-radius: 6px; margin-bottom: 10px; border-left: 3px solid ${device.is_local ? '#4CAF50' : '#fff'};">`;
                            html += `<div style="display: flex; justify-content: space-between; align-items: center;">`;
                            html += `<div>`;
                            html += `<p style="margin: 0; font-size: 16px; font-weight: bold;">${icon} ${device.hostname}</p>`;
                            html += `<p style="margin: 5px 0 0 0; font-size: 14px; opacity: 0.9;">🌐 IP: ${device.ip}</p>`;
                            html += `<p style="margin: 3px 0 0 0; font-size: 13px; opacity: 0.8;">📡 MAC: ${device.mac}</p>`;
                            html += `</div>`;
                            if (device.is_local) {
                                html += `<div style="background: #4CAF50; padding: 5px 10px; border-radius: 4px; font-size: 12px; font-weight: bold;">BẠN</div>`;
                            }
                            html += `</div>`;
                            html += `</div>`;
                        });
                        
                        html += '</div></div>';
                    } else {
                        html += '<p style="text-align: center; opacity: 0.8; margin-top: 10px;">Không tìm thấy thiết bị khác trong mạng</p>';
                    }
                    
                    html += '</div>';
                    showResult(html);
                } else {
                    const error = data.error || 'Không thể lấy thông tin mạng';
                    showResult('❌ Lỗi: ' + error);
                    addLog(`❌ Lỗi quét mạng: ${error}`, 'error');
                }
            })
            .catch(err => {
                showResult('❌ Lỗi kết nối: ' + err.message);
                addLog(`❌ Lỗi kết nối: ${err.message}`, 'error');
            });
        }
        function batteryStatus() { callTool('get_battery_status', {}); }
        function searchWeb() {
            const query = prompt('Từ khóa tìm kiếm:', '');
            if (query && query.trim()) callTool('search_web', {query: query.trim()});
        }
        function calculator() {
            const expr = prompt('Biểu thức toán học:', '2+2*3');
            if (expr && expr.trim()) callAPI('/api/calculator', {expression: expr.trim()});
        }
        function getClipboard() { callTool('get_clipboard', {}); }
        function setClipboard() {
            const text = prompt('Nội dung cần copy:', '');
            if (text !== null && text.trim()) callTool('set_clipboard', {text: text.trim()});
        }
        function playSound() {
            const freq = prompt('Tần số Hz (200-2000):', '1000');
            if (freq === null) return;
            const dur = prompt('Thời gian ms (100-3000):', '500');
            if (dur === null) return;
            const freqNum = parseInt(freq);
            const durNum = parseInt(dur);
            if (isNaN(freqNum) || freqNum < 200 || freqNum > 2000) {
                addLog('❌ Tần số phải từ 200-2000 Hz', 'error');
                return;
            }
            if (isNaN(durNum) || durNum < 100 || durNum > 3000) {
                addLog('❌ Thời gian phải từ 100-3000 ms', 'error');
                return;
            }
            callTool('play_sound', {frequency: freqNum, duration: durNum});
        }
        
        // NEW TOOL FUNCTIONS
        function lockComputer() {
            if (confirm('Bạn có chắc muốn khóa máy tính?')) {
                callTool('lock_computer', {});
            }
        }
        function shutdownSchedule() {
            const action = prompt('Hành động (shutdown/restart/cancel):', 'shutdown');
            if (!action || !action.trim()) return;
            const actionLower = action.trim().toLowerCase();
            if (!['shutdown', 'restart', 'cancel'].includes(actionLower)) {
                addLog('❌ Hành động không hợp lệ. Dùng: shutdown, restart, hoặc cancel', 'error');
                return;
            }
            const delay = prompt('Trì hoãn (giây):', '60');
            if (delay === null) return;
            const delayNum = parseInt(delay) || 0;
            if (delayNum < 0) {
                addLog('❌ Thời gian trì hoãn phải >= 0', 'error');
                return;
            }
            callTool('shutdown_schedule', {action: actionLower, delay: delayNum});
        }
        function showDesktop() {
            callTool('show_desktop', {});
        }
        function undoOperation() {
            callTool('undo_operation', {});
        }
        function setTheme() {
            const dark = confirm('Chọn OK cho theme TỐI, Cancel cho theme SÁNG');
            callTool('set_theme', {dark_mode: dark});
        }
        function changeWallpaper() {
            const keyword = prompt('Từ khóa hình nền (hoặc để trống để chọn ngẫu nhiên):', '');
            callTool('change_wallpaper', {keyword: keyword || ''});
        }
        function getDesktopPath() {
            callTool('get_desktop_path', {});
        }
        function pasteContent() {
            const content = prompt('Nhập nội dung cần dán (hoặc để trống để dán clipboard hiện tại):', '');
            callTool('paste_content', {content: content || ''});
        }
        function pressEnter() {
            callTool('press_enter', {});
        }
        function findInDocument() {
            const searchText = prompt('Nhập nội dung tìm kiếm:', '');
            if (searchText && searchText.trim()) {
                callTool('find_in_document', {search_text: searchText.trim()});
            }
        }
        
        // AI ASSISTANT
        function saveGeminiModel() {
            const select = document.getElementById('gemini-model');
            if (!select) return;
            const model = select.value;
            localStorage.setItem('gemini_model', model);
            
            // Determine model name for display
            let modelName = 'Unknown';
            if (model.includes('flash-thinking')) modelName = 'Thinking 🧠';
            else if (model.includes('flash')) modelName = 'Flash ⚡';
            else if (model.includes('exp-1206')) modelName = 'Pro Exp 🚀';
            else if (model.includes('1.5-pro')) modelName = '1.5 Pro 💎';
            else if (model.includes('pro')) modelName = 'Pro 🚀';
            
            addLog(`✅ Đã lưu Gemini model: ${modelName}`, 'success');
        }
        
        function loadGeminiModel() {
            const saved = localStorage.getItem('gemini_model') || 'models/gemini-3-flash-preview';
            const select = document.getElementById('gemini-model');
            if (select) {
                // Check if the saved value exists in options
                const options = Array.from(select.options).map(o => o.value);
                if (options.includes(saved)) {
                    select.value = saved;
                } else {
                    // Default to first option if saved value is invalid
                    select.value = 'models/gemini-3-flash-preview';
                    localStorage.setItem('gemini_model', 'models/gemini-3-flash-preview');
                }
            }
        }
        
        function getGeminiModelName(model) {
            if (model.includes('flash-thinking')) return 'Thinking 🧠';
            if (model.includes('2.0-flash')) return '2.0 Flash ⚡';
            if (model.includes('1.5-flash')) return '1.5 Flash ⚡';
            if (model.includes('exp-1206')) return '2.0 Pro 🚀';
            if (model.includes('1.5-pro')) return '1.5 Pro 💎';
            return 'Gemini';
        }
        
        function askGemini() {
            const prompt = window.prompt('🤖 Hỏi Gemini AI + 📚 Knowledge Base\n(Gemini sẽ tự động tìm trong cơ sở dữ liệu của bạn):', '');
            if (prompt && prompt.trim()) {
                const model = localStorage.getItem('gemini_model') || 'models/gemini-3-flash-preview';
                const modelName = getGeminiModelName(model);
                addLog(`🤖 Đang hỏi Gemini ${modelName} + 📚 Knowledge Base...`, 'info');
                addLog(`   ❓ Câu hỏi: "${prompt}"`, 'info');
                
                // Sử dụng endpoint /api/tool/ask_gemini (có tích hợp KB tự động)
                fetch('/api/tool/ask_gemini', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({prompt: prompt.trim(), model: model})
                })
                .then(res => res.json())
                .then(result => {
                    if(result.success) {
                        const response = result.response || result.response_text || '';
                        const hasKB = result.knowledge_base_used ? ' 📚' : '';
                        addLog(`✅ Gemini${hasKB}: ${response.substring(0, 300)}...`, 'success');
                        if(result.knowledge_base_used) {
                            addLog(`   📚 Đã sử dụng thông tin từ Knowledge Base`, 'info');
                        }
                    } else {
                        addLog(`❌ Gemini error: ${result.error}`, 'error');
                    }
                })
                .catch(err => addLog(`❌ Error: ${err.message}`, 'error'));
            }
        }
        
        function askGPT4() {
            const prompt = window.prompt('Hỏi GPT-4 (TRẢ PHÍ - chất lượng cao nhất):', '');
            if (prompt && prompt.trim()) {
                addLog(`🧠 Hỏi GPT-4: "${prompt}"`, 'info');
                
                // Use generic /api/call_tool endpoint
                fetch('/api/call_tool', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({tool: 'ask_gpt4', args: {prompt: prompt.trim()}})
                })
                .then(res => res.json())
                .then(result => {
                    if(result.success) {
                        const usage = result.usage ? ` (Tokens: ${result.usage.total_tokens})` : '';
                        addLog(`✅ GPT-4: ${result.response_text.substring(0, 200)}...${usage}`, 'success');
                    } else {
                        addLog(`❌ GPT-4 error: ${result.error}`, 'error');
                    }
                })
                .catch(err => addLog(`❌ Error: ${err.message}`, 'error'));
            }
        }

        // API caller
        async function callAPI(endpoint, data) {
            try {
                addLog(`🔧 Calling ${endpoint}...`, 'info');
                const response = await fetch(endpoint, {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify(data)
                });
                const result = await response.json();
                addLog(`✅ ${JSON.stringify(result).substring(0, 100)}`, 'success');
                return result;
            } catch (error) {
                addLog(`❌ Error: ${error.message}`, 'error');
                return {success: false, error: error.message};
            }
        }
        
        async function callTool(name, params) {
            try {
                const paramsStr = JSON.stringify(params);
                const displayParams = paramsStr.length > 50 ? paramsStr.substring(0, 50) + '...' : paramsStr;
                addLog(`🔧 Tool: ${name}(${displayParams})`, 'info');
                
                // Gọi API endpoint tương ứng với tool
                const endpoint = `/api/tool/${name}`;
                const response = await fetch(endpoint, {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify(params)
                });
                const result = await response.json();
                
                // Hiển thị kết quả đầy đủ hơn
                let resultMsg = '';
                if (result.success) {
                    const msg = result.message || result.content || JSON.stringify(result).substring(0, 200);
                    resultMsg = `✅ ${name}: ${msg}`;
                } else {
                    resultMsg = `❌ ${name}: ${result.error || 'Unknown error'}`;
                }
                addLog(resultMsg, result.success ? 'success' : 'error');
                
                return result;
            } catch (error) {
                addLog(`❌ Tool "${name}" error: ${error.message}`, 'error');
                return {success: false, error: error.message};
            }
        }
        
        async function getResources() {
            try {
                // Sử dụng cache nếu còn hiệu lực
                const now = Date.now();
                if (resourceCache && (now - lastResourceFetch) < RESOURCE_CACHE_TIME) {
                    return;
                }
                
                const response = await fetch('/api/resources');
                const data = await response.json();
                if (data.success) {
                    const cpuPercent = data.data.cpu_percent;
                    document.getElementById('cpu').textContent = cpuPercent + '%';
                    document.getElementById('ram').textContent = data.data.memory_percent + '%';
                    document.getElementById('disk').textContent = data.data.disk_percent + '%';
                    
                    // Update RunCat animation speed based on CPU usage
                    updateRunCatSpeed(cpuPercent);
                    
                    // Cập nhật cache
                    resourceCache = data;
                    lastResourceFetch = now;
                } else {
                    addLog(`❌ Lỗi lấy tài nguyên: ${data.error}`, 'error');
                }
            } catch (error) {
                addLog(`❌ ${error.message}`, 'error');
            }
        }
        
        async function getQuotas() {
            try {
                const response = await fetch('/api/quotas');
                const data = await response.json();
                if (data.success) {
                    // Gemini quota
                    const geminiEl = document.getElementById('gemini-quota');
                    if (data.gemini && geminiEl) {
                        if (data.gemini.has_key) {
                            geminiEl.innerHTML = `✅ ${data.gemini.free_tier}<br><small style="color:#6b7280;">${data.gemini.daily_limit}</small>`;
                        } else {
                            geminiEl.innerHTML = `❌ <small style="color:#ef4444;">Chưa có API key</small>`;
                        }
                    }
                    
                    // Serper quota
                    const serperEl = document.getElementById('serper-quota');
                    if (data.serper && serperEl) {
                        if (data.serper.has_key) {
                            serperEl.innerHTML = `✅ ${data.serper.free_tier}`;
                        } else {
                            serperEl.innerHTML = `❌ <small style="color:#ef4444;">Chưa có API key</small>`;
                        }
                    }
                } else {
                    console.log('Error fetching quotas:', data.error);
                }
            } catch (error) {
                console.error('Failed to fetch quotas:', error);
            }
        }
        
        // Update RunCat animation speed based on CPU usage (like RunCat365)
        function updateRunCatSpeed(cpuPercent) {
            // Calculate frame duration: 100ms (very fast) to 800ms (very slow)
            // High CPU = fast running, Low CPU = slow walking
            const minSpeed = 100;  // Fast run (10 fps)
            const maxSpeed = 800;  // Slow walk (1.25 fps)
            
            // CPU 0% = 800ms, CPU 100% = 100ms
            runcatSpeed = maxSpeed - (cpuPercent / 100) * (maxSpeed - minSpeed);
        }
        
        // Debounce helper
        function debounce(func, wait) {
            let timeout;
            return function executedFunction(...args) {
                const later = () => {
                    clearTimeout(timeout);
                    func(...args);
                };
                clearTimeout(timeout);
                timeout = setTimeout(later, wait);
            };
        }
        
        async function calculate() {
            try {
                const expr = document.getElementById('calc-expr').value.trim();
                if (!expr) {
                    document.getElementById('calc-result').textContent = 'Vui lòng nhập biểu thức';
                    return;
                }
                const response = await fetch('/api/calculator', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({expression: expr})
                });
                const data = await response.json();
                document.getElementById('calc-result').textContent = data.success ? data.result : data.error;
            } catch (error) {
                document.getElementById('calc-result').textContent = 'Lỗi: ' + error.message;
            }
        }
        
        async function getCurrentTime() {
            try {
                const response = await fetch('/api/time');
                const data = await response.json();
                if (data.data) {
                    document.getElementById('time-result').textContent = data.data.datetime;
                }
            } catch (error) {
                document.getElementById('time-result').textContent = 'Lỗi: ' + error.message;
            }
        }
        
        // Modal functions
        function openSettingsModal() {
            document.getElementById('settingsModal').style.display = 'block';
            loadCurrentEndpoint();
            loadGeminiModel();
        }
        
        function closeSettingsModal() {
            document.getElementById('settingsModal').style.display = 'none';
        }
        
        // Click outside modal to close
        window.onclick = function(event) {
            const modal = document.getElementById('settingsModal');
            if (event.target === modal) {
                closeSettingsModal();
            }
        }
        
        async function loadCurrentEndpoint() {
            try {
                const response = await fetch('/api/endpoints');
                const data = await response.json();
                
                // 🔥 FIX: Định nghĩa activeDevice từ active_index
                const activeIndex = data.active_index || 0;
                const activeDevice = data.endpoints && data.endpoints[activeIndex] ? data.endpoints[activeIndex] : null;
                
                // Load all 3 device tokens into separate input fields
                if (data.endpoints && data.endpoints.length >= 3) {
                    const input1 = document.getElementById('endpoint-url-1');
                    const input2 = document.getElementById('endpoint-url-2');
                    const input3 = document.getElementById('endpoint-url-3');
                    
                    if (input1) input1.value = data.endpoints[0]?.token || '';
                    if (input2) input2.value = data.endpoints[1]?.token || '';
                    if (input3) input3.value = data.endpoints[2]?.token || '';
                }
                
                // Load Gemini API key (luôn set, kể cả empty)
                const geminiInput = document.getElementById('gemini-api-key');
                if (geminiInput) {
                    geminiInput.value = data.gemini_api_key || '';
                    if (data.gemini_api_key) {
                        updateGeminiKeyStatus('✓ API key đã cấu hình', '#10b981');
                    } else {
                        updateGeminiKeyStatus('', '');
                    }
                }
                
                // Load OpenAI API key (luôn set, kể cả empty)
                const openaiInput = document.getElementById('openai-api-key');
                if (openaiInput) {
                    openaiInput.value = data.openai_api_key || '';
                    if (data.openai_api_key) {
                        updateOpenAIKeyStatus('✓ API key đã cấu hình', '#10b981');
                    } else {
                        updateOpenAIKeyStatus('', '');
                    }
                }
                
                // Load Serper API key (Google Search) (luôn set, kể cả empty)
                const serperInput = document.getElementById('serper-api-key');
                if (serperInput) {
                    serperInput.value = data.serper_api_key || '';
                    if (data.serper_api_key) {
                        updateSerperKeyStatus('✓ Google Search sẵn sàng', '#10b981');
                    } else {
                        updateSerperKeyStatus('', '');
                    }
                }
                
                // Cập nhật thông tin hiện tại trong config section
                if (document.getElementById('current-device-name')) {
                    document.getElementById('current-device-name').textContent = activeDevice?.name || 'Chưa cấu hình';
                }
                if (document.getElementById('current-device-token')) {
                    const token = activeDevice?.token || 'Chưa có token';
                    document.getElementById('current-device-token').textContent = 
                        token.length > 50 ? token.substring(0, 50) + '...' : token;
                }
            } catch (error) {
                addLog('❌ Lỗi tải endpoint: ' + error.message, 'error');
            }
        }
        
        // Toggle API key visibility (show/hide password)
        function toggleApiKeyVisibility(inputId, button) {
            const input = document.getElementById(inputId);
            if (input.type === 'password') {
                input.type = 'text';
                button.innerHTML = '🙈'; // Hide icon (khỉ che mắt)
                button.title = 'Ẩn API key';
            } else {
                input.type = 'password';
                button.innerHTML = '🐵'; // Show icon (khỉ đang nhìn)
                button.title = 'Hiện API key';
            }
        }
        
        // Copy API key to clipboard
        async function copyApiKey(inputId, button) {
            const input = document.getElementById(inputId);
            const value = input.value.trim();
            
            if (!value) {
                button.innerHTML = '❌';
                setTimeout(() => { button.innerHTML = '📋'; }, 1000);
                return;
            }
            
            try {
                await navigator.clipboard.writeText(value);
                button.classList.add('copied');
                button.innerHTML = '✅';
                
                setTimeout(() => {
                    button.classList.remove('copied');
                    button.innerHTML = '📋';
                }, 1500);
            } catch (error) {
                // Fallback for older browsers
                input.select();
                document.execCommand('copy');
                button.innerHTML = '✅';
                setTimeout(() => { button.innerHTML = '📋'; }, 1500);
            }
        }
        
        // Auto-save Gemini API key
        let geminiSaveTimeout;
        async function autoSaveGeminiKey() {
            clearTimeout(geminiSaveTimeout);
            
            geminiSaveTimeout = setTimeout(async () => {
                const apiKey = document.getElementById('gemini-api-key').value.trim();
                
                // 🔥 FIX: Cho phép save empty string (khi user xóa key)
                try {
                    if (apiKey) {
                        updateGeminiKeyStatus('💾 Đang lưu...', '#f59e0b');
                    } else {
                        updateGeminiKeyStatus('💾 Xóa key...', '#f59e0b');
                    }
                    
                    const response = await fetch('/api/gemini-key', {
                        method: 'POST',
                        headers: {'Content-Type': 'application/json'},
                        body: JSON.stringify({api_key: apiKey})
                    });
                    
                    const result = await response.json();
                    
                    if (result.success) {
                        updateGeminiKeyStatus('✓ Đã lưu tự động', '#10b981');
                        setTimeout(() => updateGeminiKeyStatus('✓ API key đã cấu hình', '#10b981'), 2000);
                    } else {
                        updateGeminiKeyStatus('❌ Lỗi: ' + result.error, '#ef4444');
                    }
                } catch (error) {
                    updateGeminiKeyStatus('❌ Lỗi kết nối', '#ef4444');
                }
            }, 1000); // Auto-save sau 1 giây không gõ
        }
        
        function updateGeminiKeyStatus(message, color) {
            const statusEl = document.getElementById('gemini-key-status');
            if (statusEl) {
                statusEl.textContent = message;
                statusEl.style.color = color;
            }
        }
        
        // Auto-save OpenAI API key
        let openaiSaveTimeout;
        async function autoSaveOpenAIKey() {
            clearTimeout(openaiSaveTimeout);
            
            openaiSaveTimeout = setTimeout(async () => {
                const apiKey = document.getElementById('openai-api-key').value.trim();
                
                // 🔥 FIX: Cho phép save empty string (khi user xóa key)
                try {
                    if (apiKey) {
                        updateOpenAIKeyStatus('💾 Đang lưu...', '#f59e0b');
                    } else {
                        updateOpenAIKeyStatus('💾 Xóa key...', '#f59e0b');
                    }
                    
                    const response = await fetch('/api/openai-key', {
                        method: 'POST',
                        headers: {'Content-Type': 'application/json'},
                        body: JSON.stringify({api_key: apiKey})
                    });
                    
                    const result = await response.json();
                    
                    if (result.success) {
                        updateOpenAIKeyStatus('✓ Đã lưu tự động', '#10b981');
                        setTimeout(() => updateOpenAIKeyStatus('✓ API key đã cấu hình', '#10b981'), 2000);
                    } else {
                        updateOpenAIKeyStatus('❌ Lỗi: ' + result.error, '#ef4444');
                    }
                } catch (error) {
                    updateOpenAIKeyStatus('❌ Lỗi kết nối', '#ef4444');
                }
            }, 1000);
        }
        
        function updateOpenAIKeyStatus(message, color) {
            const statusEl = document.getElementById('openai-key-status');
            if (statusEl) {
                statusEl.textContent = message;
                statusEl.style.color = color;
            }
        }
        
        // Auto-save Serper API key (Google Search)
        let serperSaveTimeout;
        async function autoSaveSerperKey() {
            clearTimeout(serperSaveTimeout);
            
            serperSaveTimeout = setTimeout(async () => {
                const apiKey = document.getElementById('serper-api-key').value.trim();
                
                // 🔥 FIX: Cho phép save empty string (khi user xóa key)
                try {
                    if (apiKey) {
                        updateSerperKeyStatus('💾 Đang lưu...', '#f59e0b');
                    } else {
                        updateSerperKeyStatus('💾 Xóa key...', '#f59e0b');
                    }
                    
                    const response = await fetch('/api/serper-key', {
                        method: 'POST',
                        headers: {'Content-Type': 'application/json'},
                        body: JSON.stringify({api_key: apiKey})
                    });
                    
                    const result = await response.json();
                    
                    if (result.success) {
                        updateSerperKeyStatus('✓ Đã lưu - Google Search sẵn sàng!', '#10b981');
                        setTimeout(() => updateSerperKeyStatus('✓ API key đã cấu hình', '#10b981'), 2000);
                    } else {
                        updateSerperKeyStatus('❌ Lỗi: ' + result.error, '#ef4444');
                    }
                } catch (error) {
                    updateSerperKeyStatus('❌ Lỗi kết nối', '#ef4444');
                }
            }, 1000);
        }
        
        function updateSerperKeyStatus(message, color) {
            const statusEl = document.getElementById('serper-key-status');
            if (statusEl) {
                statusEl.textContent = message;
                statusEl.style.color = color;
            }
        }
        
        async function saveEndpoint() {
            try {
                addLog('⏳ Đang lưu endpoints...', 'info');
                
                // Lấy token từ cả 3 input fields
                const token1 = document.getElementById('endpoint-url-1').value.trim();
                const token2 = document.getElementById('endpoint-url-2').value.trim();
                const token3 = document.getElementById('endpoint-url-3').value.trim();
                
                if (!token1 && !token2 && !token3) {
                    addLog('❌ Vui lòng nhập ít nhất 1 JWT token!', 'error');
                    return;
                }
                
                // Helper function to extract token from URL or return as-is
                function extractToken(input) {
                    if (!input) return '';
                    
                    // Nếu user nhập URL đầy đủ, extract token từ URL
                    if (input.startsWith('wss://') || input.startsWith('http')) {
                        try {
                            const url = new URL(input);
                            const tokenParam = url.searchParams.get('token');
                            if (tokenParam) {
                                return tokenParam;
                            }
                        } catch (e) {
                            return input; // Return as-is if parse fails
                        }
                    }
                    return input;
                }
                
                const cleanToken1 = extractToken(token1);
                const cleanToken2 = extractToken(token2);
                const cleanToken3 = extractToken(token3);
                
                // Lấy danh sách thiết bị hiện tại
                const response = await fetch('/api/endpoints');
                const data = await response.json();
                
                // Update all 3 devices
                const devices = data.endpoints.map((device, index) => {
                    let token = '';
                    if (index === 0) token = cleanToken1;
                    else if (index === 1) token = cleanToken2;
                    else if (index === 2) token = cleanToken3;
                    
                    return {
                        name: device.name || `Thiết bị ${index + 1}`,
                        token: token,
                        enabled: token.length > 0  // Auto-enable if has token
                    };
                });
                
                // Lưu cấu hình
                const saveResponse = await fetch('/api/endpoints/save', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({devices: devices})
                });
                
                const saveData = await saveResponse.json();
                
                if (saveData.success) {
                    addLog('✅ Đã lưu endpoints thành công!', 'success');
                    
                    // Show which devices were updated
                    let updatedCount = 0;
                    if (cleanToken1) { addLog('  📱 Thiết bị 1: Đã cập nhật', 'success'); updatedCount++; }
                    if (cleanToken2) { addLog('  📱 Thiết bị 2: Đã cập nhật', 'success'); updatedCount++; }
                    if (cleanToken3) { addLog('  📱 Thiết bị 3: Đã cập nhật', 'success'); updatedCount++; }
                    
                    addLog(`📡 ${updatedCount} thiết bị sẽ tự động kết nối...`, 'info');
                    
                    closeSettingsModal();
                    
                    // Reload trang sau 1 giây
                    setTimeout(() => {
                        location.reload();
                    }, 1000);
                } else {
                    addLog('❌ Lỗi: ' + saveData.error, 'error');
                }
            } catch (error) {
                addLog('❌ Lỗi lưu endpoint: ' + error.message, 'error');
            }
        }
        
        function copyFullUrl() {
            // Get tokens from all 3 fields
            const token1 = document.getElementById('endpoint-url-1').value.trim();
            const token2 = document.getElementById('endpoint-url-2').value.trim();
            const token3 = document.getElementById('endpoint-url-3').value.trim();
            
            if (!token1 && !token2 && !token3) {
                addLog('❌ Không có token nào để copy!', 'error');
                return;
            }
            
            let copyText = '';
            
            // Helper function to extract token and create URL
            function createFullUrl(input, deviceNum) {
                if (!input) return null;
                
                let token = input;
                
                // Nếu user đã nhập URL đầy đủ, extract token
                if (input.startsWith('wss://') || input.startsWith('http')) {
                    try {
                        const url = new URL(input);
                        const tokenParam = url.searchParams.get('token');
                        if (tokenParam) {
                            token = tokenParam;
                        }
                    } catch (e) {
                        return null;
                    }
                }
                
                return `Thiết bị ${deviceNum}: wss://api.xiaozhi.me/mcp/?token=${token}`;
            }
            
            // Create URLs for all devices with tokens
            const urls = [];
            if (token1) urls.push(createFullUrl(token1, 1));
            if (token2) urls.push(createFullUrl(token2, 2));
            if (token3) urls.push(createFullUrl(token3, 3));
            
            copyText = urls.filter(u => u).join('\n\n');
            
            // Copy vào clipboard
            navigator.clipboard.writeText(copyText).then(() => {
                addLog(`✅ Đã copy ${urls.length} URL vào clipboard!`, 'success');
            }).catch(err => {
                addLog('❌ Lỗi copy: ' + err.message, 'error');
            });
        }
        
        // Load and display all 3 devices
        async function loadDevices() {
            try {
                const response = await fetch('/api/endpoints');
                const data = await response.json();
                
                // Update device status display for all 3 devices
                data.endpoints.forEach((device, index) => {
                    const deviceName = device?.name || `Thiết bị ${index + 1}`;
                    const hasToken = device?.token && device.token.length > 0;
                    const isEnabled = device?.enabled || false;
                    
                    addLog(`📱 ${deviceName}: ${hasToken ? '✅ Connected' : '❌ No token'} ${isEnabled ? '(Enabled)' : '(Disabled)'}`, 
                           hasToken && isEnabled ? 'success' : 'info');
                });
            } catch (error) {
                addLog('❌ Lỗi tải danh sách thiết bị: ' + error.message, 'error');
            }
        }

        function addLog(message, type = 'info') {
            const log = document.getElementById('log');
            if (!log) return;
            const entry = document.createElement('div');
            entry.className = `log-entry log-${type}`;
            const time = new Date().toLocaleTimeString();
            entry.innerHTML = `<span class="log-time">${time}</span> ${message}`;
            log.insertBefore(entry, log.firstChild);
            
            // Giới hạn 50 logs thay vì 100 để giảm DOM size
            if (log.children.length > 50) {
                // Xóa nhiều logs cùng lúc để tránh reflow nhiều lần
                while (log.children.length > 50) {
                    log.removeChild(log.lastChild);
                }
            }
        }
        
        // WebSocket với reconnect optimization
        let wsReconnectAttempts = 0;
        const MAX_RECONNECT_DELAY = 30000; // Max 30s
        
        function connectWS() {
            ws = new WebSocket(`ws://${window.location.host}/ws`);
            ws.onopen = () => {
                addLog('✅ WebSocket connected', 'success');
                wsReconnectAttempts = 0; // Reset counter khi connect thành công
            };
            ws.onmessage = (event) => {
                const data = JSON.parse(event.data);
                if (data.type === 'xiaozhi_status') {
                    const badge = document.getElementById('xiaozhi-status');
                    const text = document.getElementById('xiaozhi-text');
                    if (data.connected) {
                        badge.className = 'status-badge online';
                        text.textContent = 'Connected';
                    } else {
                        badge.className = 'status-badge offline';
                        text.textContent = 'Disconnected';
                    }
                } else if (data.type === 'xiaozhi_activity') {
                    if (data.method !== 'ping') {
                        addLog(`📡 Xiaozhi: ${data.method}`, 'info');
                    }
                }
            };
            ws.onclose = () => {
                addLog('❌ WebSocket disconnected', 'error');
                // Exponential backoff cho reconnect
                wsReconnectAttempts++;
                const delay = Math.min(1000 * Math.pow(2, wsReconnectAttempts), MAX_RECONNECT_DELAY);
                setTimeout(connectWS, delay);
            };
        }
        
        // Caching và optimization
        let resourceCache = null;
        let lastResourceFetch = 0;
        const RESOURCE_CACHE_TIME = 3000; // Cache 3 giây
        
        // Playlist list functions (sử dụng API backend thay vì localStorage)
        async function getPlaylists() {
            try {
                const response = await fetch('/api/youtube_playlists');
                const data = await response.json();
                return data.success ? data.playlists : [];
            } catch (e) {
                console.error('Failed to load playlists from API', e);
                return [];
            }
        }

        async function renderPlaylists() {
            const list = await getPlaylists();
            const container = document.getElementById('playlist-list');
            if (!container) return;
            container.innerHTML = '';

            if (list.length === 0) {
                container.innerHTML = '<div style="color:#666;padding:12px;">Chưa có playlist nào. Nhấn "＋ Thêm Playlist" để thêm.</div>';
                return;
            }

            list.forEach((item, idx) => {
                const row = document.createElement('div');
                row.style.display = 'flex';
                row.style.alignItems = 'center';
                row.style.justifyContent = 'space-between';
                row.style.padding = '8px';
                row.style.borderBottom = '1px solid #eee';

                const left = document.createElement('div');
                left.style.display = 'flex';
                left.style.flexDirection = 'column';
                left.style.gap = '4px';

                const name = document.createElement('div');
                name.textContent = item.name;
                name.style.fontWeight = '700';
                name.style.color = '#333';

                const url = document.createElement('div');
                url.textContent = item.url;
                url.style.fontSize = '0.85em';
                url.style.color = '#666';

                left.appendChild(name);
                left.appendChild(url);

                const actions = document.createElement('div');
                actions.style.display = 'flex';
                actions.style.gap = '8px';

                const openBtn = document.createElement('button');
                openBtn.textContent = '▶';
                openBtn.title = 'Mở playlist';
                openBtn.style.padding = '6px 10px';
                openBtn.style.borderRadius = '6px';
                openBtn.style.border = 'none';
                openBtn.style.background = '#10b981';
                openBtn.style.color = 'white';
                openBtn.style.cursor = 'pointer';
                openBtn.onclick = () => openPlaylistByName(item.name);

                const delBtn = document.createElement('button');
                delBtn.textContent = '🗑';
                delBtn.title = 'Xóa playlist';
                delBtn.style.padding = '6px 10px';
                delBtn.style.borderRadius = '6px';
                delBtn.style.border = 'none';
                delBtn.style.background = '#ef4444';
                delBtn.style.color = 'white';
                delBtn.style.cursor = 'pointer';
                delBtn.onclick = () => { if (confirm('Xóa playlist "' + item.name + '"?')) { removePlaylistByName(item.name); } };

                actions.appendChild(openBtn);
                actions.appendChild(delBtn);

                row.appendChild(left);
                row.appendChild(actions);

                container.appendChild(row);
            });
        }

        function promptAddPlaylist() {
            const name = prompt('Nhập tên playlist (ví dụ: "Nhạc chill"):');
            if (!name) return;
            const url = prompt('Dán link playlist YouTube (hoặc video trong playlist):');
            if (!url) return;
            addPlaylist(name.trim(), url.trim());
        }

        async function addPlaylist(name, url) {
            if (!name || !url) {
                addLog('❌ Tên và URL không được để trống', 'error');
                return;
            }
            try {
                const response = await fetch('/api/youtube_playlists/add', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({name, url})
                });
                const data = await response.json();
                if (data.success) {
                    await renderPlaylists();
                    addLog('✅ Đã thêm playlist: ' + name, 'success');
                } else {
                    addLog('❌ ' + (data.error || 'Không thể thêm playlist'), 'error');
                }
            } catch (e) {
                console.error('Failed to add playlist', e);
                addLog('❌ Lỗi khi thêm playlist', 'error');
            }
        }

        async function removePlaylistByName(name) {
            try {
                const response = await fetch('/api/youtube_playlists/remove', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({name})
                });
                const data = await response.json();
                if (data.success) {
                    await renderPlaylists();
                    addLog('🗑 Đã xóa playlist: ' + name, 'info');
                } else {
                    addLog('❌ ' + (data.error || 'Không thể xóa playlist'), 'error');
                }
            } catch (e) {
                console.error('Failed to remove playlist', e);
                addLog('❌ Lỗi khi xóa playlist', 'error');
            }
        }

        async function openPlaylistByName(name) {
            const list = await getPlaylists();
            const item = list.find(p => p.name === name);
            if (item) {
                window.open(item.url, '_blank');
                addLog('▶ Mở playlist: ' + name, 'info');
            }
        }

        // Expose function for voice/AI integration: open by keyword search (fuzzy matching)
        async function triggerPlayByName(keyword) {
            if (!keyword || keyword.trim() === '') return false;
            
            keyword = keyword.trim().toLowerCase();
            const list = await getPlaylists();
            
            if (list.length === 0) {
                addLog('⚠ Danh sách playlist trống. Hãy thêm playlist trước!', 'error');
                return false;
            }
            
            // Bước 1: Tìm chính xác (exact match)
            let found = list.find(item => item.name.toLowerCase() === keyword);
            
            // Bước 2: Tìm bắt đầu bằng từ khóa (starts with)
            if (!found) {
                found = list.find(item => item.name.toLowerCase().startsWith(keyword));
            }
            
            // Bước 3: Tìm chứa từ khóa (contains)
            if (!found) {
                found = list.find(item => item.name.toLowerCase().includes(keyword));
            }
            
            // Bước 4: Tìm theo từng từ trong tên playlist
            if (!found) {
                found = list.find(item => {
                    const words = item.name.toLowerCase().split(/\\s+/);
                    return words.some(word => word.includes(keyword) || keyword.includes(word));
                });
            }
            
            if (found) {
                window.open(found.url, '_blank');
                addLog('🔊 Phát playlist: "' + found.name + '" (từ khóa: "' + keyword + '")', 'success');
                return true;
            } else {
                // Hiển thị gợi ý các playlist có sẵn
                const suggestions = list.map(item => item.name).slice(0, 5).join(', ');
                addLog('⚠ Không tìm thấy playlist với từ khóa: "' + keyword + '"', 'error');
                addLog('💡 Gợi ý: ' + suggestions, 'info');
                return false;
            }
        }
        
        // Hàm mở playlist nhanh (alias) - dễ nhớ hơn cho voice command
        function moPlaylist(keyword) {
            return triggerPlayByName(keyword);
        }
        
        function danhSachNhac(keyword) {
            return triggerPlayByName(keyword);
        }

        // Initialize playlist list on load
        function initPlaylists() {
            renderPlaylists();
        }
        
        // ============================================================
        // KNOWLEDGE BASE FUNCTIONS
        // ============================================================
        
        async function loadKnowledgeBase() {
            try {
                const response = await fetch('/api/knowledge/status');
                const data = await response.json();
                
                if (data.success) {
                    document.getElementById('knowledge-folder-path').value = data.folder_path || '';
                    document.getElementById('kb-total-files').textContent = data.total_files || 0;
                    document.getElementById('kb-indexed-files').textContent = data.indexed_files || 0;
                    document.getElementById('kb-total-size').textContent = formatFileSize(data.total_size || 0);
                    document.getElementById('kb-last-update').textContent = data.last_update || '--';
                    
                    if (data.files && data.files.length > 0) {
                        renderKnowledgeFiles(data.files);
                    }
                }
            } catch (error) {
                console.error('Error loading knowledge base:', error);
            }
        }
        
        function formatFileSize(bytes) {
            if (bytes === 0) return '0 B';
            const k = 1024;
            const sizes = ['B', 'KB', 'MB', 'GB'];
            const i = Math.floor(Math.log(bytes) / Math.log(k));
            return parseFloat((bytes / Math.pow(k, i)).toFixed(1)) + ' ' + sizes[i];
        }
        
        function renderKnowledgeFiles(files) {
            const container = document.getElementById('knowledge-file-list');
            if (!files || files.length === 0) {
                container.innerHTML = '<p style="color: #666; text-align: center; padding: 40px;">📂 Không tìm thấy file nào.</p>';
                return;
            }
            
            const fileIcons = {
                'pdf': '📕',
                'txt': '📄',
                'docx': '📘',
                'doc': '📘',
                'md': '📝',
                'json': '📋',
                'csv': '📊',
                'xlsx': '📗',
                'xls': '📗'
            };
            
            let html = '<div style="display: flex; flex-direction: column; gap: 8px;">';
            files.forEach((file, index) => {
                const ext = file.name.split('.').pop().toLowerCase();
                const icon = fileIcons[ext] || '📄';
                const indexed = file.indexed ? '✅' : '⏳';
                const escapedPath = btoa(unescape(encodeURIComponent(file.path))); // Base64 encode để tránh lỗi escape
                
                html += `
                    <div style="display: flex; align-items: center; padding: 12px; background: white; border-radius: 8px; border: 1px solid #e5e7eb; gap: 12px;">
                        <span style="font-size: 1.5em;">${icon}</span>
                        <div style="flex: 1;">
                            <div style="font-weight: 600; color: #333;">${file.name}</div>
                            <div style="font-size: 0.85em; color: #666;">${formatFileSize(file.size)} • ${file.modified || ''}</div>
                        </div>
                        <span title="${file.indexed ? 'Đã index' : 'Chưa index'}">${indexed}</span>
                        <button onclick="indexSingleFileB64('${escapedPath}')" 
                                style="padding: 6px 12px; background: #667eea; color: white; border: none; border-radius: 6px; cursor: pointer; font-size: 0.85em;">
                            Index
                        </button>
                    </div>
                `;
            });
            html += '</div>';
            container.innerHTML = html;
        }
        
        async function saveKnowledgeFolder() {
            const folderPath = document.getElementById('knowledge-folder-path').value.trim();
            console.log('[Knowledge] saveKnowledgeFolder called, path:', folderPath);
            if (!folderPath) {
                addLog('❌ Vui lòng nhập đường dẫn thư mục', 'error');
                alert('Vui lòng nhập đường dẫn thư mục!');
                return;
            }
            
            try {
                addLog('💾 Đang lưu cấu hình thư mục...', 'info');
                console.log('[Knowledge] Calling API /api/knowledge/set_folder');
                const response = await fetch('/api/knowledge/set_folder', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ folder_path: folderPath })
                });
                console.log('[Knowledge] Response status:', response.status);
                const data = await response.json();
                console.log('[Knowledge] Response data:', data);
                
                if (data.success) {
                    addLog('✅ ' + data.message, 'success');
                    alert('✅ ' + data.message);
                    loadKnowledgeBase();
                } else {
                    addLog('❌ ' + (data.error || 'Lỗi không xác định'), 'error');
                    alert('❌ ' + (data.error || 'Lỗi không xác định'));
                }
            } catch (error) {
                console.error('[Knowledge] Error:', error);
                addLog('❌ Lỗi: ' + error.message, 'error');
                alert('❌ Lỗi: ' + error.message);
            }
        }
        
        async function scanKnowledgeFolder() {
            const folderPath = document.getElementById('knowledge-folder-path').value.trim();
            console.log('[Knowledge] scanKnowledgeFolder called, path:', folderPath);
            if (!folderPath) {
                addLog('❌ Vui lòng nhập đường dẫn thư mục trước', 'error');
                alert('Vui lòng nhập đường dẫn thư mục trước!');
                return;
            }
            
            try {
                addLog('🔍 Đang quét thư mục...', 'info');
                console.log('[Knowledge] Calling API /api/knowledge/scan');
                const response = await fetch('/api/knowledge/scan', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ folder_path: folderPath })
                });
                console.log('[Knowledge] Response status:', response.status);
                const data = await response.json();
                console.log('[Knowledge] Response data:', data);
                
                if (data.success) {
                    addLog('✅ Tìm thấy ' + data.total_files + ' files', 'success');
                    document.getElementById('kb-total-files').textContent = data.total_files;
                    document.getElementById('kb-total-size').textContent = formatFileSize(data.total_size);
                    renderKnowledgeFiles(data.files);
                } else {
                    addLog('❌ ' + (data.error || 'Lỗi không xác định'), 'error');
                    alert('❌ ' + (data.error || 'Lỗi không xác định'));
                }
            } catch (error) {
                console.error('[Knowledge] Scan error:', error);
                addLog('❌ Lỗi: ' + error.message, 'error');
                alert('❌ Lỗi: ' + error.message);
            }
        }
        
        async function indexAllFiles() {
            try {
                addLog('🔄 Đang index tất cả files...', 'info');
                const response = await fetch('/api/knowledge/index_all', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' }
                });
                const data = await response.json();
                
                if (data.success) {
                    addLog('✅ ' + data.message, 'success');
                    document.getElementById('kb-indexed-files').textContent = data.indexed_count;
                    document.getElementById('kb-last-update').textContent = data.last_update || 'Vừa xong';
                    loadKnowledgeBase();
                } else {
                    addLog('❌ ' + (data.error || 'Lỗi không xác định'), 'error');
                }
            } catch (error) {
                addLog('❌ Lỗi: ' + error.message, 'error');
            }
        }
        
        // Decode Base64 path và gọi indexSingleFile
        async function indexSingleFileB64(base64Path) {
            try {
                const filePath = decodeURIComponent(escape(atob(base64Path)));
                await indexSingleFile(filePath);
            } catch (error) {
                addLog('❌ Lỗi decode path: ' + error.message, 'error');
            }
        }
        
        async function indexSingleFile(filePath) {
            try {
                addLog('🔄 Đang index file: ' + filePath.split(/[\\/]/).pop(), 'info');
                const response = await fetch('/api/knowledge/index_file', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ file_path: filePath })
                });
                const data = await response.json();
                
                if (data.success) {
                    addLog('✅ ' + data.message, 'success');
                    loadKnowledgeBase();
                } else {
                    addLog('❌ ' + (data.error || 'Lỗi không xác định'), 'error');
                }
            } catch (error) {
                addLog('❌ Lỗi: ' + error.message, 'error');
            }
        }
        
        async function clearKnowledgeBase() {
            if (!confirm('Bạn có chắc muốn xóa toàn bộ index? Dữ liệu gốc không bị ảnh hưởng.')) {
                return;
            }
            
            try {
                addLog('🗑️ Đang xóa index...', 'info');
                const response = await fetch('/api/knowledge/clear', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' }
                });
                const data = await response.json();
                
                if (data.success) {
                    addLog('✅ ' + data.message, 'success');
                    document.getElementById('kb-indexed-files').textContent = '0';
                    loadKnowledgeBase();
                } else {
                    addLog('❌ ' + (data.error || 'Lỗi không xác định'), 'error');
                }
            } catch (error) {
                addLog('❌ Lỗi: ' + error.message, 'error');
            }
        }
        
        // ============================================================
        // CONVERSATION HISTORY FUNCTIONS (WeChat Style)
        // ============================================================
        
        async function loadConversationHistory() {
            try {
                addLog('📚 Đang tải lịch sử hội thoại từ server...', 'info');
                const response = await fetch('/api/conversation/history');
                const data = await response.json();
                
                if (data.success) {
                    const messages = data.messages || [];
                    const totalMessages = data.total_messages || 0;
                    
                    displayConversationHistory(messages);
                    document.getElementById('total-messages').textContent = totalMessages;
                    
                    if (messages.length > 0) {
                        const lastMsg = messages[messages.length - 1];
                        const updateTime = lastMsg.timestamp || 'Không rõ';
                        document.getElementById('last-update').textContent = 'Cập nhật: ' + updateTime;
                        addLog('✅ Đã tải thành công ' + totalMessages + ' tin nhắn (cập nhật lần cuối: ' + updateTime + ')', 'success');
                    } else {
                        document.getElementById('last-update').textContent = 'Chưa có tin nhắn';
                        addLog('✅ Lịch sử hội thoại trống', 'success');
                    }
                } else {
                    addLog('❌ Lỗi tải lịch sử hội thoại: ' + (data.error || 'Unknown error'), 'error');
                    displayConversationHistory([]);
                }
            } catch (e) {
                console.error('Failed to load conversation history', e);
                addLog('❌ Không thể kết nối đến server để tải lịch sử', 'error');
                displayConversationHistory([]);
            }
        }
        
        function displayConversationHistory(messages) {
            const container = document.getElementById('chat-container');
            container.innerHTML = '';
            
            if (!messages || messages.length === 0) {
                container.innerHTML = '<div style="text-align:center; color:#999; padding:40px; font-size:1.1em;">Chưa có tin nhắn nào 💬</div>';
                return;
            }
            
            messages.forEach(msg => {
                const messageDiv = document.createElement('div');
                messageDiv.className = 'chat-message ' + msg.role;
                
                // Avatar
                const avatar = document.createElement('div');
                avatar.className = 'chat-avatar ' + msg.role;
                const roleIcons = {
                    user: '👤',
                    assistant: '🤖',
                    system: '⚙️',
                    tool: '🔧'
                };
                avatar.textContent = roleIcons[msg.role] || '💬';
                
                // Bubble
                const bubble = document.createElement('div');
                bubble.className = 'chat-bubble';
                
                // Content
                const content = document.createElement('div');
                content.className = 'chat-content';
                content.textContent = msg.content;
                bubble.appendChild(content);
                
                // Metadata
                if (msg.metadata && Object.keys(msg.metadata).length > 0) {
                    const metadata = document.createElement('div');
                    metadata.className = 'chat-metadata';
                    
                    // Show relevant metadata
                    if (msg.metadata.source) {
                        const sourceTag = document.createElement('span');
                        sourceTag.className = 'chat-metadata-item';
                        const sourceIcons = {
                            mcp: '🔌 MCP',
                            web_ui: '🌐 Web UI',
                            websocket: '📡 WebSocket'
                        };
                        sourceTag.textContent = sourceIcons[msg.metadata.source] || msg.metadata.source;
                        metadata.appendChild(sourceTag);
                    }
                    
                    if (msg.metadata.method) {
                        const methodTag = document.createElement('span');
                        methodTag.className = 'chat-metadata-item';
                        methodTag.textContent = '📋 ' + msg.metadata.method;
                        metadata.appendChild(methodTag);
                    }
                    
                    if (msg.metadata.model) {
                        const modelTag = document.createElement('span');
                        modelTag.className = 'chat-metadata-item';
                        modelTag.textContent = '🧠 ' + msg.metadata.model;
                        metadata.appendChild(modelTag);
                    }
                    
                    if (msg.metadata.success !== undefined) {
                        const statusTag = document.createElement('span');
                        statusTag.className = 'chat-metadata-item';
                        statusTag.textContent = msg.metadata.success ? '✅ Success' : '❌ Failed';
                        metadata.appendChild(statusTag);
                    }
                    
                    bubble.appendChild(metadata);
                }
                
                // Timestamp
                const timestamp = document.createElement('div');
                timestamp.className = 'chat-timestamp';
                timestamp.textContent = msg.timestamp;
                bubble.appendChild(timestamp);
                
                messageDiv.appendChild(avatar);
                messageDiv.appendChild(bubble);
                container.appendChild(messageDiv);
            });
            
            // Auto scroll to bottom
            container.scrollTop = container.scrollHeight;
        }
        
        async function exportConversation() {
            try {
                addLog('💾 Đang xuất lịch sử...', 'info');
                const response = await fetch('/api/conversation/export', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({})
                });
                const data = await response.json();
                
                if (data.success) {
                    addLog('✅ Đã xuất file: ' + data.path, 'success');
                    alert('✅ Đã xuất lịch sử hội thoại!\\n\\nĐường dẫn: ' + data.path + '\\n\\nTổng: ' + data.message);
                } else {
                    addLog('❌ Lỗi xuất file: ' + (data.error || 'Unknown'), 'error');
                }
            } catch (e) {
                console.error('Failed to export conversation', e);
                addLog('❌ Không thể xuất file', 'error');
            }
        }
        
        async function clearConversationHistory() {
            if (!confirm('⚠️ Bạn có chắc muốn XÓA TẤT CẢ lịch sử hội thoại?\\n\\nHành động này KHÔNG THỂ HOÀN TÁC!')) {
                return;
            }
            
            try {
                addLog('🗑️ Đang xóa lịch sử...', 'info');
                const response = await fetch('/api/conversation/clear', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'}
                });
                const data = await response.json();
                
                if (data.success) {
                    document.getElementById('chat-container').innerHTML = '<div style="text-align:center; color:#999; padding:40px; font-size:1.1em;">Chưa có tin nhắn nào 💬</div>';
                    document.getElementById('total-messages').textContent = '0';
                    document.getElementById('last-update').textContent = 'Chưa có dữ liệu';
                    addLog('✅ Đã xóa toàn bộ lịch sử', 'success');
                } else {
                    addLog('❌ Lỗi xóa lịch sử: ' + (data.error || 'Unknown'), 'error');
                }
            } catch (e) {
                console.error('Failed to clear conversation', e);
                addLog('❌ Không thể xóa lịch sử', 'error');
            }
        }
        
        // ===== MUSIC PLAYER FUNCTIONS =====
        let currentPlaylist = [];
        let allMusicFiles = []; // Store all files for filtering
        let currentTrackIndex = -1;
        let isPlaying = false;
        let isShuffleOn = false;
        let repeatMode = 0; // 0: off, 1: repeat all, 2: repeat one
        let currentMusicSource = 'library'; // 'library' or 'user'
        let vlcStatusInterval = null; // VLC status polling interval
        
        // ===== VLC STATUS POLLING - Real-time sync with python-vlc =====
        async function pollVlcStatus() {
            try {
                const response = await fetch('/api/vlc_status');
                const status = await response.json();
                
                if (status.state && status.state !== 'not_initialized') {
                    // Update play state
                    isPlaying = status.is_playing;
                    document.getElementById('play-btn').textContent = isPlaying ? '⏸️' : '▶️';
                    
                    // Update progress slider (only if not dragging)
                    if (status.position !== undefined && !isDraggingProgress) {
                        const percent = (status.position * 100).toFixed(1);
                        const slider = document.getElementById('progress-slider');
                        if (slider) {
                            slider.value = percent;
                            slider.style.background = `linear-gradient(to right, #667eea 0%, #667eea ${percent}%, #374151 ${percent}%, #374151 100%)`;
                        }
                    }
                    
                    // Update time display
                    if (status.current_time_formatted) {
                        document.getElementById('current-time').textContent = status.current_time_formatted;
                    }
                    if (status.duration_formatted) {
                        document.getElementById('total-time').textContent = status.duration_formatted;
                    }
                    
                    // Update volume (sync from VLC)
                    if (status.volume !== undefined) {
                        const slider = document.getElementById('volume-slider');
                        if (document.activeElement !== slider) { // Don't update while user is dragging
                            slider.value = status.volume;
                            document.getElementById('volume-value').textContent = status.volume + '%';
                            slider.style.background = `linear-gradient(to right, #667eea 0%, #667eea ${status.volume}%, #374151 ${status.volume}%, #374151 100%)`;
                        }
                    }
                    
                    // Update current track name
                    if (status.current_track) {
                        document.getElementById('current-track').textContent = '🎵 ' + status.current_track;
                        document.getElementById('track-info').textContent = 
                            `${status.playlist_index + 1}/${status.playlist_count} bài • VLC Player`;
                    }
                    
                    // Sync shuffle/repeat state from VLC
                    if (status.shuffle !== undefined) {
                        isShuffleOn = status.shuffle;
                        const shuffleBtn = document.getElementById('shuffle-btn');
                        if (shuffleBtn) {
                            shuffleBtn.style.opacity = isShuffleOn ? '1' : '0.6';
                            shuffleBtn.style.transform = isShuffleOn ? 'scale(1.1)' : 'scale(1)';
                        }
                    }
                    if (status.repeat_mode !== undefined) {
                        repeatMode = status.repeat_mode;
                        const repeatBtn = document.getElementById('repeat-btn');
                        if (repeatBtn) {
                            repeatBtn.textContent = repeatMode === 2 ? '🔂' : '🔁';
                            repeatBtn.style.opacity = repeatMode > 0 ? '1' : '0.6';
                        }
                    }
                }
            } catch (e) {
                // Silent fail - VLC may not be playing
            }
        }
        
        function startVlcPolling() {
            if (vlcStatusInterval) clearInterval(vlcStatusInterval);
            vlcStatusInterval = setInterval(pollVlcStatus, 1000); // Poll every 1 second
        }
        
        function stopVlcPolling() {
            if (vlcStatusInterval) {
                clearInterval(vlcStatusInterval);
                vlcStatusInterval = null;
            }
        }
        
        // Click on progress bar to seek
        async function seekToPosition(event) {
            const progressBar = event.currentTarget;
            const rect = progressBar.getBoundingClientRect();
            const position = (event.clientX - rect.left) / rect.width;
            
            try {
                await fetch('/api/vlc_seek', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({position: position})
                });
            } catch (e) {
                console.error('Seek failed', e);
            }
        }
        
        // Music Source Selector Functions
        function setMusicSource(source) {
            currentMusicSource = source;
            localStorage.setItem('musicSource', source);
            
            // Update button styles
            const libraryBtn = document.getElementById('source-library-btn');
            const userBtn = document.getElementById('source-user-btn');
            
            if (source === 'library') {
                libraryBtn.style.background = '#667eea';
                libraryBtn.style.color = 'white';
                userBtn.style.background = 'transparent';
                userBtn.style.color = '#667eea';
                document.getElementById('source-path-display').textContent = 'music_library/';
            } else {
                libraryBtn.style.background = 'transparent';
                libraryBtn.style.color = '#667eea';
                userBtn.style.background = '#667eea';
                userBtn.style.color = 'white';
                const userPath = localStorage.getItem('musicFolderPath') || 'Chưa cấu hình';
                document.getElementById('source-path-display').textContent = userPath;
            }
            
            // Reload music library from new source
            loadMusicLibrary();
            addLog(`🎯 Đã chuyển nguồn phát: ${source === 'library' ? 'Music Library' : 'Thư mục cá nhân'}`, 'success');
        }
        
        function loadMusicSourcePreference() {
            const saved = localStorage.getItem('musicSource') || 'library';
            setMusicSource(saved);
        }
        
        // Search/Filter Music Library
        function filterMusicLibrary(query) {
            if (!query || query.trim() === '') {
                renderMusicLibrary(allMusicFiles);
                return;
            }
            
            const lowerQuery = query.toLowerCase();
            const filtered = allMusicFiles.filter(file => 
                file.filename.toLowerCase().includes(lowerQuery) ||
                (file.path && file.path.toLowerCase().includes(lowerQuery))
            );
            renderMusicLibrary(filtered);
        }
        
        async function loadMusicLibrary() {
            try {
                // Determine which source to load from
                // IMPORTANT: auto_play=false để không tự phát khi load danh sách
                const args = currentMusicSource === 'user' 
                    ? { folder: localStorage.getItem('musicFolderPath') || '', auto_play: false }
                    : { auto_play: false };
                
                const response = await fetch('/api/call_tool', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({tool: 'list_music', args: args})
                });
                const data = await response.json();
                
                if (data.success && data.files) {
                    allMusicFiles = data.files;
                    currentPlaylist = data.files;
                    renderMusicLibrary(data.files);
                } else {
                    document.getElementById('music-library').innerHTML = '<p style="text-align:center; color:#999; padding:40px;">❌ Không tìm thấy nhạc trong thư viện</p>';
                }
            } catch (e) {
                console.error('Failed to load music library', e);
                document.getElementById('music-library').innerHTML = '<p style="text-align:center; color:#f44336; padding:40px;">❌ Lỗi tải danh sách nhạc</p>';
            }
        }
        
        function renderMusicLibrary(files) {
            const html = files.map((file, index) => {
                const originalIndex = allMusicFiles.findIndex(f => f.filename === file.filename);
                const isCurrentTrack = originalIndex === currentTrackIndex;
                const isTrackPlaying = isCurrentTrack && isPlaying;
                
                return `
                <div class="music-item ${isTrackPlaying ? 'playing' : ''}" 
                     data-index="${originalIndex}"
                     onmouseenter="this.querySelector('.play-btn-hover').style.opacity='1'" 
                     onmouseleave="this.querySelector('.play-btn-hover').style.opacity='0'" 
                     style="cursor:pointer; display: flex; align-items: center; padding: 12px; border-radius: 8px; margin-bottom: 8px; background: ${isCurrentTrack ? 'linear-gradient(135deg, rgba(102,126,234,0.12) 0%, rgba(118,75,162,0.12) 100%)' : '#f9fafb'}; transition: all 0.2s ease; border-left: 4px solid ${isCurrentTrack ? '#667eea' : 'transparent'}; border: 1px solid ${isCurrentTrack ? '#c7d2fe' : 'transparent'};">
                    
                    <!-- Play Button (hover) -->
                    <div class="play-btn-hover" onclick="playTrack(${originalIndex}); event.stopPropagation();" 
                         style="width: 42px; height: 42px; margin-right: 12px; border-radius: 50%; background: ${isTrackPlaying ? '#667eea' : 'linear-gradient(135deg, #667eea, #764ba2)'}; display: flex; align-items: center; justify-content: center; color: white; font-size: 16px; opacity: ${isTrackPlaying ? '1' : '0'}; transition: all 0.2s ease; box-shadow: 0 4px 12px rgba(102, 126, 234, 0.3); cursor: pointer;" 
                         title="${isTrackPlaying ? 'Đang phát' : 'Click để phát'}">
                        ${isTrackPlaying ? '⏸' : '▶'}
                    </div>
                    
                    <!-- Music Icon (default state) -->
                    <div class="icon" style="font-size: 1.5em; margin-right: 12px; ${isTrackPlaying ? 'display:none;' : ''}">${isCurrentTrack ? '🔊' : '🎵'}</div>
                    
                    <!-- Track Info (clickable) -->
                    <div class="info" onclick="playTrack(${originalIndex})" style="flex: 1; cursor: pointer;">
                        <div class="name" style="font-weight: 600; color: ${isCurrentTrack ? '#667eea' : '#333'}; margin-bottom: 3px;">${file.filename}</div>
                        <div class="details" style="font-size: 0.85em; color: #6b7280;">${file.path} • ${file.size_mb} MB</div>
                    </div>
                    
                    <!-- Now Playing Indicator -->
                    ${isTrackPlaying ? '<div style="display:flex; align-items:center; gap:5px; color:#667eea; font-size:12px; animation: pulse 1.5s infinite;"><div style="width:3px; height:12px; background:#667eea; animation: wave1 0.8s ease-in-out infinite;"></div><div style="width:3px; height:18px; background:#667eea; animation: wave2 0.8s ease-in-out infinite 0.1s;"></div><div style="width:3px; height:15px; background:#667eea; animation: wave3 0.8s ease-in-out infinite 0.2s;"></div></div>' : ''}
                </div>
            `}).join('');
            
            document.getElementById('music-library').innerHTML = html || '<p style="text-align:center; color:#999; padding:40px;">Không có bài hát nào</p>';
        }
        
        // Toggle Shuffle
        async function toggleShuffle() {
            try {
                const response = await fetch('/api/vlc_shuffle', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({})
                });
                const data = await response.json();
                
                if (data.success) {
                    isShuffleOn = data.shuffle;
                    const btn = document.getElementById('shuffle-btn');
                    btn.style.opacity = isShuffleOn ? '1' : '0.6';
                    btn.style.transform = isShuffleOn ? 'scale(1.1)' : 'scale(1)';
                    addLog(isShuffleOn ? '🔀 Bật phát ngẫu nhiên' : '🔀 Tắt phát ngẫu nhiên', 'success');
                }
            } catch (e) {
                console.error('Toggle shuffle failed', e);
            }
        }
        
        // Toggle Repeat
        async function toggleRepeat() {
            try {
                const response = await fetch('/api/vlc_repeat', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({})
                });
                const data = await response.json();
                
                if (data.success) {
                    repeatMode = data.repeat_mode;
                    const btn = document.getElementById('repeat-btn');
                    
                    switch(repeatMode) {
                        case 0:
                            btn.textContent = '🔁';
                            btn.style.opacity = '0.6';
                            addLog('🔁 Tắt lặp lại', 'success');
                            break;
                        case 1:
                            btn.textContent = '🔁';
                            btn.style.opacity = '1';
                            addLog('🔁 Lặp lại tất cả', 'success');
                            break;
                        case 2:
                            btn.textContent = '🔂';
                            btn.style.opacity = '1';
                            addLog('🔂 Lặp lại một bài', 'success');
                            break;
                    }
                }
            } catch (e) {
                console.error('Toggle repeat failed', e);
            }
        }
        
        // Volume Control
        function setPlayerVolume(value) {
            document.getElementById('volume-value').textContent = value + '%';
            
            // Update slider gradient
            const slider = document.getElementById('volume-slider');
            slider.style.background = `linear-gradient(to right, #667eea 0%, #667eea ${value}%, #374151 ${value}%, #374151 100%)`;
            
            // Update icon
            const icon = document.getElementById('volume-icon');
            if (value == 0) {
                icon.textContent = '🔇';
            } else if (value < 30) {
                icon.textContent = '🔈';
            } else if (value < 70) {
                icon.textContent = '🔉';
            } else {
                icon.textContent = '🔊';
            }
            
            // Call VLC API directly to set volume
            fetch('/api/vlc_volume', {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify({level: parseInt(value)})
            }).catch(e => console.error('Volume set failed', e));
        }
        
        let lastVolume = 80;
        function toggleMute() {
            const slider = document.getElementById('volume-slider');
            if (parseInt(slider.value) > 0) {
                lastVolume = slider.value;
                slider.value = 0;
                setPlayerVolume(0);
            } else {
                slider.value = lastVolume;
                setPlayerVolume(lastVolume);
            }
        }
        
        // SINGLE-CLICK TO PLAY (like Spotify/Apple Music)
        // Removed complex double-click logic - direct click to play for better UX
        
        // Cập nhật visualizer state
        function updateVisualizer(playing) {
            const visualizer = document.getElementById('audio-visualizer');
            if (visualizer) {
                if (playing) {
                    visualizer.style.display = 'flex';
                    visualizer.classList.remove('paused');
                } else {
                    visualizer.classList.add('paused');
                }
            }
        }
        
        async function playTrack(index) {
            if (!allMusicFiles[index]) {
                console.error('Track not found at index:', index);
                addLog('❌ Không tìm thấy bài hát', 'error');
                return;
            }
            
            try {
                const track = allMusicFiles[index];
                console.log('🎵 Playing track:', track.filename);
                addLog(`⏳ Đang tải: ${track.filename}...`, 'info');
                
                // Gọi API trực tiếp để phát nhạc
                const response = await fetch('/api/vlc_play_file', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({filename: track.filename})
                });
                const data = await response.json();
                console.log('Play response:', data);
                
                if (data.success) {
                    currentTrackIndex = index;
                    isPlaying = true;
                    updateNowPlaying();
                    updateVisualizer(true);
                    renderMusicLibrary(currentPlaylist);
                    document.getElementById('play-btn').textContent = '⏸️';
                    addLog(`🎵 Đang phát: ${track.filename}`, 'success');
                    
                    // Start VLC polling for real-time sync
                    startVlcPolling();
                } else {
                    console.error('Play failed:', data);
                    addLog('❌ ' + (data.error || 'Không thể phát nhạc'), 'error');
                }
            } catch (e) {
                console.error('Failed to play track', e);
                addLog('❌ Lỗi kết nối', 'error');
            }
        }
        
        async function musicPlayPause() {
            try {
                // Gọi VLC API trực tiếp - không qua tool registry
                const response = await fetch('/api/vlc_play_pause', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'}
                });
                const data = await response.json();
                
                if (data.success) {
                    isPlaying = data.is_playing;
                    document.getElementById('play-btn').textContent = isPlaying ? '⏸️' : '▶️';
                    updateVisualizer(isPlaying);
                    renderMusicLibrary(currentPlaylist);
                    addLog(data.message, 'success');
                } else {
                    addLog('❌ ' + (data.error || 'Lỗi play/pause'), 'error');
                }
            } catch (e) {
                console.error('Play/Pause failed', e);
                addLog('❌ Lỗi kết nối VLC', 'error');
            }
        }
        
        async function musicNext() {
            try {
                const response = await fetch('/api/vlc_next', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'}
                });
                const data = await response.json();
                
                if (data.success) {
                    currentTrackIndex = (currentTrackIndex + 1) % currentPlaylist.length;
                    updateNowPlaying();
                    renderMusicLibrary(currentPlaylist);
                    addLog(data.message || '⏭️ Bài tiếp theo', 'success');
                } else {
                    addLog('❌ ' + (data.error || 'Không có bài tiếp'), 'error');
                }
            } catch (e) {
                console.error('Next track failed', e);
                addLog('❌ Lỗi chuyển bài', 'error');
            }
        }
        
        async function musicPrevious() {
            try {
                const response = await fetch('/api/vlc_previous', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'}
                });
                const data = await response.json();
                
                if (data.success) {
                    currentTrackIndex = (currentTrackIndex - 1 + currentPlaylist.length) % currentPlaylist.length;
                    updateNowPlaying();
                    renderMusicLibrary(currentPlaylist);
                    addLog(data.message || '⏮️ Bài trước', 'success');
                } else {
                    addLog('❌ ' + (data.error || 'Không có bài trước'), 'error');
                }
            } catch (e) {
                console.error('Previous track failed', e);
                addLog('❌ Lỗi chuyển bài', 'error');
            }
        }
        
        async function musicStop() {
            try {
                const response = await fetch('/api/vlc_stop', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'}
                });
                const data = await response.json();
                
                if (data.success) {
                    isPlaying = false;
                    currentTrackIndex = -1;
                    document.getElementById('current-track').textContent = '🎵 Đã dừng phát';
                    document.getElementById('track-info').textContent = 'Chọn bài hát để phát';
                    document.getElementById('track-album').textContent = '';
                    document.getElementById('album-art').innerHTML = '🎵';
                    document.getElementById('play-btn').textContent = '▶️';
                    const slider = document.getElementById('progress-slider');
                    if (slider) {
                        slider.value = 0;
                        slider.style.background = 'linear-gradient(to right, #667eea 0%, #667eea 0%, #374151 0%, #374151 100%)';
                    }
                    document.getElementById('current-time').textContent = '0:00';
                    document.getElementById('total-time').textContent = '0:00';
                    addLog(data.message || '⏹️ Đã dừng nhạc', 'success');
                } else {
                    addLog('❌ ' + (data.error || 'Lỗi dừng nhạc'), 'error');
                    renderMusicLibrary(currentPlaylist);
                    addLog('⏹️ Đã dừng phát nhạc', 'success');
                }
            } catch (e) {
                console.error('Stop failed', e);
            }
        }
        
        function updateNowPlaying() {
            if (currentTrackIndex >= 0 && allMusicFiles[currentTrackIndex]) {
                const track = allMusicFiles[currentTrackIndex];
                document.getElementById('current-track').textContent = track.filename.replace(/\\.[^/.]+$/, ''); // Remove extension
                document.getElementById('track-info').textContent = `${track.path}`;
                document.getElementById('track-album').textContent = `${track.size_mb} MB • Bài ${currentTrackIndex + 1}/${allMusicFiles.length}`;
                
                // Update album art with music note animation
                const albumArt = document.getElementById('album-art');
                if (albumArt) {
                    albumArt.innerHTML = isPlaying ? '<div style="animation: spin 3s linear infinite;">🎵</div>' : '🎵';
                }
            }
        }
        
        async function updateMusicStatus() {
            try {
                const response = await fetch('/api/call_tool', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({tool: 'get_music_status', args: {}})
                });
                const data = await response.json();
                
                if (data.success) {
                    // Sync playing state
                    const wasPlaying = isPlaying;
                    isPlaying = data.is_playing === 1 || data.is_playing === true;
                    
                    // Update play button
                    const playBtn = document.getElementById('play-btn');
                    if (playBtn) {
                        playBtn.textContent = isPlaying ? '⏸️' : '▶️';
                    }
                    
                    // Update progress bar and time
                    if (data.current_time !== undefined && data.duration !== undefined) {
                        const currentSec = parseFloat(data.current_time) || 0;
                        const totalSec = parseFloat(data.duration) || 0;
                        
                        if (totalSec > 0) {
                            // Update progress slider (only if not dragging)
                            const percentage = (currentSec / totalSec) * 100;
                            const slider = document.getElementById('progress-slider');
                            if (slider && !isDraggingProgress) {
                                slider.value = Math.min(100, Math.max(0, percentage));
                                slider.style.background = `linear-gradient(to right, #667eea 0%, #667eea ${percentage}%, #374151 ${percentage}%, #374151 100%)`;
                            }
                            
                            // Update time displays
                            const currentTimeEl = document.getElementById('current-time');
                            const totalTimeEl = document.getElementById('total-time');
                            if (currentTimeEl) currentTimeEl.textContent = formatTime(currentSec);
                            if (totalTimeEl) totalTimeEl.textContent = formatTime(totalSec);
                        }
                    }
                    
                    // Update library UI if play state changed
                    if (wasPlaying !== isPlaying && currentPlaylist.length > 0) {
                        renderMusicLibrary(currentPlaylist);
                    }
                }
            } catch (e) {
                console.error('Update music status error:', e);
            }
        }
        
        function formatTime(seconds) {
            if (!seconds || seconds < 0) return '0:00';
            const mins = Math.floor(seconds / 60);
            const secs = Math.floor(seconds % 60);
            return mins + ':' + (secs < 10 ? '0' : '') + secs;
        }
        
        // Progress bar dragging state
        let isDraggingProgress = false;
        
        // Called while dragging (preview only, no seek)
        function onProgressDrag(value) {
            isDraggingProgress = true;
            // Update slider visual immediately
            const slider = document.getElementById('progress-slider');
            slider.style.background = `linear-gradient(to right, #667eea 0%, #667eea ${value}%, #374151 ${value}%, #374151 100%)`;
        }
        
        // Called when drag ends (actual seek)
        async function onProgressSeek(value) {
            isDraggingProgress = false;
            const percentage = parseFloat(value);
            
            try {
                const response = await fetch('/api/call_tool', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({tool: 'seek_music', args: {percentage: percentage}})
                });
                const data = await response.json();
                
                if (data.success) {
                    await updateMusicStatus();
                }
            } catch (e) {
                console.error('Seek failed', e);
            }
        }
        
        async function seekTrack(event) {
            const progressBar = event.currentTarget;
            const rect = progressBar.getBoundingClientRect();
            const clickX = event.clientX - rect.left;
            const percentage = (clickX / rect.width) * 100;
            
            try {
                // Gọi tool để seek (cần implement trong backend)
                const response = await fetch('/api/call_tool', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({tool: 'seek_music', args: {percentage: percentage}})
                });
                const data = await response.json();
                
                if (data.success) {
                    // Cập nhật progress bar ngay lập tức
                    document.getElementById('progress-fill').style.width = percentage + '%';
                    await updateMusicStatus();
                }
            } catch (e) {
                console.error('Seek failed', e);
            }
        }
        
        connectWS();
        // Giảm polling từ 5s xuống 10s để giảm tải
        setInterval(getResources, 10000);
        getResources();
        
        // Load quotas on startup and refresh every 60 seconds
        getQuotas();
        setInterval(getQuotas, 60000);
        
        // Start VLC status polling for real-time sync
        startVlcPolling();
        
        // Initial VLC status check
        setTimeout(pollVlcStatus, 500);
        
        // RunCat Animation - Multiple frames like RunCat365
        let runcatFrame = 0;
        let runcatSpeed = 500; // Default 500ms per frame
        const runcatFrames = ['🐱', '🐈', '😺', '😸', '😹'];
        
        function animateRunCat() {
            const runcat = document.getElementById('runcat');
            if (!runcat) return;
            
            runcatFrame = (runcatFrame + 1) % runcatFrames.length;
            runcat.textContent = runcatFrames[runcatFrame];
            
            // Apply transform for running effect
            const offset = runcatFrame % 2 === 0 ? -3 : -1;
            const flip = runcatFrame >= 2 && runcatFrame <= 3 ? -1 : 1;
            runcat.style.transform = `translateY(${offset}px) scaleX(${flip})`;
            
            setTimeout(animateRunCat, runcatSpeed);
        }
        
        // Start RunCat animation
        setTimeout(animateRunCat, 100);
        
        // Auto-update music status every 1 second when music section is active
        setInterval(() => {
            const musicSection = document.getElementById('music-section');
            if (musicSection && musicSection.style.display !== 'none') {
                updateMusicStatus();
            }
        }, 1000);
        
        // Music Settings Functions
        function loadMusicFolderSettings() {
            const savedPath = localStorage.getItem('musicFolderPath');
            if (savedPath) {
                document.getElementById('music-folder-path').value = savedPath;
            }
        }
        
        function browseMusicFolder() {
            // Web không thể browse folder trực tiếp, hướng dẫn user
            alert('💡 Hướng dẫn:\\n\\n1. Mở File Explorer (Windows + E)\\n2. Đi đến thư mục nhạc của bạn\\n3. Click vào thanh địa chỉ và copy đường dẫn (Ctrl+C)\\n4. Paste vào ô bên trái (Ctrl+V)\\n5. Click "💾 Lưu"\\n\\nVí dụ: C:\\\\\\\\Users\\\\\\\\YourName\\\\\\\\Music');
        }
        
        async function saveMusicFolder() {
            const folderPath = document.getElementById('music-folder-path').value.trim();
            const statusEl = document.getElementById('music-folder-status');
            
            if (!folderPath) {
                statusEl.style.display = 'block';
                statusEl.style.background = '#fee2e2';
                statusEl.style.color = '#991b1b';
                statusEl.innerHTML = '❌ Vui lòng nhập đường dẫn thư mục!';
                return;
            }
            
            try {
                // Lưu vào localStorage
                localStorage.setItem('musicFolderPath', folderPath);
                
                // Gọi tool để lưu config
                const response = await fetch('/api/call_tool', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({
                        tool: 'save_music_folder_config',
                        args: {folder_path: folderPath}
                    })
                });
                const data = await response.json();
                
                if (data.success) {
                    statusEl.style.display = 'block';
                    statusEl.style.background = '#d1fae5';
                    statusEl.style.color = '#065f46';
                    statusEl.innerHTML = '✅ Đã lưu cài đặt thành công! LLM sẽ ưu tiên phát nhạc từ thư mục này.';
                    addLog(`⚙️ Đã cấu hình thư mục nhạc: ${folderPath}`, 'success');
                } else {
                    throw new Error(data.error || 'Unknown error');
                }
            } catch (e) {
                statusEl.style.display = 'block';
                statusEl.style.background = '#fee2e2';
                statusEl.style.color = '#991b1b';
                statusEl.innerHTML = `❌ Lỗi: ${e.message}`;
                console.error('Save music folder error:', e);
            }
        }
        
        // ============================================================
        // 💬 LLM CHAT FUNCTIONS - Gửi tin nhắn cho Robot/LLM
        // ============================================================
        
        async function refreshLLMConnectionStatus() {
            try {
                const response = await fetch('/api/llm_connection_status');
                const data = await response.json();
                
                if (data.success) {
                    data.devices.forEach((device, index) => {
                        // Update old status display (if exists)
                        const statusEl = document.getElementById(`device${index + 1}-status`);
                        if (statusEl) {
                            const icon = device.connected ? '✅' : (device.enabled ? '⏳' : '❌');
                            const text = device.connected ? 'Đã kết nối' : (device.enabled ? 'Đang kết nối...' : 'Chưa cấu hình');
                            statusEl.innerHTML = `📱 ${device.name}: <span class="status-indicator">${icon} ${text}</span>`;
                        }
                        
                        // Update new device card indicator
                        const indicator = document.getElementById(`device-${index + 1}-indicator`);
                        const card = document.getElementById(`device-${index + 1}-card`);
                        if (indicator) {
                            if (device.connected) {
                                indicator.innerHTML = '<span class="status-dot" style="width:8px;height:8px;border-radius:50%;background:#10b981;animation:pulse 2s infinite;"></span> ✅ Đã kết nối';
                                indicator.style.background = '#d1fae5';
                                indicator.style.color = '#047857';
                                if (card) card.style.boxShadow = '0 0 20px rgba(16, 185, 129, 0.4)';
                            } else if (device.enabled) {
                                indicator.innerHTML = '<span class="status-dot" style="width:8px;height:8px;border-radius:50%;background:#f59e0b;animation:blink 1s infinite;"></span> ⏳ Đang kết nối...';
                                indicator.style.background = '#fef3c7';
                                indicator.style.color = '#b45309';
                                if (card) card.style.boxShadow = '0 0 15px rgba(245, 158, 11, 0.3)';
                            } else {
                                indicator.innerHTML = '<span class="status-dot" style="width:8px;height:8px;border-radius:50%;background:#6b7280;"></span> ❌ Chưa kết nối';
                                indicator.style.background = '#f3f4f6';
                                indicator.style.color = '#6b7280';
                                if (card) card.style.boxShadow = 'none';
                            }
                        }
                    });
                    
                    // Update device selector
                    const select = document.getElementById('llm-device-select');
                    if (select) {
                        data.devices.forEach((device, index) => {
                            const option = select.options[index];
                            if (option) {
                                option.text = `${device.connected ? '🟢' : '⚪'} ${device.name}`;
                            }
                        });
                    }
                }
            } catch (e) {
                console.error('Error refreshing LLM connection status:', e);
            }
        }
        
        async function sendLLMMessage() {
            const input = document.getElementById('llm-chat-input');
            const message = input.value.trim();
            
            if (!message) {
                addLog('⚠️ Vui lòng nhập tin nhắn', 'error');
                return;
            }
            
            const modelSelect = document.getElementById('llm-chat-model');
            const selectedModel = modelSelect ? modelSelect.value : 'models/gemini-3-flash-preview';
            
            // Add user message to chat
            addLLMChatMessage('user', message, null);
            
            // Clear input
            input.value = '';
            input.style.height = '50px';
            
            // Show typing indicator
            showLLMTyping();
            
            try {
                // Call Gemini AI with Knowledge Base integration
                const response = await fetch('/api/tool/ask_gemini', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({
                        prompt: message,
                        model: selectedModel
                    })
                });
                
                const data = await response.json();
                
                // Hide typing indicator
                hideLLMTyping();
                
                if (data.success) {
                    const responseText = data.response || data.response_text || 'Không có nội dung trả về';
                    const hasKB = data.knowledge_base_used ? '📚' : '';
                    const modelName = getModelDisplayName(selectedModel);
                    
                    addLLMChatMessage('assistant', responseText, `Gemini ${modelName}${hasKB}`);
                    
                    if (data.knowledge_base_used) {
                        addLog(`✅ Gemini trả lời (sử dụng Knowledge Base)`, 'success');
                    } else {
                        addLog(`✅ Gemini trả lời thành công`, 'success');
                    }
                    
                    // 🔊 Text-to-Speech nếu được bật
                    const ttsToggle = document.getElementById('llm-tts-toggle');
                    const ttsEnabled = ttsToggle?.checked;
                    console.log('TTS Toggle element:', ttsToggle);
                    console.log('TTS Enabled:', ttsEnabled);
                    if (ttsEnabled && responseText) {
                        console.log('Calling speakText with:', responseText.substring(0, 100));
                        speakText(responseText);
                    }
                } else {
                    addLLMChatMessage('assistant', `❌ Lỗi: ${data.error}`, 'System');
                    addLog(`❌ Lỗi Gemini: ${data.error}`, 'error');
                }
            } catch (e) {
                hideLLMTyping();
                addLLMChatMessage('assistant', `❌ Lỗi kết nối: ${e.message}`, 'System');
                addLog(`❌ Lỗi: ${e.message}`, 'error');
            }
        }
        
        function getModelDisplayName(model) {
            if (model.includes('gemini-3')) return '3 Flash ⚡';
            if (model.includes('2.5-pro')) return '2.5 Pro 💎';
            if (model.includes('2.5-flash')) return '2.5 Flash ⚡';
            if (model.includes('2.0-flash')) return '2.0 Flash ⚡';
            return '';
        }
        
        function saveLLMChatModel() {
            const model = document.getElementById('llm-chat-model')?.value;
            if (model) {
                localStorage.setItem('llm_chat_model', model);
            }
        }
        
        function loadLLMChatModel() {
            const saved = localStorage.getItem('llm_chat_model') || 'models/gemini-3-flash-preview';
            const select = document.getElementById('llm-chat-model');
            if (select) {
                select.value = saved;
            }
            // Load TTS preference
            loadTTSPreference();
        }
        
        // ===== STT (Speech-to-Text) Functions - Microphone Input =====
        let llmRecognition = null;
        let llmIsRecording = false;
        let llmSilenceTimer = null;
        let llmLastSpeechTime = 0;
        const SILENCE_TIMEOUT = 2000; // 2 giây im lặng thì tự gửi
        
        // 🎯 Wake Word Detection
        let wakeWordRecognition = null;
        let wakeWordActive = false;
        let wakeWordWasActive = false; // 🆕 Track nếu wake word đang bật trước khi chat
        let wakeWordIdleTimer = null; // 🆕 Timer tự tắt sau 20s không dùng
        const WAKE_WORD_IDLE_TIMEOUT = 20000; // 20 giây không dùng thì tự tắt
        const WAKE_WORDS = ['hey gemini', 'hê gemini', 'ok gemini', 'ô kê gemini', 'xin chào', 'này gemini', 'gemini ơi', 'ê gemini'];
        const GOODBYE_WORDS = ['goodbye', 'good bye', 'tạm biệt', 'bye bye', 'bye', 'bai bai', 'ngủ đi', 'đi ngủ', 'tắt đi', 'dừng lại'];
        
        function initWakeWordDetection() {
            const SpeechRecognition = window.SpeechRecognition || window.webkitSpeechRecognition;
            if (!SpeechRecognition) return null;
            
            const recognition = new SpeechRecognition();
            recognition.lang = 'vi-VN';
            recognition.continuous = true;
            recognition.interimResults = true;
            
            recognition.onresult = (event) => {
                for (let i = event.resultIndex; i < event.results.length; i++) {
                    const transcript = event.results[i][0].transcript.toLowerCase().trim();
                    
                    // 🆕 Check goodbye word - Tắt wake word
                    const isGoodbye = GOODBYE_WORDS.some(word => transcript.includes(word));
                    if (isGoodbye) {
                        console.log('👋 Goodbye detected:', transcript);
                        addLog('👋 Goodbye! Tắt Wake Word...', 'info');
                        showVoiceStatus('👋 Tạm biệt! Đã tắt Wake Word.', 'success');
                        stopWakeWordDetection();
                        wakeWordWasActive = false;
                        localStorage.setItem('wake_word_enabled', 'false');
                        setTimeout(() => hideVoiceStatus(), 2000);
                        return;
                    }
                    
                    // Check wake word
                    const isWakeWord = WAKE_WORDS.some(word => transcript.includes(word));
                    if (isWakeWord && !llmIsRecording) {
                        console.log('🎯 Wake word detected:', transcript);
                        addLog('🎯 Wake word detected! Bắt đầu nghe...', 'success');
                        showVoiceStatus('🎯 Đã nghe thấy! Đang chuyển sang chế độ chat...', 'success');
                        
                        // 🆕 Mark wake word was active
                        wakeWordWasActive = true;
                        resetWakeWordIdleTimer();
                        
                        // Stop wake word detection, start chat recording
                        stopWakeWordDetection();
                        setTimeout(() => startLLMVoiceInput(), 300);
                        return;
                    }
                }
            };
            
            recognition.onend = () => {
                if (wakeWordActive) {
                    try { recognition.start(); } catch(e) {}
                }
            };
            
            recognition.onerror = (event) => {
                if (event.error !== 'no-speech' && event.error !== 'aborted') {
                    console.error('Wake word error:', event.error);
                }
            };
            
            return recognition;
        }
        
        // 🆕 Reset idle timer - Sau 20s không nói gì sẽ tự tắt wake word
        function resetWakeWordIdleTimer() {
            if (wakeWordIdleTimer) {
                clearTimeout(wakeWordIdleTimer);
            }
            wakeWordIdleTimer = setTimeout(() => {
                if (wakeWordActive && !llmIsRecording) {
                    addLog('⏰ Wake Word tự tắt sau 20s không hoạt động', 'info');
                    showVoiceStatus('⏰ Wake Word tự tắt (hết thời gian chờ)', 'warning');
                    stopWakeWordDetection();
                    wakeWordWasActive = false;
                    localStorage.setItem('wake_word_enabled', 'false');
                    setTimeout(() => hideVoiceStatus(), 2000);
                }
            }, WAKE_WORD_IDLE_TIMEOUT);
        }
        
        function startWakeWordDetection() {
            if (!wakeWordRecognition) {
                wakeWordRecognition = initWakeWordDetection();
            }
            if (!wakeWordRecognition) {
                addLog('❌ Trình duyệt không hỗ trợ Wake Word', 'error');
                return;
            }
            
            wakeWordActive = true;
            wakeWordWasActive = true;
            localStorage.setItem('wake_word_enabled', 'true');
            
            try {
                wakeWordRecognition.start();
                updateWakeWordButton(true);
                showVoiceStatus('👂 Đang lắng nghe... Nói "Hey Gemini" hoặc "Goodbye" để tắt', 'recording');
                addLog('👂 Wake word đang lắng nghe... Nói "Hey Gemini" để chat, "Goodbye" để tắt', 'info');
                resetWakeWordIdleTimer();
            } catch(e) {
                if (e.name === 'InvalidStateError') {
                    wakeWordRecognition.stop();
                    setTimeout(() => startWakeWordDetection(), 100);
                }
            }
        }
        
        function stopWakeWordDetection() {
            wakeWordActive = false;
            if (wakeWordIdleTimer) {
                clearTimeout(wakeWordIdleTimer);
                wakeWordIdleTimer = null;
            }
            if (wakeWordRecognition) {
                try { wakeWordRecognition.stop(); } catch(e) {}
            }
            updateWakeWordButton(false);
        }
        
        function toggleWakeWord() {
            if (wakeWordActive) {
                stopWakeWordDetection();
                wakeWordWasActive = false;
                localStorage.setItem('wake_word_enabled', 'false');
                addLog('👂 Đã tắt Wake Word detection', 'info');
                hideVoiceStatus();
            } else {
                startWakeWordDetection();
            }
        }
        
        // 🆕 Re-enable wake word after chat response (nếu trước đó đang bật)
        function reEnableWakeWordAfterResponse() {
            if (wakeWordWasActive) {
                setTimeout(() => {
                    if (!llmIsRecording && wakeWordWasActive) {
                        startWakeWordDetection();
                    }
                }, 1500); // Wait 1.5s after response
            }
        }
        
        function updateWakeWordButton(active) {
            const btn = document.getElementById('llm-wakeword-btn');
            if (btn) {
                if (active) {
                    btn.style.background = 'linear-gradient(135deg,#8b5cf6,#7c3aed)';
                    btn.innerHTML = '👂';
                    btn.title = '👂 Wake Word đang lắng nghe... (Click để tắt)';
                    btn.style.animation = 'pulse 2s infinite';
                } else {
                    btn.style.background = 'linear-gradient(135deg,#6b7280,#4b5563)';
                    btn.innerHTML = '👂';
                    btn.title = '👂 Bật Wake Word (nói "Hey Gemini" để chat)';
                    btn.style.animation = 'none';
                }
            }
        }
        
        function initLLMSpeechRecognition() {
            // Check for browser support
            const SpeechRecognition = window.SpeechRecognition || window.webkitSpeechRecognition;
            if (!SpeechRecognition) {
                console.warn('Browser does not support Speech Recognition');
                return null;
            }
            
            const recognition = new SpeechRecognition();
            recognition.lang = 'vi-VN'; // Vietnamese
            recognition.continuous = true; // Keep listening
            recognition.interimResults = true; // Show partial results
            recognition.maxAlternatives = 1;
            
            recognition.onstart = () => {
                llmIsRecording = true;
                llmLastSpeechTime = Date.now();
                updateMicButton(true);
                showVoiceStatus('🎤 Đang nghe... Nói xong sẽ tự động gửi!', 'recording');
                addLog('🎤 Bắt đầu ghi âm (auto-send sau 2s im lặng)', 'info');
                startSilenceDetection();
            };
            
            recognition.onresult = (event) => {
                let interimTranscript = '';
                let finalTranscript = '';
                
                for (let i = event.resultIndex; i < event.results.length; i++) {
                    const transcript = event.results[i][0].transcript;
                    if (event.results[i].isFinal) {
                        finalTranscript += transcript;
                    } else {
                        interimTranscript += transcript;
                    }
                }
                
                // Reset silence timer on speech
                llmLastSpeechTime = Date.now();
                
                const input = document.getElementById('llm-chat-input');
                if (input) {
                    if (finalTranscript) {
                        // Append final result to existing text
                        const existingText = input.value.trim();
                        input.value = existingText ? existingText + ' ' + finalTranscript : finalTranscript;
                        showVoiceStatus('✅ ' + input.value.substring(0, 60) + (input.value.length > 60 ? '...' : ''), 'success');
                    } else if (interimTranscript) {
                        // Show interim result
                        showVoiceStatus('🎤 ' + interimTranscript, 'recording');
                    }
                }
            };
            
            recognition.onerror = (event) => {
                console.error('Speech recognition error:', event.error);
                if (event.error === 'not-allowed') {
                    showVoiceStatus('❌ Vui lòng cho phép truy cập microphone!', 'error');
                    addLog('❌ Microphone bị từ chối quyền truy cập', 'error');
                } else if (event.error === 'no-speech') {
                    // Auto-send if have text and no speech
                    autoSendIfHaveText();
                    return;
                } else {
                    showVoiceStatus('❌ Lỗi: ' + event.error, 'error');
                    addLog('❌ STT lỗi: ' + event.error, 'error');
                }
                stopLLMVoiceInput();
            };
            
            recognition.onend = () => {
                if (llmIsRecording) {
                    // Check if should auto-send
                    const timeSinceLastSpeech = Date.now() - llmLastSpeechTime;
                    if (timeSinceLastSpeech >= SILENCE_TIMEOUT) {
                        autoSendIfHaveText();
                    } else {
                        // Auto-restart if still recording
                        try {
                            recognition.start();
                        } catch (e) {
                            stopLLMVoiceInput();
                        }
                    }
                } else {
                    updateMicButton(false);
                    hideVoiceStatus();
                }
            };
            
            return recognition;
        }
        
        function startSilenceDetection() {
            if (llmSilenceTimer) clearInterval(llmSilenceTimer);
            
            llmSilenceTimer = setInterval(() => {
                if (!llmIsRecording) {
                    clearInterval(llmSilenceTimer);
                    return;
                }
                
                const timeSinceLastSpeech = Date.now() - llmLastSpeechTime;
                const input = document.getElementById('llm-chat-input');
                
                if (timeSinceLastSpeech >= SILENCE_TIMEOUT && input && input.value.trim()) {
                    // Auto-send after silence
                    autoSendIfHaveText();
                } else if (timeSinceLastSpeech >= 1000 && input && input.value.trim()) {
                    // Show countdown
                    const remaining = Math.ceil((SILENCE_TIMEOUT - timeSinceLastSpeech) / 1000);
                    showVoiceStatus(`⏳ Gửi sau ${remaining}s... (nói tiếp để hủy)`, 'warning');
                }
            }, 500);
        }
        
        function autoSendIfHaveText() {
            const input = document.getElementById('llm-chat-input');
            if (input && input.value.trim()) {
                showVoiceStatus('📤 Đang gửi tin nhắn...', 'success');
                stopLLMVoiceInput();
                
                // Small delay then send
                setTimeout(() => {
                    sendLLMMessage();
                    // 🆕 Re-enable wake word after response (dùng function mới)
                    reEnableWakeWordAfterResponse();
                }, 300);
            } else {
                stopLLMVoiceInput();
                // 🆕 Nếu không có text, vẫn re-enable wake word
                reEnableWakeWordAfterResponse();
            }
        }
        
        function toggleLLMVoiceInput() {
            if (llmIsRecording) {
                // If recording, stop and send if have text
                autoSendIfHaveText();
            } else {
                startLLMVoiceInput();
            }
        }
        
        function startLLMVoiceInput() {
            // Stop wake word if active
            if (wakeWordActive) {
                stopWakeWordDetection();
            }
            
            if (!llmRecognition) {
                llmRecognition = initLLMSpeechRecognition();
            }
            
            if (!llmRecognition) {
                showVoiceStatus('❌ Trình duyệt không hỗ trợ STT. Hãy dùng Chrome!', 'error');
                addLog('❌ Trình duyệt không hỗ trợ Speech Recognition', 'error');
                return;
            }
            
            // Clear input for fresh start
            const input = document.getElementById('llm-chat-input');
            if (input) input.value = '';
            
            try {
                llmRecognition.start();
            } catch (e) {
                if (e.name === 'InvalidStateError') {
                    // Already started
                    stopLLMVoiceInput();
                    setTimeout(() => startLLMVoiceInput(), 100);
                } else {
                    console.error('Start speech recognition error:', e);
                    showVoiceStatus('❌ Không thể bắt đầu ghi âm', 'error');
                }
            }
        }
        
        function stopLLMVoiceInput() {
            llmIsRecording = false;
            if (llmSilenceTimer) {
                clearInterval(llmSilenceTimer);
                llmSilenceTimer = null;
            }
            if (llmRecognition) {
                try {
                    llmRecognition.stop();
                } catch (e) {}
            }
            updateMicButton(false);
            setTimeout(() => hideVoiceStatus(), 1500);
        }
        
        function updateMicButton(isRecording) {
            const btn = document.getElementById('llm-mic-btn');
            if (btn) {
                if (isRecording) {
                    btn.style.background = 'linear-gradient(135deg,#ef4444,#dc2626)';
                    btn.innerHTML = '⏹️';
                    btn.title = '⏹️ Nhấn để dừng và gửi';
                    btn.style.animation = 'pulse 1s infinite';
                } else {
                    btn.style.background = 'linear-gradient(135deg,#10b981,#059669)';
                    btn.innerHTML = '🎤';
                    btn.title = '🎤 Nhấn để nói (auto-send)';
                    btn.style.animation = 'none';
                }
            }
        }
        
        function showVoiceStatus(text, type) {
            const statusDiv = document.getElementById('llm-voice-status');
            const statusText = document.getElementById('llm-voice-status-text');
            if (statusDiv && statusText) {
                statusDiv.style.display = 'block';
                statusText.textContent = text;
                
                if (type === 'recording') {
                    statusDiv.style.background = 'linear-gradient(135deg,#fef3c7,#fde68a)';
                } else if (type === 'success') {
                    statusDiv.style.background = 'linear-gradient(135deg,#d1fae5,#a7f3d0)';
                } else if (type === 'error') {
                    statusDiv.style.background = 'linear-gradient(135deg,#fee2e2,#fecaca)';
                } else if (type === 'warning') {
                    statusDiv.style.background = 'linear-gradient(135deg,#ffedd5,#fed7aa)';
                }
            }
        }
        
        function hideVoiceStatus() {
            const statusDiv = document.getElementById('llm-voice-status');
            if (statusDiv) {
                statusDiv.style.display = 'none';
            }
        }
        
        // ===== TTS (Text-to-Speech) Functions =====
        function saveTTSPreference() {
            const enabled = document.getElementById('llm-tts-toggle')?.checked || false;
            localStorage.setItem('llm_tts_enabled', enabled);
            if (enabled) {
                addLog('🔊 Đã bật đọc to câu trả lời', 'info');
            } else {
                addLog('🔇 Đã tắt đọc to câu trả lời', 'info');
            }
        }
        
        function loadTTSPreference() {
            const saved = localStorage.getItem('llm_tts_enabled') === 'true';
            const toggle = document.getElementById('llm-tts-toggle');
            if (toggle) {
                toggle.checked = saved;
            }
        }
        
        let currentTTSAudio = null; // Track current TTS audio
        
        async function speakText(text) {
            try {
                // Hiển thị indicator đang đọc
                showSpeakingIndicator();
                
                // Gọi API TTS backend
                const response = await fetch('/api/tts', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ text: text })
                });
                
                const data = await response.json();
                
                if (data.success) {
                    addLog(`🔊 Đang đọc: ${text.substring(0, 50)}...`, 'info');
                } else {
                    addLog(`❌ TTS lỗi: ${data.error}`, 'error');
                }
                
                hideSpeakingIndicator();
            } catch (e) {
                console.error('TTS error:', e);
                addLog(`❌ TTS lỗi: ${e.message}`, 'error');
                hideSpeakingIndicator();
            }
        }
        
        function stopSpeaking() {
            // Gọi API dừng TTS
            fetch('/api/tts/stop', { method: 'POST' })
                .then(() => {
                    addLog('🔇 Đã dừng đọc', 'info');
                    hideSpeakingIndicator();
                })
                .catch(e => console.error('Stop TTS error:', e));
        }
        
        function showSpeakingIndicator() {
            // Thêm indicator vào status bar
            const statusBar = document.getElementById('llm-ai-status');
            if (statusBar && !document.getElementById('speaking-indicator')) {
                const indicator = document.createElement('span');
                indicator.id = 'speaking-indicator';
                indicator.innerHTML = '<span style="animation:pulse 1s infinite;">🔊 Đang đọc...</span> <button onclick="stopSpeaking()" style="background:rgba(255,255,255,0.3);border:none;padding:2px 8px;border-radius:4px;cursor:pointer;font-size:0.8em;">⏹️ Dừng</button>';
                indicator.style.cssText = 'font-size:0.85em; background:rgba(255,255,255,0.2); padding:4px 10px; border-radius:20px; display:flex; align-items:center; gap:8px;';
                statusBar.appendChild(indicator);
            }
        }
        
        function hideSpeakingIndicator() {
            const indicator = document.getElementById('speaking-indicator');
            if (indicator) indicator.remove();
        }
        
        function addLLMChatMessage(role, content, deviceName) {
            const container = document.getElementById('llm-chat-messages');
            
            // Remove welcome message if exists
            const welcome = container.querySelector('div[style*="text-align:center"]');
            if (welcome) welcome.remove();
            
            const msgDiv = document.createElement('div');
            msgDiv.className = `llm-message ${role}`;
            
            const time = new Date().toLocaleTimeString('vi-VN', { hour: '2-digit', minute: '2-digit' });
            
            let deviceTag = '';
            if (deviceName && role === 'assistant') {
                deviceTag = `<span class="device-tag">${deviceName}</span>`;
            }
            
            msgDiv.innerHTML = `
                <div class="content">${content}${deviceTag}</div>
                <span class="time">${time}</span>
            `;
            
            container.appendChild(msgDiv);
            container.scrollTop = container.scrollHeight;
            
            // Store message
            llmChatMessages.push({ role, content, deviceName, time: new Date().toISOString() });
        }
        
        function showLLMTyping() {
            const container = document.getElementById('llm-chat-messages');
            const typingDiv = document.createElement('div');
            typingDiv.id = 'llm-typing-indicator';
            typingDiv.className = 'llm-message assistant';
            typingDiv.innerHTML = `
                <div class="llm-typing">
                    <span></span><span></span><span></span>
                </div>
            `;
            container.appendChild(typingDiv);
            container.scrollTop = container.scrollHeight;
        }
        
        function hideLLMTyping() {
            const typing = document.getElementById('llm-typing-indicator');
            if (typing) typing.remove();
        }
        
        function handleLLMChatKeydown(event) {
            if (event.key === 'Enter' && !event.shiftKey) {
                event.preventDefault();
                sendLLMMessage();
            }
        }
        
        function autoResizeLLMInput(textarea) {
            textarea.style.height = '50px';
            textarea.style.height = Math.min(textarea.scrollHeight, 150) + 'px';
        }
        
        function sendQuickMessage(message) {
            const input = document.getElementById('llm-chat-input');
            input.value = message;
            sendLLMMessage();
        }
        
        function clearLLMChat() {
            const container = document.getElementById('llm-chat-messages');
            container.innerHTML = `
                <div style="text-align:center; color:#666; padding:40px 20px;">
                    <div style="font-size:4em; margin-bottom:15px;">🤖</div>
                    <h3 style="color:#667eea; margin-bottom:10px;">Chào mừng đến Chat với Gemini AI!</h3>
                    <p style="font-size:0.95em; max-width:400px; margin:0 auto;">
                        Chat trực tiếp với Gemini AI.<br>
                        AI sẽ tự động tìm kiếm trong Knowledge Base của bạn để trả lời chính xác hơn.
                    </p>
                </div>
            `;
            llmChatMessages = [];
            addLog('🗑️ Đã xóa lịch sử chat', 'info');
        }
        
        // Load music folder settings when opening the section
        document.addEventListener('DOMContentLoaded', () => {
            loadMusicFolderSettings();
            // 🔥 FIX: Auto-load API keys when page loads
            loadCurrentEndpoint();
            // 🔥 FIX: Auto-refresh connection status
            refreshLLMConnectionStatus();
            // ⏰ Refresh connection status every 3 seconds
            setInterval(refreshLLMConnectionStatus, 3000);
        });
        
    // Initialize playlists on page load
    initPlaylists();
    </script>
    
    <!-- MINIZ FOOTER - Compact Corner -->
    <div class="footer-miniz">
        <div class="footer-logo-compact">
            <img src="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 100 100'%3E%3Ccircle cx='50' cy='50' r='48' fill='%23667eea'/%3E%3Cpath d='M30 40 L50 25 L70 40 M50 25 L50 75 M35 55 L50 50 L65 55 M35 70 L50 65 L65 70' stroke='white' stroke-width='3' fill='none'/%3E%3Ctext x='50' y='88' text-anchor='middle' fill='white' font-size='14' font-weight='bold' font-family='Arial'%3EminiZ%3C/text%3E%3C/svg%3E" alt="miniZ Logo">
            <span class="footer-brand-compact">miniZ</span>
        </div>
        <div class="footer-separator"></div>
        <a href="https://youtube.com/@minizjp?si=LRg5piGHmxYtsFJU" target="_blank" class="footer-youtube-compact" title="Kênh YouTube miniZ">
            <svg viewBox="0 0 24 24"><path d="M23.498 6.186a3.016 3.016 0 0 0-2.122-2.136C19.505 3.545 12 3.545 12 3.545s-7.505 0-9.377.505A3.017 3.017 0 0 0 .502 6.186C0 8.07 0 12 0 12s0 3.93.502 5.814a3.016 3.016 0 0 0 2.122 2.136c1.871.505 9.376.505 9.376.505s7.505 0 9.377-.505a3.015 3.015 0 0 0 2.122-2.136C24 15.93 24 12 24 12s0-3.93-.502-5.814zM9.545 15.568V8.432L15.818 12l-6.273 3.568z"/></svg>
            YouTube
        </a>
    </div>
</body>
</html>
    """
    return html

# ============================================================
# 📨 API ENDPOINT: SEND MESSAGE TO LLM
# ============================================================

class SendMessageRequest(BaseModel):
    message: str
    device_index: int = None
    wait_response: bool = True
    timeout: int = 30

class BroadcastMessageRequest(BaseModel):
    message: str
    wait_response: bool = False

@app.post("/api/send_message_to_llm")
async def api_send_message_to_llm(request: SendMessageRequest):
    """
    API endpoint để gửi tin nhắn cho LLM qua WebSocket.
    LLM sẽ đọc được tin nhắn và tự trả lời.
    """
    result = await send_message_to_llm(
        message=request.message,
        device_index=request.device_index,
        wait_response=request.wait_response,
        timeout=request.timeout
    )
    return result

@app.post("/api/broadcast_to_llm")
async def api_broadcast_to_llm(request: BroadcastMessageRequest):
    """
    API endpoint để broadcast tin nhắn đến tất cả LLM đang kết nối.
    """
    result = await broadcast_to_all_llm(
        message=request.message,
        wait_response=request.wait_response
    )
    return result

@app.get("/api/llm_connection_status")
async def api_llm_connection_status():
    """
    Kiểm tra trạng thái kết nối của các thiết bị LLM.
    """
    status = {
        "success": True,
        "devices": []
    }
    
    for i in range(3):
        device_status = {
            "index": i,
            "name": endpoints_config[i].get("name", f"Thiết bị {i + 1}"),
            "connected": xiaozhi_connected.get(i, False),
            "enabled": endpoints_config[i].get("enabled", False),
            "has_token": bool(endpoints_config[i].get("token", ""))
        }
        status["devices"].append(device_status)
    
    status["active_index"] = active_endpoint_index
    status["total_connected"] = sum(1 for v in xiaozhi_connected.values() if v)
    
    return status

# API Endpoints
@app.post("/api/volume")
async def api_volume(request: VolumeRequest):
    result = await set_volume(request.level)
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/screenshot")
async def api_screenshot():
    result = await take_screenshot()
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/notification")
async def api_notification(request: NotificationRequest):
    result = await show_notification(request.title, request.message)
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.get("/api/resources")
async def api_resources():
    result = await get_system_resources()
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.get("/api/quotas")
async def api_quotas():
    """Lấy thông tin quota của Gemini và Serper APIs"""
    result = await get_api_quotas()
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.get("/api/vlc_status")
async def api_vlc_status():
    """VLC status - MCP-style response với session tracking"""
    try:
        # Cache status để tránh query liên tục (200ms)
        import time
        now = time.time()
        if not hasattr(vlc_player, '_status_cache') or (now - vlc_player._status_cache_time) > 0.2:
            status = vlc_player.get_full_status()
            # MCP-style: thêm metadata
            status['timestamp'] = int(now * 1000)  # milliseconds
            status['session_id'] = getattr(vlc_player, '_session_id', 'default')
            vlc_player._status_cache = status
            vlc_player._status_cache_time = now
        return vlc_player._status_cache
    except Exception as e:
        return {
            "success": False, 
            "error": str(e), 
            "state": "error",
            "timestamp": int(time.time() * 1000)
        }

@app.post("/api/vlc_seek")
async def api_vlc_seek(data: dict):
    """Seek VLC player - MCP-style với validation và state tracking"""
    try:
        position = float(data.get("position", 0))
        
        # Validate input (xiaozhi pattern: validate before execution)
        if not 0.0 <= position <= 1.0:
            return {
                "success": False,
                "error": "Position must be between 0.0 and 1.0",
                "error_type": "validation_error",
                "provided_value": position
            }
        
        # Get current state
        old_position = vlc_player.get_position()
        current_time = vlc_player.get_time()
        
        # Execute seek
        vlc_player.set_position(position)
        
        # Calculate time delta
        new_time = vlc_player.get_time()
        
        return {
            "success": True,
            "action": "seek",
            "position": position,
            "previous_position": old_position,
            "time_delta_ms": new_time - current_time,
            "timestamp": int(time.time() * 1000),
            "message": f"Sought to {int(position * 100)}%"
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "error_type": "exception",
            "timestamp": int(time.time() * 1000)
        }

@app.post("/api/vlc_volume")
async def api_vlc_volume(data: dict):
    """Set VLC player volume (0-100)"""
    try:
        level = int(data.get("level", 80))
        vlc_player.set_volume(level)
        return {"success": True, "volume": level}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/api/vlc_shuffle")
async def api_vlc_shuffle(data: dict):
    """Toggle or set shuffle mode"""
    try:
        enabled = data.get("enabled")
        if enabled is None:
            # Toggle
            enabled = not vlc_player.get_shuffle()
        vlc_player.set_shuffle(enabled)
        return {"success": True, "shuffle": enabled}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/api/vlc_repeat")
async def api_vlc_repeat(data: dict):
    """Set repeat mode: 0=off, 1=all, 2=one"""
    try:
        mode = data.get("mode")
        if mode is None:
            # Cycle through modes
            current = vlc_player.get_repeat_mode()
            mode = (current + 1) % 3
        vlc_player.set_repeat_mode(mode)
        return {"success": True, "repeat_mode": mode}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/api/vlc_play_file")
async def api_vlc_play_file(data: dict):
    """Phát file nhạc trực tiếp qua VLC - cho Web UI double-click"""
    try:
        filename = data.get("filename", "")
        if not filename:
            return {"success": False, "error": "Thiếu filename"}
        
        print(f"🎵 [API] vlc_play_file: {filename}")
        
        # Gọi hàm play_music
        result = await play_music(filename=filename, create_playlist=True)
        print(f"🎵 [API] play_music result: {result}")
        return result
    except Exception as e:
        print(f"❌ [API] vlc_play_file error: {e}")
        return {"success": False, "error": str(e)}

@app.post("/api/vlc_play_pause")
async def api_vlc_play_pause():
    """Toggle VLC play/pause - MCP-style với state tracking"""
    try:
        if vlc_player and vlc_player._player:
            # Track state before action (xiaozhi pattern)
            was_playing = vlc_player.is_playing()
            
            # Execute command
            vlc_player.pause()
            
            # Get new state
            is_playing = vlc_player.is_playing()
            
            return {
                "success": True,
                "is_playing": is_playing,
                "previous_state": "playing" if was_playing else "paused",
                "current_state": "playing" if is_playing else "paused",
                "action": "pause" if was_playing else "play",
                "message": "▶️ Đang phát" if is_playing else "⏸️ Đã tạm dừng",
                "timestamp": int(time.time() * 1000)
            }
        return {
            "success": False, 
            "error": "VLC chưa khởi tạo hoặc chưa phát nhạc",
            "state": "not_initialized"
        }
    except Exception as e:
        return {
            "success": False, 
            "error": str(e),
            "error_type": "exception",
            "timestamp": int(time.time() * 1000)
        }

@app.post("/api/vlc_stop")
async def api_vlc_stop():
    """Stop VLC player - MCP-style với state cleanup"""
    try:
        if vlc_player and vlc_player._player:
            # Get current state before stopping
            was_playing = vlc_player.is_playing()
            current_media = vlc_player._player.get_media()
            stopped_track = current_media.get_meta(0) if current_media else "Unknown"
            
            # Execute stop
            vlc_player.stop()
            
            return {
                "success": True,
                "action": "stop",
                "message": "⏹️ Đã dừng nhạc",
                "previous_state": "playing" if was_playing else "paused",
                "stopped_track": stopped_track,
                "timestamp": int(time.time() * 1000)
            }
        return {
            "success": False,
            "error": "VLC chưa khởi tạo hoặc chưa phát nhạc",
            "state": "not_initialized"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/api/vlc_next")
async def api_vlc_next():
    """Next track - MCP-style async với immediate response"""
    try:
        if vlc_player and vlc_player._list_player:
            # Get current track info before switching (xiaozhi pattern)
            current_media = vlc_player._player.get_media()
            current_title = current_media.get_meta(0) if current_media else "Unknown"
            current_index = vlc_player._list_player.get_media_player().get_position()
            
            # Execute command
            vlc_player._list_player.next()
            vlc_player._list_player.play()  # Đảm bảo phát
            
            # MCP-style: trả về immediate response + track info
            return {
                "success": True,
                "action": "next",
                "message": "⏭️ Chuyển bài tiếp theo",
                "is_playing": True,
                "previous_track": {
                    "title": current_title,
                    "position": current_index
                },
                "timestamp": int(time.time() * 1000),
                "note": "Track info sẽ update sau 500ms qua /api/vlc_status"
            }
        return {
            "success": False,
            "error": "VLC chưa khởi tạo",
            "state": "not_initialized"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/api/vlc_previous")
async def api_vlc_previous():
    """Previous track - TỐI ƯU: Không block UI với sleep"""
    try:
        if vlc_player and vlc_player._list_player:
            vlc_player._list_player.previous()
            vlc_player._list_player.play()  # Đảm bảo phát
            # Trả về ngay - Web UI sẽ poll status để update
            return {
                "success": True, 
                "message": "⏮️ Chuyển bài trước",
                "is_playing": True
            }
        return {"success": False, "error": "VLC chưa khởi tạo"}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.get("/api/time")
async def api_time():
    result = await get_current_time()
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return {"data": result}

@app.post("/api/calculator")
async def api_calculator(request: CalculatorRequest):
    result = await calculator(request.expression)
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result


# ===== GENERIC TOOL CALLER =====

@app.post("/api/call_tool")
async def call_any_tool(data: dict):
    """Generic endpoint to call ANY tool from TOOLS registry"""
    tool_name = data.get("tool", data.get("name", ""))
    args = data.get("args", data.get("arguments", {}))
    
    if not tool_name:
        raise HTTPException(400, "Tool name is required")
    
    if tool_name not in TOOLS:
        raise HTTPException(404, f"Tool '{tool_name}' not found")
    
    try:
        handler = TOOLS[tool_name]["handler"]
        result = await handler(**args)
        return result
    except Exception as e:
        return {"success": False, "error": str(e)}

# ============================================================
# 🎯 VLC MCP ENDPOINTS - Hybrid System
# ============================================================

@app.post("/mcp/vlc/call")
async def mcp_vlc_call(request: dict):
    """
    MCP endpoint for VLC control (JSON-RPC 2.0)
    
    Xiaozhi-esp32 style protocol:
    {
      "jsonrpc": "2.0",
      "method": "tools/call",
      "params": {
        "name": "vlc.play",
        "arguments": {"file": "song.mp3"}
      },
      "id": 1
    }
    """
    if not VLC_MCP_AVAILABLE:
        return {
            "jsonrpc": "2.0",
            "error": {
                "code": -32603,
                "message": "VLC MCP server not available"
            },
            "id": request.get("id")
        }
    
    try:
        response = await vlc_mcp_server.handle_mcp_request(request)
        return response
    except Exception as e:
        return {
            "jsonrpc": "2.0",
            "error": {
                "code": -32603,
                "message": f"Internal error: {str(e)}"
            },
            "id": request.get("id")
        }

@app.get("/mcp/vlc/tools")
async def mcp_vlc_list_tools():
    """List all available VLC MCP tools"""
    if not VLC_MCP_AVAILABLE:
        return {
            "success": False,
            "error": "VLC MCP server not available"
        }
    
    return {
        "success": True,
        "tools": vlc_mcp_server.list_tools()
    }

@app.get("/mcp/vlc/status")
async def mcp_vlc_status():
    """Get VLC MCP server status"""
    return {
        "success": True,
        "mcp_available": VLC_MCP_AVAILABLE,
        "vlc_available": VLC_AVAILABLE,
        "total_tools": len(vlc_mcp_server.tools) if VLC_MCP_AVAILABLE else 0,
        "protocol": "JSON-RPC 2.0",
        "architecture": "xiaozhi-esp32"
    }

# ============================================================
# 🧠 INTENT DETECTION API ENDPOINTS
# ============================================================

@app.post("/api/detect_intent")
async def api_detect_intent(data: dict):
    """
    Phân tích intent từ text input
    Trả về suggested tool và confidence
    """
    text = data.get("text", data.get("query", ""))
    use_llm = data.get("use_llm", False)
    
    if not text:
        raise HTTPException(400, "Text is required")
    
    try:
        if use_llm:
            result = await intent_detector.detect_with_llm(text, GEMINI_API_KEY)
        else:
            result = intent_detector.detect_intent(text)
        
        return {
            "success": True,
            "text": text,
            **result
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/api/auto_execute")
async def api_auto_execute(data: dict):
    """
    🤖 AUTO TOOL EXECUTOR v2.0 - NÂNG CẤP
    
    Phân tích THÔNG MINH response từ LLM và tự động gọi tool
    
    IMPROVEMENTS:
    - ✅ Ưu tiên phân tích USER QUERY trước (chính xác hơn)
    - ✅ Phát hiện câu PHỦ ĐỊNH (không, chưa, đừng)
    - ✅ Phát hiện câu HỎI (có phải, có nên)
    - ✅ Context-aware patterns (xem trước/sau)
    - ✅ Multi-language support (Vi + En)
    - ✅ Better logging và debug info
    
    Args:
        llm_response: Text response từ LLM
        original_query: Câu hỏi gốc của user (QUAN TRỌNG - ưu tiên cao)
        auto_execute: True để tự động gọi tool (default: True)
    
    Returns:
        {
            "success": bool,
            "intent_detected": str,
            "tool_suggested": str,
            "confidence": float,
            "tool_executed": bool,
            "tool_result": dict,
            "analysis": {
                "source": "query|response",
                "matched_pattern": str,
                "is_question": bool,
                "is_negative": bool
            }
        }
    """
    try:
        llm_response = data.get("llm_response", data.get("response", "")).strip()
        original_query = data.get("original_query", data.get("query", "")).strip()
        auto_execute = data.get("auto_execute", True)
        
        print(f"\n{'='*70}")
        print(f"🤖 [Auto Execute v2.0] NEW REQUEST")
        print(f"{'='*70}")
        print(f"📝 User Query: '{original_query}'")
        print(f"💬 LLM Response: '{llm_response}'")
        print(f"⚙️  Auto Execute: {auto_execute}")
        print(f"{'-'*70}")
        
        # ===== BƯỚC 1: PHÂN TÍCH NGỮ CẢNH =====
        import re
        
        # Phát hiện câu phủ định
        negative_patterns = [
            r'\b(không|chưa|đừng|chớ|thôi|ngưng)\b',
            r'\b(no|not|don\'t|stop|cancel)\b'
        ]
        
        # Phát hiện câu hỏi
        question_patterns = [
            r'\b(có phải|có nên|có thể|được không|như thế nào)\b',
            r'\?$',  # Kết thúc bằng dấu ?
            r'\b(is|are|can|could|should|would|do|does)\b.+\?'
        ]
        
        # ===== BƯỚC 2: PATTERNS NÂNG CAP - CONTEXT AWARE =====
        enhanced_vlc_patterns = {
            "music_next": {
                "patterns": [
                    r'\b(bài tiếp theo|bài tiếp|next song|next track)\b',
                    r'\b(chuyển bài|skip|bài sau|bài kế|sang bài)\b',
                    r'\b(tiếp theo|next|forward)\b',
                    r'\b(phát bài tiếp|play next)\b'
                ],
                "keywords": ["next", "tiếp", "skip", "chuyển", "sau", "forward"]
            },
            "music_previous": {
                "patterns": [
                    r'\b(bài trước|previous song|previous track)\b',
                    r'\b(quay lại|back|lùi lại|trở lại)\b',
                    r'\b(bài trước đó|bài cũ)\b',
                    r'\b(phát bài trước|play previous)\b'
                ],
                "keywords": ["previous", "trước", "back", "quay", "lùi"]
            },
            "pause_music": {
                "patterns": [
                    r'\b(tạm dừng|pause)\b',
                    r'\b(dừng lại|stop playing|ngừng)\b',
                    r'\b(tạm ngưng)\b'
                ],
                "keywords": ["pause", "tạm", "dừng lại"]
            },
            "resume_music": {
                "patterns": [
                    r'\b(tiếp tục|resume|continue)\b',
                    r'\b(phát tiếp|play again|chạy tiếp)\b',
                    r'\b(mở lại|bật lại)\b'
                ],
                "keywords": ["resume", "tiếp tục", "continue", "phát tiếp"]
            },
            "stop_music": {
                "patterns": [
                    r'\b(dừng hẳn|stop completely)\b',
                    r'\b(tắt nhạc|stop music|ngừng nhạc)\b',
                    r'\b(dừng|stop)\b(?!.*playing)'  # "dừng" nhưng không có "playing"
                ],
                "keywords": ["stop", "dừng", "tắt", "ngừng"]
            },
            "play_music": {
                "patterns": [
                    r'\b(phát nhạc|play music)\b',
                    r'\b(mở nhạc|bật nhạc|chạy nhạc)\b',
                    r'\b(play song|start music)\b'
                ],
                "keywords": ["play", "phát", "mở", "bật", "chạy"]
            }
        }
        
        # ===== BƯỚC 3: PHÂN TÍCH ƯU TIÊN USER QUERY TRƯỚC =====
        detected_tool = None
        confidence = 0.0
        matched_pattern = None
        analysis_source = "none"
        
        # Priority 1: Phân tích USER QUERY (chính xác nhất)
        if original_query:
            query_lower = original_query.lower()
            
            # Kiểm tra phủ định và câu hỏi trong query
            is_negative = any(re.search(p, query_lower) for p in negative_patterns)
            is_question = any(re.search(p, query_lower) for p in question_patterns)
            
            print(f"🔍 [Analysis] Query Context:")
            print(f"   - Is Negative: {is_negative}")
            print(f"   - Is Question: {is_question}")
            
            if not is_negative and not is_question:
                # Chỉ phân tích khi KHÔNG phải câu phủ định hoặc câu hỏi
                for tool_name, tool_data in enhanced_vlc_patterns.items():
                    # Kiểm tra patterns
                    for pattern in tool_data["patterns"]:
                        if re.search(pattern, query_lower):
                            detected_tool = tool_name
                            confidence = 0.95  # VERY HIGH confidence vì từ user query
                            matched_pattern = pattern
                            analysis_source = "user_query"
                            print(f"✅ [Query Match] Tool: {tool_name} | Pattern: {pattern}")
                            break
                    
                    # Nếu chưa match, thử keyword matching
                    if not detected_tool:
                        keyword_count = sum(1 for kw in tool_data["keywords"] if kw in query_lower)
                        if keyword_count >= 1:
                            detected_tool = tool_name
                            confidence = 0.7 + (keyword_count * 0.1)  # Càng nhiều keyword càng cao
                            matched_pattern = f"keywords: {[kw for kw in tool_data['keywords'] if kw in query_lower]}"
                            analysis_source = "user_query_keywords"
                            print(f"✅ [Query Keywords] Tool: {tool_name} | Matched: {keyword_count}")
                            break
                    
                    if detected_tool:
                        break
            else:
                print(f"⚠️ [Query Skip] Skipped analysis (negative or question)")
        
        # Priority 2: Phân tích LLM RESPONSE (nếu query không có kết quả)
        if not detected_tool and llm_response:
            response_lower = llm_response.lower()
            
            # Kiểm tra phủ định và câu hỏi trong response
            is_negative = any(re.search(p, response_lower) for p in negative_patterns)
            is_question = any(re.search(p, response_lower) for p in question_patterns)
            
            print(f"🔍 [Analysis] Response Context:")
            print(f"   - Is Negative: {is_negative}")
            print(f"   - Is Question: {is_question}")
            
            if not is_negative and not is_question:
                for tool_name, tool_data in enhanced_vlc_patterns.items():
                    for pattern in tool_data["patterns"]:
                        if re.search(pattern, response_lower):
                            detected_tool = tool_name
                            confidence = 0.75  # Lower than query but still good
                            matched_pattern = pattern
                            analysis_source = "llm_response"
                            print(f"✅ [Response Match] Tool: {tool_name} | Pattern: {pattern}")
                            break
                    if detected_tool:
                        break
            else:
                print(f"⚠️ [Response Skip] Skipped analysis (negative or question)")
        
        # Priority 3: Intent Detector fallback (nếu cả 2 đều không có kết quả)
        if not detected_tool:
            print(f"🔍 [Fallback] Using Intent Detector...")
            try:
                text_to_analyze = original_query if original_query else llm_response
                intent_result = intent_detector.detect_intent(text_to_analyze)
                detected_tool = intent_result.get("suggested_tool")
                confidence = intent_result.get("confidence", 0.0) * 0.8  # Giảm 20% vì fallback
                matched_pattern = "intent_detector"
                analysis_source = "intent_detector"
                print(f"🔍 [Intent Detector] Tool: {detected_tool} | Confidence: {confidence:.2f}")
            except Exception as e:
                print(f"❌ [Intent Detector] Error: {e}")
        
        # ===== BƯỚC 4: TỰ ĐỘNG GỌI TOOL =====
        tool_executed = False
        tool_result = None
        
        print(f"\n📊 [Decision]")
        print(f"   - Tool Detected: {detected_tool}")
        print(f"   - Confidence: {confidence:.2f}")
        print(f"   - Source: {analysis_source}")
        print(f"   - Threshold: 0.5")
        
        if auto_execute and detected_tool and confidence >= 0.5:  # Giảm threshold xuống 0.5
            if detected_tool in TOOLS and TOOLS[detected_tool]["handler"]:
                print(f"🚀 [Execute] Calling tool: {detected_tool}")
                
                try:
                    handler = TOOLS[detected_tool]["handler"]
                    tool_args = {}
                    
                    # Extract arguments cho play_music
                    if detected_tool == "play_music" and original_query:
                        # Trích xuất tên bài hát
                        for kw in ["phát", "play", "bài", "song", "mở", "bật"]:
                            if kw in original_query.lower():
                                parts = original_query.lower().split(kw, 1)
                                if len(parts) > 1:
                                    filename = parts[1].strip()
                                    # Loại bỏ các từ thừa
                                    filename = re.sub(r'\b(cho tôi|giúp tôi|giúp mình|nhé|đi)\b', '', filename).strip()
                                    if filename:
                                        tool_args["filename"] = filename
                                        print(f"🎵 [Extract] Filename: '{filename}'")
                                    break
                    
                    # Gọi tool
                    tool_result = await handler(**tool_args)
                    tool_executed = True
                    
                    print(f"✅ [Execute] Success!")
                    print(f"📊 [Result] {str(tool_result)[:150]}...")
                    
                except Exception as e:
                    print(f"❌ [Execute] Error: {e}")
                    import traceback
                    traceback.print_exc()
                    tool_result = {"success": False, "error": str(e)}
            else:
                print(f"⚠️ [Execute] Tool '{detected_tool}' not found in registry")
        elif not auto_execute:
            print(f"ℹ️ [Execute] Skipped (auto_execute=False)")
        elif not detected_tool:
            print(f"⚠️ [Execute] Skipped (no tool detected)")
        elif confidence < 0.5:
            print(f"⚠️ [Execute] Skipped (confidence {confidence:.2f} < 0.5)")
        
        print(f"{'='*70}\n")
        
        # ===== BƯỚC 5: TRẢ VỀ KẾT QUẢ =====
        return {
            "success": True,
            "llm_response": llm_response,
            "original_query": original_query,
            "intent_detected": detected_tool or "unknown",
            "tool_suggested": detected_tool,
            "confidence": confidence,
            "tool_executed": tool_executed,
            "tool_result": tool_result,
            "analysis": {
                "source": analysis_source,
                "matched_pattern": matched_pattern,
                "is_negative": is_negative if 'is_negative' in locals() else False,
                "is_question": is_question if 'is_question' in locals() else False
            },
            "message": f"✅ Detected: {detected_tool} ({analysis_source}) | Executed: {tool_executed}" if detected_tool else "⚠️ No tool detected"
        }
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc()
        }


# ============================================================
# 🧠 SMART CONVERSATION ANALYZER v1.0
# Phân tích hội thoại thông minh & tự động điều khiển MỌI tool
# ============================================================

class SmartConversationAnalyzer:
    """
    🧠 SMART CONVERSATION ANALYZER
    
    Phân tích TOÀN BỘ lịch sử hội thoại để:
    1. Hiểu INTENT thực sự của user (không phụ thuộc từ khóa cứng)
    2. Phát hiện tool phù hợp nhất từ 50+ tools
    3. Extract arguments thông minh
    4. Tự động thực thi tool
    
    ĐẶC BIỆT:
    - Dùng AI (Gemini/GPT-4) để phân tích → HIỂU NGỮ CẢNH
    - Không cần regex patterns cho từng tool
    - Hỗ trợ TẤT CẢ tools (không chỉ VLC)
    - Context-aware: hiểu conversation history
    """
    
    def __init__(self):
        self.conversation_history = []  # Lưu lịch sử hội thoại
        self.max_history = 20  # Giữ 20 tin nhắn gần nhất
        self.last_executed_tool = None
        self.last_tool_result = None
        
        # Build tool catalog từ TOOLS dictionary
        self.tool_catalog = self._build_tool_catalog()
        
    def _build_tool_catalog(self) -> str:
        """Tạo catalog tools cho AI prompt"""
        catalog_lines = []
        for tool_name, tool_info in TOOLS.items():
            desc = tool_info.get("description", "")[:100]
            params = list(tool_info.get("parameters", {}).keys())
            params_str = ", ".join(params) if params else "none"
            catalog_lines.append(f"- {tool_name}: {desc}... | params: {params_str}")
        return "\n".join(catalog_lines)
    
    def add_message(self, role: str, content: str, tool_called: str = None):
        """Thêm message vào history"""
        self.conversation_history.append({
            "role": role,  # "user" hoặc "assistant" hoặc "system"
            "content": content,
            "tool_called": tool_called,
            "timestamp": datetime.now().isoformat()
        })
        # Giữ max history
        if len(self.conversation_history) > self.max_history:
            self.conversation_history = self.conversation_history[-self.max_history:]
    
    def get_conversation_context(self, last_n: int = 10) -> str:
        """Lấy context từ conversation history"""
        recent = self.conversation_history[-last_n:] if len(self.conversation_history) > last_n else self.conversation_history
        context_lines = []
        for msg in recent:
            role = "USER" if msg["role"] == "user" else "ASSISTANT"
            tool_info = f" [called: {msg['tool_called']}]" if msg.get("tool_called") else ""
            context_lines.append(f"{role}: {msg['content']}{tool_info}")
        return "\n".join(context_lines)
    
    async def analyze_with_ai(self, user_query: str, llm_response: str = "") -> dict:
        """
        Dùng AI để phân tích conversation và xác định tool cần gọi
        
        Returns:
            {
                "tool_name": str,           # Tool cần gọi
                "arguments": dict,          # Arguments cho tool
                "confidence": float,        # Độ tin cậy (0-1)
                "reasoning": str,           # Giải thích lý do
                "should_execute": bool      # Có nên thực thi không
            }
        """
        # Lấy conversation context
        context = self.get_conversation_context(last_n=5)
        
        # Build prompt cho AI
        analysis_prompt = f"""🧠 BẠN LÀ TOOL ANALYZER - Phân tích hội thoại và xác định TOOL cần gọi.

📋 DANH SÁCH TOOLS CÓ SẴN:
{self.tool_catalog}

📜 LỊCH SỬ HỘI THOẠI GẦN ĐÂY:
{context}

📝 YÊU CẦU HIỆN TẠI CỦA USER:
"{user_query}"

💬 LLM ĐÃ PHẢN HỒI (nếu có):
"{llm_response}"

🎯 NHIỆM VỤ: Phân tích và trả về JSON với format CHÍNH XÁC:
{{
    "tool_name": "tên_tool_cần_gọi hoặc null nếu không cần tool",
    "arguments": {{"param1": "value1", "param2": "value2"}} hoặc {{}},
    "confidence": 0.0 đến 1.0,
    "reasoning": "giải thích ngắn gọn lý do chọn tool này",
    "should_execute": true hoặc false
}}

🚨 LƯU Ý QUAN TRỌNG:
1. NẾU user hỏi câu hỏi chung (thời tiết, tin tức...) → KHÔNG cần tool → tool_name: null
2. NẾU user yêu cầu hành động CỤ THỂ → tìm tool phù hợp
3. NẾU LLM đã nói "đã chuyển bài", "đã tạm dừng" nhưng KHÔNG gọi tool → cần gọi tool
4. Confidence < 0.6 → should_execute: false
5. CHỈ trả về JSON, không có text khác

VÍ DỤ:
- User: "phát nhạc" → {{"tool_name": "play_music", "arguments": {{}}, "confidence": 0.95, "reasoning": "user muốn phát nhạc", "should_execute": true}}
- User: "bài tiếp theo" → {{"tool_name": "music_next", "arguments": {{}}, "confidence": 0.95, "reasoning": "user muốn chuyển bài", "should_execute": true}}
- User: "mở chrome" → {{"tool_name": "open_application", "arguments": {{"app_name": "chrome"}}, "confidence": 0.95, "reasoning": "mở trình duyệt", "should_execute": true}}
- User: "hôm nay thời tiết thế nào?" → {{"tool_name": null, "arguments": {{}}, "confidence": 0.0, "reasoning": "câu hỏi thông thường, không cần tool", "should_execute": false}}

TRẢ VỀ JSON:"""

        try:
            # Thử dùng Gemini trước
            if GEMINI_AVAILABLE and hasattr(genai, '_client') or os.getenv("GEMINI_API_KEY"):
                try:
                    api_key = os.getenv("GEMINI_API_KEY", "")
                    if api_key:
                        genai.configure(api_key=api_key)
                        model = genai.GenerativeModel('models/gemini-3-flash-preview')
                        response = model.generate_content(analysis_prompt)
                        ai_result = response.text.strip()
                        print(f"🤖 [AI Analysis] Gemini response: {ai_result[:200]}...")
                        return self._parse_ai_response(ai_result)
                except Exception as e:
                    print(f"⚠️ [AI Analysis] Gemini error: {e}")
            
            # Fallback: dùng OpenAI
            if OPENAI_AVAILABLE:
                try:
                    api_key = os.getenv("OPENAI_API_KEY", "")
                    if api_key:
                        client = OpenAI(api_key=api_key)
                        response = client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=[{"role": "user", "content": analysis_prompt}],
                            temperature=0.1,
                            max_tokens=500
                        )
                        ai_result = response.choices[0].message.content.strip()
                        print(f"🤖 [AI Analysis] GPT-4 response: {ai_result[:200]}...")
                        return self._parse_ai_response(ai_result)
                except Exception as e:
                    print(f"⚠️ [AI Analysis] OpenAI error: {e}")
            
            # Fallback cuối: dùng rule-based
            print("⚠️ [AI Analysis] No AI available, using rule-based analysis")
            return await self._rule_based_analysis(user_query, llm_response)
            
        except Exception as e:
            print(f"❌ [AI Analysis] Error: {e}")
            return await self._rule_based_analysis(user_query, llm_response)
    
    def _parse_ai_response(self, ai_text: str) -> dict:
        """Parse JSON từ AI response"""
        try:
            # Tìm JSON trong response
            import json
            
            # Thử parse trực tiếp
            try:
                return json.loads(ai_text)
            except:
                pass
            
            # Tìm JSON block
            json_match = re.search(r'\{[\s\S]*\}', ai_text)
            if json_match:
                return json.loads(json_match.group())
            
            # Không tìm được JSON
            return {
                "tool_name": None,
                "arguments": {},
                "confidence": 0.0,
                "reasoning": "Could not parse AI response",
                "should_execute": False
            }
        except Exception as e:
            print(f"❌ [Parse] Error: {e}")
            return {
                "tool_name": None,
                "arguments": {},
                "confidence": 0.0,
                "reasoning": f"Parse error: {e}",
                "should_execute": False
            }
    
    async def _rule_based_analysis(self, user_query: str, llm_response: str) -> dict:
        """Fallback: phân tích bằng rules khi không có AI"""
        query_lower = user_query.lower() if user_query else ""
        response_lower = llm_response.lower() if llm_response else ""
        combined = (query_lower + " " + response_lower).strip()
        
        print(f"🔍 [Rule-Based] Analyzing: '{combined}'")
        
        # Extended patterns cho TẤT CẢ tools (HỖ TRỢ TIẾNG VIỆT KHÔNG DẤU)
        all_tool_patterns = {
            # === MUSIC CONTROLS ===
            "music_next": {
                "patterns": [
                    r"bài tiếp|bai tiep|next|skip|chuyển bài|chuyen bai",
                    r"bài sau|bai sau|bài kế|bai ke|sang bài|sang bai",
                    r"tiep theo|tiếp theo|ke tiep|kế tiếp"
                ],
                "keywords": ["next", "tiếp", "tiep", "skip", "chuyển", "chuyen", "sau", "kế", "ke"]
            },
            "music_previous": {
                "patterns": [
                    r"bài trước|bai truoc|previous|quay lại|quay lai",
                    r"back|lùi|lui|trở lại|tro lai|bai cu|bài cũ"
                ],
                "keywords": ["previous", "trước", "truoc", "back", "quay", "lùi", "lui"]
            },
            "pause_music": {
                "patterns": [r"tạm dừng|tam dung|pause|dừng lại|dung lai|ngưng|ngung"],
                "keywords": ["pause", "tạm", "tam", "dừng", "dung"]
            },
            "resume_music": {
                "patterns": [r"tiếp tục|tiep tuc|resume|continue|phát tiếp|phat tiep|chạy tiếp|chay tiep"],
                "keywords": ["resume", "tiếp tục", "tiep tuc", "continue"]
            },
            "stop_music": {
                "patterns": [r"dừng hẳn|dung han|stop|tắt nhạc|tat nhac|ngừng nhạc|ngung nhac"],
                "keywords": ["stop", "tắt", "tat", "dừng hẳn", "dung han"]
            },
            "play_music": {
                "patterns": [
                    r"phát nhạc|phat nhac|play music|bật nhạc|bat nhac",
                    r"mở nhạc|mo nhac|nghe nhạc|nghe nhac"
                ],
                "keywords": ["phát", "phat", "play", "bật", "bat", "mở", "mo", "nghe"]
            },
            
            # === VOLUME CONTROLS ===
            "volume_up": {
                "patterns": [r"tăng âm|tang am|volume up|to hơn|to hon|lớn hơn|lon hon"],
                "keywords": ["tăng", "tang", "up", "to hơn", "to hon", "lớn", "lon"]
            },
            "volume_down": {
                "patterns": [r"giảm âm|giam am|volume down|nhỏ hơn|nho hon|bớt to|bot to"],
                "keywords": ["giảm", "giam", "down", "nhỏ", "nho", "bớt", "bot"]
            },
            "mute_volume": {
                "patterns": [r"tắt tiếng|tat tieng|mute|câm|cam|im lặng|im lang"],
                "keywords": ["mute", "tắt tiếng", "tat tieng", "câm", "cam"]
            },
            "set_volume": {
                "patterns": [r"âm lượng \d+|am luong \d+|volume \d+|đặt âm|dat am|chỉnh âm|chinh am"],
                "keywords": ["âm lượng", "am luong", "volume"]
            },
            
            # === APPLICATIONS ===
            "open_application": {
                "patterns": [
                    r"mở ứng dụng|mo ung dung|open app|mở chrome|mo chrome",
                    r"mở word|mo word|mở excel|mo excel|mở notepad|mo notepad",
                    r"khởi động|khoi dong"
                ],
                "keywords": ["mở", "mo", "open", "khởi động", "khoi dong", "chạy", "chay"]
            },
            "kill_process": {
                "patterns": [r"tắt ứng dụng|tat ung dung|kill|đóng app|dong app|close app"],
                "keywords": ["tắt", "tat", "kill", "đóng", "dong", "close"]
            },
            
            # === SYSTEM ===
            "take_screenshot": {
                "patterns": [r"chụp màn hình|chup man hinh|screenshot|capture screen"],
                "keywords": ["chụp", "chup", "screenshot", "capture"]
            },
            "get_system_resources": {
                "patterns": [r"tài nguyên|tai nguyen|system info|cpu|ram|memory"],
                "keywords": ["tài nguyên", "tai nguyen", "system", "cpu", "ram"]
            },
            "get_current_time": {
                "patterns": [r"mấy giờ|may gio|thời gian|thoi gian|time now|giờ hiện tại|gio hien tai"],
                "keywords": ["giờ", "gio", "time", "thời gian", "thoi gian"]
            },
            
            # === FILES ===
            "create_file": {
                "patterns": [r"tạo file|tao file|create file|viết file|viet file"],
                "keywords": ["tạo file", "tao file", "create file", "viết", "viet"]
            },
            "read_file": {
                "patterns": [r"đọc file|doc file|read file|xem file"],
                "keywords": ["đọc", "doc", "read", "xem"]
            },
            "list_files": {
                "patterns": [r"liệt kê file|liet ke file|list files|xem thư mục|xem thu muc"],
                "keywords": ["liệt kê", "liet ke", "list", "thư mục", "thu muc"]
            },
            
            # === CALCULATOR ===
            "calculator": {
                "patterns": [r"tính|tinh|calculate|bao nhiêu|bao nhieu|\d+\s*[\+\-\*\/]\s*\d+"],
                "keywords": ["tính", "tinh", "calculate", "cộng", "cong", "trừ", "tru", "nhân", "nhan", "chia"]
            },
            
            # === CLIPBOARD ===
            "get_clipboard": {
                "patterns": [r"clipboard|đã copy gì|da copy gi|lấy clipboard|lay clipboard"],
                "keywords": ["clipboard", "copy"]
            },
            "set_clipboard": {
                "patterns": [r"copy vào clipboard|copy vao clipboard|set clipboard"],
                "keywords": ["copy vào", "copy vao", "set clipboard"]
            },
            
            # === BROWSER ===
            "search_web": {
                "patterns": [r"tìm kiếm google|tim kiem google|search google|mở google tìm|mo google tim"],
                "keywords": ["google", "search web", "tìm kiếm", "tim kiem"]
            },
            "open_youtube": {
                "patterns": [r"mở youtube|mo youtube|youtube|xem video"],
                "keywords": ["youtube", "video"]
            },
            
            # === BRIGHTNESS ===
            "set_brightness": {
                "patterns": [r"độ sáng|do sang|brightness|sáng hơn|sang hon|tối hơn|toi hon"],
                "keywords": ["sáng", "sang", "brightness", "tối", "toi"]
            }
        }
        
        # Tìm tool match nhất
        best_match = None
        best_confidence = 0.0
        best_reason = ""
        
        for tool_name, tool_patterns in all_tool_patterns.items():
            # Check patterns
            for pattern in tool_patterns["patterns"]:
                if re.search(pattern, combined):
                    confidence = 0.85
                    if confidence > best_confidence:
                        best_confidence = confidence
                        best_match = tool_name
                        best_reason = f"Pattern match: {pattern}"
                        print(f"✅ [Rule-Based] Pattern matched: {tool_name} ({pattern})")
                    break
            
            # ALWAYS check keywords (không chỉ khi chưa có match)
            keyword_count = sum(1 for kw in tool_patterns["keywords"] if kw in combined)
            if keyword_count >= 1:
                confidence = 0.6 + (keyword_count * 0.1)
                if confidence > best_confidence:
                    best_confidence = confidence
                    best_match = tool_name
                    best_reason = f"Keywords: {keyword_count} matches"
                    print(f"✅ [Rule-Based] Keywords matched: {tool_name} ({keyword_count} keywords)")
        
        print(f"📊 [Rule-Based] Result: {best_match} (confidence: {best_confidence:.2f})")
        
        # Extract arguments
        arguments = {}
        if best_match:
            # Dùng combined text để extract args nếu query trống
            text_for_args = user_query if user_query else llm_response
            arguments = self._extract_arguments(best_match, text_for_args)
        
        return {
            "tool_name": best_match,
            "arguments": arguments,
            "confidence": best_confidence,
            "reasoning": best_reason,
            "should_execute": best_confidence >= 0.5
        }
    
    def _extract_arguments(self, tool_name: str, query: str) -> dict:
        """Extract arguments cho tool từ query"""
        args = {}
        query_lower = query.lower()
        
        # play_music → extract filename
        if tool_name == "play_music":
            for kw in ["phát", "play", "bài", "song", "mở", "bật", "nghe"]:
                if kw in query_lower:
                    parts = query_lower.split(kw, 1)
                    if len(parts) > 1:
                        filename = parts[1].strip()
                        filename = re.sub(r'\b(cho tôi|giúp tôi|nhé|đi|nào)\b', '', filename).strip()
                        if filename and len(filename) > 1:
                            args["filename"] = filename
                        break
        
        # open_application → extract app_name
        elif tool_name == "open_application":
            for kw in ["mở", "open", "khởi động", "chạy"]:
                if kw in query_lower:
                    parts = query_lower.split(kw, 1)
                    if len(parts) > 1:
                        app = parts[1].strip()
                        app = re.sub(r'\b(cho tôi|giúp|nhé|đi|ứng dụng|app)\b', '', app).strip()
                        if app:
                            args["app_name"] = app
                        break
        
        # set_volume → extract level
        elif tool_name == "set_volume":
            match = re.search(r'(\d+)\s*(%)?', query)
            if match:
                level = int(match.group(1))
                args["level"] = min(100, max(0, level))
        
        # calculator → extract expression
        elif tool_name == "calculator":
            # Tìm biểu thức toán
            expr_match = re.search(r'(\d+[\s\+\-\*\/\(\)]+\d+[\s\d\+\-\*\/\(\)]*)', query)
            if expr_match:
                args["expression"] = expr_match.group(1).strip()
        
        # set_brightness → extract level
        elif tool_name == "set_brightness":
            match = re.search(r'(\d+)\s*(%)?', query)
            if match:
                level = int(match.group(1))
                args["level"] = min(100, max(0, level))
        
        # search_web → extract query
        elif tool_name == "search_web":
            for kw in ["tìm", "search", "google"]:
                if kw in query_lower:
                    parts = query_lower.split(kw, 1)
                    if len(parts) > 1:
                        search_query = parts[1].strip()
                        search_query = re.sub(r'\b(về|cho tôi|giúp|trên)\b', '', search_query).strip()
                        if search_query:
                            args["query"] = search_query
                        break
        
        return args
    
    async def execute_tool(self, tool_name: str, arguments: dict) -> dict:
        """Thực thi tool với arguments"""
        try:
            if tool_name not in TOOLS:
                return {"success": False, "error": f"Tool '{tool_name}' not found"}
            
            handler = TOOLS[tool_name]["handler"]
            if not handler:
                return {"success": False, "error": f"Tool '{tool_name}' has no handler"}
            
            # Gọi tool
            result = await handler(**arguments)
            
            # Lưu lại
            self.last_executed_tool = tool_name
            self.last_tool_result = result
            
            return {"success": True, "tool": tool_name, "result": result}
            
        except Exception as e:
            import traceback
            return {"success": False, "error": str(e), "traceback": traceback.format_exc()}


# Global instance
smart_analyzer = SmartConversationAnalyzer()


@app.post("/api/smart_analyze")
async def api_smart_analyze(data: dict):
    """
    🧠 SMART CONVERSATION ANALYZER API
    
    Phân tích hội thoại thông minh, tự động điều khiển MỌI tool.
    Không phụ thuộc từ khóa cứng - dùng AI để hiểu ngữ cảnh.
    
    Args:
        user_query: Yêu cầu của user
        llm_response: Phản hồi từ LLM (optional)
        conversation_history: Lịch sử hội thoại (optional, list of {role, content})
        auto_execute: Tự động thực thi tool (default: True)
        use_ai: Dùng AI để phân tích (default: True, fallback to rules)
    
    Returns:
        {
            "success": bool,
            "analysis": {
                "tool_name": str,
                "arguments": dict,
                "confidence": float,
                "reasoning": str,
                "should_execute": bool
            },
            "execution": {
                "executed": bool,
                "result": dict
            },
            "message": str
        }
    """
    try:
        user_query = data.get("user_query", data.get("query", "")).strip()
        llm_response = data.get("llm_response", data.get("response", "")).strip()
        conversation_history = data.get("conversation_history", [])
        auto_execute = data.get("auto_execute", True)
        use_ai = data.get("use_ai", True)
        
        print(f"\n{'='*70}")
        print(f"🧠 [Smart Analyze] NEW REQUEST")
        print(f"{'='*70}")
        print(f"📝 User Query: '{user_query}'")
        print(f"💬 LLM Response: '{llm_response[:100]}...' " if llm_response else "")
        print(f"⚙️  Auto Execute: {auto_execute} | Use AI: {use_ai}")
        print(f"📜 History Length: {len(conversation_history)}")
        print(f"{'-'*70}")
        
        if not user_query and not llm_response:
            return {
                "success": False,
                "error": "user_query or llm_response is required"
            }
        
        # Thêm conversation history nếu có
        for msg in conversation_history:
            smart_analyzer.add_message(
                role=msg.get("role", "user"),
                content=msg.get("content", "")
            )
        
        # Thêm message hiện tại
        if user_query:
            smart_analyzer.add_message("user", user_query)
        if llm_response:
            smart_analyzer.add_message("assistant", llm_response)
        
        # === PHÂN TÍCH ===
        if use_ai:
            analysis = await smart_analyzer.analyze_with_ai(user_query, llm_response)
        else:
            analysis = await smart_analyzer._rule_based_analysis(user_query, llm_response)
        
        print(f"\n🎯 [Analysis Result]")
        print(f"   - Tool: {analysis.get('tool_name')}")
        print(f"   - Arguments: {analysis.get('arguments')}")
        print(f"   - Confidence: {analysis.get('confidence', 0):.2f}")
        print(f"   - Should Execute: {analysis.get('should_execute')}")
        print(f"   - Reasoning: {analysis.get('reasoning')}")
        
        # === THỰC THI ===
        execution = {"executed": False, "result": None}
        
        if auto_execute and analysis.get("should_execute") and analysis.get("tool_name"):
            tool_name = analysis["tool_name"]
            arguments = analysis.get("arguments", {})
            
            print(f"\n🚀 [Execute] Calling: {tool_name}({arguments})")
            
            exec_result = await smart_analyzer.execute_tool(tool_name, arguments)
            execution = {
                "executed": exec_result.get("success", False),
                "result": exec_result
            }
            
            # Cập nhật history với tool đã gọi
            smart_analyzer.add_message("system", f"Tool executed: {tool_name}", tool_called=tool_name)
            
            if exec_result.get("success"):
                print(f"✅ [Execute] Success!")
            else:
                print(f"❌ [Execute] Failed: {exec_result.get('error')}")
        
        print(f"{'='*70}\n")
        
        return {
            "success": True,
            "user_query": user_query,
            "llm_response": llm_response,
            "analysis": analysis,
            "execution": execution,
            "message": f"✅ Tool: {analysis.get('tool_name')} | Executed: {execution['executed']}" if analysis.get('tool_name') else "⚠️ No tool needed"
        }
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc()
        }


# NOTE: Đã xóa duplicate endpoint /api/conversation/add (line 13967)
# Endpoint chính nằm ở phần CONVERSATION HISTORY API (line ~15208)
# Giữ lại để tránh conflict với SmartConversationAnalyzer


@app.post("/api/smart_chat")
async def api_smart_chat(data: dict):
    """
    Smart Chat với Intent Detection tự động + VLC MCP Integration + Google Search Grounding
    1. Phân tích intent
    2. Nếu cần tool → tự động gọi tool (REST) hoặc MCP (VLC)
    3. Gửi kết quả tool + query đến Gemini
    4. Trả về response hoàn chỉnh
    
    🆕 VLC MCP: Tự động dùng MCP protocol cho VLC commands
    🆕 Google Search: Tự động tra cứu Google cho câu hỏi realtime
    """
    query = data.get("query", data.get("prompt", data.get("text", "")))
    use_llm_intent = data.get("use_llm_intent", False)
    model = data.get("model", "gemini-2.0-flash")  # Default model hỗ trợ grounding
    use_google_search = data.get("use_google_search", True)  # 🆕 Mặc định BẬT Google Search
    
    if not query:
        raise HTTPException(400, "Query is required")
    
    try:
        # 🆕 STEP -1: Kiểm tra có cần Google Search không (câu hỏi thời sự, giá cả, tin tức)
        realtime_keywords = [
            'giá vàng', 'giá usd', 'tỷ giá', 'giá bitcoin', 'crypto', 'chứng khoán',
            'thời tiết', 'weather', 'tin tức', 'news', 'mới nhất', 'latest',
            'hôm nay', 'bây giờ', 'hiện nay', 'hiện tại', 'today', 'now', 'current',
            'năm 2024', 'năm 2025', 'năm 2026', '2024', '2025', '2026',
            'vô địch', 'champion', 'winner', 'kết quả', 'score', 'result',
            'tổng thống', 'president', 'thủ tướng', 'chủ tịch', 'ceo',
            'iphone', 'samsung', 'tesla', 'apple', 'google', 'microsoft', 'ra mắt',
            'là ai', 'là gì', 'ở đâu', 'what is', 'where is', 'how much', 'bao nhiêu',
            'sự kiện', 'event', 'lịch', 'schedule', 'khi nào', 'when', 'giá xăng', 'giá dầu',
            'covid', 'bão', 'động đất', 'tai nạn', 'cháy', 'chiến tranh', 'xung đột'
        ]
        query_lower = query.lower()
        needs_google_search = use_google_search and any(kw in query_lower for kw in realtime_keywords)
        
        # 🔍 Nếu cần Google Search, ưu tiên dùng Gemini + Google Search Grounding
        if needs_google_search:
            print(f"🔍 [Smart Chat] Phát hiện câu hỏi cần Google Search: {query[:50]}...")
            try:
                google_result = await ask_gemini_with_google_search(
                    prompt=query,
                    model="gemini-2.0-flash"  # Model hỗ trợ grounding tốt nhất
                )
                
                if google_result.get("success"):
                    # Lưu vào conversation history
                    add_to_conversation(role="user", content=query, metadata={"source": "smart_chat_google_search"})
                    add_to_conversation(
                        role="assistant", 
                        content=google_result.get("response", ""),
                        metadata={
                            "source": "smart_chat_google_search",
                            "model": google_result.get("model"),
                            "google_search_used": True,
                            "search_queries": google_result.get("search_queries", [])
                        }
                    )
                    
                    return {
                        "success": True,
                        "query": query,
                        "response": google_result.get("response"),
                        "intent": {"intent": "realtime_query", "needs_google_search": True},
                        "tool_used": "google_search_grounding",
                        "google_search_used": True,
                        "search_queries": google_result.get("search_queries", []),
                        "grounding_chunks": google_result.get("grounding_chunks", []),
                        "model": google_result.get("model"),
                        "message": google_result.get("message")
                    }
                else:
                    print(f"⚠️ [Smart Chat] Google Search failed, falling back to normal...")
            except Exception as e:
                print(f"⚠️ [Smart Chat] Google Search error: {e}, falling back...")
        
        # 🆕 STEP 0: Tự động phát hiện và xử lý documents/database với Gemini
        doc_result = await auto_process_document_with_gemini(query, model=model)
        
        if doc_result.get("activated") and doc_result.get("success"):
            # Đã xử lý document thành công với Gemini
            print(f"📚 [Auto Document] Success! Documents: {len(doc_result.get('documents_found', []))}")
            
            return {
                "success": True,
                "query": query,
                "response": doc_result.get("gemini_response"),
                "intent": "document_query",
                "tool_used": "auto_process_document_with_gemini",
                "documents_found": doc_result.get("documents_found", []),
                "model": doc_result.get("model_used"),
                "message": doc_result.get("message"),
                "auto_document_processing": True
            }
        
        # Step 1: Detect intent
        if use_llm_intent:
            intent_result = await intent_detector.detect_with_llm(query, GEMINI_API_KEY)
        else:
            intent_result = intent_detector.detect_intent(query)
        
        print(f"🧠 [Intent] {intent_result}")
        
        tool_result = None
        tool_used = None
        mcp_used = False
        
        # Step 2: Nếu cần force tool, gọi tool trước
        if intent_result.get("should_force_tool") and intent_result.get("suggested_tool"):
            tool_name = intent_result["suggested_tool"]
            
            # 🆕 CHECK: Nếu là VLC command → dùng MCP
            vlc_commands = ["music_next", "music_previous", "pause_music", "resume_music", "stop_music", "play_music"]
            
            if VLC_MCP_AVAILABLE and tool_name in vlc_commands:
                print(f"🎯 [VLC MCP] Routing to MCP: {tool_name}")
                
                # Map tool name to MCP tool name
                mcp_tool_map = {
                    "music_next": "vlc.next",
                    "music_previous": "vlc.previous",
                    "pause_music": "vlc.pause",
                    "resume_music": "vlc.play",
                    "stop_music": "vlc.stop",
                    "play_music": "vlc.play"
                }
                
                mcp_tool_name = mcp_tool_map.get(tool_name)
                
                if mcp_tool_name:
                    try:
                        # Call via MCP protocol
                        mcp_request = {
                            "jsonrpc": "2.0",
                            "method": "tools/call",
                            "params": {
                                "name": mcp_tool_name,
                                "arguments": {}
                            },
                            "id": 1
                        }
                        
                        mcp_response = await vlc_mcp_server.handle_mcp_request(mcp_request)
                        
                        if "result" in mcp_response:
                            tool_result = mcp_response["result"]
                            tool_used = mcp_tool_name
                            mcp_used = True
                            print(f"✅ [VLC MCP] Success: {mcp_tool_name}")
                        else:
                            print(f"❌ [VLC MCP] Error: {mcp_response.get('error')}")
                            tool_result = {"error": mcp_response.get("error", {}).get("message", "Unknown error")}
                    except Exception as e:
                        print(f"⚠️ [VLC MCP] Exception: {e}")
                        tool_result = {"error": str(e)}
            
            # Fallback: REST API
            elif tool_name in TOOLS and TOOLS[tool_name]["handler"]:
                print(f"🔧 [Auto Tool] Calling {tool_name} for query: {query}")
                
                try:
                    # Tạo arguments dựa trên intent
                    tool_args = {"query": query}
                    
                    # Gọi tool
                    handler = TOOLS[tool_name]["handler"]
                    tool_result = await handler(**tool_args)
                    tool_used = tool_name
                    
                    print(f"✅ [Auto Tool] {tool_name} result: {str(tool_result)[:200]}...")
                except Exception as e:
                    print(f"⚠️ [Auto Tool] Error calling {tool_name}: {e}")
                    tool_result = {"error": str(e)}
        
        # Step 3: Gửi đến Gemini với context từ tool
        final_prompt = query
        if tool_result and not tool_result.get("error"):
            # Thêm context từ tool result
            context = json.dumps(tool_result, ensure_ascii=False, indent=2)
            final_prompt = f"""Dựa trên thông tin tra cứu sau đây, hãy trả lời câu hỏi của user.

📊 THÔNG TIN TRA CỨU (từ {tool_used}):
{context}

❓ CÂU HỎI CỦA USER:
{query}

📝 YÊU CẦU:
- Trả lời ngắn gọn, chính xác
- Dựa trên thông tin tra cứu ở trên
- Nếu thông tin không đủ, nói rõ và đưa ra những gì có"""
        
        # Gọi Gemini
        gemini_result = await ask_gemini(prompt=final_prompt, model=model)
        
        # Lưu vào conversation history
        add_to_conversation(
            role="user",
            content=query,
            metadata={
                "source": "smart_chat",
                "intent": intent_result.get("intent"),
                "tool_suggested": intent_result.get("suggested_tool")
            }
        )
        
        if gemini_result.get("success"):
            add_to_conversation(
                role="assistant",
                content=gemini_result.get("response", ""),
                metadata={
                    "source": "smart_chat",
                    "tool_used": tool_used,
                    "model": model
                }
            )
        
        return {
            "success": True,
            "query": query,
            "intent": intent_result,
            "tool_used": tool_used,
            "tool_result": tool_result,
            "response": gemini_result.get("response", ""),
            "model": model
        }
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}


# ===== 🔍 GOOGLE SEARCH GROUNDING ENDPOINT =====

@app.post("/api/gemini/google_search")
async def api_gemini_google_search(data: dict):
    """
    🔍 Gemini với Google Search Grounding - Tra cứu Google tự động
    
    Tính năng cho phép Gemini tự động tìm kiếm Google để trả lời
    các câu hỏi cần thông tin mới nhất, real-time.
    
    Args (JSON body):
        prompt (str): Câu hỏi cần Gemini tra cứu và trả lời
        model (str, optional): Model Gemini (default: gemini-2.0-flash)
        
    Returns:
        success: True/False
        response: Câu trả lời từ Gemini
        google_search_used: True nếu đã dùng Google Search
        search_queries: Các query đã search trên Google
        grounding_chunks: Nguồn website được trích dẫn
    
    Example:
        POST /api/gemini/google_search
        {"prompt": "Giá vàng hôm nay là bao nhiêu?"}
    """
    prompt = data.get("prompt", data.get("query", data.get("text", "")))
    model = data.get("model", "gemini-2.0-flash")
    
    if not prompt:
        raise HTTPException(400, "Prompt is required")
    
    print(f"🔍 [API Google Search] Query: {prompt[:100]}...")
    
    # Lưu user message vào history
    add_to_conversation(
        role="user",
        content=prompt,
        metadata={"source": "google_search_api", "model": model}
    )
    
    try:
        # Gọi Gemini với Google Search Grounding
        result = await ask_gemini_with_google_search(
            prompt=prompt,
            model=model
        )
        
        if result.get("success"):
            # Lưu assistant response vào history
            add_to_conversation(
                role="assistant",
                content=result.get("response", ""),
                metadata={
                    "source": "google_search_api",
                    "model": result.get("model"),
                    "google_search_used": result.get("google_search_used", False),
                    "search_queries": result.get("search_queries", [])
                }
            )
            
            return {
                "success": True,
                "prompt": prompt,
                "response": result.get("response"),
                "response_text": result.get("response"),  # Alias
                "model": result.get("model"),
                "google_search_used": result.get("google_search_used", False),
                "search_queries": result.get("search_queries", []),
                "grounding_chunks": result.get("grounding_chunks", []),
                "message": result.get("message")
            }
        else:
            return {
                "success": False,
                "error": result.get("error", "Unknown error"),
                "prompt": prompt
            }
            
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": str(e),
            "prompt": prompt
        }


# ===== 23 API ENDPOINTS MỚI (Tool 8-30) =====

@app.post("/api/tool/ask_gemini")
async def api_ask_gemini(data: dict):
    """
    Gemini AI endpoint with Knowledge Base + Google Search integration
    
    Flow:
    1. Nhận query từ user
    2. 🆕 Kiểm tra có cần Google Search không (giá cả, tin tức, thời sự)
    3. Nếu cần realtime → dùng Google Search Grounding
    4. Nếu không → search Knowledge Base + Gemini
    5. Trả về response
    """
    prompt = data.get("prompt", "")
    model = data.get("model", "gemini-2.0-flash")  # 🆕 Default model hỗ trợ grounding
    use_google_search = data.get("use_google_search", True)  # 🆕 Mặc định BẬT
    # 🔒 BẮT BUỘC search KB - KHÔNG cho user tắt
    use_knowledge_base = True  # LUÔN BẬT
    
    if not prompt:
        raise HTTPException(400, "Prompt is required")
    
    # Lưu user message vào history
    add_to_conversation(
        role="user",
        content=prompt,
        metadata={
            "source": "web_ui",
            "model_requested": model,
            "ai_provider": "gemini"
        }
    )
    
    # 🆕 STEP 0: Kiểm tra có cần Google Search không
    realtime_keywords = [
        # Giá cả, tài chính
        'giá vàng', 'giá usd', 'tỷ giá', 'giá bitcoin', 'crypto', 'chứng khoán',
        'gold price', 'exchange rate', 'giá xăng', 'giá dầu', 'giá cao nhất', 'giá mới nhất',
        'stock', 'bitcoin', 'ethereum', 'btc', 'eth',
        
        # Thời tiết
        'thời tiết', 'weather', 'nhiệt độ', 'temperature', 'mưa', 'bão',
        
        # Tin tức, sự kiện
        'tin tức', 'news', 'mới nhất', 'latest', 'breaking', 'sự kiện',
        
        # Thời gian thực
        'hôm nay', 'bây giờ', 'hiện nay', 'hiện tại', 'today', 'now', 'current',
        'năm 2024', 'năm 2025', 'năm 2026', '2024', '2025', '2026',
        
        # Thể thao, cuộc thi
        'vô địch', 'champion', 'winner', 'kết quả', 'score', 'result',
        'world cup', 'euro', 'sea games', 'olympic', 'bóng đá', 'football',
        
        # Người nổi tiếng, chính trị
        'tổng thống', 'president', 'thủ tướng', 'chủ tịch', 'ceo',
        'ai là', 'who is', 'who won',
        
        # Sản phẩm, công nghệ
        'iphone', 'samsung', 'tesla', 'apple', 'google', 'microsoft',
        'ra mắt', 'launch', 'release', 'announced',
        
        # Tra cứu chung cần thông tin mới
        'là ai', 'là gì', 'ở đâu', 'what is', 'where is', 'how much', 'bao nhiêu',
        'khi nào', 'when', 'how many'
    ]
    prompt_lower = prompt.lower()
    needs_google_search = use_google_search and any(kw in prompt_lower for kw in realtime_keywords)
    
    # 🔍 Nếu cần Google Search, ưu tiên dùng Gemini + Google Search Grounding
    if needs_google_search:
        print(f"🔍 [ask_gemini] Phát hiện câu hỏi cần Google Search: {prompt[:50]}...")
        try:
            google_result = await ask_gemini_with_google_search(
                prompt=prompt,
                model="gemini-2.0-flash"  # Model hỗ trợ grounding tốt nhất
            )
            
            if google_result.get("success"):
                response_text = google_result.get("response", "")
                
                # Lưu vào conversation history
                add_to_conversation(
                    role="assistant",
                    content=response_text,
                    metadata={
                        "source": "web_ui_google_search",
                        "model": google_result.get("model"),
                        "google_search_used": True,
                        "search_queries": google_result.get("search_queries", [])
                    }
                )
                
                return {
                    "success": True,
                    "prompt": prompt,
                    "response": response_text,
                    "response_text": response_text,
                    "model": google_result.get("model"),
                    "google_search_used": True,
                    "search_queries": google_result.get("search_queries", []),
                    "grounding_chunks": google_result.get("grounding_chunks", []),
                    "message": f"✅ Gemini đã tra cứu Google và trả lời (model: {google_result.get('model')})"
                }
            else:
                print(f"⚠️ [ask_gemini] Google Search failed: {google_result.get('error')}, falling back to KB...")
        except Exception as e:
            print(f"⚠️ [ask_gemini] Google Search error: {e}, falling back to KB...")
    
    # 🆕 AUTO-READ ALL KNOWLEDGE BASE (BẮT BUỘC) - Fallback nếu Google Search không dùng/fail
    enhanced_prompt = prompt
    kb_context_used = False
    
    if use_knowledge_base:  # Luôn = True
        try:
            # ĐỌC TOÀN BỘ Knowledge Base - KHÔNG filter theo query
            kb_result = await get_knowledge_context(
                query="",  # ĐỂ TRỐNG để lấy TẤT CẢ documents
                max_chars=50000,  # Tăng giới hạn để đọc nhiều hơn
                use_gemini_summary=True  # Bật Gemini tóm tắt
            )
            
            if kb_result.get("success") and kb_result.get("context"):
                kb_context = kb_result["context"]
                docs_count = kb_result.get("documents_included", 0)
                
                # Thêm context vào prompt
                enhanced_prompt = f"""📚 KNOWLEDGE BASE - TOÀN BỘ CƠ SỞ DỮ LIỆU ({docs_count} tài liệu):
{kb_context}

{'='*60}
❓ CÂU HỎI CỦA USER:
{prompt}

{'='*60}
💡 HƯỚNG DẪN TRẢ LỜI:
- Bạn đã có TOÀN BỘ nội dung Knowledge Base ở trên
- Phân tích và tóm tắt thông tin liên quan đến câu hỏi
- Trả lời DỰA TRÊN dữ liệu có sẵn, KHÔNG đoán mò
- Trích dẫn nguồn cụ thể (tên file, phần nội dung)
- Nếu không tìm thấy thông tin, hãy nói rõ "Không có trong cơ sở dữ liệu"
"""
                kb_context_used = True
                print(f"✅ [KB] Loaded ALL Knowledge Base: {docs_count} documents, {len(kb_context)} chars")
            else:
                print(f"⚠️ [KB] Knowledge Base is empty or not indexed yet")
        except Exception as e:
            print(f"⚠️ [KB] Error getting context: {e}")
            # Không có context, dùng prompt gốc
    
    # Gọi Gemini với enhanced prompt
    result = await ask_gemini(prompt=enhanced_prompt, model=model)
    
    # Thêm metadata về KB usage
    if kb_context_used and result.get("success"):
        result["knowledge_base_used"] = True
        result["message"] = result.get("response", "") + "\n\n📚 *Trả lời dựa trên Knowledge Base của bạn*"
    
    # Lưu AI response vào history
    if result.get("success"):
        add_to_conversation(
            role="assistant",
            content=result.get("response", ""),
            metadata={
                "source": "web_ui",
                "model": model,
                "ai_provider": "gemini",
                "knowledge_base_used": kb_context_used,
                "token_count": result.get("token_count", 0) if "token_count" in result else None
            }
        )
    
    return result


# ===== TTS (Text-to-Speech) API =====
# Global variable để track trạng thái TTS
tts_is_playing = False
tts_stop_requested = False

@app.post("/api/tts")
async def api_text_to_speech(data: dict):
    """
    API đọc to văn bản - Ưu tiên Gemini TTS, fallback to gTTS/SAPI
    ⚡ FAST MODE: Chỉ đọc 500 ký tự đầu để response nhanh
    """
    global tts_is_playing, tts_stop_requested
    
    print(f"🔊 [TTS API] Received request")
    
    text = data.get("text", "")
    if not text:
        print("❌ [TTS API] No text provided")
        return {"success": False, "error": "Không có văn bản để đọc"}
    
    # ⚡ FAST MODE: Giới hạn 500 ký tự để TTS nhanh (real-time feel)
    max_chars = 500
    original_length = len(text)
    if len(text) > max_chars:
        # Cắt tại dấu câu gần nhất để không bị cắt giữa từ
        cut_text = text[:max_chars]
        last_sentence = max(
            cut_text.rfind('.'),
            cut_text.rfind('!'),
            cut_text.rfind('?'),
            cut_text.rfind('。')
        )
        if last_sentence > max_chars // 2:
            text = text[:last_sentence + 1]
        else:
            text = cut_text
        print(f"🔊 [TTS API] Truncated from {original_length} to {len(text)} chars for fast response")
    
    # Loại bỏ markdown formatting
    text = clean_markdown_for_tts(text)
    
    tts_is_playing = True
    tts_stop_requested = False
    
    try:
        # Ưu tiên Gemini TTS (chất lượng cao)
        print(f"🎙️ [TTS API] Trying Gemini TTS ({len(text)} chars)...")
        voice = data.get("voice", "Aoede")  # Default female voice
        result = await gemini_text_to_speech(text, voice=voice, save_audio=False)
        
        if result.get("success"):
            print(f"✅ [TTS API] Gemini TTS success!")
            tts_is_playing = False
            return result
        
        # Fallback to gTTS/SAPI
        print(f"⚠️ [TTS API] Gemini TTS failed, falling back to gTTS/SAPI...")
        result = await text_to_speech(text, save_audio=False)
        print(f"🔊 [TTS API] Result: {result}")
        tts_is_playing = False
        return result
    except Exception as e:
        print(f"❌ [TTS API] Error: {e}")
        import traceback
        traceback.print_exc()
        tts_is_playing = False
        return {"success": False, "error": str(e)}


@app.post("/api/tts/stop")
async def api_tts_stop():
    """
    Dừng TTS đang phát
    """
    global tts_is_playing, tts_stop_requested
    
    tts_stop_requested = True
    
    try:
        import pygame
        if pygame.mixer.get_init():
            pygame.mixer.music.stop()
            pygame.mixer.quit()
        tts_is_playing = False
        return {"success": True, "message": "Đã dừng TTS"}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/tts/status")
async def api_tts_status():
    """
    Kiểm tra trạng thái TTS
    """
    global tts_is_playing
    return {"is_playing": tts_is_playing}


@app.post("/api/tool/open_application")
async def api_open_app(data: dict):
    result = await open_application(data.get("app_name", ""))
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

# MEDIA PLAYER CONTROL ENDPOINTS
@app.post("/api/tool/media_play_pause")
async def api_media_play_pause(data: dict):
    result = await media_play_pause()
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/media_next_track")
async def api_media_next(data: dict):
    result = await media_next_track()
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/media_previous_track")
async def api_media_previous(data: dict):
    result = await media_previous_track()
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/media_stop")
async def api_media_stop(data: dict):
    result = await media_stop()
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/media_control")
async def api_media_control(data: dict):
    result = await media_control(data.get("action", ""))
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/get_active_media_players")
async def api_get_active_media(data: dict):
    result = await get_active_media_players()
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/list_running_processes")
async def api_list_procs(data: dict):
    result = await list_running_processes(data.get("limit", 10))
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/kill_process")
async def api_kill_proc(data: dict):
    result = await kill_process(data.get("identifier", ""))
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/create_file")
async def api_create_file(data: dict):
    result = await create_file(data.get("path", ""), data.get("content", ""))
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/read_file")
async def api_read_file(data: dict):
    result = await read_file(data.get("path", ""))
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/list_files")
async def api_list_files(data: dict):
    result = await list_files(data.get("directory", ""))
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/get_disk_usage")
async def api_disk_usage():
    result = await get_disk_usage()
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/get_network_info")
async def api_network():
    result = await get_network_info()
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/get_battery_status")
async def api_battery():
    result = await get_battery_status()
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/search_web")
async def api_search(data: dict):
    result = await search_web(data.get("query", ""))
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/get_clipboard")
async def api_get_clip():
    result = await get_clipboard()
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/set_clipboard")
async def api_set_clip(data: dict):
    result = await set_clipboard(data.get("text", ""))
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/play_sound")
async def api_sound(data: dict):
    result = await play_sound(data.get("frequency", 1000), data.get("duration", 500))
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/set_volume")
async def api_tool_set_volume(data: dict):
    result = await set_volume(data.get("level", 50))
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/set_brightness")
async def api_brightness(data: dict):
    result = await set_brightness(data.get("level", 50))
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/mute_volume")
async def api_mute_volume(data: dict):
    result = await mute_volume()
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/unmute_volume")
async def api_unmute_volume(data: dict):
    result = await unmute_volume()
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/volume_up")
async def api_volume_up(data: dict):
    result = await volume_up(data.get("steps", 5))
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/volume_down")
async def api_volume_down(data: dict):
    result = await volume_down(data.get("steps", 5))
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/minimize_all_windows")
async def api_minimize():
    result = await show_desktop()  # Sử dụng show_desktop thay vì minimize_all_windows
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/undo_action")
async def api_undo():
    result = await undo_operation()  # Sử dụng undo_operation thay vì undo_action
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/toggle_dark_mode")
async def api_theme():
    result = await set_theme(dark_mode=None)  # Toggle bằng cách để None, hàm set_theme sẽ xử lý
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/change_wallpaper")
async def api_change_wallpaper(data: dict):
    """Đổi hình nền - endpoint cho Web UI"""
    keyword = data.get("keyword", "")
    path = data.get("path", "")
    result = await change_wallpaper(keyword=keyword, custom_path=path)
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/set_wallpaper")
async def api_wallpaper(data: dict):
    """Alias của change_wallpaper"""
    path = data.get("path", "")
    keyword = data.get("keyword", "")
    result = await change_wallpaper(keyword=keyword, custom_path=path)
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/paste_text")
async def api_paste():
    result = await paste_content(content="")  # paste_content với clipboard hiện tại
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/press_enter")
async def api_enter():
    result = await press_enter()
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/find_on_screen")
async def api_find(data: dict):
    result = await find_in_document(data.get("text", ""))  # Sử dụng find_in_document
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/lock_computer")
async def api_lock():
    result = await lock_computer()
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result

@app.post("/api/tool/shutdown_computer")
async def api_shutdown(data: dict):
    delay = data.get("delay", 0)
    # Sử dụng shutdown_schedule với action="shutdown"
    result = await shutdown_schedule(action="shutdown", delay=delay)
    if not result["success"]:
        raise HTTPException(500, result["error"])
    return result


@app.get("/logo.png")
async def get_logo():
    from fastapi.responses import FileResponse
    import os
    import sys
    
    # Tìm logo theo thứ tự ưu tiên
    possible_paths = []
    
    # 1. PyInstaller frozen EXE - trong thư mục _internal hoặc cùng thư mục EXE
    if getattr(sys, 'frozen', False):
        exe_dir = os.path.dirname(sys.executable)
        possible_paths.extend([
            os.path.join(exe_dir, "_internal", "logo.png"),
            os.path.join(exe_dir, "logo.png"),
            os.path.join(getattr(sys, '_MEIPASS', exe_dir), "logo.png"),
        ])
    
    # 2. Thư mục script
    possible_paths.append(os.path.join(os.path.dirname(__file__), "logo.png"))
    
    # 3. Thư mục làm việc hiện tại
    possible_paths.append(os.path.join(os.getcwd(), "logo.png"))
    
    # Tìm file đầu tiên tồn tại
    for logo_path in possible_paths:
        if os.path.exists(logo_path):
            return FileResponse(logo_path, media_type="image/png")
    
    # Log để debug
    print(f"⚠️ Logo not found. Checked paths: {possible_paths}")
    raise HTTPException(404, "Logo not found")

@app.get("/api/endpoints")
async def get_endpoints():
    global GEMINI_API_KEY, OPENAI_API_KEY, SERPER_API_KEY
    return {
        "endpoints": endpoints_config,
        "active_index": active_endpoint_index,
        "gemini_api_key": GEMINI_API_KEY,
        "openai_api_key": OPENAI_API_KEY,
        "serper_api_key": SERPER_API_KEY
    }

@app.get("/api/endpoints/status")
async def get_endpoints_status():
    """🔥 NEW: Get detailed endpoint connection status with stats"""
    status = {
        "endpoints": [],
        "active_index": active_endpoint_index,
        "total_connected": sum(1 for v in xiaozhi_connected.values() if v)
    }
    
    # Add detailed info for each endpoint
    for i, ep in enumerate(endpoints_config):
        endpoint_status = {
            "index": i,
            "name": ep.get("name", f"Thiết bị {i+1}"),
            "enabled": ep.get("enabled", False),
            "has_token": bool(ep.get("token")),
            "connected": xiaozhi_connected.get(i, False),
            "is_active": i == active_endpoint_index
        }
        
        # Thêm stats từ EndpointManager nếu có
        if ENDPOINT_MANAGER_AVAILABLE:
            try:
                manager = get_endpoint_manager()
                stats = manager.stats.get(i)
                if stats:
                    endpoint_status["stats"] = {
                        "total_connects": stats.total_connects,
                        "total_disconnects": stats.total_disconnects,
                        "total_errors": stats.total_errors,
                        "last_connected": stats.last_connected,
                        "last_error": stats.last_error,
                        "uptime_seconds": stats.uptime_seconds
                    }
            except Exception:
                pass
        
        status["endpoints"].append(endpoint_status)
    
    return status

@app.post("/api/endpoints/reconnect/{index}")
async def reconnect_endpoint(index: int):
    """🔥 NEW: Force reconnect an endpoint"""
    global should_reconnect
    
    if index < 0 or index >= len(endpoints_config):
        return {"success": False, "error": f"Invalid index: {index}"}
    
    ep = endpoints_config[index]
    if not ep.get("token"):
        return {"success": False, "error": "Endpoint has no token"}
    
    # Trigger reconnect
    should_reconnect[index] = True
    
    # Cập nhật EndpointManager nếu có
    if ENDPOINT_MANAGER_AVAILABLE:
        try:
            manager = get_endpoint_manager()
            manager.should_reconnect[index] = True
        except Exception:
            pass
    
    return {
        "success": True,
        "message": f"Đang reconnect {ep.get('name', f'Thiết bị {index+1}')}..."
    }

# YouTube Playlists API
@app.get("/api/youtube_playlists")
async def api_get_youtube_playlists():
    """Lấy danh sách playlist YouTube"""
    return await get_youtube_playlists()

@app.post("/api/youtube_playlists/add")
async def api_add_youtube_playlist(data: dict):
    """Thêm playlist YouTube mới"""
    name = data.get("name", "").strip()
    url = data.get("url", "").strip()
    
    if not name or not url:
        return {"success": False, "error": "Tên và URL không được để trống"}
    
    return await add_youtube_playlist(name, url)

@app.post("/api/youtube_playlists/remove")
async def api_remove_youtube_playlist(data: dict):
    """Xóa playlist YouTube"""
    name = data.get("name", "").strip()
    
    if not name:
        return {"success": False, "error": "Tên playlist không được để trống"}
    
    return await remove_youtube_playlist(name)

# ============================================================
# KNOWLEDGE BASE API - Quản lý dữ liệu cho LLM
# ============================================================

# File lưu cấu hình knowledge base - Lưu vào AppData để tránh Permission denied
def get_knowledge_data_dir():
    """Lấy thư mục lưu trữ knowledge base data trong AppData"""
    if os.name == 'nt':  # Windows
        appdata = os.environ.get('LOCALAPPDATA', os.path.expanduser('~\\AppData\\Local'))
        data_dir = Path(appdata) / "miniZ_MCP" / "knowledge"
    else:  # Linux/Mac
        data_dir = Path.home() / ".miniz_mcp" / "knowledge"
    data_dir.mkdir(parents=True, exist_ok=True)
    return data_dir

KNOWLEDGE_DATA_DIR = get_knowledge_data_dir()
KNOWLEDGE_CONFIG_FILE = KNOWLEDGE_DATA_DIR / "knowledge_config.json"
KNOWLEDGE_INDEX_FILE = KNOWLEDGE_DATA_DIR / "knowledge_index.json"

# Các extension được hỗ trợ
SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.docx', '.doc', '.md', '.json', '.csv', '.xlsx', '.xls', '.rtf'}

def load_knowledge_config():
    """Load cấu hình knowledge base"""
    if KNOWLEDGE_CONFIG_FILE.exists():
        try:
            # Sử dụng utf-8-sig để tự động xử lý BOM
            with open(KNOWLEDGE_CONFIG_FILE, 'r', encoding='utf-8-sig') as f:
                return json.load(f)
        except Exception as e:
            print(f"⚠️ [Knowledge] Error loading config: {e}")
    return {"folder_path": "", "indexed_files": [], "last_update": ""}

def save_knowledge_config(config: dict):
    """Lưu cấu hình knowledge base"""
    try:
        with open(KNOWLEDGE_CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        print(f"❌ [Knowledge] Error saving config: {e}")
        return False

def load_knowledge_index():
    """Load index đã lưu"""
    if KNOWLEDGE_INDEX_FILE.exists():
        try:
            # Sử dụng utf-8-sig để tự động xử lý BOM
            with open(KNOWLEDGE_INDEX_FILE, 'r', encoding='utf-8-sig') as f:
                return json.load(f)
        except Exception as e:
            print(f"⚠️ [Knowledge] Error loading index: {e}")
    return {"documents": [], "total_chunks": 0, "last_update": ""}

# ============================================================
# VECTOR SEARCH ENGINE - Global Instance
# ============================================================

_vector_engine = None

def get_vector_engine():
    """Lấy hoặc khởi tạo VectorSearchEngine"""
    global _vector_engine
    if _vector_engine is None and VECTOR_SEARCH_AVAILABLE:
        _vector_engine = VectorSearchEngine()
        
        # Try loading existing index - kiểm tra nhiều vị trí
        vector_paths = [
            Path("test_vector.faiss"),  # Trong thư mục gốc
            KNOWLEDGE_DATA_DIR / "vector_index.faiss",  # Trong AppData
            Path("vector_index.faiss")  # Backup trong gốc
        ]
        
        for vector_index_path in vector_paths:
            if vector_index_path.exists():
                try:
                    # Remove .faiss extension for load_index
                    base_path = str(vector_index_path.with_suffix(''))
                    _vector_engine.load_index(base_path)
                    print(f"✅ [VectorSearch] Loaded index from: {vector_index_path}")
                    print(f"   Statistics: {_vector_engine.get_statistics()}")
                    break
                except Exception as e:
                    print(f"⚠️ [VectorSearch] Failed to load {vector_index_path}: {e}")
                    continue
        else:
            print(f"⚠️ [VectorSearch] No valid index found in any location")
            
    return _vector_engine

def save_knowledge_index(index_data: dict):
    """Lưu index"""
    try:
        with open(KNOWLEDGE_INDEX_FILE, 'w', encoding='utf-8') as f:
            json.dump(index_data, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        print(f"❌ [Knowledge] Error saving index: {e}")
        return False

async def summarize_with_gemini(text: str, filename: str) -> dict:
    """Tóm tắt document bằng Gemini Flash (optimized)"""
    try:
        import google.generativeai as genai
        
        # Configure Gemini
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel('models/gemini-3-flash-preview')
        
        # ⚡ PROMPT NGẮN GỌN - phản hồi nhanh hơn
        prompt = f"""Tóm tắt tài liệu:

File: {filename}
Nội dung: {text[:6000]}

Trả về JSON:
{{
  "summary": "[2-3 câu chính]",
  "keywords": ["5-7 từ khóa"],
  "key_quotes": ["2 trích dẫn quan trọng"],
  "category": "[loại: technical/business/etc]"
}}"""
        
        print(f"⚡ [Gemini] Tóm tắt: {filename[:30]}...")
        
        # ⏱️ Timeout 12 giây
        loop = asyncio.get_event_loop()
        response = await asyncio.wait_for(
            loop.run_in_executor(None, lambda: model.generate_content(prompt)),
            timeout=12.0
        )
        
        # Parse JSON response
        import json
        result_text = response.text.strip()
        if result_text.startswith("```json"):
            result_text = result_text[7:]
        if result_text.endswith("```"):
            result_text = result_text[:-3]
        result_text = result_text.strip()
        
        result = json.loads(result_text)
        print(f"✅ [Gemini] Done: {filename[:30]}")
        return result
        
    except asyncio.TimeoutError:
        print(f"⏱️ [Gemini] Timeout: {filename}")
        return {
            "summary": text[:400] + "...",
            "keywords": [],
            "key_quotes": [],
            "category": "unknown"
        }
    except Exception as e:
        print(f"⚠️ [Gemini] Error {filename}: {e}")
        return {
            "summary": text[:400] + "...",
            "keywords": [],
            "key_quotes": [],
            "category": "unknown"
        }

def extract_text_from_file(file_path: str) -> str:
    """Trích xuất text từ file"""
    ext = Path(file_path).suffix.lower()
    text = ""
    
    try:
        if ext == '.txt' or ext == '.md':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        
        elif ext == '.json':
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                text = json.dumps(data, ensure_ascii=False, indent=2)
        
        elif ext == '.csv':
            import csv
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                reader = csv.reader(f)
                rows = [', '.join(row) for row in reader]
                text = '\n'.join(rows)
        
        elif ext == '.pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
            except ImportError:
                # Fallback: Read as binary and extract text using basic regex
                print(f"⚠️ [Extract] PyPDF2 not installed, using fallback for {file_path}")
                try:
                    with open(file_path, 'rb') as f:
                        content = f.read()
                        import re
                        # Simple extraction: find readable ASCII/Unicode text
                        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f]+', ' ', 
                                     content.decode('latin1', errors='ignore'))
                        text = ' '.join(text.split())  # Clean whitespace
                except:
                    text = f"[PDF file - Cần cài PyPDF2: pip install PyPDF2]"
            except Exception as e:
                text = f"[Lỗi đọc PDF: {str(e)}]"
        
        elif ext in ['.docx', '.doc']:
            try:
                from docx import Document
                doc = Document(file_path)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            except ImportError:
                # Fallback: Try reading docx as zip
                print(f"⚠️ [Extract] python-docx not installed, using fallback for {file_path}")
                try:
                    import zipfile
                    import xml.etree.ElementTree as ET
                    with zipfile.ZipFile(file_path) as docx:
                        xml_content = docx.read('word/document.xml')
                        tree = ET.XML(xml_content)
                        paragraphs = []
                        for paragraph in tree.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                            if paragraph.text:
                                paragraphs.append(paragraph.text)
                        text = '\n'.join(paragraphs)
                except:
                    text = f"[Word file - Cần cài python-docx: pip install python-docx]"
            except Exception as e:
                text = f"[Lỗi đọc Word: {str(e)}]"
        
        elif ext in ['.xlsx', '.xls']:
            try:
                import openpyxl
                print(f"✅ [Extract] openpyxl loaded, reading: {file_path}")
                wb = openpyxl.load_workbook(file_path, data_only=True)
                rows_read = 0
                for sheet in wb.worksheets:
                    for row in sheet.iter_rows():
                        row_text = ', '.join([str(cell.value) if cell.value else '' for cell in row])
                        if row_text.strip():
                            text += row_text + "\n"
                            rows_read += 1
                print(f"✅ [Extract] Excel read complete: {rows_read} rows")
            except ImportError as ie:
                print(f"❌ [Extract] openpyxl ImportError: {ie}")
                text = f"[Excel file - Cần cài openpyxl: pip install openpyxl]"
            except Exception as e:
                print(f"❌ [Extract] Excel error: {type(e).__name__}: {e}")
                text = f"[Lỗi đọc Excel: {str(e)}]"
        
        elif ext == '.rtf':
            try:
                from striprtf.striprtf import rtf_to_text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    rtf_content = f.read()
                text = rtf_to_text(rtf_content)
            except ImportError:
                text = f"[RTF file - Cần cài striprtf: pip install striprtf]"
            except Exception as e:
                text = f"[Lỗi đọc RTF: {str(e)}]"
        
    except Exception as e:
        text = f"[Lỗi đọc file: {str(e)}]"
    
    return text.strip()

def scan_folder_for_files(folder_path: str) -> list:
    """Quét thư mục và trả về danh sách files được hỗ trợ"""
    files = []
    folder = Path(folder_path)
    
    if not folder.exists():
        print(f"❌ [Scan] Folder not exists: {folder_path}")
        return files
    
    print(f"📂 [Scan] Scanning folder: {folder_path}")
    total_checked = 0
    
    for file_path in folder.rglob('*'):
        if file_path.is_file():
            total_checked += 1
            ext = file_path.suffix.lower()
            if ext in SUPPORTED_EXTENSIONS:
                try:
                    stat = file_path.stat()
                    files.append({
                        "name": file_path.name,
                        "path": str(file_path),
                        "size": stat.st_size,
                        "modified": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M"),
                        "extension": ext,
                        "indexed": False
                    })
                    print(f"  ✅ Added: {file_path.name} ({ext})")
                except Exception as e:
                    print(f"  ⚠️ Error scanning file {file_path}: {e}")
            else:
                print(f"  ⏭️ Skipped: {file_path.name} ({ext}) - Not supported")
    
    print(f"📊 [Scan] Result: {len(files)} files found (checked {total_checked} files)")
    return files

@app.get("/api/knowledge/status")
async def api_knowledge_status():
    """Lấy trạng thái Knowledge Base"""
    config = load_knowledge_config()
    index = load_knowledge_index()
    
    folder_path = config.get("folder_path", "")
    files = []
    total_size = 0
    
    if folder_path and Path(folder_path).exists():
        files = scan_folder_for_files(folder_path)
        total_size = sum(f["size"] for f in files)
        
        # Đánh dấu các file đã được index
        indexed_paths = set(config.get("indexed_files", []))
        for f in files:
            f["indexed"] = f["path"] in indexed_paths
    
    return {
        "success": True,
        "folder_path": folder_path,
        "total_files": len(files),
        "indexed_files": len(config.get("indexed_files", [])),
        "total_size": total_size,
        "last_update": config.get("last_update", "--"),
        "files": files
    }

@app.post("/api/knowledge/set_folder")
async def api_knowledge_set_folder(data: dict):
    """Cấu hình thư mục knowledge base"""
    folder_path = data.get("folder_path", "").strip()
    
    if not folder_path:
        return {"success": False, "error": "Đường dẫn không được để trống"}
    
    # Kiểm tra thư mục tồn tại
    if not Path(folder_path).exists():
        return {"success": False, "error": f"Thư mục không tồn tại: {folder_path}"}
    
    if not Path(folder_path).is_dir():
        return {"success": False, "error": "Đường dẫn phải là thư mục, không phải file"}
    
    config = load_knowledge_config()
    config["folder_path"] = folder_path
    config["last_update"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    
    if save_knowledge_config(config):
        return {"success": True, "message": f"Đã lưu thư mục: {folder_path}"}
    else:
        return {"success": False, "error": "Lỗi khi lưu cấu hình"}

@app.post("/api/knowledge/scan")
async def api_knowledge_scan(data: dict):
    """Quét thư mục để tìm files"""
    folder_path = data.get("folder_path", "").strip()
    
    if not folder_path:
        config = load_knowledge_config()
        folder_path = config.get("folder_path", "")
    
    if not folder_path:
        return {"success": False, "error": "Chưa cấu hình thư mục"}
    
    if not Path(folder_path).exists():
        return {"success": False, "error": f"Thư mục không tồn tại: {folder_path}"}
    
    files = scan_folder_for_files(folder_path)
    total_size = sum(f["size"] for f in files)
    
    # Cập nhật config
    config = load_knowledge_config()
    config["folder_path"] = folder_path
    indexed_paths = set(config.get("indexed_files", []))
    for f in files:
        f["indexed"] = f["path"] in indexed_paths
    save_knowledge_config(config)
    
    return {
        "success": True,
        "total_files": len(files),
        "total_size": total_size,
        "files": files
    }

@app.post("/api/knowledge/index_all")
async def api_knowledge_index_all():
    """Index tất cả files trong thư mục (parallel processing)"""
    config = load_knowledge_config()
    folder_path = config.get("folder_path", "")
    
    if not folder_path or not Path(folder_path).exists():
        return {"success": False, "error": "Chưa cấu hình thư mục hoặc thư mục không tồn tại"}
    
    files = scan_folder_for_files(folder_path)
    print(f"⚡ [Index] Starting parallel indexing of {len(files)} files...")
    
    # ⚡ PARALLEL PROCESSING: Index nhiều files cùng lúc
    async def index_single_file(file_info):
        try:
            text = extract_text_from_file(file_info["path"])
            
            # Check if extraction failed
            if not text or len(text.strip()) < 10:
                print(f"⚠️ [Index] Skipped {file_info['name']}: No text extracted")
                return None
                
            if text.startswith("["):  # Error message from extract_text_from_file
                print(f"⚠️ [Index] Skipped {file_info['name']}: {text}")
                return None
            
            print(f"📄 [Index] Processing {file_info['name']} ({len(text)} chars)...")
            
            # Tóm tắt bằng Gemini Flash
            ai_summary = await summarize_with_gemini(text, file_info["name"])
            
            if not ai_summary or not ai_summary.get("summary"):
                print(f"⚠️ [Index] No summary for {file_info['name']}")
                # Still index with basic info
                ai_summary = {
                    "summary": text[:400] + "...",
                    "keywords": [],
                    "key_quotes": [],
                    "category": "general"
                }
            
            result = {
                "file_path": file_info["path"],
                "file_name": file_info["name"],
                "content": text[:50000],  # Giới hạn 50k ký tự mỗi file
                "summary": ai_summary.get("summary", ""),
                "keywords": ai_summary.get("keywords", []),
                "key_quotes": ai_summary.get("key_quotes", []),
                "category": ai_summary.get("category", "general"),
                "indexed_at": datetime.now().isoformat()
            }
            print(f"✅ [Index] Indexed {file_info['name']}")
            return result
            
        except Exception as e:
            print(f"❌ [Index] Error indexing {file_info['name']}: {e}")
            return None
    
    # Process files in parallel (batch of 5 at a time to avoid API rate limits)
    documents = []
    batch_size = 5
    for i in range(0, len(files), batch_size):
        batch = files[i:i+batch_size]
        results = await asyncio.gather(*[index_single_file(f) for f in batch], return_exceptions=True)
        documents.extend([r for r in results if r and not isinstance(r, Exception)])
        print(f"⚡ [Index] Processed {min(i+batch_size, len(files))}/{len(files)} files...")
    
    indexed_count = len(documents)
    
    # Lưu index
    index_data = {
        "documents": documents,
        "total_chunks": indexed_count,
        "last_update": datetime.now().strftime("%Y-%m-%d %H:%M")
    }
    save_knowledge_index(index_data)
    
    # 🆕 BUILD VECTOR INDEX with FAISS
    if VECTOR_SEARCH_AVAILABLE and documents:
        try:
            print(f"🔨 [VectorSearch] Building vector index for {len(documents)} documents...")
            vector_engine = get_vector_engine()
            
            # Prepare documents in correct format: [{"id": str, "text": str, "metadata": dict}]
            documents_data = [
                {
                    "id": f"doc_{i}",
                    "text": doc["content"],
                    "metadata": {
                        "file_name": doc["file_name"],
                        "file_path": doc["file_path"],
                        "index": i
                    }
                }
                for i, doc in enumerate(documents)
            ]
            
            # Build and save index
            vector_engine.build_index(documents_data)
            vector_engine.save_index()
            
            stats = vector_engine.get_statistics()
            print(f"✅ [VectorSearch] Index built: {stats['num_vectors']} vectors, {stats['embedding_dim']} dims")
        except Exception as e:
            print(f"⚠️ [VectorSearch] Failed to build index: {e}")
    
    # Cập nhật config
    config["indexed_files"] = [f["path"] for f in files if any(d["file_path"] == f["path"] for d in documents)]
    config["last_update"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    save_knowledge_config(config)
    
    return {
        "success": True,
        "message": f"Đã index {indexed_count}/{len(files)} files",
        "indexed_count": indexed_count,
        "last_update": index_data["last_update"]
    }

@app.post("/api/knowledge/index_file")
async def api_knowledge_index_file(data: dict):
    """Index một file cụ thể"""
    file_path = data.get("file_path", "").strip()
    
    if not file_path or not Path(file_path).exists():
        return {"success": False, "error": "File không tồn tại"}
    
    try:
        file_name = Path(file_path).name
        print(f"📄 [Index] Starting index: {file_name}")
        
        text = extract_text_from_file(file_path)
        if not text or text.startswith("["):
            print(f"❌ [Index] Failed to extract: {file_name} - {text[:100] if text else 'Empty'}")
            return {"success": False, "error": f"Không thể đọc file: {text}"}
        
        print(f"📝 [Index] Extracted {len(text)} chars from {file_name}")
        
        # Load existing index
        index_data = load_knowledge_index()
        
        # Remove existing entry for this file
        index_data["documents"] = [d for d in index_data["documents"] if d["file_path"] != file_path]
        
        # 🆕 TRY summarize, nhưng fallback nếu fail
        ai_summary = {"summary": "", "keywords": [], "key_quotes": [], "category": "general"}
        try:
            ai_summary = await asyncio.wait_for(
                summarize_with_gemini(text, file_name),
                timeout=30.0  # 30s timeout
            )
            print(f"✅ [Index] AI Summary done for {file_name}")
        except asyncio.TimeoutError:
            print(f"⚠️ [Index] AI Summary timeout for {file_name}, using basic index")
            ai_summary["summary"] = text[:500] + "..."
        except Exception as e:
            print(f"⚠️ [Index] AI Summary error for {file_name}: {e}, using basic index")
            ai_summary["summary"] = text[:500] + "..."
        
        # Add new entry
        index_data["documents"].append({
            "file_path": file_path,
            "file_name": file_name,
            "content": text[:50000],
            "summary": ai_summary.get("summary", ""),
            "keywords": ai_summary.get("keywords", []),
            "key_quotes": ai_summary.get("key_quotes", []),
            "category": ai_summary.get("category", "general"),
            "indexed_at": datetime.now().isoformat()
        })
        index_data["total_chunks"] = len(index_data["documents"])
        index_data["last_update"] = datetime.now().strftime("%Y-%m-%d %H:%M")
        
        save_knowledge_index(index_data)
        print(f"✅ [Index] Saved: {file_name} (total: {index_data['total_chunks']} docs)")
        
        # Update config
        config = load_knowledge_config()
        if file_path not in config.get("indexed_files", []):
            config.setdefault("indexed_files", []).append(file_path)
        config["last_update"] = datetime.now().strftime("%Y-%m-%d %H:%M")
        save_knowledge_config(config)
        
        return {"success": True, "message": f"Đã index: {file_name}"}
    
    except Exception as e:
        print(f"❌ [Index] Error indexing {file_path}: {e}")
        return {"success": False, "error": str(e)}

@app.post("/api/knowledge/clear")
async def api_knowledge_clear():
    """Xóa toàn bộ index"""
    try:
        # Clear index file
        save_knowledge_index({"documents": [], "total_chunks": 0, "last_update": ""})
        
        # Update config
        config = load_knowledge_config()
        config["indexed_files"] = []
        config["last_update"] = datetime.now().strftime("%Y-%m-%d %H:%M")
        save_knowledge_config(config)
        
        return {"success": True, "message": "Đã xóa toàn bộ index"}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.get("/api/knowledge/search")
async def api_knowledge_search(query: str = ""):
    """Tìm kiếm trong knowledge base"""
    if not query:
        return {"success": False, "error": "Vui lòng nhập từ khóa tìm kiếm"}
    
    index_data = load_knowledge_index()
    documents = index_data.get("documents", [])
    
    if not documents:
        return {"success": False, "error": "Knowledge base chưa có dữ liệu. Vui lòng index files trước."}
    
    # AI-powered search - tìm trong summary, keywords và content
    query_lower = query.lower()
    results = []
    
    for doc in documents:
        score = 0
        matched_in = []
        
        # Tìm trong summary (điểm cao nhất)
        summary = doc.get("summary", "")
        if query_lower in summary.lower():
            score += 10
            matched_in.append("summary")
        
        # Tìm trong keywords (điểm trung bình)
        keywords = doc.get("keywords", [])
        for keyword in keywords:
            if query_lower in keyword.lower():
                score += 5
                matched_in.append("keywords")
                break
        
        # Tìm trong content (điểm thấp nhất)
        content = doc.get("content", "")
        if query_lower in content.lower():
            score += 1
            matched_in.append("content")
            
            # Tìm đoạn text chứa query
            idx = content.lower().find(query_lower)
            start = max(0, idx - 200)
            end = min(len(content), idx + 200)
            snippet = content[start:end]
        else:
            snippet = summary[:400] if summary else content[:400]
        
        # Chỉ thêm vào results nếu có match
        if score > 0:
            results.append({
                "file_name": doc.get("file_name", ""),
                "file_path": doc.get("file_path", ""),
                "summary": summary,
                "keywords": keywords,
                "category": doc.get("category", "general"),
                "snippet": "..." + snippet + "...",
                "score": score,
                "matched_in": matched_in,
                "indexed_at": doc.get("indexed_at", "")
            })
    
    # Sắp xếp theo score
    results.sort(key=lambda x: x["score"], reverse=True)
    
    return {
        "success": True,
        "query": query,
        "total_results": len(results),
        "results": results[:20]  # Giới hạn 20 kết quả
    }

@app.get("/api/knowledge/context")
async def api_knowledge_get_context(query: str = "", max_chars: int = 10000, use_gemini_summary: bool = True):
    """Lấy context từ knowledge base để cung cấp cho LLM - với Gemini summarization"""
    result = await get_knowledge_context(query, max_chars, use_gemini_summary)
    return result
    
@app.get("/api/knowledge/context_legacy")
async def api_knowledge_get_context_legacy(query: str = "", max_chars: int = 10000):
    """Legacy endpoint - không dùng Gemini summarization"""
    index_data = load_knowledge_index()
    documents = index_data.get("documents", [])
    
    if not documents:
        return {"success": False, "context": "", "message": "Knowledge base trống"}
    
    context_parts = []
    total_chars = 0
    
    # Nếu có query, ưu tiên các document liên quan
    if query:
        query_lower = query.lower()
        # Sắp xếp theo độ liên quan
        scored_docs = []
        for doc in documents:
            content = doc.get("content", "")
            score = content.lower().count(query_lower)
            scored_docs.append((score, doc))
        scored_docs.sort(key=lambda x: x[0], reverse=True)
        documents = [d for _, d in scored_docs]
    
    for doc in documents:
        file_name = doc.get("file_name", "unknown")
        summary = doc.get("summary", "")
        keywords = doc.get("keywords", [])
        key_quotes = doc.get("key_quotes", [])
        category = doc.get("category", "general")
        
        # Ưu tiên dùng summary và key_quotes thay vì full content
        # Điều này giảm đáng kể token và tăng chất lượng context
        
        # Build compact context
        compact_content = f"📝 {summary}\n"
        if keywords:
            compact_content += f"🔑 Keywords: {', '.join(keywords[:5])}\n"
        if key_quotes:
            compact_content += f"💬 Trích dẫn:\n"
            for quote in key_quotes[:3]:
                compact_content += f"  • {quote}\n"
        
        # Thêm header cho mỗi document
        header = f"\n\n=== [{category.upper()}] {file_name} ===\n"
        full_entry = header + compact_content
        
        if total_chars + len(full_entry) > max_chars:
            break
        else:
            context_parts.append(full_entry)
            total_chars += len(full_entry)
    
    full_context = "".join(context_parts)
    
    return {
        "success": True,
        "context": full_context,
        "total_documents": len(documents),
        "context_length": len(full_context)
    }

# ============================================================
# TASK MEMORY API - Ghi nhớ tác vụ đã thực hiện
# ============================================================

@app.get("/api/tasks/recent")
async def api_get_recent_tasks(limit: int = 10):
    """Lấy các tác vụ gần đây"""
    tasks = get_recent_tasks(limit)
    return {
        "success": True,
        "count": len(tasks),
        "tasks": tasks
    }

@app.get("/api/tasks/search/{keyword}")
async def api_search_tasks(keyword: str):
    """Tìm kiếm tác vụ theo từ khóa"""
    results = search_task_memory(keyword)
    return {
        "success": True,
        "count": len(results),
        "tasks": results
    }

@app.get("/api/tasks/all")
async def api_get_all_tasks():
    """Lấy toàn bộ lịch sử tác vụ"""
    tasks = load_task_memory()
    return {
        "success": True,
        "total": len(tasks),
        "tasks": tasks
    }

@app.post("/api/tasks/clear")
async def api_clear_tasks():
    """Xóa toàn bộ lịch sử tác vụ"""
    success = clear_task_memory()
    return {
        "success": success,
        "message": "Đã xóa toàn bộ lịch sử tác vụ" if success else "Lỗi khi xóa"
    }

@app.get("/api/tasks/summary")
async def api_get_task_summary():
    """Lấy tổng hợp thống kê tác vụ"""
    tasks = load_task_memory()
    
    if not tasks:
        return {
            "success": True,
            "total_tasks": 0,
            "by_tool": {},
            "success_rate": 0,
            "recent_tools": []
        }
    
    # Đếm theo tool
    tool_counts = {}
    success_count = 0
    
    for task in tasks:
        tool = task.get('tool', 'unknown')
        tool_counts[tool] = tool_counts.get(tool, 0) + 1
        if task.get('result_success'):
            success_count += 1
    
    # Sắp xếp theo số lần sử dụng
    sorted_tools = sorted(tool_counts.items(), key=lambda x: x[1], reverse=True)
    
    return {
        "success": True,
        "total_tasks": len(tasks),
        "by_tool": dict(sorted_tools[:20]),
        "success_rate": round(success_count / len(tasks) * 100, 1),
        "recent_tools": [t.get('tool') for t in tasks[-5:]]
    }

# ============================================================
# CONVERSATION HISTORY API
# ============================================================

@app.get("/api/conversation/history")
async def api_get_conversation_history():
    """Lấy toàn bộ lịch sử hội thoại"""
    return {
        "success": True,
        "total_messages": len(conversation_history),
        "messages": conversation_history
    }

@app.get("/api/conversation/recent/{count}")
async def api_get_recent_conversation(count: int = 10):
    """Lấy N messages gần nhất"""
    recent = conversation_history[-count:] if len(conversation_history) > count else conversation_history
    return {
        "success": True,
        "count": len(recent),
        "messages": recent
    }

@app.post("/api/conversation/clear")
async def api_clear_conversation():
    """Xóa toàn bộ lịch sử hội thoại"""
    global conversation_history
    conversation_history = []
    save_conversation_history()
    return {
        "success": True,
        "message": "Đã xóa toàn bộ lịch sử hội thoại"
    }

@app.post("/api/conversation/export")
async def api_export_conversation(data: dict = None):
    """Export lịch sử hội thoại ra file"""
    filename = data.get("filename", "") if data else ""
    return await export_conversation_to_file(filename)

@app.post("/api/conversation/add")
async def api_add_conversation_message(data: dict):
    """Thêm message từ Web UI vào history"""
    role = data.get("role", "user")
    content = data.get("content", "")
    metadata = data.get("metadata", {})
    
    if not content:
        return {"success": False, "error": "Content không được để trống"}
    
    add_to_conversation(role, content, metadata)
    
    return {
        "success": True,
        "message": "Đã thêm message vào history"
    }

@app.post("/api/chat/log")
async def api_log_chat_message(data: dict):
    """
    Endpoint đặc biệt để Web UI log TOÀN BỘ cuộc hội thoại
    Dùng cho các chat không qua MCP
    """
    messages = data.get("messages", [])
    
    if not messages:
        return {"success": False, "error": "Không có messages để log"}
    
    # Log từng message
    for msg in messages:
        role = msg.get("role", "user")
        content = msg.get("content", "")
        metadata = msg.get("metadata", {})
        
        if content:
            add_to_conversation(role, content, metadata)
    
    return {
        "success": True,
        "message": f"Đã log {len(messages)} messages vào history",
        "total_messages": len(conversation_history)
    }

# ============================================================
# USER PROFILE API - Hiểu người dùng
# ============================================================

@app.get("/api/user/profile")
async def api_get_user_profile():
    """Lấy user profile"""
    return {
        "success": True,
        "profile": load_user_profile(),
        "summary": get_user_profile_summary()
    }

@app.get("/api/user/context")
async def api_get_user_context(max_messages: int = 10):
    """Lấy context từ lịch sử hội thoại + user profile"""
    return {
        "success": True,
        "user_profile": get_user_profile_summary(),
        "recent_conversation": get_conversation_context(max_messages),
        "hint": "Dùng thông tin này để hiểu người dùng tốt hơn"
    }

@app.get("/api/conversation/files")
async def api_list_conversation_files():
    """Liệt kê các file hội thoại đã lưu"""
    files = list_conversation_files()
    return {
        "success": True,
        "storage_path": str(CONVERSATION_BASE_DIR),
        "total_files": len(files),
        "files": files
    }

# NOTE: Endpoint /api/conversation/today đã bị xóa (không còn file theo ngày)

@app.post("/api/endpoints/switch/{index}")
async def switch_endpoint(index: int):
    global active_endpoint_index, should_reconnect
    if index < 0 or index >= len(endpoints_config):
        return {"success": False, "error": "Thiết bị không tồn tại"}
    
    device = endpoints_config[index]
    if not device.get("token"):
        return {"success": False, "error": "Thiết bị chưa có token. Hãy nhập token và lưu lại!"}
    
    # Thay đổi endpoint và trigger reconnect
    old_index = active_endpoint_index
    active_endpoint_index = index
    should_reconnect = True  # Trigger reconnect trong xiaozhi_websocket_client
    
    # Lưu vào file
    save_endpoints_to_file(endpoints_config, active_endpoint_index)
    
    print(f"🔄 [Endpoint] Switching from device {old_index} to {index} ({device['name']})")
    
    return {"success": True, "message": f"Đã chuyển sang {device['name']}. Đang kết nối lại..."}

@app.post("/api/endpoints/save")
async def save_endpoints(data: dict):
    global endpoints_config, should_reconnect
    try:
        devices = data.get('devices', [])
        if not devices:
            return {"success": False, "error": "Không có dữ liệu"}
        
        # Lưu token cũ của thiết bị đang active để so sánh
        old_active_token = endpoints_config[active_endpoint_index].get('token', '') if active_endpoint_index < len(endpoints_config) else ''
        
        # Cập nhật endpoints_config
        endpoints_config = []
        for dev in devices:
            token = dev.get('token', '').strip()  # Strip whitespace
            endpoints_config.append({
                'name': dev.get('name', 'Thiết bị'),
                'token': token,
                'enabled': bool(token)  # Only enabled if token not empty
            })
        
        # 🔥 FIX: FORCE SAVE khi user bấm Save - không skip
        if save_endpoints_to_file(endpoints_config, active_endpoint_index, force_save=True):
            print(f"✅ [Endpoint] User saved {len(devices)} devices (forced)")
        else:
            print(f"⚠️ [Endpoint] Failed to save to file, but config updated in memory")
        
        # CHỈ reconnect nếu token thay đổi VÀ có giá trị mới khác rỗng
        new_active_token = endpoints_config[active_endpoint_index].get('token', '') if active_endpoint_index < len(endpoints_config) else ''
        if old_active_token != new_active_token and new_active_token and old_active_token:
            # Token đã thay đổi (không phải lần đầu nhập)
            should_reconnect = True
            print(f"🔄 [Endpoint] Token changed for active device {active_endpoint_index}. Triggering reconnect...")
        
        return {"success": True, "message": "Đã lưu cấu hình"}
    except Exception as e:
        print(f"❌ [Endpoint] Error saving: {e}")
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}

@app.post("/api/gemini-key")
async def save_gemini_key(data: dict):
    """Save Gemini API key - Auto-save endpoint"""
    global GEMINI_API_KEY
    try:
        api_key = data.get('api_key', '').strip()
        
        # 🔥 FIX: Cho phép empty string (user xóa key)
        if api_key:
            # Validate format only if key is provided
            if not api_key.startswith('AIzaSy'):
                return {"success": False, "error": "API key không hợp lệ (phải bắt đầu với 'AIzaSy')"}
        
        # Update global variable (allow empty)
        GEMINI_API_KEY = api_key
        
        # Save to file
        if save_endpoints_to_file(endpoints_config, active_endpoint_index):
            if api_key:
                print(f"✅ [Gemini] API key saved (ends with ...{api_key[-8:]})")
                return {
                    "success": True,
                    "message": "✓ Đã lưu Gemini API key",
                    "key_preview": f"...{api_key[-8:]}"
                }
            else:
                print("✅ [Gemini] API key cleared")
                return {
                    "success": True,
                    "message": "✓ Đã xóa Gemini API key"
                }
        else:
            return {"success": False, "error": "Lỗi lưu file config"}
    except Exception as e:
        print(f"❌ [Gemini] Error saving API key: {e}")
        return {"success": False, "error": str(e)}

@app.post("/api/openai-key")
async def save_openai_key(data: dict):
    """Save OpenAI API key - Auto-save endpoint"""
    global OPENAI_API_KEY
    try:
        api_key = data.get('api_key', '').strip()
        
        # 🔥 FIX: Cho phép empty string (user xóa key)
        if api_key:
            # Validate format only if key is provided
            if not api_key.startswith('sk-'):
                return {"success": False, "error": "API key không hợp lệ (phải bắt đầu với 'sk-')"}
        
        # Update global variable (allow empty)
        OPENAI_API_KEY = api_key
        
        # Save to file
        if save_endpoints_to_file(endpoints_config, active_endpoint_index):
            if api_key:
                print(f"✅ [OpenAI] API key saved (ends with ...{api_key[-8:]})")
                return {
                    "success": True,
                    "message": "✓ Đã lưu OpenAI API key",
                    "key_preview": f"...{api_key[-8:]}"
                }
            else:
                print("✅ [OpenAI] API key cleared")
                return {
                    "success": True,
                    "message": "✓ Đã xóa OpenAI API key"
                }
        else:
            return {"success": False, "error": "Lỗi lưu file config"}
    except Exception as e:
        print(f"❌ [OpenAI] Error saving API key: {e}")
        return {"success": False, "error": str(e)}

@app.post("/api/serper-key")
async def save_serper_key(data: dict):
    """Save Serper API key (Google Search) - Auto-save endpoint"""
    global SERPER_API_KEY
    try:
        api_key = data.get('api_key', '').strip()
        
        # 🔥 FIX: Cho phép empty string (user xóa key)
        # Update global variable (allow empty)
        SERPER_API_KEY = api_key
        
        # Cập nhật environment variable để rag_system.py có thể dùng
        if api_key:
            os.environ['SERPER_API_KEY'] = api_key
        else:
            os.environ.pop('SERPER_API_KEY', None)  # Remove if empty
        
        # Save to file
        if save_endpoints_to_file(endpoints_config, active_endpoint_index):
            if api_key:
                print(f"✅ [Serper] Google Search API key saved (ends with ...{api_key[-8:]})")
                return {
                    "success": True,
                    "message": "✓ Đã lưu Serper API key - Google Search sẵn sàng!",
                    "key_preview": f"...{api_key[-8:]}"
                }
            else:
                print("✅ [Serper] API key cleared")
                return {
                    "success": True,
                    "message": "✓ Đã xóa Serper API key"
                }
        else:
            return {"success": False, "error": "Lỗi lưu file config"}
    except Exception as e:
        print(f"❌ [Serper] Error saving API key: {e}")
        return {"success": False, "error": str(e)}

@app.get("/api/serper-key")
async def get_serper_key():
    """Get current Serper API key status"""
    if SERPER_API_KEY:
        return {
            "success": True,
            "has_key": True,
            "key_preview": f"...{SERPER_API_KEY[-8:]}"
        }
    return {"success": True, "has_key": False}


@app.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    await websocket.accept()
    active_connections.append(websocket)
    try:
        await websocket.send_json({"type": "xiaozhi_status", "connected": xiaozhi_connected})
        while True:
            data = await websocket.receive_text()
            
            # Parse và log WebSocket messages
            try:
                msg_data = json.loads(data)
                msg_type = msg_data.get("type", "")
                
                # Lưu user messages từ Web UI
                if msg_type == "chat_message":
                    user_msg = msg_data.get("message", "")
                    if user_msg:
                        add_to_conversation(
                            role="user",
                            content=user_msg,
                            metadata={
                                "source": "websocket",
                                "msg_type": msg_type
                            }
                        )
                
                # Lưu AI responses từ Web UI
                elif msg_type == "ai_response":
                    ai_msg = msg_data.get("response", "")
                    if ai_msg:
                        add_to_conversation(
                            role="assistant",
                            content=ai_msg,
                            metadata={
                                "source": "websocket",
                                "msg_type": msg_type,
                                "model": msg_data.get("model", "unknown")
                            }
                        )
                
                # 🆕 SMART ANALYZE - Phân tích thông minh với AI (MỚI - ƯU TIÊN)
                elif msg_type == "smart_analyze":
                    user_query = msg_data.get("query", "")
                    llm_response = msg_data.get("response", "")
                    auto_execute = msg_data.get("auto_execute", True)
                    use_ai = msg_data.get("use_ai", True)
                    conversation_history = msg_data.get("history", [])
                    
                    print(f"🧠 [WebSocket] Smart Analyze: query='{user_query[:50]}...'")
                    
                    # Gọi Smart Analyzer API
                    analyze_result = await api_smart_analyze({
                        "user_query": user_query,
                        "llm_response": llm_response,
                        "conversation_history": conversation_history,
                        "auto_execute": auto_execute,
                        "use_ai": use_ai
                    })
                    
                    # Gửi kết quả về client
                    await websocket.send_json({
                        "type": "smart_analyze_result",
                        **analyze_result
                    })
                    
                    print(f"✅ [WebSocket] Smart analyze result sent")
                
                # 🔄 AUTO TOOL EXECUTION (Legacy - vẫn giữ để tương thích)
                elif msg_type == "llm_response_check":
                    llm_response = msg_data.get("response", "")
                    original_query = msg_data.get("query", "")
                    auto_execute = msg_data.get("auto_execute", True)
                    use_smart = msg_data.get("use_smart", True)  # Mặc định dùng Smart Analyzer
                    
                    if llm_response or original_query:
                        print(f"🤖 [WebSocket] Processing: '{(original_query or llm_response)[:50]}...'")
                        
                        if use_smart:
                            # 🧠 Dùng Smart Analyzer (mới - thông minh hơn)
                            result = await api_smart_analyze({
                                "user_query": original_query,
                                "llm_response": llm_response,
                                "auto_execute": auto_execute,
                                "use_ai": True
                            })
                            result["type"] = "smart_analyze_result"
                        else:
                            # Legacy: dùng pattern matching
                            result = await api_auto_execute({
                                "llm_response": llm_response,
                                "original_query": original_query,
                                "auto_execute": auto_execute
                            })
                            result["type"] = "auto_execute_result"
                        
                        # Gửi kết quả về client
                        await websocket.send_json(result)
                        
                        print(f"✅ [WebSocket] Result sent to client")
                
            except json.JSONDecodeError:
                pass  # Not JSON, skip logging
            
            await websocket.send_text(f"Echo: {data}")
    except Exception as e:
        print(f"⚠️ WebSocket client error: {e}")
    finally:
        if websocket in active_connections:
            active_connections.remove(websocket)

@app.on_event("startup")
async def startup():
    """Khởi động server với endpoint manager cải tiến - ghi nhớ endpoint mỗi lần khởi động"""
    global endpoints_config, active_endpoint_index
    
    # Check music folder config and notify
    config_info = check_music_folder_config()
    if config_info.get("has_config"):
        folder_path = config_info.get("folder_path", "")
        print(f"🎵 [Music Config] User music folder configured: {folder_path}")
        print(f"⭐ [Music Priority] Will use play_music_from_user_folder for music requests")
    else:
        print(f"⚠️ [Music Config] No user music folder configured. Will use VLC music_library as fallback.")
    
    # 🔥 NEW: Sử dụng MCPEndpointManager để quản lý kết nối
    if ENDPOINT_MANAGER_AVAILABLE:
        try:
            manager = get_endpoint_manager()
            
            # Đồng bộ config từ manager (đã được load và ghi nhớ từ lần trước)
            endpoints_config = manager.endpoints
            active_endpoint_index = manager.active_index
            
            print(f"📋 [Startup] Loaded {len(endpoints_config)} endpoints from saved config")
            print(f"📍 [Startup] Active endpoint: {active_endpoint_index} ({endpoints_config[active_endpoint_index].get('name', 'Unknown')})")
            
            # Register callbacks để đồng bộ trạng thái
            def on_connect_callback(index, name):
                global xiaozhi_connected
                xiaozhi_connected[index] = True
                print(f"🔔 [Manager] Device {index + 1} ({name}) connected")
            
            def on_disconnect_callback(index):
                global xiaozhi_connected, xiaozhi_connections
                xiaozhi_connected[index] = False
                xiaozhi_connections[index] = None
                print(f"🔌 [Manager] Device {index + 1} disconnected")
            
            def on_error_callback(index, error):
                print(f"❌ [Manager] Device {index + 1} error: {error}")
            
            manager.on_connect(on_connect_callback)
            manager.on_disconnect(on_disconnect_callback)
            manager.on_error(on_error_callback)
            
            # Vẫn dùng websocket client cũ để xử lý messages, nhưng thông tin được ghi nhớ
            print(f"🚀 [Startup] Starting WebSocket clients with remembered endpoints...")
            
        except Exception as e:
            print(f"⚠️ [Startup] EndpointManager error: {e}")
    
    # Enable WebSocket client with error handling
    try:
        # Khởi tạo 3 Xiaozhi clients đồng thời
        for i in range(3):
            asyncio.create_task(xiaozhi_websocket_client(device_index=i))
        print(f"✅ [Startup] WebSocket clients started for {len(endpoints_config)} devices")
    except Exception as e:
        print(f"⚠️ Failed to start WebSocket clients: {e}")

@app.on_event("shutdown")
async def shutdown():
    """Save conversation history và endpoint state on shutdown - tránh mất data"""
    try:
        print("💾 [Shutdown] Saving conversation history...")
        save_conversation_history()
        print(f"✅ [Shutdown] Saved {len(conversation_history)} messages")
        
        # 🔥 NEW: Lưu endpoint state để ghi nhớ cho lần khởi động sau
        if ENDPOINT_MANAGER_AVAILABLE:
            try:
                manager = get_endpoint_manager()
                # Đồng bộ config hiện tại trước khi lưu
                manager.endpoints = endpoints_config
                manager.active_index = active_endpoint_index
                manager.save_config()
                print(f"💾 [Shutdown] Saved endpoint config (active: {active_endpoint_index})")
            except Exception as e:
                print(f"⚠️ [Shutdown] Error saving endpoint config: {e}")
        
        # Lưu endpoints vào file cũ để backward compatible
        save_endpoints_to_file(endpoints_config, active_endpoint_index)
        print(f"💾 [Shutdown] Saved {len(endpoints_config)} endpoints")
        
    except Exception as e:
        print(f"⚠️ [Shutdown] Error saving: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    import uvicorn
    import webbrowser
    import threading
    import time
    
    # ============================================================
    # UNIFIED STARTUP BANNER - PROFESSIONAL EDITION
    # ============================================================
    print("\n")
    print("╔════════════════════════════════════════════════════════════╗")
    print("║                                                            ║")
    print("║          🔐 miniZ MCP v4.3.0 - PROFESSIONAL EDITION        ║")
    print("║                                                            ║")
    print("╚════════════════════════════════════════════════════════════╝")
    print()
    
    # Step 1: FREE EDITION - No License Check
    print("🔍 [1/4] Kiểm tra phiên bản...")
    print("    ✅ miniZ MCP FREE EDITION")
    print("    📦 Loại: FREE (Không giới hạn)")
    print("    👤 Người dùng: Community User")
    
    # Auto-startup check - Lần đầu chạy thì bật auto-start
    marker_file = os.path.join(os.path.expanduser("~"), ".miniz_mcp_installed")
    if not os.path.exists(marker_file):
        print("    ⚙️ Cài đặt khởi động cùng Windows...")
        AutoStartupManager.enable_autostart()
        try:
            with open(marker_file, 'w') as f:
                f.write("installed=true\\nversion=4.3.0\\nedition=FREE")
        except:
            pass
    else:
        if AutoStartupManager.is_autostart_enabled():
            print("    🔄 Khởi động cùng Windows: BẬT")
    
    print()
    
    # Step 2: Check Firewall/Internet Permission
    print("🔥 [2/4] Kiểm tra quyền kết nối mạng...")
    firewall_status = FirewallChecker.check_firewall_rules()
    internet_status = FirewallChecker.check_internet_connection()
    
    if firewall_status['rules_found']:
        print("    ✅ Firewall: Đã cấp quyền")
        print(f"    📌 Rules: {', '.join(firewall_status['rules_found'][:3])}")
    else:
        print("    ⚠️ Firewall: Chưa có rule (Windows sẽ hỏi khi cần)")
        print("    💡 Tip: Nhấn 'Allow' khi Windows hỏi cho phép truy cập mạng")
    
    if internet_status['connected']:
        latency = internet_status.get('latency_ms', '?')
        print(f"    ✅ Internet: Đã kết nối ({latency}ms)")
    else:
        print("    ⚠️ Internet: Không kết nối hoặc đang kiểm tra...")
        print("    💡 Đảm bảo máy tính có kết nối mạng để sử dụng AI")
    
    print()
    
    # Step 3: Initialize Server
    print("🚀 [3/4] Khởi động Server...")
    print("    🌐 Web Dashboard: http://localhost:8000")
    print("    📡 WebSocket MCP: Multi-device support")
    print("    🛠️  Tools: 141 công cụ AI sẵn sàng")
    print("    ✅ Server initialized")
    
    print()
    
    # Step 4: Open Browser
    print("🌐 [4/4] Mở giao diện...")
    print("    ⏳ Browser sẽ tự động mở sau 2 giây...")
    
    def open_browser():
        """Mo browser sau 2 giay"""
        time.sleep(2)
        webbrowser.open("http://localhost:8000")
    
    # Khoi dong thread mo browser
    threading.Thread(target=open_browser, daemon=True).start()
    
    print()
    print("╔════════════════════════════════════════════════════════════╗")
    print("║                                                            ║")
    print("║              ✅ miniZ MCP READY TO USE                      ║")
    print("║                                                            ║")
    print("╚════════════════════════════════════════════════════════════╝")
    print()
    
    # Fix logging error when running as frozen EXE
    import sys
    if getattr(sys, 'frozen', False):
        # Disable uvicorn's default logging config when frozen
        uvicorn.run(app, host="0.0.0.0", port=8000, log_config=None)
    else:
        uvicorn.run(app, host="0.0.0.0", port=8000)

